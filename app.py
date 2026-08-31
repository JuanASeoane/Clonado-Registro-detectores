# -*- coding: utf-8 -*-
"""
APP DE RADON - VERSION STREAMLIT
=================================
Migración de la app de escritorio (Tkinter + OpenCV) a Streamlit.

Conserva TODA la lógica original:
  - Gestión de Centros y Detectores en SQLite
  - Marcado del punto del detector sobre el plano
  - Generación de informe PDF (idéntica, con logo y cabecera)
  - Ajustes (técnico/empresa)

Cambios respecto a la versión Windows:
  - La cámara ya no usa cv2.VideoCapture (ventana propia de escritorio):
    usa st.camera_input, que abre la cámara nativa del navegador
    (funciona igual en PC como en el móvil Android).
  - El punto sobre el plano se marca con un clic usando el componente
    streamlit-image-coordinates.
  - Al generar el PDF aparece un botón "Enviar por WhatsApp" que usa la
    Web Share API de Android para abrir el diálogo nativo de compartir
    con el PDF ya adjunto.
"""

import os
import hmac
import io
import re
import json
import html
import base64
import sqlite3
import zipfile
from datetime import datetime, date, timedelta

import streamlit as st
import streamlit.components.v1 as components
from PIL import Image, ImageDraw, ImageOps
import hashlib

try:
    from zoneinfo import ZoneInfo
    _ZONA_ESPANA = ZoneInfo("Europe/Madrid")
except Exception:
    # Si por lo que sea no está disponible la base de datos de zonas
    # horarias (raro, pero podría pasar en algún sistema muy pelado
    # sin el paquete "tzdata" instalado), se sigue funcionando con la
    # hora del propio servidor en vez de romper la app entera.
    _ZONA_ESPANA = None



st.markdown("""
<style>

/* Checkboxes pequeños: borde naranja visible en toda la aplicación.
   Cuando están desmarcados, fondo blanco para que destaquen sobre fondos oscuros. */
div[data-testid="stCheckbox"] label > div:not([data-testid]) {
    border: 2px solid #F5A623 !important;
    border-radius: 4px !important;
}
div[data-testid="stCheckbox"] label:has(input:not(:checked)) > div:not([data-testid]) {
    background-color: #FFFFFF !important;
}


/* Checkboxes: casilla vacía blanca en toda la aplicación */
div[data-testid="stCheckbox"] input[type="checkbox"]:not(:checked) + div,
div[data-testid="stCheckbox"] input[type="checkbox"]:not(:checked) + div > div,
div[data-testid="stCheckbox"] [data-baseweb="checkbox"] input:not(:checked) ~ div:first-of-type {
    background-color: #FFFFFF !important;
}
div[data-testid="stCheckbox"] [data-baseweb="checkbox"] > div:first-of-type {
    border-color: #FFFFFF !important;
}


    /* Checkbox del Anexo II: mismo aspecto visible que el resto.
       Desmarcado: fondo blanco y borde naranja. */
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-anexo2)
      + div[data-testid="stElementContainer"] div[data-testid="stCheckbox"] label > div:not([data-testid]) {
        border: 2px solid #F5A623 !important;
        border-radius: 4px !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-anexo2)
      + div[data-testid="stElementContainer"] div[data-testid="stCheckbox"] label:has(input:not(:checked)) > div:not([data-testid]) {
        background-color: #FFFFFF !important;
    }


    /* Acordeones de la ficha del detector: fondo oscuro, texto blanco y borde naranja. */
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary {
        background-color: #262626 !important;
        border: none !important;
        border-radius: 8px !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary p {
        font-weight: 700 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary svg {
        fill: #FFFFFF !important;
        color: #FFFFFF !important;
    }

    /* Portada: texto blanco solo en estos dos botones principales. */
    div[data-testid="stButton"] button:has(p):has(span) { }
    div[data-testid="stButton"] button[kind="primary"],
    div[data-testid="stDownloadButton"] button {
        color: #FFFFFF !important;
    }
    div[data-testid="stButton"] button[kind="primary"] p,
    div[data-testid="stDownloadButton"] button p {
        color: #FFFFFF !important;
    }

    /* Botones de eliminación: texto blanco. */
    div[data-testid="stButton"] button[kind="primary"] {
        color: #FFFFFF !important;
    }
    div[data-testid="stButton"] button[kind="primary"] p {
        color: #FFFFFF !important;
    }


    /* Ficha del detector: acordeones compactos y con un solo borde naranja. */
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) {
        display: none !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] {
        margin-top: 0 !important;
        margin-bottom: 0.45rem !important;
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] {
        border: 1.5px solid #F5A623 !important;
        border-radius: 8px !important;
        margin-top: 0 !important;
        margin-bottom: 0 !important;
        overflow: hidden !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary {
        border: none !important;
    }


    /* Tipografía general uniforme y legible.
       Se excluyen expresamente los títulos principales de las pantallas. */
    .stApp,
    .stApp p,
    .stApp label,
    .stApp li,
    .stApp span:not([data-testid="stIconMaterial"]),
    div[data-testid="stMarkdownContainer"] p,
    div[data-testid="stText"],
    div[data-testid="stCaptionContainer"],
    div[data-testid="stAlert"] p,
    div[data-testid="stButton"] button,
    div[data-testid="stButton"] button p,
    div[data-testid="stDownloadButton"] button,
    div[data-testid="stDownloadButton"] button p,
    div[data-testid="stExpander"] summary,
    div[data-testid="stExpander"] summary p,
    div[data-testid="stTabs"] button,
    div[data-testid="stTabs"] button p,
    div[data-baseweb="select"] *,
    div[data-baseweb="input"] input,
    div[data-testid="stTextInput"] input,
    div[data-testid="stNumberInput"] input,
    div[data-testid="stDateInput"] input,
    div[data-testid="stTimeInput"] input,
    textarea,
    table,
    table th,
    table td {
        font-size: 17px !important;
        line-height: 1.35 !important;
    }

    /* Mantener intactos los tamaños especiales de los títulos de pantalla. */
    .titulo-centro,
    .titulo-pantalla,
    .titulo-principal,
    .subtitulo-principal,
    h1, h2, h3,
    [data-testid="stHeading"] * {
        line-height: normal;
    }

    /* En móvil se mantiene el mismo tamaño base para evitar textos demasiado pequeños. */
    @media (max-width: 768px) {
        .stApp,
        .stApp p,
        .stApp label,
        div[data-testid="stButton"] button,
        div[data-testid="stDownloadButton"] button,
        div[data-testid="stExpander"] summary p,
        div[data-baseweb="select"] *,
        input,
        textarea {
            font-size: 17px !important;
        }
    }


    /* Ajustes finales de navegación y títulos. */
    /* Botones de navegación: Inicio, Volver y equivalentes -> texto blanco. */
    div[data-testid="stButton"] button:has(p:first-child) {
        --nav-text-color: inherit;
    }

    /* Los botones secundarios usados para navegación conservan su fondo,
       pero muestran el texto en blanco. */
    div[data-testid="stButton"] button p {
        color: inherit;
    }

    /* Títulos naranjas de las ventanas: algo mayores y más visibles. */
    .titulo-centro {
        line-height: 1.20 !important;
    }

    @media (max-width: 768px) {
        .titulo-centro {
            line-height: 1.20 !important;
        }

        /* Reducir drásticamente el espacio superior reservado por Streamlit.
           Se mantiene un pequeño margen para no pegar el contenido a la barra del móvil. */
        .stAppViewContainer .main .block-container,
        [data-testid="stMainBlockContainer"] {
            padding-top: 1.0rem !important;
        }
    }

    /* Texto blanco en los botones de navegación (← Inicio / ← Volver). */
    div[data-testid="stButton"]:has(button p) button {
        text-shadow: none;
    }


    /* Todos los botones normales de la aplicación usan texto blanco.
       Los campos, selectores y desplegables no se ven afectados. */
    div[data-testid="stButton"] button,
    div[data-testid="stButton"] button p {
        color: #FFFFFF !important;
    }


    /* Informes: respetar siempre la capitalización escrita en las etiquetas.
       Evita mayúsculas automáticas en acordeones, botones y pestañas. */
    div[data-testid="stExpander"] summary,
    div[data-testid="stExpander"] summary p,
    div[data-testid="stButton"] button,
    div[data-testid="stButton"] button p,
    div[data-testid="stDownloadButton"] button,
    div[data-testid="stDownloadButton"] button p,
    div[data-testid="stTabs"] button,
    div[data-testid="stTabs"] button p {
        text-transform: none !important;
    }


    /* Espaciado vertical uniforme.
       Criterio intermedio para:
       - acordeones de la ficha del detector
       - acordeones del informe final
       - botones del menú principal del centro */
    div[data-testid="stExpander"] {
        margin-top: 0 !important;
        margin-bottom: 0.45rem !important;
    }

    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] {
        margin-top: 0 !important;
        margin-bottom: 0.45rem !important;
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }

    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-detector) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] {
        margin-top: 0 !important;
        margin-bottom: 0 !important;
    }

    /* Menú principal del centro: misma separación visual entre botones. */
    div[data-testid="stButton"] {
        margin-top: 0 !important;
        margin-bottom: 0.45rem !important;
    }


    /* Navegación: fondo más oscuro para que el texto blanco tenga buen contraste. */
    div[data-testid="stButton"] button[kind="secondary"] {
        background-color: rgba(8, 48, 66, 0.92) !important;
        color: #FFFFFF !important;
    }
    div[data-testid="stButton"] button[kind="secondary"] p {
        color: #FFFFFF !important;
    }
    div[data-testid="stButton"] button[kind="secondary"]:hover {
        background-color: rgba(6, 38, 53, 0.96) !important;
        color: #FFFFFF !important;
    }

    /* Títulos principales naranjas: 50 % mayores que los 24 px anteriores. */
    .titulo-centro {
        font-size: 48px !important;
        line-height: 1.15 !important;
    }
    @media (max-width: 768px) {
        .titulo-centro {
            font-size: 48px !important;
            line-height: 1.15 !important;
        }
    }


    /* Color del título de los acordeones según su estado:
       - ninguno abierto: todos blancos
       - uno abierto: abierto naranja y cerrados grises */
    div[data-testid="stExpander"] summary p {
        color: #FFFFFF !important;
        transition: color 0.15s ease;
    }

    /* Cuando existe algún acordeón abierto, los títulos cerrados pasan a gris. */
    body:has(div[data-testid="stExpander"] details[open])
    div[data-testid="stExpander"] details:not([open]) summary p,
    .stApp:has(div[data-testid="stExpander"] details[open])
    div[data-testid="stExpander"] details:not([open]) summary p {
        color: #9A9A9A !important;
    }

    /* El título del acordeón abierto destaca en naranja. */
    div[data-testid="stExpander"] details[open] summary p {
        color: #F5A623 !important;
        font-weight: 700 !important;
    }

    /* Criterio general: todos los desplegables llevan borde naranja. */
    div[data-testid="stExpander"] {
        border: none !important;
        border-radius: 8px !important;
    }
    div[data-testid="stExpander"] summary {
        border-color: #F28C28 !important;
    }

    /* Botones de la página principal del centro: texto en mayúsculas. */
    
</style>
""", unsafe_allow_html=True)

def _ahora_espana() -> datetime:
    """La hora de "ahora" tal como la vería alguien en España (con el
    cambio de horario de verano/invierno correcto), en vez de la hora
    del servidor donde corre la app -que normalmente va en UTC, y por
    tanto puede ir 1 o 2 horas por detrás de la hora real de España-.
    Se usa siempre que se captura una fecha/hora "ahora mismo" pensada
    para que la vea o firme una persona (colocación, retirada, fecha
    del informe...)."""
    ahora = datetime.now(_ZONA_ESPANA) if _ZONA_ESPANA else datetime.now()
    return ahora.replace(tzinfo=None)


def _parse_fecha_ddmmyyyy(valor, por_defecto=None):
    """Convierte DD/MM/YYYY (o date) a date para st.date_input."""
    if isinstance(valor, date):
        return valor
    texto = str(valor or "").strip()
    if texto:
        for fmt in ("%d/%m/%Y", "%Y-%m-%d"):
            try:
                return datetime.strptime(texto, fmt).date()
            except Exception:
                pass
    return por_defecto if por_defecto is not None else _ahora_espana().date()


def _date_input_texto(label, key_texto, valor_inicial=None):
    """Calendario Streamlit manteniendo internamente texto DD/MM/YYYY."""
    key_fecha = key_texto + "__date"
    valor_texto = st.session_state.get(key_texto, valor_inicial or "")
    fecha_actual = _parse_fecha_ddmmyyyy(valor_texto, _ahora_espana().date())

    if key_fecha not in st.session_state:
        st.session_state[key_fecha] = fecha_actual
    else:
        fecha_texto = _parse_fecha_ddmmyyyy(valor_texto, fecha_actual)
        if st.session_state.get(key_fecha) != fecha_texto:
            st.session_state[key_fecha] = fecha_texto

    fecha_sel = st.date_input(
        label,
        key=key_fecha,
        format="DD/MM/YYYY",
    )
    texto_sel = fecha_sel.strftime("%d/%m/%Y") if fecha_sel else ""
    st.session_state[key_texto] = texto_sel
    return texto_sel





try:
    from streamlit_image_coordinates import streamlit_image_coordinates
    IMG_COORD_DISPONIBLE = True
except ImportError:
    IMG_COORD_DISPONIBLE = False

# Favicon personalizado (favicon.png debe estar en la misma carpeta que
# este script). Si por lo que sea no se encuentra, se usa el emoji de
# radiactividad como respaldo para que la app no falle al arrancar.


_carpeta_script = os.path.dirname(os.path.abspath(__file__))
_ruta_favicon = os.path.join(_carpeta_script, "favicon.png")
_icono_pagina = _ruta_favicon if os.path.exists(_ruta_favicon) else "☢️"

st.set_page_config(page_title="Detectores Rn", page_icon=_icono_pagina, layout="wide")

# Imagen de fondo de la app (fondo_app.jpg debe estar en la misma
# carpeta que este script). Se codifica en base64 para poder ponerla
# como fondo por CSS sin depender de servir el archivo como URL. Si no
# se encuentra, la app sigue funcionando sin fondo (solo el color
# oscuro de siempre).
_ruta_fondo = os.path.join(_carpeta_script, "fondo_app.jpg")
_fondo_css = ""
if os.path.exists(_ruta_fondo):
    with open(_ruta_fondo, "rb") as _f_fondo:
        _fondo_b64 = base64.b64encode(_f_fondo.read()).decode("utf-8")
    _fondo_css = f"""
    [data-testid="stAppViewContainer"] {{
        background-image:
            linear-gradient(rgba(8, 12, 18, 0.48), rgba(8, 12, 18, 0.55)),
            url("data:image/jpeg;base64,{_fondo_b64}");
        background-size: cover;
        background-position: center;
        background-attachment: fixed;
        background-repeat: no-repeat;
    }}
    [data-testid="stHeader"] {{
        background-color: rgba(0, 0, 0, 0) !important;
    }}
    """

# Ajustes visuales generales: texto en negrita, títulos más pequeños,
# texto pequeño (captions/etiquetas) más grande, botones destacados en
# naranja, y fondo gris oscuro en todos los paneles/tarjetas.
st.markdown(
    """
    <style>
    """ + _fondo_css + """
    html, body, [class*="css"] {
        font-weight: 700 !important;
    }

    /* Aprovechar mejor el espacio superior en ordenador.
       Streamlit puede aplicar el margen superior en distintos contenedores
       según la versión, por eso se corrigen todos los selectores habituales. */
    @media (min-width: 769px) {
        [data-testid="stHeader"] {
            height: 2.2rem !important;
            min-height: 2.2rem !important;
        }
        [data-testid="stMain"] {
            padding-top: 0 !important;
            margin-top: 0 !important;
        }
        [data-testid="stMainBlockContainer"],
        [data-testid="stAppViewBlockContainer"],
        section.main > div,
        .main .block-container,
        .block-container {
            padding-top: 0.25rem !important;
            margin-top: 0 !important;
        }
    }

    /* Títulos más pequeños (pantalla de inicio y el resto de pantallas) */
    h1 { font-size: 1.55rem !important; }
    h2 { font-size: 1.25rem !important; }
    h3 { font-size: 1.1rem !important; }

    /* Texto pequeño (captions, ayudas, pies de foto...) más grande y legible */
    [data-testid="stCaptionContainer"], .stCaption, small {
        font-size: 1rem !important;
    }
    label, .stMarkdown p {
        font-size: 1.02rem !important;
    }

    /* El texto de "Selecciona..." (placeholder) de los desplegables:
       forzado a un gris oscuro siempre, sin depender del tema, porque
       estos campos suelen tener fondo claro (blanco/gris/rosa) y el
       texto blanco del tema oscuro se volvía invisible encima. */
    div[data-baseweb="select"] [class*="placeholder"] {
        color: #4a4a4a !important;
    }

    /* Texto e etiquetas en blanco SIEMPRE, sin depender de que el archivo
       .streamlit/config.toml (tema oscuro) esté presente. Si ese archivo
       falta (p.ej. por no haberlo subido a GitHub, al ser una carpeta
       oculta que empieza por un punto), Streamlit usa su tema claro
       por defecto y estos textos se verían en negro sobre fondo
       oscuro. Estas reglas hacen que se vean bien en cualquier caso. */
    label, label p, label span,
    .stMarkdown, .stMarkdown p, .stMarkdown li, .stMarkdown span,
    [data-testid="stCaptionContainer"], .stCaption,
    [data-testid="stWidgetLabel"], [data-testid="stWidgetLabel"] p,
    [data-testid="stFileUploaderDropzoneInstructions"] span,
    [data-testid="stFileUploaderDropzoneInstructions"] small,
    h1, h2, h3, h4, h5, h6 {
        color: #f5f5f5 !important;
    }
    /* El fondo general de la app y el texto normal, por si tampoco
       está el tema oscuro aplicado (fondo oscuro por defecto). */
    [data-testid="stAppViewContainer"], [data-testid="stApp"], .stApp {
        color: #f5f5f5 !important;
    }
    /* La celda que se abre para editar un valor en una tabla
       editable (st.data_editor) NO es un input normal: es un
       elemento de "glide-data-grid" (clase gdg-input) insertado en
       un "portal" aparte. Sin esta regla heredaría el texto en
       blanco de más arriba y, sobre su fondo claro, no se vería lo
       que se escribe. */
    .gdg-input {
        background-color: #ffffff !important;
        color: #000000 !important;
    }

    /* Botones "primary" (acción principal: Abrir, Guardar...) en verde */
    button[kind="primary"], button[kind="primaryFormSubmit"] {
        background-color: #F5A623 !important;
        border-color: #F5A623 !important;
        color: #000000 !important;
    }
    button[kind="primary"]:hover, button[kind="primaryFormSubmit"]:hover {
        background-color: #d68f10 !important;
        border-color: #d68f10 !important;
        color: #000000 !important;
    }

    /* Botones "secondary" y "tertiary" (acciones neutras: Nuevo centro,
       Importar centro, Nuevo detector, Añadir plano, Generar
       documentos...) con estilo neutro de contorno, no relleno de
       color, para que solo la acción principal (amarilla) destaque. */
    button[kind="secondary"], button[kind="secondaryFormSubmit"],
    button[kind="tertiary"], button[kind="tertiaryFormSubmit"],
    .stDownloadButton button {
        background-color: #262626 !important;
        border: 1.5px solid #F28C28 !important;
        color: #ffffff !important;
    }
    button[kind="secondary"]:hover, button[kind="secondaryFormSubmit"]:hover,
    button[kind="tertiary"]:hover, button[kind="tertiaryFormSubmit"]:hover,
    .stDownloadButton button:hover {
        background-color: #333333 !important;
        border-color: #999999 !important;
        color: #ffffff !important;
    }

    /* Botón "Eliminar" (marcado con un div justo antes, ver código
       Python): con borde rojo y texto rojo, sin relleno, para que se
       vea claramente como una acción de peligro/secundaria. Se ancla
       al nivel exacto "stElementContainer" (sin ">", el marcador queda
       varios niveles por debajo) para no afectar por error a otros
       botones vecinos en la misma fila de columnas (p.ej. "Abrir"). */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-eliminar) + div[data-testid="stElementContainer"] button {
        background-color: transparent !important;
        border: 1px solid #F5A623 !important;
        color: #FFFFFF !important;
        width: auto !important;
        min-width: 0 !important;
        display: inline-flex !important;
        opacity: 0.9;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-eliminar) + div[data-testid="stElementContainer"] button p {
        color: #FFFFFF !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-eliminar) + div[data-testid="stElementContainer"] button:hover {
        background-color: rgba(245, 166, 35, 0.15) !important;
        border-color: #F5A623 !important;
        color: #FFFFFF !important;
        opacity: 1;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-eliminar) + div[data-testid="stElementContainer"] button:hover p {
        color: #FFFFFF !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-eliminar) + div[data-testid="stElementContainer"] {
        display: flex !important;
        justify-content: flex-start !important;
    }

    /* Acordeones de la pantalla del informe final: fondo rosa si le
       falta algo, gris si está completo (igual que el resto de
       campos de la app). Entre el marcador y el stExpander real hay
       un "stLayoutWrapper" de por medio. */
    /* Los desplegables del informe final usan el mismo aspecto que los botones
       de la ventana del centro. El estado pendiente se indica únicamente
       mediante el punto rojo del título, sin alterar el color del botón. */
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-rosa) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary,
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-gris) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary {
        background-color: #262626 !important;
        border: none !important;
        border-radius: 8px !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-rosa) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary p,
    div[data-testid="stElementContainer"]:has(div.marcador-acordeon-gris) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary p {
        font-weight: 700 !important;
    }
    /* Desplegable «Profesionales en esta sala»: mismo criterio visual que
       los acordeones del informe final. */
    div[data-testid="stElementContainer"]:has(div.marcador-profesionales-sala) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary {
        background-color: #262626 !important;
        border: none !important;
        border-radius: 8px !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-profesionales-sala) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary p {
        font-weight: 700 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-profesionales-sala) + div[data-testid="stLayoutWrapper"] div[data-testid="stExpander"] summary svg {
        fill: #FFFFFF !important;
        color: #FFFFFF !important;
    }

    /* Botón "Generar documentos": lo más pequeño posible, alineado a
       la izquierda (ya va en una columna estrecha). Fondo negro con
       borde y letra amarilla, igual que el resto de acciones
       principales de la app. Se ancla al nivel exacto
       "stElementContainer" (sin restringir a hijo directo, el marcador
       queda varios niveles por debajo) para no afectar a otros
       botones de columnas vecinas. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-generar) + div[data-testid="stElementContainer"] button {
        font-size: 0.75rem !important;
        padding: 0.15rem 0.6rem !important;
        min-height: 1.7rem !important;
        height: 1.7rem !important;
        line-height: 1 !important;
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-generar) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-generar) + div[data-testid="stElementContainer"] {
        display: flex !important;
        justify-content: flex-start !important;
    }

    /* Cabecera del bloque desplegable actualmente abierto: texto en
       amarillo (igual que el logo) para distinguirlo del resto. Se
       usa ":has(div...)" sin ">" (descendiente, no solo hijo directo)
       porque el marcador queda varios niveles por debajo del div que
       realmente tiene como hermano el botón. */
    div:has(div.marcador-bloque-activo) + div button {
        color: #F5A623 !important;
    }

    /* Botón "Importar centro": fondo oscuro con las letras en blanco
       (igual que el "+" de "Nuevo centro"), con borde amarillo. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-importar) + div[data-testid="stElementContainer"] button {
        background-color: #262626 !important;
        border: 1px solid #F5A623 !important;
        color: #ffffff !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-importar) + div[data-testid="stElementContainer"] button:hover {
        background-color: #333333 !important;
        border-color: #F5A623 !important;
        color: #ffffff !important;
    }

    /* Botón "Nuevo centro": mismo fondo oscuro de siempre, con borde
       amarillo (igual que "Importar centro"). */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-nuevo-centro) + div[data-testid="stElementContainer"] button {
        border: 1px solid #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-nuevo-centro) + div[data-testid="stElementContainer"] button:hover {
        border-color: #F5A623 !important;
    }

    /* Botones "Nuevo centro" / "Importar centro": cuando su formulario
       está abierto, el texto se pone en naranja (y vuelve a su color
       normal en cuanto se abre el otro, ya que solo uno lleva este
       marcador a la vez). Se restringe al nivel "stElementContainer"
       (con ">") porque, al estar en columnas, un selector más genérico
       también coincide con la COLUMNA vecina y pintaría el botón
       equivocado. */
    div[data-testid="stElementContainer"]:has(div.marcador-activo-naranja) + div[data-testid="stElementContainer"] button,
    div[data-testid="stElementContainer"]:has(div.marcador-activo-naranja) + div[data-testid="stElementContainer"] button:hover {
        color: #F5A623 !important;
    }

    /* Botones "Volver a Centros" / "Volver al centro": letra en azul
       claro. */
    /* Botones "Volver..." (Volver a Centros, Volver al centro,
       Volver): fondo azul claro y letra negra. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-azul-claro) + div[data-testid="stElementContainer"] button {
        background-color: #083246 !important;
        border: 1.5px solid #F5A623 !important;
        color: #FFFFFF !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-azul-claro) + div[data-testid="stElementContainer"] button p {
        color: #FFFFFF !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-azul-claro) + div[data-testid="stElementContainer"] button:hover {
        background-color: #062735 !important;
        border-color: #F5A623 !important;
        color: #FFFFFF !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-azul-claro) + div[data-testid="stElementContainer"] button:hover p {
        color: #FFFFFF !important;
    }

    /* Casilla de "Nº personas" en categorías profesionales: ancho
       máximo para que solo quepan 2 dígitos (en vez de ocupar todo el
       ancho disponible, sobre todo notorio en móvil). */
    div[data-testid="stElementContainer"]:has(div.marcador-num-personas) + div[data-testid="stElementContainer"] div[data-testid="stNumberInput"] > div:not([data-testid="stWidgetLabel"]) {
        max-width: 8rem !important;
    }

    /* Todos los botones dentro de "Imagen exterior" (Subir archivo,
       Activar/Tomar cámara, y el de capturar dentro del propio visor):
       fondo negro, borde y letra amarilla. Se restringe a la columna
       exacta ("stColumn") que contiene el marcador; un selector sin
       esa restricción coincide con TODA la página (cualquier ancestro
       que también contenga el marcador en su árbol) y pinta botones
       completamente ajenos. */
    div[data-testid="stColumn"]:has(div.marcador-imagen-exterior) button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stColumn"]:has(div.marcador-imagen-exterior) button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* Dentro del formulario de importar: el botón de subir archivo y
       el de confirmar "Importar", en naranja sobre fondo negro. */
    div:has(div.marcador-uploader-importar) + div button,
    div:has(div.marcador-btn-confirmar-importar) + div button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-uploader-importar) + div button:hover,
    div:has(div.marcador-btn-confirmar-importar) + div button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-btn-crear-centro) + div button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-btn-crear-centro) + div button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-btn-guardar-ajustes) + div button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-btn-guardar-ajustes) + div button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-btn-guardar-centro) + div button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-btn-guardar-centro) + div button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* Botón "➕ Añadir" (categorías profesionales): fondo negro, borde
       y letra amarilla. Está dentro de columnas, así que se restringe
       al nivel exacto "stElementContainer" para no afectar a otros
       botones cercanos (p.ej. "Eliminar seleccionadas"). */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-anadir-categoria) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-anadir-categoria) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* "Planos del centro": botones "Añadir plano", "Subir archivo" (del
       selector de imagen) y "Guardar plano", todos en fondo negro con
       borde y letra amarilla. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-plano-amarillo) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-plano-amarillo) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }
    /* "Añadir plano" iba con el texto muy pegado al borde: un poco
       más de espacio a los lados (afecta a los 3 botones de este
       mismo marcador: Añadir plano, Subir archivo y Guardar plano). */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-plano-amarillo) + div[data-testid="stElementContainer"] button {
        padding-left: 1rem !important;
        padding-right: 1rem !important;
    }

    /* "Detectores colocados": Nuevo detector y Abrir detector, en
       fondo negro con borde y letra amarilla. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-nuevo-detector) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
        padding-left: 1.1rem !important;
        padding-right: 1.1rem !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-nuevo-detector) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-abrir-detector) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-abrir-detector) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* "Retirada de detectores": Capturar fecha y hora / Guardar, en
       fondo negro con borde y letra amarilla. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-retirada-amarillo) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-retirada-amarillo) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* "Guardar detector" (al final de la pantalla del detector): fondo
       negro con borde y letra amarilla. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-guardar-detector) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-guardar-detector) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* Títulos de cada apartado de anexo (I, II, III, IV): todos con
       el mismo tamaño y en amarillo. No vale hacerlo con un "style"
       en línea porque st.markdown sanea el HTML y le quita los
       "!important" a los estilos en línea; por eso se hace con el
       marcador + esta regla, igual que en el resto de la app. */
    div[data-testid="stElementContainer"]:has(div.marcador-titulo-anexo) + div[data-testid="stElementContainer"] p {
        color: #F5A623 !important;
        font-weight: 700 !important;
    }

    /* Casillas de "Categorías profesionales": el cuadrito de marcar
       apenas se veía (sin contraste con el fondo oscuro), así que se
       le pone un borde amarillo bien visible; y el texto de la
       categoría se pone al doble de grande para que se lea mejor. */
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria) + div[data-testid="stElementContainer"] div[data-testid="stCheckbox"] label > div:not([data-testid]) {
        border: 2px solid #F5A623 !important;
        border-radius: 4px !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria) + div[data-testid="stElementContainer"] div[data-testid="stCheckbox"] label,
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria) + div[data-testid="stElementContainer"] div[data-testid="stCheckbox"] label p,
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria) + div[data-testid="stElementContainer"] div[data-testid="stCheckbox"] label span,
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria) + div[data-testid="stElementContainer"] div[data-testid="stCheckbox"] label div,
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria) + div[data-testid="stElementContainer"] div[data-testid="stWidgetLabel"],
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria) + div[data-testid="stElementContainer"] div[data-testid="stWidgetLabel"] * {
        font-size: 17px !important;
        line-height: 1.35 !important;
    }


    /* Punto 6 del informe: tabla continua con celdas tipo Excel. */
    div[data-testid="stElementContainer"]:has(div.marcador-tabla-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"] {
        gap: 0 !important;
        background: #C9DFF2 !important;
        border: 1px solid #7897B2 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-tabla-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"] > div {
        background: #C9DFF2 !important;
        border-right: 1px solid #7897B2 !important;
        padding: 0.34rem 0.38rem !important;
        margin: 0 !important;
        min-height: 3rem !important;
        display: flex !important;
        align-items: center !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-tabla-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"] > div:last-child {
        border-right: 0 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"] {
        gap: 0 !important;
        background: #C9DFF2 !important;
        border-left: 1px solid #7897B2 !important;
        border-right: 1px solid #7897B2 !important;
        border-bottom: 1px solid #7897B2 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"] > div {
        background: #C9DFF2 !important;
        border-right: 1px solid #7897B2 !important;
        padding: 0 !important;
        margin: 0 !important;
        min-width: 0 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"] > div:last-child {
        border-right: 0 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"]
      div[data-testid="stTextInput"],
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"]
      div[data-testid="stNumberInput"] {
        margin: 0 !important;
        padding: 0 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"]
      div[data-baseweb="input"] {
        border: 0 !important;
        border-radius: 0 !important;
        box-shadow: none !important;
        background: #C9DFF2 !important;
        min-height: 2.7rem !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] div[data-testid="stHorizontalBlock"]
      div[data-baseweb="input"] input {
        background: #C9DFF2 !important;
        border: 0 !important;
        border-radius: 0 !important;
        box-shadow: none !important;
        padding: 0.45rem 0.4rem !important;
        min-height: 2.7rem !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-fila-punto6-excel)
      + div[data-testid="stElementContainer"] input:disabled {
        opacity: 1 !important;
        background: #C9DFF2 !important;
        -webkit-text-fill-color: inherit !important;
    }

    /* Tabla de resultados del informe (Resultado/Incertidumbre): en
       móvil, Streamlit apila cualquier st.columns() en cuanto el
       contenido no cabe (cada columna pasa a ocupar el 100% del
       ancho y salta de línea). Aquí se fuerza a que estas dos
       columnas concretas se mantengan siempre en la misma fila,
       repartiéndose el ancho a la mitad cada una, para que de verdad
       se vea como una tabla de dos columnas también en el móvil. */
    div[data-testid="stElementContainer"]:has(div.marcador-tabla-resultado-fila) + div[data-testid="stLayoutWrapper"] div[data-testid="stHorizontalBlock"] {
        flex-wrap: nowrap !important;
        gap: 0.6rem !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-tabla-resultado-fila) + div[data-testid="stLayoutWrapper"] div[data-testid="stColumn"] {
        width: 50% !important;
        min-width: 0 !important;
        flex: 1 1 0 !important;
    }

    div[data-testid="stElementContainer"]:has(div.marcador-btn-informe-completo) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-informe-completo) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* Descargar PDF / Excel / fotos, y Seleccionar todas / Quitar
       selección: fondo negro con borde y letra amarilla. */
    div[data-testid="stElementContainer"]:has(div.marcador-btn-descarga-amarillo) + div[data-testid="stElementContainer"] button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div[data-testid="stElementContainer"]:has(div.marcador-btn-descarga-amarillo) + div[data-testid="stElementContainer"] button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* Título grande de la pantalla de inicio ("Detectores de Radón"),
       al doble de tamaño que un título normal. Usa una clase (en vez
       de estilo en línea) porque Streamlit elimina por completo el
       atributo style="" de cualquier HTML si contiene "!important". */
    p.titulo-home {
        color: #F5A623 !important;
        font-size: 48px !important;
        font-weight: 800 !important;
        line-height: 1.2 !important;
        margin: 0.5rem 0 !important;
    }

    /* Título del nombre del centro (pantalla del centro), al doble de
       tamaño que un título normal. */
    p.titulo-centro {
        color: #F5A623 !important;
        font-size: 48px !important;
        font-weight: 800 !important;
        line-height: 1.2 !important;
        margin: 0.5rem 0 !important;
    }

    /* Subtítulos de sección en amarillo (p.ej. "Centros registrados"),
       con el mismo truco de clase por la limitación de style="" con
       !important explicada arriba. */
    p.subtitulo-amarillo {
        color: #F5A623 !important;
        font-size: 1.1rem !important;
        font-weight: 700 !important;
        margin: 0.8rem 0 0.3rem 0 !important;
    }

    /* Botón "Abrir centro": fondo negro con letras naranjas (en vez
       del amarillo/negro normal de los botones "primary"). */
    div:has(div.marcador-btn-abrir-centro) + div button {
        background-color: #000000 !important;
        border: 1px solid #F5A623 !important;
        color: #F5A623 !important;
    }
    div:has(div.marcador-btn-abrir-centro) + div button:hover {
        background-color: #1a1a1a !important;
        border-color: #F5A623 !important;
        color: #F5A623 !important;
    }

    /* "Ventanas" (tarjetas, expanders, formularios) con fondo gris oscuro */
    div[data-testid="stVerticalBlockBorderWrapper"],
    div[data-testid="stExpander"],
    div[data-testid="stForm"] {
        background-color: #262626 !important;
        border-radius: 10px;
    }

    /* Títulos de cada ventana/panel en naranja */
    h1, h2, h3 {
        color: #f5f5f5 !important;
    }

    /* Cabecera de los desplegables tipo acordeón (expanders) en blanco,
       para que combine con las casillas (p.ej. "➕ Nuevo centro") */
    div[data-testid="stExpander"] summary {
        background-color: #ffffff !important;
        border-radius: 8px !important;
    }

    /* Ocultar el texto "Press Enter to apply" que Streamlit muestra
       bajo las casillas de texto mientras se escribe */

    /* Salvaguarda final de tamaño de títulos principales. */
    p.titulo-centro,
    p.titulo-home {
        font-size: 48px !important;
        line-height: 1.15 !important;
    }
    @media (max-width: 768px) {
        p.titulo-centro,
        p.titulo-home {
            font-size: 48px !important;
            line-height: 1.15 !important;
        }
    }


    /* Tamaño único para TODOS los títulos principales de ventana.
       titulo-centro: ventanas vinculadas a un centro.
       titulo-home: inicio, detector, datos de empresa y acceso. */
    p.titulo-centro,
    p.titulo-home {
        font-size: 48px !important;
        line-height: 1.12 !important;
        font-weight: 700 !important;
    }

    @media (max-width: 768px) {
        p.titulo-centro,
        p.titulo-home {
            font-size: 48px !important;
            line-height: 1.12 !important;
        }
    }


    /* Categorías profesionales: mismo tamaño que el texto general de la app. */
    div[data-testid="stCheckbox"] label,
    div[data-testid="stCheckbox"] label p,
    div[data-testid="stRadio"] label,
    div[data-testid="stRadio"] label p,
    div[data-testid="stSelectbox"] label,
    div[data-testid="stSelectbox"] label p {
        font-size: 17px !important;
        line-height: 1.35 !important;
    }


    /* Textos auxiliares/comentarios: blanco para mejorar la legibilidad. */
    div[data-testid="stCaptionContainer"],
    div[data-testid="stCaptionContainer"] p,
    [data-testid="stCaptionContainer"],
    [data-testid="stCaptionContainer"] p,
    .stCaption,
    .stCaption p,
    small {
        color: #FFFFFF !important;
        opacity: 1 !important;
    }


    /* Título funcional de cada ventana principal. */
    p.titulo-ventana-principal {
        color: #F5A623 !important;
        font-size: 48px !important;
        line-height: 1.15 !important;
        font-weight: 800 !important;
        margin: 0.35rem 0 0.45rem 0 !important;
    }

    @media (max-width: 768px) {
        p.titulo-ventana-principal {
            font-size: 48px !important;
        }
    }


    /* Criterio definitivo para títulos principales de TODAS las ventanas. */
    p.titulo-centro,
    p.titulo-home,
    p.titulo-ventana-principal,
    div[data-testid="stHeading"] h1,
    div[data-testid="stHeading"] h2,
    div[data-testid="stHeading"] h3 {
        font-size: 48px !important;
        line-height: 1.12 !important;
        font-weight: 800 !important;
    }

    @media (max-width: 768px) {
        p.titulo-centro,
        p.titulo-home,
        p.titulo-ventana-principal,
        div[data-testid="stHeading"] h1,
        div[data-testid="stHeading"] h2,
        div[data-testid="stHeading"] h3 {
            font-size: 48px !important;
            line-height: 1.12 !important;
        }
    }


    /* CORRECCIÓN FINAL: títulos principales visibles y categorías uniformes. */
    p.titulo-home,
    p.titulo-centro,
    p.titulo-ventana-principal {
        font-size: 48px !important;
        line-height: 1.12 !important;
        font-weight: 800 !important;
    }

    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria)
    + div[data-testid="stElementContainer"]
    div[data-testid="stCheckbox"] label,
    div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria)
    + div[data-testid="stElementContainer"]
    div[data-testid="stCheckbox"] label * {
        font-size: 17px !important;
        line-height: 1.35 !important;
    }

    @media (max-width: 768px) {
        p.titulo-home,
        p.titulo-centro,
        p.titulo-ventana-principal {
            font-size: 48px !important;
        }

        div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria)
        + div[data-testid="stElementContainer"]
        div[data-testid="stCheckbox"] label,
        div[data-testid="stElementContainer"]:has(div.marcador-checkbox-categoria)
        + div[data-testid="stElementContainer"]
        div[data-testid="stCheckbox"] label * {
            font-size: 17px !important;
        }
    }


    /* Profesionales en esta sala: etiquetas en formato normal, no mayúsculas. */
    div[data-testid="stElementContainer"]:has(div.marcador-profesionales-sala-sentence-case)
    ~ div[data-testid="stElementContainer"] label,
    div[data-testid="stElementContainer"]:has(div.marcador-profesionales-sala-sentence-case)
    ~ div[data-testid="stElementContainer"] label *,
    div[data-testid="stElementContainer"]:has(div.marcador-profesionales-sala-sentence-case)
    ~ div[data-testid="stElementContainer"] div[data-testid="stWidgetLabel"],
    div[data-testid="stElementContainer"]:has(div.marcador-profesionales-sala-sentence-case)
    ~ div[data-testid="stElementContainer"] div[data-testid="stWidgetLabel"] * {
        text-transform: none !important;
    }

    [data-testid="InputInstructions"] {
        display: none !important;
    }

    /* Casillas de texto y desplegables: SOLO el recuadro donde se
       escribe/selecciona queda en blanco con letra negra, un 25% más
       grande (1rem -> 1.25rem) y con esquinas rectas (sin redondear).
       La etiqueta de encima (el texto que describe el campo, con
       testid "stWidgetLabel") se excluye a propósito para que quede
       sobre el fondo gris de la ventana.
       Se usa "*" para pintar TODO lo de dentro del recuadro (incluida
       la parte del desplegable de Streamlit, que anida varios niveles
       de <div> internos) y luego se fuerza la etiqueta a transparente
       para que no quede afectada. */
    div[data-testid="stTextInput"] > div:not([data-testid="stWidgetLabel"]),
    div[data-testid="stTextInput"] > div:not([data-testid="stWidgetLabel"]) *,
    div[data-testid="stTextArea"] > div:not([data-testid="stWidgetLabel"]),
    div[data-testid="stTextArea"] > div:not([data-testid="stWidgetLabel"]) *,
    div[data-testid="stNumberInput"] > div:not([data-testid="stWidgetLabel"]),
    div[data-testid="stNumberInput"] > div:not([data-testid="stWidgetLabel"]) *,
    div[data-testid="stDateInput"] > div:not([data-testid="stWidgetLabel"]),
    div[data-testid="stDateInput"] > div:not([data-testid="stWidgetLabel"]) *,
    div[data-testid="stSelectbox"] > div:not([data-testid="stWidgetLabel"]),
    div[data-testid="stSelectbox"] > div:not([data-testid="stWidgetLabel"]) * {
        background-color: #ffffff !important;
        color: #000000 !important;
        font-size: 1.1rem !important;
        font-weight: 400 !important;
        border-radius: 8px !important;
    }
    div[data-testid="stTextInput"] input,
    div[data-testid="stTextArea"] textarea,
    div[data-testid="stNumberInput"] input,
    div[data-testid="stDateInput"] input,
    div[data-testid="stSelectbox"] div[data-baseweb="select"] * {
        background-color: #ffffff !important;
        color: #000000 !important;
        font-size: 1.1rem !important;
        font-weight: 400 !important;
        border-radius: 8px !important;
    }
    /* Los campos "disabled" (de solo lectura, como la vista previa del
       tipo de zona en el informe) reciben en algunos móviles (Chrome/
       Safari en Android e iOS) un color de texto propio y más claro
       para los campos deshabilitados, que las reglas de arriba no
       llegan a pisar del todo (algunos navegadores usan la propiedad
       -webkit-text-fill-color específicamente para esto, que manda
       más que "color" a secas). Se fuerzan aquí las dos, y se quita
       cualquier atenuación, para que el texto se siga leyendo bien
       aunque el campo esté deshabilitado. */
    div[data-testid="stTextInput"] input:disabled,
    div[data-testid="stTextArea"] textarea:disabled,
    div[data-testid="stNumberInput"] input:disabled {
        background-color: #ffffff !important;
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        opacity: 1 !important;
    }
    /* La etiqueta de encima NO se toca: se fuerza a volver al color de
       texto normal de la app, sobre el fondo gris de la ventana. */
    div[data-testid="stWidgetLabel"],
    div[data-testid="stWidgetLabel"] * {
        background-color: transparent !important;
    }
    /* Menú desplegable del selectbox al abrirlo: se renderiza en un
       "popover" flotante fuera del recuadro, así que se cubre con
       TODOS los posibles selectores que puede usar Streamlit/BaseWeb
       (testid, data-baseweb de popover/menu/menu-item, y roles ARIA
       listbox/option) y con "*" para llegar a todas las capas
       internas de texto/fondo, sea cual sea la estructura exacta. */
    div[data-baseweb="popover"],
    div[data-baseweb="popover"] *,
    div[data-baseweb="menu"],
    div[data-baseweb="menu"] *,
    li[data-baseweb="menu-item"],
    li[data-baseweb="menu-item"] *,
    ul[data-testid="stSelectboxVirtualDropdown"],
    ul[data-testid="stSelectboxVirtualDropdown"] *,
    [role="listbox"],
    [role="listbox"] *,
    [role="option"],
    [role="option"] * {
        background-color: #ffffff !important;
        color: #000000 !important;
        border-radius: 8px !important;
    }
    div[data-baseweb="popover"] li,
    li[data-baseweb="menu-item"],
    ul[data-testid="stSelectboxVirtualDropdown"] li,
    [role="listbox"] li,
    [role="option"] {
        font-size: 1.1rem !important;
    }
    /* Opción resaltada al pasar el ratón/dedo: un gris muy claro para
       distinguirla, siempre con texto negro */
    div[data-baseweb="popover"] li:hover,
    li[data-baseweb="menu-item"]:hover,
    ul[data-testid="stSelectboxVirtualDropdown"] li:hover,
    [role="listbox"] li:hover,
    [role="option"]:hover,
    div[data-baseweb="popover"] li[aria-selected="true"],
    li[data-baseweb="menu-item"][aria-selected="true"],
    [role="listbox"] li[aria-selected="true"],
    [role="option"][aria-selected="true"] {
        background-color: #eeeeee !important;
        color: #000000 !important;
    }
    /* Campos vacíos: coral suave; placeholder oscuro. */
    div[data-baseweb="select"] > div {
        background-color: #FFFFFF !important;
    }
    div[data-baseweb="select"] [class*="placeholder"],
    div[data-baseweb="select"] input::placeholder {
        color: #4f4f4f !important;
        -webkit-text-fill-color: #4f4f4f !important;
        opacity: 1 !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ============================================================
# BASE DE DATOS  (idéntica a la app de escritorio)
# ============================================================

DB_NAME = "radon_data.db"


def get_data_dir():
    """Carpeta donde se guardan la BD y las imágenes.

    Se usa una carpeta junto al propio script para que funcione igual
    en local, en un servidor o en Streamlit Community Cloud. Si prefieres
    guardar los datos en el perfil del usuario como hacía la app de
    Windows, cambia la línea siguiente por:
        data_dir = os.path.join(os.path.expanduser("~"), "RadonApp")
    """
    data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "RadonApp_data")
    os.makedirs(data_dir, exist_ok=True)
    return data_dir


def get_db_path():
    return os.path.join(get_data_dir(), DB_NAME)


def init_db():
    db_path = get_db_path()

    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS centros
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  nombre TEXT, zona TEXT, fecha_medicion TEXT, imagen_exterior_path TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS detectores
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  centro_id INTEGER,
                  planta TEXT, sala TEXT, fecha TEXT, detector_codigo TEXT,
                  plano_path TEXT, punto_x REAL, punto_y REAL,
                  foto_situacion_path TEXT, foto_detector_path TEXT,
                  fecha_creacion TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS settings
                 (id INTEGER PRIMARY KEY CHECK (id=1), tecnico TEXT)''')
    c.execute("INSERT OR IGNORE INTO settings (id, tecnico) VALUES (1, '')")
    # Planos del centro (sin límite de cantidad), cada uno con su nombre
    # (p.ej. "Planta baja", "Planta 1"...). El punto rojo de cada
    # detector se guarda en el propio detector (punto_x/punto_y),
    # apuntando a uno de estos planos mediante plano_centro_id.
    c.execute('''CREATE TABLE IF NOT EXISTS planos_centro
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  centro_id INTEGER, nombre TEXT, ruta TEXT, orden INTEGER)''')

    # Categorías profesionales expuestas en el centro (p.ej. "Enfermería"
    # -> 8 personas), con número de personas expuestas de cada una.
    c.execute('''CREATE TABLE IF NOT EXISTS categorias_centro
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  centro_id INTEGER, categoria TEXT, num_personas INTEGER)''')

    # --- Migraciones: añadir columnas nuevas sin perder los datos ya
    # guardados, por si la base de datos viene de una versión anterior. ---
    c.execute("PRAGMA table_info(centros)")
    columnas_centros = {fila[1] for fila in c.fetchall()}
    if "tecnico" not in columnas_centros:
        c.execute("ALTER TABLE centros ADD COLUMN tecnico TEXT DEFAULT ''")
    if "direccion" not in columnas_centros:
        c.execute("ALTER TABLE centros ADD COLUMN direccion TEXT DEFAULT ''")
    if "tipo_centro" not in columnas_centros:
        c.execute("ALTER TABLE centros ADD COLUMN tipo_centro TEXT DEFAULT ''")
    if "propietario" not in columnas_centros:
        c.execute("ALTER TABLE centros ADD COLUMN propietario TEXT DEFAULT ''")

    c.execute("PRAGMA table_info(detectores)")
    columnas_detectores = {fila[1] for fila in c.fetchall()}
    if "codigo_sala" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN codigo_sala TEXT DEFAULT ''")
    if "profesionales_sala" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN profesionales_sala TEXT DEFAULT ''")
    if "hora_colocacion" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN hora_colocacion TEXT DEFAULT ''")
    if "turno_trabajo" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN turno_trabajo TEXT DEFAULT ''")
    if "nivel" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN nivel TEXT DEFAULT ''")
    if "plano_centro_id" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN plano_centro_id INTEGER")
    if "fecha_retirada_real" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN fecha_retirada_real TEXT DEFAULT ''")
    if "hora_retirada_real" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN hora_retirada_real TEXT DEFAULT ''")

    c.execute("PRAGMA table_info(categorias_centro)")
    columnas_categorias = {fila[1] for fila in c.fetchall()}
    if "turno" not in columnas_categorias:
        c.execute("ALTER TABLE categorias_centro ADD COLUMN turno TEXT DEFAULT ''")

    if "resultado_bq_m3" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN resultado_bq_m3 REAL")
    if "incertidumbre" not in columnas_detectores:
        c.execute("ALTER TABLE detectores ADD COLUMN incertidumbre TEXT DEFAULT ''")

    c.execute("PRAGMA table_info(centros)")
    columnas_centros_extra = {fila[1] for fila in c.fetchall()}
    for columna_nueva in (
        "superficie_construida", "superficie_util", "num_plantas",
        "fecha_comunicacion_trab", "medio_comunicacion",
    ):
        if columna_nueva not in columnas_centros_extra:
            c.execute(f"ALTER TABLE centros ADD COLUMN {columna_nueva} TEXT DEFAULT ''")

    # Logotipo personalizado (dato global de la empresa, no de cada
    # centro): se guarda como archivo en la carpeta de datos, con la
    # ruta y el nombre original guardados en "settings".
    c.execute("PRAGMA table_info(settings)")
    columnas_settings = {fila[1] for fila in c.fetchall()}
    if "logo_path" not in columnas_settings:
        c.execute("ALTER TABLE settings ADD COLUMN logo_path TEXT DEFAULT ''")
    if "logo_nombre" not in columnas_settings:
        c.execute("ALTER TABLE settings ADD COLUMN logo_nombre TEXT DEFAULT ''")

    # Datos de la empresa (globales, no cambian de un centro a otro):
    # nombre de la empresa y CIF, con sus valores por defecto.
    c.execute("PRAGMA table_info(settings)")
    columnas_settings = {fila[1] for fila in c.fetchall()}
    if "empresa" not in columnas_settings:
        c.execute(
            "ALTER TABLE settings ADD COLUMN empresa TEXT DEFAULT 'Área Sanitaria da Coruña e Cee'"
        )
    if "cif" not in columnas_settings:
        c.execute("ALTER TABLE settings ADD COLUMN cif TEXT DEFAULT 'Q1569009B'")
    if "logo_laboratorio_path" not in columnas_settings:
        c.execute("ALTER TABLE settings ADD COLUMN logo_laboratorio_path TEXT")

    # Migración de planos antiguos: los detectores que ya tenían un
    # plano propio (plano_path, de antes de este cambio) y todavía no
    # están vinculados a un plano del centro (plano_centro_id vacío) se
    # migran automáticamente: se crea un plano del centro con esa
    # misma imagen (agrupando detectores que ya compartían el mismo
    # archivo) y se les asigna. Así no se pierde ningún plano ni punto
    # ya marcado.
    c.execute('''SELECT id, centro_id, plano_path FROM detectores
                 WHERE plano_centro_id IS NULL
                 AND plano_path IS NOT NULL AND plano_path != ''
                 ORDER BY centro_id, id''')
    pendientes = c.fetchall()
    if pendientes:
        cache_por_ruta = {}  # (centro_id, ruta) -> nuevo plano_centro_id
        contador_por_centro = {}
        for det_id, centro_id, ruta in pendientes:
            clave = (centro_id, ruta)
            if clave not in cache_por_ruta:
                contador_por_centro[centro_id] = contador_por_centro.get(centro_id, 0) + 1
                nombre_auto = f"Planta {contador_por_centro[centro_id]}"
                c.execute(
                    "INSERT INTO planos_centro (centro_id, nombre, ruta, orden) VALUES (?,?,?,?)",
                    (centro_id, nombre_auto, ruta, contador_por_centro[centro_id] - 1),
                )
                cache_por_ruta[clave] = c.lastrowid
            c.execute(
                "UPDATE detectores SET plano_centro_id=? WHERE id=?",
                (cache_por_ruta[clave], det_id),
            )

    conn.commit()
    conn.close()
   



def fetch_planos_centro(centro_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT id, centro_id, nombre, ruta, orden FROM planos_centro WHERE centro_id=? ORDER BY orden, id",
               (centro_id,))
    rows = c.fetchall()
    conn.close()
    return rows


def get_plano_centro(plano_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT id, centro_id, nombre, ruta, orden FROM planos_centro WHERE id=?", (plano_id,))
    row = c.fetchone()
    conn.close()
    return row


def insert_plano_centro(centro_id, nombre, ruta, orden):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("INSERT INTO planos_centro (centro_id, nombre, ruta, orden) VALUES (?,?,?,?)",
              (centro_id, nombre, ruta, orden))
    rowid = c.lastrowid
    conn.commit()
    conn.close()
    return rowid


def delete_plano_centro(plano_id):
    """Borra un plano del centro. Los detectores que lo tuvieran
    asignado se quedan sin plano ni punto marcado (ya no tendría
    sentido, al desaparecer la imagen sobre la que estaba)."""
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE detectores SET plano_centro_id=NULL, punto_x=-1, punto_y=-1 WHERE plano_centro_id=?",
              (plano_id,))
    c.execute("DELETE FROM planos_centro WHERE id=?", (plano_id,))
    conn.commit()
    conn.close()


def fetch_categorias_centro(centro_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT id, centro_id, categoria, num_personas, turno FROM categorias_centro WHERE centro_id=? ORDER BY id",
              (centro_id,))
    rows = c.fetchall()
    conn.close()
    return rows


def insert_categoria_centro(centro_id, categoria, num_personas, turno=""):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("INSERT INTO categorias_centro (centro_id, categoria, num_personas, turno) VALUES (?,?,?,?)",
              (centro_id, categoria, num_personas, turno))
    rowid = c.lastrowid
    conn.commit()
    conn.close()
    return rowid


def delete_categoria_centro(categoria_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("DELETE FROM categorias_centro WHERE id=?", (categoria_id,))
    conn.commit()
    conn.close()


def _propietario_actual():
    return st.session_state.get("usuario_acceso", "")


def _es_admin_actual():
    return bool(st.session_state.get("es_admin", False))


def usuario_puede_acceder_centro(centro_id):
    if not centro_id:
        return False
    if _es_admin_actual():
        return True
    propietario = _propietario_actual()
    if not propietario:
        return False
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT propietario FROM centros WHERE id=?", (centro_id,))
    row = c.fetchone()
    conn.close()
    return bool(row and (row[0] or "") == propietario)


def crear_centro(nombre, zona="", tipo_centro="", propietario=None):
    if propietario is None:
        propietario = _propietario_actual()
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute(
        "INSERT INTO centros (nombre, zona, fecha_medicion, imagen_exterior_path, direccion, tipo_centro, propietario) "
        "VALUES (?, ?, ?, NULL, '', ?, ?)",
        (nombre, zona or "", _ahora_espana().strftime("%d/%m/%Y"), tipo_centro or "", propietario or "")
    )
    rowid = c.lastrowid
    conn.commit()
    conn.close()
    return rowid


def fetch_centros():
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    if _es_admin_actual():
        c.execute("SELECT id, nombre, zona, fecha_medicion, imagen_exterior_path FROM centros ORDER BY id DESC")
    else:
        c.execute("SELECT id, nombre, zona, fecha_medicion, imagen_exterior_path FROM centros WHERE propietario=? ORDER BY id DESC", (_propietario_actual(),))
    rows = c.fetchall()
    conn.close()
    return rows

def get_centro(centro_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT id, nombre, zona, fecha_medicion, imagen_exterior_path, tecnico, direccion FROM centros WHERE id=?", (centro_id,))
    row = c.fetchone()
    conn.close()
    return row


def get_tipo_centro(centro_id):
    """Función aparte (no incluida en get_centro, para no tener que
    tocar los muchos sitios que desempaquetan su tupla de 7 valores):
    devuelve solo el tipo de centro guardado (p.ej. "Consultorio"),
    usado para saber el prefijo correcto del código de sala."""
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT tipo_centro FROM centros WHERE id=?", (centro_id,))
    row = c.fetchone()
    conn.close()
    return row[0] if row else None


def update_centro(centro_id, nombre, zona, fecha, imagen_path, direccion=""):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE centros SET nombre=?, zona=?, fecha_medicion=?, imagen_exterior_path=?, direccion=? WHERE id=?",
              (nombre, zona, fecha, imagen_path, direccion, centro_id))
    conn.commit()
    conn.close()


def set_tecnico_centro(centro_id, valor):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE centros SET tecnico=? WHERE id=?", (valor, centro_id))
    conn.commit()
    conn.close()


def get_datos_informe_centro(centro_id):
    """Función aparte (mismo motivo que get_tipo_centro): datos que se
    piden en la pantalla del informe final y que se guardan para que
    también salgan en el Excel -superficie construida, superficie
    útil, nº de plantas, fecha y medio de comunicación a los
    trabajadores-, sin tocar la tupla de get_centro."""
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute(
        "SELECT superficie_construida, superficie_util, num_plantas, "
        "fecha_comunicacion_trab, medio_comunicacion FROM centros WHERE id=?",
        (centro_id,),
    )
    row = c.fetchone()
    conn.close()
    campos = ("superficie_construida", "superficie_util", "num_plantas",
              "fecha_comunicacion_trab", "medio_comunicacion")
    if not row:
        return {c: "" for c in campos}
    return dict(zip(campos, (v or "" for v in row)))


def set_datos_informe_centro(centro_id, **valores):
    """Guarda uno o varios de los campos de get_datos_informe_centro,
    p.ej. set_datos_informe_centro(cid, superficie_util="120")."""
    columnas_validas = {
        "superficie_construida", "superficie_util", "num_plantas",
        "fecha_comunicacion_trab", "medio_comunicacion",
    }
    valores = {k: v for k, v in valores.items() if k in columnas_validas}
    if not valores:
        return
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    asignaciones = ", ".join(f"{k}=?" for k in valores)
    c.execute(f"UPDATE centros SET {asignaciones} WHERE id=?", (*valores.values(), centro_id))
    conn.commit()
    conn.close()


def delete_centro(centro_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("DELETE FROM detectores WHERE centro_id=?", (centro_id,))
    c.execute("DELETE FROM centros WHERE id=?", (centro_id,))
    conn.commit()
    conn.close()


def insert_detector(data):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute('''INSERT INTO detectores
                 (centro_id, planta, sala, fecha, detector_codigo, plano_path,
                  punto_x, punto_y, foto_situacion_path, foto_detector_path, fecha_creacion,
                  codigo_sala, profesionales_sala, hora_colocacion, turno_trabajo, nivel,
                  plano_centro_id, fecha_retirada_real, hora_retirada_real)
                 VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)''', data)
    rowid = c.lastrowid
    conn.commit()
    conn.close()
    return rowid


def update_detector(detector_id, data):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute('''UPDATE detectores SET
                 centro_id=?, planta=?, sala=?, fecha=?, detector_codigo=?, plano_path=?,
                 punto_x=?, punto_y=?, foto_situacion_path=?, foto_detector_path=?, fecha_creacion=?,
                 codigo_sala=?, profesionales_sala=?, hora_colocacion=?, turno_trabajo=?, nivel=?,
                 plano_centro_id=?, fecha_retirada_real=?, hora_retirada_real=?
                 WHERE id=?''', data + (detector_id,))
    conn.commit()
    c.execute("SELECT detector_codigo FROM detectores WHERE id=?", (detector_id,))
    codigo_guardado = c.fetchone()
    conn.close()
    return codigo_guardado[0] if codigo_guardado else None


def actualizar_retirada_detector(detector_id, fecha_retirada_real, hora_retirada_real):
    """Guarda solo la fecha/hora real de retirada de un detector, sin
    necesidad de tocar el resto de sus datos (se usa desde el bloque
    "Retirada de detectores")."""
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE detectores SET fecha_retirada_real=?, hora_retirada_real=? WHERE id=?",
              (fecha_retirada_real, hora_retirada_real, detector_id))
    conn.commit()
    conn.close()


def actualizar_resultado_detector(detector_id, resultado, incertidumbre):
    """Guarda el resultado de la medición (Bq/m³) y la incertidumbre de
    un detector, tal como se hayan introducido en la pantalla del
    informe final, para que también salgan en el registro para
    laboratorio y en el Excel (antes solo se usaban para ese informe
    en concreto y se perdían al salir)."""
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE detectores SET resultado_bq_m3=?, incertidumbre=? WHERE id=?",
              (resultado, incertidumbre, detector_id))
    conn.commit()
    conn.close()


def fetch_detectores(centro_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT * FROM detectores WHERE centro_id=? ORDER BY id ASC", (centro_id,))
    rows = c.fetchall()
    conn.close()
    return rows


def get_detector(detector_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT * FROM detectores WHERE id=?", (detector_id,))
    row = c.fetchone()
    conn.close()
    return row


def _generar_codigo_sala(cid, detector_id, nivel, zona, tipo_centro=None):
    """Genera el código de la sala automáticamente, con el formato
    PREFIJO/NIVEL/CORRELATIVO (p.ej. "CS/S-1/01"):
      - PREFIJO: "CS" si el área del centro es de atención primaria,
        PAC o atención primaria + PAC; "CO" si el TIPO de centro es
        "Consultorio" (aunque su área/zona diga "Atención Primaria");
        o "HO" en el resto de casos (hospitales, centros de
        especialidades, etc.). Se mira primero el tipo de centro
        guardado (más fiable) y, si no está disponible, se recurre al
        texto del área como respaldo.
      - NIVEL: según la opción elegida en "Nivel":
          "3 niveles bajo rasante (Sótano -3)"  -> "S-3"
          "2 niveles bajo rasante (Sótano -2)"  -> "S-2"
          "1 nivel bajo rasante (Sótano -1)"    -> "S-1"
          "Nivel de la rasante (Planta Baja)"   -> "PB"
          "1 nivel sobre rasante"               -> "01"
          "2 niveles sobre rasante"             -> "02"
          "3 niveles sobre rasante"             -> "03"
      - CORRELATIVO: orden (01, 02...) de este detector entre todos
        los del centro que compartan el mismo nivel, por orden de
        creación (id). Es estable mientras no se borren detectores
        anteriores del mismo nivel.
    """
    if tipo_centro == "Consultorio":
        prefijo = "CO"
    else:
        zona_normalizada = (zona or "").strip().lower()
        es_atencion_primaria = any(
            palabra in zona_normalizada
            for palabra in ("atención primaria", "atencion primaria", "pac")
        )
        if es_atencion_primaria:
            prefijo = "CS"
        elif "consultorio" in zona_normalizada:
            # Respaldo para centros creados antes de este cambio, cuyo
            # Área/Zona todavía dice literalmente "Consultorio" (ahora
            # los nuevos guardan "Atención Primaria" ahí, y se
            # distinguen por tipo_centro en su lugar).
            prefijo = "CO"
        else:
            prefijo = "HO"

    nivel_code = NIVEL_A_CODIGO.get(nivel, "PB")

    hermanos = [d for d in fetch_detectores(cid) if d[16] == nivel]  # d[16] = nivel
    ids_hermanos = [d[0] for d in hermanos]  # fetch_detectores ya viene ordenado por id ASC
    if detector_id and detector_id in ids_hermanos:
        correlativo = ids_hermanos.index(detector_id) + 1
    else:
        correlativo = len(ids_hermanos) + 1

    return f"{prefijo}/{nivel_code}/{correlativo:02d}"


def delete_detector(detector_id):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("DELETE FROM detectores WHERE id=?", (detector_id,))
    conn.commit()
    conn.close()


def get_tecnico():
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT tecnico FROM settings WHERE id=1")
    row = c.fetchone()
    conn.close()
    return row[0] if row else ""


def set_tecnico(valor):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE settings SET tecnico=? WHERE id=1", (valor,))
    conn.commit()
    conn.close()



AREAS_SANITARIAS_GALICIA = {
    "Área Sanitaria da Coruña e Cee": "Q1569009B",
    "Área Sanitaria de Ferrol": "Q1569011H",
    "Área Sanitaria de Lugo, A Mariña e Monforte de Lemos": "Q2769003A",
    "Área Sanitaria de Ourense, Verín e O Barco de Valdeorras": "Q3200085C",
    "Área Sanitaria de Pontevedra e O Salnés": "S3600011E",
    "Área Sanitaria de Santiago de Compostela e Barbanza": "Q1569014B",
    "Área Sanitaria de Vigo": "Q3600373I",
}
AREA_SANITARIA_CORUNA = "Área Sanitaria da Coruña e Cee"

# Logotipos específicos de cada área sanitaria facilitados por el usuario.
# A Coruña e Cee conserva el logotipo UPRL/SERGAS ya incluido en la aplicación.
LOGOS_AREAS_SANITARIAS_B64 = {
    'Área Sanitaria de Ferrol': 'iVBORw0KGgoAAAANSUhEUgAAAWoAAAAwCAYAAAA4ujDGAAAQAElEQVR4AeydbXIcN5KGE9W2diNWPeScgPQJLJ+ApHQAak5A+gSm/+3SGyE6YsSZf6ZPYOkElg5gfpzA1AmGPIGbohSxIZuFfR8UUERXF1tNSZQ1bFQgG4lE4iuBSmRlobur4fa+/wTg0Ga4Uj9nYC0sRQJFAkUCt0YC1a0ZSRlIkUCRQJHALZVAUdS3dGLLsIoEbp0E5nhA76+onT2dY/mVoRcJFAkUCdy4BN5XUZ96554489/feE/foYHFnYPF4f8cPPzL9i87xKRTNf+1fXDv7v8erHYBHqBLT+VS3uJ/HywnWorbPLVLPnWQl3DSXaAf8CSgn/QX6OYlnhIXCRQJzJcE3ldRL7na73hzjz41saE0L373h+b8z57+Kb54439Lyq8yv6e+H3TB/s8W//jD7nXpepE5QsmOdtZGynt2Ufk961z1m3pLeQfUUVf1ZsDF83tli+B9UDm/Zbro7/C7/ePUX68+q4+/ivZE2SUUCUQJlGgeJfAuivrInL0wc1+bd38zZyfO+W995da8+R+VfhHA/tyr/r3eNG9fmvp5vnvf1ea+Ur/G3TQaR+g3fY8w+ufaicXLxXERQ5KiDUpV40R5rmMpQ08g5bop/CivQ2l7vbt23G1H9FPB2eDC7Sg2bSLPTP31kqHTEwog+nPRNv7y3S+hXaVLKBIoEphDCcyuqJ09RdkJtgafu1VzHthMikZKbMdZdXz++P692rtNFJOZoYwU/QnB26LpGtyxZ6YLZam+bRIr2QRvo1d/XzvMocloPmtXHZP38vGDPSnQoJzJ8VaBWz2oH5IGcFkoXjJzIc86F/UkqHx9T9lL2gB2UOrRyl9RGz++2n2w9XL3wQ6gDYb6T713W+IvoUigSGBOJVDNOm4p3z09ih86qzcv3thDLD0zW//jM1uWIkGhrJj54A4Q36/O+01nvldpqdyNh7qqDmlElurxcPsAS5fkODhbxJ2RoGshJ2boztxDkwUO7bUsZHCN+1KBVv6h8s7Od9emjhkXh8rtiPeIDUCxVWYobrOqCpuKZVeU4RLlMvK8o2X8RQJzJYFqxtGeOvlZxbvgzN2Tkh4J7w0otZDhbTUpy5D+yB9Yr4ZrJrTrfxpu759EyzVQwodcDXoSOEiAXznQ44c2mz35iI/lj/6XSItWOxSsUAXv2JSWqHNRLw/N24afYWPSxoEiXxjUrt08nNXhxWTos6rOg/fVMWn85sQFigSKBOZPAtWMQ1767A87wYo0s5WBt+OIW37xKJ+UnfA9q2uszJzlo+Ln/1h7JvfBspPPVw0vytI/bDcSERiDj75p4qquUKLkXIKUuRLPB3fcMvUJD2HQuFTOeBkYnjBE/ayuUN7C+kN0j6zTH1weXa6g8LtEWf1dUkkXCRQJzJcEqlmHezHwvOw6gT/HZXViGQaLUFblojf3SDxn4M7cN8L/9PBSPl+T31wdWbiobNXS1fFRd5Wndw7XxnOxrydlLDyE0c7aSIoemTyUv35L+Itu+cAYP4ISdn4PvtCfSCfyVjVyxaUEIQfeBSj92Wd2bJ2rJIsEigTmQwLVzMNsLMv1wD+Ob4i2JDDfKGnQhQwn/dEB63URl0Rs2bk6+IFrs2spvMEdt4lyNfnfcXNYdtWN+2NBA//SGjzLHUf/eFPviLIUlb/QyzBorHO9ePU7w+hPp++cpTZvG2r/6YiN4bJIwYoEigTmSAKzK2oJxcXjargJ3gZSLi9U5M8Lzm/JH/zbcHv/UHASN46j8CLwslf3lEd+CyjIy2wzFGRzysVOcJ3kyjrWdSr+s6hshU4GXla65uniTP5wKePQp9Amypg2al5Wmi1qQwj+dPW96bNeYKr9LStXkUCRwNxK4FqK2nv3gxRN+/JtGm6N1X3zgr2ihYFe1nnzP8ZsKT3/4+COexjTihzfqNyTv/gwB76sgj9etO+JxRiUtaxnTrHsSWDBMofegNsR796oY/HWVUW93zc8pn3Lfw+fAHoLPro9UPq1uVUxctYbV8gx/ZeSXu3WbeUqEigSmCsJVLd1tPiLX+0+2NLLxFUAPFd457trT17Kd90FygHQiZN8Xu+uHUM7V7lEIyb9UvWA58AJjkRPOOkunGf10QZnvekv8Er9z/uc11/wIoEigfmRwNWK2tlTX7m1t4ki8Ih3Kp8e3wPfVKaSWSRQJFAkcKsl8M6Du0pRH5l34Rt7qvlIgB9W0ViAdoS1CK9y4FM0EY6895d1SWlPcBRCkUCRQJFAkcCVEuhV1Dx2p0dycPlVJ84XQyOPmuFNOOkcoPMIDw3cN0feSBaYIgFeavIScgpLySoSKBKYEwn0Kurh9r7nNAIyAI8nJki2AI08lAm84G1mhkAXhK9zK/a8gMyybxzlmJ7a5YRF+sux0fC7/SdjX3yJvWAs4vV9eaJTRxhHZG+jVE48qY0U81Iw8KF4aVc8J4Im/7v9ia+3c6pE+ScXb/xvyEq4v7v9y16opHwUCdwCCZQhXF8CvYqaamT5LqOAwKcBPzAE7zQec81vakzluYHMIWeSnf9ZVS9zgoLjheoLX1LZSN+gVF4bpBjDMbg/qjrEbcasiHz1etLgdEcAuYRCPSjp8JOr3jgTfUg/6I+FH47yP6Xq2SA4Aqj0ogBX0pH6+8KZ+wYlL1oJRQJFAnMogSsVtZTIhhTXwdtk4r37Ad6pfN6+nKWuqXVcMxOlx5lkFN3gjruH+4UfQeJUxaB2X9RVNWYdN/y2rmbOpBg3FV87aMMaO0mSvnIu6/iJSQbGT64+vr9JP+iPXEHL0CxeF5XHcl7wlXuovHBaRf29pzE8VfkNrO3IWqIigSKBOZJA9Z5jPcWCDIrE7Ow96/qgxdNXxaU8t0adM84cuwsvQbMWo4V9hpIUeSFY40I+UFhVPc/x5SseCx3aPWU2L2iFpFDHbz0OXE09iVzif2cJlL4XCVxDAu+uqPWYL6tv+eXugx1ZfZu1OZTIJ6OsvdUoPUu/kbG4czD2k6a4daAlWXlzW+bsWVDgnExxfivlzRrz1DCUfz8B5aKlvqAN7Zg0QNs5wBP7siS+MUsf/te7a6Gs55cLIRQoEigSmCsJvKuiPht8LsWWiapRJuO0LPujo86c3Apmo2hN//7GllGkOaSfDh1uH+DqWEiWq2HBylWBMrXrXNq8pGiDf5qYovxeN3EO6sOO4CCBXB5P6F/O04PzQ1dhTD15hVQkUCRwiyUwTVGHc9JXjP04KcA831fWnnKw8QtL+2icdLMpKcpjWkjKlo1ETwAOkOXM17Rba9ui9awXeXvD7f3DlI6/DEg1M4F3bsxHTaFk0efWsPqA/9nhK4dHfT2kf+CqY0IZR2t7wZsPY4LvdkAZRZFAkcAsErhSUUt5PEGhXFHJhDKBz10YpxWs5zqeUlcP+/uT6viyUFbrTl5bUHreNkQ7HcnaDi/oZD2jvDXmwwDe808rR+ZtI/CL+V3DSG2oLJvUemhLiRQuBj70zcc/BxCd/0h82OWLv7xn1vMPMCpTQpFAkcAtl0CvovaVW6vij+iDmx7pO3JY6v7halBoVaN4cl7K1/h/RQR3zn8r9MYDvmYpXX4UaQUrmf5yxO3ijY9WvwtKkh/+V2eCKwd/e4LU5/pNnfuqlzkzngP+ZZUPAQs8z+MMNxl1M/4zWeyHd7d/2Yt9OWYjQLbpdMiAn1RVAfH9OpRlT12KT5y5b+BjTMouoUigSGDOJNCrqGeRAcfypESeoYxQPFKAKJ4vZyn7sXhQulFZL9PfoBj5PWrv/naefgzJ24aU4LNRY/m2XQuuCL1U9ObwXyf6ktKPcsj/4Na8beR5cqHsUZC6pKxXhctid9/Evix68z+eP77f1j9SHyIfbqcV6lIZtTnOJ1oJRQJFAnMkgV5FLXfBQTyuZuAooCtksm7O/xwUj9lSHw/lZSEGhQUeeftYb4T2klMpu/eXB7X7Qu4XfNSryYKlwUB7fL9VltASnD++f0/5wc2jeFVA+TEIZ6L/vnbYlydaKEt9KGulOR8dfNPCl1/tPsitddjs9e4abqJlWdd/5QlEfO5VD19gLh9FAkUCcyGBXkXNyGXNbcpiPgRvQS4QLNRZQGWwChWFEH6gP2B/0gdnp/+kpieaTX1Z3DlYlIyfAeA5I9Z1cXXkErkejjwlV76uP8rdU9erpXAXCXwaErhSUat7WMgritvgvD/BQp0F2kINsqBorC6lx8I8JuJLwnVZznsjuT3mUQY3NeYo2yVcSWljvKm2Sr1FAjctgWmKeqJt71z2DykT2S0hnlpA0be0gkxKoHKmF5vu62I524e/OCGjdxGv5Ur68JWXGosEPq4ErqWozduXw+bLIVN7mXzSU5k+QianJnJIZ6rzpqHlPAmHnvMlHDqnR4bb+/yantwW4csyKbuNqSduWC0NZCj5AeD4t8/jS034qRs6j+qkW/jul62UR34CaC3P9i87CYeeeOY1rup61cU/NM5lwJwMm7lj/g6buRyfw7vbv4Tz9MOcTzhl87qYJ3iIc3ofzkt3eAW4Y8K6wT3T5Q1zqPnu0mlbZQ+7c5voxJQhho+YusCnAQcBKAfEPnI0leQYdOp4Rrm8/8OwrvfHXaWxhuH2fjh0EJNXRnc7cm/S43ND4aHmog/65mEY+zWcKNPUO0nfP4SG/HraiuM+aI8hDzv139UYunNEPdCGyGF7f6Q4rIHUBvnAXcoKwLtwPUUdSvs9JjSgnQ8mjoUv8ifh5pCf/ZFgU7AqeMTLzKEElQuIG5o8ATwtaAwTAWFThzas9GRxjx9+gp4zk1Z9j7Rh7eT0Bvd6cQk0qfTpzT2iL6T5NiNpnmACePcD7Uq2x8gYHgD+wGeu7bfSq762kc3xhYwkhy0APBeFq4ybbMVciDV9tqqPnyTb9jfXXfNVfc2tjV2f3zE9AVl7xV9ZXEkv3tuMDsJ64KW7yPyKI4pQdfuf0jdjRQ+BvqrPj/TCfQc8EONH6rfWwZM8r6U347I8nZ3Pj7WY2g1g6apddZxw9ZGX2+uhvy2xRVaENS/HnS2rjz+EX4QUkeCsJg8eki3EutZj3S29D+nKXWnuFeZmbN2rLO3QntA2jPpcXFf1y19+OY+6kEtbEUj2JEZ+09YM46bPmqODXEcOpcyhqV7dp/6J1t4huPTDYa6LVPYeoLyJUE1QEkEvDn3l1lIyixck9J9Z2Kkz7GTssPGI3kbG26Acc+uvq8m/wU9nPnxxh1MfLp7hRkD0OW/2fPc+pzpa6HVH1HVQ0KqL0yDwLiOj9Ce4qT758rXAQmq9206gzvjhndvi5Ina+0Lj+N68LXNzLOolZF5Ft+/ZIsvZ5gavf6+R/4IGvBBxoeMhyLaZc25Cvmi00ZErp2+Y4xZGnfcIuqloR9PiQjzewmWKDZWU5nH11e4D/sdzuTb3VXfdZGf2Fy7eWFhrNnkt6T57MkmepHC6KV8b4jgGclpa51FhoJRMSgWFLdbx4OK9xJoM95KesKMiWDhMSwAACutJREFUHmfMUlldK7GNLLcXzeW+6M3/KAF/Gd85tAXavjRzyBxdJa9QJh8zeBp3yDTL26Qubd4xR1Hb1uP7UqT+e/rTHTd1AoM7jvV0Zs4HGTZryu9JOb8gL8z/4/vpt5FMuoiNW61MD1cqaimbk85gxmvytqHO/AwRi4IdVviSYDJ4G02ta7LEB6ew2+JqkGJlQhfSr+tdpyEJK1iqF87aHZhxUbfFKyhmyUaTy5dtDNnErHeOqP/l7oMd44szujm6lpiVa0wCWovcJM9FPIq40KuD5uqY3Bl+bwW2AFhJQhaCwjJbimmReoKzsG7ytfBavnPm1bLLm9vUDf1UpCMzr/kWNhmUZ+tT25ss81ZK+uKXZMG6fauBUbsqyMzVhmKyvivcC2brsc6z2EYf65W0V9rYlHnk+NKXkD879DyljHVpxGYe3j015LjhLljtdkJeQw7HcL02PiWXZtnApHvE2gnaGVyjGMzAo6A7XLMlKS8IO5RiJ0XZZ6XPVtkH4EKxUo2Pv64HDuDPS3DVTVA139Zkt/x5KH9TXIgUbyHdjNWdil30qUe5trnvh5xHf3ay0FJtqd/EQz1mJfo8xtHSWTLvZHUCtpSe/HJ58IcX8CKvOEenr6U8M55wpHTY+jYPNrM8k5GyJaX6gs3fzE6lWMfzRUyh+rx6YjxVyr2l+k5oN+WlOPZxyTv6DNhSH5/unx2VkbL2e33rT3nXDsHq8/bQm38S17ildZxXpr6FPxOhrzLk9sgb1HZoV1zRNWTxXuAPOx6Gtq7gv4os/RPayBWa5mxTsgz+ZOI8r68e7o0Ew8l7ZOpc5+O2+O3r7tMQcwXQhsmY0nyf0A/XuIRs8B8WxmD5xQtvpSs9KSuaGqq+3GFYnM3CBEcofXwtzVmwGKznovzd6CAfqt40wT2sH4+km8aZa61iGtYYVy+hHssjH8ACGtQqJ7eQ0usXlf/XsDPpqmNLk/SU3VMT/ER8C10e0d4rqI2x/in91r6/V4P/RoXjo/Zp89i/hvzZWCeUqCzt4Pc38/zDzqg2x5PWTCPlhgw3I7+yqBKyqlFaK1cpTtbC4HO36sxjrS6qjwdDNvrchdU8Kp9iSMQN+Uz3ykS/1ZxpDQb6xcDP9NhMmWkQ3UMLn9XVHmtcvM99n4HhbYO+a5P62YJycV9HfhWZDC4+IYxkZdaNrK50RU2WnqRE//tkxgwUjYf3VQE050F+MxRrWLy9ddzIBaAdFTrSfG8pbsNIMmgTXSS9L+nSs3Svolb+iiZjNSxIsxWllwSTQQoPYp2/kIAwDivOucu62G3G8z9+ij5kjyd0QNb+aoJXzeMW5AlgYcpHt6mb5Qsp5Bea9J/Sbj5slPaCeVtkZ20t3+YmnKjruoSkCHTDH+dlU7+Jp/U9L3Mb8Sifdc3LCPkDwrFsJh/lm80WMTyX3JZfj1vT0Dt+y6D0oevW8OFG91bfo406KC0zKc6dwNDzMdKN+lLuqwE+zKbt9aggLfab++yE+oDQb28bMW+sRtagebdpWsdSDle2OVZoSsJ7t6XsM6zo0LYSCn0GBu6kM+WdMo7z+ISn9ETo3gsDq9kIz2JbE/zTCLXZoulKv0QpVOJp3j1p7sJ92zN/sLUgvvwbxeEJv82c8FFfznXkefu4mzk1bdrfqq3VkebbsivpiIxkLrqN/OWLzTx7DL9KUUv/WLOLjLGPJ7z3wZzHClAOE6ioJzQL6qAn56OTeOkZGvUu9D3g7/DBzeL1so+iEmK0cH24gUVb9+bi7q2Uxt+ZqJXFzJrq5KlAf+BGIsdf/toeyQJRAkk+Jnm38hdOdnoMBwc0d7K23dfC13kxrnimEObN2wbMTn5T2iEmbXIfhPyQ6P8Y6QZmow+52tCJs76tUB+gusLv5rRjgjEDnhisUQ4o+CzneihuDJVYEizQLiB8XWDWMTCc+ePaHEpuse+ltuXXZdn8XlgQS69LR/TesKj7xGGZm50iu16mGybOMu44p7wT+SFuUqFXdfwVz7hRBdrlhw/6It+ALvPGMemYccJ1UjwqJX4Nhse/lPxkYt2QjV9N1q5uyGPt6D+wwLvWAE8PObBAuoO4u8352oNN8lCusma2xHM2uGPPSAtfiTtqu3sP7ri/im7pRYryw+OqFnrwMVJOb36D7LxVWH+wB8j9qMPt/RNv7lHo+z/WQh2BSR95v8Hpn8hzFyQf5uO5LJpW/uASBC+jwk0hvA2sAd+cKti4q7ltM0D0OIosc4CcTmbwREXdCXxzqqn30R7jgPqZa+ogTay5HDFXQRFJ6aa6Uqz8FxrTRL9DWX3Ex+tToe8enKf+s9RminU/f89mkfqcGmgsV7dFntZwr7ETymiD1FrHumznIt4LZ/G+SVWOx5nch7pnL974YzEsWGP1C21Cuq/T/IQ2m6zrf2Ztpvq6lYyN+41/1s0nrfE9ZM7MfPuUHYxYeR687t0078x5fHJZEf/TkTZvygfo6Qv8Vch8lw8tLKzKVLTi5ZnZ1VZ1YvzYsbcNLQxZ8/4nax5Rn8cFPtYTeHJIj6U5kwv/GuN/0uL5Tcr1V+Wt60Z/MpKgkyKueHGkjBTIM8lKbW8EgZOviSONjzvWEyYMxZHKEft4fpqJVzocVYo7t5KXIe83eHzTfMlw09gnUD83tbrBDS1LWdhYcNAWIs9YzivcXJofJ+t4LB9FU/sD5JkAZaAbDsV2lK99Kgw3pNmp7ygU8vQENKJ+5nq4va8HUYex8IK1EudqQYqHPsJ+CY1f98oTJayt+hq+9cuKGyy6VcIabiiXn1Xz4rw1MC5zzJp16r7WGv5Sxs9Ev99yL6DkJl1RqYFM7mndm7mvwxOEZZe3cF+nuUG23F8Zx+xo1maqj7nuVtCO22ylb9wj6YGB3kWo3Kn6056RrnFT6Z7X2pC1ve/RH1pHj8Q34cu2nr6gi95dUXvbYNEloHE1vCD4ZAJWTgIt6K9kLSwKHo4k0NRJFmTiyeMLX01YC5SlHiwFLA7wcKOrMk3GHum8bpFDGFy4HeomQb6U7T3z7m9OL5cA8kRDAcAS/nkGWgLqVduLqa3ApI+r+j6Y8iZexW5lYMzIa+KG1mi5wciDh8dMcGJlhYDsofnoK5S8t7ws5C58XttINN5PtHMVKogfg9qtkh+TbUT7g/Bk5b52mnPmXm3eYy0MNFcqsxYVfVsGZKAnNfLgob/gxOQleC3fuvr7VaIT9/HBL74tADzAf4bxrH12p9oJ6eyDjUi8X7GuIVNnFZU3acYELW0w5JEmjzKUZXykc5AiC7I1tZ3TwVUm5FFPAtY9bZGfIOV145Sfx3m/cnrCu3Wk9Ofxy02kqSPx0xdofeOGhzEPmnXwkPUCjTlivpl35h+gDo1tFX6LV9/44UMXVZHnVkYs/gQIq2+QLMjEk8dX8UPnSNZLvRgCT3WCAymdx6mNfFJQKNQB0O4YvzYSaAneVm/iSzHt5fXNA86YGf9VYyUPHuYg4JJxzgsNgIa8wbtAeWjE8HUBOvldOmnaPdfLt5daN8w9NOBtZagPHsoHvNNv6qC/5IMTT+ODFz5gGi/58ALgoc5/rp2AJ4AGkKaPCacMAL0L09qkDHXk0C1POs/PceomP4e8Xzk94Xn5HE91QaOOxE8MDQAnL+GkgUQjJp2AeWf+gW4ZePrGDx/0W62oGfzHg9JSkUCRQJHAzUigKOqbkWuptUigSKBI4INJoCjqDybKUlGRQJFAkcDNSOD/AQAA//80qCxfAAAABklEQVQDAB0QjKKn2VLfAAAAAElFTkSuQmCC',
    'Área Sanitaria de Lugo, A Mariña e Monforte de Lemos': 'iVBORw0KGgoAAAANSUhEUgAAAZQAAAAwCAYAAAA2P4zsAAAQAElEQVR4Aeyda3Ict5LvE9W27kRc9ZCzAlIrEL0CkvYCSK+A9ArM822GvhGiIq5155upFYhagakFmI8VmFrBIVdwmqQUcUM2C/P/oYBiFbq62XpTVFUgG0Bm4pUAMgHUo4vh9oG/BXBkM1ypnjOw9iy9BHoJ9BLoJfCJJVB84vL64noJ9BLoJdBL4I5KoDcod7Rj+2bdAgn0Vegl8JVJ4P0NirPnX5nM+ub2Eugl0Eugl0CHBN7XoJx55/ac+ccdeX921PzO4fzwvw7X/337jx184qlS/3v7cOn+/zlcyQEeIMendIk2/5+HiwmX/JqmcqGTB7QUJp4D9YAnAfWkvkBOSzy930ugl0Avgdsogfc1KAuu9Dve3KObG/dpOVDuV3/5I3P+d0/95F+98f9KSrowv6u6H+Zg/9/m//7blnL8cPtghDEY7ayORNu/KvyuZVf5ptwS7ZA8yqLcDGHx/FXYPOEuKJzfMl3Ud/jLwUmqr1edVcc/hdsTuXe9BHoJ9BK49RJ4F4NybM5emrmfzLsfzdmpc/4fvnCr3vxTxV8GsM97lX+Vm+btoamel0++d6W571Sv9vGc2hHqTd0jjP579dTi5WK7THmYLhmEoPzVTpT8GjsPoWsnI7CpyHEzD8Xt9ZPVk7wc4c8E54MrtyPfZOz2TfX1kqHTjg8Q/oVwG//+yx+hXMV710ugl0AvgVsrgdkNirPnKGXB1uBbt2LOA5tJIUrZ7jgrTi5//X6p9G4TBWpmKE15n8F5mzddg3u2b7pQ6qrbJr6ilfM2evV/V4+aUBGq39IVJ9Aun6zuSdEHIwLFW0HYykG5ThzgqEr+gpkLNMsu8klQ+HJJ5AUZrB2MT9w1LauMp6+e/LB18eSHHUCGkPzPvHdb4u9dL4GJEugJvQRugwSKWSshI7GrI5gjjmiu3tg6K2czW/v7G1uUwkPxLZv5cAwkvj+d95vOfKdyVbqP7sqiOKIQrfxPhtuH7ByItsHZPMdYCfIdR2IG78ytm3Y04F5rx0FY7b5W9IVfF+38UsZH/kTH0ZbS7Yjh+OLXH6K8DANjVhTB+FnjijJcIF0D3Qd7CfQS6CVw6yRQzFijM6f7AOKdkyFZFIwU7nQo30DwtpKUeoh/4h92A8aRXCjXPxtuH5zGnUDAhB8dMWlnVd9H4b5HwMcfGcVd3cM40f2Sfwo1b6XDECgo5x3GYAFjFJS9tw0/gwGVgcPIzg1KVxs5Z2W4wR/qrKybzvvihDj3dfB76CXQS6CXwG2VQDFjxRa++dtOWZWb2fLA20kMW/PiCCcpZYV3rSxZtTdZPk04lnL5/1b3dWy06HRPQqh57ZyOaoMnBG3w8d4JflEWKHso1yCjo8iLwT23SH4KBzeojtLOZXQ2w45N2G/KAiOjULeLx2Jr1IejrpwrGKYcqV1UjurjvQR6CfQSuI0SKGat1NXAc9P4FP5mGIUqXFhhm+5beHOPFD8n7Mz9rPBndxe6J2G6r6OKzF0VtmLpyu6h5EreO7cl1heCtWQ0FA5utLM6kkFCJuu6n7Sl8Ms8fWCMP8FYOL8LX6hPxON5Kyq5cpQIogncq1L8m2/sxPqrl0AvgV4Ct1gCxcx1q1bqa4G/Hd4QbkFgvjImBOcaYeKfHNgNzO8chhvzFO5cGe5TlGZvpZgH99wmRsB0fyg/MiurY685NfyhVWGbdP39ptwRbSEaKQWv3aDa7ZypjJ1hvN9D3XkXxbxtqPznIwzYdZI+1Evgc0mgL7eXwEQJzG5QlIWLj9FyPHQTSAm+VJLP55zf0v2Kfw23D44Ep9HAHYcb6te1WhINeg0o8muyGYq8eqrNTjkyaxqVmNeZ+M+jUVBw3HGfxVW7tXPds5HRCHUKZWI0KKPkpr+ZDGB1v0d1r+rs7KXK37L+6iXQS6CXwC2XwFsZFO/db1KI9U3saWGrdjGfrfkD3fT25p/GCkg5+6eDe249xuU53vDf1f2MoybwUiL3i4R7jC/GYFS0G+GptV0JLOx0wFfgdsS7O8p2EGVRkO/jisdkX/1j+ATga/DxuAvjVJpbESPvynAEdkL9ZUxW8rytv2aSAMYfY54vEmZK3DP1Eugl8NYSKN46xReSgPsZr578sKWb8isA4aZivnyyuneheys5kA4Aj5+a+/rJ6gm4S6VLOHziF8qHcBN4YivhU5h4DpeN/CiDd2WoL/BK9W/WuZl/H54uAYyJdpR/atGz08twuqx6ai+BDyWByQbF2XNfuNWbCgo84p3Kp2ObwDeVqSf2EvhwEuB9Ke32XrZ3pR8u/z6nXgK9BMYlMMmgHJt34Q1yJTkWcJ9AXsuBO2b1Da8o8Mkbc8fe++u8ZFzGOHpEL4EPLAEdFW4Nrtz6bdqd8Mg6R3AfuKl3OLvb2TSOUOlHgPCkWqb+hm8Sz13DdxoUjlvSUQxhnfuPvZ8BDhoCgTeFiTcBPEc34Aj76lFcoj1MkQAD9WsaiFNE8U6kqzeeD3iOjVtkOtwe+5fSfY7IUkGi88DE2D+Zkjbx4BMXr8cnPg2Gvxzs8YKsjuAOSSNovxMVEwt/Cm+M1h5liOZ5erFGKpDw+IoaD3nAR1j+vmCsHU1c/p24+9t/7A63D0akz0H4dl6/HJykcuFtlk08AWOZtHlZid70xZfLfoQ8UM4ZX7susU+bPCk8qV6JrjI9PCme/Bwf2kE/vvH/oh8BjTMe/GmNn8C3fXCU+hs+5TUaxic4U/530e80KEN1ThIw4fiEVKv94KAxoOAl3GKIEfCCI6LyPcIl/KmACahym4N01DVAqQ9tEa/PBy804ckjtIN4E1I68eSDnJvrgTUMMg1G8ZwKKj5NyGE2yFBsw+2DUwYqslLY39ckD5n0PzNJIPbfspiXY1jBzOmYVouixzz4IMqa7rfkn705g96EsBsXc3Lqn/D0XXwXK6HH/CF97C08/s3Rbyxzkc8WNZkDn9mCeVtnvDRpddj5vYltqpmqgMrabda/wtpxE1dcFa12O3Ob4puLdVFwzNXpVc9FyWD/pvqED7UqGz/7N+lq2Zsz3vXakHI+YW5Y+6rrQpvM3E/2kS76I3y9vNGPxpc4NI5U5MrAyvqBn8Bntkyd1AerVvFJF/hn6CPx31lXTGqZdhKLKMpJ9ITnQ4fwpnin76pvZnXSPiIyTArnf1cRi0xiHntOAzS90S9a7TQ5goL4uyiDXxNmDWhwMYgSaCCFfLLBeEQ9qI95mzfzzyxe8EmxHSkqvHGEeGw6InTmfh7+cjC22hZf7zokcDXwO0KfAzGsYNtpzIaHMtg9++ppwIVMMZ5ePPkhfKQz+c0cIu+acGfqx435xjtPwrWci5/W4RgOo0SZ2q0vEm4xmkeZc5Q8l79I2+Cbk3KdaSyQf6o7Pnk480eEEzQfPAnzxWxOfOfmfBi7CrdcM73Fl4Vzw9hKoIivDAntWphRoday5yGVki+Fq16aG7vKrnbNutCey8YDLjXTBwoEo+jtIXOXOiFbvpxBeFC6pbIomLcWZAif+ccXGj+JT32/oqog11YbhLtTrpjYGlliKdjDifRI0GD5zcQbo90eAi79jXl1J343bDXh/TMUsm7MLjGJ+RhjHAAP0gBIuVf8tqb4uRQ4E1vBt3NJSTGQAAYcOWi3sWeSgWkFRfnUg/qgVMBZvAKfJo4v3Lpo4ek08S+pDc+VfqNjhRZT9l6SwDyKXSt8GYk9yY3V7XrAJYYO3zk3H9D/ZiOb8UoLEim7sDLlv3AmJfXxe2ysXCf1YRx/2lW5HeVzrIUGvoKZ06JFmGVOBeR/WFcZEZXttjTeHt60oHRFWX0hY0otogFZ0JhmTp3JUOFPSTFOev1k9UR9yeP0ks84/VNgvHf08zlzNy8Po4zhCPj4ZYviXtEyHKOd1ZGvvvWXL1xCsrvyU7xnQ8LWNHb2+Xvm9UGTp0+sSMlv0ZnNzFsDIBKigjjXwGfgTNvyxxRv5bE6eXHZsYLKcEvKtXrQQYHkyvgW/sCV5JPQvd8hgbiyn/umLHaj3ObC6jLj5ZgKpTzUWbeU54bIx9k4WRatOprcPhg7Y/fmthj3QdmZvVB8oqJkYeG0YlU5Dwvz4U/TogFRsZWLO6nzQfhqgtsTNnx4VH7LaTxDo7xHk4xTK8GMkZBXtejZq+pg58goT652rgS56fhWSvY30c+mfhaoMlJnKFzJgLqP/Y+Q8rjRqS46MjJrGjnV5VGzj5q0GzPMGW6OMzdPEhv9R3lNCDRvwciOZEAsu5wVJ6brph2dWL5YV7xzzbVS0ip68ULbOq2iN0teyjO7NUbFW8kAsDTYWaU2O58wuNR+DU4UxD4D33TMZNVESOSZfHZ0zQFOIgae/DlNpjCYFA6TgvITwBPrsiC+sHWGL0FUWtJHLrQp4Xu/SwJ+R9hjFg1BbupLKb4t4dou7GLcIyEXJfPHg9ZLr8KahcUSNKCMRxpQhtwTYSfpHArSzAd/IeJhGQPmiY5GHpjmjXkL9wTodxhD36s+ou2PpIguq4XHucbTFvQcVFeM15mM0767smp3Ze93hcesNX8pmzpQF+oZ6tbOesmb5FbV97natBL4bfyK7VvWMVFYrRfx46t/v+uRsorwpY3s+mrdQ0kvIl+TP2ioWa6xYFX/HDbhptKcKz9IX91Uzuekv6tBOdeZYGuwh8nLqu1ztqZRtjPXWin89ca4gdgaAOmT8FERzMUVLQpiV5Ppxi2/5ZeUBconAeSu1YgG4Y6grgtn4tQP/inABzdDm6bwfNUkDLQEsCBYHGrnAagfmcTjq/3q/J8F0PyVFfsdSrE+x7/QoiksNJRxcHGxoT4Mn9G5Xnx4FH1g6frByLH4Mu9+FH0u7oot7aqEqz8FpDCuczVPXX3hKGtBZQdlDfO7QjAa3til6R7AQXr4JCxe8qM8je1dGRs+qzTnZVBp06Ryk+Hw5jbpC8a5eMORcihTkVmd8gi780rPVKlUlyP6JsG0ulQp3uuXHVKQCblcyuhrQe0AyYPjONDmzZ8QiMaUYA3eXEifFrk14Q4FphmUM7VT56n6HXcnI62kcrQvDKFbx8XEnZRXB/v7ozTYTsglKpnwN7x0PpAGQN2xUUFoxccjk0eapMFYOu+ZtGQzEzDBLqR8EpAoleHjYAKnOnB/xA1K94C46nqUJoryGDMacfLNpcFKmo8LX2buUvD0mxSW57M6R8hVwC5ibLXvBzYqzQUlpdX5XpTxjQ2/PhqyF8o7lXGkhHyVuvOpsly5DP6XwW/Xl99R+Ezjbb+R565wlowO4SZg4MT7WLgFwXu5ZDSUH0+FVW3ynqe/zryMQZ65FpMrpp2f5N16XLbJhzxdlfa42S6lI99pDx00swnhYbUjXFbaWnEHwof7OdO846i7zjHe+zEf738550O980efQ99WuzWMrFn8k7y/s11Yk2/UoTvrgr/wwESDosG1h+Kb0L4xpQefm7z9PpmSF0k/OJTxiEKDnsla589A1zKC1dgZ6J0zOAAAEABJREFUHVsrCGfP1ebmZDqGL/DXqd8+QBlKhTFdC2UpktxV9TRSPWiFf6Ey13O++KXierCKr3eZBMKENVvzuvGZDHryjUdPRYs8dUqMeDAqunfADfOsrxe5V9AE6DI+GC3j2Cnlj6/FQcCnPq0LUQCjMNw+GN3f5h2Pw03KEtpQVnHBI6PgWk+Ukad4jv2UXX/gkWIX33s5lcHC6Zj8mmDmdsxs7MksxvTgygUFrEXYUT5elabedfmi3a6wQzM7S4s2eDuglv1w+0CK3D8Tz7nKpD4KVk71Dvdz6j765Y+titL9W/Nt/7FDOMreNO/DQzOUhSEJRsP5sBBJxr/4ttgT40vv3W+8dkBa+K4Kf6LS5mSQQtkYesV1FOd+buYnPhYRNZ942A3yCoHaR+xuQNHVDA2C1SKedxKWIJ9nfAsIs4ljslnhWx0OnfSatEHYhGXp/wH+YwMdq4HCCo6bq0fUl4Fw9cafVmW7UNeoIMIRXnMypTqn1VuVxuqBzoAEmkpKK7FNcAkYnKSLeZ0z+e5LqcS6nMh4bCBbbtrCh5LCF9+fQx3ZkI/8U2fuZ/hoE/QexiVQDsqg4L7RzficWsaHGlDsOQ2jYuZ+Ul881NhoTu4Fb+5RE8LRlLfQZ6NslRmPW8KCwLKrrBY3Wqw4KW7/zLwtMjbpdy14mBvxZrxll9sTYm5YrdAVHHdSsrSbE4Bx4gyYOEYxaJTVSjEIDwjYuZS/6t0iGe0tza2A1RwaS6s0tCvcjIenCWp7UOBdhijy1bI3Z+GR/8E9t0iZkZ685Wb/eCn75nxMTMlv8apvJfugA5j3vnp0fE31/p18lCbsYFM/4w++De2ljzeU9jDxmY4wm3NTdV1XvdGZzfws8SnvcB9VPnJvjjmhvmzXaVBmaRLCHG4f7DMgUZCajCjIh7Ok/VQ8DBQNXozKIvU1bxsqW/V0P17qDFRhdbJwWsGOMgURFI1WfxqEzclUD3Thg7JJiizllfD4Gpzh2IK84uRDqfwc6zLPII4rtpCcOkS+MyHCZJGvMv3TJp9wvcskwOOc2gW7DqVTH3de6DiSiQ8ffsriUmMBnCAoSHxBOB9v+6vs2t2kvhAvj3tzzyZlHXzKSjT55Dt/obpAVDykoe+JN+GyrtfqXszD4Td5aK/ymE/4C+WruGvypDB46CmOj1EDf6myiDeBOok2L1gHL99dKH/CAONauHnJI9wbgKZ4KBucwp0nGYmP9OTTBKVZESCjCn79PjzyT10yvor+5PuWjzyafIRTea18q3Shv+F59eQHPiTrNP++k0H4D/Euvn6yegItwUg6Qnj6y/nCrSY+ZJh48APfr99vireVX5PPVcfp55cdciePLxU6DQrWN63mCJs3FHFXG5sWWNZ2nIX0WnEHxUo4KNNxto+GudAEUMcuDkr3QL4TrDQ7VvHJCkKDWfQwKeSvCEjfAhQZk7mLJlxIS+MYnIqHwRjrsvhKgxhaE14/WeV4cJHByqBVGveqg6+Zpg/3Eugl8GEk8FrzbyTDcVNuzPlZ+Kbkt2RayN5UzpdG7zQoNEIr7PBkBuEaqvsM/K/HjaA0rLLlBReeXgmhz/TTtXL5TFUJxwWUPb9zOB92edrpEQaXgMHKoE3xO+v3DWtJ4P52dZ+FM/qhjj1bxD5yZyTA8ZmOK3fuTINiQyYaFNHZcSzLr522aa1HKVn9T4I6URWYk9fKS/Gv3sWb7WvaieyOZlgVffUCu+MCuC9j4szpiNXvsJvXCvb0jjf5q20e8/02LXI/VEdMMyhjZXjnwlnqGCFDxJttGKSM0kebEnDhzVn3U78Tsf6SBLwVezoOXdIx56IWGau6DyHjIkLvegm8mwQ+eaq3Mijm7eFwyhMnqfbpnkmKfy6fp6SawDFCXhdwTZ4UBp/zEgfP02JDHUcI9ocT5EE+0bCSrAb4ARCXuiEHEIafvAnzpArxGn75YyvRoCcAV/PERyGJg088X6uPHICb2k9fDNWXyLyLl4dOoHfR7mtH0VUG/T5Unk3gwZVmHhXtsDYYxDUZl9KqNS0ywA87xlis141PCHXVZRjr1qxPCt9Xm4Y6gk3x3KcdovMIbE2KZWgu1H8LcDrM6kx8qHzzo92h+IaqT8rsflV+eLlyKHwCykg8yFx4lRHL++Vg7AvMorfyoN6p7PsTyhg2ykO+lNnCRfp9pU91afrgh5En+pLJdR/DOylP+KHncF9lDZVnjk9xaF3AeK7L+q/Dzo0APFXaw3ockm8m386vs2c84Sla5KsxTBZvA34XYXelIEOUrWi34njLm+NJrE35K4JHHCNIgCMErToGV5RloImOX0MgZj8oavKQYU0dtGTmn4FvshJXfo9kWDvOSL06D2imUC6qK3UBy9v1XnF2hAG8+41yJdsTZAwPAH/gM1fXW/GV7PMUsH5VgIwkB/r+EeFpjXfVl4CX8xfR6jTO7yq8LGg5JqMz97M3t5WX4QrjSa9lc8E3XeEpw/tSDgontxzLruMaBc+aYzMScr4KXb2MuzaUQq4Q3b8ddQmMPr7RHSKNH7VpSdHOfGmn925H9GVBcNRX4/xPRTQG/VNn/rHCcpoXjfa6Ss5rV3955Cl65SK+zs9VLwBTh4oh/n57z8Lx31BGyZt7ZGanjrJ0X1fzsfUpG6su8qweilE/qN6/qexgCAtnIS+7vq55I27g7WSS7Fw4WbCxq6Puaod/pjofITsSTMrTT+8P6kfyLoBWtfOaOmJhcl2Wb8k8scV3ppZd1TcBTV29teS7n8t3qDGXeFTvp0oYxjdfHplsUNRRbLvFnLs5c/53Kbe9ZFiYXKwA4qPD40+E8fht4VbzjD5F3JnnUc8VHSU8SO/AaAIcUedm+TpmWGlCWiE2eawsgyFRXhxLwB+OJvJvCLnqkUCSdn46A8Is4J3b0rHHksp7oHY8VscuMinS4Ex5NOtNmCdLEu1r9Jsfg2y+RzRNFlIGm7lcmThK03l0G5+CPBd9LryfokDuQv89+T6ME9FeqIyf5U91vNOR1yNPgBIXDkUiz2uBIu8G16gL9Vl5deOTg34nzzLKcq6J11xil3RealFDnhfxqUrxhJf78nmmMbyR9IZ4JjmedAz1ZDwD3HOI/bEmQ/0cXCjr1+83o56aiwqyzlNzJsx95pDnPROdsCA7nswkfQIS1LxVf62gkMEDuewudbIAfgI0675IXcW3HGWnYOXyPJFdRXn737zualfQU42cxr4zN79zOK++aPFF+S5TZ+WxckFfdsjXRQMED/WWH56i5asgEw2KlOJpp1JNtfS2gWEhyuRiBaBw5+Qzb6OpeSnhx3YMEAaSBh9CnOPjbm9bpoQ1Is2VM608LFy0i7xDRD9hAkk2jtWT4shG3ns58qdzTathyfIhKwHrr4kS0FjcEpFPofBV3pkUrvjn8klvFpQqRkPktvP0RfXY53HkazNkMSm0fFWccYTocejfN+VOiE34kdGhfaqX4w+llsOYm8D7jmjlbS0lhAIKbW5kGMvlPam9109WTxokK5GPEGPjX4tL6Y2xIyqx3uycD++NDL51tL/mZw4qEl44lN/pXNxVxFW7fapLxozxd+bDwxafqtSxctSfvjWmxsc6aTy60WaVL5sIUgHoqNHO6kg6kmgbZHHqF5gIJ+XY5potRnrBCtzynRT6KuHPBXHwmY9fI0714EwwQbTUiVT7RfX1gHNNiN+H2nrHCVXTCaQJVPB/COzy4sSC9r5wGVdGHHU180r1xh9qO9qkfW3huPpdML4ADJgU44Qz5JZsYl+hOMFHOfL157HjgkiLHxN1vCU+/vFJZcKfz3H8Sb84HY9p5Vd970m0Lqd5duS1koY3tmOMLdTP27r46s/M5yvzsURCaIEYvlM31Hk8wEpd6EnuRISWoYwKaE5t4A1wkc04miXgfBGOkggnSAZGK/HFhMMfxE+2XBUeuYHqgtYXKZBhYPK2qPJfjqS4LLtcPDJqzknKJu0wzAlPP56n+Z8lnxrNZTeVuYOovmIX11ps53ne0B8duV6jvIzVMPZr5R9uXlNNIvO0vV4gMIa89JLqlffBvJiny1c7m6jbqs/QbB+cVmPVwlWE3+ynWSnCKrxVwYzdVImwcreOi/T3t/+gQXy75ghBdrB9WpRWSZq09S6DwtXGlWsoWzToAFZ4ULoltZdJtaZJ8c9hGKxQK1AeW9AZ9BrQdNhczlNxvvuvylhqplb8xro3+e90uLq3cMbLq4DaeqYFwJb8qU6KbkcMjV2KJ35cVp9NEanhKENjCKUZjfzYxyfh1k4p3Pvy1Zn0cVl94RjSRHjFUZTyVp33mPg5YzzOC//3whjTWOOMe72LN0/7NnFfONoflBB5qw1bKuu55u9png8f2sxxddzbvDUu5pBJmZnZMoZWfpfTrseFe2Aq91FrAaXTjq4ECZeMXIh72+Deo+k+p+p+WpoLC9tA+4Q/UrIT9eOnqEZZjeEzyYE+tbQ46PpMkd0gX76KPpJB185rSUac+2XzGqu/cwuEtqiteGOgczS/gnUXZVmwIBh3DHxhS1ecyJvklp1z13npHHMS4yfDU4fsxpx2T/WZbZjUEyrDhJAwN2VYHmiQasXp6xup0XDMqVPmmSyFbviHbFBAIfB+P2n1pY5syXvWur9f6bc/dZQP4/UU+QPqIybzjcdC9Kt4n0uBbcUVF/8yGCZgs+VhJcn40cQjf0DpULJj98u8dhsx7Qv66HV2LBRpY14yPLonycq2RZeR2hLinJ1wKFsRubloaBTsdp77cfH+wCx1iSv5oISSAopGd6wAdmI5MvaFRONbYxW+aIQ5jsRotBZH0AXHqmP9RYoLneULl1wXv6l9izDEehMEOPbUiYK91DHOyqzyJ2ETlDefZan1Q5M2S9h3GLI8z3etG+VLH1T3iur+XWUhC6kBbkeR8KFP1WdLHfM8jHkhMzdVvs160i+De26RvMzbBvqvyDK7joqhsu7XqDzkvQ9b3diJ5zm9jmsC3pRXzfuRA/W5n3eh7u9aHJ3BoCC9hBg7waed3Jq3aoUFXcJ+GBRRiISfZVZ9IaSfjCZMt0OJQPHxk9qEe7iWQPNJLck/7NokewyKXcUvO19zj4eiwtSxjmdCjv1zJimu71+YsnatMlL/mBmsZuFT5o77HGtpBVcRpv8yaZ3zKACMY82cDJ0m8KmXkgJUiUUxnEdDo+CHdI46sFvoVEBp3neVnWQxaaxKETFXZLBsbdYaSyYY2LEd/7yOYUzHgJKLFnjXuTkdg/nCrYv2sOthlmvOjxeKc5t+xLh9vIJuyPmyOi4/026CsT0Xx3or1dvKl8Qj7VZSXk4366ULQb8bNLdMznw41nq3nD5eKin9RXZaWE9N6hMN/t808J5HAdcFw9OEMEhrahW4v/2HzqEPN6ExUGQkt0Q5H9yz9L8Qy+qUfzRXV4N77j/EY1ERmehMCtMA32UVRz5FlJ23gpUu7AFY+VGnoY7Vhjqr9IET4LUAAAUPSURBVOYemc7641FO4OEHniZQP/BfE9BmZ24zyKdeqYUnrDjm4IbtOjzTZMIigfTimTNzTDxrXiG9tw3x7KuP6xUrYfF1fmqecebNPzWl0/gby1PpOh0PkIjQVkLObwp3rh1yesow1MGqI6TO+zgWrzSW0jhh3EXSRI+6i3gm6FRAwksUnnmvo7GDfcYzMmLR5qux+jIfq6QBRlJEUva0h+hMUPAJeTMtXH2YhySiHZpLLA7nrHQ74JpQGT1XfU06e2S5yTctnMuOdk7kdzafZMwOUnObuvEUXKtueZ60Y2KeIqQ8k9+sg3cu6LhEw1eSDhfkM2fSIWGsZxyzyle6aJ+xTDnUIy3WWDwUWZ6zR7NKFdyEDp09exafhNPbhhT/oZl/ZtzUM3sxyJ4SoR7wNKHrCMGFf4H0z3QU8S8NlD+Vbs2b3xtpciSDUVSDXqTKQaMDVfYGky3QOSr0xrPz/4z5LMMTJ3CVUL8+vn8S6m429nVisQTXrDfhSY+xBuY7+hPbzP9NjCvt6ub8XOSZKoG42nqR9wWJ4tGPlfFz+OCuwVHu3FDG/xpXhV5V90WeMwa66BXX+O/gnkPhSoGq8//zcFEcYbzJb7mBFjQg1Pdb+F2QxpJ4DgGNOxRdF2sb592WFouPuxQQjBc6jtIceKpwuKfI3KAsjeeXkhP1F6nboezJu5s6jh1pnpXmWCBo1+l5v8OrHX9Krg/JZ5LxutTqHLr4NlCE4zlPx9AeZJZA904xot2J4mkMvNGohns3r7PjzjxPtWNqf5BfE5ISD5XwttGkEUbRB1rjBzko+iKOcQXbbjRFvmbupyRf9fepeQtlShb/NIXV38+hv7tBUSbD7fimqnwGkpnNCcbc50L4wq0m0ED8TivJ8CnukQSX6lSUxV7iafpXHU+uKP06+Tjn/8EAJfwKZaHMNHl2iTfzFjo4OpC8iUBnhWne/UgeADTh6snH89zgEpCvyp5PZZEPMKnug9KmDk67gxdtRl4oqbx5DHRo8OS0IvZ/wqM4Jevw+CS41BeEr6zYJ5/XmXKAxmSFNpDsUxp8aAD9C700OzFdhAuVrWBweRzkSON0ULqlAr5/sxE839wrdiy74CvNfScYo1EH0nVA3cZmdspjC0g4ZHcho5Hi1IW8Uhz/lebAoHQPTEqnMZ6XXj9ZPbF4FWpDng4SeTfxlA1A6wLyVP8skoayTGUOSveAfKxxQS9UZkJBB+ed20u45INv8ib8JNmpfmNyJo3wW+SVYKDTCfV7Sw6T8lSaif0hWq3HUhh9Q5kpnvum8ZLKwocXkOzWGeOEAdI12z5JvpcyyvADr9Tfykfidz+a5K92f6d2Bv1V2B2+UC4JXj+5HtzNJiPcxNP0J/GD5zjiQpOMcMqLMJDiTT+VMZKCSPg0UcmHchMeHz5wCW7KN/Eln/LI52sC2kz7J7UZGjw5HRy0HJ/iqS+I0w/TeKGRX0qDT7oE0MmDOGF4CQN5HBwATwCNncAjH3wO5AvkeOpAui7IeYmTB0C4C6gLeeU08JdSOheaF5PoXXjyaeIpGwA/DUhDWZcqk7JzXug5HhwwCy88k2Q3qX7gyT8B6cmnCeASPfebfCmc55nSgIcnxXOfcgDw+PB2QaD/9+ppTgN/ob68nCBf+NFh0FNdwN1pg0IDe+glcLcl0Leul8DtkUBvUG5PX/Q16SXQS6CXwBctgd6gfNHd11e+l0AvgV4Ct0cC/wMAAP//zQoztQAAAAZJREFUAwBk+5sLvtAdUgAAAABJRU5ErkJggg==',
    'Área Sanitaria de Santiago de Compostela e Barbanza': 'iVBORw0KGgoAAAANSUhEUgAAAXwAAAAwCAYAAAASCsFpAAAQAElEQVR4Aeydf1IcubLvU9Uevxvx3BfuCuhZgZkVAJ4FwFkBzAqG89+7zIswjnjDPf8NswLjFQxewBhYwWmv4MAKTgN2xIs5Q+l+PyqpUFVXtxv/vnZVKFtSZiolpaTUj/rRxXDvhf8M4NQWuFI5F2DtWXoN9BroNdBroKWBohXvo70Geg30Gug18IVqoDf4X2jD9tXqNfDOGugFfHEaeHeD7+zZF6eVvkK9BnoN9Br4AjXwrgb/wjt35Mw/+Rx1s7x/sjz8z5Otf9/7fR+feCrn/947WX3wf0/W2wAP0MandIm2/H9ORgmX/JqmfKEjA1oKE28D5YAnAeWkvECblnh6v9dAr4FeA2+jgXc1+Cuu9Pve3OO3yfxDpsH43vzLn5rzv3nKJ//mD//PZEQL84cq+0kb7P/b8p9/2mobrxvGE4z1ZH9jItrxTeEPrXWVf5S7op0goyzKnRAWz78KWybcBYXzu6aL8g5/ejFO5fUqs8r4d+GORP4f4Poi9hroNfC5a+BtDP6ZOXtp5n4w7/5izs6d83/1hdvw5n9V/GUA+7RX+a9yx7w9NJXz+uCRK819p3I1j59Uj1Buyh5h8reNc4uXi/XCByWDHYyz6okR3mTlDj6BjPSOwme5DMXt9cHGuJ2P8BeCy8GN25dvmoyOTeX10qHTjgkQ/rlw2//+0+8hX8V712ug10CvgbfWwOIG39kzjKZgd/CNWzfngZ1ksGQM950V4+ufH62W3u1g4MwMoybvEzhvy6ZrcN+OTRdGV2XbwVe0ct4mr/7fxmkOFaH6LV0xhnb18/eHMsTByEPxVhC2clBuEQc4ipG/YuYCzVoXchIUvlwVeUUTyT6TQ9x1rCmPX18dfL97dfD9PqCJCvkX3rtd8feu10CvgV4D76SBYtHUMuKHOmI4dVbu3PxhW6w8zWzzz3s2kkHCMK2Z+XDMIb6/O+93nPlO42dmSvphXVkUp+SglfN4uHfCyptoE5wtc0yToL1iT8zgnbkt044A3Gut2Amr3reGuPBbol1eH2zMrTNHN0q3L94zJhL5VpgxAZgVRZicLLuiDldIl6H7YK+BXgO9Bu6sgWLBFBdO59DiXXLmVmXsJwp3OoxjIHhbT0Y3xD/yD6tp48gp5OufDvdenMeVdMCEHx2haGdykoBz94CPP5q0DnWGPtZ5/T+EWrbSYagVlPOOyW0Fmcu6SWvetv0CE5wmICaEpUHp6knIWRluAIcyS3TuvC/GxLmvgN9Dr4FeA70G3lYDxYIJV+79aeesas1sbeBtHMOWXxxRJKOp8KGVJavenOWjhq//a+NYxyIjpzNxZbysncdpPSEJQR18PLvHL8oCYwzlFjQpKPJ8cN+NkKdwcIPqqOiSm65hxyPsvbJgElCo28Vjn03Kw1FOmytMHG2kdiFtVB/vNfBFa6Cv3AfTQLGo5JuB56biOfx5WKtgVqphhapV7rI391g8l4SduR8V/uTuSmfipvsKKsjSTWHrlq7WGX7bCHvnOLJ5LvbNZNQVDm6yvzHRhIFOtnQ/Y1fhl+30gTH+BGPu/CF8oTwRj+etqPTKURmIHLhXovi9eza2/uo10Gug18A7aKBYOG210t0M/M3wtnArAvOVsSe4lIWJf3RgNb3MUUvM2bkynJOXZncynIP7bgcjbbo/wfGNZVdZHessqeIPrQpn1Gbwzz/KfWFW4iSi4K0bxN2C8tgfxvsNlJ1n8c3btvJ/NmGCuU3Sh3oN9BroNXBnDSxu8CXaxccUOf54E8hIvVSST+ec39V5+T+Hey9OBedxAjoLN1xvS7UqGvQaMLS3ZDMMbfVUkp1zJJQb/SjrQvyX0WgrOO24Keyq3c6l7hfIqIcyhTwx6uRRmltXymUZ/XC/QWWvyqwbxcp/V7TedWqgR/Ya6DWwqAbuZPC9d7/IYNU3OeeFrdoFLFqO98430E1Rb/7XKFjG0/86uO+2Ylye4w3hQ52nn+bAS1PcrxDuCb4YLRhkHQkJdyiFhZ0C+ArcPnh4qnj1WxYFcp9UMdP855/AJwBfg4/HOUweJUbfGe8KcMQzpvwy9utt2dZfC2mAyZnJtj2JL5S4Z+o18AVqoPgC6xSqxHn6q4Pvd3XTdh0gnBvO64ONoyud7beBdAB4/CBMP68PNsbgrpVO0doRv5KcGhEDPHGT8ClMvA3XmTzy4F0Bygu8UvnzMkfRvbeABjD22pHxePDhpD8OW0BjPcvXoIHZBl8rTV+4jTcpIfCIdy6fjiUC31ymj0LsM/lKNMDTU6Z+xw7pK6lyX81eA2/UwCyDf2behTdQJeFMwDm1vIYDd8bqFV5R4JM35c6897eyNAinOHpEr4H3rAEZ+t3Bjdv6nFb3PBLcHzG954ZeUFzQe8cHDxdM/tmwpT7EDnZeoaBTZ/ycr9Pgc5yQjhoI69x56vl0cNAQBm8KE88BPEcT4Aj76lFHoj3M0QDnzjTYHJaeNEcDuunNB+6m+i061Y369r+8HecDQ3RuqLd5PGnzLImLdwqf86Tw8KcXR7zAx30vlS08TMDgTfTkSx55T71xnfLiJn/ixU94/Hb8wd7vh8O9+f9o90A8pAOQPdx7MaHvEc9B+KY+fnox5km4xBPT+hRPPrKUdgI94Wb54qPueT4T4RptQ1rhcp46TF7Qc6Bdxc8HD0/Qv8LnXWWhLkPpCv48PWGloVznbfkRfxp5KGtdFtGmwsgmb2ikaYPwkwdZe+R00g73XpxTB/oQx5WKT+hXebkSH/SMr37ptNPgS5CnYGRI2N8+bgkqADhodDR4CQdC6we84BS0fE8hCH8siA1Jg6UGqJTUMdtTF8o4ZyCGerTLntKRtgXcfA3sNAqNI/q5oCoLgyY+hhmY9JMaDKOAruB9MKMTiL13HRqI7bcm0loMK9hyOobUouUJN8ZF2dQAOZafuwvoOYTdbMah9glPT8V3UTJKMzikjb2Fx2s52ox5jvgsSc5J2ytOuac+zCd8cIy7yBfi837ulQUPJfCwQADxsitv1Ase4YOT7B0FltrvnAiX3FnSh3kbmfO/zdSvVVeUxWPayK6Q83/r8ulIjjahbU476lyXhTKZuR8mHfdq6nYNb927H8xsrHquy28650NbhqPAJiXFVjQmj1Jk2ne7lANQuXnwApZGGV/rPiDILgh9xIwvGexgK3Ie6q56YHuWkU8fMtXXVBfztp3ewqctEp+JnvHVL50WNuPSSnyEIZtBrtF8CAzeGtEVcNU3a7pIHxIXlKhOqTxGDDIeK1Vj8LLUdnojWLTapQH8Z1GGxq8JiwaiEaFRAIsfPaMBw6ea1TjK/5RyUB4LH3jzT5P4ZoMZR2RnpiMwZ+5HJovE1/vzNXAz8PviuARiWMGmU58NN+1f6ca4r57mWkH/Gdf5lW7G55DRLPJuCnehdtxezt75EK7hXPx0BsdMTBrkqd3uiHDOGI0N5bZ5fVCD+nhefkkmDx20ys8CpFEveOAPY8VsRWFNCkF/CjadM3+a5Jl3wYDfOGs9tWatK8iSTFuJebToU9G6fDzAUJr7ThxLqvOh/NrlZaFM19nDDzVTFVhR+x7xljw80vuWoGHwY1sy0dKWW3N0uzmrDsi+iv2FvkXW7TKCmw0efaKnpThJ1qyx7kvSxTp50G+uVV/qgVFPTxPGvr4knOq4cZT4iEvYEvSZBt9knGQAT8Q413nvfoF3PpM9XETWXBl3JFaN6J9iMAf33SqDjI+V0YkGpfu2LApmzFpqxW+bQlw6cyhfwbs5GvoqNjo+nQwJYWUQHlN1P5A/5aA8arCRmWPVYVzartGpY4M9Ck8XiX9VdXhmag9mevh6mK2BZQyvty0GufTG5D5vAAdBzrnlEPg3m9iCV1mUoY+UfFRPafgvBHmdzsfvITHpz2rDWO7qe0xh4eB2Aq4lUUaER31XJIu+0qK+S1QGR4sLM7dvMvxvXOw5q3Rms68oY8WqhY+MmfKYzd5JeX2wMVY7smJe69JHZ6ImUjsG9+O8T4xjCJXkMrblUvi0uhAtd6a4wB9GW6Ho+3GxT2jCke5pg7jbyKSvoYOgiwxJEKOeJm3zti6+l+Asu0IcuWarRYZ/m6CU6Z8oExokrEzeRsiHSJM+oeB1z2DS2uqhoKCELOO44r/0heNZ/aVZM3mW5C7BdTE/v9asLL/hWrhVEasb4QokV8a3eAeuRE5C936HBuLqaImjiqi3zgHsvN+JR5GnGijbEnXW6idrw73b8294xVM7b26Xfh8H4XPFd2piK8DEHwy1Jn2t1sKf2rSNRjIy93QMoz7L0cHUSg+xpRYqUdY2x5Xg3hVqg6N+FvvjJfppy1Ud19FDtdvUYsrscvC/rLFwsuzSIm9X0YtQfz4rYjb7iE2Ms5zKws7E0tEFfCrL47x94uQCqQFlNSHzifFfxK/j1JNGO4VJRAsEteVxbMszX01QDTlEtFAMaTVBcNQE6r1A2tkF3asN1B8fpvokP+mADGkv8AmIgxesCMaCaedtgty3N/hahbBCvdKKVqvQHSkWY/TZGH1vJcbT0jdoaNikoOSDs3ipA+3S6GEiYDacnmUj52xPHfxEnao6n9974eGMA1tnc75uiJR/8uGJZVnRYJ4aQLEjqr1cqBNye5ilAb8vyhmTetCb2rJzAGuQ++re1Eg6f6JdIBO9ktYuLGagARjaRBlyJq/z1miYzbzDQM89smCcyGB8qz4WdmvazY1p9yQzljGUO/RBswudke8meu4jy1Qv0Y9caSN7xysZnEH1iQ/z5o/0s52XL2ax6tGZdKf8n6k+q5PWYiry1Ude0h26seKbIvg31XFbYnsXv3E+no412gLpA9gpq3fS/umDvd/r3VFaIMTFgZlVbcnYtNZFnzKOsjRxM/G1yG8VDePe335pN7ZB54SbMqC9sDUJtIio65N4ZvnFLMIb8JeDb9xuzoNijVWPfR6XMxcGwiR2yH/9YaOkoOSnFcMwDuC60VuzrC16aRJUBw83yPBJFm/MEaxB+e8LThJo8B9RvpqhO8AH6UKdusk9Ng5SVjmjYfVJjVPzxtHDSqRZfTFwtUJVfPnGiuPUTxRPrj5LxsBGI1zR4mJA7Rc+kyHDG8eCDyvAimn6F4PB4sjCDURbirtKiyv1ZrlJLsMyVW7wAh45lSfnFx7sYp5y0eCEyU5Hj8fozTnH4s1S+VIi9Wny0rGGLVnpjqlPorX9lNY7t4VMZIuHPrwV8lRkUefNhYVO3gYqS30/gfaZVxbyudbueqCjXU1UL53uiYGrwO/jYzQpp2bv0IZq213wbWC3IhnPPBOfjkja9LvG084OnZN/1JOKYdvoKdVZ+QUdIJ8+pEnMAYpfCJIjXPMlZPRHKvfLeQafxDRu5G94444BYr6wsPWy6YuV/yxZ09zvAaMOMUZMGjBMSCgIUMU5gqpX/2nA1o2eBrS2/chYFNS5w43AK+16ANKlHUajwQ7C+bzTCulbeFTWU8pHWDJG+DnQVcp8lQAAEABJREFU8IovefOhTgr3rkMDcZDqHoznsxmn6FXAyvIy0upUfmCT0lwwbFoxHUUd1/RZgbB9liEW/blkpzzYlfFV1c4ji/ZKeeoYpOpv7Cja5VbX9MEAKb+GqwycwygtNQh3jESDQ9+6zdv7Y42Rl75jASejuQVNBTsKupiRX0irXYiOIo6TntR/aYvOI7YZYmxYLcY2RUe/8hZ3y7qfk+t+osUf7wQlCdE2rKg+z1IZ8UXHVs18UioudrGP76R75WNhZ9etJ8vuC1H3zba+o27qUwHpl3tWD+MCAvEBaj6160yDr4ofyTiGARFSNX+mjBJkd2PL1n2N58jqTvGO2FJnnYjQQN/HT0AnMG+c2V5M1AGCEhnA1eq8GsBSjPjPzFezrMJv7SbKQ4lDBwp5KZJc2t76eFNPeP7DdqvNF7+0adbxj1hK0ztpIA7sTa/jCCbbHDSgj8UyNYCZZIPRV/vrJuhp6BtijG7Etj0H6JocMLImw7eT56HJO+BTm0YZwWO1O9x7EZ6xHsqAkRcE2j2We43xlssjrHKH45/IQ5IGXGvVKgTGQN7buWRwXh18X/+15pUWLMYu12yJ8uaS6c8yeMEuaIHU9bhkMtJL2gWEv+pEHkAeksV5+q78Wa7W+3DvhdrNPxXjZdKvwsFpQlnP26brpiw7eO2e/4GcofT+YO/3Q1et7hmPmrM8k2k4raB8CZQX+JlPSqED9ZutUJB3+ImGecWk65Q3ftCTJgFvLpRDee2TDfqmzkxU1MWs0k1xvziEfu9+Ad+lKnYEvYuvgLENvnAbRfwzEMImY9jiWWkrmMFghSfDBivpVeDQwISd839tMHygCFshDaInEs/Nt1PKy80mbZnOhZNzoaxxAE81eipzNssqjdWdEcUD+WDUaibcCAQPxAa1KOuSBqMhYlnG5m0b3YZtoqQP+BSzfPH9fagjiSBj78W5o5OqDaiTyL3r0EA5KMMAvKebnm1yqQEFDsOLn8NrngIx94Nh9P/wMjA1dcWbe5xDOO+NbTapJvKauVpxW5iwa2QMlNXiQ4sJpwGsQeptRN+k3VOZ0niz7PLOsSK2VLeMVAdjn2G1WeMWDcSFxYpF/eTprqvJ5FJGJeg1p1H30lxt9HNaFfY78sPNWvkNp/GPcQr/FNcg3EZqvZuzkeCZ6jiK+rXsWvN5+3j3Sz4W4WN37atHbnXM4Z86xpHZ2aB0O4HXG+Nv6jgv5qV7BG4HOV1Av6ENu2gL45xH/mXUdTNZ1SbhvhB5edlkMUy86qxF7Imr6vJcdVmdxL6IH9tlDD3jo86Br9PgS/AbnZeCZZSOMWoYMBlSDNjDNyb8iAxXWqnERhlRXvO2rexVTveX66pDqz8L52yq0VGyOhvbWhpFyYKrOyOKBxqDUfLBJdBMS+c2ZMWGYND/GMuyTGfkPC5I1s9EDRf5GMChQwutPP2vOZ9wvWtpgEddtYt0cbA2qOgf2pX6A5MmYfzEdK2+AE4QjBi+IJyRNv0Ndr1uVluId0swtcslL+EDTT5yl69UFvLHB9dV7pjOUbcUxiddAvqM0tfP9ENX3OEnnuQLvw6keNLLteqfcLkv3mVBMPjy3VUsMzwxLfRQX2jwQJNPPiPCbaAuojvSt2nCkw79VPDzo1V0TR1zXvFV9INHDb+tQ9K9OvieDyiOsjTr8AEB9/OjfHzX2YhGWULdYjj0jZpBgVjnKTy6V5qGvsQeXEzjiIgn9AnCbbhWm4iu+m2EST/KHGny+w+Mf0V7tEU98rSvtYARjbK7jC/UGb5Og8/MkFYehM0bhhL+NmzKqP0WDJjZSptInPRasQbDRzjyQvoocKVOKgWMNBN+K18KfLTOyiplHnCzGp0Od/AodFzxrQtI3wA6cGyMBj7yhrTkFRuCBg5n96KPXqkzQsvhddVgjYZ91cGXp+nDX5MG+rp+zRqYaFGIvVlEB118nQYfYVql7mgFz80oohXoWIEV8yKgBKxS5QUX/mgkhD7RT3sm/ETFCNmmsizrppJ0zJMRx4QDMf7cpWFjkt7rNdBroNfAXA3MNPhKxYp9TX7tnPeNR9VYPc+COlEVWJLXkKX4V+/izdhNbb0OJ5q5v3qF9AroNdBr4INqYJ7Bn8pYN5HCed4UoYWobwa18H20qYHCmW4gux+6tl72fq5eSq+BXgO9BmoN3Mngm7eHw72TzpsctUQF0pm9gp/U8ZRLDjym1C4QuJwnhcG3eYmD52mf4d4LvsCp45hufSAnTnwkq2Eo/QEgOP+/1s0ZwvAjmzBPEBCv4affdxMNegJwNc/e7/spDD7xfE0+eh1W7ULbBOCBgvZx2YO93/lscKAPM/6u9kJ/pEe34j0X8Cb1OX2AdoIOwANO9IkAnglx8NAB5It2Oqt9kAcdvpwfXITwkAS0BMhvlW0qX3gznolkHSe9EBZM6SLHUZ5h6Lehzzd4kYv8B9IpQPhNMNx7MVWPrjQPJHOYtU8d/s+TqYVnVcYXU2VDP12yaZuh6tSmdciZGuPDVpkeqJy0XVsWcfQDEM4h5dPuCwmPD/+wlVceh54DD9CInj9plpND+G4GPyTxhwgOwdYPykWRQn8Wxze6D8FjdTvy1wWPuWkshUySMlVOK8oy0ETHrwFaG2gcZGjiSx0uPO4FPuclLnmPNfHt5/gq7DVhAlUs/XpzjykLcd7OJc6OKoB34b+EpdsxOoYHgD/wmavLrfi6L21iX+Hlqq9S3vY9Z8s8JHDzhz/P29xVb26q7ZpK+ua+acdlUxfPzUuvj0U4d+afyB+bt22e8VbYaBN4wCk+DjzOeAlmGzx00+UK46mPNfWhjn5hFp/hX4t8Fv01Uz2sulZ5SCI3IMhvlC3lqzqnfDHuiceb56mPVfTCc+pu+mW+NWVVP2ygcHjKzLV1C0Hg4zsk6BQQaq6LtmPTOh7hbieM8lbb+IG3sbWuKV1JZ15jSm0/ZQAZn1Vb+f2WmG6dm3+a61xpGjpSOX9UX2je7xQTk4BXGQS7qS2EDi6VV33hOKfV+KqvmNq+2SdVLzMjf3ktV32vaDPUr0VK0YbBT8jg6watzpY3Qrj5s0Snk/E5io0XvptBp5Jyw0BosivGSwTVc6SKfFzn1MGvDx6t85SOi+8AyBCf0hh5SeDJofOYpSyDoZes1cg7Qkft73i42zd0p172yfN8U9g7t3v986NV5fet6vHEvI0Y4HkHQUYsC08RBXh9sDE1IOD7WqDWh3RH+1BvtXl74IeXAWte9ZFJx32UIatA7Ww18J7Be1U99bVFm0hueOkpvK0qHvpXzfPzox3iarOHrXc5lMzW2oMy9EdvnU/DhX6g8g34NIDZBQYEIalsTpNQnq/xXoHZkvrKoemScaffXsLz6uD78Jhiae47nlOP9Qn9BrrYVVXP46cNHHgAngb810Zbr7DNBud3A1H6yifhgOv+mW6nv200jWCWLumKcaOKPBNpyji62/G5kmyY+BqullM9pVfrPDFJ57WOfPWs/9TnO9KTjkoz7+3iJdnNmTpUPXZyfRsfQTPjRTQtHCU5utifVohqEql0TKQFMw2+lHLeafSSADpn9a15o2LqVL+IFDKU33Qq5FxZTe4PEuPJGI5QZADo/Evpa5p3yUzKmsCffwOceiEbPJAGrjoEK8GgG/DvAshnYBqvumugsDKz/lpIA7SP2gLDNzUgFxJgnv5ig9a3o2gTDcRAU9/Hv6B/5TJjHGPRGJzwaFDu4yeIq/vLFO/yJ/sbE28e47BU0T1yL4v4pmWFM7vmmFCLLBkHymUyfBjIJRZlFq/XWhQgL0Y/ihfGhpkMsOOT4JfxpccPkneYTPjIm1l4qzZlEsrgbVt9gvF5YdXLT4k8yw/66yKSj1b4O6JdMIHKr503t2NaOAtxpj4yywizaFjL20b8nS5OTmsq+yH9L2eSvd5RnM9zUK+ZC03ZMLG1nDpy/dIAYWWAkBbXYlHSC9bhlu9kcDcIfyrAAJC3j1/TJAywZUsQVk4gW1BUbx9fqpP8NtQ5ZOg8LR4mP1BhEKqx1eizGhq2O0EYyErBUY682qVy4w9ZkdaUPoAGyupNV50ilKvEI4RHhYf1GekJAyaSGh7HMGcTGdsGthmRETOMgk1dztjqr1h+qV8oWq/yQz/yhhFiYhKp2+XGJXJw9DKedJQtfjMmTAyDG7cvfj5p8It25uO4GhTqLq7ivdVXdV5eYRf7TZPa4L4de+28zS/06ZJF2ykUQoYv3J/Rjg69Lw1K12jXenxqLMuucbzVaRz5Yyf0xJiSYNoXw6xg5TSud9BFyse8251k7TCsxuGKdgrKwwmsczfhC3coiUwIv9C+Cne6sKt3/kgTyMsr7TJzpqz/HBWqF7Q/Z/yJU6fBH4ZBUA0Awp6ZCimzwFlY+VrHRfoHuqkBaSi5NAjhTwpa/Thzq3kZVMf1W2gYhpqNWXVQKl01YDd1bvePYdWwNY9k7KpRntH4VWPbUpunZn7LgPK4c9nfMqsvK1n15cyPWicZ3qmxEfvFRVrlJ0NYtFbqqaCMGQy1jMvfhVvCuMjHMRnhT4EGdp1v6Lc6DpKBe2I6FlS+J5J3NJXoAyKCwWLFrXsMjI30CYxwHPZh8h0jVro9DnkTEWjsMAE8RyfFHOOoFXm4byb+x0r2fBA/e6Jw24V8tAg8Coa3pnryuWCBec2OS8cw4gFXc6SAZLMTY8dz5G6ss0117EN7LZXVV14tv/JJjHqJ9tzNsNnqFyJPuzUVbp0ZTqQ1QXOFIkRwMpz4pSuqShOZhjXn3K0sHUlMs3xkDGUIj0Te5qvdR31m+UrnnLeUZgiFcq42KN23MuwvTTd00swcDbsGpC2zMqhX4uncsinqzrHUoTRwG/petOx3zvALSeDiN+PLuNKP1WqdDW8woCJpymtMsFNUxkF1M62D5Eh7OUUwty/cWugz3sLqfpKtEEXL3VhGn69OPqHfZW+Kj8XEKl9e0/nWgmYi2VdaGcq4jNRvw0fZ4hFBM+EbYnlfI/wG9pocvkNkFnYcjI1opC5lWHdrpu7AXdrJPPe9dL+DsqV7KDHv9FG3YMuyMugYxO3kk0IoRrWoM3QlWVvoz7JLY7A+wy/Nhb9hTKvqaA+wm+fkA0gOO8DO3QSyPX+8hF3quJkd7fCm8nzy+mBjnBUjBNXW6PACnYa8ArZ7oVlUtI5fb9usBDooNUqrF7ZNxiwmZEenFhanirxJFmwfA+qzMu9C2d82Tww/nYv0UiKDWkGfZvBNb44nhB4Laab6x04QovpZyztYiyZyt6NBofj4ZAThzxQ+m2IFPVcT7mX7jHWRQspohDPzLuOY4XhY4WG7HeMEzcCf6mvX1arvwsyHr0EWM1b3lFH9rP7sNv0OHOCrJ2ym7k3EfDfF0ziGUNwmMvzII+xcGfstsQ8Nfl858D38kTcXdtPRCHYedYj3nV1xU9B2yqYcVcI84xM7texvyzARTTdOjVW2pQsdeW7GetvOjD6YfnAAAAV/SURBVGgiN/xkhF38D4783kSdj+5jkihNCoRzwIbG/BrfI1veP1mW7WQxctHVR8KCgYlUpyx1XmbLVu0omAgUvHXFbfDuobQtI6Uzz1kUwc8K1HAjZkgUo23s2Hv3i3rAszjg6rLCkwOKrokx8GCP57dPdqAxuNUQKPRyoDNJ4mJbc87/VSuC+rs6g/vuP4S31AlED52QJygYmKTTVj3ozlvBKgD2AOkccahjo+Hei3Nv7nEoe+vJiLzchClfEPCV/qADYCi9Sc+nTLj0z4mMXa0SZ8vw5FDTskA0Gpfa8R4hDxKGfqjjSeF+Iz6ozshN7Vg/X45cjhOga/UW2pdwE9w+8amygVwA4vjjW/9z81VZj9X3w/fr6XOx39qNFaEvLpBVzUK9ckBeTezQKX0RfvFoZe14QqjeSWunzISDbhlHYulwM2R2cAZUGjPkmfRfamfHOBPDGrrW+GyX4UJtOVWGV+z0tdJn3A3Vl5S+dsmukM9Q9/MgaGyfUl/1t60wTuNOI+Unnrlf4Iz56dRAnNHFt/H55v0RD2uQX4KKxTOJXaDLlA8+9VQ5phYhRZXoLX6liHy1UVQrlMu3kPRhk3jbVgc/MVZSOr9UZs/bT1wIZ/Dk0HW26MIM7p/qPO2fGtycp2563XyayJAkg17Ev3JDJgCNxjdf3aAKdI4AFOceQJSzBk97EmJyokyh7Jq1lVfnVzPhySFtYcn/a4Ski6A3tbkG4l+vdJzR0IV2XYkv+RjyBo8i9PHS3LqCcv6pBre36um0NdpDSIPHzP1gaiNo8CBTbT4y4Vm9Wcd1Xa3yn8ex08ExH0W+ddlUpln5qpznKss2fY0+J6n0219fdxwPiDbXUa8cZFT36wQdOsVgiR9jGhZGNW8K6Exfwc4/jhFeTWgPlf4kh3n9O40Z+E3l8eZ/Rf/1+KzsVBCdfmQc+UvH+n9kEx7/+udHOxqbMsK0/ckOuAAav+QBKB70yVNZ0W4saUJgVS5S7hy4pWFr8sg5yuqMPthRJlNXfQZZVXGPySsHDL/SMokhV8FbV8T7E6neifL2Bl8VHu7d/skzRlBCwxmd/Du5D8XsC7eRQAPjO818fM51ayIDnfJEMYkn92980bENf7SFHKdVvDrJE8KvWAVImBrqkHguW+jgWAEimwh0daJV8+4vyACgCVd3Jo4ewCVALmVPeSEHmFX2QWlTZbev4CrUyZPO8Ael4wupywzEvPrS5y70NnS90EO61wcbnCMvw5/aa6Cd26vY9vBgvIUbpXbFJw4eOpDaFZ84oHZt9EfySHT8PA5/G16rbORDfpQNn3ie7yuVU/mI7P4S6KX7FlxbFnHyK6RHwjmAg9YG+j18s3TqrTgqdb9CsE7ft9Y1+MaFtrB/M45WLL+UJtDaeQ46+nfSVc5LnVM9KafkfddVhkKTAOl4nybJwU9lURnXoZdmY9NFOIdB1hcKHSNBY5IRa8PRJtAoP/IJ4+dMrw82xqUWGAEvncAzCygvtELlz2UQjouB76g38QRFCnyJPkpPgCK76ohiEk/uz+IHjwG50oqRcJJJGEjx3E95TLKJhhtvyADIt8EvPnAJ3iQ38SWf/HJ5X0uYeicd4BPvqjv6hN6GWfxJBvypvfK2THRwqV3xiScaPnFk4BPvgpwOXx7v4gcHH/lRNnzi4NsADZhXz5Df3zYaR4vIIQ20NqBL6PhtGnHwCeBrA2WFD79NIx20NlCWNi/p23w5D7KAHJfCKS1y6/D+Rj0BJVxK384Hei3rbxtz318i7ax8kgzyQSYA/yxADjT4UtrcRw6Q475og59XtA/3Gug18D9RA32Z36cGeoP/PrXZy+o10Gug18BnrIHe4H/GjdMXrddAr4FeA+9TA/8NAAD//+FgT0oAAAAGSURBVAMApLKGGiIKCfUAAAAASUVORK5CYII=',
    'Área Sanitaria de Ourense, Verín e O Barco de Valdeorras': 'iVBORw0KGgoAAAANSUhEUgAACAAAAAFOCAYAAAASbf3/AAAQAElEQVR4AeydB2AcxdXH/2/vTsXdYGNMb6Z3DNiWZGN6r8F0QockQCA0SzLhiC3ZlNBJqCGQBPKZ0EKvBkuyDZjeS+jVNHdLV3a+/5wkW7Zv907S3el0eqN52t2ZN+23s7OzM7N7DqrqoqiuNyrKQOtA3taBKNQ0E6iauT3r6etpSVXdPTjrseLmgPpfCSgBJaAElIASUAJKQAn0eAIKQAkoASWgBJSAElACSkAJKAEloASUgBIofAJwekAZtYhKQAkUCgET7cuibJOWiNkEwQHaxhGWWiWgBJSAElACSkAJKAElACgDJaAElIASUAJKQAkoASWgBJSAElACSqDwCUAXAPSEk6xlVAJKQAkoASWgBJSAElACSkAJKIEeTkCLrwSUgBJQAkpACSgBJaAElIASUAJKQAkUPgGWUN+OJQS1SkAJKAEloASUgBJQAkpACSgBJaAECpmAlk0JKAEloASUgBJQAkpACSgBJaAElIASKHwCtoS6AMBSUFECSkAJKAEloASUgBJQAkpACSgBJVC4BLRkSkAJKAEloASUgBJQAkpACSgBJaAElEDhE0iUUBcAJDDoPyWgBJSAElACSkAJKAEloASUgBJQAoVKQMulBJSAElACSkAJKAEloASUgBJQAkpACRQ+geYS6gKAZg76XwkoASWgBJSAElACSkAJKAEloASUQGES0FIpASWgBJSAElACSkAJKAEloASUgBJQAoVPoKWEugCgBYRulIASUAJKQAkoASWgBJSAElACSkAJFCIBLZMSUAJKQAkoASWgBJSAElACSkAJKAElUPgEWkuoCwBaSehWCSgBJaAElIASUAJKQAkoASWgBJRA4RHQEikBJaAElIASUAJKQAkoASWgBJSAElAChU9gaQl1AcBSFLqjBJSAElACSkAJKAEloASUgBJQAkqg0AhoeZSAElACSkAJKAEloASUgBJQAkpACSiBwiewrIS6AGAZC91TAkpACSgBJaAElIASUAJKQAkoASVQWAS0NEpACSgBJaAElIASUAJKQAkoASWgBJRA4RNoU0JdANAGhu4qASWgBJSAElACSkAJKAEloASUgBIoJAJaFiWgBJSAElACSkAJKAEloASUgBJQAkqg8Am0LaEuAGhLQ/eVgBJQAkpACSgBJaAElIASUAJKQAkUDgEtiRJQAkpACSgBJaAElIASUAJKQAkoASVQ+ASWK6EuAFgOhx4oASWgBJSAElACSkAJKAEloASUgBIoFAJaDiWgBJSAElACSkAJKAEloASUgBJQAkqg8AksX0JdALA8Dz1SAkpACSgBJaAElIASUAJKQAkoASVQGAS0FEpACSgBJaAElIASUAJKQAkoASWgBJRA4RNYoYS6AGAFIHqoBJSAElACSkAJKAEloASUgBJQAkqgEAhoGZSAElACSkAJKAEloASUgBJQAkpACSiBwiewYgl1AcCKRPRYCSgBJaAElIASUAJKQAkoASWgBJRA9yegJVACSkAJKAEloASUgBJQAkpACSgBJaAECp/ASiXUBQArIVEHJaAElIASUAJKQAkoASWgBJSAElAC3Z2A5l8JKAEloASUgBJQAkpACSgBJaAElIASKHwCK5dQFwCszERdlIASUAJKQAkoASWgBJSAElACSkAJdG8CmnsloASUgBJQAkpACSgBJaAElIASUAJKoPAJJCmhLgBIAkWdlIASUAJKQAkoASWgBJSAElACSkAJdGcCmncloASUgBJQAkpACSgBJaAElIASUAJKoPAJJCuhLgBIRkXdlIASUAJKQAkoASWgBJSAElACSkAJdF8CmnMloASUgBJQAkpACSgBJaAElIASUAJKoPAJJC2hLgBIikUdlYASUAJKQAkoASWgBJSAElACSkAJdFcCmm8loASUgBJQAkpACSgBJaAElIASUAJKoPAJJC+hA5Fgci917V4EAgfCkVHLCXAC1CgBJaAElIASUAJKQAkoASWgBJRAzyKgpVUCSkAJKAEloASUgBJQAkpACSgBJaAECp+ARwkdGBPz8FPn7kRgsXkVE8tmLieO82Z3KoLmVQkoASWgBJSAElACSkAJKAEloAQ6T0BjUAJKQAkoASWgBJSAElACSkAJKAEloAQKn4BXCR0vD3VXAkpACSgBJaAElIASUAJKQAkoASWgBLodAc2wElACSkAJKAEloASUgBJQAkpACSgBJVD4BDxLqAsAPNGohxJQAkpACSgBJaAElIASUAJKQAkoge5GQPOrBJSAElACSkAJKAEloASUgBJQAkpACRQ+Ae8S6gIAbzbqowSUgBJQAkpACSgBJaAElIASUAJKoHsR0NwqASWgBJSAElACSkAJKAEloASUgBJQAoVPwKeEdgGA+PirV/4RmAPIW4B5czlxipqwoolhyXI6iTB4h2o/U9QqASWgBJSAElACSkAJKAEloASUQIER0OIoASWgBJSAElACSkAJKAEloASUgBJQAoVPwK+EDkR0AYAfoXzzM7gfbvxIxHHIcvLn4T+ulNXipo+X07FhHHMMDJ5ZSVcdlIASUAJKQAkoASWgBJSAElACSqC7E9D8KwEloASUgBJQAkpACSgBJaAElIASUAKFT8C3hA6McX011DO/CAjmojH4KaZUfLKcIIkJj40tp2PDBGKfQ7AgibY6KQEloASUgBJQAkpACSgBJaAElEC3JqCZVwJKQAkoASWgBJSAElACSkAJKAEloAQKn4B/CR1/b/VVAkpACSgBJaAElIASUAJKQAkoASWgBLoFAc2kElACSkAJKAEloASUgBJQAkpACSgBJVD4BFKUUBcApACk3kpACSgBJdCWgBGEpxbh3BmlOOuxYhw+NdDWV/eVgBJQAt2CgG2/zn+jN8Kz+iE8bQDCM1ZpFrtPtyvoZ3W6RWE0k0pACSiBZQR0TwkoASWgBJSAElACSkAJKAEloASUgBIofAKpSqgLAFIRUn8loASUQE8lEDYOql9YG5V1O6Oqfg9UTv8VqhpORnToH9DLTEDffhdi4zV+g+qG4yiHYXzd3qiuH47x09bDVTNKeyo2LbcSUAJ5SCA8rQ/bsM1RVT+W7dQxifareMEkxGKXIxq8DjFzM6LuLYgWXYdo9ArMpZ9t46rrjsL4F3ZNhLVx5GHRNEtKQAkogTYEdFcJKAEloASUgBJQAkpACSgBJaAElIASKHwCKUuoCwBSIso7hX3QK34jqhtuWU7sG2wrZvWSunWW07FhIsFrqVZOUasElIASSE6gatYQVNYfiEj95UDwVjhyGwS3Q5y/cXsrIJMBU8X9PwG4Hsa9g8d/o97tPL4NgdCt+Mm9FVUMX1l3Iia8MIzumbfGfo1g2iBMqN8kf4Vlb+9iCPtVBbvwIlPlqqzbGPacXvdRceZPgkeM9s3qCQ0bpjwv9msSHlF0qfO5U0tZb9f3zX91w7oIv12UVj4rp20KuzjGSy6u2yateNoq2ft+pupIV8QTfm1A2+JkZd8uYrLtT/X0MzjJfwvbsNuYzt8SIrDt1zkwOB2Q42DMrwAcxrbsOEBOA3BOcxsnbNsCtyfCRoPse9WfgQmzhsHGDTVKQAkogXwjoPlRAkpACSgBJaAElIASUAJKQAkoASWgBAqfQOoSOqlVVCPPCHCSQE4EzKnLySKn10r5dAOrLqdjw4gcT71NKGqVgBJQAssTGD9tI1TV3QiJPgsHt0HkbLYhe1FpS8raEPTldmUrYn8GoB8Ea9CTbRR25/YYHp8DR66C6zyBqvqHGPdJqHyR7RJ9M2Gv/7gI0dAJnMB7Im/FdW7F9+4G7SruFsPIOTg5Y2USeRKI1eP7Oa+huuFFnotHUF13C8/H71BZX4HwtJJ25S8d5eCioZxQ/VvKMkTX4oRrOhHmWKf3GmNhAlN98w+3GpHvBqaVMwlNIo8HPCUut6cVT1ulSHwD3/wZ5O91YfMWXXwA7GKXtmXK5P7FdTsg2sA6GHgccC4D5Ci2SSMp6wEooqRrSxJhBCNh4wAug4k/hkjDHbALOqBGCSgBJZBHBDQrSkAJKAEloASUgBJQAkpACSgBJaAElEDhE0ijhLoAIA1IqqIElIASKGACkpgArm44DIHgbIj8FpAtAAymhCidsTb8AMa5AQQHQjjJ6US+QlX9nxBuWI3pBjsTOT79wd7D7ASsndDLT3GcH1Aa/ald5WyKBKi/OiUzZRKsB8FGgNmMshP39wPkVIjcAAfTEQ0t4UTmm6iqOw9V04fi5tn2vKFTZsqojzk5PZdxrOcv7mmcFBfq5I8967FiQEYD2IHilf+16PcJinb7gdvUVng9iaxF5skFicUzqeNpq1Ez6g0e9qJ45THf3fvi8MOZ/UxZI7Dnzn5NoarufrgymzH/mrIhpR8lU5ZxmY0gOJ51/CVU1T2Ai1/cJpE2mIdMpaLxKAEloAQ6QCDnQezXUKrrH2U/wvjIYlTXHcc+SH7d73MOSxNUAkpACSgBJaAElIASUAJKQAkoASWQGQLpxGInT9LRUx0loASUgBIoNAL2ze/q+p05Afx3TsLeC0h/ZN1ICZOYgIipQzR4NipnbMrjwrTGNME1LwG7zOkGBdwKIlcAzix82TQZldPH4KKnO1cfxNzGchuKtzUYgz82bI58Mn36rsvrYQyzJBQv+wlcaUBYXHSVEbFsLeOuykHn0hWJ4142O52LpSW0EVw8Yzj69rsBLl6AyCEtHtndCPgnB8ONTke/fjeisn4n2Mmw7KaqsSsBJaAEvAjk3n3xrE2Y6E4UD2vmwOBsbFN+N5C4b0GNElACSkAJKAEloASUgBJQAkpACSgBJdApAmkF1gUAaWFSJSWgBJRAgREITwtyAt6+fms/PX64ncbKWQklkdrGMDIFjnttztLNdUIiXzHJ2Qh34SQx2mXsmVmHk9/nQeReBEsnYfzzO8LWlXZF06K8auAZ7n1I8bY2xbg5y1sh1z5GANmSsjW8jIH9extxvI6uNq7zD+Ym2tXZaHf6gi/g4mPce3jnF1CcO6MUVTNPhWtuZj5ORk4WMmFF048OJ8GRm9DUcCqvGbvQiU5qlYASUAK5JNAFaYXi+zHVVSkrW4NveI+qxY8ld2KcxFdWUBcloASUgBJQAkpACSgBJaAElIASUAJKoP0E0gvhAGI/NQw13ZxAILo57Gdv24rrbtzNS6XZVwJKIFsEosU7JQZlAfv2Ne8F2UrIJ16B/dT8dj4a3d3rcxQFX+mWhRAZzHyfgUDwdkRDZ8BOstKhXfYPo5ZwUvYvKcOI7A/70wMpFXOgEH4+AAcHAsZ+Wj95goJG+j+Ky8sXJFfIoWtx08cQdLc61sS2ZyrmNr2Izr4NGq5fA71iV0HilwPYDsI/7nSRFaa7LevPFESKrsWFzBsd1CoBJaAEckYg1wmFZ6zCZvxgJmvbP27aWvkJIteiKHQHbhne/RaqtS2K7isBJaAElIASUAJKQAkoASWgBJSAEsgnAmnmxaGeCM+V/wAAEABJREFUrsYnhG5vHfMEXHlpOQH+0e3LpQVQAkog8wQqp3Ny132AA7NrZT5yjbGFQAyC5xEeMR/d1wSZ9a1gzDXoFb+Gk95JBvip4WeLB/6N3v4MDFaBE9ibel1vfyq1E/+Hp8jIAhTFpqbQyZH3Cy7PzyM5SixTybyPeOBG/GXswk5FaCf/o7gOcE5D17z1Dw8zgBNipyCIvyI8ay0PHXVWAkpACWScQM4jjMW3hsGWSdK1faC/I1R8A7p3PyhJ0dRJCSgBJaAElIASUAJKQAkoASWgBJRA1xJIN3WHA/rp6qpePhNofpO2iFlsK/btWjqpVQJKQAm0EDhn2gBOttrFQau1uPhvDId2IfMA+QzAu7xnvAS4j3Fy+y763A/jzqT7W9z/kPIN9xdya7hN07ZDNc0Y80StEa77aJ7kpXPZEPulIDkN1fW34vwne7cvsh8aWR/uThGmBK7Zv0NfGUgRcbu9+0RPYRi7CIAbD2twJ8KdnLz2iLrdzuGwXQDwPIy9RtsduisCNMGYqzBlpG1POp6+fbs+gssYwWEU9mX5P31rf3ZgPvPxFYOwTcMsbh8DxLaL/6L740DCzfp9yf35QLt/xsOB4EDEYpdBFwFAjRJQAjkhkNtETpvN50zZlfef5e+ZBktgcCtWkYsRHr44t5nS1JSAElACSkAJKAEloASUgBJQAkpACRQ8gbQL6AAiUJNPBH4CzEsQ8wIHoZ9kxh7iIMp/AXkKgjoAb9C9458dNuYXxvE6ZTrlqUTcBozfPAmRF+g2m2J1uFGrBJRAwRHoHTgWcMvSLNcvEEwFzHggfgRCMhY1FTujZvR+mFT+a9SWH8bjMoSC5ZzA3YdxnkC5mHIT5Qm2L5zkE8P97Nj+TenEzQlHvA/IS7kVPIltv3sD2TM/MOqnKA+kFCO2jX+G944Gnks7qckJTYZqt3WOQnGfXyP8tl1oll7o8C5xKjJ9eKcprGViNkMvdxvqdp0NTythTk73z4BZhCLnFn+dHPs65gvAvQ85r+Ntryn8jHSM4P/wY+k96ah66oyvG8i2aArPFdsyT62VPYyxdfFttkt3AaYSDo5DJLobaspHUvZDTdnx3B6L2op9uR2Z8HOd4xK6xjAM3makMUr61uBoRKO1CE8bkH4g1VQCSkAJdIRAjsOsERvCfsVOTNV+LYibhF3E/7ehKDgefxi1hPtqlYASUAJKQAkoASWgBJSAElACSkAJKIGMEkg/MocDqIH01VUzSwRiHJB+mTKRA80nIuaeArinIo7TOAB9BvfPQDxiJyVORdycDDHHcbLNTrLZyZ90PqE7n3E/zHguAjiY7TgnMfypgJyGoqLTEwLGH3dPgwmeDNc9CTBTALxGUasElEChEKiaPpRtwJ4wks5b3O+wHTid7cFZnAy7iZP+LyFcNmclFMIJfvt51ykVn6C2/GnqXoOib85BzJwO457AOC5kGLvgqJHbJFaSuKXrtF46ij8yD1cgFjslpwI5D+PG2QnHdPLYAR15j4GqORF6RkpB/AyIbd8Dp8DFSbzvH8c6cD6MeZBxpDdxS0Vy7EU5G5G56S4gYSjWDyf+Pnfs4jJuvKysB3F2xtSpXdcniYTGQmRjrxy2uD+F8KiPW/bzYxNyvyW7P+W0fre9nowJw+CbNGC8jSAqO/Vb0PaN04AzgWkdQ0nPJib+zQye2zPhyEmILzkHNRV/waSK53HF2O88I7F+k0e9kNAtCv4ecYYFzuJ100Bpx7UtxyAaugThaW0nyTyTVQ8loASUQIcI5DSQETQ2bQ7IMAj/kDBL2Ee4EyiZDNsvTDjpPyWgBJSAElACSkAJKAEloASUgBJQAkogowTaEZlD3RhFbVcRMOZdODgGxt0PRdFaDjQ/gsvGvIVJYz7CZRVfJAana0d/iyljP8Ok8g8wpeIVTsQ9hOKBlyPkHAXIfhxsuQfgtA7/rWAXc1D+NsDZA0W9j8c6va5GbdmjmDjqNUyu+BA1ZZ8jvNN3CanZ7fOEW+2IN7l9CKHYpWiSvRn2aMC8t0K8eqgElEB3JGCwM7M9kkO1wq23NaYBrjkUoWfvw+TR9k1zb91kPuFxkUT7NXn0C/ih5FqEoocB8b3gundSfTElM3bet/7lsKkYRNnGfpJoV23bmiux7atNP2tiIpD4TwiXzUkp9h4yacynmDzqfbbvL/Je8l/8WHwdimIn8j5iJ/NvaEc2h1H3JFROH8xtejbo2s+o2y/M+PU3SnkfrMDLg9OPN73U09cSc6Kvsv05jID81VenKzzDY2OJ+3mu6nbbdErd99jP2IJtyvq+RTemCeJcjHB5OgsFvKNatfF4ILFI0fZfvfWW+sj/uHsSTOhQfFh2KyaWvYzL9phHt/ZZO5k1hWG3KbsVaDoMIr9mBB9S0rE2ryejKeBfv9KJSXWUgBJQAh4Ecuocfr4YAWd7pjmEYm2U/6YiKjWo3cF7YRWV1CoBJaAElIASUAJKQAkoASWgBJSAElACHSfQnpB2UDLYngCqmwkC4nLA/BtO2Z+LooHbYWL51MQkW3isfUPWpJVCeMsIwqN+5qD/dNRUHM2JoC1hjJ1gWcTwCylPQbA9astPRc2olxDebi5OH24HZ+iVyoqBzcuVnFiqGXUPQqERzO8fGep7Snr5o6JaJaAE8ohA+O0+kMCuzNEgipfl9S2vAeZsThR/iHDY9VJM2/0WtjvhsT+iZsx0tnMnYEl8U4a1CwF+YbsSZ1o8VNt+ApJ68YNfpM3nZS7vI++jpvws3o/24z3kK54T4xeMfg7vLQdDnC25n561E9Qm8RM2X/gGMNgFJSXroivMxfVbAGInM+BpxLyMRrzj6d8TPaKhbeHgBAC9KV7WZZ25G0HYL4F46aR2r64fDsecBUg6n9OPsy7PhokfitqKu1A74nvcK2xv0DkzjnHU7v49r5l/wXF+xchmUdKJty/1z0b1jJ2or1YJKAElkGkCuY0vEhgIkR2ZaC+Az7XgPT5m/ojL7SIvSdWPgBoloASUgBJQAkpACSgBJaAElIASUAJKoEME2hXIaZe2KichYOxk+wecxHqJns9RHqY8QHmKbjMAeYvbHwFOr/Bfwhq3Dq45GcXRG2An8hOOnfw3acx7iJceBMF5cOQ3CA04APaLAZ2MNhHcvvlWVDaZJfgt47flTDi3/IsA5lvKR4C8CqCeYgfE3+bg+2fc/wXNA0NQowSUQBcSiDRyAspsmyIHcyHmRhQ993oKvY57XzXmS4Sip8B1DoHI3YB8DTVdT2By2eMw5lhm5A2KS/GzfdjmH+mnsJJf0YCX6WYnzw23ya1gVbjx3egplNxZYwSu2HRX80mUE8rOY+gV/dlHp2d5haeV8D5/Ogu9BcXPfgg4f4ddtOin5ecXntaHaR0NyFZIbZqo8m9EowegdvSb3M+OnTjqLRj3UMD9NxNgX4j//e3m1D8W4Vn9/NXUVwkoASXQXgI51hezPtvkbZhqHDBPIR44OfHlJzqoVQJKQAkoASWgBJSAElACSkAJKAEloASyRaB98TrtU1ftZgKmkYMeswC5Dsaxb06eiqh7QmJgN1RyJH545wi4oaMRj/0acZwMI7+DoIYDJE9w/+8I8HhKxROwb0Uig+ay4fM46X8zJpb9E5laWNCavbDEMLn8fk4O2sH+O+lsFzhcx+15HNj/DVxzKhz3JIQCx7GsJwDuyZzIOR0GZ0LcyTDmSerOoahVAkqgKwg4i0rgYm3fpA3eR8zcn5E3//0SCo9lezLqBcTd30PMBX6q6udBwLgeHh11FoPieAPb7EsYg/1kPzc+VuQAnPVRsY/G8l7hLe1iuccASfUlmqMQnhZALk1V3SDALWeSfSjJreBr3uMaeN9uTK7QA10jgYN5v+ekvG/Zycu5E6H+7DP56vl7xoObAOZwKqXqt8aodzf7YL9P/IQSA2TV2p/XCMXPZBr/oDBt/ve2zLs5BE2RrbxV1EcJKAEl0AECuQxiF80hYL9mMpTJcvKfz8JTRn7GfbVKQAkoASWgBJSAElACSkAJKAEloASUQDYJtDNuhxOz8XaG6bnqhlMjwNuAnMEJ72MQNRNQW/Z3TozX4fIx78EOBIeHL8Ytp0cxeeefMGXsx5hS9jJqy6ciGK0FAqegSM7DxPJ30F3NxIo3EHL+gGj0ZIRi1ZhUdiNqyh7C5NEvwPqFR34G++WBmtEvYcqYp1j2uxGMTYIEToZjDmSxr6HMp6hVAkogpwSKSjhZt4Zvkg6ewZSKX3x1Mulp05pU8Uwmo+wxcUkWSmoXZiwc8iTrydQ0+gZDMeDHVF+UWD6Trnsf4KZq/zdDJDBm+YBZPgo6mwKyA8CSw8MYvIYlsex9GcMj2bx1rp6xJiB/BtCL4mPNqwiV3oTwlum8Ie8djyu/gcha3gqtPvIcXExhH+znVpesb8Nj5yLuTGE6qdsykTUhzinUzcYVzGjVKgEl0BMJ5LTMZz9eBIM9meYsOM4ETB75P+6rVQJKQAkoASWgBJSAElACSkAJKAEloASyTKC90TsM4P05XnqqXUogBsFliGIUasrv5CTZJ7i8fMFS31Q74bGNqBn1NcKjcjconSpPHfW3Zbh87FcIj10IkdT1p7XsEyteJLtz4ZjRTNp+Zjp1WCqqVQJKIAME4vH1GUsRxdsaecXbU33yi0AabW9HMnz9sCbEzGQIfkgZPO7unlKnrcLk0T/AyD1tnZLsB+AEzkninh2n8LQgXLMdI9+A4mUXwZg6/Hnsj14KPco9bBzAvRYi/guKgCjEPQHh7eZ2ik/l9M0Z/mRKCitfM183YHLFh1TMbf9iyqiPWY9uZLpfUvysQHACJszcFGqUgBJQApkhkNtYrt+3CbXl+/KZbjdMHPUa0nkWzG0ONTUloASUgBJQAkpACSgBJaAElIASUAKFSKDdZXLaHaKnBTCwf59AzD4c6Khs16R/T2OVbnntlwJCZdsD5jJK6kmmdONVPSWgBLwJiLuht2eLj2uiLXu6yXsCIlnLov0yg5E7Usfv2rfmU6u11TDRv3Ayvamt00r7xuyLCS/aBSsreWXcoalXf0AOga8xcxCXJ3xVeorn4VMDiDYcxeKOpfhYiXBC/I+YNOYjH6XUXvZT045MSK1IDeP+G6HyR7nXNfbjbx9nn+Ye1u80vqwVD1M3e9dw1xDQVJWAEugSApqoElACSkAJKAEloASUgBJQAkpACSgBJVD4BNpfQgciTvuD9ZAQhpP/gjfh4Cx88O20nJX6nGkDUN0wGlV1x6O67mxuz0Nl/emoqh+Hqpnb49wZpR3OS/jhXqiuH56Iq7r+DFQ3/AFV089iOsehcmYZLnqakyHIvgmLi6Z+k0h4AuUTQHL7th7UKIEeRkACkrLEjqyaUkcV8oSAm918GNSnTEDQ/reYJ499H4JpKeJ2EI8el0InQ97upjDY2T8ymYHSsvf8dXqI70ZrDQPkNAADKT7WPItI7DYfhfS8Ln5lA/IiqBgAABAASURBVED2R0pjvoMErobtW6CLzL3j4kDxDRD5JmUODA7AhOnrpdRTBSWgBJRAKgLqrwSUgBJQAkpACSgBJaAElIASUAJKQAkUPoEOlNDhwLdOvHqBE3wC416E4DfPIDGw66WYIffw7F6Y0HAiegX/wxhv4yDytYBcwe3lcHAV58g5sOz+Hb3ceziBfzTC0/ogXWN1qxpOQGTg3QxyBwQ3cPtnGHM5xLkSkGvhxG9HsPQeVNcdBZsXZNlcuc0iLAn8AyJ/YD7008pZxq3R93ACbjz1pBTcdXs4JS1+KwEXn7fu+mzX8PHz9hJzMz1dird1ZC9Uvpj9BSlO7NcQlHpnxPo4f0FXTiwjT4ztFwTco2Fcu2DCZ0GRfA4jN6LvLp3/ySN3yZEwpldKAgZXJn5mKaVilhVqdvoSLq5LI5VSxJ0j0tBTFSWgBJSALwH1VAJKQAkoASWgBJSAElACSkAJKAEloAQKn0BHSuhw4DvQkYAFH8a+/W/c36Oo4mmEx0WyXt6q6UMRabyZA903QGQ3wAxjmgMoRRSHwgFwGUz3rbh/IMTcgqbQbQg3rMZjfxt+aXVEgw8wzF8gOJDKW1IYF3rx2J5/m4Z9m28TQPam3IbIkqsQnrEKsm2uHrUEH379CPN2craT0viVQI8mEAx8nUb5yxCeFkxDT1W6nIC9LWQxE4FgGouypB/vST4TwR75M4HX6PMGxce668OJjPJR6LxX+O0+gBwDPyN4A6GRs/xUeoxfPLIJ+yhnQaTYp8wxGDyIosiz6OyiifOmDYJdCCJi+ynwMZ+gaOD1Pv659SoeYBcAfJoyUZGDEZ7FayilpiooASWgBLwIdK171TNDUF1Xh6r6GZjwgn127dr8aOpKIBkB+wXEqvqLUV3/MSpnjEmmUvBuYeOgqmFLXq/jKU+jqv4dbueQyQ8Jqar/CFW8lidM/xOqZ+yE8Nt2fCg1lsq6CYznRcYxIrVyjjWqX1qb+XqT8k/YsbYcJ6/JKYG8IpBoA+pO4/XwGarqzuvQM3xeFUgz02UEdpkWRNX0O1iX3sD4ut27LB+asBJQAkqgZxLoUKmzPIPQoTzlQ6AlgByP2tGPorMD2EjD2Il2caogOIravSiprADSGw4OQzR+ASfsOImB5Cb8WD9EGidyQH4sFexbjgzLPU9rrH8viJyOqFvZqZ8b8ExjBQ/7dYWaiochOIMd0cYVfPVQCSiBTBCIyS+Mxn8xk5Gt0RTcm5N8th2gutq8JWDc7GatqLgpjQRYT6T9XxEKFf/A9v4JQHwKIXZx2y64eXYI2TLReScx6l6U5NYuBHSd6xD2yyd6hrlqRilc9ypABsDfvA0ncCPCYzt/Ly8ObQcja/snl/D9N8Jb+rdtCbUc/WvOyz9SpiayOiKxHVLqqYISUAJKwJNAF3qEww6c0n0B2YH39K3gysE4LYv3bHgY+3WaqnpOWtYv4mB0tEVcbk1mpW4uJ02/RlXdF4z3WVQ33E35Ayc/D4CdWB0/cz2EXxuAqVMDHjntuPP5T/ZGVf1VqK5vLZ/dZrh89ZmLb0L9t5wU2KHjBc5wyMgqG3CMYT/GuiGc+HjWU+++H5UKwxpJLDK8eMZ2rDvXIVL/M8S8Bchkyu68ZjfndjCAQQkRbASRchjnYsB9EZG5P7Ku34uqhv1w4ay1cNZHKy/+PHfGKnDkUMAZCAfZ66+jA+ZwXocmcgJDbkXZheUaCTsByoOc2Mq6/RPXa1VdjFuX0p7raz7sgpUL6/tmNa+HmwAq60/GhPrvmT/bpkSZbnvzavWjS8NXTz8okWe7gGTCjMvpvphi/S2H9jBIrltVb+Ox8Vn5J8dBSxLptf6z9b267h2mmTx8dQbbuUzGVVW3EJUvVrQWIzvbF9eAOAcAfLYS8LxPXw9dYcbP2Ihtiz1H9hxaiWf8fFXVR3hv/ob1mffs+re4/zBlMiqnn4rqWSNQWbcx7KKg8DTWH7aVmeYQnlqE6oaJWHbPzkwZm9sTyyzKMkzIaZvWllFZ8fZs9w+j06YI4ICc/YwwE1SrBJSAElACHSOgCwCScTO4C7Vl/0zmldTNds6rGrblTXh/VDUcyc7ECQkJ82EpaYA2jvZt24g5CzCc/Obts41XGrtBDo6vh0hpb0/dxaX92NFjZ086MhhxPkriv+dDcuoHugvr10Dl9F+hqu54dnSO4IPEvrh4xlbt+imBYPAeluc/MPzzLJB6KAEl0CECpokTcuZz37CCIRz8ORWX1q/vq6eeXU9Asnz7boqslkYhWafS0FpRJTx8MSeTG2DMdyt6tTkOsC7uiC8bN2jjlrld+zaYMcf6Rij4AibwkK9OT/C0g6g/2X4KdoGfMYjCOFdg0oiP/NTS8zPCgerNANcOTvsFmQ8XT/gpdImfI48w3fkUH2tWBczWPgrqpQSUgBLwJ9CVvtE9h8K4nPyTYj659ebz5h5YJdIFPyX1SYwYZjIPN8I11wHuDTDmH5QGui+kZMhKfwjsM7VdmLYrYI6i/BmQ/3L7JgKxmYgsvBFvrMEJBvsGtZ1YQGZMnzV5f3WnJ8pny2jMjSzvXSzjs4D8hHwzBuykOpQ8yJgdaxF3V4hskciNkTFYLTI6sV+o/2y/rbq+DNHoFXDjT0FwFkT6tymuYZ1lH958yXr0cUJgvqU/3XjEHQj6QuRXgPsQgtFH0ee7CzDhhV05Sbkq7ET6+LqB6OUeDxg7wW5D5JdssoZd0MBrlNkyZnU4sjca69syoEcWbTD6aeJ6Fcdeq3eS05O8Xr9iii4llSV7/I4jg+NgJ+lTaXfUf4tLDYKBt2HklkReXdzAvvffGd1TlFTtys/NZQLLxnC2XbJiAv9jWOCdLeKIu/Uw9APbZWNuZvn/Tb/XKVFK+tagifHMZIB/cfvX5rzaOOUJutn2n5sWG4v1AnhPgprlCEw1AURio3gOOHFr2DYLJ/+dA2Hbx+UUc3AQLZoLOHctO4+4mak+QPkf8xfntvNW7IIkGYrEPRtbArI/ZTwc5xaA92oR1kN5EJFAFSobDkH1C2vDLmpEhszmiLMsM5aWEeYm1t2pgLzF6yadlyzQxixkmBksyz8Zx41L4xR5EZeAbTlya+yiS8O2394jgCJAxsIp2ZR5FKhRAkpACSiB7BPoYAq8+XcwZOEGexcBx3ZCUpcwzM5T1Yzd2L24mje82+Dgb+w03wFxrNyKSBoTKNHg4Qx7HiBB5KNxpBqDmw5MmbUiszE7VH+FyB3shlj5Gzv9tyPSeFXan9q7ZOcFgPN3CD5JmZ4qKAEl0D4CbugnXptPpwgU4PW3J2JyCR8IB6TQVe9CJuBENkmjeF+noZNcJVryCuvaG8k9W1wNtoFxNucDtLS4ZG4TGzCK6Q/zjdDgfhT35n3JV6vwPTcZat9I+UPqgpp/oOir/6TWS0Mj/HxvGOFggvROoT2bAyGfp9DJvXeT+Zr19mXfhA1YNtkMdiDFV1E9lYASUALJCXSZa/OkASdRZQc+xzq8n9r79LZwYuVpLRzPZMbtT/XVlj+O2vILMbniPNSMPhc/lp4CR05mOzye8m07krOD6YtZph9h8Aq3MxJizKs85oSYWX6yqTViIQHI6hA5mk7XAe7tiAYv4zMw72N06ay1X5aZPPrBRPlsGWsrzsFHZSchjpOY1rmM/n1KKjubZbit82L+zTjsZNynTDBCSWZZJ4QTccm8cu0WWg1GjmSqfSiAoBSuexzOmVaYzzn2p4WGDT2fBb2ZciplEFqNwQLuPgSY8+HiNI5ZncJ+douYU3mtnA7jnkdGd8EYO1kNiPDZULaGI5fCBDjeFb0V0Rk3c9znFnqOp+TfGNZZjxXD4DCWs3lBuy2DMXsggOZFIMiBmbjrO4nrtabs9ygqOxkhOQkOeD7MJUydk4+8ernjY4dA8EcMm3Ec7IIOH8UOe4XDLiaOfBE1ZRcn8jq5/FyEyk8BnJNg3DMh8trKcYsL2x6KcyGiLJMtWw3DTWbba6W27O1EmHsljsnl/020yzX0q634HYrWOhFxw/hlPHVs+8FNSvslYCbCiZPhgJMY51mw6VixL2yFx7Zpk40wb8WMkZOS/N+trBRBYv2yluV3XxnI+ncQBEOa05BiXiP7Y1HRus3HOfz/5+E/orbssmXnsfy36FtyIozD9sfwOdK0Oacp8xXjOZ/HsnzMetJyv0br/WmhZ2h7HxDZCeJcTC63AoGbENntWNgvV3gGaofHuHFx1v3Hl5axxtZ/lhFi29vJjCmNn1qkFvA1y/YnuIGTsXYJ63/F75fGWVv+NERMQiuX/5oad+Ccx55tkhzGfOyJ8DuhNm66qwSUgBJQAlki0NFonY4GLNBwjYDciUC/d5DKnDCtBJGGKyHuP3hTPgGCHQCxb6qVwBojn6HI+c7uesoF01YHZCLD9kX+mj4Qcy6qZrV0Fj0y6jpzyWEufR2Wx/7UwBBud2RH7BSIezeqGqYg8YklanhZ24FZfXA9wzzEuNq3MtgrTnVXAkqgmcDlZQt5TT7DgzkUP1vC6+8IRIv+AfvbiX6a6le4BIxJvfALSGfAOTmjy3f8nvWsjrIkuULCtR+MuwcufcXeUxIOGflnjMCVQ2GM973XmHmA8yTCW/bse1Hl9MFkdS652/4NN572Q/Z5KmEnYjxV2uPRZzVqp544MWYmFi/4nrr5ZXtFf4Y4dgDKO1/CFlkwDLGmod5K6qMElIAS8CTQdR5NgYF8vtubGbBtNTe0BpxolP3Qd0F/HnWtvWV4FJPKP8CPpbewLT6L93u/vobN68MQcwAnp3aExIfDRRmk6BAgMC4hUnww4xjFKbsRVP4NIK/C24TotSUgv+Gk1X9RVXdaVhZF2Am2yyq+wIff3g1xK/n8zHEMeBsxTyJecn6nxTSdiVh0HEJmdxh3JBP8E2Uhpa0NwIkNaOvQZfsR2ZN32x2WS194fnsVL++2nEI3PThv2iBE4pMgUs0SbE4RSos1rLPuIXDdUxGKXcdJpH9gUtlTmDzqhYTYn76srbgL21Zch5g5B25gLOCw/uNdRmAoDmV9Cq8Lw/Ed8yvWuSE8zj+76qrrQmDbp+KlmROsA8hh6AoTFhfh8m8wqeIJzF9wBdsYy/VWZiVO8bY2z+Jeho2G7uOtlGEfm9eaUV+j6LmpMLFzAFmhj22+JdtJmFN0Fy5nmaw+0jTh9RsxpeIVhPrfACN/ZKgV4qbL8nY+7NdOBva7BpPGvIdwip/7StRSse1vskUpjTB4nNH/hvx343ZHGNnKV2B+B/8vrNRDzB6+cdg0HDMcIntBMJnXjL2eXKxoBDbf3s+lK+q39zjSuBVc2HIHmoMaXs+yNVMdmZWfrWlOJP3/44fPQ+3I5xCPn0+eL/oHNN9R52wgWMHtdnBkZ7hmL7Ter1vvTxLnBL89P+ZmxvczxcuuAoN9IHI1onMfSHwNwEuzM+7h4YtRM+olDJLLmZat/7EU0dl6MglF0evZRr+P09m7pZJUAAAQAElEQVSvSREg697254iErCBrtUnLzn8ciuj8VGMFbYLorhJQAkpACXSQQIeDOR0OWZgBP0TcfTZl59L+XtAaof9CYAfFh0LQ0pHCMiPmHTQa/8GGYOggdjbWWBYoT/cMNgbifAj0yZ80/sIO8P9W0hAJQOznEs35iAbvwgV20cNKWssczh7WRP2H2TlO9UCwLIzuKQElkAYB4WNx4CUqcgCI//0tB0zM/kDTW6iafmTz4h0j/kHUt2AI2NXv4hyUujwyK7WOh4awPsadR/jQ/pOHRouzORiRxf1bDjKzuahhY95ntoPw/uQVo2AmZIld4GC8VAre3X7lyHGOgIADRnagyLPEjXA52Bwe9YOnRns9GiOrMt31/IOZhXCc/+H6fZv89brANzy2ESb+AVNeRPGxZhhMbE0fBfVSAkpACXgQ6ELnQICTCWZ35iBAabbCVltkNxQH8mdi1S4E+KH4cYi80ZxJj/+COkyqeCQxOWUnmiZXfIianb7kYP3XzcJ962Ynr2rKb6LbcBiX5TcfMUavQXxO6Mgwpn0lBi25kH3pPtTNvL13XBzBvs8DsvJzONoY1zThMk6ydFYm7/4TLh/7FcIVn6B29KuoKb8ETnAXpvQaxzUMt9YKIF0/zjT+zYEcn+AkJoqwvBkKxz0I4cf6Le/cjY/sJ/mLgxNY3tNZir6UVmu/0nAN1ikdwfP1LCaP/gHhtm9Nt6q1bMdJnNfBL5gy6mPW83sRKuFEqXsex2bsiy2t57dFOQ83p80OIRrdlXVxK+aO9ZD/my3rozka4bqNmw+76L/ts04a8xFCZWcxj5fC8Lr0zYqsxjbkakyo2x3hMMvgq5w5z3DYRWj+bNi2sW2sxryPxlgdbNva1r09+3YiX2Qay59iohefIxB/Dhdsk6Iv3ZL4pZcKTIB13/RqcWG1RZSMn4Rxd0Zt2X5sr27CpDHPcTubx2/7ijGfM54YxcOahYz3A9847BcRJla8klhsM6m8CqHYNjByIIxJshDA91nLIw9pONt6I/g1BCsu2BkMxxyA11cZlEYsOVARk7i3OLiDiRmKh5U5iBffhZoR9Qn2drHfFN6P7MIVK633J3svn1TxDGorzkBjySZkfhcjXEwxlOWtkA6wCh33heH4RHX98Kxdb38YtQR9iu9mfmwdYJIe1uALLAw+wPbaf3GfR/CsOJf22ojx7kkppbS1rNfxX9FBKGqVgBJQAkogawQ6HrHDG0+848ELKqTtCHzKQewPfUsVnj0IIjUQ7AI/Y/ANSmLebw7agXVBOUQ40eYXUSf9eiXuzZ27ERuzCiS+A8LTgp65CTnsmMuPnv4AB4jY2Q2FLsRFs/v76AHB2OvkYh8yfdXUUwkogXYSsA9FRh5jKPsJSG5SWenP8bu/IVr0T1Q17I7z3+idKoT654qAm52E7G8ERueexshTvbnlwgk8Tr2O28tGvQXX2EUp3nGIDOXtYy9vhXb62Lf/QxjOUN4DgInBOKnDhz81f/6Uyj3SLqnn4KmcyLL7Xfcx+v8fh9emg6POyJRx4nZQfjXf6Iywn+B+6qvTlZ4Bxw4cfu2bBSOrQwKDfHXUUwkoASWQjEBXudlFgsY9ks9qvD+vlImBcOQE2E9wr+TVRQ637LCE4x3pLH5tRwbFJCZTYd+Sxk0MOJfiZTkhhYvYl/4tn6Xtm3Jeeh13D287j2W0ixE6HkdnQr7/xetwMYGzbcwD2cAE2SXo2nubHWsJzh8Hg62TFK2I7mMQ7bc58yxJ/LuXk/1EfECO4TVp+2xtFzssgpG/YlH00g6/PWrfWi2quBZwjiCUaRTv8S16drkdumQoBAcwHwMoK1hZDRH8Nitf5FghpZSHYYlx8nISRHjdINXX+TaEQTXie26T0/r6zdAoXPdtlqXtQ+f76LuLX3tH9TRsbNFCluVLatrnCG5WsIYlBr5AU9C/H71CMLY7DggV1hhEefxPpnMi2+s3wQN0tQmPtef9UaBpV7bZU5mdRkqzFWn7VnWzWyb+x/fYmdEk+7KfMA+jgV7b0j9/rOu+zHM2P6MZ+rP92YHyE+DicMb7OiV5vaMHIHz+Ndciumv2FgGM34Hlc/znHQRv4JsvfkS+GNv3s19aEPC+uUKmBPbvdFxYP3QFHz1UAkpACSiBTBLoRFwORKQT4QsoqImwA/QyLi/3nhizD1eRpr1Z6IMpIYq3FczDvGLvxRWLG3rDwSB2bhzvSDLgY9c4djYaEU7eB9aF/dyjZ1yl5Ad25D0VrEcxBMcj1FhmDzwlPHYuz0WDp796KAEl0HECRfIvtjv/aUcEpdQ/jNfunShacBnG2zcQHu7VjvCq2p0IvFW/JYdcTmaW/e9xBp/g+yAHU6jZGeuav6YM7pgzYN9eSKmYhsL4V+zEcjk1V6EktyIckDLTYd+qS65R+K6Wd0B+w3vxdikK+z+2D7fhyrJUg5cpomnjbRdpwLFvqfgtPADH8eYghi+QryZqvobIt77ZE4TgmqGcFAr66qmnElACSmAFAl122PjTFoDYCbbkYwhG9kLvviORV0a+y0p2Egtrg5PYb7qD90Lvr9GI9KX/eYiGjoAdT8h4ZjjpLvgm49GmG6HtLxWXPE8OD7KcUQbjxBKKue06u/j5NZifYyH8S56LYfTfA+HnuzafyfPWPtdN19oakN8D6EdptRyHMnVw2M++hmMrra4d2YbFRU3ZdIQCJ7NfeA+jsF8V4CYPbVRGwIV3+yPyKwxazAm+fMi72Cnqm2HkL4DMg5dJ1GEpQzx+KcbP3NBLLePuu+/gwjhzYNCmbZOvERafyVOkZyK9IxCW2RjW0yRBhGcRmAspWYR0zTubC1XtNWD71IZt0SMIOeM5+e/fF2egnNva3b9HkcNr1jxAvtHm9I3/c0+zUvv+nzujlM8Zv2OggZSVrZEhkPgRebEopjV3RQ4nvSUTo9itMbZsxaA4+hSMewGZv0FhHWnxarsRCCA7wUglzP5rISuGeUl5z5bP82osIrp4COAcQhwcG+T/laxswqfacdnp46yUmDooASWgBHokgc4U2oGxN7jORFEoYZ1GOHjFtzQbrTKAA87jAEnegUJbY+ahf7/kHVqr1gt9yb6X3c26mEykYErgBIu8Y1oSYSc71QIAG3xVdudPTjmZY0y9VVZRAkogwwTCo36GMdWMdRalPdau6P0tAnILIgNvRDY/jdaeXPVY3SysHbM/b2PkD+wVbJYSq8GDnfr8I1rMogV1rI/+CwkMdkDTXju2hOjcJrRoDUawO8XL2ndO3sXigH9/wCt0obhHdzuI/R2/gXPe8sEBK/N3rFOa6hOe7aNy6fMBRr5+6kAyFyWxOan1ukjDlP7E/s4vqVN31qaOHazkRq0SUAJKIC0CXaNkP6/tOOfwGZYDwZ5ZGMhn6tMR/rTEUyOnHnagXX4ChM+qyLypHfE9UHIFmTycIvLV2N+ZiE3XWvntuRQB0/I2zEFaillSCg9fDMTtAoCfYIRjTNK7yyYC7CKLYGhv9iW29CltKf0PQrTIry77BM8TrxOmlcB1L2Ld2nD5HMlCuOZ+bF328fLunTgKj/wMbvASxmDretu3wumUB7b5yyPH8zlmgE9uhkCcw2DfZvVRypmXffnIMem88RyCyD4IuNfAfpE0Fxl8FwYO+9pAbGlyjnBydulRx3f6NxkGjrFMdsvdFayxziaC0gXxFXy8Dw8/HBCU8FqwzxFfAaEwwhlcoIwMm7D96TS5jnm2PzkH5jvzCwBK3RHM9b6U5FaYusH+GLgkO/el5KmmcO3NewnSeH5KEU0y7/DYGIpW4diDTOFzLu/dyZQSbkGI7It402k8MZJwyfi/eIo+iVmY8SQ7FWHULvz0ezHAIdMjsfEQO9bSqZQ0sBJQAkpACSQl0ClHh6HT71RRuXCtiUMCfp0A3vt7DYTB7mkxsOt5v2k0nrpRhx0JWk+FDHnYJQaZSSaCkOFgv0e+hvYlP/H2Xy6Y2Q/z9ixezmnFg2Dg6xWd9FgJKIEMEagd/S2CJUewPfuiXTEKhPrrQ+R4bhsQ2f12XPDS6tzvGmsHD4QDu36pC4pZzq1R/cLorErl8xUYX7eBX1Yy65cYGclclImBMOcGsjqWkfq3zzA/wsH/Ua/z1v4WpshfU0TEh/D4KSl00vR27EDICoOkywW197L/4OpRS5Zz7UkH500bBMjtFP+BKME7mL/g6g5/VhZeZrDDNmZ9L9+Eu53oEPyM8Ng8GxxJ5K7532XD53EgZC4PvPuC9ATctbkJUtQqASWgBNIk0EVqgxt3AORwCP/gY+zAeewre7/1Ucqhl5FGwLhZS7F2+LeIyQTG778oTWRtxOO/y8rbluJ8wvS71g4Ivs0MvMPawRri9MHuGzg8zr1df/BQ3n8PgpH+vomLbAsTT/ZpbN9geeW5dnAbGBycYI42xpgfeNSAcRLnNnN28ojPIfIXXk//y1ykGYqpb59dmC87Vic+MbK/ZfZDfP4mPjpd4JWiq9icoyDP9b6ILPkLnznTeCGpOVDH/19qg9pJ+mVtp7iZrU82hWQiiXrLOjw4xSRpm8DvPm/Pewnrp12A9BiKlrzbxjcPd8UgFGQezTSeVzuOap8HMpfP8LQ+jMyO2aRqBwcjgPHYZRqvDYbocjs3Bth7NrJjwltGsF7xQ2T+UIoEimDM6QjP3CSFXna8XXT9Pb21ZIlFR+6ZPPQfI3Jlc8A5AGHTNfd+ZlCtElACSqBwCXSuZGyY0+psdi6VbhGagwKu+dk3qw5GQVDqq7PUU4qwRoksPVxxp2kBB4VN+p+0WjF87o9/AYq9B9q/XRAAjM8XAtpkWKQYvYz3SlSrGhB2+O2OihJQAlkhEN7hS4g5hm2afXvXPnS2I5lEp76IYU9AUeQdVNdfisrpm+f8N1/nFQsfzLzb2eYSDQXkWiDwQlbFCT5HHr9DzoxIp5MyRhB+bQAmNOyJ6C9vMv+HMk6H4m0N4oBzPxYvzNxDqevex4dw/wV4MHvDfqHAO2epfWx548YucBAf5Z8QCqQaEPAJ3s297Kcii4N/ZilSDCyauZDA0bALOKicWfuDrYODfOMUicLFl746+eH5DbPh374KVgFKbZmpqlYJKAElkAaBrlC5akYp79UnsK+Q+lnYmN5w5Vc4/8neXZHVLknz8rIP+Sz8TzIyvukLRmNQU+YnFIzxT9c3UxnyvCjxU4ovJPrmYgL41mcsJENJrhSNffs/VLQzjGzPuurX37NBAxDnRFS+uKo96HZi2I93zRiWM8mkjPkJ0UBm3tZuC8YuvG7s8yL5PkjnRkp+2Avr+wJyAYTjTEhp1kPM3S1vvgKQMrttFAT8k/0RwASEZ7D/2MavsHYN6xifOTk/nm65vulrxwUGsB1eDBfTcMkuDJ9u4K7S23khjPMyx2T4XMVzm7FssG2IBrdljCMZZTrPGPugLLQtdXuGPX14lNfQX1lX/Mf/gUGIuEejKya0xfg/cTCnrgAAEABJREFUP+bqTNmyxyKHwcimKZMU9IbBHojO5PhbSm1VUAJKQAkogfYQ6KQuOwMinYyjQIJLHE19/B+SJL5GOwrLwZHEIHbyIFfutYg3x18oXf+wnjyHbV3jMPITwjv4vREZpE6Sh8+20bTZN+5abY5W3l3sppgMWjmIuigBJdAOAnYAp6a8gQ+cpwPyb8B09CHZDj5MgONMRd/+56OqYVvYwTeoyS6BZS9jdCid8Mz1UN1wOKKLroRr/gNIuoPRn0Lce3D13qkemJG2KY4v4ADFvb76IgPgBPb31UnlWVW3GcSp8FUz5jaER8z31SlUT/uA3yu+P0T2SlHECAymYNLI91LodczbLuyBSTFpxPZK5LuOJZDDUMbMYdvqP4BjpBhzl7A/nsN8aVJKQAl0awK5zzwnE35whzPd3SjgPcA+v9rP9CZv30QCELMrinttR2VJhCn4f2IgeJzyi39RZQjZ7JCYJPdX7K6+zzPjEU7AcSKuMcj93NqN1mW67kE8D0NaEo6QtfdLFwZbQyJHwvaBWgJ0m82lr5TCONt45DeIonjAw69zzldus4j96QfZDthFjp2LKyOh2T6F3D15zstborMPSQsAsVusbKQfHLMfGuevs7JfF7sYUgUWMRcxipflOKMciyhOQ2Lhg5daD3PfwZbXCcLI95zc/R/sWId1ymcJiwvjfgQIn6ulb8a+DnNhQx/YT/sDzXXc8LkJsC9zGSQzIv1hzIk9atFeIPYBGdkXYZIRaeu2C5peGtjWoUftR2esDdc9GcK/5oKz7yeR5t2V/vN5VnaCccsw1WTn/rNSkuqgBJSAEugZBDpbSjbQnY2iYMIb9AkkH8RoLaLrzGvdTWNr315L9dD7Pm+j/mmmkVD2VcxCPiR9CkjyDiNomgJ8EEH6HSORuQzlbTcsyZ8V5d65VB8l0M0J8JqeWPEG5/4vgpFKFuY7SkesvZdswSYizLbiNgxbo4qTy+t2JCINkyYBI73hBjbChJmb+Yr9MkN1/XBMeGFXnpPDUD39QkyovxOx2N8huIGp2Qe6vtymY5cA5iYEY7PSUU5b55JdmjiQ+CT1fQbNpRcfPndFeFY/6nXMinM2yxzyDGywABL4q6d/oXs0Pr8BIKdQBsPPGDwNlN7lp9Ipv/4/CcP3ovhZAxesj34q+eAXYB7FYwC6NX+mFCVBW+ZWB90qASWgBPwI5N4v/E5v3j8PoKyVSFzMIu7/BwZv8tjj+VDYD5Q9cP5TqdpzRlEgNhb7GMZ85VsaA07aynq49B3v/ohvBHnuGQq9yeeBExEwdwILvCYJsleIQMxOiO/BBAIUwLivAPJPAMnzIgiwLh/fPd9YjJQA8OqzDYEE16d/dmyo6U3CfSk7kbcz1vArq8LIsYAUo9n8wGeG2wF3TvPhSv8FkB3gxEYgbBzkkxFEmZ3nKNMpMYqXXQ1wz0XQOQosvJdSj3J/e47DtmdVOPgeIj7PlHlGpTj4Pc/hT5R+6PdTUUZyV+xsRAZ7Ma5SCq38DwZsB433eLY4Y1Hceysq9wwbHttIRvZa8y+vmLVQHPdqZ/3DdnffcOJnIfbjdbVZS1Fs+/QwXFPfcrzyRsB7D/bFzJn9V/ZUFyWgBJSAEugggU4Hc9jR8Hho73Tc3S0CB42/9PbNtJiPff2X91wT85pSdODkYQZZRMlz63wLBOxDiHc+TYjszBBvhRV9HP83B7+N9F0xhB4rASWQJQK1o79F0TfXwgT2YwpPUfwGHOjtaYN8uNwBgireWx5GVf0pCId5n/HUV4+OEhBsz4exf8HEn/UVx3mGSTzKc/t/gHsb4IRh5DjKGLq372HW4G9YEvgL7AMzA2fM2jc0YvF3Gd9rFA+bGKDbGk3xLTwU/J3H1w1kvTzGX8k8ippRX/vrFKhveGoRnNAhgOwCJFgjqTHyFQxuQdEjHKxKqpEBx6GMQ3rxn581ELejXy3xizezfk4ij8Y3UpESlATEV0c9lYASUAJLCXTBTmTeRkx1V4qdcORtQj5DLHYD76uPAmIHhJHE9IIx+6GoT/YmIZMk2qVOJX0XQhw+N/vkQjjh7BgOjH8aRCEa+xWlmtH3YWLFi+wvdvR5ooNkjMA1to/bOiYRY5/lKojczQi/oHhY2QxwD0B3e2aJRu34S9/khRJOisf3wFkfFSf376RreOxCwMmPn8yKLhkJx7FfKAHPt2G78xJi5i+sC8/6lHIAIMdizvOp+pvIsXEBeQtwK1mWt+FvVoPEazC+4WiAdd9ft/B9r98nAtetQrBoHIKRz7tPgRv5bGV+BVP6K1yx5+JO53vq1ADsT1wAmy+Ny5j7AbmZ8j48jVmP9Wg/hGfn2zXhmeNOe4hrx/eTLw5bGrn0RmMPHZuOOkPZnu7P5qV3Aocx3yJm/oWAa1+YYFuVcF3xX4AOe6KX2ZJbtUpACSgBJZARAp2PxOEDUWE+fLafjSBQ3Ad+xpFP6O01yEGvNlawIUJ9mgdJ2jgvt1tb9joMZiznlumDxUtsjMb+67i4M1N+6jcU68+6tHZ6aUgUoSL/B5pYtH96camWElACGSEQHhdB7chXMWTIgYjjBAC2vWvktiPWtn1bQXArIru9yIGJHXHujJYV6B2JTsMkIVAEyGAAdrY0laxGvUGADADA89DuQaJFfPj7M2rLz8TVoxI3FcaTWfu/bzlQY+z90OceK8P4wLk9OvK2TjAwjvfb5ofX5DmPwXXuSO7VA1yb1tqApbwIMKxX3EtuOZDvPoTi6FMcJPd66E8esj2uC38SnivbhviEYrdGJP8XALg2j8LM+hTFMm9aLH4a6qcElIASWEog1zvht4vYn9uT94eWwVxx4eBplAx6g5Mt/0d3nwlv2RxiynnPcHKd7S5Jr3EJ70uGE6N+qbO5NyQK23Xz01O/dhO4eIadBD6UdAmZoQ1eRHTJ4wgJxx3EvunJfgzdV7a92e/YG5FdWxcOrKyRjy4hdwHLmvyrioJSiPwK/ebsjI70m9Mpr+u+RG7fp6OaNZ3wjFUAOYjPKavDGoF9w/lJ9HY/hRO8DjD22PqsKLaOlGFAYOcVPfLg2MU6vV6DyO/J9xuwEYWnkUEI4DpMmLkXevxP8InB5NE/ILzTdwiP9brWPUl2mYfNq30Zo3b4tzznptP5eHnwYFaZXzOeIgovAfwABO7E6qu9x4Nn6OY1vlMKuxA8Fl0XPcaEFvIaS1VXBAj1GCJLC2rbkzhGsy4Nh/APcCHyMuLOiwgWPcU2d+ZS3ZV37MKBXyFb956V01MXJaAElEBhE8hA6Rw23KlueBlIpltEwbt6dH3fnAYj8wBjHx591RKexmyKYkmj8+T+kfo/UbJje7EfBwg6bn6Cg2v9gxuBkQ1Zlzb012v1dZ8FHvHqeDYruYFNm3f0vxJQAjklcPawJkwp/xdCvXdgN/9ipm0/955iMJNaXlZkOAcmnkYvE8aEF9dnO9GZ9sgrFXXPDgEOZMt7HOA/F0WxCdlJoiXWe8fFeR95DhDvN/CFT99GdgNmDEJ7zPlvcGDXtb8H6133DF5GsevzVkR7EuxmuuFpfeDEr4ZgVf+cm3fhuDdyUM3//u0fSWrfPqsa5iXiryg8l26ig+Ov18W+xtZZny8q2OwZxFDcy9hdFSWgBJRAKgI5918wbygHgH8FkZY3ic2PPP4XwltGUBz/kPfu/8DwL1nGhJOQBsdj3p4DknkXnJv9motIwLdcxjUQaQIWuL566tk+AnahsTHnQND6U1ELETC34Mq9FgFPzUU8/hQj9JqsdjjGsyNrcVn3mkTty76S+I0jbc5y1aCprgLZmIgpKbU/HfcxuXaNNUYQi2/NxO1XzYLcgufwPfZVn2JfNYYfil7j8aMUk/Bb8Z9IH0DOw/lP+i0QRpeZmrLpMM7RfHb+n28eDAbCjddg2BqjcdrNIV9d9SxsAuGwg1DoKACbUayN86K4Cx9/9RHsGA/Mw7wevrEeHrItjDsWp83uGfXINfZ+zWdKDxrW2SDK8XCfFxSsUgHKFmuuCsfZH3BaxgfMfJbyAVwx6gfYL/1AruOx9xcrBEcgUteycJSaapWAElACSqDDBDIR0IGIk4mIun8cUgS4qW5QcyHOQ4B43+jQYkQCiLuHtRx5b2pHv8pOWIoJdu/gKX1S5zRVFJMwseINX6Xw8wG4sheEZfZVpKftQCFwF8LhVIMew6mtVgkoga4iEN5uLorLroJrfs0HQftJ/ycB08EFY6Y/w54NN3oNKqdv21VF0nTbQSDRVuNBAGeg+Ls7kenP/jPilWwMs1nX3lvJfXmHCsQDaX5tpiVgaN5w7nkvKjOIQ/A0JDaHej3McvA0EvoNDO/hfiVP1AepxaQxqc6PXyxp+n3L5sIsSaHswATyc8C2bcYl3od1K1U/ewka46ZtMN1XAkpACXgQyL1ziTkABlu1JMy2Sv6JtUrfTByHx8Zg5Da2c94TCoJt0cs9AD3BRGLFMGaQf1ElAgPy2iXir9cNfC98YTNMaLgS1dN36vLclrojyHWPNvmYiUDoucRxmOMOTqSenYvXADFIZkRWg8g+2GiVbrRYZYdGluntZMVpcWNfCWVwnGsRmXkCMv1p7/BwO9LU3Ba0JJjTzaXP83qT/Xjem1+6aV6IdC+C7qeJfNwyPEo+f+Mp914kITIaRaW7JPTz8V/xyDqOP/6RWfuQktwK+Cfb8H8YQ7ZKNZ6ZPA51LQwC0X3XZJ0/DdIyLmvwKQxugV1ob0s4fyHHdc0TdtdD+jD8IRgUW8XDv8Cc40NYIM4D8L+XFfwIafzBy7tg3WPYFgK2jab1OXY2QiVPQqT5HhrD4wBeonjZwWy7ToX9ipSXhrorASWgBJRAOgQyomMbcysZiaxbR2Jc+wDh32G2gxxuvIEP9u+kVVaRE3FR3TopdSX+d+rwoZT/M24TY+jNN+n2x/0MmhbdmjJYrGhdPliNS6lnFcS8BTf+it31FcGOvv7qqQSUQPYJhMXF5IoPURT/K6KxU2AC+zNR29nvyEKAEgj2R8D5P1TWH8p4Om/7Nxm0PoR4x8YBVvMRvV/Lqhi8xrx8xTS6uTWNMPgvXNkfoaIzUVM2HfbnIXJRqsvLFwByPzjyQPGyg+DG9vPyXMndvvEkzmi6r0lJbgVfwbgzkItFDslz0HWuF88axfv3WeDQoW8mRO5A6JsHfHUy5TlvVYPUCy3Zd3X7It+NiF2kYN8u8cmpLIYpTbUo0ie8eikBJdBzCOS4pOGX7Ge1T2eqJRRrv2P/4FacbifW7CFlu5Ef8R5yN/fYdvP/yraUTucg3LAat4VtA4nJf/9FioKFiLsfwvax0c1NMLAd+4znQmSLLi3JhfV92Zexb722TlothDGPAn3mLM1XUb2dwGE9deNL3ZbfsW+Q74lAcetil+V98/HI1iFjnmJZ53lmr7l/x8lhczkiS67HhIY0v9joGePyHhKfzDowDk14Hbk28ZAty6EQNL+tLPgCbtr5B80AABAASURBVPQB9ueXPaeaCPPlPM2sebVPvSFyMsKv5efCD3uOF8MuyL4awLL6zIMVbIDHo+DGL0fV9KHcV9vTCNhnXhM5lvV5WEvR+WxhHsJ6pc0LYqzj9fs2IRC4jbu2PeQmqS0DInsk9Sk8x50gsNcOfMxHWFz6s49/4XmFpwXhuicCTnO/zRjeN+U2hHdYtpjq8rKFcPFPDtt4fSlUCGYcovO35VatElACSkAJdJhAZgI6fGBY1kHOTJzdMxYR3vjN5qh+1nuSwJasdrRd5Xwvd+2KZ2587SoIyp98NaznpNFfIW7+wnOxyB5mVDr6EwDGLOKDdE3zZ/NS5Mi4E6kxgOJvDR8NDabipzad0GQhxtdtQOdNKGqVgBLIBwLhsTFcPvYr1I56EuuUHMTO/hEc7PmY4jWY4pVrTtrJMDjyN1TPOMRLKW33ecXCdtM+XHgHsW9ZCU5DTfn2WZXa8p1QW24HZ7zzkq8+zW/MxHg+pyLm7oaiskMwpewp2N9QzHWeI4vu4Tlt9E3WgA+kvhrLPKMz1oZgFB2KKB5W3kJR/EUPz8J1Ds/qh3j8XBhZK0Uh/4eQVCJXC0ESC3uwJEWeBM2T6ynUutjboA8gbPfgbYRlLY21ty31jk99lIASKFwCuS5ZtMku8F62QN7g35g86v3lsjFO4nyOvY9u9nPg3CS12yJqCn9CISYj2Ifxv6ca8zlcseMJSUF1H0fD+zD4vM5tV2c6ENsWEPsZeI7nwJq3EZBnEN4yYg8SEg67cJueApzl62/Cs+WfyFpwAwe2HHWPTVH8TRj8N3VmzarsN53I+vkIJjTsifDStzpTB/XTqBnzJWrLXkdiEa+fYhb8jBzHWO24ETfWOtdgytjP7N5S2W63uTDuE4DxXiQBXreNC4cvDZNvO1ePWoJQ9O9wnJuYtSjFywYB2Z3dzn/D9vGhpkcRiE1fl+W19+zmdtDgM8Td+5dbsAeaiSNfZ5vxOPe8bCkgpyM8rXXhHwrSnP/kamwT90pZNuP8F/YaTKlYQArxou0g2J/tppMolchMtkEPgRMEWGrEIBafDiPvLHVaeWcVIH5Yj/lJiZXLry5KQAkogc4TyFAMzQ16hiLr9tGIbAanuAzGiG9ZYrgJBs/46iz1NAejqn7s0sNkO/YN1pLiabzJPkZvl9LVNgbIPxDDG0hlJjTsCiNHplIDxJbrSXYi7kLic2xIbsJhBwHsz3MwNLmCuioBJdClBOxbX5PL70eRszMgF8DgdQBeK3/plcya/jBuJca/sB1MivY2WXB1swTYTsO+MW8HtGz7at1SCSfWzXdU+h/lLbbHMyi3QuRwhAYMRG35EbhszAyEE+01usTY32oVudE3bcF6vK/ug9RGWL82Zhl9BvVkMYz7PMJj56aOroA07Mr+WPwICHal+PV55gPxS/DOV7ae5QbAKmsZ5sm/TTGw/VdOrucmSx1OxS5SMCbgG96w/VysCwB8GamnElACCQI5/We/Ymfk+KVpGvkMYu5fetx2x23kpKo8TKc4xcPKiZxQsF8U8PDv5s7haYPgyCkQ8WvzozDmLk6WftPNSwtUPsuBfbMdyyGUrrPhaX3gBPdmBtamsMuHJfz3DBY6nySO2/6bvPtPcHE5n128J1ElfhQmvLBZ22B5vW8XaQOTmMcPKKmsPVebsg4+imjDg6icWYbEm+/GuqcKm1/+E+o3YTmObZOpdxBp+neb4+Zdu0DJRT0gL/G8GyQzBrx2nUNx0ez+ybzzws1+pWwhprAMN7HcTSnyNBqR2OVo/oJLClX1LggC9rnOdfblPXr9lvLEuP8snJIk7QInbsW9kXXp+xbdZJuRiIQOZFva/dqGZKVZ0W0qn82K+/yKDDZc0Wu5Y4OZcOO2b7Occ0EfnDujFHH3D4C0PmcvIqe/wrZBWMHEQl+xjlg+Xi9H2kVJe2GVyLKFpCtEoYdKQAkoASXgTyBTvnYANVNxFUI8g/lQuA8unTnQtzCJFc5STZ3PKSms9OVN8WxcMM1/wMO+aWnkDkbm1xGjdy6sfMwO478wpeIX39Ts58WMuRjCP19Fehr3G8TNFagd/S2PfOzeawCyJ8RygxoloATylUB41M+oLfszr9WDORBhByTeZlZ9Bn7p29YKtoPjXIhLZwxu66z7aRP4mNwvY1t9Ee8xj/HBLPmg1nLRyZeAM5ETusezPT4Yv8T2Qk3Faagpuw/hLf0nXJeLJ8sHIfNXlmmRbypifs/yi69O+O0QBLsA4v17vMb9mezsgyt6lImFNoRrTmWZ/fo7MbL5L0zxM0t/O5IBsm5/fs3w3NqFKn5J2f4r+1d+KvngZ3qzDtq8+mXme0SdmJ+C+ikBJaAEAOQOwmmzQwjKgexjbNqSaJT7zyAUeLflePnNJ78s5D3lSbbdPs95ZktEQ2Nw+FS/CfLl4+0uR+HZvVi2s3jP3MEny+ynmedQHL/XR6ebeHHCOFC6DYxs0eUZbiwexjzYRaGl3IL33C/Yh3zY843NYnkUwGyKh5UhcAOnIjy1yEMh/5wnl38EkctY7lR9p9a8c1IGB8CJ34/ooskYX78butPbvtd9VIw4TmFh1qSw2KaJdfEeXDE2efkXLfiaek9DsJDbla2AzwumHMHIFmzD/J8tVg6dOxf7FrJrLgGcvzFRrwk3etGKORaRpvMRnu39DEQ1tQVCwP4kK3AA60brs9G3kMBjKN4p+Xhu1HkPME+y9MnHbkQCvOefiPDzq1KnsKwxgtfqh7P8Z0D45126H2BwZcoxce/w3dOnt6lgxu09lZuEnYkiPI9k5uqRjRB3Gr3syyXcJLXrw4nt3q3uMUmLoY5KQAkogS4hkLFEUw1KZiyhbhKRMJ8HoSk2mlt/W1v2NsQ5hx2HRn9FOLCfpAuG9sPUFAMeqw95Di7uQ9eaCMs0FaGBL/lmww7eSOAo6rLz5KtpPV2ITEDpszPsgafYOKPRsYyz3FNHPZSAEsgvAjVln6NorT8D7q8B+TMHTlK9lYAWE2S7cCCi7rEtx7ppFwH5Go57NyaV3wxHxkNgH75SxGCG8vxshMWh1/gw+wn+Mjb5QFiKWLLuHa74BJCn4GukDOOn+68mb/ylNwzspxC9YxJ5FsXlH3srFKiPwZmsM9unKN3nEHMzakd8n0Ivs96bIw6Rz3wjFQhg+iGfB6zthJCBHYhjXn1L8yUWFekCAF9E6qkElACQQwarRO3E2n6A0zypCnzDZ9SHgZHJv5Zz77g44kH77GgnVl2PnHIiQfbDFmty66HRHZ1vnh1CrOkY3pN+C+GkiXcZXodxxiM89kdvlW7iE35zMPuTR0GwVpfm2H7GPoDdmJfNl+bDmEcwf6j3VwztAmbB3dRfQklmhX2foxBZc1mcybTyyk0MosX3A3I9gJ8p6drV2E8+FQG5CZHgHaiuK+8WC3R++HlTOOagpYUUeRcIsH1a6rL8zvX7NsGB/eT5p8t7LHc0jPVoD/xhZn5/9nxKxS8QqWHO7+W5i3LrYaU39Y5HU+NpsG/0emipcwEQSLz9Hx/J883nOmPHnl3W5RcRL2lAWJLfj3uVLYLgQd63foC3GY6mol0Yl3irdEOfC2dsAEcuYM5bFzhyd0VrFrHctyCOp1f0Kejjsx7rB+MeyXkO+/zKohr7BcD7MLTEo57w3hNy34Qx06mc/FnWPguLszcai7u2v8AMqlUCSkAJdD8CmcuxA5Fg5qIriJgG8oZXk1ZJgv3sJ/uvS6krYJz4Nd4cso6v7tnDmoDGMDtimRsYWJx4tjW+6bb1NPgCbug6tP3NvLb+rfvDVt8CcI/Ask8DwduY2xCK/gvhcPIOaGvAjVZfhbtnAtIfapSAEug+BMLrN6J29KsYstofYWQP2HYkvdz3gpGqvP7kYnrl6FqtP5W9Cxe8d4ADYPAzfSD4HUrjf/BTyg8/cwcM/7wzU4pA8DBvb/oEZSz/b0jxtm7grwh7DI6gQI0d4AVOY+kClOTWsnfN3/HhtzOTK2TRddw49hVcv0Ha5sSNrIKYM7T5IA//Ny5ZHXAGp8xZIPA5bhmefNAkZWBVUAJKoMcQyFVB7YLsQLwMMDtQHCbLyQR5DcXRl+B3v+y18zcw8giA+ZRklmMOphyN7taMV5IpdDs3O/n/OSf/jbkEfl8bAj6HQTWKnn4T3d2MrxuIyPw/sTzHsChFlC60r6wCEz8OIsWJTBgzFzB34Xo7ppJwSf4vEn+Sem8l97Sushr/Hwk7scadbmEvGz4PTX2vhev+mflNDABxm9oKbF9wQ4gzDpD/YtjQe1E1a2uEw/baR96Z8NtFMLFDAGmZTJII6+IzKOrzCfzM+9/Ynyn5r49KCevEwSiWIT46+eFVM/IbSLwGArtYOu6TqcFwMB6lhmN2Plrq1b0JzCvux+v3YBZiEIXW8JnCuQ+Tt/MeU07cywMv8dqZRTEMtLI14Bi2OQipvo67csj8damaPRQhdzIzeCAlRElmLb/72R7ciMSXf5OpFKhb337s98lolr25/TfOu2xD6mB/AtSryOGxC2HMPZQFSVUEwjq2IxwzIm/vK0kzro5KQAkogTwgkMEsOGyoeYPLYIyFEJVgc1TV3QT7gOFXHjtJHiq6nQwbeFNL3nFqDS+ogAkc1nrouZ28+0+AnEmJIBOmV6mNRey/NKQJYv6AyTszDz7a9o07cQ6DkR19tKyXnUJ4BYsDf0Lz79NZt+RiH7AlwAkJ2Sm5groqASWQ9wTsIqbJ5XVwo7sxr49TUt9fBKvCWVJFXbUdJSBiYLkbmcgHtuSfv1wWdxFEJqGq/hSc9VjzgOkyv/zZEz5wCl7xyVAAjrs7Jr85MKmOHbh0zW8h/EuqQEeD1zF55IvoSaZq1hBA7JtvRfAzgtexcPUrcvrpfyw17E/JXB55TSLRK2EHIxps/s3fxGGe/QtgDfapyNs3X78gHrVvV7DMvnrqqQSUQA8nkLPiD1tnEPsSB2HphLZwUDd+H5/l/PsXdkLBRJ9lWJ+v6si6cORAhJ/vje5s7HNr9Yw18XlTFdv5a1mUNSnJrO0Hv8Zn5t/ho2+e4uC3m0wpJ27irIIJLwxrv9RvgosbdkRVw5Goqr8TDj6FyOkQJAYZcpL3ZInYfl508bGAbIWEEbKVf6C2wmdiP6HIKtrICVR5kkeNlGTWodJ+iGDLZJ5563blNotQXDEFxj0NxnwBA7/J4RWKYVhmDITIIZDYLET3uBWV0zfHhfV9yUJWUO66w/j8TZifvZmB5vpn3M9gnEeR6qfM7h0Xhxu4h1x+YFgvuxWc6CGMP3/KmzSnYjBpzEcwwSp6p1pUxPOHa1BZvy/sgiWoKTgCvaMjWCY79tJcbwUvYnHkCfDmBD8TGvk9nMAzENjnkJU1RQIQlCES45ivaY57Za1u4MK823asesYooPG/EDmcmQ5RklghC/k73Ph5SPnTtUmCd2en8Kx+zP4BlNZn68WsQk9g8BCfPh21rS3u9RqtzpiFAAAQAElEQVQg0+BlBGyH3BMwZ0wvLxV1VwJKQAkogZUJZNLFdvQzGV/hxCXOqYjMGwf7aTnfUu30MVzcCEE6n8g9HxNmDfONznrOn/8gYB7lbq4HhP+Nmgrvz6cxQwkbDW3L/J3GMkvi2Oufwbd88KzFV19956Wy1L0pdCDgXrz0WHeUgBLovgSmjP0YsfhFLMAs8MkBqYzIOOjnCVNRSu1fO+r/IM6fAMOH1xTqgsvQt+8RsAu6Uqh2iXewmPcP8yzvId6Dl66sj4ULRifN35LdNwNkF3gacVk1r/L0LkQP+0l6J14JyJrwNeY7DiqekPINOt84OunpmvmM4VuKtxUMRtC0DlJ463WVj5E1mPTqFG9rzA8wYhc7eOuojxJQAkoAyCGD2PaAY++tkkjUuG8hGn8msZ/qX8kunHQUu8jMTa5qHD4/7oN4cMPk/nnuGn67D6rrRyAWPB1wH2E/4hLm2A6ac7Oc5TO8cGISt3LC/DjUjnqsixbUtc3UOTCBD9sveB+ueYllvYfn7niI9G8baZftR8eyL+OcAeFfIhPu14ib28GMIpW5cs/FnAh+imr/oyS3gnWA0B55209OnmsgzP5t7eh/IoBDAdxH+YXSXsvJdXMSHOdFhGQKxtfvhubJofbGk1l9u3DZjY+FwWYtEcd49htQXOr9kw8tionNx198wO1dlOTtk5CayK8Rfn4IuoOpHfEmq/sfAGPLzzbHI9Ni+pPTLfgysp8uAvBg1F2d7XO8CZzE7A+gWLsEJn4jrhmb+tkiLDEgatvB921AD1kTEtgTZz3OCVwPjXx1tuP41S+sjQl1ByGIK8jlKV4Hwz2yy3sCptH/XMxf7UxMHv2Dh16hOguisY1Z/l1ZwCIKrfkc8djjsC/48MjXhocvhsi1bIvs83tyVZGdMbDILlRJ7q+uSkAJKAElsCKBjB47GY2toCIzZONegGgDB0F8ChbmQ5Yrj8DI/1HLe6KCnrRDYGLX4rqPirnvba/fJwKD66nwGSVX9gM47BilSu202SHm7Y+A+A9qA43sBPwb7pJnUw54jK/bgWnXUt+fC9QoASXQbQhcNuYtXtOXIL3J6LXQCy1v8HSbEuZhRsVg61G3QHAJYBalyOAqEPkjosF9OLgZTKGbe2/7IGlQD4j3JLBgNQ7gliVdPBLEKRAE4GWM+ZiDms96eReeuxE0Ld6Tg/iHs26wf+NRQoMojFwFO6jooZIT57j8yHQ+ofjZ/nDddWHfxPTT6go/myc7eWDQ3zd5kc9QHExnAalvNOqpBJRAoRPIUfnscx5wLO+tq7WkGIE4t+DyXdJrp8LiQuL/gcGHLeFX3hizPuLOUbA/NbCyb/642IWpldMH4+K6bVBVfzSq6ycjMvdmsrkdrlzNjG5LEcoya1hymC/pcCOMnIGiAedgYvk74Cwd1GSOgK07LuxnzZtfrDAmTt4PonQtO8GbRjpiECt9i+eS/Uz2e5KFMOgLkb3QFFwnmXfeu02seBXRot/D4ELm9XVuvSeIqeBh+9D9DDhyK6LRa1A5fQzCdoyMrl1hi4pXZzkOZNJ9KbRmLgRTEd429cJnaifGpEzob4zDXqPWJYnIFuQ2rkvLmSRXnk6TKp4ng3NYJv+6LzKUfeYJ+HTJGEw13s9HngmpR14SiATG8NzvuTRv9tm5X+Cxpcepdtbu/RnD/9dHjZPB7j4YMGBDH52u97J1+pxpA3BR3TqomrEbJtSfi0jDDYBzKwxFcBpEeq+UUWPbf/M83S+Ca07BpLI7unQBPDPSJfa02RwLEvtllU2Wpm+cx1Bi3lp6nGpnsbwMiP0CKDxMH8CcqT/96UFHnZWAElACKxHIrIP3IHBm0+mesYlszo7CaQjPWMW3APa3gUxpDW9o3oMdrREYsye+m3NU62HyrRgU/fIi47uX/lFKdm2i4yO3IPDNRykTGtT4awj2SKlnwM6CXIXL9vB/IDtv2iA+VNoH041SxqkKSkAJdC8CH3z9Atuxm9LItANxt0tDT1VSERgncQRLbwXkcgARirc12ACGeo2Sp+yNvQ++510ABHk/GomS+PJvgV8wbXWG2YfiY+VhzHM68laUT5x57HXJK2tDnFMgYtl4Z1TwNCdv/u2tkCOf0tgc1s1UfSoHjrM5czSAkl/W/h4nzNYQ/vnmTD7GT4O8F7n4hlVPJaAEegyBXBV08JKd2WrtT5HmJM3rWBR5GMJn02aH1P8njf6KSrfAIPnCeJEA4B6DTdZcj3r5YY35A6ob3lxOermzeY+p42T/QxBcz/Kcx+3RSIwPILRcxo1ZRH/79b4j4WJ3RFGF2pFPIbxlZDm9Lj2QJ5nHC9sv5lKGuZ9Zf5+S/XEJJpLS2rojgXHUcyiAI19z+zAuWa+J2/TsZTvMB+Q+lm0ukhmBwLjD+X9E3i9WSZZ/sGRX7PQdfiy5E5CDyej3ADrwVqtxyIDXqhzPfuTdiDXckfgZCUaWc1tcNAYiI5muUGwJZyIYqwfa0T4V9/mYAR8EeKXyXxIbhDEnIjpzaBK//HSyDBznAsD4jOXxPALbwsFEvNWwKdR0fwKJr3I4Z0HQsiAGi3gp/AMXlS9Iu3CnD4/CkalsLb7wDiPDEHMP8PbPuc/GCDbOWO5+/UbD6+gdmomgPAtx/8HJ/InM1RmA7AVgEKW5zeAOrcvyfszrZRLcWBmicgx+KLkZUypSLTxn0AK1g5vWgHHtHEVJSwl/RkD+hfDYxpbj1JurRy3hPeLvgPwEbzMazuKx3t7qowSUgBJQAksJZHjHyXB8hRZdkJ2HYxFx90aq1c6Tt/8BEjwMMP43SeGgh5hzYD9H5EcrfMBixHA34/Ob/PCLIX0/QT3i8l+Ex0V8A9nfOgT+RJ0gxdsaRBFwTkfNqK+9lehj35ArDR4DwQE8ClDUKgElUEgE7O8tDlk9zHbsB99iGbYCRuxEnq+aeqZJwL49Hzf2KzJ3wZjkA/A2KiF3wUZwgv9BuN5+rty65o/Yz+8Z167K976vGrMjH1A3Wi7ToeCePPb7fCcHe82zsA+qVCx4a9+UizXuy7O9D69Fv37fd/S/DTVjfN6MyhEtO+DgwL7N5D+IZUwFIoH8q7u93cEcXBqbgtZC+n/QI980YcHVKgElkD6BnGie9VgxIH8A0I/C2wGf51zcgHQ+JZwI0PLPLhYQZyrEeC/iElkLxj0N4cSkVEvArtyI7TNswRwsEwP7mfGN6bYuZRXeQ5ef9Kcj7SxOIIxG0cBVUFu+P2UqJld8CPtyAAHQP3+siTcwf1e0XyrCDHMYaso3Q6j3ary3TWah2I/iHndybqdODbDu7EW89tzY5DmZIy/Q7XXYumdd0hJOGheVNFB1JiW5tT93IHIU1l+jV3KFbuB6Cyf4aso+x6Sy6xEzw5njOym2b2W4bY8NQMBJIhwP13kOVXUn4fw3Vn6jtj0xtkc3/GkJRE5mkOZzYUwTXPcahMfavhSd07ThLaJ87nkQMN7jVIJNgPhB1JE0Y+1atfDYGOYUPQnX2C+T/OiZGQHPoYxAHFfgsvdbJ42hppsSiMVHsB0sX5Z78zIQeG7ZcZp7k8r+R03bLni1CUE45gQ0jwVTtcttCe8+9v68BXPSIrI53Xjdwo4JDGVb0RvCPyossybG/WqY/mugtmxj1FRcjCm7vMz79Tew7SQ9e6Y1wrbuEEAsU9AY3k+nYuLI17nfPhsPsA6aep9ARXDYjodfy7/F+z6ZVi8loASUQFcQyHSafgPBmU6re8YnKGXX4Y9oajgYVdP38hU3aj8R90waBd0YJlDpG5dNK2BWh5HPAHHhZRwOWEhsF8+4orExELGrHr1iiMCYL+HENvSMw+bFCuKTIfCLqzkNwVOIx1ZLGV+k6FgY/I6BSilqlYASKEQCZw9rghE7Ge1dOhE+d7hDvRXUp90EplT8gnigBpAnkMoI1kHEPI2qhi1TqXaBPwfp8ItnuiLFvIcdjMTEBbXCs3vxPrU7IP3gacwsSNO7nt6F5rHJ2uxzmGoWy6F42RgvwvsRKn3SSyH37rEPmCfvgUybIZF1OJA70u7mlbiogHCCyy9Tgp9h4u/4qaifElACSgBADiBwALh/3wq2uS1feRPD/RkoLn66Q4mvUzQHcO5j2AjFwzqHIN7QOuDsoZMrZ/Ma+xJ3LyfAv8jgWeZgPsXLbgonxHvQD0VeCgXlHt5uLorKJvD+dgQM3qGwnuS4hO+uvTYg+7J+tUxgmu953h5B8XM/ob3GLpg15lae50WeQQ3KEDBjPP27k8dlFV/gw7KTYdxDYOQhZt0u+PQeZ6JCUpvo38hNKFpwE58dtoX9/HZSxUw5sn2Kfn0QjNk5ESOnp7h9FCWr1nPbTsu2zZX3AJnO+DwWSUsxgANwUT3rGve6g7UTmMW97uR5vZLZ9V8UIdgH83/8d8oXkhiR2jwlML5uIOvv0TDo05xDmQf7RZNQvx/QEeOavzOuTz2DGlkPxj0++9e6Zw7aevzC4fGpMGb5e7aI/YLdO1T0aNMkSD9OdM8fivClwn21lkD1jHV47n8H4Z89NvgSjrkOYFuJdpqfAvMZ1yMM9TMluRXZHpHFZQiH/cYlkodVVyWgBJRAzyGQ8ZJqo5se0k3g4D6I80RKgeyfRpSlvL3+JmVcifRwIB9Kvc+TAQdr8G/PuIB/snO0o0+eiiBiP+mWTtmOYzwhSiq7n2d+EmVq5WjuAGQY1CgBJVDoBGb7F5ADOwKfCVv/0OrrQWDKyM8AGc97wCtIaWQzPuf9GVXT8msRQG3Fe8x/nW/2DfbDqv37J3Qam7YAZBuf++YSAA348Cf7iWLuFrg9/8nenGS+gvf5tVKU9APqXAc7GJ5CMWfebrEdxEn9eXw3bvsmOctWmgmdlFLPYA6KAm+l1FMFJaAEejiBHBQ//GJfuHIMjPRuTs3YCaQHgF7eg7jNisn/n75DDIg9znuxfbMwuQ7M2ojLsgV8Hlo5cbYTB7Xlx2FFQeAEGExlHrwWMgxgH+U3aAweAPtlOyoWvA2Li3nzXmCf8VYI7JcAkDNj3/6PREcwPY5tGDs+YgB5DUWo52SCx6QP/E108TRA7NemkNQI+kLkt7iwvm9S/+7meK/EUTv6WSxY7UiO1/yW9ftvrMO2r0WW7SiMIATBMRD3Trw245ClC3HbEUXaqpUvrcI8ngVJTMyDW7ZLzt86/BMb2+78I8f2HoMjHpOlibq1LQJS3q2ua9uHL+p/I8+pXXhvn3fgY/aCcarz6K1un6yq1/IEjHCCdnveX0fxWgiAO7w+3oUJTu/wNZH4/L25l3Hx3s3/K1qB0OkkvPniBtx2tf0SsV6nr3S/DkWOh5GLmLn3KcnbM4NtWJIqxPZYnzpq7VcCjct7ATZMwDDGLop6CMFVffpuCc3k/+xCpBjqYPAGxeMciH2hkHMce+pXAJJTVFcloASUAIDMQ7APTpmPVWNUAkpACSgBJdBKIBDgQE3rgcdWUOLhSmSVQwAAEABJREFUo86dIVBb9jYHBY7n2MCbvtEIhP67QIr+iIvq1uF+flj7OVcJXOubGZE1EHEPTvxGq+NuR92NKV72O8Tc52B/nsJLo5Dci/qcyfO/d4oiuawflZhU/kEKvdx61+w8h4OTdvGHHYzwTltkR1w8YytvhRz7XDxjO15N9jO7qRJ+HuFRP6RSUn8loAR6OIFcFD8S24HJlLPtsn0BcPs27wvTOJkQRYeMGMSd9xh0GgeBk8dhUMw09kKfAfk7EF8z8hu4zmUsRz0luRWsiwAuRlPQMkyuU2iu1+/bBCfAiXNj3yDPXelmDO7LSeuDmaCdQOAGMf6bCpR9x23H7JV7LQLkVgB+X3rYBUVOYXwFgAVN2OuHNWHSyEfhmgvJ9Ehei7exv9iU8Ev/nwCyFScjJ6FP/wNw2uwQsmFMdG+IbNsStWFen4eJv9Ry3P7NOInDjdtJqjdh+JcsBiODmebeaAoMTOadt27hLRcivsS2WVczj40UL8uJYzkacM9BeMYqXkrqnocELmzoAzh7wcg6Lbmz5/kZFA35sOW4YxvjPASDz+Bt1oUbPRSHG9Ydb6Uu8wmPjaFojWdhfxoE5qek+RCwjTL7sq07Exc93fzyQFLFHuK4yZrrsaTjKEKhlS/57wHYn0rhTofskvm2Dj0NmMXJw5si+o1Gk7sFty3pJtdUVyWgBJRAjyWQhYLrAoAsQNUolYASUAJKoC0BOz7X9jjJvmvmJXFVp0wQmDz6Xbg4jVHZhzpuPK19IDsMQanEzVkaxPNM2sdj0siZfEB81UfDep2DdQbbQbo9eOC9mMSY91E6ehZ1Ct9Wzdwegt9AJOBfWHM7J//t5/r81XLtaxd/uPJ/HKRJ1YAUcbDnqlxnzzM9N345/TjAxP9+1si/AE6SQY0SUAJKwJtA1n3sm81i9odg2WSCi+cxKPgRwOkA/uuQnVI+l23zo3Bg3y5eOQphijBbI2DGZm3icOVU2+nCNnrKqI9hAhfwXuT1trsAzqYs5w2wn2VuZwrdVn1B05esHZ/nNP+9AjuwSu7FNJvHsIx5D6GSRxEWF50x7hK7wMMuOEwei6CUE86nIzyrX3KF7uoqBlMqfkFN2XT8UPo7RENb8Zw+Skm+aCd5MYXOG8NxazE4siP3M2vDDavx2jqKeerVErG9Dh9GcfyXluOObWpGf8Vr+j6IJC+rgH1nszsEW6O7mcv2mMe29xqWbyqz7n1tCPrS/ww0mSOQ9Z9xYEpqM0MgKBvy8cHes4tbIpwDN/4AwuvbhQAtTh3YFEfsoj37U3DJn7vETp7LQdjw+fxdtGcZFPfi85XzsA+B3vQ7DcHe+3Lbc61dyGHiB0Ck9Yu8LgQvoCj6FljB0FFjFwga9yEIvPsHgo3Yru+Nc2d6j9l0NH0NpwSUgBIoAALZKIKTjUg1TiWgBJSAElACSwmYwNCl+8l3DES+T+6lrhkh8PG3sxlPJQdOkw/E07PF2n7BGfi8qQr28/Etjl26sRPBYm4CfAZ4RTZDSdE41qOx8DMid6KzA8XoBiY8bQAkdg7P91r+uRUO9gQuJTfjr9dFvlNGvcaUbd3lxtfujqr6UxJfgfBVy6Jn2Diomn4K4IxOnYp5FbVlr6fWUw0loAR6OIHsFz8YGAbIrgCKKNbaBZmf4EezCaoatu24zNgGAScC13j3O0Q4oeruiTUaB9uE81ZqR74KyOkAPL7aYgSQ4ZwyvKLHvE07AAtZZu9PLSPD5qzHitlXORmQAWg19qcFI01rdbyOttRvp2R9wGV/CH5vwA9HJDKSk6rSmnxBbe2nm68Y8RHWLTkEjnMoy/Y8ZT4lHUsmshFM/LyMvlVr+1VRtwyCbSlMA9Z8AZj5iBRt2bnzzvYp6HzN8/mLjdRDVieLg7rl5Pjk0T/AccNk9RDLFqN42T58vLoCr9Ufn1eLv71y29PdDe81Jr4vDK+3VhYGbyMQCHTuemBbGC9mOwh7TSxojXqlrWBjSGBsXv80hv0pjCaM57U9jfn3+opcLyD+Z1TPGAXbzlCxx9lNpq8DyJ4UuyACNPPhup8gElin03XJcXqxjtpFpIw2qQ3S9QD0dlKMU1BLrRJQAkqg5xHISontQH9WItZIlYASUAJKQAkkCBh3o8TW858YGMzx9FaPzhOwn7wPRe9jRFdSOGjL/35WzB9R3Pv3efObp1F5EnA/98syYKrovyoluTXmK4SCjyf3LCBX+zvE0dB+gNi35ALwNj9zQPUKhJrye/GNI5ezCH6Dl/SmFfweGw7ZmntdY5fUbwU4v2M9bJ1ES54PYzgY5fwpuae6KgEloATaEsjyfuKT3e4+TGVTSqsdwnvDrRDzWqfFyNMQ2bk14iRbgXEqEAuMSOKXX07fRP/L9v0qZupnipc9Ak3x3+Ki2YX/aeHw2Bgc+Qi5Wj7Yv28Foa/wxqa5sNN11NZz4GXev89n/MWU5NbIqpz4OhSVby1bgJBcs3u7nj48ikmjHsH84EG8dsezzj/BibRFaRRKANkHgV6HIFOmqW5VRrU/ZTVKq90GkPszct7j5hEIhsDbCFzzK7xet4W3Sh77TBrzKc/dJJZxJnPpfaWK9IaDy/FV4yGwC22orDZPCVz64ppw5Bie09DSHAr4zIeXO31NuO6rjLcWIgOXxr3yzkBIYF8sCa6xslceuVxZNof3pjMoMygedV+GAu6NiM/YgdcJ2688yn+2s2LHCkywnMkMBzth3Fo7AI5zKcR5pdN1yd5TBQfZSH1kU7ixXyEc1jkpH0jqpQSUQE8kkJ0ya2ObHa4aqxJQAkpACbQSMObg1t3kW0Nn52P+U5tNAuGxjQjFboHhIA/gpkiK/QNzPkI4GefOKE2hm33v+OKfAHkU/mZNX2+DmxEeke7bTL5R5bWnSQzKnMk8th0w5WEbaxBnPXgCQTwKO4jfxivvducNeRIwL6SRr40RCJ6B8OxBaehmVsX+fmrQOQ2CzdKIeDqGrPZEGnqqogSUQE8nkO3yD25agwO9hzOZrvsMqzGrwI0fg/BU/8VTzGSX2r+zD2VCd/DeeQ/zsYSSxEof3gdOR6CpZ0ykRd0nWd7jAGd6EhiZc7I/U+HidEC6bmGFgBNubgUCC7Zin0RQ6OZ69pcnlf0VsQT3S1ncT8HGAn5GwOcF81tUPmMn7v000/AzgkCArGG/7BVClxlZHZDf5vUbz/AxNRWvwQ2cw3brJR8teskgPhlORL++B3bbsrIUBW3DxkEkdiLP5SZdWE6HzcBIPj/ukPcTt5PLPwLcSyF424fXVoibCaiqH+ajU3heC4tWYaE4Qe9koK1mTB2ztl0/HtGxa3YsuIZSAkpACRQogSwVy8lSvBqtElACSkAJ5AuBCS8MQ/jTrhncraw/EP5vf1lKLlzHfmrS7qtkk0B47EIsca7k4MFlqZORgdSpRKk5qEs/rc5M4Mo9FwPmOYr97U/r0l5ZjMZFf2lvoG6pH8Pvme8dKd5W5EsIrkd41A/eSnnic/2wJohzOc99Y4occfLIHIvokjMQfpv7KbQz5W1/NzUaPw1wf80ovd8epCeMaYLBVTibZbLHKkpACSgBHwLZ9zKHsF3aclk68gggvIdkXOziQ7vaEysZ4d1IzJ5oWmvkSn755lC78xyIXEFmMzyzJsLBbLcS/frZyUtPtYLwuKziC9SU/wuTyv6X1fIUBXaGYBcsNfIWjExAxusqpgCYQ/GwsiGMuyfCr3Ci20Ol0JztOV4w5Dq4rp14TPElLhZesB1QZL/WwINO2NNeCXJi7gBA1kWrMbgbGT/nwvYO98F+iw5JjG2fgCMRCWydxDf/nUQMake8xn702TDmQ98MGwyj/3g0FfEZwgj31eYVgVnrMDunQiTALa35FoI/ItPXhGsqWVe+gJcxWA2QI/DZmCLktWHdL/quDiJXA/I9khuyNHtR54we8eWeVgYlsW0BM5bitDi9A8h5FLaHkjkRTATg/fKFYGNAjmM+tL2BGiWgBJRAM4Fs/Xd4s8tW3IUbr7Gfb4WdhLCfAFQB0mVgb/5xqFECSiC3BExgP0S/no2qhv1w/hu9c9bJDtfbN8vuYGGDFB9rXsWUkZ/5KKhXJglcPWoJinpfDsHdfMBP1SavBjHXYqPVR+Ss3iQtqxgOXr0JyKvokDH3of8se9/uUOhuE6hyehmMnMP8ckCD/5NZO8hp8HcO2s8CTy66g4nHX0PzFyAM/E0vQC5FZN7pOVkEYBcavDnjNDKvZbpsW+FnmHd5lMhf8VNSPyWgBJRAC4Hsbi6cthbv62dDpPV+MQcmXo2asusyLr1LJ7Dte9+7QNIHEj8HJ0zrmsWq3hlbwYd9kZqyzxHgfcbgpxU8Ww8FIhzUxkSEXyvsz8W3ljib2/CsfjCxowFpZbkYcO9GUeTPGa+noQE17Bc/B29TBMH+WLRoTW+VHPtMaDgR1Q1xPuOdlbWU7ULMyaNfgInbT41/ANuP9E4sCMfZDfYT0946qX1WXTyMrMdR0aFY+yVcc3HGz7lt7+DYZ6IvbCJJRaQ/EDgah+f5V0qSZt46ikHoqdksw9k8sot12B/l3opWSByyPRz3GkyYvilwyYoaetxVBA6fGkBT9HieobVbshCHkfuxyLky49fE5PLLIHI7IC6SmUQ9we5YsxssigmPiyBafD/g/pdFiVJWtiLFEPwagcb9EDat7c3KeoXiYn/6yXV+w+KsQrE2ChdXsB5dRcls/y+4Jp+P8TQTSd7mAOTtHImL6lvrNVXVKgEloAR6NIGsFd6BQSBrsRdsxGJ/724cb1ejVZA+A4NTAeP9cFWw9UULpgTygsAWEHMPihfcyEGiXbL8WXdB9Yw1ETGTARmIlEb+kVJFFTJLILzdXMTdGog8y4hjFD+7Ghznn6ieOdJPKet+W4/6jHXY/o5lUzvTWgwJTEU47LYzXPdSt9ecyC0Q9kzgZ8zMxMC5n0q++dVW/AgXdzNb31FSWYf15Bo0/Xw+ql/ggIKRVAHa78847eRZZO7vOVlwA5mnk8b3HLz+N2or0ilD+7OkIZSAEigwAlksjp2cCxUdDSPrtaTCfoD5Dz6qeKflOLObqu1/AORvfA70vn8L9sBawXJ0BzOxrJ7Z/C3lW0pya7AXIgv/jMrpg5MrqGtKAob32khsB0BGAbCLie3U8zu87z6F8NhUXwVikHba8JYL2dfgxBcWeoY0sjlCzoGe/l3iQU7ibsVJ9+wuoJk8+l1ATgPwKcXbGlkLGNzxvNw8OwTI6exbDYU1Bpy4c2/ClHL/dK1uRyRW9AHbpkcZlOnwfzIr7t4YtvqWyby6hZt9BipqeoHXzuWUr/3zLDvCDUxBbPq6MK7466pvTghstPomfBY/oE1an8ExD+KqkZlvB/kQBcTvgP8XI1aFcX6Psx4rbpOn/Ny9bPg8RGN/gsFTzKDXWICdDK9BbNZojhc41Ctcu8qisRDZbWkBjZmB/iUPLj3O6M6dEcC1P5vEPgDEeWcAABAASURBVKBXxLIhgs7BvH/Ze7yXkrorASWgBHoIgewV0wFMPHvRF2jMgkbeND/ExPJ3VNrBIGD+ByPeAz8FWl20WEogjwj0ZV5+DcHt6BW/EvaLAOEZ9oGHzhmydrCuqmEbGNd+RvMwppVq4OB7Duvcn6HUNZr2EJg8+j3AvRgwb6YRbD3q3YSLXhidhm52VMZJnIMNT8EYn4fIpEm/imCIg5ZJ/QrDMWzfmoyfDZFh/gUynKwIngX7UxD+ivnlK2LgLnmW5/5uSjr9CIcDZRNgAtdgQv1uSAwmZ6hIduKsqm5XhELXQHAJY2Vfmv/9rGErB3M3SkJPIjGwBjVKQAkoAX8C2fSNOesD5nC2Yc19NGO+4ITPVNwrXoPjnc+N6z4Gg7e9I5JenHw9CvaNb2+lPPERg22/uY/luYIZSv4lACFd4Bg4gbN71KeFCSRj9qKGPrxl7sn41qVY28TjJ7EkyP6rPcyCNAUaGKudDOYmiRWEADkW4ZdWR/4YgZFNEXWyv9ikaeErELELMpd4Fl9MH+CHjk/mfBXZEo4cwfiFYu07WGL+AZ58ZMOU7rCAZbLn/Cvv6GUdwNkL4S76ST3vjKXvYxfNFJX+Ew7sM/pin4BCv/1gAn8m8g24r7YrCdivjYnswXt26zNelNmZhmDsddZbw/3M25oxXzK9vyPx/OIRvWAf9Ou7o4dvfjlfPvYrPj/+gVLHjCVnJlgPbvx6xHfrHmVCB8xZs/ohGDidIftSeIph275bMX74vMRxpv+Fwy7b7pfI3f5sklf/soQZ2QcxZ+1MJ6/xKQEloAS6HYEsZjj1oGUWE9eolYASUAJKoEsIrA8jv2Fn+xZE4vejqv5cVD6zaqdzEn67D6rrqyDuvyEYB5HeKeM0cgUuL/8mpZ4qZIOAwaSKl+HEz2Bd+DGNBLZAIDAJ42dsl4ZudlRCkZdYr/x/w3L5lOOA8xxQ6jOot3yAbnkUC+0K+3uMsAPT8DL2wftKFD35updCXrtftgcHJ0JX8Py/nGY+SyE4CK7cji8bL8VpD/dKM5y32rkzShELhiHO7VQ6mJK6jaMSr69XYUKXIzzC/hRSwkX/KQEloAT8CGTN77TZIRjnAMa/CcXaGESeQjHe4oGhZMdGl3wORx5j5I2UZFbouAsibvcYfB83Lg7TeBfb9zuZ7+SToSLFHPg+FaHF42AncajY46z9Ek91/aOorrus3Ys7ikOcfJT9yYwTBPxv8DV53gv7U1Y8zIq1cTvu9Yz7F0pyK9gS0cbjee5tnU2uk3vXddnP2DzryV65JyeO3SdZ9s+yktZ1HxUj7p7MuIdQmIxpguBf+HJ09p4Vw+IiKLMh5kXWLz43JFJe8V8f5mNfxL5bf0WPbnUcHr4YwdI72De+FAbJ2y1bIEEAwAHUOYsnIc2+LkOozTyBxvnrQGR/GOnXEvkcPoc8BLxgf4a1xSkLm5jYt8I/8Il5ICCn4Ar785LIfzO54kP2fSqZ0bcpya2wbXfFjnVslFyhm7v2ie0Cg7FLSyFoAJznlx5nY2dh7+9g5AlG7fFTjMah3w4wgdGwi+x5oFYJKAEl0FMJZLPctrHNZvwatxJQAkpACeQjAYH9WwMiYyC4ClL8PSfvH6EcgXOmDWCWhZKOFVS+uCqq6s5D9JePIHxogthB5aI0AtdD8J809JKr9G8ysG8GJ/dtdhWsDoNrmL/ncyqV059EZd3GyHdj+U3c5WUYZw9mdRHFx/IBTVCBgHsxqqYP9VHMnpd9ewXggHvaSXyFeHwmwltG0g7R3RTtW2gGJzHb61K8rUETr4WjEdntOVTV5fZ6WC69+r90+JORtSO+B5wjOUC70Lugy/kEIFiH5a7EoIFfoLK+BhMaNmR4WU7L78B+1WR83QaobpiI0vjnMFJN9XUpAUoa1sxHNHYwEnlPQ11VlIASUAJA9hisFl+Lke9D6UPh3A6+gcF/ER6V3cmEK/dahLg8zrQ+SaSb7J9gTUh8T5z/ZPeYcJq8+08Q9wYWxS5MM9yubAVD4Dq/RWRe9idnV069a10Sg/nBM3jO94KRuQjvvKBdGYpH94XBZm3CTEVthV2o0sYpC7vvfzeLsT5G8bJBQH7NvlT+fAVAMIjXzo4d7l8hXSMGkdjHgLxLcZFp8813W7JRsu1Ta8xvw3495F6JtzpkZRsum4MY7ofIfI/4he7bAO4ITDUB7ndfaxcBFI+6EjCXsxAxipcNksca9AxS1HYFgcOnBuCYETxXwyH8a87DLMybOx3hcOavv+b4m//3in7O+9u9bIPtFwea3Vb8L2Ys5s/nNbuiR54eF0fsmMc15OkxGZ3IdzkC5lSc300WNiSynMa/i57uzxp0ODX7UWhlHs/tfQg1fc+D7NlbhkfhND1B5rxneCYziD4ci+prxyC5q1YJKAEl0CMJZLXQTlZj18iVgBJQAkqgexAQsYMZ+zGz/0av4Georn+YE9iTUVV/CmUPjK/bAVXTt05sq+tHoLruAMp4VNXdDyf2DkQ4kCDtGAgzPzKtm/Dh1x1/M3tesaQxkVcCyDbM35iciuOMRUAOQncxRaPehJEzmN05lFT2EIj9EkDdwFSKWfEPRe9jvOnk01DvXUiTHZjnbiFaI4hE9mfJ7Nuc3PhYQSmvgR0oub0WEouMpE2a2BN9+m/vk1N/r5pRX8PEOSlgPoThn7/2Ml/BqnBQBde8xfbtaco5bNvGYULDrmzrdsbFddugatq2qJ6xU8Ktqp5+db/HhPqnEBBOOJgJZNeOz+uKHZj7AG7gQFwx9rtlGdE9JaAElEAqAlnytxOy8Ug5Y9+WIhTbTs1CUcnL3M++jZu3OQg8jS23x4SCFAOyJ4K9NkV3MZPGfMp7w6Usl9/CBvZD3Ym48FU7mdZdSta5fIaNg0hoLPvpR0PwC+KGE+pi+2XpxXvBNPtMcSzD2ucTwOALxvV3oB1xoIPm3sNdGPcehvaZGJENIXI47AQdFbvcGvbxjIxAv/6WW3azI6FFPB/sh5vkk/IGH+GbvovanQn7lSVH+OwkzYuME2+oO0+gOP5Fu+PqSABpeo7B/NrCvuxDHoV3n++a5x9mLmM2LGz7zc2Mz37BZDG3avORwBZr9We2DgOkdWJ0Cff/gev39VqogoyZ8C5NcAJPs8n1/vKeaxe4ya+6zWR5eGwMbvwBGLkDifYlKa0SwD0RRfOPyf6CqqTpZ97RLmaXkh15zxrByAMsu4HgdRQ59QiTCbJsQrt+CSP3wyDqkZLQfXfEmrbnVq0SUAJKoIcSyG6xdQFAdvlq7EpACSiB7kdAxD5s7gdHxvPh4FbKU5wAewpwHkYA3OIFQP5LmQyRgwEzBO0xxnDgyLkGoeDDuHdc8sGj9sSXn7pBPuTsggvr++Zn9lbIlR0IipmHYOQq+nh/9pSeCWvMiawfk1H54qqJ41z+C49dCGNuTZ2kibBuPgP7hl5q5e6pMeG5DSAYz8yn88UNquWFHQJxd+K5kQ7n5uM5M+Di9yz7G2jPIgCboHCQXGQ37l7N8P/HuvQs6/JziOMRSOgRwJ2WcBPQT64BZHcAvSjtsGIHVt+AmHNQ3NTQjoCqqgSUgBIAssUgUroq27tDYKT13j0fxrkHl+zwU7aSXC7ey8sXAOYBuv1A8bIbM4+7wv5UgZdGvrkHRz1PplXM1reUZFbovw+Ci2tQOb0dC8mSRdVN3GLT1+U99nTKWjB4AZGA9yTSikWyk+qh4MkQtC4EiXL/dhSXf7yianaOxSDmvMa+QB3jT/6cYsB+l7MvNl7TflGDal1shYQEw5nnnbO+KEGWOOzfBFlioaxoXebkRdi3Plf08T8WlJpN2D7sxfClLar/Y915FOFdFrUcZ3djnxcc9w4m0khJbgU7oyk0OrlnN3OtHf0t4qaWjNn3ZS+4m2W/R2Q3aiog7rJPthvzJKLGLlTJQfHZDi6IvMf68TwTi1BWtpJYtDcawUX2q4fJ2oOVw3S1y5QKuyCNz3fmYWbFUJJYGQyRC9G3v30GTOLfzZzOfb4/HOzP9nXtRM4Ftk19GD/N/TRxnO1/YT4Xu4EHIMbnHk7mbvxw2IWq2c6Pxq8ElIASyEcCWc6Tk+X4NXoloASUgBIoDAKrQLAOIKsAdtCL/ztmXUBqUBS4HuER2V+9ji4z9iF4QwSk+3wWzw7MmyW3wTX/4ABi8gHPVpzC2iA4BRK9BOFZ/Vqdc7YV51YOSHAiwS9FmY8gOv4TE35R54Of/T1hU8QBDGyYD9lpRx56Q2Q7jK9vfZulHUFbVO3CofVLn4XrnsPBBL+3tVoCpNz0Yp7WotaalF6UTlrzKuwChTnMYzgHb1Z0MrcaXAkogfwikLXcSHQU494F9h7OHdoZKHaeg3Cgnwc5sUW9ZjIdu5iUm6SWbbA5AoPd7L/J3Jp83ARadzu0tYPbP5Y8wLB/4gD7XG5XtoIABEfDLq69sJssDkUHTXhaCdzAr8iCkycmxvv0nbh61JK0Y9torWHUPY3SOlb1EaL4P1jOyJHpFZ3DlOwEUfLFMYlryGzHvujY/JmwkMHM8wnYYs3WBT48zIIN9u4DyOosewArGoNPEAo+u6JzyuPrPipifdmDeltQrLXPIdNQXPwGWIGQK9PkPMpnILvwwytF+8zzuy559vHKUWfcp5R/CuNeyiimUdTmE4HE87V7JiD2xQzw+pgHY66BfV7nUU7sNbvMY7oPMN3vvdMzW8Bxd0fiGvbWypiPgYPSBdKp+C6r+IJ5nsA2rN4nHj5fm8sxvmFHH53u4VVatBkgewEoptCaT3heH8H1+zbxIDe2ZMQXiJu/kHk8eYLGgX2xKBqyX6hKrqKuSkAJKIECJpDtojnZTkDjVwJKQAkogS4mYLCwi3OwLHnXvRhFA/6Mwp78by3vOny47F4Pjfbtl+jiKkCeQGoTgOAkRKLn8qG/5YEydaCMaPxQ/A0fIB/zjcvIw7i0/Etfne7sGZ3LwXHZtxsWwQ7abIMQ1ulU3k8fHkXxc3UwoQMBeRL5YgweR0j2Q/EzDR14Ay5fSqH5UAJKoOsIZCflEzgpC/kNI2/+fLUxSzglfWXO+2Ph4Ys5n3c787GYktyKbAMT4QRycu+MuzrYoNNx2jeeo/gXRP4KSPK3JWEX0MqJCJqTcPPsEDJpHFk3k9F1Ki776X+YC2AnrgzeRVTsG6RIy9i3/xz3OHJcq0U/yvpyL3pFP285zs0mPDbGPjz7GOY9JmgoSawMYt72R6SvnXhP4p9rJ2PH9nZHk8v+oZHspM54XWzEyaNNIfxrm0jii0zm77hk56/bOqe1//2c1RnbOIj0btGfi4C5E7a9aHHIycZOropzM9NqpCS3ggpE4t2x/52kPGIwueI9hIIn8rnq9SQKXeUUgGA1LNwks+1kRksjveA6rV+ryGjMichsHTMyJrFv/xnncWz3nd+EtdXKsIhBzHkJIi/5RNwbgnGw17CPUsa8xKwKlPbpdHyTxnyPkGvqAAAQAElEQVQEEzwDkI/gZUQ2g2NuxPi6DbxUsubuGLazGYjd/hyPmEN4r7IL65ojNPJvBMu8y92sldn/YXFZT/6P8rF3xLynwh0P+5KDt5L6KAEloAQKkUDWy+RAnEDWUym8BNaEMWFU1l2r0g4GcZzPG/6QwqsOWiIlkOcEastvA+RgPti/yAEb+8kvg9yaGJN7l2mPQ3H8coS39BoYpVpBWQ4KyHY4d4b9akL3KdiVey3CgvmH8T73AjMdp/hZ+9B/LuZ8f2xOP9d7yw4xOHIP63TUI3Osc86N4NMuCtGE31gNIhzghdMti2fMZnDBwRQjncp/OOyidsT3qCnbm6f6UrYx9rPS8U7F2bHAcdbFH3jNXIqi6IEIl82BzVvH4tJQSkAJ9GgCWSi8HQAeGjyKMS/7lLAjr8IpYb+Qrrm2oYGvs718yyfZIATnonLGpj466XoJTHwNKpdQPKxwkqmT9yMbs508jDt38V7k138aCJEqfNl0LM76KHOLJ103aLPQpWI/3V/94nCW/xaW0U6Ku3DwcPpvrPIcREPDeT89lOVo7t8YfAtXGnDJLrl7U5GJJ+y8Rd8Azmzue/Qr7IS72Q3SuCvrs1Cv660gBEcuRtWMP+C30zo/SbZiiS5s6ANx94OR9ZbzSkz+YxaKSm/muTfL+aU6sPXGdY+DwXZLVQ0aEIi9s/Q4pzuhl3g+P/FM0jIW91xcUreOp066HrZtdoXPidLaPgls/G/Paa7/6cbTKT1O8oZHfAXXnMJo3qC4lOzaecUspwmQsyRNyLAtAK+vPj8n908aKMOOhuPkYvPhFa8pQdDNXBveNplE3TLnJOpCs/tCtqsPY1wX/HSiva8Z8xSvT9OclST/7bVrr2F7LSfxbpfT4lA/wB3qHUYcNJqgt387fIqXfMg6eDnLlvxLL+D5F+yAgNyIyrqNqZvB+uikKkMq//QK2lhv29XjAQmg2fwIJ/AgwpL96xwrmMmj+Zwuz9LVpy45B6Lp530RDuewDWSO1CoBJaAEupRA9hN3YFyPB5rsJ96NUxjMztgJfLg6W0XSZyByJM/5QIpaJaAEck2gpuwhIHgQDKopzzB5dsD5P9vW4BcY3A7HORI15f9BeCwnZrOdaF7FvyP6wH5WPK8ylTIz9pNwTuhUnrvp1E3RT5D+1KvB4CUHA0aonwMrBlHzPhN6lbKyNXgRtSOT+62s3f1cGn/uxf5bsPtlvCXH9jcjRcYi/HzmBs6CsSmIww5w3M1UcvflB2O+Yp/wHqb5aywJXNYD2zgWXa0SUAIZI5CNiOIzt4GDSkbdet9wYfAUcv12LTOQsOEtF0KcpxP73v/WhhMbj4tmt3z+2FvR34fjzCKcUPDVGoJzZ7ZOwPkqpvSsHfkB2U6m3rsUL7sajPkj+s85GufOKPVSape7SB9/faevv38nfcPTSrDJGnvBRP8BkbUSsQm+ggn4/dxDQm3pv8rXBrEfeRIEG7S4Ge5/wInJ/yGXP1PRknji88jGzOBhhOJhZQDzfDYuenFND4WucC4C3EswoGgCxr+8AcIZmkixX60I4VAI+1qCUJuCGcB8BPsTb5fs8BPaazZdqxyOnMl4Ay1B4zzmBNVY77fwWxSzsgk5P0HkZf+4ZStEcSLCs3v566Xyfd4BXBtHM09jgqzvQ9Frtcy0R6mSb+u/ffnrMFINmA/QvKCjrW9m94vjJUxjNUbaek/ibhsrYg9K0Bhs5mKPcil2AULAGQojPs8p0guuZKb9blu289/ozeeZk+i0NaXVfoaA47dorlUvO1vHrYPgF8/IBQFes2di07VHoLOmKN4bIt7n3aAXQu7gziaTCG/HpVxzH/dvgcECbpNZh45jAfkTqmdsgUyYxEIJ2cg3KtfhvcVXI7Vn5auDEZDJELR9Ce8lzPvFe4FT6lg7p2FiL7CNWewZifDeIs7ZiO2xvqeOeigBJaAECo1ADsrj5CANTUIJKAEloATygUCtfVO2/DrEzSl8EDiDg483sgP+KSCZXgHs8iHqM8b7NwBnoCh6PiaO4kOrcICILj3LDiPvLcg6MZLRrYoe3JkDrvEJPJf2k5Cpzt0QQK5Gdf2vkCvT65cvIeYFsl1xgQLzKjfkKhuaTgcJGLMnQ5ZQMmPDHCieUvEEQpGzYZxTWW9vpHyTmciTxWK+pesNcAKnIlpyJmrLH2/XbxwzsFoloASUwIoEMn5cNX0o3Pi5MNI6qWqTiMC4H9qdLhTbt4j4p+/sD2fx0bCTy/6K3r5nPc7JUNN2ImVlXWM2QH9n1ZU9OuBiJ6qLonWAcwlD+y22XQ/GTEEv+7nbaYOo23FrJyHtl3X8YhDZAtn6rG6Y+Y8GfwuDq5iFTSjgvuFE5ktoMj6f+01oNv+zn/6XhQcBQgHPGazh84n5DlI0zx50ibgOJ7ax0D9t2R6h6J8QnjbAXy+HviJ9Ie6ZcBr/gujuh3f6K13haX3wRePvWIIaCvv8/L/MfgcjV6Ho52mw9X+Ze+q98dM2QtydSMVlcRrTCHH8FtBQPYt2KKf24bzPShzzScX2X49GdMluPO/JJ7F9Ai/1WhgphuPYLwn0anYTB4LVgSWdXPjUHFu7/o+TOJbIc5x85vkwX7UrbHuVi9EXkDV5nh0kNw4M1kDQWSW5d5ZdL32+GMbleTE+CwCwKvO/esYW2Ngi2XawePGucOUY1oO2iwu+hRuda1W6RKIRPk+ZL1OkPQTx+ESMn+E/sZ0iEiCwFc+9X/3vRf9tU0aTrsKUil8g8b9C5EEG4RgC/69siyE4FHCvwUV15YCRlVXa4bLWWv1Zv9r2yZIENlviqk4sELQLSWThWTDYdYXI300sblvBMWeHInaM8Gff9MTsCGNOQvi1/Lmn+mZYPZWAElACnSOQi9BeHa5cpK1pKAEloASUQK4JiBhcVvEFJpU/gJhUQoL7wbgXMhv2qwC/cNs5a8yLfCjiw4azD0KR8zgpNhXhsSkGzjqY5MivInDM/ZST81h+w4f4D9pVwqGrLIQxf/Ytk7hXIhL4sV3xtlc5LC5C7kt84D2DQU/0zY9j7Dn4I3UXZG2AmZlYzoYPsKvH6+lmJ2K5abXmPeZjZutRXmxjchUHic7wFrmoXfkswY8QmZDGObHnJT8FMpllbqRk1obHzkXtqCcRQyUnvfYHnAvYJr2UkUQMh1FgXuL/CxEP7Mch4ipMGvUELhs+LyPxayRKQAn0dAKZLb+dtDNyMuykqiCApUaKeD/aeOlhV+wIJ3fYCU2R9CpwnHMQC+2KxBtzKbSTeffuuyUgw+FnRNZGJM5BfT+ldvjZtwp/KHqE/b/z2J9b5BNyNfqdh2joFlRN377DE0nRJdtDsDnj8rHuaETnHcD7ofgotc/rnGkDmO9TEAk+CsglADZmPprjFyyCg3r06Z/epFVTyc6AU804BlNarDiADIWJ9EdXmWJ3EO/5bSfikuWEE8CcsIuFLszYFx2SpdJuN+kNyJ6AuQGDG59Edd0BnKwuQXtMOOzgovotEA3dxHjCDLompY01fGZxJiDa959o7pe38UuxWzl9MALBKmqNoCyzIqWIxzLzpu2yWNPf+3ZBgPV4NRgJ+AQSiHCiU6oQ772Jj56/V7DXELhg29PyxjNjhcFGCDmbkrf4B86C79WjlqBkoH0jOsy2K5qFFFgsIzCBYYyb55j73EliJcE3ju063PYniTRtp8aS1ZnR7Xkuij3DGNi2YQdgTJGnTns9YiXDYGJVEGywQtA1EAsNWMEtd4elffuSxappJFgOx61G1axlC3rSCLRU5aovS5nO0TwOUZJbYV/GmKNwYX1fZMrUjPkKceePrPMve0YpCLHfNBZB+SeqG87rVPqlZiRENvRMy3qIbIMfYjvZ3XZL2ARRNP8I5vcUCHmhjTHYImdjNW2SXbprMJDnuM/S46Q7Yv3Zt1j0K96zgklV1FEJKAElUDgEclISPlTlJB1NRAkoASWgBPKLgEn8Juekke+htuLPqCnfA6EBfNiN7shO+YWUh/gQ9CazPIcPwIu4tZOtSxJbY+zx19y3v4vJQQJzMeJmH7ihQYxrBGoq/oLJo97nxH96g36MqEPW/g7exIpXMLHib3ktU0a9BrvwIt1Cnj2sCfZtYr9yTap4InH+0o2zo3rhsTHUlM+m3JkWY5uv8JYp3ujraGaShIuVTgeci+FysKFVjHMh5s/7Lol21zldVvYQJ4pv9pSasn+0K3N2Uc2k8vvTOid+9agr/WrK72QbkfkFAK0g7W9WThnzGmpGXYmaip0hcTuY9ht6/4vyGcW2aTZ9K/a3hW29tWL36WYosDqfU9eG+Q2c+IaJuGrLr4C9rm0a9FSrBJSAEsgMgQzEYj9dHJ4WTAxONwXPhCPjGWs/ShtrHPbx/sBB7IMSE5bNk+vSRiE7uzZv9jPiVQ3bwrh2cZaTIiGbp41h5G/YcI1f4bTZvTB1aiBFGOst1A1h/PM7wsGDEKSauOgFuJWonjEKp80OIRPmluFR/FT8b0Z1G4yJc+tlOUmLQwCnHtHdbuI5WTdxTmw+LC+vUNbvrMeKcdELo2FwD5AYMIe3sf7mblQ1zER1/aWY0HAi909on9Sdx7Cc/Ki3k8nfo3foR04w3AqRnZiurWPCbbM1+AxxzIJfn9DWO/sTCFXT9+I9+iEI1gP4H60mMTnICSU5hUxWga3Xttyt3tncTjUBVL+wNsvwO+bInqNUqRXxPJyLUvfPuHDaWs31KJH/VOEy5W8YUYR1bVFCgEZu48w7z4kMot9YGHkIkdD/eM6nJCb1z5rVD/brEfYLG9d9VAwrdj/h9toAVD4/EpHd7kPAvMLwxwDSH8tMHAafwJV9UTvqb7hyG/tsuMzXa8+eP1u3q2esCcf5EyBHM48rXnO2XZiEyhljcMK0EoTZXiEHxi52sF/JaAyWw8SPYb7Izjddm88RiEceQ+X0MtbR0rQnq23dD7/dBwFUQUwF0Kau2OvANb/F+Nnrw9Z53yxkwdNes7Xld7DtPBcG9tk/U4lIojwXz9iA4wY891g3RcSsI+6fsNEao2DbOlt3UgTotLeta5XPrIpA/HcwGAvhn1ekwglhmN8gEtitU/mzadpr78L6TVjv/gURuyDG1q1lKRtsjmC8EpUvrgp7/eSChU3d1tOLnu6PaPQC5mst6+QrzUyOA6J/wQSWx15P6eTVXnvj6wbip89tvTjIN41mz7EImksQtm1YRtoHgykjP0NATmL031M8bCItW28vZ/qz2Gc4KvEzRbZ+WlaAwNPwGrfnrpr9DJhrqFZM8bP9IM4tqK4rb56wZ3g/bTBt217YhZ+RGUfycCKAoZTlrWAvROee38yOfUWkjHf58B06Yho2b3bRlwjrEgamEc1qbBdvQDR0YYKxDZ+TvKaRM1VRAkpACWSUQG4iW75jkZs0NRUloASUgBLIRwL2gb9m7GxOPl9BORi1FdugpnwImhYPQTy6BQcCdoQp2QgLT613jwAAEABJREFUVl+V7mtRdqT8CjUVk2A/vT1555/ysViapwImYN+8ri37OyaPrl0qtWWPdumn7QoYd7cu2qQxn6Km/CbKsZT1OdmwJVyzC7f7cYDhcJbteI4r/DqxL0K3wBj6bUnd9Sg2zE2wcVBRrRJQAkogKwQ6E6kdWL541g6objicA6aTOTXxHkRqGWXyyUuxk3nmP+jlPoVN1jyH4Q7EhLrdcfGM7RITgQyYEWsnDyqfr0BV/R4cLD8CXzTdAHGnAbI20jZmCPug92DQkpfxxtALOCB+ACrrK2DjtoPu42euR7dyTtiOTkwaVtYfQN0r4ASeJIPUkxY2HyJbA+5/MKhpMqoa9uOE3phEfNWcYB9ft4FVabfYRQBugOV1HmdY/7dpBaWAnMpJ209RGn+V+bgCExqOpuyJi+t2QGXdxkvzU1U/FhNmHol+fe9AMPBw2mUEiiDYGcAfmc7fIOaO9olcybDHUPYEZDUAAUoyG6Pjq2gMfMTtytZODFnGG695HifMnwUclgGrrqyYcClhnsejV3w2YsV/QvWswxNM7KRIJt8AtRM4VdO2RNX03ViXDsDr9eOZL3veDmEuHEo61ub1DISCL2G1JVewvh+K8S/sigkvrJ/2xHA6qSTTMcYuzD4NRbFBWBIYzOvFnqPLqfoEDN7ndgk5CmUNnvOLEMTb6Bv9EdHGDxEJPYPvv/9vQqKhJ+n2LqKLvocTnAGRgyltJ6rigPmU/acbURQcg8nldYw7tbXn6qIXRvHcHY1BjdcB7qsMdDqlbdw8bLEiQ+HEH8WaRfciOuOMxDU5oWFXVE7bFLata1Hr9MZO9lZOH4OqGXshstuxiM39B9k9gub6jbSMYB048gKv26ex8dDfoap+H1TOLMP5T/ZOhE+k0TAycf22pjVsjbMR+eUVQE4G4FDaWgeCgxBomo5o8cWJ+C6eMSYR3rZHE15cv61ydvbFYN7qt0Hkz4x/PqV91nDCz7abtp22eR5ftwtsuxwJTmLbMxswo9KLUDYjnafQr989qJ5xBFmMXcqhcsamsJPn6UW0staU2f2XxlXdMJpx74HIzFMgJY/C4DyWPXndXC4mGQzIg+jb7+5E/mwdteW9uL6C7fY2iYUB8DDhhtUwvm5vxGb+Bt/PeYgTym9QczvKylYggBwFiX6EQU3XYkL9sZgwY2/es3fBhfVrIFPG3k9t+zyB5yvRRq9xFgKlbAflt2knIRKAyKFw8TIic29H4j42c3e2g5sl2kG76Ku6fvhS9lUzdkN8zxMQkAeZxjmUICW1FTkb0djDiXOWaLd5r06wb9gRdhI8dQwra0wsf4eOF/H8f8+tnxWWcXPAvRuBxi95/u/jtf/7RD/K3p/tz+Ik+iUteark9Tth5kEY3HgVELeL3Tbyi3ypn2Bj7j/CCfu/wf7Moj0vtoxWLqlbh37CsgZhr4Xq6QeiKTie/b8G5uvvEHjVC4fh/kR29YgFJ7DeHpqo+9XTd4JdUEHPjFi72LOybmfY81vZcAjzdSkceQOQ3ZCuEbHXYA0CTW/zPlULG88E3lMvnJVevy7ddFRPCSgBJdCVBHKUtm38c5SUJqMElIASUALdksCVey3ClLGfwT4U1Q7/FtcPs2/IdsuiaKaVgBJQAgkCdjJ/csWLnNR/DjUVD6Om/P9QW/Zv2P1JZXQb9RL9Pk3o6j8loASUQA4IdCqJVef2gxu9koO+/2Y851PW5L5w62ftQHs5J2SuBMyDHLB/ArH4Vfj+Fzu56xcufT+n1+acSHyAAR7nxOM9TOc0QNr/KWOxpREOuMtk2DeZBffBxr3FBsWcKDyOcT4A4/wXjvsINR+EyO8pA9E+M5R55MSPeRiO8whsfFaCzknti6aNdsnITzhZaidiOfAN08Yn+W5zOTdlPn5PhX8yLCdw5VoOnJ+1ND+CR2AnHuyEENCPevllDZYA5gX0fyr5l8CMOZv+D1Mug2AkJZSiAALI+jBuJST+fzDmcQD3IugO4zYzdtV+q0KK/ghxHgPkvxCZxO0WANo7XiaADIURnj/zHwQCD8N1fo11+5cge8Yw33WILXkQ4bGNsJ9xn8iJ+dryKqzqHApxf82kqyg3AfIqgEYKWEbL3bYTZQD2bJHR3K5LKaK0tfaLSPV0uAKC09AYGI/wiK94nJ4tNhsj4NxM5X8y/Bnc2jaGrLjnaYUT6GZ/1pMbIYbXo/soJHQebFvnGaadHlKyExy7AMV9DCJ3wmAct8XtjIXqEmA4cpRrecB8xqci1N9yBCeUmYb5F2xb0pzWIxBcBRE7uQcfsybgsk7yeo/HOSnt/DcRB6In+oTJnJd93g+ZvzLCuygLKenbS98JcUL3ZF49Le0yy+DgAZb5IkbS3vbfXjuHsA7cDcGjQAsHx1yAnz62dZhRdsAuWLJVgqc9L4b1S8A64N7MNHamSNoxigSpfyjzdw/bJuYv+F/EeQ9x5XL0GzDIM56I2YWMHoZxbwDMXhBJXe8EA5nObwC5i+EegyuPIIh9kSmzxeBSuDiXwjpKJsDVsG00mEp70xD0ZdhjGeyfMLHHyfoCrLtJCYrjbB/NTTxuqc/u43Dd26ln254gt+lae+5HQ3jOIMuuD9fcgnjx+ulGspJeKMh6ir/wnMxdyS+ZQ3M59wPsYhnDvod5DNHgyES/JFG3WF9tv8S49wM4ExDvOoFkRvrT9RhApsK0KWfMORJhY+tpH/Z7eL9x7ocj9q3/rSESgL+x/lsxvksA8x8I2M9g+xyNsY/lHzBt38+cAYz3bzw/j8HBfQxXBXtvRAeMmLUY1wWJeIzzKILRwzoQiwZRAkpACeQlgVxlyslVQpqOElACSkAJKAEloASUgBJQAkpACSgBJbASgc45rDdgPgdyL4LI3pS9OiSOsy+H+ccjNPiHzmWmTWjjvg8JHAlH9u1QnkRWLovDMsI9ihMg7+OdT5rgOvdAzFFwHDt5dxjT2ruTadnwjIfx2Thh/tGmRO3bDYuLn0pmMa8nw2DlsiQrX1s3W9Z47ALEcFOifDY/IpwMS8Klbbiu3AcOgil6COGwmxSWI3/p5DnaB+Icg6Kij5PG3yHH6M/M02TWmwMo7T9PnrzNQZyc+Rc+n9c86d6hvKUKJAYGMzFld7YBK+j+YdQS1Ix+CTXl1yCKCxEzx7N8BwHCCSNwchbvw5h5sDGgxRjEufcz5QO6c/LOnMntwZDgSRgyJIxJFc8kFhlQIW0bafyY6ZzBtDvO1siBiMevw3q2rUs7ZX9F485mnn5FYZ3K0DVlr1knfhyK+3yRSNymYcDJcLYnIjat/Zheezjsw2v/UMq4hMTNPxPx5uJfuPwb1pkreO6eYHKGkqb9TwwCTlLjqESeHedQiHSW8d6M42Be+80cTPxarLpRNM0Mray2OPBuS94YX4fOS7JzeBAEjI/nGu7FCDX9tHLCrS5uHewXxzzbjpT1cW84Qq7xJ1tj7PR2812WAHID5RDmLVn5Oujm7Mc6dA122roRJaXfIWB+D3svS4js3+m0HOfApfHBnIUm9zN01IRHzEco+he4bG86dG6cgxByXoTtlyTKx7ogchjLaOtvB/ktrQuHLC1n3P0PwmLwzYeL4OAGxr8PpaPx74NA4EyEer+PTJnionkQ9zfM0wGUjuYrSTjyLQo+lKlsajxKQAkogS4mkLPkdQFAzlBrQkpACSgBJaAElIASUAJKQAkoASWgBFYk0Mnj04dHExN9k8qeQmdkYsWLCK/f2MncLAs+peIXTBr5TKfylKw8taOfhY373nFxTBn1cWJSMpleRtzKP1hWoA7s3cJzUzv6TdSWP90hDpN3mYnLyt/pUNiMlL+ddaq2fBr8fhbMflGs8/l6DnaipgOnI2mQ8NhGTBz1WuYZc7J80piPYOtp0oQ76Tjn7X8iFOmPH9++iZMs/hO0l5cvWFqPPvz6RsyffxSWONsjsnhNFEX7IRTtm5Ci6AA0LVoHqzrboWjAySgq/ysSdXfERzi7g1+Bu2yPeZg8uqFTfG0eLhvzFmxb10lsS4NPHv1Dp/LkWY/HsH5uuTCRjk3DXhOeuu28viZXfJiIN1f/Lqv4AgsWHMu60Q/fl/4jrWTDYReT2G7ac5apcq8Yj21Tw+KmlZ9kSleP+jk7577lfNqFN7ZdSZa2dasd/W1G0q8Z86WNLiMyTuKoLXu7+XpvKceK3Dt6bM+XjT88fDEmdrIt8MtDTUU9Lmdb1xkg4bE/Jn7exC8dP79w2ZzmfknZUxk5x8nSmlLxCYtocMvp0cRXOpPptMdt4sgGhLebyzgzY8NbRlAzZnpWyh8e+VlmMqmxKAEloAS6mkDu0tcFALljrSkpASWgBJSAElACSkAJKAEloASUgBJYnoAeKQEloATaS8BO/oTHLkxMArUnrF2QcP2+TYk3+e1Pvdk42op1s18PsJM4nZlkbU+eVDd/Cdi6YuuHXcyUv7nUnCkBJaAElIASUAJKoPsQyGFOdQFADmFrUkpACSgBJaAElIASUAJKQAkoASWgBNoS0H0loASUgBJQAkpACSgBJaAElIASUAJKoPAJ5LKEugAgl7Q1LSWgBJSAElACSkAJKAEloASUgBJQAssI6J4SUAJKQAkoASWgBJSAElACSkAJKAElUPgEclpCXQCQU9yamBJQAkpACSgBJaAElIASUAJKQAkogVYCulUCSkAJKAEloASUgBJQAkpACSgBJaAECp9AbkvoQCS3KXY+NReQSFpiTByFbVi+NFkA1C1UGGJYshjSrRcQ1iGoUQJKQAkoASWgBJSAElACSkAJdC0BTV0JKAEloASUgBJQAkpACSgBJaAElIASKHwCOS6hA4NAjtPsZHLmLcC9PD2RJ2BMUycTzNfgnPB2X0iPA3kZvJSvBel8vswinucH02fhftj5NDUGJaAElIASUAJKQAkoASWgBJRA5whoaCWgBJSAElACSkAJKAEloASUgBJQAkqg8AnkuoQO0N3ekndeQU3FxWmJgwcItJFSeLb56wZPpsXB8hKpKzwIS0u0AIK702YB5+2lIXVHCSgBJaAElIASUAJKQAkoASXQNQQ0VSWgBJSAElACSkAJKAEloASUgBJQAkqg8AnkvIROzlPUBJWAElACSkAJKAEloASUgBJQAkpACfR4AgpACSgBJaAElIASUAJKQAkoASWgBJSAEih8ArkvoS4AyD1zTVEJKAEloASUgBJQAkpACSgBJaAEejoBLb8SUAJKQAkoASWgBJSAElACSkAJKAElUPgEuqCEugCgC6BrkkpACSgBJaAElIASUAJKQAkoASXQswlo6ZWAElACSkAJKAEloASUgBJQAkpACSiBwifQFSXUBQBdQV3TVAJKQAkoASWgBJSAElACSkAJKIGeTEDLrgSUgBJQAkpACSgBJaAElIASUAJKQAkUPoEuKaEuAOgS7JqoElACSkAJKAEloASUgBJQAkpACfRcAlpyJaAElIASUAJKQKstPFwAABAASURBVAkoASWgBJSAElACSqDwCXRNCXUBQNdw11SVgBJQAkpACSgBJaAElIASUAJKoKcS0HIrASWgBJSAElACSkAJKAEloASUgBJQAoVPoItKqAsAugi8JqsElIASUAJKQAkoASWgBJSAElACPZOAlloJKAEloASUgBJQAkpACSgBJaAElIASKHwCXVVCXQDQVeQ1XSWgBJSAElACSkAJKAEloASUgBLoiQS0zEpACSgBJaAElIASUAJKQAkoASWgBJRA4RPoshLqAoAuQ68JKwEloASUgBJQAkpACSgBJaAElEDPI6AlVgJKQAkoASWgBJSAElACSkAJKAEloAQKn0DXldCBSKDrkm9J2eCfMDg1LXHM31tCpbOpp9LZacUr5nzqdr0VuTSt/Dr4DfWeSDvDDv5D/fQYC25MO95sKqZdJ3A+4rHX0s9K/Ka0WQCPpR+vaioBJaAElIASUAJKQAkoASWgBFIQUG8loASUgBJQAkpACSgBJaAElIASUAJKoPAJdGEJHRgT78L0m5MWNKC27HbUlt+WUiaW1zUHSuP/pPIPUFtxF9KJ1zV3M0aX0rXWDT2YVn4nVdyB2tFvpp3ZiWUvpxWvZWXk6bTjzaZi+nXibkwZ+1naWakd/WzaLCCvpx2vKioBJaAElIASUAJKQAkoASWgBFIQUG8loASUgBJQAkpACSgBJaAElIASUAJKoPAJdGUJna5MXNNWAkpACSgBJaAElIASUAJKQAkoASXQgwhoUZWAElACSkAJKAEloASUgBJQAkpACSiBwifQpSXUBQBdil8TVwJKQAkoASWgBJSAElACSkAJKIGeQ0BLqgSUgBJQAkpACSgBJaAElIASUAJKQAkUPoGuLaEuAOha/pq6ElACSkAJKAEloASUgBJQAkpACfQUAlpOJaAElIASUAJKQAkoASWgBJSAElACSqDwCXRxCXUBQBefAE1eCSgBJaAElIASUAJKQAkoASWgBHoGAS2lElACSkAJKAEloASUgBJQAkpACSgBJVD4BLq6hLoAoKvPgKavBJSAElACSkAJKAEloASUgBJQAj2BgJZRCSgBJaAElIASUAJKQAkoASWgBJSAEih8Al1eQl0A0OWnQDOgBP6fvfMAkKK6//j3N1vu6KBgjb3E2AsKXgGxa2I0MaLJXxMTFWNNbMDdgYzCHdhiTYwl0USNCRpLoiZGDcIVwGDvYi/YUOncbZn3/749UMru7O7d7u3u3e/dezczr/7eZ16b997MKgEloASUgBJQAkpACSgBJaAElIASUAJKoPsT0BwqASWgBJSAElACSkAJKAEloATyT0A3AOSfsaagBPwJqKsSUAJKQAkoASWgBJSAElACSkAJKAEl0P0JaA6VgBJQAkpACSgBJaAElIASUAJdQEA3AHQBZE1CCfgRUDcloASUgBJQAkpACSgBJaAElIASUAJKoPsT0BwqASWgBJSAElACSkAJKAEloAS6goBuAOgKypqGEkhNQF2UQGkTMEZwXksvuDMGombWEExoPhB1zWNR23QDamfdweMDqGt6gqYRdY3/4PF61DaOQ23zCahrqaDZHOMbB8Gd1xvTTaC0Yaj0SqDECLgzgrjw+T6JOjhh5jasjz/AhKZLUdd8E81fUNv0EM1M1t1GHv+G2lmX05xDt6MxfuZeiTrvzukPd0Y5bFtQYtlfT1zLw+andt6mmNB4MPM9gXn9J/P+LOqa3kZt4xc8Gl5HaD7j+Zs0z9D+fl5fRP9VCSZjm/qRSRCqlIASUAJKQAkoASWwNgG9UgJKQAkoASWgBJQAkJh7aJrHOQXjY77gXMMhiksJKAEl0FECxbIBYA9MmH00J02PSWsmtuyVcWbrmrdCbfN308ZZM+sYOHI4DCTjuPPmMTIqI3nrGo/ChJk7ZCxGzaydM4rXshAzLON48+mxJsMyMaHxcNTO2jRjUepm7Zc5C+yYcbwd8qiBlECJErAD1YmNe6C28Vj08q5GJPg429F3uQj4BGAuY2t6FsQ5kcejmcMDaaoAOQrA2RCZBjF3A14zzYcISBOirb/Fc02/ZF9wMMY17ZJYkKTnktau6zAf2yYWBesaq74+1syqTGx4KLbM2cXPmsYdv5ZzTZlXn9c271mUsqdiaRemx83czTdPq/NWqGNt8664aV4oVRbyYm8360xsHIZocAzKltzGOjgbJvAa4N3HsdBEwIyh+TEE36UZAUgVj6NZpy+iuY5uDyAQeAaO8w6isYcRCTVgQtMJ7FtZthu35eJ3OUpFHTc9ADuGsPc/Gj6X+XkA0voajDwGyGTm9XvM+54AtoHIBjwCghDNEADb0ewFkWMguByQRjjOmwiZ6YiVnY4aMq6ZuyFUKQEloASUgBJQAkogQUD/KQEloASUgBJQAkqABCKxfTjfsDPPkmuDdwEZg4bKx6FKCSgBJdBBAsWyAeCXMN79nDRNbzxzdsZ5FXMw4N2VUbxG/sjJW8k47nx5FPwmM3lxDzznhxmLIfLTjOJ1nPthUJNxvPn06GRYJhL3ToZnLIoJXJQ5C/OjjOPtiEcNowRKjYBdOKxtPhHR+BXwnNtYl+5i23k6RDhwld4dzI4d8J7MOG6Aif8VAfwBDq5BXdMvUTNnZ9iFaZSiOmQw8zKOkv9lLSPOXxBdcQTtikwP7AtHzqBQa8sLfHMtph6x1i3opzR0NLwxgoE6CvtNHtbMT1Gcm1/g41gvyph/PWH2d1DbdB4X8W9AXO5kgtcCchyAb9OEabLVfRjAbg44D5C/JNqDqHMzIsFpmNg0GmNnfAvFrOpmboEdNj8TcNj24G7AXEVxR9H0p+moZlg5nOPaa8jjDjiR61E761S4TZt1NEINpwSUgBJQAkpACXQTApoNJaAElIASUAJKQAnYeU5Hvst50BRzQeYdGHMpQgP+CYiBKiWgBJRABwkUywaADoqvwZRAaRNQ6ZVAyRCwn7SuazyTg9P74ZiruVB2Ks1elL8ji4YMllJvCMEwiJxEH5fBid7LxcQ7UDPrR3Bn9KVd6eiobE9hvw8IF8zXMIItgcBYFJtq+yIAwQZYV961rs0m8CSEUlEC+xn2jbBWHta4F4W33xQib+Ll95cjf0rgzt4aNY1TYeJ3QzCZZjQNy6dYPrlMeSu2CwdB5Bx4uB7B0AOobbSbeYZjTBd/5cAvV+68wahtHgc4D/BZegoEP4TIt/yCdMCNbM0OgPyYcV+BqNxNFmNg21KoUgJKQAkoASWgBHoiAc2zElACSkAJKAEloAQQDW/OBX77tdRkMD4D5HJE+0+Hu2sEqpSAElACnSDgdCKsBlUCSqBzBDS0Eih+AnZXam3zCQjhDUBu4EJWJQwGgytnNPnUwsj7A/IdpnkCxJmOSOhl1Lb8FOe1pNghi+JRdrHTxI8EZBMkVWY/TJxdmdRJLXsQAfMOHPMi7hkdz1um65ovRST+EhyHC97Yg+n0ocm3tuPLjSDYh3X3XBi0YEjrv+E275rvhNPGP6HpSERan+bC/1RA9gbAdob/86plIGBGQOT3CJo5mNBsfzogrylq5EpACSgBJaAElEDREVCBlIASUAJKQAkoASUAeNiXGLamWVd7gLkeoeituHKP5es66rUSUAJKIFsCdoI22zDqXwkogZwQ0EiUQBETOG56ABObduGiOxf9zR8p6SY0QlMYLbB/W0K829E7Pgt1jT/G+Blb46Yieqt4TTKbLR1AiU9Y02q9cy9+HqabwHr2atFzCBh5FQF5PecZPnNGXy4yH4japjmAmcCyyEV/IzlPJ6MIma5QAuBARM0LmNB4D+U6BPbnREC3jOLIgSf3qU2Ybg0M7qQ0WzJGoelqLRDZGca8hAlNp8OdYTdTSVcLoekpASWgBJSAElAChSCgaSoBJaAElIASUAI9noA7rzccY18I6rM2C1nMuYLJCMWmwR0VW9tNr5SAElACHSOgGwA6xk1DKYHOE9AYlECxEhjzz97YYZOT4CUWyk6nmMX0xj0Xy2QojNyFQPAOfNB6Nmoad4TrFld/Fis7GpB0b/kegBdn7w5VPZVAKwRPw638LKcA7OftBwYvhOf9ifEPy2ncnY9MWHd/RLn+CpHrUNv8Y4xvHNT5aNPEMLFxH0QjdjPTpfSZ//SYSBpdDoMrEQ02YELztvTLdo3/VSsBJaAElIASUALdl4DmTAkoASWgBJSAElACbSvtzw/uSRAhmlXaLALMtdhkk6m6+L8KiR6UgBLICYHiWjDJSZY0EiVQGgRUSiVQlATGzRuADQeeD3GuoHx2QMpDEWqBAFLFRbQpCDg/BfYpR7Goc+aXwYuPTSuOQX/EvR9zkM+8pPWtHrodAT7ged7MnGbLbdkAkdZrIXIBjX2ozGn0OYxsA8pnv5BxHRznsBzGu35UNbMqYeRaOhwLSBDFo/oCchKMuQLuUxtDlRJQAkpACSgBJdCtCWjmlIASUAJKQAkogR5OwDUO50C+w/mQ7UmifS7QYCnEuQmedwPO3aGN9qqVgBJQAjkjoBsAcoZSI1ICWRFQz0qg+AiMe2wAAq2T4EgdhbOfpuahg9rgKy5szYHBg1zgvo3nXIAzdwAyi+cfIneqN+Pj4lnf4unP+n58FER2SJtFgd3tW4UJc+3AP6139dDdCDgfY9mmc3KWq/Gzt0Yk/h/GdzwNF5f5v6PaiK2j/wXkHtbh3wO4BsDfWdeeZn3+knaG153Xgg0h2KjzEaWIYWLLXhC5Ckb2T+EjE+vPYHA/PTYw/78G5KftxlxAFtMAPMLjlzx2RNsvAXwf0cg9cJ8d2JEINIwSUAJKQAkoASVQEgRUSCWgBJSAElACSqDHE3i6nPMHe3NuYfU8SJRI/oxYv8swdcTnPFetBJSAEsgpASensWlkSkAJZEhAvSmBIiMwtqkfguVXQHAeJeOAlP8z0mIXAlvp9VMOYh+HeGfCcXZHeOAmaKjeHw1Vx6C++hdoqP41jz9FfeVInm+BJUvKES+zn8i3bwH/jYNfLjiaxYwnTlO6+qZ5Ieb/dGagfScvT3y14DuIR6uY/8z8+0amjh0k0AYIyy8+BrrQGHMPrs/B7m67g3x84z5wYvdAnL0hCCBjJR69LoXB+zzexTo8GiFsjobKLVBfdRDr62g0VJ3B8/NofoSG6qGor94Q4Y0Hspzbn684ByIzGe4TmuWMx7YHjKoI9IQ5O8B4N1K+YZQt0/Eu5ZfFMOYVGKmHxHdGfdXGaKj6IY91aKi+lkzuaDfVv0F9dQ3qq76L0KBNAY/p4Abm/D2apTSMi//TaUncryrEVvwD7pxvQZUSUAJKQAkoASXQDQlolpSAElACSkAJKIEeTyAaGcT5hkpyCNPEYPAIwuUupu3+Fa9VKwEloARyTiDTCdGcJ6wRKoEeTUAzrwSKiYD97H8QVwByGrJTEcA0w8gUxONHIBQ7AlNG3IjJFS/C3TXiG9X1R7Zh2r5vc/HsbzQnIFy2LyCnw+BWAE3Otl8fAAAQAElEQVQ0S2hKT3/UWkWhd6HJVA9EQA7EhU9umGkA9ZdzAv8D5IeQwEFdauIrf4dcqFjzDghgGhe6h7I+SoZR2sXpz2C8f8Bu+ombatbDE1FffQ/cqgVp43B3WJKo5/VVN+CzskNgHBpTw7j+wrAv0thNQTwUSJ/XsgFM9Fq2J3ZRPhMhLI8P6PE+GHMmwrGRaKicgCkjX6Vdem3bu/oRT6G+6hy0yX7wcBED/YuGjCkFT9JqY4YjGr8ENXO1LUgLSz0oASWgBJSAEigxAiquElACSkAJKAEloATs10JFdgUM50zMPTChU+AOXahglIASUAL5IqAbAPJFVuNVAj4E1EkJFA0Bd0Y5Aq3nQrJc/DfGfgr8InjOaVwoq8e0kc/CHRXrcL7c/T7h4tnfML/yLIScnzOecwG5G8AXNKWh3ZfCiMl3AcluAc+Yg9E7uB1UFYrAMoTCb2DK/q92qbnsEPvFi87l2dZfI3WAHIBslDEPcln6HC52j8GUqj/ismr7BYBsYvjG781Do2wDXkLDiOuxZMkpgHcqQJmMzEQmyniSibeM/YyZF0JvY5kcjsxUhCzuB8xZiK08BQ1Vf4E7quMP4FdWfkYeN7NMkYU5myJk+jMPIcpwNJzocTjnkTKGU60ElIASUAJKQAl0EwKaDSWgBJSAElACSkAJwHgHkcIAzkE8ADg1mDqsdOY8KbhqJaAESo+AA0gWn4pF4ZUxg1HXNDQjY8zWEJRW/pChEhFANs+Ig+UF2RjdV4VgzPaZs/AGFRiFJq8EioSAEUTCP2I7eRYXntgfZCSW/Vz6HYBzMvao+i2mVryWUahMPd0jcbgVbyL0+B1ojZwLxL/PgfFvGXw5TSotwJBUbl1nH1m8IzlW0dhPeWWRLtvnuBydRQD1qgTaCUSDl7G8jeZFkCYTbb+scQ4Q+iXCj9+76jfmTCYBM/Jjv+xh34QPfXQDYoETYbzvMdwjNK00KXSOh2kbtZ0IeKcwMbYL/O+vW9n+3YJw8FeYUvUQcrEpw6YnYmA3NYWe+DtEfkxzp7XOwGwImLHoN+A7GfhVL0pACSgBJaAElEBpEFAplYASUAJKQAkogZ5OwJ3Tn3Mkh3KO8wmYwCWYUtHxFzF6OkvNvxJQAhkTcDjRWFq/tyz4PnP3v8yMTACkL7qnCkPASXxkxkJwMrqvGgxxLmf2MmQhdrcdvRdKa7pKoEgI1Mz+NsQ7g9JktkHIGC7+x12EBpyKhsqXMJqL9QycF+26Hq4atRD1I1vQUHU2+potOFC+gml9CgqNYlPGcLFR9mafumcHRBO256fggBmZLuJ2IAkN0u0I1DYdAci5ADJ5W9ywbL4KL3Ak6qt+i4bhn8JlHWPgvGh3dASXD/8QDSMeRqjyKKZ7ACAP0URoDPKlxs/emu2E/QLBgPRJiEcmf8brC34Fl7LaRfv0gbLz4ZJxfeV7CEZOA+Q6AEyT/321bENvNYBtU3w9qqMSUAJKQAkoASVQEgRUSCWgBJSAElACSqDHE3CHL0FD9TA0VH0v8TJVPuYgejxkBaAElMC6BJx1LfRaCSiBPBPQ6JVAMRAYc1MITuz7gAxHZuoDLlKfxAX5aXB35SJeZoFy5qum+isuJo6FFzyQC3y3w2BBzuLORUSXzB4E8Y6BiN9i7FKfpDbC8OAYH3d1UgLfEKiZNQSC2m8s/M640G1MC5eef46p+zfTp6HpOu0y/an7z0V95VEw5nuA9wQTX0STW+3OCCIQZ/yyWwYRr6CfW/F5r7Nxz+g4z/Or3VGtQOvlbLf+zIR4zv9+WuSHcGfv6+dF3ZSAElACSkAJKIESIaBiKgEloASUgBJQAkpACSgBJaAECkBANwAUALom2bMJaO6VQFEQ2GinrQH5JYBM+oEPuHA1HvXV99B/YfXU4a8gHDsLYs7mYuIDFGY5TeF1m/k2hRhFk0p/CSN/TOWYsA/g1xjb1A+qlIAfgekmAMc5HjCZLHTbmF6GiV+EqdVz7UVBTUPVYzChEwGZBMHTADya3OhoeHNAjgHgX4eMicPgTwiVnYebh0bpv2t0/UEfwXPqmdiTNOl0EFFvGs5r6ZXOo7orASWgBJSAElACxU1ApVMCSkAJKAEloASUgBJQAkpACRSCgFOIRDVNJdCDCWjWlUBxEDDBGsB+ahrp1BJ6uA4rnft5LA5t36atr74fMTkLxpwPOK/hs88LK5tjfkoBBtIk1wZ3wYv9ifKm/nKBke0R9L6bPAK1VQKrCLzQvD1EfsRy33+Vjd/hY8CMwdQDZvt56lK3huGfIjTg9xD8DAZXAV5uNvEYM4xxVabNi0gzHFwLd6j9CkBa7zn1MK3iTcY3kSZ9no2pQB/vEPpVrQSUgBJQAkpACZQuAZVcCSgBJaAElIASUAJKQAkoASVQEAK6AaAg2DXRnktAc64EioDA+NlbcyHaLlinEyYOmAcRxU24umJlOs9d7n551QLM//gPCH14HX43Kv2CWr4EdGcMhPF+njJ6Q4Lh4OUoD70Hceznz5N7FfBPfoYx80LJPahtjyfgug6M2Q8eF4dhxJ+HaWN5G4P6qsK/+b+uoO6uEUyuehkLy+uwVfmf1nXuwLVlcQIg5fBXnxPbLZhc+Ya/tzy6Tql8Gga/RlolYd7n/4O952n9qgcloASUgBJQAgUi4M7pjwmNB6OucTJqm/6DuuZXeP4x6poWf2MaP6Md7ZseQG3jONTNGI5z5vv9bFaBMpOPZDVOJaAEckrgJj4r1zQOY5tyGducJh7f/bqtqW1cxPO32c48yWMtapt3hf2ZsHQC/KalF8PcRzMXdq4knX91VwJKQAkoASWgBJRAiRDQDQAlcqNUzG5CQLOhBApNwC4mBeIXcKk5kIEo70OCU3F51dIM/BbGyz2j43BHR5i4oSmMjgbJ0/GbxPwn8P5ncCu+AswTXMBdnlpQ+Q4GrxiR2l1dejSBRSP7w8iPWH/TbRKJ09+tCDqzyKtwdYOJ+2r7Cf7Tc/AZfvf5IaxbR/qmRQ90n4PIsvshUjgmNu2GqlthMI/ypNYC/smuaD10j9Se1EUJKAElsAaBuqbhXPB4CbVNkYSpa4rx2uTF1CbiXsi4F9C8z0WTJzGh+W7UNk9CXcuPUTO7EhObdqHddqiZNQTnPGLHSbKGtLk5dY3DNE9AbdNHNBHKYk2cx/zku66pc/HWNi5DXfPYjBakckMot7HYTap1LZujpqma9/xKMn8HkdgiGHkMkAnsuQ5hd/sdnm8CoP83RoYAHOMCR0NkGhCajX6ffsU4Hkdt84kYO/s7ZDIwZ5veJszcBrVNz9JEEiY3dcFDXVM0EV9tYysmNN2R0U93IYWqa/p1e1ysrzbeuiYbf7blK5qIo65xKWW7hAzLeW+OTNjZdqCOcdd2KN50cixnGp/w/i2geZYLsNNR1zwZ9l7ahdkJTd+mLJvAfSmcIveds65tOgQ2f9bkL4/pGGTm3s7f/uze2nm2P/tW2zQPVv665uzarEScjf/GhObt1o40D1c1s3amjM8leCdkbcpO1romw7AxmlV1sfmfyPonrowk+pEJzQeirvkveL91ERyZA8hYCCp53Aqr2xuRATzfBiIjeayHmBcQCS5AXdMtmDDzQIyd8a2k5fIr7ATIMIbrA8Qc+Cnb79Q1Hc88ab/jxymXbrWz9uY9fJHG5NY0s91tfBq2zcqVvHYzXF1n293mZlz4aB+483qzfb2WeV7O8hbhMUpDmVmv6orSvIkJsw9IibKueSvKPzuRF9t+1zbFeJ3je5pooz5lvAuYzqvsox5B7azrUNN8Bq9HwW4KqntiK1z4POs625aUwnbQwW44qm0+mel/AZvHuiZ7zzqfR8vKxmdNXeM/MWHuNkzjQxobf7Zlwoa5leVrcAdzmXmwCTMPpYzvrmIRQbb9XZ2to2Ro823NhMZrMk88S5+2b6prmoq6pqWr5I3y6PG68/evLh/1tfldTGg8PMtcqvceSsB/YNNDoWi2lUC+CGi8SqDgBFoP3RYG6RbKQD+GC9XXYsr+rxZc5mIW4KIZnOCUk3xEXMmJh3sx6bgo/RhI4CkeX6JJpTcF5BBOmqV7kxmqeiCBPiFOTpmDMsj5m4ibv8IdviQDv6XvJbLkfE7YlaXJSARO4E5ceZjPBpw0MeTSWXADo4vQ+OktEDRVsBONfr7UTQkoASVgCQTNAhi5DYKrEgbmt7y+l07vgKM6HnOnBXYj6YYAOG7BFrCLLcacwDGPC3h/gRNvgodnOJZ8AuLciv79LkJN008xcfYwjnH6Mlxu9CTmLBB8leneCJtvI7+hjT2fDsh8gFf8VzRa7B4IcbCgnz0pGrHSCmIXImpaRmLIivNg4n+Dg8chcgEEW9OszksMwKe856/wOAuQJxLGmGbehTcBfEXj0bRrQS+IHMR7dweC8UZEQ1cietCxqJ2zcbuHTvwPhhZD8CfYMgHzG8C7nrH9CZAWmnR9L9ZRbbz+L80fmLdr2+NkORN5GCuXpI2L4ZJrD08n4rIyGlwNg5tp/gXIF/BTxsTp/DLNXYC5LhEHnGsg0kw73gPnbdg4rYG5mna/h5EH6PdjnhuaXOjeTGNjiGxKsycgxzH+CYl76chseGYuIqE/I7roQk5Ofw81c9lW5HChxQm+C5s/axJ5NDfC4C+8P7bskQGKRwklNca2l2vLZMuOyJ/Jjfee8gvlB54DsJLGX9s4IQczz5PgNm/k77mTrhL5gun8ibm4Cu3t628BuRsGr9GkLk8GUQBzaG6nuS4R3t4v8e7D8nDm98guyNQ0Hg3H4T32HiSvHzO+3jSrtLBNMct48SbleYNHaz7l0dZbHpiyyBCenArjPIpg8F5EFp/HRdWqxAYew3LpNm2GWOwX9GP98ZBGT2JKjnmN5d32Nau52HPtd5AnJUHeU3MXyd/KFP7OcjCf57aM8TJbbRYx7GyWjL8w5C2w5TnufYlcKcN+xrbnAOXEG4yWZZT/0+tWenmOstm27F703Zz5e5t1xdj+h208rqL7NWznbqGfh3m+kKbItHEAjyaFWHFvCV3uIvurVpnrmZ+/AvIqAOaV/zuvbfq2XWT/hJ0gcgTEOQeO+R0E/2W9nQeE/4GyJZeitvnHmPDEdjhueqDzya6O4QCPabzEqxuY3lUwci0Mbmc+bR9t2ypkodoY1o4V/gzhOCbRhrIciHMfx02LYYTtIq5mfDcBcg+A1wHJZFwSBPB/iLSen2gHeZE3HXfehcFNCRa2vzbs79rrxgdM09Ck0vbnGv/L8vRHsmsfe9n8Q2iXKkgn7W3fZDDr6/QMLNubGes/aT6jKTLtOYDQQJUSSEtAC0paROpBCeSMgEakBApPIGB+yMEXF63TiSKzEI79IZ2vHu8eDh1DBvYBg4ck2pgXIPIMTfvgNvjBW/RlJ0P4QMez9XUYjoxArHyb9Z3UpscTMDiLZalPGg528uDviASeTuOveziPf2EQmZycQWbmI9D6+8KX8QAAEABJREFUSAb+usaLSBNgnkuTWD9OolchMjezCck0kamzElAC3ZzAJdXvo6HyKtRX1awyv4bnjUFcTmU7yQlWeNkTMLZPsZNeLzIs2y20G2Oe4bVdULTuPE2q7du/W0HwfUAmw8GN8OJ/4MLgH1A76xzYN8nBxRd0Qtmvqkwe9jzqq6fA5ruhcjwZnINo+Rh45hTG/CdO5NkFU56m0mYREpOz+D2PnTG3MvxjNHbid0mK1BxO4vZG2WdOCvfisp7OSfG6pqEo6zONcjN/Do9SSSHtvWU3xtzaRUuRehg5jff3FxA5BRJnmYudBqExAZ7HT4Hn0Q1jGOg3DP8yDL65LwIuEOMXtOfiYuwmTGj6PuxbdLTokHYrvmR5uIamBg3V41E/4jwsWXI650lPZfrTGOfnNJnoD+jpQsTNaQhFz2BcF7bHWVWLyZV/xfVHrl5kpLekOrXl1KrGRFyJcls1FuHKMykfGXoXMdCz5GN4XFfH6OcmWp6Ktn6no776Atjw9ZUTMaXyP3BHxTC14rV2O9sOVI9DQ+VZMPExvC9nsy7YifMYw2emDezzykIye4EBmhLGYC6vP2BcqfLO5WkZAMEh9D+F+bgFgcgNXGg5HLn6IsCU4fPXzmP1WQiXnwYTHEO5/sx00yyyyEeU6y80nanvv2c6fFY2jwPC9tAsZnzJ7hkA6Y91lS079ZXXod7eo6qz0dpvDMvZqRDh/cJ99G4XP3hIqblgZY5Hm7mc7emmKX111qHh4E/RUHU1bDlLtK9V5yI0gOWJdVpwE6NP1gd8Ckcuh4NTWe9+yfDnJ8LbOKZU3wb7BTAGTKvrntoC5bFL4TiW9bGArLl57EsAf4JnzqT5OcB2J27bfBoB2xLzc9aVX7M8TIfdmMRCC0gQIsN4OhlwbkEIt6Gu+XZEzR+Zxv9BaIMMlIjB5GrtdzJAlTMv9RUfob56GuYv+CX6sj32eL8NXAB2c1mKekfXNbVhn2PQCINfwXinJupc6LEzUF91JaaOyLRPWDPG5OcNVf9C+PEzscIZw7J2Cj1dQPMeTWpt8AqMsXKdgsjyMawzV8P+ZJ47OsL+60HKWEdTQ3MR5n98JuCdBpgxsPnhSeqI6WLMK/RiNw10vL0Dx0fAXZRxJmN8i8cU7T/rGEw50xP6W19Pq/6KebiBxubFmvMQkzEcO7D9BttTZD5Obd8Mt4iJvEPT3j8BPIrdHLKUdqnKRRkgu9Oczzp/I0z4VuywyZmwm42QA+WKx/zNo5lEU8M++EKEB54OOxYy3nlkYzenZpCQ+Yr39zqA4/hg9HSWg/MS8SXa0crb4HKc01BZz7IyFvWVZ2J5hBwDp8DI+TCYwQSiNH66HGLOQRg1GPtaPz+PnXKbWv0GZZzaLjv7u/rKcxDvb+/3KeR/P2y9XD8Bu6FrAiR0KkKxX6Khun3slch71T/W954jG9s32frbUD0hIW/DqrGZCZ5GWU9jKk/SpCpXdKI25g1AOlffjK1vcgeAJ8nHtnHJ+llQBWnKWR+T1zc6qlYCqwmUxoPnamn1qARKmoAKrwQKTCDxhoCpAqQX/JWHgEziJFKayRP/SLq9q/vsQObxKJpymiSag3+RGQgOfOdrR/sQ5+EhXqeaGAYHefvAxPaE62ofTVCqVxFw5/Tng8ePVl2lPhh8Bgnejqsr0r9BlDqW0nEJLhlFYTem8dfG3FlUbVow8hFEnvAXmq6CofBat+SZaiWgBJRAlgTEYBonW6dV/peTq79i4P/S+GmPY5B5EEyGke/BcXZHPLAbJD4ccI5AJHrc1yYmRyEeqEDc2Q3G47jF/IwDmPsZeerxDWDHn7sw/tFcmGkAvEdR23R+7t8+Yr4vG7oYdoHV88ZSpnk0fvpTeNFLEMPYTplw8AKy4CKSczhE9obgWBi8wIQNzWod4PR2X2wYD6y2KNrjmHkhPLfpeMp3L/NxOvOzPY3wGrw2NP/j+WFA2/ewXDgJXfEnTB35CCdN52DKyPk07yRMYkF65CxMHfEAF83/iLYVFwPxIzjx/GNGxIVTjpcZETXjliE8fp9x/xGx0B8wvnFbXudG2wVX+2WzuLmG6dqJ1UiaiBcxv/UIDbyZ9ehtjiFiafwncc7CKrFowEWuL3rdyVCX0XxM8402iAMyDf3N+ATjK/fI8ItGrA92cev1BQ8yBpdsOTmN1MqYufTzf6zbe8OJs77G9wMCR35d9+H9AHCquWiyOwztjdzLyFItVgsgm8AI6zy40LroEtzEcoU8KHfoCkwd3gK78CzmaaZgaFJo8xbgXdap+h5jexEKno825/+4gHwUxNsXAZzABJ+nWTttMYNp56/t/ZxW/TSCrEcmNAYGP2cAu8jNQ0odhhimKeNy346mTBNcmFyG+pEtCIUvoZz/WsfnSsD8CcHIlZhc9TJsvVvHQ0aXF87dBl7kCt7Ps+l/zXG+xzS5aCQjESv/Fcoqb2Hbci8XwGbhsuom1NNMqWI7VH03gm2/ZXv7S5blKhjnHAg+YVzgMUQZdwJwLM9/CgjbMQxCp5QYrNnviNj20SdG+RTa7/jwSeFkf/6xhuMa279vuvFV8LyTAXkNGSnzMTxzKfb8+C6WmVdg65zrehkFzdaTjfdqLtDa8hiK3gAjZ8GY95NEw77UPMFy/n3sWfUHNIx4Ju3X6iyDhhEfo77qASA4GQZr9xVrJsLYITIbrbHaTrV3UdveRc9mf388YuZAOOV7MF47Pvh0zeR4HmY++8AF235kpuzPjU4d0YyVTi2M/MM3kDFLYXA+870/TGBXxM0+lOeAr/snO05F7CA4she8QDXjupLmM5pUuj8gB0CcevTyHk38dBXyoOxmDjsWCsdv53jcjpfb2yLfpORvCAenJMqEOyrVAvCqGMTgmlGLMHX/ZoTbbkI0/BO2fWPJiu3xKi9JD9KXfs5A6IvTkKsNeknTWcNSKOu03b9CQ9VjiMQnsA97eg1Xe/ol4F2LcPRGTBn2DtxRMWtZMOOKh4bhn2JK5T9Z7iayfH+URpZ5CJV1rr7FbH2LnAvTegIkfiCi2BOQCwC2YVhTSQhG+mD0Pc6atnquBJIR0EKSjIraKYF8ENA4lUChCcSxHUWgMcKjn25CYMmLfh7UjQRiyyv4aLMrz5LzNOZDGMyCu+vaE4zlVf+l/dsMl0qHATkeGNkbqpTAagKR6DE8Lafx18JJBPtmlL+v7uNqZGTazBg+NoUDt6T115UeEg/y8gyT9JuUoDO2RiC0GR/4krczUKUElIASyICAXQAUYz93m3oizU6sBnAyplRdjIbKhzG54kVMq3gzsYhbz4XJK0Z9gtXm8qoFmLb/uwl3O2HdUP1n1Ff/EJ+XD+ak47Fss57iJJl9AyvV5HpfALsAcgWC5jGMa9wjL5OPU6sXcmL3MaZjkFJJHNJ7KewkdGeM/dkdy9mymlL5FjneB9P7YI75bmbSqyZvbVMunIPZlFZFqqebAGoad8TglTMgMoVSbgVBiEerLUfbb41DZPmoxARu/UEftW86FOtm/aQ2Qj/2p3jqR37A8nIPVgSqyOdilpfFDNQeXpgasCHtT4KD+1HbchA68zUARryWnsbFIyP/pl26SdxnAXl8vXE8slAd8WrfQAuhkRS4qEVe38TBZ7PYzRhXZevVN7aZntlFo169OO4w8wBJVS8B+ynmsPNv1u1nWffbN3LYMr267icWnirfw9Rhb6Bh/3+hoWI0Ys5w1vdHeR9XlXOsqxxabEQzHu+3Pgj3qU3oV3idY01eweh7XBTiveNyV+rY4wjLspzU+SsrP2M8CxKsJldNRxTVgNwGrJV+kNeZaZf3ZuqwL1i3psN4o8hpPvzul0gZ7Fc3AjiPbahtV9F1avlCpvUIZVzGY7s2+Jxt7hNwRy1qt+jA/5pnhqA8cjkEowFZ89nnSxhzPhaWH4+GypcSC+6WF1IolwtHtr7bfqyh4rf4MroDZbVt2hcw/EsRrNPWiX4Hafod4yGK3JTB7tDvdAT6uTu0cSG/mXfyJ7yvtiz6xyJi51+exujRcX+POXa15TBe1gTIwwBHSPyX0LYMGpmFcPkYLi6+hdEciyQcMv0nhk3DK6wnL/uGMBJBZOVStlOdM7ZO2y+CXFb9Pqbs+zqmVFwOJ/hd8p8NCoJ2xXZdbHvffpXNf7thwondyPiiKYM5spQt699QP3wO7IL6tOq3YeW5Yo3xqR1f2DGYXQyvr7oIbf22hZGrAMM2SbwUcfcjx2rEcS/GNx6c258EWCNFWxamHmB52bKwhsN6pzHYn6aw40pkqWwaV+z3CaZWXcPbYjdafsIYDE0qPZBsxiG2+P/Yh3AeMpW3PNj37vMeubOtNMzv1/G/BxN4kn1IqvHE1x679kQMvMjr7N+e901XnAjeXZqDtp19qK1vtjxfXvU66iuuhgkdznv1JNNfXY5Z30zH6hsjUd2zCLCgCAtMz8p0t8mtgf2zD4Af8cx+FoQPheYpQKzhOawd3YzdHe7X4ENV/gloCkqgoASMEcS9HfjQvFUaOaKwvym1YLBtW9J47cHO7rzezP0BbHu/xWMyzTbXvIqwM3c9R5cPHoLfrme/poXBoYj02X5NKz3v6QTk6LQE7EJ33Lkxrb/u4sF+qs++eZouP4LHE5/JS+evq93jnn0z1E5G+aXMsToqMObpoJ8ndVMCSkAJpCUgeIV+/H9fnB46pe3i5dSq+7Bk6Qg4zo84TrqX8X1Fk1wLp/5EhiEgsxFZNA5j53BcZSS5547YCsdj3huASfMWVEfiziDM1L0/R8g0cPzNBWeO/8CJOvH6o3VlWQahu97Lhc/3wXPNx3Lh/QE+D1SsIwBZ4jkuX5yBcPTqtG8pIgNlv1bUUEE+AbtpxL4xG1sjlEBkd0j8NkRCJ8N+CWkNx06d2t+QN2aBTxwxlt03EOzdqfriE38ap+iX9MBySyl4Qm0gmIOQl7ou0VNa7Q5dASMvwngmrd+MPYjBZRUvIhw7BpAaAO/CXx2BSOtNmPDU1vQmNLnVdoOlE5gPg1huI84wNruJKFR2Dn3/jWa1DB3b8dNQ/SK8oP0SwDy0f+6aUSbVfVlnz+QCznHI5WaZpEmtYekeEIcj70PEltfVDp8DsXdWX2R9tGN7WVHDduYHEP59E4Fd/L+c5ewPGf+EwDdh289+N2oZQoMmw/5cAFif2DEgL4p1wtN+Jy9ok0XaUMEFOef3rPPRZM5f2wkCkD6FeZ4yrR5E7DjEw9fKfADHXIedh773tVW2J8vb7Fz/W8x7Dtv0TIVgOZ88/GmmPY5VaXWdL+eic3/gSfv8mmlE3/gLBl6CGLtg/Y1dZ8/slx4aKi8EnKPZjjYDsvbLOVhDCXZiKbkS3960mn5lDZccnxrOAcgaZWGd6O34xMTnr2Ob/eUenzwEOGfzHtk+0a+MbMT8jkPEflkxl2PwdCIvjbCtt3MhthzTM5kYfAAT/5QXxafL4qzDhveFcna5dKxvDcNfgAxQEZgAABAASURBVOfUsr691p684bOEGYCDt+1YfWuPRP/3EAIOhH89JLPdKJsc2JhnmJ/f0ZzPyn8GBw6nwXgnQ7wT4XknwfF+zkZ+DO/umRC5EEZuod8X6TfGo+quJ6ApKoHCErjoP70hsJ/q6pNGkHcQjzwFO4GbxmOPdm6LbAljDiMDhyaJlih5P5Jy0TG28gGGfz9JwHYrQS9I27ntF/q/xxOonbMxy9O+aTkI5mJaxby0/rqLhzI+pIuzUdrsiPNoWj+F8FA2yE6Svwe7ndMvfWP2R59IYSas/ORSNyWgBEqMgCxje9M1GzztZ5/tb5KHo6fDOL8mqP9y3JP6rbvEuAdjEYpNQU2TfUtTGCY32hguTpnUE765SSV1LJdUc7wnfG43i+lJ4EkY0ViA58Wl7edfw8t+QaEuA+Q7nDdY8x7YieNmOIEzMLXyfrijcjinIAbh/8zgpOavYYST81innMgWEHM1ItFa1MwaglwoL7KM46qvAPGQTLUvtn6JxSvbkjlnaNcJbwewvJqFvAeWO+MxcXiJzzzHeNE57ZhPISny3ZmY7cL7GwuuZxQTaPwXtMQ5CCbya1zYnJv7yQTX0obsZN1ytJaP/F7YjRZxllngrVUJ9Vp1zPLAulHWOpfFdBLEsRtC/MJvzDZ2IlpDx+ft7dX1Uqd8RhbBYNnXTmK+gle25OvrbE5c10Ev2IX/UyDyTRuZGCfLw0DodrY936SVTdyr/dov84Ue+yc8YzdpPE7rVXWMZ7nUjqP9Ti55+sbFcuh5d0CE890+Hg024fzKFj4+8ucUDPYGzDZMYPXzXIR1+gEsi/4XWb/5z1hW6+AGUca7iJed7xsYSYf0mwtaGO73NLY/dYBAEAv6Ca+z1yvLuRAsNj/Zh00X4o2PmiEYS15PgY0qUipnN3hyMSbOyuPLOPIF5YikFAHyOQY4nWvrbOT2axehCNtObyrz/qm18jE7EstU1DYf4OMnx04zPcplx2Ir2iM2bI/Zf0dXrrputy2i/zH2d7ZtL2B9+4jlFzeRiQc2IpQnhI/LBaqUQBoCDguLFpQ0kIrHWRazcbwTRn4AL/xjhIO1aKj8A+qr/4kp1U8mfhvG/t7e1Oo3MLn6aTRUzcCUqn+gvvImrIiMg8SPh8jxzM/f2dmwY+WZ6i4ioMkogQIT6BvuB8ieSKvMy5Byuwszrc8e7EEQQCUgOyOl8pbAhOybH8l9XHbIYjrcRuOj5Ti4s7f28aBOPYWAE9+DWU3/WU9x/goIH5zQM1QgygdVM8A3s8ZEEY/P8fVTKEc7CemZ5yBIs8Bg9saAtm8mQgslr6arBJRAaRPwZBnbm6VdmonEJ2P3v4OLLWcwXTth5TepZ/u54yFyBWrnpt/cxQgz0o5jUOi+MbKcE+SmGQYOHNMbZYEQikm5M4KILh4L8S6GYP2xpzF2YeUiTB7Oicc8jDNc10NZ5H9M+2LOU6x6swlrKLFl42yIUwO3ZYM1HDp2Gg4vh31r2RiTNAKRCO/U51ge7sQkb9KYM7N0xWNZWePzt5RHzBdcVIlnFoGPL3G4CJCHe2iTtD8z0Lb8ARixGwE4qW8tk5o+gJyAsDk4L4vVDvmhwKo8NJ8SPMZFed4zll+7wYYWWWu72cYEWP/MRwybvLzSYZXeBgFzJXbY/PBV1/k/ONGVAPsWJBTlk+WIB9couwn7zP6tOHwzwDuXnvvTfKPFLGfb9Ec0DLM/P/KNfUfPXLY306qeQcg5k1HMpcm9tjUYeapnmUpb7P1OpvnIxN+mm7zHsnMPvbI88n9SLYM5J75fXtqcpOmtaRnZHDC7fm1jzEdw5HZcM8rOCX1t3aETU+ByZtv9UJBzEHiH8gfY5vVB2WcOz7PXX2EFDBYhHyohZ3Qe4uYSpuHzBSBD2c0IGKeGZSWcD1EYZ4wyGB6Ta8FKrCxn35HcOStbuzkPgXksf359so1S+G8v+rsBE1t45FW+tcu2WLCULFbPhRiOzZagLV6YsVcm+RUxmXjLm59EOe5zL5m9yTRYVqU3Fi8J8Fy1EvAlwMJictOo+Cajjp0mYMxMOIGDEHz8Z1z0fzjxm2uJ34PJpPGhn2tGLcKUka9iStV9CD0+GjEczo75uU7LpRFkRkB9KYFCE4j17gXItvBXHGjJq2gYzokhf4892tU1AvvFFWD1Du71cRjclZZjLHYrA66gSaX7Iho7PZWj2vckAt4OzG0ZjZ/2EHQe9PPQ7dw8Z3M+/HAS2S9n8jrKgsXbpon5Hx+0Vz/0Js+ISB+sDKX7+ZbkYdVWCSgBJfA1gdYYBD5vHH3tMccnfBa1G9TDsXGA9ztG7tfmlUPkKCB6Ly57rR/9dl7HhZOe0vlJ9s5IcuVhy+FhOoR/QG/EJdSZ6HIadroJIBo4n/fGBWQw1lNmISCXYY9K9leSv4lPu9BZX9EIz7sCMDGsq2xfCHMmIvHzYX+qYF33rK5XcnlOYrwbKfJj6I4oNluawj2DxDrrxciHjGJ1+pyzc5bj4H082nVOxw0XyFZH27mokoa2ZT0W+RvHZy2A70L8RrzPP8UmW6YZxyF7ZYxlF8s+YA5D7Pz+cubvKQjbHoN++Ozzji8khcqzuGGyCbFfjZqWkTnMjU9UAZZJ1pfVPowx6BXLQt7VAXks8w6FmF14trb2OD8RDD0DOiJnim2ZW8EFFJnGKL+gya0WfAUD7XdySzV1bOfu0Abj/JvMX0npyXh9Yb+qtsu3/DePp4ygEw7iVFC2b+biDO7B5P3tfHzH6spqUfq+HoXgM5oCjO1WC8GjO5xtrszkWYCM+yI40OF59vqmfWKs5m3ZB8wwhB1nTKt6gmnc4htChAuqzv9h+82G+frLl6ORj+Huw746Xwn4xCvOdxCP34Cxs7/j4yt3Tp7EIRwhfx0j+5MNt+xcvfg6rpyfxGC8zxhrgevb3gvI7L9sUwIsy331RRHeEdVpCXSsUU4brXrIEYGVMKYZ8H6EcNWBsL+v47ocYHcydhvHZdVNqK/aG47zMzYcfJhP9/ZZJ9Ps4cE1+0qg4AQk0h8w6T559iU8ebbgsha7ANHZQznYqkwtpqxAOHZNavdVLr1hP5l4H3yVHIfaWZv6elHH7k/Ak+2YyTQThzIPK2KFnWiikF2mXdfh+GUz1sXe/mmal9DmdO0br/4Cre1qInbyJ/0DfhDfvDWydgx6pQSUgBIoDQKu/f3lLSZyPHotBV5Gk0IbLpNLFZYsvBruswNTeMrc2uGCFHXmAfLk0wRmM++fAML+K+agGJR9K/nZWScCciFNAOspidDqfpjQo+jMZ4oZSWZaDMq8v8HIX+h//XkPkTKInInw8uNx3PQk8jJUkehOixHkpPiakUgxFOI1BfI5733AAhj8k+V9uY8v8F4Ow4Do3uiOyn56GYHXmTU7UR9AoJfD867RRrbnosAlGD9zL84nStck2slULny0Dxed9oZhHV83Kgfv4+Vh/mVp3TCZXnvBJnp9gCZG0/10MfY7+aIcLnuDbcq/YJD8uUokAJH90BbbrUvrhTujHJ45mmkHElk3ZgHi8kewkqKzatNtTWejyFl4Y55gXBEOcQIYEC7idofjjGiQ/I39YgFFTqVNGIL/wzmPlKXykT97L45clA90RCXG4PshFJ+E8bO37kgMGqYLCCR+4tJEYJwAVoSKuL51AQtNIiMCXTcIzUgc9bQGgS9h5GaE8BPUj/g7XFn/ARidVPbTJZP3vwOI/4wPZ3dyEJSfQXUnxewGwTULSqDwBLzAjhSi/aGDJ8m1LGG781ZyN7VNEHCNAxN3+TDgM8gy/4I76t2Ef99/B6xgPA+x/fWZBMdGcHAMVPVcAufM50OnsZt3gmkgvITe0WgaP93I+fC+EBkM4Z9frgTvoyyw0s9LQd2mHvwFIOm/UCBY/40oqFICSkAJlBgBd5tWhGKTYHA9xz9pnj3NjxBdfiLOa+mF7qCWL/oAkFeZd/ZfgTSb15B/ZTfSRZcM5WT9OUj65j94i8z7/P9nTB3GvopnXaHdUa2cl6hnUqkm5wcB3iTsuOn+9FOsumfLZeeuguDCqrFvyvmxsBt8DvbzUNJuIe9Dys9ybAZgw3g5z3Ol7Sf2U49t7dhYTAWcwERMnLt9rhLNazxOeBBE7IbnZM87YQx6Oj9z12X7fQWYh9nmfJLX/BUq8mLrd/LJwR26InEv7bNf6nS2hSNH4qIXuq4PjgYPYZqr3yTnvL78Db0r56cWsURd4qGXYWQh69IArGhNM+9Y4Dy+8/7HlOBxmjTa7Iv+AzZJ46nUnO2mEdYVRHwEt+3wUQjEx6Fm1hAff+pUKALx+CtIzCFxfOEYe7+gSgn4EcjPIMovRXXLhMBn8MyFMCsn45Jq+9CdSZgO+hGT+GkAE6pjBOM4YOo5bw8yw12jNRUlUAwEvF3TS+GtQLSNk5PpffZYH9Gn9ubkxEE++Y9wItV+2t/HyyonOznmOc8zvudX2ax/MOgL4xyiA+/10eTQZntEWseittHNs7kANS07ZS13ry/40CUbpg1nzJtY8EY8rb/u4iHaZj/fSDbpMuR8ip33aUvnq6DuxrydNn3P7JbWj3pQAkpACZQCAbvAi+C1gNi3LjkZjhRK+vPZ9FSUx/bmUVJ4Kh3r64+0fVEt4E2E12Z/z7uwsi8+lIuv3hiOQ3enIMn5ivkHQrF5dO9aXfbEm0zwLprkGxsFW8BIA86faTdI0luxaZUHgeAHMLD30R+Gkb0wZl7I31Opui5YCDhTuCB2DhBZkrNcCF5mXNMB47eJPATBUfBi01AzN/1zBCMsqA6H+jD9gZRZeFxbG+yMIZGN1rbM0ZV9Hodn27iXWV7twliOIi6SaIqt38k3lhWcW4E86pNMGJDvIrS4a35abWxTP6Z3NgAe+R/gXFvgTiTKHbqX6rXiLTg4HQjcjkXldpNS8ebvntGcM3EaKWCEJrUWzsF40c1SeyhBFwM77m5hv/QI27yoTw7sJpmfwZEJ3WYjrk9mS86pNcC2BL+EJzejfFCaDdUllzsVOA8EHIgE8xCvRtlxAp8jFj8YZbE7kHgrrOMRZRWyYfinWNjrZjjmu3yQ4INKVqHVsx8BdVMCxUBAzNbpxZCFuPKArnvDJ71AxefDRE+lUGGaVPopDqRfTOW4nn1Z25sw3myG4UPIeq6AJP724f+9oSpPBGRb8v0VRGryaiBnw4nvkHUmgnE7GbZ60iB1cM97AzePiaX20M1cJPH2JBeHfPJlTBvr1pfoks8W+8iRzkkkky+vbJkuGnVXAkpACZQMgYZhnwHO7/jc+ZqPzMI2fFcEAv+HcY/7t/c+kRSVU33VHDSMeAKXHVL4Tfe9zKFk8yOaEE0SbRbBwZ+R2LCRxDmfVm7iZw//ySTep0mmBYL90StwKtwZxTeflUziHme3cgVEPswg2xug/xd+z1YZRFGkXtzREdRXPMU6/3BO65HBVzByJc1taXLOumGOBqK/xfjGQWn8FtbtOq9rAAAQAElEQVTZJBZI7YLT+nIItoIXt+3V+m65sAl5H7OszqXxWwzLRUqFiaOY+p18E7i6YiVC8T/A1pHUaX0HjmP7vtQ+cuUSMCMZVTWN1YZjnv8g3O8Ne9HtjP2ZpymV/0HD/s/g5qHFX5fE+QDGpP6SSuIGmTKOQQckTrvNP1sM8SFMoIblMd1XEHoB8kv09i7WTQAoLmXbuoaqxzCt8n9wd40Ul3AqTTEScNjg9ZzJ4mK8A2vL9CUEJ+OykS/yAaHr74vtpHf7ZA7sbwAasxiqckJAI1ECRUHAiN1V7y+Kl3j4Nf6eerDrhKZvs422D3GpIMQB+Se2DHNSG5kpd1QMEniMg++PUwYw4MKfVHOCM5efjkyZXM9zMA7zHO4a49i0mFQ22rGypZicXxWPwUoEA0tRsN+KQ9crQZCTO/5cBK1EkubBvutFXz9F+Xx9u3VsBHwAX8dOL5WAElACJUtADFbgWRhMZxaW0yTXIgEY71iU9dONkMkJdcz2ghmDIR4nf5H6+cDIE2iM2TeNO5ZGZ0OFnLcBmQnAvq3Gw3qai5s4FtHQnuu5FNhCkycB+4xjfN9QpydqAwG62xeWma98a8FShAfWMpm/sx1NvdgmbEPtz8k5MhHj5hXvQlbcWcR88FmGOVpfByFyBsbN3A32p0vWd++cTaKsOi2AaetcRBq6KAi41S/yXv6Z5cmkkMeBMb+AO3vrFO65sbZv/4scz8ja53CM+Yjl+AHg3hW0U11oAsasgCOp287V8nnCPmr1RTc5ijEow2cIy6msJ7OZK85j8n9ybeeizkRv73Sdj0wOSG2VQCkQcEpByJ4ho1nEfE7FcmcGj4XTo0fHEVvJQYmxb2SknowpnISllrLKqwSKhIBJPcG3WkLHSfXQvdpHzz26hv2lczgBbE6TSr+DuPcUTs9yx3PEzObD4KupIk3YixyF1mD3+vxYImP6Ly0BiduHrqCvP7ETVrGYr59u5+g4EATgqyQOSPFzMV4b0qv2yaP0/tSHElACSqA0CNi3V+D8ncKmeRtOhiAWPRk3ddfPhJNAV2o7pi0P/wRGdvFJNkL3v+LJUYXrQxeEl7Kfn8eFnKU+cm7DhZzvwZ2X/M1hn4B5dNKoEwSMXTTxH78m/EkrFn7J8VriQv9lQ8DddRlC5b9kPbmPwfzeACyjnxMRbOMCzpz+9Ft8OrByJev6Ehj+JZNOsCeCTgOiB++dl00AXttbAL6kUV3yBISlKHQzs/EeTQotWyLinZq/nx9h+xfEUNa7/SmA/ZpRHOLMhnGeZ/lNtamNXlV3GQHPBNjaOP7pOXG6t9J0T+1WLYCDn3Mc1cQM2rzykFTbfuMcRIPH4Zz5ZUl9qKUSUAJFTSBNY1fUsncf4YyxDe2/EJI/IzERUuCs2U8SxuK/A2Q2VHWSgAZXAkVCwDgZvD3q+f2OYMcyMr7xcNQ0HdmlZvyT++Z+d2rzJjDeKEJI/il2PmZy0uJZuj9Pk52+vMpObNqJG79wO8NxRvh5ULduSsCLhZgz/wlUI22IB6L0lztt3xKa0Hhwl9Zd21ZMbNkto0zE+dAO+HMRxODF7RgroygL5slw8jtd4ga6ASAdI3VXAkqg9AiEP7KL/4+mEZwLiXIE3mvbJY2/4nE+r6UXahvHobblF7kfk3YymyvsmNacCPHbRGfeg1f2FAqp7NcJ4+YlQD5BSiV8vpED0RbZMqWXLnfQBBME3Ce5SCBDEue+/8x83D6q9BdYxj02ABNmnYa6xpNY5/3Hp748snSctM8XMDIFxjwKg6hPaN4L7zzE4scDhm2qj89COLUNti9Evc2kU2064vOQ2M34VyN64FE4bnqaTcCMKRtdvtJ+wc9nwTibyArkt5j7na5G8kVwPiB2g2Hy8iTsAcU7ARtHdkI+lPtkH4g5lFG3v8AhWMw6+h+E2j6lXenrC5s34hjHRU3LSLj2RZlSzJJhm5ju+dpbAuN8XIq5y1jm1xe8ybI5CQZ2zOe3OWVbABPR79MDeVTdlQQumDcYdU2XoLbpkJz3fV2ZD02roAScgqauibcTEPkSnrkJbqUddLbbFfr/5aM+5IDlKnYCptCilHT6KrwSKBoCHifI0ghjJPcbAAI4Cw5u7VITCNYBvQanyW12zhHswgCVNMm1YCnb8ccwrfqr5B7S2IaifwWMH/8g4J2VJhZ17o4EJMh7D/9JLgEnTmPJJzg6ykQiQzgGGNeldde2FXFvbGYix+0Y1prU3o3x4EjxbwCA8Xtra1X+xH4JYtW5HpSAElAC3YSAOzoCE7iHk4/+Pz8n4Lgu/kOgCBeukt2K8lhfiJzG5+nDsSLEhatkngpkF/AqKNeO/qnLa+iFhf5+usDVib/LVN5DYqMtz9bTLA9ix+hmV+Tj0+DrpZeBhXppJxAJDAKMXTBov0723yDOcVpLMqeSswv2HgjjHAfBwfhipf+4PZeZEzEIV7wCBF1Gm2YjumwCz7sE45tO4L0R+i8efcXuKwCZBREekUoF2RbweVx+jx03/Q1cLoqk8pm1/VPLmPYvIIHjsHzpR1kHL4YAxdzvdDUfu4HMOPczWduH8JBMyxYw3nFwX8r9M1Y8uB3L6iGArNrAbV5h+XoMbgG/qoMcqiC2AuRUOPE9gCcdlKJynJ1hONLxl/19mHj32LSRKp/3jI6jLN4Mx1xKHu+n8tZuL9uz77gG4xqr2q/1f5cQCEXYVpnTyH5vDNq2NOtbl4DSRPwIaMHxo9N1bo9g6gj7+3Zdl2ImKU2p/jcHKQ9n4lX9JCegtkqgeAiYeFpZRHLfJxg8znQ37WKzK2Jtyd/UpyBZa7ubHjiC4QbTpNILEQ08mMoxrb07yv7u4U2+/kSGJnZZ+3pSx25HwHgeB/vGN18GAuR4faHcfgZTPmS6XV1/K5hmeh0Q26ZZ4+NXQvBofHwUhZOTyRdakMnPBBRFdlQIJaAElEBWBJZ+9TJ7sf+lCcN+To7kwtXANP6Kwzko9q2/IcUhzBpS2DGt2C9KSe81bNc+NZwCBuYBSzPYnLZ20JxftbV+SWneZ/nw6+834MQ1FwV/FMx5+h2IUIOsIuA4e0HEfwMAzEeIe82rQpT2wQQGMAO23vPQxdoVDw37PwMT+DVg/F8qEtkUDm5BbfPB9Mt2tYtlTZVcYiND9L8w5oVUXhL2AgFkE8A5G5GVD6FmViXG5ODnYVzXw5TKtzBl/1dx/ZGlOeYu1n4HBVJl5S+z/3iMqcdokukwy9t3EV/y7WSOHbZzZwThgX2S7Nxex8T2qg+hvuL9DsdZbAEdswNrYu7m27o6f2Pm9ebC/tHMg6RO2t43mdnhF3xSR1x8LnZjypSqRwGpYZ1YjtRKILIjAnIn+5BdU3tTl5wSCMS3Z3z2Zxh4UK0EOkbA6VgwDZVDAp8jGp2Qw/hyG1XcOYcR6m9hEUIHtAZRAsVDQMRvINcupzFl7Sc5/B9x7uaD18ocxphJVJvAk/RfPMgkJuunt9mIA90f2NOURuR+XNnJr7jEQtcw/mU0qbUTH9+ln5VMLYm6dBmBACfgJeqbnIMgHC+AXCp3f/s1i6cZZdfWX8EQuHMyeMCRGNsWsqGEqXUYjlOW2rlIXDyTfkHLfmWkSMRVMZSAElACOSVgF1vEPMI23fjGK9gWISc/n+v1TThLR2MEdvLfoPgmx8vim3JBYnfmKPViuZgVnAD+EDjAo7/C6vcWt0KM/W3uVAs4q+Ubjtav+qy+KOBRk7YE3JYNWIaO4alfHYjy+Wo6ymKlvyhm6zxibJvM5sxz4fTU/ZshOJ4C+H/KXqQP/V2GmuaqonqudEctBAJTKX8G84/GgcgwOM5/MLj1eoxt2Y956cuwPVPbMlis/U6h7sikPRcD8jAA9mf8n1TL9ojHD2XZKU/q3BHLtsAgjmf+D4L2+Shj3oZ4fwHEf4yDElHtG26OLhFp1xfTdR0MaT0KInuu77iGjfHegMTvXMOmm5+yfDZU/A0OzmH5/cI3s4ItWJyvQM2cnVGyPwHhm8PicUzUN+eo4hFIJSlVAnbQ5JSq8N1E7mtw+SifAUmBczlt/3cBuRWqOkBAgyiBIiJgJP0GAEg45xLbRXExzTmP1z9CTmqYcn8vWbiKxwkss3XqEKYV8Ry0k5cN+4iD7UdSp2NdZASiZXvbMzU5ImCwAJB7aO7Ir8HfYZwPkK1yTBQw/hPfBpzId3I8nhPDB0D7KdHsZc42j2v6N5zQi7VtuKZV0nODGOVLtwGgDF48d21BUkFyYCkYiPRqaXov6kMJKAElUKIEDJoh8N8ECekDz+xR9DmsfdZ+Mao4J+tENofYt2chqTnKMvavC+ByVIoCK/tpWnh2HMKxkJ8sZjsuBG7l56Nr3DSVxNvYrZ799PWhpBGiSaYNx7Yv0uHvmHRAG4+lrS95soyLIYfDyICCZ2RK1UwuaNaR7/w0suwBkcmIlO9Ovz7tQZpYcu0cbn2M8lzNaNP0B/TRrntDcDqC8b8jEnRR11iFfHzSvT2t4v1fzP1OoaiJGJj4UzBmDkXwaJLpfoB8F61l30KulMgRENmjPTo+wwv+iNCI0vxZifZMrP1/45U7QzBqbcsSuoocsjOlvZDGb45gCd0vR/1IO/7gaU/RYhAcdBdzexnNJzSptJ13OhASm4TYLB17paKUC/vBbd9mn3hgLqLSOHo2AYedIQffPRtC4XJvliFu/lq49DNN2fsnferEMyFkpdWzEiguAsvTiiPI4K3btLGs78HgAXS1knDvnCTpziiHwa984zJ4BFP3f93XTyaOwgG3yF/o1W9Rs4z99s84CBT6U50LAiKvQgJ1CMmv82rCfS5G+IOXshbZeBEY8Z/4NuhFP6kmWbNO8usAgaidnH336+uuOBEBPK9P2qRiJso8+9UVwIBMJJw2rkJ7EJN+wthAx2GFvk+avhJQAvkj4OEdNtp+k402bQcevl3UizuJhadlPwJkXxSbco0Dx9hFjkFpRFuEmPMFuKKJYlAx2An4dBshN4CB/URqYSXW1IGNY7vDwUVEsQVNKv0ZxLkBoehzsM8/qXyVhL0RxMoOB+RoCP9QaCUGbaH7IHIVJfmUJpV2WMUrgNgU1M2y7UIqf11rbz9F7bXdyOfdm5jwcprMtMi3mOdz2Q78AdGvboQ7d0eA9yaz0KXtq5j7nUKTbaheCIGdX1mRQhRba/eFeCNzUl7GPTaA6Z3OtPrQUMtbiAb/DFdSbUBASanxjYMQd86lzHajIw8lputmsl/yanivU7/9zxaUTcctiMk9JZa73Ijr7hpBOHoL2l8E9Zt/CEPkezCcR7vstX5QlXsC7oyBEJzBiDehUa0EOkXAYWjdAEAIhdFOM8qXLixM2lmkGo99CIPnoCorAupZCRQVATHpH6BNYlIwD2KzrQPs58TzEHeKKOPR3HwCMBo6FnBSv/1vYBBwpgGSm75UwAViMw+pVYBJVaCm8TupvahLdgS4kBwMfgW34sv8mr0WMz2yiAAAEABJREFUwR3tv2CdTHATZBjjP/Et0hcStxt4JFkUHbZzRy1i2Dks5VEeu0oLJJB+A08ZyMSQjY9YjImLHX0wfXrAx1cxOA3KQAi/B/AMgqsXJaAElEAREyir5jhRPvKXUByOgTYDlvm9teUfRT5d7SJMZNFP4TgTAdkARadeDgLOphRr1cIEz5LrKEKef/+aPFx+bMOBL7gY6P8cIxJAwOyeHwEyj7XH+6xp2Qle9G5y2JvGoUmmPRi5EYs3+gvcUa3JPJSMnf0074SWk+HFb6fMxVPnr65YiQ2cP1Om62j8fsorBJHD2C78EWObimcBZ+rBXyAcaIDn3UL5sykjNj9c+JeTEY3+DxOa/4zxs7dmHN1XF32/U2D0wjmaFYHHKcUsmlSaz9Hm57jwhfTPn6liWG0f6jsSkP3QrjzYrwz2HrYA3UHVztkYAUzhOOwEZkdoSktb+Y0zBZDjaDgewvrKIA6YfyAcuAaXV/XcZ287B7RCrgbE/pwrmSCFMrbO/AyLF14B1yRnmiKkWqchUDN3Q0RCFwPmp/QZoFGtBDpFINWgvFORauCMCHCxyHsSCzbyG5BnFFHePcU2+JwPBv9jOpSZ/1VnQkD9KIEiIyAfpxdIBsO+8Z7eY5Y+yuzbB7OzDNQ57450/qGkncXZHHSljkvQjECg82//r87t4sUfAvI4Jzp9BtrYDk7gALiu9uHoAcrBUoiz3D+nRmCcb2G6yUOZMI8x7Wwm3+i9C7T0WoFM3or3pC8a+xb5A6mk/8kDwL4B2QVgNQkloASUQEEIeBCkmSRnXwdsiKiX64WqMqBtJ0xs2iV789S+mND4PdQ1TkZk0cvMg31rlYvsCVkLAjJlosviIY5ph7DvDKf0k3AwCzmmyPTz24kQef1nhGMQa9KkYrBlGh/5du6Z8dsvS9iFlQlNp0DiLYDswHqw/rOTMXGWvY+4vPIrhCP1uH6HNhRMmT7wAjugI3V+/JO2zh+OuubJ2HDls/DMLRAZULCspEr4/IqVWOFwAcfcRi9+G3r47CAHI2iuwAUziuetXrsxuzVYyzJzKeX/lMds5iKZJ/RnmBMRiD2LuqYrMK5pF7jz7GIVoysGLWUIl3+7Q2Vw4qp+Z0LTpUXf7xQDarshxuAaAH792jCElvyQfeT6bRcDZqTdZwfCxO3b/8F2/+YtOM5DcMVDoZTNjZgN0KfPztmXtZbdUNdSwfpzPM2tkNgLgNi3kYuoHiGdksTmpvHN+8JE/8q2+iQIQskDyWK63QnxLoA7nHNyyX31GNurK75EKFJHJo+yLfWbmwzSz2mINJHbjL49hk+qjAoGYouy72Rf32buhonN+6O2aTRqm2+CRO0zxa8BKFNCUN15AnZg1PlYNIbsCRgOPiTwHG4e2pVv1WUvpw1x5R7LYeRlnqZZgKAP1asI6EEJFBuBwBtpJRLTC9Hwxmn9Zeth6WdfMsi/+ED1TOeNPMu4ltH467QLpv7BE66R0Cged6RJpWOwu7qxMtUn5ZC1uv7INog00/gt9tmJ7wMQOXxI1vFrgNIjsAKsP2ZxWsHF+xZeeTmQ1l+2HkKxeRA0dr7uGtZ/eQaGU7/+MhiYQAZ1qnwJhIsU/nGBsm+G3qHifFsUVOfMLyOTbXnmrx2x4zB/P+qqBJSAEihZApewm8GSNOILHOnNtj/Xbfp2nKRvgoeXsjeRp2Dkn4BMYH9jP0FfvPMrfVcGYNAHQPHKSOHW0158Bdlm8tLEwMJujl1P8u5tcdO8EGrm7IzY7J9zYegWlq0bIJLqi0YLIc50BJz/w7TK38J+5r2gdGQ/xL1/Z1/f2UYEgrbO2+faCXBkFwhyP/ZGjpRd+AwFJrJx/RMgK+CnxDkJZcHxXCQf7OetS92s/A2V02Ccn7Hdfxgmgy8ariegDKTVhQiaxxFrnYBxM0cwj71pV2i9LZx45/odg4ksf8Xd7xSa8ur0l248CyL/Xn253lEkCMG5uLClY/MrxgjallUxXvv1Ex4QhZHHEY2ln4OzvvNnhFEfj0DgGXhsv7IyHhf8vWaG/yvNKTQb0dj4eChy7boOxs74Fia0HI2QTEHAzGB7fQClTiZ/K9uWmfDiNbxr56B+pN88HKPoQdodtRCQMwE8QuO3duVAZCyioVPhzujZC9bGHAPP60B9C7wAz7RA8Df2d2N4tPPyycorb4VqJZA9ASf7IBoiRwQ+hfEW5Siu/Edj4p9BpHTkzT8R/xTUVQkUGwEv9mZ6kYQPw7J5en9Z+rCL2p73Nz68n9JpI875gLyDdMpI+k0CfnHYRTngaHqxn1XnIal+E/Ho7JxPYsU8+8WVl5mioUmhzQHsQ3ZK4ajW3YmAnfwy+IxZ8tt5zbk92Q54J0h/udX286wS/zWM0/n66+FsSDrxDOB5/pOUNgr7+3SQTwGJwE+JbA/Tq4+fl4K6Dfh0azKxm3r8xTCwbYK/H3VVAkpACZQqgUmT2PgLJxrTZMCgHHFJ8fZWmrA93tkLsr/pT1O8C5bJ7lE42EZra3jw0SJc6BsZ9vGRX6fuHLt9LrJvt9pP/NfMOga1jS4+aL0VTuwPMMZ+Zv4oZj/JxhzD5zFzN8d1Z8CsPA+T958FzmxDVdcRsG/Sm16TeA/+DsPlLaRSpjdETkJ05ak4s5gWcMSgoeJRPhycCXEuZh5eo2F/kSofqexlE4Ybh6BzC/N4BSa27JXKp9p3QwL2iyMebgdM6nGGYE+EvWNgF4+RpbrkyQEQHMUytuGqkJ+wqXsM5VWLV13rIV8EXONg2rwBGNe4JS5uGoWaxvMRPej3CIXYP8VvZh91DpNefy6gvT1sAmQcnNBpmDrixh792X+kUPWV78F441m2nwTEQ0olg1i/zkUk+EPYDYIp/amDElAChSCgGwAKQb09zUXo7AJVezxd898B5e3Ijlv0SKWZVgJFSOArymQNDyn1AJh4fhaVp474HA2Vz3XaAJ/CGEmZg3aHGOD5Lwq2+0v9v88nu/Ihbhg9BGmSa2NaUF6eu8//r05lWrW9T/8GjM9kpwzm+PtouDNSy7c6Pj12AwJid6JHfTMiZm+sLM/PosiUkfM7XXdt/U98xjfdm1Kc6IuXpd8AYGEY8wnrib9fY7aD4/S23ovSGLMz5Uq/YBEqe4X+VCsBJaAEui8Bz/MZ9+Q125+xL5kIg/OzMh4njY25iWHmcGzKhc68yqiRpyNgTBmWRQq2uSGdeCXjLvIDRL25qGt67WvT/7MXEV0+D473H46pbmVe7GLASTwOZ91Zd4zF5zB5im5nIBrcD3GchbL/3oeGgz8FOMZDESiDNyFSz7qbZZ3HRZT+DzRcCDH25RjD8+LXDft8goBzCQX9J9uqOI+p9EZkMhYDg6NTeSiYvX0rd8ni38Iz32WZu5xytNJkqx1AdqQ5DV78QdTNuh41czdEYdRnzMdE8s6yDK7udwz7HSwtjOglmqqJs12SJyh9inornFeRExEdleULMZybagvswTZlBAR8Fk8skj4DE2qBmzhHAZVhGZsNccbymGVZwyS2F38B8AqP/vMQ9NQl2mAjEp71dd9k+6lI86tY2vo0gs5M9jd3w8GlgGO/WHAoIEN4TwSrlaVh8D4vp9G+AqHwcfi87EZMGT6fdqpTESirfo28agHPtjsmuTcjMLI13abgnbZqHnumNmYeYNtpZFffDC6GwZ0AXqKJ0KhWAjkl4OQ0No0scwKCRXCCpfNJ/Zi3mI3RMqjKhID6UQLFR8BDK+vwR2kE6w8j3+nQruc0EefMORbbHCL+n5USfAnjdXzQZHcRB2R/ymwX5nhIoUV+gWhsMR9ATM4NcD0gSd6kwSrFAbZjTsLKfuvvZl7lQw/diIDAPqimefCWbRDus2VR59rx9slAvgje3ffzDPwBBgto/DcAAFsh4A3MKL5CeBLZCcak2wDwEZZ9sbIQ4mmaSkAJKIHiImAXEHO+T2AhvPCNaKi6OisztfJyNFT/Eg1V+2N+1UCIY9+CfpG8uADK/8Wo7eR3McqVC5kMooj383IRVQfi6EZBTH9mZluaHb4xZjtAtgGwBc2GfBYr41Fo1tDmA0B+hFD5ANRXDkN91e9x+f6vYlr1V3y2LK77IvgAQe+PrLtZ1vmqK5mvU2lGYcnSTZjfH3AM9wIAv0V1Ohdas92cUvkWELyE947y8jqVSCKD6PQbjG8+EMdND/C8eLT9quC06rfRUD0eEtwdBv8g/+U8pliQSim6XaRlWXbOhkRmo6bph3DnrbuRJWXg3DiYL9Aa+z063O9Us9+pHIRS6HdyA6zzsez1if1ZvQcZUernTJGdYZxDsir7580uh+OMZDm0bSSfT00bjPcQGoZ/yrQKqxM1Q57D4kXXZV/Wqi5FQ/X/sb3bBTHZEJ6pYR4XAAXd1GA3aVjO3/RPgu0B0M5sxePGgHB+zDhYVxlzFWLRLRF+fBvmqYZmHtz9PsHNQ6PretXrdQi4vOf1VfNgApdCYOel1vGw6lLoKrIFAuaPqJm1M1x3/fuwymu3PYjzEj6KZF/fGqomo6HqJJbL3TiOGsS6dj6M+ZCcimv8RIFUlyYBVkaRkhLdGPvE/wVlzsQsY6VJdHn0X1zamEUIxkpnA0Cw1yI25UUqLzsjYClvcCZlwvrp+MIgE0mv1YcSKEICTpj117yWRrIAHOyAyIEcOKfxWShnB5vxqaqfb/IG7yMsHd+wFGm0+T+SaYRpilcbGYzAylOLV0CVLGcE4vIS40r/posX/y79Fa8WZLIB4CXcI5lNpAZkPkQW+2ZYYP/28PVTKEeXD8UG34YIJyL9hPCeR99wZkz8olE3JaAElECxErjkEttW+2/wtLIbbyUCjp0PsFfFY2y/NaXiIUjoaI5T76Fgxfe8uXI5+5HEF/14pITdTYv3GezPJhUkX90pUXmTubkbkNvXMgZ/4dwaF3+QXBn5Fsv+D7CsLX09Th5Dadnaxej6ygcRWXEI8z2dwhf/IlLD8BfgyVmA5/9VKZEBcMxfsOOmR2LMvDRjVOa8ENq+sRuOHg8xPwPMQzR2UddkLYrIDnBwO2JtDZjQ9G2UkiqFfqeYeI4eHYd4cyjSPJoUi2pmEJ/LvodtttyUfjLT5fEtWP6OhsBujGIY8wrC5SyTPO0u+vKqpSiruhxe/CTm9fmCZcuYlUyfY6z1+qc7af8s5Uq9AVPkUITC22LnnYX+VHeEwPwPH4eH81ctTPvFsBUcmY7ooUNhX67y86lu6xNwh65A+PFrYeI/AewXBTj6Wt+X2iiBrAg47KRKrfGznxw5l8U/vYHczfyxg8iKSdd4diSCtpXxrkksB6kk3qaVIn2oMUsA7w8ZlQkDlhvzXA6IpI5CXZRAMRJY9sUytoeZDNZ3ZV2yb30UXy4Sg0fZioL1o0mt7a7USGh5ag/pXJxd6aOSpgS0nI1xjw0oAUFVxM4QWD7ETtR9lj4K5yco1p+FsIvdQPp6ZfB0+nyu8hEc+A5g3zhbdX8erGsAABAASURBVJ3qYFCVyqmg9u2brbakDA5Nam0CLXRMPaFBR9VKQAkogdImMAkQZwj8VOLtdVmE1mAnxnh+CeTAbcqwdxB3rgbMWzmILbdRxL0Y2jfN+c9BGOEihgnmNvEuic0uAHZJQusl0p0sjGlGyDkX9ZWnrGXmf3QyEp9WRnLOAuGiwAkIeeNwXssG3QmJb16uPOwzRGNj+fxcGnNMUytnw8g5MLAvBhifvG0EyMXYaGUlppvi+hIAVil3VCvqR/wd4b4/BZgnmDuZL/vCD7JU/Vh2z2LY32NC84FIzDlkGUMhvRdzv1NILsnSDo74gNb/pFlGk0w7LAvDEYoPy/grAA4XloH2L0caRCFyK/BIBs/tyZIvYjtXPJSPeJIScmGS+eRJl2tHvmLK56/VN9m+av6CX0BwAevw8zSp2rXvAF4NXtxsJ8BIl8veHRK8Z3QcZZUPcLxeS4Yf+2bJgJy9SYjP3oV1Snz9quP6BFzXQ9nIZhbVa+lYnOuaFEx16RCwnZtXOuJSUnHeQkPVXzIyYuayoYkyVPFpD33hiM/nnYtMZEd6k2WxysvGUGZlVCZs2YHzfj7patxKoCgJ2LcUHHmDsi2i8dNbwXH2KspFxMSb+WZPCu//Zn7iCwArOzE5bE5kGv1pil8Ltkaw1zFQ1b0JXL9DGx+y7CJwmnyaXdAm+6bxVBjnyIG7MmE+ePO/r5ZnfJ3XdHR3jfBhf+6aVknPDY6CO78I63Rgb8q7PY2fboVjnoU7Ku7nSd2UgBJQAqVPIL6Jbx64vEj3r9B3wAoei1cv3+gFCjeHprh0WzwG49nngJivYII+8LwyXz9d6bgy3hfG9E2fpPN6ej/58dEjYrUT/8Hg3czrjRyTJn/OEglwjutk9DI/x3ktvei3Z+ihn38MgX3j179uFQuN+QtmwWAczfyUIglzBOwJTy7Fc3N3QTErd69FqK+ajlDsfIr5Y5q/s83IdtwcZLgDGO43iLUcgFLbBFCs/Q6KTLnCOhp/mGX/5ZSSiWwM4/0Iu2zZJ6Wf1Q7uvzdgXCcDq+b2Ba+hNXovXC7eoRsqVzyY4L+Zs+La5Jjonz5u4X2bxpbrc8qXTLOOyyjEMRbjni7CeYFkIhehnS0Docg9MAGX0n1Gk1wLxwPGHATPTMQlczdP7kltfQlY1mF5nH5S99V0VK0EMiHgZOJJ/eSBgKA/4r1K56FIIv0h6OLfxcoD9/xHqSkogeIl4DnvUjj7O0I8pNRBPsQcgpX90j/wpIwiTw7ibAXIUPirKMS8CfeA5BNT/mGBqY2D2Nb9JJ234nI3Y4r284zFBaq0pfGc/2SQAQcStBNfGXjtai/yQ6YYpEmjA4+l8bCOc2DWOhbrXwo2RPQzu7FnfbdC2ZzzSBkX9ocx+S1o/PRbbNM+ogdDo1oJKAEl0D0JfDGXi7zOZmky53GMxnHs561p/BXW2W7ac1B8GwA+WRaDOHaydmUaQAMQDwykH6EpvJZAGGl/Kse0sa+0n64vhLz5STNmAvmJuBOxusOXsGzcxhga+byYalyyAeCdiT5mf8AURxmiwHnVxx3nwXNeYH5TvVWc1+Szjtwuln1R/i/AuYphv6RJpYOse1WQ2NQS2NBh4I5aiIaqxxCqHI0A9oHBPJpU5TRVnnfnIuJUxJt2S+WhKO2Ltd8pRlj1Iz9gub7TRzSHbkegtS3dvBMQ73s4BHvRP6s/l5Y970ZcxXKYsOim/xqGfcb+1n5uv7gy6I6OoCz+ABI/X4NUqgyCExFs+1kqD2qfAQH79ZVIn7vo81aa1GNKkTJWjGMRjY3jUehXdbYE3Aq7oaX46lu2+VD/BSdgO7aCC9EjBRD0gxMrnQX1mPTnfSodeSlsYbSmqgSKmIBx3qN079D4a8GBCKzgYru/ty51PW56AMbswwH7lmnSfQfGeRF8KkFH1HKZCEgQJaVkW2zYekhJiazCZk+gVf7LQJm89XgUJjWmqyeMqgv12KZ+XHTIpIw2oWH4p1lJFuw3iw+UH6cNY7yz4M4oT+uvqzz06bc5PJOOCSctzTOIhbng1VWCaTpKQAkogQIQGLB8CMd4m6ZJOUb3l+COskeeFrH2zHPsm4pro4Jd9DPOh1wQW5yGXC8E4v047k7jrYucJb4hU/L/uSsjXwChAn0BgNLlQ4v5FqMtvgnzafu/y/HLFNZXPm9RwmRaZBsYrwbjZ2+XzLnb2YkYOB7HapLdGLaQIG4eGsXSIX8CvMsoxlKaFDqxieNI9PLugfvUJik8FZe1Kx4mVz+PcLSa5fQkwLTQZPpyAOucsyficgqfG+xGqOLKm580xdjv+MlbSDcvfg8Mn7FSySAyAI5zGs6ZX5bKC9x5gxE3v6K7QwOWtZcQcOzzOrq3EsNxxJOsU6bo8mnHh219L+W9/StljKaQj/fLTEZd89Gs4yU275ciR4WwvnKP5QhFrwC8P/iwtpI5dD8Ltc3Xwp3T31qoyYYA65uYWbDjjGyCqV8lsA4BNnzr2Ohl1xAwGMxGsF/XJJaDVAKeffDWxjodSnVXAsVMoH7YZ4A8BWAljZ/uC3Hs7zrxAdjPWxe6Dd2sNwc9x6ZP0byDNtjfNEzvdV0fdS2b82GBkwTrOhT99WCIOZIPob2LXlIVsOMErq74kuMG+3lR/zgEQxAxJ2T8u4X+seXGNWQOAMw2aSMz8kBaP+t6cHeN0Op3NOn0doiGD0vnqQvdh7FNsz9p4pOkLIEnzZg6zO/tLJ/w6qQElIASKBECptfulHQQjZ9ejFj8BT8PReMWCi+gLJ/TFJc28Q8gsLKZlIIZ9IMT2Ayj7ymOuSKRDTiGCKWUt93hLQTLMue9uCwO49nnoRQcRBhtEAvesEee+uie6FRWPRueM4njUluWkESRm4ziovglqH184yTu3c9K5CtmqnQ2AFBY2LfGo479SQcaLLdWPuZIRCIXY2zTZj5+isvJvqVaX3UXonIcy+okCvckj6kWBem8Wpswz45FazD9G+D0WDS6WPudogG0hiBTR9j+4iba2H6AhyRacAT6fToiiQvgug4irYcCsgsSyrRBvH8gWP5B4rK7/xPvNRgprk2Oq5nbhWlxxkPkUVql2jDKtQ1zGSJlI4tqzoQCl5R2Ry1CqLwextwLg9Rtq0CYrzFoi5wN99nS2lhFwQuujfc6GaduqwouoApQCgSK46GuFEjlXsZN4ThbJwYOuY87tzFONwEYZ3s26INzG3H3i01zpASKmoDdNejIgxD7lkw6SeUo1DUPS+ery9wXeScwrQoaHy0RQJ7FFYnPJCF7FT8eRtJMPpt3OPh6pWsN7CdN/SZlghBnf0Tadsw+zxqipAg45kb4PVy1Z6YPy8Ph+PbmW7dfFvi/27wR25zRrDMbpZFkGcTMTOMnuXMkdiu5+Ly9ZIMJJ/O8X6JuZrpP7lvP+TXujL4IyFgmUk7jo837CKKJHlIsUNBFtRJQAkqg1AnYrzzBO5LZYDvN/6n1yzCBd1M7F5NL38/g4HSaq9E7WjyTdmWhT0jJvikf5zG5FimDZ3bAoG2d5B662FawFeDzdS5j4hA0A/Mz+UoSEmrAh3FAlgDwaNbXBiEY9EPZgWkZrB84VzayZtpcaJFl4BNIrmLvVDyueCjb/yHA1DCeVJsUAxD5EVA+CePmDaC/7q2DX37M0jSB493fYsMjUi+EFBuFy6uWIhSbStmvp2itNMm1wP79H4L4NS6YUVrzgpdXLUBD1W/Yrp3OXExlubXtINIoztfK/5XW4uAa/c6i8tT3Mk3Ge4xzNPYPtvPzUudXBrC8nJG0/Wo7cEOWox/QvVd7eHmLdf9xTNqnePr7dsHy83+59wIgp0C8f2HSAXEUmwo99gHbtAaK5fczINtT/lrstNWuvJdCv6o7QsDd7xOU9bkQYv7K4KnLgqAMjnMOostPxoXP96Ff1ZkSCJW9xLbqFzCBB7HZQ6kZZxqf+uuRBJwemeviyHSYFbgKiw9N/Umh4pATeLZxAw5m7Ke30+28LxaJCyWHpqsEip/Aax9y6sjMTS+o6c2H5Kkohk9m187alAua9vOEadpLs4wD/UdgNzqkz+DaPmrm8iHOORqCAFIq8yoncn/M9vCALjVe9BCK9G8aH212gxPft7QmKXyyo07JCayMvcQymkH9xXB48SMxZl5h+21jBDFTxfHO9yHiU7cS2Z2JkLyfOMv23xWjPoXII77BBEL3aiBwbMHrSTR4GZnsQXlSa/txRXBy5bUFr6X2pC5KQAkogW5AYIdNNgLMgWlzYnA3Lq9cltZfMXiwX6eZUv1vTK6cDftJ2mKQycrgDlvKMeyzgN00ixTKOOwx90CfSDCFh66zvm5+GeXdHpDU45nExmbTgknHZbfoKuyJkUqRgWPKEHzWSeVjlX1+Du5LYUY8hEZowPsRhUgbcAmKRrkSwyab/I3yXEWSlI1n62kT5qLASQisPJXPlIUvT+vJl0ML96gVmFrViPoRT8GV5BtLUKTKvsUZlavZDt9NCVO9MUsn9IfgNJSHToA7PWwtSseIwdTqN7DCmQYjv2K78pGv7MKcAkfg21tu6+uvmBzX7HfsTzwUk2zFKEsfLITxbJlf7iNeNQIrqtZ2N8I5oQqWEPuyjO0jWGfMHMSdF9GReSiUoLpm1CI0VD6MKSPnF2WeXdfDmxVPAc7lvFepfirQzk1UIx45F+fNTvMSUAneo64U2d17AeKBS2HMf5GYw0iVuNgvAl2IsmUHw+U4K5U3tV+bgDt8CfuvhzC14jXYsr22q14pgYwI2M4qI4/qKQ8ExIzEAKzaMZiH+HMX5RAuvu2Tu+i6a0yaLyVQAgTuGR2HxG+mpHxQ4X8/7cgwRELnFXTC5ryWXnCcKyHIZFA+F3bixS9Pqdyc2AGAt2MqZ7QPZB9Dm7wC+8m4rjTTRr0LkSco2xKaVDoE4/wA2281MJUHte8GBJb2WwzI/QBiNKm1oBf77QsweMV3UnvqApdxzZuy7tjfRuyfJrWVEDBfFQvT+EvlbOhgJ3AiPPpo6QOYk7HD5rv5eMqvU23TKEDOYH4FfkrQChO8A7bN9vOnbkpACSiB0iYgHOftD8jW8FccC4UfBVcToaoTBMTA8f4DE0/16fbVce+CPo79DfrV14U5vv/ZACa8HfvuII9JNPNj8DLigVchPE/iI6mVOzoCD5/TLU6TTAdhnCEo27Q8meM3dvk6a7X57QcDSaRgsAzx2FdwuaiRsCiSf+fuwIV/uQvwHuXEfwqW0pe5OA2twQOgE/4oWnVl5Wd8unBZ5v5BGf2eMwbyXk9DZNMTYL8USs850XZDwYTmKahrXIS6mck/u56LhK6uWImGqumIO2cxH7YNSB2rYGOYWEVqD+pS0gTcUSzn4Zlsn+zb7CZFXgbymfoHGPeY7YvavYx7vD9MYBSMbNJugS94fADTqu3PgPBUdVEQuEfiCLX9i309+yi2bsmFCkHkBJSb4ws655lcttKynVZbvWXPAAAQAElEQVTxJuCM5TB9NtrnTpPIb4SWmwPmJrTNruaYRtckCUS1EugKAqxsYitgV6RVmDSyeRDsagmN7IRofOeuTjb79Jyd2EDzwTv7kMUZwsuPWBqrEigVApNHPAZw8i+9vL1gF8tiwYMLMiAe29QPvbyzOYD8XlpRE59Fl2lp/SXzYNMx5iDA2SiZc8JO7CShMwv2M4kJiy7/9x+mmGr3Mp0SeiQcU9gF34QY+i9vBOzbHCb+OON/jiad3griXAb3qdWTE+n859a9ZtYQhHApIx1Bk0abZxF3muF24o2pqPkfgH/RGBofLXuwXTsdv54x0MdTfpzGt2zPh2LLRNImYPBfhP/DtjqtT/WgBJSAEihdAue1DIKHnzADnJfg/+Q6yoWav6F+3w+TO6ttVgSC3juAPAh2hjSpNMfEsQz671TBc2QfjHIsI3YeIkX5MMthvMfRho+yTlG8ZQyTfGJAJMB4N0K59KKf1DpfLm1xm+4QfDNaYD7Ninwl16l4p1S8j0CAz2BiP6WdfAxmZEcEpAaRubt2Ki0NnF8Cl1W/DwTqmcgsmuT3kg4Q6QPBlXi+6fucz8zxlwCkHCa4r00mr+ayin+wCTyXZpFvOsYUbtOwr2DqmBMCrewPPXmY5WB5iviCgIxAqO++Xy9WBnvtwOe5A1kHQmhXTyMUm91+qv+LioA7qjXxVVOYuzmnGE0hW2843qWIhY+BO4P3O4UvtU5PIPzYC5DAZNaNl9J43hj2pbToocMRdFKM79LEoM5KQAlkRUArWla4cuxZ2CxC6nCcsZ+eyXHkOYrO7tJ28EvGpmWFEPy0uimBkiEgwgd6ZxIHwXa3MvyV7ABPLkY0tKe/vxy72p8eCDo/hsivGHO6t4fphQt/oYome5K1KXO4KIdRfPBL3c7Zt4uikblZx52rAJMr3mZUM2j8dG8gZttrPz/qVuoEli57nXWXk1ZIPxFscBii0eswdkbXvsV3YfNGcKSecv4iA9wrWfceRVnbmxn4Te3lnQWfAvJ3AOm/IiA4Gb2DU2A3/6CL1PjZWyMQ5+K/DE2folmMAMZxoin5wkT6CNSHElACSqD4CdjnzPKYnUSvSCPsK4jjbnDGHao6TyDx1mPozzDGb2NpiH34D2G/xNX5FDsWg3272AntwsCb0STXBq8jKA/DvtWb3IePbfADOqZaEADE2QCxYB/4qLw5mbj98tnGjF9orF6IIBahGJV9rtyt4imWpyk0ycuUQDjWGwFEJyLxs2vFmBGVKUGgYfizMHIBz5+m8dEyhI6XQeQoHnO7aCamGu48Ptcy5rxpMQhv8QCjvx+Q1ONtkYFQ1X0J2L4jII8wgz7PoWYrmPh3sbKyH+y4xcD+POO3GcbqCAz+hEkHZDCvZr2r6XIC06q/Yt80DjAPMe0UdV0Gw3iXI1Z2MJ+/U88JMgLVPgRc18PrHz6GuFzIevGuj09AZEcyvxIxzz4DCFQpASWQVwJs2IzJawqFjtyYYm9IDsaOLccUGlPK9KPNx9HtAJpupFnsc58bjVEJlBaBjTd+kQL/hcZfC+zfMHr6AyY27sFj1+i28M/5LD6FA/XN0yZosADwJsL1eXhHCmUnF018BIzZIYUPa20/bzkTvWcyHXtZAGMn12Lm9xmk/ANMatwyA3/qpVQJXH9kGzzHLoS8nzYLtvbCHItQ6DrYBei0AXLgoXbWpijDZEBOSrQeSKOMmQ8E/4DEgkQav37O9lP5oci/APMMwEk9+KpyTuyfhpC5qks2AdiNBk78MkB+AKCcxl8b+SMmV73s70ldlYASUAIlTiA6e1M4gdHMxUY0qfQyjtHuRiTwRioPpWtvBLWNk1Db9HeMn7lXl+YjvPIVOLiDE7Sp54IMqlFmCvfG9jNN/SHe/uSSagEuzj7/Hnza6xX6yV4H4nbDIMf4KYMOQSyeKm0bKH8maLbkGGrDRALGxJnPD7E8tDhxXYz/RkscX/R6FJDzWF+XIamSIMT5PpzYFbAbvZP66eaWFz7fh/X9ItQ1PYJxxfq8xjF0fcXzbJvP5t3wf4vTyPYQ+TXveW6/QGdkD7St2Jrp51e727RCAtNZvz7Pb0LFFPuqfqeu8T6Mb+nafqeYMKwpy6Ihr8Az//HpD8N0+x6Cfb6NSKPdmPUTBg/TWD0H4eC/WQ9S96XWV0809mcTamfdgbrGywq+8athxMeAcynbqtm8FSnulWwNE5+I+Kium/OkMN1O2zmZtz56gnWCfYOdp/XLodkPYs6lj01pVHeGgB1X1c26HbWN16B2jm2nOhObhu2GBBxWymDB82XkKngeF5llJJDGeEFOoCIzFQw/zAHdkUgXZ8Ld/BCFUeTPhaszZ/QtTPI+qdoHFINr6SNEUwCdpiwk7pv1gx8iFGvMWEDHuRhfh7XhfYzgTmSk1JMSKDECid9s9G6n1JksMjn0tzviMpuTFofAde01rXKuBRc+2oeTIrfDMb9j7PbNAh58dZyu1yD8yWs8Zq+feqGcD3Q/gUjAJ/BSBOQ+5tvz8ZN/p8uqn6es6dq63og65+dfGE2hoATsb6wZ72pAMimTtr7+AIHYLZjQbD+ji/woTiiN/9+2EIf9phnDNNIvdFv5PYxDfUX2n+5lAutpd9RCxHENgCU0abThxI38DEEzGe5LeRqDkcmYeSGE8CcI7CJXeibG2EmoG9IIr85KQAkogdImYD+z6nmHA96hzEjyMVj7b4i2wAlP79gb3oy5mPXEFvtFmPMgJozyXpyc7kJh7aY749zIvum5lKkK+nE8PoZ9JPvLlL7y5+A4WzHyUTQpyod5Dv0G3Az780j0lLWOhz7kuPqT1OHMRgjGN+f430nuJ0+2tm4YszNgNkikIGI3KbyNAf1TLKwnfBX+n70PDVXTOQ68jsLY5zMe1tWGZcn8DJHgT+CaruW6rihdfs0xYWjZ7qzvPweY97c/zs3YNx/5sBvPLx3+FMW0Czip3+IUtiAig2B/EiCXcog3hHFW5TLK1HEFXqebNTwk1Sapbalaru53jIRRHuzafqdYmV2/A9vYwG2sm1/4iMhnaHMUHHyPfuyXaViNEUXc3AB3eAbPnQzVs7Qg1OsEwDkanizGsg0K33+FK17gutP1vA2p+n2BkX0Rd07F+MZB9Ke6owTsJoBQ5GEIahnFlzTJtdg5WBlMx/RzJPSkOiUBQSR8AuvbsfSxDCs9bZMIQvXaBBwYE1vbqgBXYt7E1OpZqK9Mb6YOy3z3v7vfJ6ivaEEm8RozhznPZCKd3nKut8Og4C/gTucDUc7j7liE5zxShtASTkjA742MjsWdaaj6isaM7l191RzYSf9M452y/6uZxcvyaCT9G5Y2XTVKoBQJ7GkXlI2dpEk9KFszX4JeENyB6MHnoq5lc/YfsqZzp87dGQO5ODkKZb0fZDwn0WSiOblkHocTfwDu6EgmAdbzE14+PDHQX89hLYvZmFxhv5iwlmVBLhzvN0w3TuOjvdFwmzbz8aBO3YFAQ/XNzIbdxJPh2EUOhmfuxoTGg3P6UHvc9ECiPZgw+2gE2h6jTAfSZKKjnDn5HaZV/zsTzxn7ScRnfgPDSZn0gey465eILL4NNU3V+E1Lr/RBMvIhCca1zYdjSOt/GcK++c9DWv0p4EzGWx+/l9anelACSkAJlCoBw4WwtrJKjiknADIAqZSYd+nntwj+q/u1iWM5TvPMpRApA2QmXn63699A/bxsAYxzA/ti+yY8kirBIWj7cv+cjvmTJrSOpR1bwIxmX77tOi6rL9+D55yBmt2/Wm2R9THc9ysu+DzrE24ATGAfbPq9QFI/+bMcyHuyNyB9YJVgEQTPwd21Y8866Golv2N5uR+QVPI6AOoQ4xjJbpLkRY/Q7nMD4HinAbIVx+MP4p7jMhy/ozDKbgLYZEgTy+JU1pMPu1QIg3KIjOiSt4aDseVsZ5IvCBq6eN5TXZr3fCa2Vr+DwvQ7+cxfZ+J+88P58PBnRpG8XgoE8H4JI5fwjOf06WAOorEZPFO9FgGO8SY22j6MYzx8Sl6zkdhksZanrr9wxUNsoP0ZgFvYRy1PKoAgBJETEXR+gevmc3yW1JdaZkLAbjSNrXyArO2mi68yCaJ+OkLACGqb92A/fSkMvoSHpm65abojaDTMWgScta70okAEpC8r6ulo3XQEEg+7BRJjdbJ213n/fkey4+MDCtoHN6vd9JiUgFoqgZIkYD/XGO/1N0DuZRtkkJnamN7q+QB0I+qaj+v054Xsp8Hqmo9GNHQlPO9PgBwIINO+aT4CgSswecSbDJO9do0D8WrZyolPYI9srvZx71qniPMEYPy/dmDfxIjKCfkRTLhgajblvd+qy03t4xvD/mRDrjNmTC+0Lt+i6/OTg09zrZCLAHkSmSrBvizP98KRqahtOQhuZ74+xIcNd863sMOmpwDejTDxuynGtjSZaFuv/olQ+SWZeM7aTyh6Oev19IzC2cUXMT+i/7/ii/jFiZ86sW1DRoFTeKprGoYApjDO2+ijiia9NqYNMHchHHgEdtd8+hDqQwkoASVQmgRqm/bj+IuLSrBveKfIg1kEIw3YovxfcN3kE/IpQha9tf3KXsA7mXKyfzCfw3j/LUi7b9/YhvMw5bD95Uoek2n7Mw0/wSVPtn+OPpmPfNjt+K1h7EN/CpFAkujtm9OTMK1iXhK3zK3aF9TvZ4BWmmQ6yH55BD6O9EvmmBc7uzkmFtwTcKoZf5AGnMx9h2Xk+cR5KfwL/edjeM7lMGYuxfVo1tciWzNf47Fx2570J+t76GY2dlzZtvx4QI4FwPspTwFiUOzKfjFwZeAOGOcK3qflXSZuot7LfgjE9oFll8+EI/EABKGkSQi+QjDGe5XUtbQs1+134D1RkH6nWKnZZy+D3wLmndQiymBAVn+qfAX9/h39sAiq1iZQN3szxKUWBpuR0UsISyZfHF07jnxdXbnHcs4/cEFa7qR88RTJ9IfxLsSnn/4I+W5/0M3VZYcsZl9/LVnfCmNS8e7mEPKcvdrGTdiHjWUqW0DkVZSXpf6yFz2p7rkEnJ6b9SLLuTg7cbJ4Mr69caaT53nKACf0oyH7OcJJTGALGtXpCagPJVC6BC4buhiITQHEvqWKDFVv+jsKDq6FxP+CusbxGNe0C+zmITqk1b+eMZCLj4ehrukSBMv/Chg+bOEUiHyLYTOcBDKcrDNjsev+TzJcxyZQorOHczBqJ9iYbApt8DQWljelcO1666GVfNjEnWkSLiPTw+A+tUkafx1wNttwvupKxm/vW9caKb8ML8zKfR8psgcc55Yuz5PEroU7z9alDtyHVUGurvgK8CbxKouBvgyA2Prm3YxI6E+onXUy6p7YnHGk13aT4oS527DO/xh1LdcgEpsOEfvTTEcBUo5MlUEzAuZSTNrni0yDZOXPHdWKuOMyzH9oMtNiJylwITy5DdGmqzGh6fsY25ThpL8R8tg8waW2yZalPwIyBoDdMMVDBlqcdVS03QAAEABJREFUxxANXQ39jGQGsNSLElACpUmAbeW4mRUQXEMzHOB/JFEGKzk+m4Jw9E6cPjSaxEfpWrnGQXhpBRw5kZnow3w2Y+mmhZscbxj2GcT5DeV4kvIk02FO2n4PkcDBGY/zk8WSjd0FzVvBeJczyPpjE2MWw8iNiOI+cECKzqqoMN9mpk80eyNmDkzinh+r8Y/3hydnwJhtvklA7kfI+/ib63yfGQdtK6TDqbiuh/LIsyzjU2Dwfop4HNb+4YibCaibu1EKP93HOjp7KIsrF8SMrfOzUR7wWWQssmxfXbESUXMb5b+U97PrNgEAW8HzjsfSJ9t/CiNfWJxwX5ZFLlQmScCYWVjU6r/xPkmworMqtn6n6ACtEqj8iXcBuQtgD8N//tq8RF9Pwh1V+C8q+wvata7ntfSCePbnNQ8CTIzji8eAitRfGcpWOmMchJyO90+gcocuRDTsst4/wCtDk0TLJuyHpyHSfFgSR7XKhsC06q9gvCsgHLsBsWyCqt80BNwZ5ZDAcaxrtpyyLHv/KcgXxdKIqc7FQcApDjFUClZYB0aGwQRuhR2gFQrJ+BcHwpjrKcvuEP4VSo6SSleFVQIlTqB+5AeILDuadd9vAmz9TBpwgdlwUszhQj5mc/HqNdQ1/5HmXIxv/DEX+Q/ChOYDUdf0M9Q213Fx7Caax9An+C4fDDhxhxrAsYOV9Sf41k9tHRs5GaGqh2G/YrCOS8aXJn4WW7mQv3/zeyTekPL31WWuifwGZsFggU+aQvddEY0O8/HTQSf7OVLZnYHtxH0XG/NdmETaTD6neiBE9mGMXZwfjESktbP3yLAetMB442HwNvOQqQ7So91M8QOIcyMQfgm1jfNorkTNrDNQN+tYTGg8gPX2u6y/v6SpR13zn7HjZv+DiT4HyB+Ypq0/+wMYiGyUwbtwxMWlVS+QOx9UkB81bf+3KONF5DI78wTEctmLYc6CkbsRwvtk8TBqm35HNnU0v0Dt7CN4fQTbttPIyEVd481k8zjgPQ/LRfALHr8DwMbFQwbamCcQKjselw//MAPf6kUJKAEl0DEC0T4CD07HAncylOs6qJk9AsHAbwGxfZ8gqZIVbH9drAz8jhPrrUm9ZGtpMAiC/tkGy4v/1jlbUpZawPk2448zr3cU9tO4YjCl4m04HBMDL1Km9bVgU44VxiMW2m59xxzbuM8ORJnXQEZ2fLFu5FGI3AsT/D0ur1q6rmOHrm08xlzBsKne4uwLYybBfoWKnr7ReTizn8MPltlFiaMgCMAqY95AOHQX60LXTZqLDIYj5Tb5DhuXi2LBCjs2qoFJOeEfYj6/C4lOhc17hxNLElASmzrbGSZx7lIru3HWxK+HCOu+fAVBE17+cHGXytDZxGw9+bz31SwXVzKqLiqLJsy0jkVZ6FAe86dNZGdGvgPNOtosY7t3J64/sm0dh8wuxQzkvS7tfsfWy9qmI1DX9ASfdz7n8TM++zyJmuYf4YAZmT/nZEas8L5c14PjPMA2/500wrTByEz0ir2Vxl/+ned/EYZtsyGh/CeWJgW7jtFbDobBWYAZQI6fQQKPwZXkX4JBB5ThXJSRzDfYp0riiv0+IbdxgPwPqVTi5SRzOWqb90zlJX/2sjnsZor8JdC1MU8d8Tna+o1nubibCRuaUtVs95whFD5MU1htjCAWHsF69itABrHefQkv8Ih+2QWqUhBwUtirdSEICITJjkCUA6zxM7bv0p8DsG/uunN3hLP43xAZCuEfhVGdAQH1ogS6A4ErD1uOcPBkZuW/NDGaLLSxA6B+DLAdYH5Ocy0C8hcu8j/OAckTtL8dYqYg8UasHMzjAAC9aUL0Kzxmpg2HNTAcrONUhB6/B650/GFi4szdAKmCv3oPMfm3v5cCuBrHPmw2+aacmPgyB6BTn3hHkSkZDJG9mKfOTUoWT64GwjEjWUcyrwPJZLf1oGHEo6wdF9H5I5rMtUDomTyFk1SyD0Qu4MTH7wCHE+wyg/X2IbrfSMPFCnMSj3vR9KfpRb9ZTqwm6ivLrvdLTKn8L8MbxpNHLQYNI16AJ/ah6JmsEhJh3oxtowaSxZEQnAGRKTR/gMQfgeAR3rebYWQSIKcBOJBmQ5peNJmPrQ2igPk3vAHHwR26gmFVKwEloATyR6AsNoDt16D8JZAi5nHzBiBy4E/Yfv6ZPvakDMJjMv0xvPgliPb9Lexbp8l8dMTOYYpIGBRUXTBjMALxKZRhJNt+h8eX2ZfM5rHAWgymjHwHQfM9GDTSsG9aSyR7v/aAh+sTX7uBsddrecjJhduyAaIrLoDIUYzP8uFhtRbbR/4ZoeiFmDrsi9W2OTlGVs6BwW28F6kW+naFlLurfo+8Pe85SXhVJK7rYGzTZtio7QrAOZe2fD7if+BzjsnOgMuFisRljv6JZ396Q3xiCyLUx8/dJ+gaTi7HfeHqe2ljx5Gp2HK8JT/HkJW/Rc3jdhxF77nQYstP5/PQKVGMwJ3zLZjoZAj2XhXVu3Diz2Q9Qd8W7cvyOYhxrM5TkHEGed112m6IX1l2Awz+wERtfeQhz1rE9lnXombWMXlZDDtzRl9I4ETmaYN1chIDnPvhBZ9cxz7zS0+Enq3hoYC6o/1OzawhGNL6G5az+yg9n3NkMI9DYGQknw+nozI4I7Ewmo+f52NCBdNt3psQ+QfTj9Ck0PI+/TwEd9SyFB660HoTwDMO24cClzW2d9HmvWFiLjO/FQ21/BefBd/lSebafgUH4vNVGOZVTDDzCH18BiPvweBq3ssPwUKNpEp2odO1mNjCOSj21Un9dMDSOAEI/1IHdTAgLKmds3RxvMEwYucpANs0GQkBbbafRJcp+/ML4lwCmIdgsO44s8vE6HRCAul0HJ2OgPXt4tm7wvPs+GJbMqVM5gmUfZjNC0GdlkIjKC0CXVvhS4tN4aS1g6pA8C7s8K2Dcc78srwLYj9B3Bb6PqKxByCyX97T62YJaHaUQLch4O7/LkzgIj5A2EU/n4eeAuTYcJgoeIkpn4vFS+6E63Z88d9+xjzufJ/tnX2QZZRJtU3xPvQuT/VGUNJAXWJZbz/XKjOZVpq3n+QwRMJb01/30Z6MQDSx67Y75KkcRvbGRS1DcpKZqVWcoJGzYH+2Apymz0mkOYuE7Un83/Cc73FR/tGcxZpJRNMq/4c4fkmvdhIvxmORaFkBwX2Iyq8xbfevikQoFUMJKIFuTSDQl9nrQ9M12j5j1jaNQqD1SohzM+wbsMlSNiZO6zmImwtQFr8OdpKQFjnTdtLTIPWEsUEY8Wj+nrldThrXtOyEXsHr2Uf/hPkSmhiM3INp1V/xvDj0JdUfAHImJ7sfALCcZl19MIcXN2Li7D3hchJ+XdfOXNuF0og3ATB2AdxuKv4mNmM+pP2VCJWfC3dU7sfldhO0Mb8HEpt+U01KnwSJTsGEOdvDdR3kStnPt7aOOgBB81s+f53BaFfPz30Og6l4/SM73qd1DrUE7EZOWwZTRboBYm3tCwWpfGRq7wrLee/JvH8co4JjwVQB5SdwyidgUuOW9OsnW6oI1rY3xrZzPvFIGdp82oS1Y+vYldu0DefXJgHyAxrb/sR4T2cj0OctZKsk1hvihNcItgHLi90QsIZVF5xeNXQhPHM582HnCmJdkKJNYjD7jtvRyzs7pxuQ7E8SDgydB887yiayhrFzDPMA+T3K/tXx9lmkNznZ+44UKoxAaM17msJbB6070++4L4XhBM5lGbObnMvXk0AgsC9SiPcnPN9S2aUvriHP6rLKZRD5F1NJ/hUAO0NkTDNiZc/TT+G1s5SL/2B7naJP5p0CTC8EB67uW3Ivs32ZsLb5AEZ8I9nZlwWYpFkOx9yW9Zc0nTJbJ/owrhRaWB7jnf8CgI3dHRVDzDyMuPdbwEv+VRaB0GsF24kpiB+yD89zpONbMKIymhTa9MNKhFI4Zm/tOX2Yk9Xx2c0HG6G13D4TZB9XZ0IkvjgVuoSyzAbEtrUoLdXP1iPWN9hjctGN1xtbJ/wld++sra1vE5urYLxbILJve3RmOWBuhjs60n6t/5XA+gRSF9r1/apNVxEQCAwrsnjXov8n52Fq46C8JV3XsjlirePZfP2GDcZOeUun+0asOVMC3YtAw/Bn4eECGNMAgzQLzF2VdRODmHsp11kIxe7v8Kf4Vou74+bfgshItnm9V1slOXLizfwX7j4rk7gV1krE8IHKTggmfzD9WjqzIxxhPr+2KP0TwV7MEycHSz8riRwYfAdlyF3fW1/5IALmdMZ9B80KmmLQn7GuXYNQ6CxMrSjM72gmNgE4p8EIH/CxpAigtLKNnQbIOFxe9TpUKQEloAS6gkA8vh2Tyd9zJSNPaNcEUdMyks+YXPiHnaA6hfZ2woyH9fRHtJkGB6diWtVf4Y7KzWf/GekqLYDYzzv7jPnMhhxXbYp8KLvAGz3kKEj8BvZBx0H4Z9MxWMAJ5X/b0yIyBvUVLyOO8YBcC2AZzTe6XfYj4cWvQ7TpuzjnEZ/J62+C+Z7Zt0fHNx/KhdKbSeYM+u1Ps1p7gDwDCYzFCmca3Dx+Kaes6k2I1wBgNk0y3YeWP4OJ3YDIIUfAyk2LDmu7GXl8876IBq/iQttv0f7Vg/Cq+OwzyLUIO3/K+k3xVRGkPCReLvF2oXvqeUAjmyDq5GaBhQmhYa+FrN+X8XQWTZwmmbZ8T0FU6jFh7vbJPGRsd9M8u8ixI/3bIw9JtDEbwZENkrh03sqdEcSExsMRkRsY2YkcA69qe2Q5x37/yb4cG84Nmi1hvDXb7kF8Lt4Ihm5MpEv1tKp3EAhfyjT/Q9M1WmQABBMA73pMaDoMlnFnUq6dtyl6h2oBMw4iZWtFZfA+2U7F5+H/oeMvHAjb9zT9DsufY/LY7xx4FBzyMpJ9vxNbyrGCOW09NmuB4oUB2xIzFt/eeFtedQ8tYhCMPMPM2K8txnhcWwu40Ob8FZcNTb5gvLbv/F+ZaB84sg3vVSBFYkL7zdHLWbNvpVWOtP1SRDR0NgQ3MsahNDY9e3garf3/x5PsdNDbnu3aJikDCZhPJ3cvLdqfNynrdSvHZ3fAbu5AUmU38hzGOn0Damd9t9PtD9hui8O6Q2pJk6Ol/fmBUGtqDvSSsU6MV7ytybXvGmE2Yvu3Zp+yhlM+T1m/Xnv/ObaxbM/NC/lMKT9xR8ohxo5RbJlIlcTmiKwYkMqxU/b2S1mR4OnkdyMMhjGu9vpm8BwiA+ZAlRLwIZB64O8TSJ26gIAkOoNvM6UJWCpPYsLsH+Z0Z6V96Kxt/CngPcrG40KmsxWN0KjOioB6VgLdjQAHZdOq38bKwOVA8Fjm7mWawmmDNwA5EV7bGZha2QR31PoPYshSeWYfwOybJtQzbBs5KCWPNB4L4rxb5WscxNu3Ezyf9AOcpB3j416KTn1h5IhSFDypzAI+jHm75bR/n1z1DEJ9fg1jJ8qksJMTBrPh4EcIxSbDfmEEBVTTKt5EzEyE8Y0pIlIAABAASURBVC5g/f+kYJIY8zZl+B5WOleivvK9gsmhCSsBJdCzCIyxi2LmOGa6nCaFll5c/D0U4xu3RcJ/Cm/JrO0b4eNnb43axnGItszlwsNf2Q+dRq/bsc1N9owZ4STaX+l+DMK9GjC5iuPNPIy5zmsph4kfyidrThozteS6PxxzKuwbocnds7e1C1R1LRWIBO9m/rm4LaMYyTcyOJiDtkCajZwM0dXaLn7Y54BQZCpEjmXyb9KsqZkHqQDkVvTrfzXsZCQ6ojgBPqF5OzzX8jsEzO2MwY7t1i6bnvyN49iTuBB3b05/EoKJradd8RCqnoeQcwrdZtCsrwW9ADkYYm7D880PYWLLSGRbTy58vg/qmo/Gjpv9g/m+n/GNgWAnAOTK/wZfsSxeiHDsWrgVX9Imt3rAQi7UOXaRxkkZsZhNEJQRKd2zdbBl6rWPX4JBA+vC2z7B+5HFaJjY/ahp/B5sm+LjOaXT+8s3AbyhTC+U0g+wFQx+gpPfWbvM+QRI6+ROD6Nu5ohEnTdiy/ThDPNN/MZjXTJzaZeddp/sA7F1DkPWCDgEgn1w/uxv4l/DMb+nbKcn7/sKQmXnAOb5/Ka1Vux2AfP7MPInREIPoGZ2JRILW2v58b+wZcqGk9Y/0iMXLaUPj99oYxaw/v0cW5f/K+s3l7+JBbD9jiOH8B611+s13VafG9MfATk1f/2OczOTOpDmGxkcZNbvxGK2fnCBkKH9tIiNm/1BsAq2z/PzW0pu7gGLYfA3lu/1vzhjzAyEw41Fk53y8DaUc09feQR7ANHvZt1f+UVqNzfWNP0MjjxAb3ZDkF274Cm14UjSwQNZf83JliETOgni9IavMt9H7ZyNfb1k4+gOXYhwuc3DP32CsazLfpTtFkSDkzhOHuTj19/pgv9uCWNG+nuSgYib7+Zkjui5uYMB24dIX3yjtoYT3y3rNvSb8B0/u2d0HGWPNwPO6TBY0PGIChAyYrbkvfOfRxb5DkR+ADsmyJWIdvNmbctPEPEeYNwcS2HnNaL2AHk46/oGVT2NQOqBf08jUbT55aBUsDtM/G7suOndmPjULrjw0T4d6rztA6r9PejxT+7Lh85HIM5tzPYuEPBhlmeqsyegIZRAdyVwdcVKNAx/DF6Ig9PEGzGfcoAW76LsRpnOpzRXsp06EPVVf8PUg78AZ4rRWeXO6c8HlcMAGYhUypg25rUFl1W/n8pLwe1HSxyOYx9MI76yiOyO2qZDfP2UmqPBD+G63WX8EoRINXb51oDc3QZOzLl7LUJD1dUIOXtC8BfGvYyGDwf8n19tGH0rzet8ODqdMlRwUacRbjH8RiKlurxqKRpG3IqQ7MM6/ltOmNgNEjG65FvbND6AERcehlKGJ/K+mJHvHGn8SkAJlAaB6dMDsIu0g1uvgsj3fYUWhCH4DRcl3sLg1qWonfUO6mb9leOIGi5aHovxLXthYtMuCVPXNBS1zSdwoW486hofQ7R5IQLxtyEyjW3r3kxnE5ogTbtuf7OK/YNZCHh30d/OmFL9Y47x5iHnb3Ybgf2Esf1d9V7xOwE5AH5KJMDx5hj0Dj3LvPwRE5pPQ23TqVmZmpbTUdd8E8M8SDMfkeAK5rMZIscw6Y1oHJrVegkgj2NAG48oTmX77SmV/0G8/36A+Q2F/IJm9XOAzctGEJyBSPxj5vcG1DbvCruBwv7sg530tAttdqxmjX3xwH4twM5hjG8cRL97ktXdHCe8xKH9GMa75luwEV6/AYPjUVZxIqaOeKVTC3GMLGPtisdF9zfxeeVhMN4VMGYx5TBrhzcOgCG0OxxxbwYGr3wPdU23o27WsbA/81D7+Maw9c2amrkbYvz/toX9GYyaWbWoa2xG2VKWf/MAwx9JsznN6jrCfJsXEEA1Jlf/OefjJruwctGcHeDFLuH9tPWTSafUVqZxmDDrR3D/2RtuDsbcdsK/ofJJMr2eqdoxEQ9JdZi2nJ+SfyDaPJN18dAEz/Naeq1aCBG6J9f2zf9x87YEghMAZyTLpyT3SFtBiO7nYPOP3ub9uQ21zecg2zpf23IWapquR23jPTRvIboZx5SBmRD5EVPYGMI/nqzSbAGd6SzPn6+6Tn+w3G2diYaOhXinYO031csh5jT08g7I2T1CNso+a+z7NucJfgLB8+vXk2ziSuKXtFhOl0PkSx7tIuxy+rLlJsCjbXu+C8ebheeaX0Ntcx3GsV+ybcu4xwbAznfajTbW2HPbLtXO2pT17weINDdx0ct+ieJwCHoxrtU6AiMzWT9Gsl96EqcPtXMRq90yPJq1+x2DUb4Bhf0OMAal3+8MgmcOZl7XXFzkZSlrlu9weTMMmtbJBccwcn3uxyzrpJLu0rYNtk20/a5nbqT3rWj89EawP2kxeOVstlXTkHVblxgPncewf6J5nHVpAaKhlXBwOxILy+iHNZXgXXjBJ5CpsuMFW0+joVr2EaezzkuaoLsA0QbYem37tjSeM3K2mwBEzqVf+2U+w2MqzfGK1MGReZjQdDoumDcYtq2xay1gG5AqFCCJBeEJzduhvOzPvGJfBR9lHKYxFjtsfhLs/KVl5OM7qZMde7nPcs4zfirsWFQ4wsDXaiDtzsOLs3dnm1lO7umYfx0wJyeu66G+4inGdTLT/hDs0FCsyuUYKLHhZc7OEO8PgGwDf7UBna9FZLO5qJt1RYfqW13zubBjy9qm//D4Ifp/upJp3wVBNeO2m+G+uV8G75Phf2mvWgn4EnDo+k3B4UVBtAcDlmQUWpliECIlhDAgx8Fr44Njn/sxpPUi2M/l1c7YNfFbafbTO2Ob+rHxDiY2B7iP9IfbvBHqmrdix7g7JrQczgfUOsSCDyEQsI3DoexY7f2Hqo4T0JBKoNsTmDrsC9SPYNsROYhtxk00zzDPX7HVNjzmUBs+1JsPGG8jxDQAMoyTwhdxYPgRcqqig5mHYYzSTmQmNyKvwZNi+ywrRV5Hb1E2g7zsmxfJ8wG029sJ2nWCJr00ZR482AmW9nCrwxfdUTZC24E7Jc3DepYxj1ZLaIo5TzshFh1AGXOv7Vv3wcqTWKeOAuRvNK8C4OQF/+dWr2B0b9A8wgeQX8ELVaKh2r55Qqsi1G7VAjRUnQ3jVHHS6nesR09Tyq9ocqlt2fuMcc+FES68BQ9CQ+UlmFad63SgSgkoASWQIDBu3gDUNO/PZ7+DEqam6Ug8v9kZiMb/BZGz6CfzZz9BGcTZGnCOh4DjMnMvAp79OtJLHCu8xLj+x77lbjgyFfataGAQAMFqlVjAwQokJqbQCONNh51cNWZ31FefhCmV2f8O9uq41zpywtW+CVbTVJ3I84TGgzG+8YeILL4UYcxmmj+E8A9plUNfzK/8HMbczPNbsjKO93uOL8cwzPdptodIyCfFtxAjy1x81conkZw4Tdv9K7zx8VjE5QjGdzvv53M82rEiD9QiYQhYtrznEgtZkZV3Ibr5RLS1nILowSclzPabj0G//lMR7n0/AvImy82zDHk8TTmN1faZ4ktyn8myVYc2qWYfPR2ueCiEulmi2GTTiRDnZ7ynD1KET2msjDwAX/8T5lyEiwGgP+deLki+Cin/BFHvi4RxogsRaHuLvv4Lx6lH+0LJ6jxjlVrGfD8FAxde+EAkvoSBzil3xuBEXaiddRBsfahp+R4iwV8hHH0Y7YvTwQwS6A/P+SsiA+9E5NATGd9hNAetMnvDLj5lEMnaXsRgZfBWMv0djf9Y1LIF7Bjt34jEX0Sv+K3YcbPTOKd1NCbOrkRdy+arZGnPo/2qwvutExBo5aITF8aBsrXTTnm1KQQns0xex2N2dV68G+DgbFimItsCWPfe0mqVNviYc3D3r7pa/2AXj+pm7deep5aDmM/D0cb6E+7zZ3q+GZDBWE/JFrS6H5FBdyByyE/Y7h2OWoa19338k/vCLljQQ171m++/jjhqmcbrLMOGx1zpt2DkBwhGNsby2DYIBH8Ag+tZVx5lAvZ5I8IyZNtstrVmCoJ4CY58hECvF7gw+S+ULX0gYWLBf6NP8FWI8x4g90GwPwCHZrXm/ANe5MUliGA0po16k+cZaCOJN5C7a79j8Ar5vpsBiHYvIjthibP2InC7S+n+d4euYEm5lhwWf5MJ8zjCA2d+c90FZ3YRd0LTtzF+5oHt7cOswxL1vY93M9utuZRgD5r0WhCCyD404yDIsq1L+P8NRH5KcxCQ6PeQQtl+uwle9IOk7rZdmti876q8sK2zP5fS/Av0Cd/LOj6BYdasn7xMqgOU46esyg/C/vxAbdMRX8dXN3ME7IJ50mBpLEOPfcAxiG3PWPbZX6X2LhBsS3l/j/LWDxJtzeC2sahtPo597uGom7kFxjZthpqWke1ysV2uaTk6MTYy5glA7CIuMlAbAxxbRqN/R4SMxpOVbd8TpnnPtcLbzZd1jVWJ9MZzHGzH/zt86xRElv8DMJNpeq/lv/1iT8TjjYgGr8PFVvaZhybC2/gntuzGPiTY7i2P/5cumQX7goQxH4MFGoVWdqPF+JbtYTduWg4Tmg/lOPbHiIZuhMTmALI3MlISZBnhPXIu5DH7+gbDtgc/Y9hDmNzmNEKTRIsHx8xBmfNOEke1UgJrEXAgkqIgreUvvxcBfJ8L1VNo6tOaCU0/zFiYxJsJTRPTxlnXXA8JjIfwL+PIC+VR7CKBbQTqETD3AMHbEZPfIxC4DmG5CrHwNAxpm4Zo/6sQMdezof89xPkTjHcvRCbByEhASmN3Zl1LZmWirmV84k0UZKhqm36SUZmw5QLeAT6xqpMS6DkELjvwZU7GcZIjeCLbkfPZttzCzL8Cgw7skGfIdu3BmA8Z10MwcinPzwC84zGl6hLk7dPYzjImfQXNr1Iaz0xA+UfP0724tX07ISDnUcjUeQHoZvsK+kqnly9dAYnfSW8MY8MVsQkE2ihneh0JLOSzxA30WMR5il+BYGiNyQVKm0vt8sFgSvWTCEXGAOYXMDIWEI4fzAesc3F0WDFemFcBuYOmBh5ORSj4E9iFf7txCCWgGipfwhe9LkQsfhLlPw8GN1Hq52hW0HRUM6w8RdbXML5z4MR/zIX/8ZgyfH5HI9RwSkAJKIGMCARXfBsOF6XE4UIHjYP7GY7Pg7If2ySH5/nQbPOwCAYfAfIMzYM850QZXKZ5EYBfAsHjUFb9Ey7634KGEbmd5HPBJ/joCP6/HTbfRu5DwJnOvn8c5diS6RebjlGgZ2jepikNbd/enlb5P4QqxwByEoDxNLcC8hoAmx9AhBPy2Br2LTNwHO2YmwHcDmsc8zsIzoOIncPYgHZWG/5bynFIMwRXsqz8GrHQiewvr8KVlZ/RrbD63B3aUF/5IMcGv6Bs51CYy3h8AgZfsVxZ2WnVYc3FS1h2f2AMF8GRn2L+gsuRq7FTpGyfRF0QtgEG97NNeAAiZCw7MD2hyUwLeE/lBxCP99G5Hza+dnMpepuN0BFlvzAX6jsJkL8ik2dIAf9kM4hqrqDYAAAQAElEQVT8BMCN8Lz74MV/x7CHfS2PkfsBcx/dL4Zgex6Fpri0cNEnuDL1BP0GKzk/51zRnifmxXgPwsHtzM8PAYRpkmtBWcKPmDvo/wHeq/sScTjByxB1hiQPlENb2za0OjMAj/OpeD9nMYv5K+vEf+GOiuGaUYtw6fDH0FB1PsKh0fBip/L+j6e5ie0HF6q5dG8TFvQiC9vmV/HSvpF+MIxUArIJgBDNas36axbx4hHKPQnwTkUoenlW7Y7td0y8munxHtl6Jvch0I36nbc//ojl6PdkZDfR87CWjsFgAc0ac0BmAMqke20AsFkOfmsO7/GDPDWAsWXmL3B3te03rbpI77JtGTxzGgIBjmtY1uDcDzF/Jv8TKUFvmmLTiylbE77stSypYCuDGyHurWrrmB8DtlvgvKI5iKzXrKdIo4L0vy/9XE3zIIRxWYPA7YgY2w/QOkvtuh6i/R5lqOt4vxfymIkup6eDeU+mUJ6/Me8PwJPvsT4cCPH+0i5X/H72w39nnHaDw1Y8Zt5HiZTBbrK1mzYCskY+zUSsqWLLN4eRP8AycNgeOSBX7ybKZNup1OmJ9AHkNMTxV3gOw5Aj5O/wvHFYEeqFfKvrj2xDWfTvMOB9NF/lO7m08S+eXYag91OIaa9vxtwPA8514WSyLMI2zlsKI83AANs+QZUS8CPgcNDk+XnoEjeDI9kI1mZkjHw3Y5kEezDOC2gyifvcjOMtHo/9YXfxAUfwPp5AcxqNzS8X53AqBKMp6uE0e9KwYef/ktImk/tGP9658MyOmWdNfpBhmWDcUpE6XnVRAj2NgBhM2f9VNFT8CSti4+Ak2hhOwAjbT/kj25+nYMTunP+U58sBae9fDIeUwBe0ew/AizR2IGUnaWxdPBoI/BILy6dxUvhhmtxODDOxtbTLScX6qrvgZ6ZWPwR3dGStcMV6Mblytm9eEvmsvjsj8e0AvH5kS/r40vBLpJlnP5m+OXh51VJMqX68uPM08i7k43de173p9nO+9VVz0FB5PcxKLkwHjoYTOIb9Zw1E7qX352HwLvvHhTStWK0MVvL0U5q36D6X5nYYYZ3HKDhyHLzgeQhV3ICpVY1whyebKGLQItY3D43i8pGvUv4/IRwcC8/5MTzvSIiMAeR3AGbR2In6j9iGrd6oEYMBFwHMh4C8CiNP0FzF858y7KGQwIloW3ExJyunY8rId6BKCSgBJdAVBELh1+A4v2Q79P1V5lAeD8irgXMIHMPnTe/7EPMTRM2ZiGEc3lhQj/rq37Ed/Bcahn8KuxkNeVAux6YhZyYkfhLzafN9FI8H0uQ3357XwfjBSeFQPcqrVvcneYCSpyhd8TiGeAmhx3/HBbOLYMAxgHMw+0b7ltL1TLWJdh/zeinP19HG5vctWj5Ccyk8/AhwDmd5/RmC0YtZVu7A5cM/ZN9r6F48elr1V5TtHixZ4lLeU2DMEXBwDAScJ8BdzOvTNB9z3LQM6ytDHpbFB6wbM3l+Pf2ewuPhMBw/haIXor7yJkypeh12IXX98B2zMSvmsfzbukBjvsfzg2g6WF4T5XwUjHcE42B8rOeOqUMoYseFHZPP/kxVKFxDJtnXU+PZseepiHn//loez7N5zD6uDtfhBJPseMZjF8PlYnYqYuWxpTCB89rzFLecs2+7jTmsPTzvUcC5ABtv1jUbaeymjnD8XojUM3trLArzqkParIQj/0laJ+xzxtQDGtkOXYP4ynEIyGi2/UfBYCwgDwF4C8aQJW14kdDtG02+4Pk82t4IyM8A77uIRk/HG59chvoRT/neGyRRrhiUYRbT7p79jm2PYmU3MefnkudcHjkvYuKAaeG5nWc+mvWXi7a8atflfPZJvVGl3U/p/Xe3sc/DtzHfCwH5H/sAm390qXr57TaEA78H4j9M1O/2tngUz7Nrg7qqvQOOhAndD/t8nQxUrxjbJXM+5bftHI2x7Vbn8mK8Q76OD/ETEYvPT5Z0RnZX7rEcsXLe80S/0gG5mB8n+A8EQo+zjhzfLpc5isdc9FG2XyCzRBvvrpWf4BLOU8ROZjrfh/FsepZJ5vLbvtVwfOMxbmOOhoOpGFq5Yq008nXhjlqEZcGbGT3LOUeHPCmYHrB/G2LmdnjSXt88zglZNl1Vf7JNx3GOQJtMh9vFG5MKdoM04c4QcDoTWMMqgR5NQDOvBHosAT702t349jOVDVUzMJmLf6GK0xCOVSIc+Q5C0W9hYa9B+LysHJ+Xh7GwvBftNqH79jzuzUnh4xB+vJ4TXv/ggv8zqK/4KOVDQo9lrBlXAnkk0HDwp6y3z2JKxUNcnLkMr390AuvmUISjOyAU2xSf9+qPz1l3rVlYPoBu36LZie5VCFeewjp7A+orZyU+VWvfWHO5KJBHcbss6sTEYsVrmDpiJtncilDFOdij8kDmfTearbGw15BVXHpj4csbs03bBqHI7tjzo8MQrhiL+oo7GbYZ9m3/Kw9b3mVya0JKQAkoAUvAtmGTK//Hdmhml5n6ihZMrp6bGM/ZhczLqxbAbr6zCwhWpvwbg8TmzpEtXZZn20d02FQ1Ysqwd1DK/abrelwwWwT7FZ2pFTM5jrgDwcpfs788gOOELdlXbriqrwx/c2T/GYruxGeA7yP0+CWYWnUf+8wW2A2d7ii7yIKiVolNspXvYSrL+pSqfyD4+GXMy884DhhOsyU+77UB7JhpbVOGheUbIhTdFrtXHYRw5a/J6jaaGbDs7KQ3Vwhynu+pB3+R17owufp53v/O3TN3v09QX93UITltG2fbmQ7XQY7xujrstFHv+t5nuzmgYf9nOsQjWV4mVzwL+xUL30Rz6Gjr8OSKW1kH+mLPqh/B9kXZRm8337/+0SR83msjTK5s8g/OuYjLDlkMOxcxZeR/UF95Jd746BjWtZ2wsNeGWFhehtV1cWF5H9pvQrM/wpVnw47V7Yb3y0d9mHSTgX/Cq127f79z2dDFqK/8M9u3KrLsi8979cIeVSMSdvVV87hI9p/VMHjkoqlZyWP301OqZvLZeDP2b0cw7+91eQbtWMqteBP1I2flrH1I1mbkys6+bGDnBlKBsm1Fw4jctXXryp2o21V2410qCdLbJ8r+iKc6zNvObXamj1s3T8muJ1fYl6q+yYt71ApMPWB2h2VeNw3bto6W+DcJ5Pns+uFLEKqciM/L+8L+7FShFrTt2Hxa9dscozbmjOW6bHN5bV8GK4avZuW5eGj0uSGgGwByw1Fj6YEENMtKQAmsIiB8CLeDJTt5sdrYXb9rmtX29ph4kOHk4argelACSqCABGz9TdTJUTFO6LabNeuuPbf19mvDxX4bpoAid03Sq9o1+/C7Ou+Wxdfm9OjXvEaPjsO2gfmYyIcqJaAElIASUALFTGCd/vLrfnJoNLHBd/W17UsT443SfQb4+i7YTRCJvKwaO63OY7KjzXdiLMHxk44TvkaoJ92MgH02sOXflvWOZs3WKRtHtvXEpm3D2rpmw69rrH3CaB3M7tbYtp1t3GqeiXtLOxtJUNZ84/8TBIOLrHX3M8yvLTuJvHe/3GmOlEDRELBzKbatsW150QilgiiB7kNANwB0n3upOelaApqaElACSkAJKAEloASUgBJQAkpACSgBJdD9CWgOlYASUAJKwHXtOsLQBAhj4hA8jUG9kv0MSsKL/lMCSkAJKAEloAQKS8B23IWVQFNXAiVJQIVWAkpACSgBJaAElIASUAJKQAkoASWgBLo/Ac2hElACSkAJAIcOhMHIBAnBZzDSjKdeb01c6z8loASUgBJQAkqg6AjoBoCiuyUqUEkQUCGVgBJQAkpACSgBJaAElIASUAJKQAkoge5PQHOoBJSAElACQNQ7AiK7EYUHyFNA7CnoZ7uhSgkoASWgBJRAsRLQDQDFemdUrqImoMIpASWgBJSAElACSkAJKAEloASUgBJQAt2fgOZQCSgBJdDjCUxo3g4wp5JDLwg+4fldCM34iNeqlYASUAJKQAkogSIloBsAivTGqFhFTUCFUwJKQAkoASWgBJSAElACSkAJKAEloAS6PwHNoRJQAkqgZxOomTUEHs4AZCgX/lcAuAGhQQ/CdT2eq1YCSkAJKAEloASKlIBuACjSG6NiFTMBlU0JKAEloASUgBJQAkpACSgBJaAElIAS6P4ENIdKQAkogR5MYNy8AXDkQog5OUHBmN9gWfRGuLtGEtf6TwkoASWgBJSAEihaAroBoGhvjQpWtARUMCWgBJSAElACSqBjBNx5vTG++VDUNo7JyLgtO3UsIQ2lBJSAElACSkAJKIEcENAolIASUAI9lYDbsj2CrXcz+2fRROB5YxFzrsA1oxbxWrUSUAJKQAkoASVQ5AQciASLXMa1xTPe5pjQeHBGxpidGTj3+TNmKeNtonlcDTJnYPA/GKwks1zrMgj2yKhM2LIj2KgzAmhYJaAElIASUAJKoIMEom1DEfD+BpGb0ho4p2Ax3utgShpMCSgBJaAElIASUAKdJqARKAEloAR6FAHXdWDf+q9r+jUi3hzm/TAYmQ+D/0NZ9U24vMrOidNatRJQAkpACSgBJVDsBBwYEyt2IdeST8QOPB6DkfQGcj5E+qwVPhcXIgsgoZ+ivuoQNVkwEMP7gS9zcQvWiWMDlodJNOnLhC03MCPQcaUhlYASUAJKQAkogY4QsBNJMJcAMhB+ysDQ+XkgfgaursjHxkFGr1oJKAEloASUgBJQAmkJqAcloASUQM8hUDtnY0QOqUFg5TPM9MWAeYfrBpci3GcUGqpmwBUPqpSAElACSkAJKIGSIeCUjKTFJehAxCOh4hKpBKQx0otS9qMpYa2iKwEloASUgBJQAlkTcF0HgdZxDHcAjb8WvAvPcxGOvwBVSkAJKAEloASUgBIoGAFNWAkoASXQkwjE9oZ4P4M4s+ChBuHQD9BQfQncvfST/z2pGGhelYASUAJKoNsQ0A0AHbmVBhsgKLoBIFt2jiljkL40patVciWgBJSAElACSiB7Am0HHQnBuRkEbIOHemy66b/gjiqtr1RlkDn1ogSUgBJQAkpACZQQARVVCSgBJdCTCBhvHiA/Qdw7H1Mrb4Y7/EOoUgJKQAkoASWgBEqWgG4A6MitE4TgxQd0JGiPDWPf/IvLhsx/SZc5yq9aCSgBJaAElIASyIaA+9QmcDCRQdL/LJNBPcqif8K5O7TRv2oloASUgBJQAkpACRSMgCasBJSAEuhRBKaO+Bz1VfMwrforQAxUKQEloASUgBJQAiVNQBdjO3r7RDbqaNCeGW6kAyexAaCUs6+yKwEloASUgBJQAtkQOGd+GSJtZwOyO/yUQRSQ29BQNVnf/IcqJaAElIASUAJKoPAEVAIloASUgBJQAkpACSgBJaAElEDJEtANAB2+dc7GHQ7aIwMOcSBmcGlnXaVXAkpACSgBJaAEMiZgv/4z4JNqiJzAMOU0qTQX//FPhCK1qTyovRJQAkpACSgBJaAEupaAE2G4ggAAEABJREFUpqYElIASUAJKQAkoASWgBJSAEihdAk7pil5gyY1sW2AJSiz5d4IUeAua0tUquRJQAkpACSgBJZAFgZH9YZxfMYDPmEk8GMxBwFyKSQd8Sr+qlYASUAJKQAkoASVQeAIqgRJQAkpACSgBJaAElIASUAJKoIQJ6AaAjt48Y0Z2NGiPDLeobxie2a+U866yKwEloASUgBJQAlkQiIZOhTGHMYTQJNfGWwDBWASqXoSI/s5kckpqqwSUgBJQAkpACXQxAU1OCSgBJaAElIASUAJKQAkoASVQygR0A0DH795+qJm7YceD97CQ/UKbc2J/xxLOtYquBJSAElACSkAJZErAbdkJBg0QhFIGMaYNMKeivmoOXPGgSgkoASWgBJSAElACxUFApVACSkAJKAEloASUgBJQAkpACZQ0Ad0A0NHbJ5zSlugPOhq8x4WL44elnWeVXgkoASWgBJSAEsiIQM2sIYh5N3OklHrxH2Yx47oADSMe5VG1ElACSkAJKAEloASKiICKogSUgBJQAkpACSgBJaAElIASKG0CugGgc/dPF7Uz5meOydhrMXpUmZSAElACSkAJKIH0BM55pAwSGAODoT6elwLyW2yyya1QpQSUgBJQAkpACSiBYiOg8igBJaAElIASUAJKQAkoASWgBEqcgG4A6MwNFNkPFzZv1JkoekTYSfO2BGQ3lLBS0ZWAElACSkAJKIFMCZgV9BmlSaat/XS0ybU4d4e2ZB7UTgkoASWgBJSAElAChSSgaSsBJaAElIASUAJKQAkoASWgBEqdgG4A6MwdNF5flJvhnYmiR4SNrjzY/zPARU9BBVQCSkAJKAEloAQyIXD9kW2Yv+A6SHAojJm7XhBj/odQeAKurPxsPTe1UAJKQAkoASWgBJRA4QmoBEpACSgBJaAElIASUAJKQAkogZIn4AASgKqOERAJwcNRcF8KdyyCHhDKnVEOkeNLO6cqvRJQAkpACSgBJZAxgXtGxzFl+HyEB43gGGAMwz0LgzbAvAMneCrc/T6hnWoloASUgBJQAkpACRQhARVJCSgBJaAElIASUAJKQAkoASVQ+gQcTsbGSz8bBcuBA0E1oov2LpgExZ5wLHQgy1hp8yl2xiqfElACSkAJKIFiJODuGsHkilvheSdC5Cp48jNM2f/VYhRVZVICSkAJKAEloASUQIKA/lMCSkAJKAEloASUgBJQAkpACXQDAk43yEOhs7A1J7UPh33TvdCSFFv6Y5v6wZhjYTCo2ETLRh71qwSUgBJQAkpACXSQgIjB1BGvIDTgEkytauxgLBpMCSgBJaAElIASUAJdQkATUQJKQAkoASWgBJSAElACSkAJdAcCugGg83exDMY7Cm2BbTsfVTeLIWz2ARz7+d9S/pmJbnZTNDtKQAkoASWgBApAwH4NoADJapJKQAkoASWgBJSAEsiCgHpVAkpACSgBJaAElIASUAJKQAl0CwK6ASAnt1H2QiAwEq5Rnqt5ntfSC8Y5CDDbr7YqzaNKrQSUgBJQAkpACSgBJaAElIASUAJKQAl0fwKaQyWgBJSAElACSkAJKAEloASUQPcgoAvWubmPAmNOA57sn5voukEsfWRrwBtd8jnRDCgBJaAElIASUAJKQAkoASWgBJSAElAC3Z+A5lAJKAEloASUgBJQAkpACSgBJdBNCOgGgFzdSIM9EQ2Pgap2AiZ2Kozs0H5Ruv9VciWgBJSAEsgDAXdGEDWzhmBC83YY27QZ3Hm985CKRqkElIASUAJKQAkoASWgBDImoB6VgBJQAkpACSgBJaAElIASUALdhYBuAMjVnRQI4NWitumQXEVZsvFMaPweIGNIhExQykplVwJKQAkogc4SuGleCHUt+7F/PA+1jdfQ3IdIaBYc53EY/BMhPIpoWyPdZ9HtftQ1/5bnp2Jiy24455Gyziav4ZWAElACSkAJKAEloASUQAYE1IsSUAJKQAkoASWgBJSAElACSqDbENANADm9lTKAi96/4ULH5jmNtpQic5/ZDMbcTpH70pS4VvGVQDcmMKHxYOTD1M2swMSmXTB2xrdwXksv2De9p08PwHVtf1N8m4Jq52ycFw45YzvzQIxv3LbkSqJrHJw5oy8X9Mfg/da5gDc30T+K/AoiP4Bgf+Zpd8B8h8ddedybdtUQOYbnZ0JwCzzvBfTvvwC1Tf9AbcsYXNi8EY5jWTJGGKb4tduyAetC9Xrla1LjlsxjceVhXOMe68m5dhk+oKDAxzb1SyMf27O523RaRrdps/TpdKDtHN98IOqahmLCzB3gshyPmRdKtI3FWp6tXBPTlgky7wCLtctVx+OoaRyGcx7Rn97qdKHXCJSAElACSuAbAnqmBJSAElACSkAJKAEloASUgBLoPgTsgkz3yU0x5MTg2zBeHcbNG1AM4nSpDL+eMRCRFVcAzqAuTTdfiWm8SqC7ErCL8kYeQz4MAs3w8BJCoQ/Q21uBaPADPL/pQ4geXI+Js4/DxDn7JDZJ2bfCi4JvbGReOOSMbeBeBAL7FAWqTIRwZ5SjrnkEoi03YlDoXYjcxGB70XRUbwDBURDvJpSZ+dhxswe5qH4m7GKqO2MwinUzgK1jkfhJrAtPrle+YnId3Id6dRRIXsIFpH49Odcswx4eyku6mUZaZnb0lc/KGo+elGl0Kf1FcGTadGxa2ZqAeYJp/g8m8Aai5lMMXrkSkdCz2GGz21mWz4f9etSEmd/B+MZBaN8sRe8F1Ft9u5xld0peWGTLLpV/Ry5Gr/662bSAxUSTVgJKQAl0OwKaISWgBJSAElACSkAJKAEloASUQDcioBsAcn0zBSEuVhyLUNsPYd/wynX8xRqf+1IYfUPHM++HAegW5Yr5UK0EuimBA7yuy5hsAsjhAMbDi/8NXuxBwLsR77fWoGbWMXBn0J2uqlMR+ByhYGMqx6Kyr2vZHJHgOC7K3waYMZRtQ5pc6v6M7LtclLyBx4cQC12Fi5uPQ80zQ3hdXHopBpLB8UjaH5qDERmwE91U91QCIgEIdqU5ESJXQvAITOBuBJwrET34TLaNlT1qDJltOTCmjfXrGfR+/JNsg6p/JaAElIASUAKpCKi9ElACSkAJKIH/Z+86AOMorvb39ookd2wwtuk99GZwk4xNL6GFYBJCJ5TQEsCALRlYgiWb3kL4Q0noBJPQe4mxVWzA9N6rTcfdkq7s+785SbbK7d6ddCed5B290e7OvJl58+3UN7NzPgI+Aj4CPgI+Aj4CPQkBLtRKoCdlKE/yMhiOno8167fIE3lyL0bkl22gOIMJZXvBh1H65CPgI5BVBN59QLIaX2aRmZ9IOZBBLoRYN3LB+N8orZqM7njMPTPRCTQb9q7fdUI67U/CnOYwueYQqHMXRCZDsHH7I0s75NpQOQoOroesuCfxUwNpB82IsX3MBeERLN8uJzdIb8A6s30R+6F6KAJB5mt7LmqfwOt0WNZtWLPubpRVHgp7pv+VO0FpQSKLWP/ZNtqduJmthQT+g4+Aj4CPgI+Aj4CPgI+Aj4CPgI9Az0PAz5GPgI+Aj4CPgI9Aj0LAoqI+3qNylC+ZEWwB0f+gdPbQfBEpZ3Jc8aZZzLgWiq1zloYfsY+Aj0D2EHhgQj60+0H2P8MgshsENgKoxuTK61aLNhMZGJH7MuDufFZbLXxVdyQs5//4HsdRgALaTiKmDazNdPeEyHWY/FIONqC1MyvimAX+sGtowZE4v2qYq7/vsToj0JuZ34L2cEDuRjT8AhKnpSTKO3yTQOB71C6oStz5/3wEfAR8BHwEfAR8BHwEfAR8BHwEfASygoAfiY+Aj4CPgI+Aj0DPQsCCas/KUf7kRijKFhBrJi6s2RaHz+iBJy2o4OLK9bFo6YsQlNCaPDPbPvkI+AjkNQLm98nzS0AulMoQWHJWos0srT4A9swB+SViF0ij8g0uHT2zC1JOP8lo1WgobIAL8WAvgC4xpu8pBOqtrKfenggvmLUtVPZOETSMoDk1h/1oCkbfezVFQBL1qRegu8Ky/otI9V2Y/NLmOPPJgtUUkVXZVjyBaybUrnLw73wEfAR8BHwEfAR8BHwEfAR8BHwEfAQ6iIAf3EfAR8BHwEfAR6CHIZAfyvIeBmqL7Cg2Rzx+CzYfOooLWsEWft35wbYtTKraCTG5i9kYTuuTj4CPgI9ANhDYggteDyEavAqTKneG+cI8G7F2xzhE74ZI/u7SMz/boHIFBBuih5qMs2XKazBwLjGRlGFFD0HZHP8UgJRA+QxEwGKZOhISfQr9+/4Rk2evRbfVlRxAZsA3PgI+Aj4CPgI+Aj4CPgI+Aj4CPgI+AllEwI/KR8BHwEfAR8BHoKchYPW0DOVdfoQqW5HhUFyJaMGueSdfewRSFdTttTMCciWDj6b1yUfAR8BHIHsICEJsM4+FJTcgNnf3HrV5Kl2UVOuhcn+67J3OZxYgA3ILe7iRGaS9CIo5tPdC5DJAJ0NwKiDnAqjg8828PkT/d3iN0GZGBb00swApuTNniFdvyTztlV5AWQdwDkyP1+fyESACgo2h8leI2KvvJgB9Gx/Pf5to+OQj4CPgI+Aj4CPgI+Aj4CPgI+Aj4COQLQT8eHwEfAR8BHwEfAR6HAJWj8tRfmYoALF2gTp3YErVFvkpYgZSXfjSprD0JqiWMFTPOdWAmfHJR2C1QMAeF0+dTzWLta9AkZ4F5rFN+Ih2Ueq40+AQYbuJEXDit6A+tH8aIXLIoh+njUO6eKXiE8xATD/NYaY6FrXIkYxgd9p0aCEcvR5x2Rth/Jb3ZyD4y1+xftFVmFp8M0KR67HCmopQ6DyEoicjrgcgHh8JtU6Byt1Q/JxOItnnaUeMDiZQ3kFphVT0Je9esOetmRa/z5QnCGhm7QH0NQr+NW0dbTZoICM5EZbci3Nn5rDsiDIdL4qwvX+NZTi9PiJVm5euvyM34IEJafRhXqL7fj4CPgI+Aj4CPgI+Aj4CPgI+Aj4CPgLNEfDvfQR8BHwEfAR8BHoeAv4GgE57p2pBsCkVpa+htPIY2DMLmbTQdhcSXP9xASZX7Q8nWsO87AyRQHcR3pfTR8BHoDkCKRd2yCyVqCjeNW1bXrwLKkq2oF0DoQUFCOoGcJxiOHohI5tHu5S2HkgnbTQZ025uCAuPoHT2HjA/PdLk05lXkYnIBIts8JaXHIPLiw1mnZnT9NK6YF5/Mv6eNhVxkU5qYMrBtJI/Y/qYV2AXL8D0koWwD1yBU4ZHGYHCHh/DNaNrYY9cwvufcFnJV5i+2+uoGH0zKsYcjQ0Kh0LlaEBruOC4nJbxMmRrql+R3T61dfypni+uXB8iu5GtgDY1CczfzojU7Zia2efIIwTOzag9KC/ZGeXF69MWIRTsDwlsxTpxKNS5m3maT7uCNkabPomwjMmeKAw9B7t6cPoBM+FUScH9PcJr7JYRFtloG6cV35ZCLt/bR8BHwEfAR8BHwEfAR8BHwEfAR8BHIDMEfG4fAYYsSq4AABAASURBVB8BHwEfAR+BHoiA1QPzlO9Z6gWRmxAN/h2lc3bE4TMC+S4w7BlhXDh3J3z//Y1ciHuI8ufwi7O8R8MX0EegByCQcmGHeUyHh2zJyJ4QwSVcxJ02thrTSqbiowUjIXGzaehMLuI+BdVvoPxLFtbNTax7EdtrXLdoM93y0CPcWS4CdQewH9g4RXbqyHMffiwYh2lj30vB6+1tNgpUjLkbaw/ZHYJ9aW9kgNdpzcIpL42U5Z8AaIw1vYsSl6iOZtneKr0AK7nMpoGxOLumaKWLf9NzETCbXKaOep914mFUjD0aoeBWUOcgZvhy2nlsH5fxmglti6hej/OrhmUSyOf1EfAR8BHwEfAR8BHwEfAR8BHwEfAR8BFYhYB/5yPgI+Aj4CPgI9ATEbAgYnWvjOkyLht9kZYFfmLeHNp8o16AHAvEb8bmQ0+BPbcf8tUY2SLDTocTuxWqx1HMMG1+kcJ8DfpDWmVCwbKjy+EbH4HVGQH7xUDK7Kuk+vozZRQrGcxxzVN3+xhTx9yC+uUToHI6F7r+RfsteZQ2HVqLbdBl2HSdXXnNnmzppOzzrELAnrMG39s+gAyCmzGbOxRP8j1Pws2Jr/zdODNzP2uzepSXVNH+GdHAHxj4YtoXaWtps02ZxXfJi/0hsjsDrUmbCbEs66HopTn6ijsTUXzetBBwstg2mg0BFWNfQHlxGeJRc6rGxVBUAojQpkNsy/UAhPEXTKpk3YRvfAR8BHwEfAR8BHwEfAR8BHwEfAR8BHwEMkPA5/YR8BHwEfAR8BHokQiYxX9ju1Hm5AUErN+kZYG/M2P5uthrQbAzFb0ViMbuRNms0bBnBilvftDJ80IordoL0fiDlNOmUDtAhIpm3uUbiS7kguDlSLdcCObkWxZ8eXwEVhsErtxnOaaNeQzhwHmst8fTmt/HTif7XCjFDrD0XEyZ43/tii4yddgAIjsAarlKIPiUfv+H0LNmgwdvc0CXj3ofPxZeh7ieSFFOZQpVyOpPADDGTCgSXheK/RhEaDMjxa8YdnxmgXzuHofA9PGfIPTt32DJ8YBcyDqW5s8CSB+Wn5Ng4UD4xkfAR8BHwEfAR8BHwEfAR8BHwEfAR8BHIEMEfHYfAR8BHwEfAR+BnomARaVhmgrGfAFAfsalo19Pywq+4uJSPp4AsApMkf58OAgIPIdo6D+4sGZHPncdmZ8kmFI5DmvVPQ3Bo1RA70Fh8veEAgoHSJSyfpZWmTBlR61F8I2PgI9AFyIgCnv0L6goeRZO/70gMitNYYJs0w+B4+zJq6QZxmfLJgLBOPusVF+ry5sIF1bDtnPb/948PIrpJZ+hfNRdWL9wd0wba079yU5uM41F5Ncsx+tmGizBLwgAzunwjY+A+fmUqWM+RWj0lXACexKQX2jToQGwZBqu/7ggHWafx0fAR8BHIGMETpvZBxdWlaCs+lKUVT5G+xHKqn6mrW1mF6O06lM+z8KUqr+irHosJr7ZG77xEVAVlofDUFb1Be0VAJ99VLofAuY92tWDUTb7MJRW/o31fR7t57RL+F4b2oLSqhUoq/yW9k2Uzb6d7sdiUuXGaX/sYj6KKa18iOEqMaV6k04H6fyqvpR3X+apgjLMQlk127rKhSirashfw5VtXSXdq57A5NnnY1L1LrDfyb+TMjsdPD9BH4FujEA2RJ9ceRbbjiW0k2F069mI04/DR8BHwBuBKbM2Yp37I0qrb+PY5G2UVn3JPns5bVO/3TguqX6Nbv/CpKo/oOzl9WB7fNTknaLv6yOQOwTKqkayHH/JMfB0mDFpllKyAM1SVH40HUBAGLYX7cFc2JqHssqnMLnmENhz18U/5oXonls6kwrjsuoNcGH177D5sP9B5QUmaI4yLuTVJx8BH4GehsC74/Kk4RfF9O0WIhgxC/qXEeaFtN4kEoCFS9hODvFm9H1zgoBj9WYf0d8j7lqovgp7+AoPnux6CcvRKcOjjDRr5ZpxpU/2zEKoc7xrAEUcsI6k/1La5CQynGW6OLmn77raIWCLg2mjZrPc7AfFHFpTvr1hUAzDd9/93VdAe8Pk+/oI+AikgYA9I5yYhxoFRFnVJVSkvY01QgvhgO2STgHk17SbARhIa+aLTbYfBBvTbSzbrQupZ5iFgiWLqGx7hX3cWbhw7s5cLFsL3UHhVlb1F5RV16+0pVUOyqo0Z7a0so5xf0+cvqXS51Xae/h8EcqqD+YcfRdMqdqiAbs8OjGQLzptKn2ZZUXZp+m6DHMMSmu25rXzyX6nD8qqWB4b321pZYzP2X6vUb7HH2j5Lqs+ZvzP0N7Ad3ouplTvjdI5O/G6SYNST6XzQcggRaN4PN+UvarfMA93o6z6R0T1e8D6D0ROh2Bn2g1p+zLWhnZAUATIENrtyHcsgNsRkE8QDbF8Vz3MOA5jWd6Ki+xrwLYtNDembYgU7AYRswmyL+KB3G5sNG3dpDkbso3bC2VV19B+gRCWUN6nIJhM0cYCyrZOBvC+IX+AubKtE7pjf1jWZQjoy4gsZFtX+Rzf7fEonb1dIn9mwwQD5i2ZTQtlVVegrJoLJaZOVEWIRe7aukQ7V/ltom6UVb3DtJ5EWdW/UFZ9DspmHwzT5zRgtzHs1wdghgZygt0F8/qjrPJRrGzjq1hnc9zGd6T/KK2ah0k1m+YEi9aR2jUDictLKMsJHlFMYhvYOs32PJvF9dLqoygn29qqDrThle+xjx1lROiwLZu1HkR+B0EfthsHYKNhnfPOWgteVrMOSqsfQsvy3QGMWuAbRynbibKqel7/y7Fiv9bJZ/158kubw7QXq/KTi367CZ8424aFtA3tVGn1iyirnoEpNRUsJ0cxz+MxadaOXK/5FeyXhyCXm+/Nicxl1d8x/foGW0XZWryLJpnz41paaU6Ozvrrbxkhx0wXPNcfpdXb4MLKY1BW9ShKq1ZAA5+x3t0C0RMgsg0E6zOcWWMz/bWxjeMS3ZHuxyGAu6GRLxGp/prv+kYYrBs2KhpesmSJSudux7h/wqqym/13mOhXqzi+Yt9aWvU+8/Mo7VVM8zSYj3svqNye9xvAfszgkaWMNYtm0qzdGX9DGS01bUN19vNYVmU2cjTksYxjh7KqB2HGLqVVf4TZbF7KMZdpf+3Ps/v+mmWz027txDzvRIisD0cPRlCythG25YC703LkJ+SBAN+J7AvLuRfR2F34qv5sdja/RunMbRomiR4hM/GyZw5gg7k94z4Efb8/H9B/s3DdySg4yQFl4J1PPgI+Aj0TgQcknjJjop23mGqPjyEcm0pl8RVsixallA0wx9BfljOlQBoCrJYsqgKxBiCh2HNDQCIQ+cbNt5u4ZyZmJHwIBJu7B9K38WP4P1D8152HPopSNAz4+OBT3iJgqXaObKIoH/0y4JzK9Oay/KROV+QERBfuQ36ffAR8BHwEMkfA9EGTK0cgus5fEI3exQhm0V7EPm4bXoO0JHEA/Yk379C+SPtswjac5vQu739me9VsnCkm3HBAroMTexZi/Q311cfDLLqB4wrkq7FehuplCQvnckr5Nyge5fV7XlO3x2TMiEQKyD8YZuFUZCeImI2DlxDrhxF35jDNZxAgdpHQ2Zhc+WtMfmkQ/QTdwvA9ByLbEsvxzJdZ0BsMcY7uog1rEeL2L8rS8G4FV1GmfxHGNwGhH7JhWOZlLSTeJcwC0N4AzmA6VzLdZyDxGl4fRhDlKJ1zfGIhIVcLnUw4c+L7Kp27Ni6smoCgXoqQPAYLDzCeP0DAcse7BmJbAFMfXiGmjwNi2ozbAdDKfXR7GpDXAPzI/Bregbw/mO7/YTvwHCy5GpE9jluZ/7NrihCp3g1wLiEfF9BQACtirnzMMpmNIFMq90Rk6Hmw4jOYryeYwl9oN6Btoijr3QI+mLbueV6fbbSz6f4R7xfSrmrrRMwCg9nU/k/m738IyBWYUj0hsUgEYkrmvKN334tD8T/m/wq+o8soH+fi+DsgzwBYRptdSrRzYjaH0GJrprsfEzgO0KsA62He1xC7agTwIKLLr8Sbc87mQsbhmDJrs6zOkYqW1jOtBxJ5Nu08cCUgNwF4kta8V17yiAQWglFBZ5h+qIU6Ri98I5O7B2ZjP+jGh3aQaVNfZ7j7Wc7+j+/5ahTKZ3zuOM043EFAzJjjFsZ9OyOcC8gKpGdiZPuQ4e6FI3dT9z4fAJ06QGb85Mi+EN2KsQjL8a8QwniYTUZ06FQKWUspx3/47i5rtFfx+k/KMI/W5J2XNEmJEvA5RB6G4h98h9fynvGC7YX1XwwcZOpSmpG1ky2y4hfKz/e8ckx2DWO6h5bvUFe1wXTIAlmADKBlGyVDiCP7JD0c6kzmesldEPwPVuAl9hvPIhq5Fd9/XwqzCDn5xRJMn9cf2TSO9TlUr6Ml3sy741zP6JlvfZ/XzN4jA+ScBGZ8h5wY27ZwQeX6HC8chUDRZXwvTyMutzOtA/lO2PfyjoWT/5fTfgIIdSho6rN5FTMW+ZJluBYNfGA4oR3Gm9N4fZrSP4ho6HxMripBthaSw9YPjP+axDsUvkPotXy+CwBlhBkX8baDlOhXMRgNY85f8Xog7TnM5o1QmYmg1ECd+xEbeBESJx/MWg+qgmyZQusLxtdQRkUuY7rX8/nfjP5LWqXNBhUykoY8gmMH4FA+T4TgFqY3CyLVkMA9iH0zhfVxAswYlpknT/ej+uDOFPoQWkObEMuDzE02rJWNSPw4coKAacTGsTCXAw472tA/EYQ5Qq2cDdJvMlaamN3UU+ZuhtLq36G06nJEQnfAsf4JdcxA7GIAI2lDtD75CPgI9HQEzI7pVHlUyd6gIFVaxt8ev4zt0a2APAJFGgN5OQJvVo2CbzoPgUte5KA+voZ3giqcHHXzsYV3Dlv4mi9nRKe0cGv9oHozbh4e40TlQXqZSQkvSUis0Zx0cJEkiZ/vtPoiUDH2LbaNF7L8pKeQVJzeKV9jrL5vxM+5j0APRIB99wWztmUfdA0s4VhMpwLCeSjCMKZhXGYU7RWAHgPB0YjHTkA8ehIgJzfY2EkMezzHcEcDOIb2r7znwiDMAgAfEzSQYSfAwrUIxG/j4tgJWd3gjiya8tE1qCi+KGHLSy5E+Zg/I7EhyzqTeXgjzZRM3t8C5HaoTILIaVDrD4Bycb/Rqh4L4Ez6l/PKMbCahYiWSkERjr+wARQTIJgOkX/Ait6GsupjYL5kZcC8ppNfDVL2wyljs8VV+TUii7eiW+eSvU0E5SV/T7xX837LSy7AcjmdQvwRMHoRLOF9emTqhepi5s0ocmdDZRYD8ooPALrzwYXMZo9tIHzv4tzAhYTbOKf5M86uGejC33nO5ss6U64kfisc3ASRs5iXzShA09ieZVM+5rNRYh8HiR/DunwionIKQpHTMMhqsLGCPwHOyVxYYzvBNgM4gfhcx3Cf0wKCYbQMj+vI8y+8UX0vesXvgTBNqNFLgSZEGYxOjLdZIlWB2eQUWUSFvNx0RQt9AAAQAElEQVQMsf7KNHdh7E16MAeKV2kvAvRYwDom0dZJvLGdY3vn6ElQOY5js6Oh5mtDXMHwH0D5x5tGMhslTmAcf0M0egv1dwfA6OMaPfPm8sCEOOvCU5g6xub1IpQXl2EHtnXBAuYXZgPq25Q1UwV+HKb8K74iIpUw9aLJAh8yvl+IG3l415aETn0A2R7AiVCz+cq6CWrdjkjwbyidvQ/rScfLhD2+DuUldyXynGgHiifjo/ks65bJ95mAvAbAK99RQGrIwQUW3MBr+yzwNyZzHwCzme5LAKbP4KU1aQGcoGk3Wntk//mc0bWoKLmRfd5ZiIL1OHQCVAwuc5mYFyb0bkEfQeRUhj0Bcf0TPl5wBjG/APZo0162YGzXg4gmfgK2fMwU1FqnEZ8TWV747rSG8bmVL3rp52xXLgbiJyAc/BOmjZmGS0q+AujVEYqG14ElXKiRvg3RKHUmuj+wPhetGlw67b89cgnKi++BKdsJWzIJ4XVPJz4nUoaLoViOdIyCc0+5nnPQE+AETkU4at7huSwbFzbEPfpenLVZ7jcAXDX+J1SUXIdEXjguKx97HkLB0wBlvy2m/sTSyU4rnhV8/oH2bTS1T4krXqIbx2Fax2tyErC/kPXoeQDtRZTjb7CCt2Bp3R0oqzyH49phdO84TWddqSietjLf08aek6iTjsWyLneiZZ/TNj3VHwHr/8jXvvZJm9o2/Tvz+AwTMGPfRYwveTug0huHa4B82aUz5/ZLbBYMgn223gjBKUxgHV6FV0PLeT+DMp5BezQch/0zWNZNu9Volc+OcyxlPwqqZ7JMmwXqX0zgRmuhod+ZAgu3IDr/CpgvytFBY+/6HSrGlCfe4VRTdkvOTYyVFKZNZdn1KGdtkzbl/BfK/x7z0dC3AlW8N22qKc9tQzS49ILICIa7AAHwXVq3cP51DMtpY1vVwNTu/3bJZ4n8mfpZPsbMl85BOPYnpnkyZXuIlv1l2rGbPP7EMO8wxGw01Mlq3n9GW0vrQtIHihKolELAsWvsFs6RJiBXpx64SNFhZ3teL1jyJ8azJi2YlxDt7xObX9Bxw0Le8Uj8GHKKQBCKYXzpu9A27HKxcBsCzrO4sLo47ZQj1ftAY7MZx/8xzF945eBEd+L92hDJfiPNiH3yEfARyFMEtl6rabCUXwJOG/sjtHAyB2Smg08lm1FKc+Cfis33zx4C4xwOSFJ9DRKEhbWyl2YXxJRJkrFF+5J9C9rkpGrKslEIK+J4l0xmMYSXZKQcuDpm0SSZp++WLwg40vntZ0FJJbN/KseDyqs3WcJFvOhu3ky+r4+Aj4CPQCMCE5/pjSlVZyEYeJguZsFnG15DtA2kVJZDqDBz9kIoeikVv/diasnTmD7uFUwf/wmfv0zYqbt9jEvHvEIl0FNUdt2HtdeugPaiclT2otLJLKg1b7+4wIPxcPQaBPVOKqHc+9EGKbr+v1lsqBj7LT765kGITKSScXkKoaoQD2yBEPMf1bMQnn8NPpx/MyqoMC8vuQ9NtqL4LuL6fwhHpvJ6HGKyK+IyErD+xjR+SpKGBaFuQMG5PG5AoPYRnDdrsyR8+eM0KLIJVA6BNNc56AbM336d8nOHqZC4hgte5cXzsCJwCUSe9GZXvhO5AJYO54LIpoiFtkE4XAIn9Bto8LDEFcFxCNFdZUdwiYrlfz7jbF7++biSegHWKEAuRS/nCXTWMd9IYibP/BWCRffS53rKfACvZkOC8NpAqlz4lwmIKvNXWIapo+/G1N2exaW7vY3LixfAHr8MZvHQ2MuGL0b5bl/j0pI3MZ3tRTnLebigFLBKiMkRfPeMi3eQPox8JAQTADG6ri0gEoAxgiJYYmQwTx23Z88oYls3nXE+AsFxADaiHBavhsyCYRUksAfU2Q9Ll1yO8uJ/s76+gOls66bu9nminSsf8yWmlXzERcM5qBj7BHnuwnLrYoTC4yFi5H8dEAcNRgBZE6oHMM07Ea2+F+ZobuS5mSBxXDL8K3y0gAskzBvwnbfEiTpxHyycRL5dEA9smij/cEaiPtpQL0zdMDYSHYdodHs42ByIjwGcMuLzFsM1YcbbZiQQQAcBMhrAH2G+8OsdfwrmZzTokFUymyHKR8/H+oUzmCrbX2m+ONQ6qQjlfhpx1oPI8slor61fPolYncp247dQZxSc4I7M7+W0i1omKIXErOMbH1pGmuKJC+yXFy9Fxci3EGa/FcVhlOtuqJq64h1WsRCKS9jn3YmKMW+wDVgIg693qHb6Uk7Thk8b+R4+/vYOlr8zmLbZuJIkPv2cZeg0LFtyFdunGpiFcjC84eyINR/XqI5gFGYjf1ObYq6jWN75TsFyTN+uJHujOrZZb2EJ8y64gRhpCnGWEJrLEQ5cxDHfiywH37ONj6UI03ne5t2Vl3BhsG4iE51L60UO8/sKy+4lvO4PKfgVEKa1doLW7QXTNq20zoGIYlfA2RyOszVEDme4GYzcvT1o+AJ7C/IcBNOXh1CJKdUXw359ALJrlH3t0kT/g/gUCMwiqXsK5lTOer243e3Tynat3/kIBY5ESPZBTLdnHToY0PfaJCzohyFPBdu4d8RhyrxN0C/KPluuhcrezPOqRWvV5Xw3VyGu2yFaezK2L/4HyksewrSx1Wx33lnZZ5t+27RD08bOwrTiB7FD8f8hWncqIvGRsGQq87KwmYgh3pt3eQo0UIWyqlNx5scFdMsembFS+ZgXIXGWR5kD740cH9L/JOK+A+dLm7A92R4I7Q4NHQpTZuuihybGoDHdkgKOYl7+zmvz/PCxBfWDwRG4HiHcAntmtssoExNlW7EIl45+ju/mr2xHqP+kG31c6E2+x2MRi2/HfG4CWDsAdXvCCTWMH8LRQ6BOMSDMo+wGxe0AltEmI6HjQPL8mlj8HyJrXJWDesgkckT1K9hf6D4tYhdsDgtHtnBr54PVznB+sK5BwBTmMJNmJXWGIuYM4n2aJGuTcQgrQX82mqZR46NPPgI+Aj4CeYZAxfBv4egh7LSXppZMRnJQZiZaqVl9jo4jYAsnT2omP16TvxDf3TDYdrcdX6QNVOKoPz0enAUhuVFAnoPTaz6MKYx9RWz+x1s3/AKMagQmz96KPD75CKxCwGbdKy95gOO3qwDew8M4GAqV/WCO0vVg8718BHwEVnMETD99zsvrIVT0DzhyBdHYmLZJcUd1lC6GajmWDtmSijQuhHHx2x5fBwj7NqQw5DFfhU3b6Ucq4GajongXtl/mZIDPGTBOa0gg0pf2ECqhZqJ09gFURBUaj7y2ZhEj2L+KMpqNWby4kS7F9FFfwB7zQ0Jha0+IJF8AIVb2+BiVZXW0i8i7ANPHvILy0WdyPLwV38GNnL//Ar6RFikRPT4b/HajYvZVLiwenpf4mYURiZ8BwTDK24ykNxVa++GLCJV9zZy78vaa0b9QWflAG6xbyGQthBN/nAvbrybe7+Ujv4H5wmvaiJ/RZCtGfs+FpW9Yb95g2b8I8aKtAUyj/RFI1oerAOhFOxKBeDUm1ezLd9lUF+mcUxKc/0FfmCOMrSCVtNifqfXj+zIy8RZx4vEFIKdjwbrboXzMfxNl1B6+AsKyizSN4TVhzAJrRfEMxIp2YbwXA/otr25tiplLdByHk+eF2L7shF5DnwfkfABGL7YqXsUCynAeQoX7YOqoF7l48CNu2L8eSCN/Jl9m8dGUgfIxjyCK3VhnbQA/0zbkSyC8X4N2AtR5EqU1+8CcIEaHvKZEWxd/lTK+QetGMZjT+0KRk3Fp8a0wG2kS7R7rRQX7jKvG/7SyXpj6ccX473D5+G+4IPxZYgG2fGwFKkq2B6z1YY4OVnwGoJa2LYnZGKKDoLIbEH8FpdXXst1bB6aNacvdfpdThkfhWG+ybH7tGomaV6sRLMQKXLnP8g5Zs5BpsDF4mUXs8uJJiONwKD5qlr4Fcaxmz517a4uTqPfx4EUQeZayGQDcZRBwzmu9lrzPcw/WYR9TZqePYnl1zIJtq+h0EVROxNTRTzfU75beHXradCD18zgQSOjcscrImlAcw/oeWuXWxXembXMCD0PwQwpJ3oLqHezLlqTg60pvRcWe38PB7cRZ4Wr0B4h1KNsaGxXFT2HqLh+ifNevkeiPGN7Uv5V27I+Jsm42sE0b+x6mjvkPKkqOwJIlwxCQQ5iOOSFgIbFpGss2T1X4YPryjTl+uxiRZU9jclUJrs/yAjITgWkvFE+bW1erRKZPZEWH2qdE+7b9ctgcH5nx7GUlX6F81GMIsQ8HrgXYBvJfA4mFQUUGg4bH9v8X1pk+xO4UaO0baDiRjONdllokTJT4VyImeyL8wvmJ/uSyvRbDbF5LeKf4Z/gM/xWJjcsXQoJjGOJ5tvns93nXQCGmtj5vr0G/727AxZW812zkjVGSzNgh6HzONGcxnShdXEi/QTjwPC4reTNxUonpP80Ys6m8mj7W3vU7+vO9FM9FecnpQHh7RnYb7ULatiRMEehHjyMQDVZhcs2vYKvF5+ySyWNR7EPWvTmAuudR8SXf5fO4bLe3G/Ixen6iXjfl0eY4wpT38jFfonyMmVMeDyncgXHeR4GTt0+JPMoACE5FdPlDuHje+iwzQv78JbMhX6wjoNJ6ndeCJadiyqyNOio8X7LkNwgdzaEpdB2NIy/D871ZVu+0RRMdkDZvj2d0enwO/Qz6CHRrBMxg25Lr08jDEPLsn/WJPyP1yQWBmK6gz2JaNzKD5R0RG7+BG0Oeu6cvXn1gBES25eDTbRzFQbe8CPMVlInVKPjh1PD2K1oX0k1hWWP9Mu0Cz+ruHNebWN4+9oRBYP6KUSRbevL5nj4CPgKrLwL2zELU7bE3iiKPss/5A1uNUDMwHCjepdufsEHRJbghG8e7iia+JnWco5nOC7StFUFDAesBRAPnJxZ0kEUlGxPLOtnbRCii6c+zHnWLCM3JWBXFZzKt37Ltn02/5spJPjaSoC/M8fXR4Gl595MAmw/bESK/a5S05UV1W0h8LMwCbUufrnsKBN+GwGx2zZ4MZhxYPmYKLN2Pkc6kdTnmmz6QwQg416E+uG/OcbFtC1PmborQj1cyz9cDia/xsdIozALHsxA9Gh/N/wdu36hupV9Hbwwm4eg0lu0TGdVc2rYKGtVCgIpQs+EW7TTmONW16n8HkbsBGYUWhoua0JdhjlMOj7kOZoNCC/92PCS+mB4zFYqjaefRGgybIhIItmOZvwnRRSfAyNbkk69XCS1i2/IpxYvS5o7KqfCfOmYSwol3dBFx44IB6lwTFFgsl2dStv9g86H74zj2aciikciPjG0BrdJ2MrG/nF7yPDE4mwl/T2uoN/Pbx9x0qS2s+4bp389y/AOvXhSkvEEvhtz5ET+I0Vc0T2IFZZ6KijEvNndsdt+xWyncifHvxUiEtiWJ7I7o4pEtHbv4KS7zAfkSXkb1PWzQ6ycvlrzxC1ovA5p8ITBbQpqNE38d8whqosVwnEMBMYuPSGMqHQAAEABJREFUTfUTbYywRIiMYEv1JL7/thTnVlM3luWxrWIuFzWb9zFtxMiZg+kvo7iCuP+XaTSOaXQgikIhPneAiNGkyo0QWVxB7Bh/q3EJdDnMz2qJdTQu54K3bbcdOyBDM3XU+6jvewjf2JXMz3etQheyLT4BcetGlFWPgM1xUyuGdj/a42MQ6w2Gz27/Wr7r1whFT2VeOL7Styi/Mg0Xkq1hOXch+vJOyMUmALNxPK7vA+Len6MdZurwTxFa4ziW/7NpP0oRwzhE6+5GaSV1YyxfKZi7xptyFRRuzzI4njZZHVoXCPwWHRkPM2MWIxdefep+CAgrcgbHQGnXDxi7H8a+xD4CPgJdhYDI/ezMzSTTSwLTBu6Erdday4vJ98siAmItZ9/jtQHAJLYT1NovL451NdJkZNNktmeEYVl7kHsdWjd6h0pFo9Rc5R/vZZSB7xBDt4F4P5b7PbChX6ZXgebfrUSgsOg7lp0H+Zxqsr8ZzNGF2ZykMlGffAR8BHoAAuYLg0jBkbDkRuZmB9qWpHgPFi5A6Nv/4pTh2VNKmS/zKkpq2HdOBNRsAmjZDwo4ppPz6HcRJr24QUuh8vDJUfObm50gmCjCz89C3DEbAR5jgo2KVt61pIGAnINA/aEwYxTkgTEbTVSPoiStv2ahE0mlP0QOwjAnuT9ZOp2CcY5xdUX20xXFpSWvQgPncJw3k9arH98UllyAteq2pxy509VFdt8GToQKbzmG6bD+8X8TqcYhoFIYZ6O8pConX/Ia5bf5GhPBUzm2eZhJt1bkB4lTbwzt2z4M7M8LEak7HepUIHF0LJrHw7T0ecRwMqaPfhq28BlZMqIw+RKcxfbMnBbSKm4xX3FVIFLLtq5yjSwlmqNoaiMQMafBtMpDjpKzx/yA8uIrYeF4iFlcwpceKVn0G0F7FdYJn4AZGuB9dsjhYoWqWVDsnHwnk3rHMeY3t/9FLyMDF9ODtHzqSjJ1NoTnWF9fAjzrzGCoro+uMInTNYQLPY2JK6VVPI1ggFiybjY6t7x04OnkeSHW86OYymCXWAZC2X+bsZcLQ6c7ByL1lNl7cV/k46yOAXOZSQe/QFCbyyRWxv0iF23NUfKOnsV3bqypDy3HsyuZEzdch7FKUeBMR+lL2xL35v1QgqHd/4L4EdleWEUG5rIx38KcEqFo0tmG8UuhaZcziKQV63mzN0XAmg7RP9KnL21zMm3yNYitOA/lY7z6huZh0ru/cvvlCBaZscJlfEfftggkEmA+D6DbzYjudWB2P9QRvkOYNp7RZ5FMW21+DkGCf4bAbDKAq1FsD41ORmz2Jq48HfJwfiamsQ5FkSyw2YwdLr4dlpzN+FN8ICMjyDcJF1StlyyqLnezXy2CE9iTcmxMm4TEgvnpt1iRGT8m8U/PiZFo9gtbeml3Dpdq9hrYzpE43VSEjWI4XWaopM+bdqTdlbFjfVJ3zbUvt49At0Kg3jHHTRpFZyqxt0A0tGEqJt8/SwgEYj9xEOnxBXsinX5QnIuvIrsknrrTv3RlrVvXKDXGk72QNhnVs4+uxKdftzw+0nzxBH2Cg9RIskAJN/OlQDi4aeLe/+cj0BwBs9sezgusX6kmvaZcjsTivQuaB/fvfQR8BFZzBM58sgCFfY5kH3QJkWirRFBdDDiT8EPhczBH1pMpqyRUvl866h3EwUU5zG0bt/Shku1oBEI27Ln92vrnkUvAmg9oXadIZNsOpo99B5Z1EdOc55HmMMA5G/Xr58cYIlqwHSDmq3ckNdRm8H0XI1afPz/nVaS17GONQjapyB12/Pird5nnaRB5zyMui37mi9FTcfK83Cz6TXlpI4h1I+3+TMuMGXhZScq7R1BXeB4uL/6Q97mlipFvcSG+lLjMZEImbV4MUekO6YV3+hs8jEP61vwMUvSbS4hzGe26bQIqXqXbxbis+C1ec0PbL3iF+LLOIlkaa1CuP8OSS/K6rVvQNw7HWkhZ3ectuUBvKstd3fLLIDiN0b9D60YCyGaAXoo3qiYiW5sACnop01Z0pTHHVGvQfGU8n21SkPNKLiR2pUBNaT//HWW5m5jXN7kkuQ5i2R5FHr6fJL65dIouHkr5dmyWhFmk/BcuHrWwmVvL2448rbHc6DsOZnlxz6vI/igo3LkjyWQ3bMThu/F6f/TGInQXEyoyX4X/0qniTi9ZyEXoBxAOngALV7KO1nqkz/orv4HEKzD5ZbZXHpyZeCWWVLuwmRKO6c3HLYJXKLYRpAiBWCHv20eTajZFKHAHAx9C23JTIhAhxrdhReAamOP7yZB1MnqWUPCfULkJiqWt4jf1e1u6XYFNhh4EWzMflzBwG7Iczrt0WRv3bDn8EKpmXqYzui9pk5MgxDZzH8Stw2Bn+TQdk6IlP/GSKK28ZpdscbAe56xqVTBi9xM5oFwPlV8joAfDzIXJnFdUH1mf7ciBlKkXbRJSU/62RDw+ju8oiHaa7BTadia+MpjZlYfErhd2RKmuahqWlUE9bxxLIVacPGnEm0iXrN2GhBU5lIG07S4kGaSRbdZ035vhyyRtJbMJk9o2lE2y++Qj4CPQqQiYIxRFnmE797NnugqzU3EDTx7fM3sIBMd+CkgVAO9BnMjG7Nb/i/K5a5PXDFh4yX9KT0IOwCS2DXlH0brREqg8nPSLKQcPQIQTVbegGEifQ5GtiQUj86kHIRCp5aKZei0ANWRWdCwKV7hMIhpY/P8+Aj4CqxkCvfvuz0W26VT0rAvhX8vsc14kE1Ex9gncnMUv/1umwScqDKeXfMbUz4JiAR1akpiTAHAs6mM3I/ElX0vvvHlyLCrPxXsslE1hjaLVHFMqUg6I22Ic36psB4mdha42ZgFWnUP4ns2GSXdpRPpTC3MGF7pD7kyd6FNbGIclXor8jgljTsL45Lsq1sP7GZHbe6QXjO7mcKy1ovlClnHvuD2/ahg0YjZZFzMykw4vzUj1FWJQiquGG6VtM48c3l4+5iPEnb9C0XzDQYDlpx8GfZNZ2bA/L0Qv5zzG9WdA+yeR2hxffi1Cz78MNobIlZkwIY7Q6Co48ncmYb5c5KUFFTJ/JyEaLYM9M08Wd1vIBwxbqnRhORWH184l89vTU4ufRLzgYL7LVCeuDITIX/F69dnEsm2ZzlzyFVD5CYrOz3dzWWOxz/n4Gosp64CVH2XEth0M6Pck5XqJ1o0stnH7YeJbvdwYcuaucTM/px6CKSjfoGAm24GXWD6ULkmp3Y5mISlgncu4k7UzzaMlDoGTcfXXrRc1m/N03n28MAKV77wT7Nqi7y1ba99QDBCOydC5RkRhj/wGwUKbZayciXst5HIBUg+ARP+WtZ9qUp3PsueVJkXKMZmPWxyZBTZSrPO9EAoWoD1m4qyNYOmdELD+Ktu7FpGw7spMWPEbcM3oX1r4ZPvBHrkEYetGZufF5FHrZhwfXY66qt2ZX0nOk4FrTKOsi9EMQmTGauZz4eDTDPQUrUc60pv5OgqxXkPJl2WSOr7X3DUo5rS6cORxCm3yGefVjdYA5AiE+w9BvhmJ70ORdqD1Io5z4r/GCrRbfgswO2vRtcaSCljWcFbo1Damf01bWA08Ao2PTytexb5px5sXjKYNBP+lKYyIV0VIM5JOZku3TKjsi75F/0tbOsc6P60yYcWHQ/RfacfrM/oI+AhkFwHHeZ8Ren0hAw4mQmwJR3SL31FEDzA2FUCWZQZXqU4BMJkdghXR11FWdSTfz5qcGHV8kGxizZ1NL2b78SIIB4+AxwRHX0V5cfJF2uklC1lmze5mr/SOxNIXB3ox+H6rKQJGIQqZy9x7bSKhN7ZFoLfbcZTG37c+Aj4CqwsCtm1hypxxCMAsRCXrW+JUZD2IcPTuToNkyZK3Oc+6gWOD5G2ZYAKiiy7G+VV9O02m7pDQhwueJ24zPEUVPRiTZ2/lyZNrzz7YFBbGASik9SbBGKy5Yqw3Uw/yNZsArBDfobY8JaptFvsB8kdk66tm0NivDUNIrwSsVUdk03klCb6C4EpsNybVguvKINm5EUXR93OZ9o1sixY2xmkWEXuhtm+g8Tn1xSzI1c//Axn/BJFk43TqxOQRxGufgM2FTDLmlMy8qSBgNnuYDRfJkirknOB4xMK/z+sNT8kk7yy36bt8hnBwPMsFF5kQ9Ug2DAtnoD50UJY2AXgk1UlevRbUQ52XmFoMEg/ymh903vbL2Q9dBtXlrgKJboPw4gPZx4srT7Y9zMlBYo0BpB+MEfwIlafxyXdei4aGsz1W0L/vKLYzu6UZeCR+/HynTsXDTbBeA5Ry5G5Bzi3dnupuD1+BNfpdy+xdQbuI1ov2RKD26h41tg06NSxPZiNCAPHl6ffXTSjZc9dFQbCcbcouTU6trgvgxG9FcKb7V+ytAnTo0R79C/uSS5in5JtkBJsgoBdiUnXXjrPTzaTZ1KDWw1C2h95hNgNie3qz5KnvxeN+Zv7MxrQfPSUU2Q6Fzq6ePJ3tOfmlQUzyGNpUfbxArF0QCu7I8atF/oyJgZSD4IzDZTeA6je4dNQbuHS311Pay0q+SjvxaSN+RsXYt1LGadKF8w7jdWi7ByknSU4Gu9wUdd0jY82kTLdMVIx5B5OGL24W0vt2+qgv0ioTiXJhfecdme/rI+AjkDMEov3mwyiiOPryTEOwJRb/lFrB6BmJ75k2AoE6c2zmK+nxy1Dy/R3RuisxpWY8bOW4gy55SWkKFRm4MUT29eBWDs6u8fAHrPhN9F9B60IyBIWh41w8fefVHQFL3oTAXenWgE8QTsxtIt3A4f/3EfARWD0QiOy1FTR+NcC+BUnNF3Ccm2CP77wvqG7Yvx5O6FFAXqXSRtHaCFs5xSkIy2/9hbFm4DwwIQLE/wV4LLxA1kYgsF+zUJ17axZh1dmTi0PmuFSTtnm/i5K+Z+ML6QNYZ2Xtq7hEnHn+L9iHimypSimlYne898paKfnSYTCbaaLLTwDkAJaftuNxRRQOnkCoz3Mwx4+jk43Nsl0vM5jqTFqKiAAslo3Y0lDiOZ1/ffoMh1kEBoYkZVd8DUtuytkRwskSNcr3UOBqZii5XklkLdaViahfukuy4L4bETBf2cL6C++qab301+vBvP94wZbk6wH0Hhf+A+8xI8vYfq7Ba/5QsGg2RJ6nQKZ956U1SQEgEzGxJjvtF1IZFURjW7GelZAzTGvkegeCyqQn8pGhgdr5/4Ln+vGdHEk7oDEG9s1YzmeTbqNTi8s6sKxDcH51nxau/kPPQMBsiglFTTv/N2bIfe1FWCJFjkJIToKdg+PWmXinUyD0Ocv9hxDpDQn0yij9ic/0Zr09BnDMzxG5LIDqkygIP8tFz85br/vg2zeYJ75LjSXJjwBWMYKYjAueS3X6R5LgXeAUDhrdUaqNnWE4ui/Lpct76AK5001SzPpo9DVAv0oRpC8gY5EvxrYtWLE/ANI0XzLlbSHcjHJsqzIB2C2zetYYn9V49S/dDQFhcyRO+koaQeIc+wQAABAASURBVO6Osutu2Pny+gj4CHQPBK7c3kyiPqSw3m2d6sYIFpiJHll9yjkC9vg6OIFrqKwyO33TSa4fmY4i/z8QrboVZVXD+Zx/lLZEejJZmyb7vG1DbyHY3+XYsEbeqbt9zF78mcan5BfVP8Moa5P7+q6rMwJxTuIU6dS/7VZnmPy8+wj4CBCBSZVcNHDO5F2TcoG3zcgs+qmyP7Jep6ub4ppeOaBlgz5lX/g4xKU9E1mDyqiTEFm6DZU6kgMJumeUwaL3iNt8D+GpKdBRXfY7l4MGrsX3dhgSC/ugUfNlzg18z8kXQMlBv2IEVhSb29XC2ttEuNjesNDtlWGRoYjXr+PFkpafcnEsaO0KcOEBQgUokhj9DoHAHbB3SP/DiiSxdMjpytE/wnFuBGQxywTLsRShoCiAdIz52QmxzuZ8w62tU4jchsCot9OJLrs8dW8xvjtplbYtKTaDxC6GPW/Ntp6+SwKBcOQdQK5lX+Cl4Df67TFQ56RuuYiBVsa2HYQi1RAcDYk/0sq3ax/NV8/QOyiE+2KFyHYIx83CHtlyTGd+Ql2QjgNkCxijahYLHyJ+Xn0lYHgztioI9dkRZiFJEAALJS0Xn/Bvviu3DdrmtI/xCAY2I69PPREBe/wyLI9dBUcfYPaU1o1MWT0NkYIxbgzdy33EMpb7s9junoFAMNXJRi2zFuy1C6vPHwHpj+TmF/rfCrORLrl/blzNSU3hgtsAeRfJjcVx+G8QKDoBh88wbUByrrxxXf4TIOxD4W1E2X4O6J6blGIDvwXkY6Q2WyNffmKudo91Wb5PY/1p2uj6BcWvoF1G25aEnKL7IBrivLitdyoXM0BKxeP75yMCqqZDqctAtEx4M4jWZ/UR8BHwEcghApZDJY0m7wBXJisbwApyIL3Swb/JNQIFz7zCwcrvOPA1fVE6qZmB8aaAHMtwj6K06mwquXohj0xaokzjQorgJG9e5/9gFLveTEDQutETP5F1EZTDUkXj+6+GCEwb8QvLzvupc07FW2omn8NHwEegpyJgTt0JOHtTXXAgs+jyRYd+D5VnML14EXk6l27YrB5xfRYKFwW9WhRoOBCbgIlvdb8xA4XPDS1fBJEF3nHLRhg0oLc3T45866PjKR+Vuo3xO3iI7/lWqJqjrN3Gjf0Z5ojV6hQA4DMilEpH0wtx6Ue+jtE5cwohzp+g2JzjcHGJ7EF8MHIeIG7vCLk3THuuM5syPtyQljMQVn16db9Qj2KYg1iOzJyDt61I9Bu2JXfAlhg629jjYyz/9zDZ5JtgBPyT3RGpPYV5F/L51BoBg2Eo8gydzeKa908BqB6LSKiEvN2f7PE/YWrJ8zCbx/MtNxqqgQrrq6tgXNSQo2DPHOLKkS2PfgsGMyozby7kFaxRXyFY+ChMuYG7aZePGY9ofC/W6Q0bw9fy/nFYegOfXcYz9BFsDSu+B/Jl8Yki+ZRlBK4dv4j9jFnAMyd3uEeuyrITPwETq025defrDj62OCgvnouKsS8gk4X642YWwkIpIMQCyY3q4/ioxJw+mtw/p65P/oC4818m4dC2JUER++vzsNk627b1zDMX0w6qfpFSKrH6A9GGNjQlc54xfPlhHRScG3Ec6S3aAODH/Fg7COIQANSR8z/gsB95BHG9jU81tC4ka9LjL+xXXeb19HUhM6l28fKd8xoBYdFWpJowrsqC+CcArALDv/MR8BHoPgjou4B4bwAQFHCSR4UWcmNUC3B1TVHOrTkyNTc5yH6stu1wkP8ExLqUAxXvExpapm4BMpTv60pEa2eidPbvUDpvKP4xL4SuNemlvgSnk9F9UKy6AJHCRqUlOb3IcriAqy97sQB6DMwXTd5Mvu/qh4Ayy9W03iTKCamKN5Obb4rJk4igPpL7dtG0vTM04Cal7+4j4CPggUB0zlBoYALU5eh/hWlL3odlUbkm5t4jshx5FcU+hEila+yCEP1/h4LFW7vyrG4eP9ead2WsR87VQv2Kdrb/HtGm8jInTgj+TLZG5Zr+hGDgBhQVUymHp+ju9rVoAJBdEazfnmOfzpcbXWAkYE4681rIbBDKiXcQD44DCp39oNgfwr+GWFv//wUB/AsPSLy1R6c/vzg+xjbpOih+BCQEBC2kMlOqN+HiWxlzR/6kzKwv8gTKd/s6qW9nOIZ7fcI8vcCkKAv/tyXWAZyA0kqO3dp6+i5EwJxCFw/exLlniq/8pB9E/9ptjmdm1rolhWt/Js5m3vuzq/zmFIBIeE+YDYmuTB31YBunwT0YS1PdiUPlJlwy3Ou0CLKjfVaWma/4D4BIQSICBdPRx7Bt8TssmwYPtzpeCEcPxopFZgEnEdT/1wMRKFjrM5aDW5gzdx2mSACC8Qhpcff4gpy5ySYp6+ywwO8gMp7RCm0yWsZ6fCu6alxiU98ZKnge0C+TCdfgJoPZBp6C02Z2g6/mdXGDzB7/zXvx8M5rL3NqQ+Ln0dRbTGFxWxHiP2+2nPtOmbURFBOYToDW0Jew8DSmFS+CgxvZhiw1jkms0G9fxGXHJH6eTqkH057Bfc8uQ0CppAkE098AoE76vF2WKT9hHwEfAR+BVgj84nxLl9QLzMHAFuTLDakcgZ/iF+fc9ul/TG4ykMNYI/U3MfZ7aN0nOPRMQhx/yK4Q6z6g7lF8XX8qpszaEvbMYBLeTnBKI4mJ1YNhye88OUVmoPf3Szx5VnpGfwGojARicDMWtkCRM87N23dfrRH4NHXuZQgueLVfar52cDjaH9bySTlvF39xLsRrs/2fMmjHK/KDrO4IKJUbcfazKIHwLxkcwv5H8BLWD/+QzLtT3MyCDpznmJZ7XwisB7WOWi2VpASmDQ0qClFpVNjGvaWDoqBXCi1YywAdfrLNb1nKYRDhIn4itjj/P4gPvnkPtsQ4hpoF4D1aN7k2AHQ/rC6/kRxwqPRTjoeJiBspolyY75gex65ZiwrqUgiK3JJheZqJ74vMu3Fl6VSPZfgA6vyPaa4BiXmfAHDyvBAX1Y4kL8sP/yclXcI5xzNJvTrNcWmEdWA2y/hy1yRFhgHWb/3Nv64IAdNHfUEcrydHhNaLRkIK9yfe4sXk+3UAAXt8DCGrGqqvAeIgmVEM5Ds4ECuqc3cKwOQXBrJ9OwmQpgX5DxEL3oeUph0M5rjvEPsp1V81hlZY8gjCAz/ABIkjLnexPTX6q0bvVheRnRHGfq1c/ceehIA5CVIDpr+Zx7LgNt5hjmVt9u8HY5Oh/fiwetEF1UMB6wxm2kv39w6cWvPzOWTrInKWfQWVt9CwYbqtEMI3qBiHNayd2M4J8tlYCCG14XjTiqVmy0cOznsFXuWpQWjVevSKRhseuuh/Qucd3BMiWzZIoDHAmY066y26ceQvcyEwH/skbz9E+iMePDnTsaL3hKNBEv9/PiIg6iAecZ88tJZZrSWtnfxnHwEfAR+BvEdg8KwVlDH1BgB1BpMvNyQwyswLIJJjq7/PTQZyGOsV47/jXNtmCn+j5YCR/zMlwXAq7i6DWrciGpyYOBEg0zg6yp9O+HDcKJHW9WD9kfl4Afavaz14VnmZRQ9LzMDO/WskBZW22Bv256kU/avi9e9WEwQc969tmiMQWt7OtlG9J7GCvoCck/N20cH5VKz5X/7CNz4CGSJgvxuCym+gcP/STFEHsSpxyvCuVYQg8jIXEL7zyCF1Fnootlh/Yw+e1ccrFl4b0PU8M6z6A5YtTT1+9owkU889zQLPCQzVpID7BnDuQeKrHLr+UGiOvH+KZZKKLj63pUL67YVgcKO2Xj3QxXEMXt7jO9GFiDvm9IT2A1DvHMC+2msjXYS4P4Wbu7odaJbF/qPqIXiaLjE44r5xgQwYuHQD8nKMbh7crHyGeLxrNzjY4+NQeReJI2rd5EQB8zIehbHVow64wpDCI1j4X3K8S+tBEoSFCZhUNcCDKb+8Squ3Qensv+LCyu0B9R6H54vkC8JfUpQnKO8KXtuSIEC/0fw/HLk6BcAq3BfQndFgYqxD9+Pykex/Ghxc/7fHY6NhazP+CWj6+l/Aebx1B8yir4nvswUfQHAvb5Mv3ACFcHAiJr80iDw+9VQECuq/YJl8gWUheb1oyLdF3dHuCEjuPmRqSCd7/8+v6ss2ajKmVJ/Z7hM8lW1bAHsRmy3dBROH+M1GUfVSd55O8AmtWAhLXqesda6pCdaHWrvDftV7rOIaQSd4mA26jqyTMqXEaSaF3XPt8MynwoCYOa933+no57DHu79PdIKJWkM5792PKfWnNWROvHoEfUb/ZB6wpGAhNPAQ773K/+7ohYxOAbAYoU/dEgGJIxDM4GuNuPsuxG6Zf19oHwEfgdUCAdt2OP9dmDKvahajUnL5DLlAoHzXrxFFBQfHx3OgnulJAA0SCThgltGATIHUvoiyqjNgJhjoHJMyFSOLZe0FWFz0dOV+FWG8BYgiXROUN8hqLC9JiQs4WoLab7rPxDBpNnzHrCMQkIYJQqqIHfEqs6lC+/4+Aj4C3RWB2GIuEOvuEP655kF/RNT5wNW7szyqrfmAvA4vI9YwaPxQL5bVxk90BN+qWTx2z7LiFVy5t5fi2T1se3xs20IMB3IItG0iuELhOC/CKVy1SGcWmCX4AGX/KsGT7J9gW0h0fyrGJJl3z3KzdoUIF8c8cqXyIQr7ZqDzaRXX2TUDYcnhdA3RutH3EHnVzbNL3G0q/2PWfynX4Qj1esdVBlWWk9AI+m9C606qHwLh9DZOusfSUR9FLM5FSfnUIyLmB5tDrB38E088ULJ3Nu/yXrYycQ8usPxsj0DAY5EJeWZ0K777I6GyAexLTFnIM/mSiGPadct5lD7uP8tgTrYQPQzL3uJ8n5zZpHNnrsl+5yRGWUBL0o9h6SO8SUntYgjiWIbbhraB4vpvfPj1Rw0P/G82vDnBf/HOo5/TXWBFTbtMNp96JAKJBcbAsxzLePffpm6o7tFtMLACg6DW76HYFx9Egu2Se9Lz/WCBejW4b4BUpxYqb8K+OH29WruESRHIfNjj6PsQeB2fXwSRMUB0cIrYus6796GUEb9KKYDoK2jazJSSOc8Y+vTpx3lHqo3iDsvevC6VPLERLrgr9ebFlCNAS2JZd7QKNse+oDH9qmI279z1xKLrAM6hmPhmb/KlRVZaXD5TPiIQR7TQ60uJljIrvm7p4D/5CPgI+Ah0FwScX1JKKtoNfncpZS66L8PlxUtx6Zh/w8HOgD7NjNSDs3FkbjiAkc0Z7AaEOOgpnb0H7JnukwMyZoFSR2HpzlBQwajJx03mKCmVKlxS4j7ZT5aKPfoXqDOTmHlsnJAdEMDOSAwWk0Xiu62WCNSHjPIzddY16G8ASI2Sz+Ej0PMQUN2LCqlh3hmTd7FiyPfePJ3g++L4GKDV3imx/1XnINhz+3nz9XDfic/0RhynAuKueFVdCnMkcfvGYWifOWgIx0kHQYXjuEQM3wLWoyh4quUm3qkjPqGCznzVogmutv/CjOOY/AwrAAAQAElEQVQITMrRz9e0Ta9rXGwuWqmar3880qciUDAD9vD2b+QodLaGYnuPRIzXBwhLZuNXEyrX1swtpo751DP/l7zbm2V9Fyi8vvKuh8r7qIu3H0dkyxRxEcExX4ayzXOJU3UQLAzHRsN6uXD4zqZti8fMT4qk6r/WBJxdOZd0by/zC02zkaX7vfepY7+Awhx979auWxyPHICC5anaokzfhiBcUAyVHRMBlb2jyvMIhD5PPHv/y9zXnrsuoGazQVN5+hIB6+aVp9w0xVhQ23DaDZmbnFpeTf+tR8NsXmjp4T/1JARC1gcs9+4b2BryakGwP05+rHvU+0B8CPsntqsNwrfrf6DXRgy3Ha1F60LyPTTwKUxbj640olD9EgovfbRQwq0Ri23Kq7nnJc/o50WbMx8p2l9dRJ5n80zy9MUJyNYQy+iRvcJ8C8ua48WQc79lzxZBnN8CshYajMPy9S9MG/tjw2Pj/4ZTRMwmgNpGl1YXKWC43RCuNeOGVn7JHz0qXPIAvmueIGAm9pcN5wQiTXlMYVK4FJw04/DZfAR8BHwEugIBsVIvdInlL3J1xbtpnea0ko8Qih0OdU6FaBXMwnhrnvSfd+Ag7jFEgv/A5NljYM/L0cQohUBnPlnAiVkxucxkhZckJDKf+TW/9ZbEM4WTE3wKKt+5cgkCACagrqo/rz75CDQg0CtoJqLa8OD1P+63jV7w+H4+Aj0RgYaNc4enzJrgQ9ywWX1Kvk5h0HmpxwyyOSIRo2DrFInyLhF7Rhih3kdxTWG0h2xxjlkeRmjAWx482fUyv4lct7QECo7bmDrgQOQVxKMvw7YdtDBUpgb07+T1WnDeFoHlR7E85KcitUV+2vEwQwOIhveDIMXXWPoaNPhAO1JoCHLyvBDTKKZ1Hz9StQ3o+53/cxENInb4f2zx2hx/b8c8hlzjUq2DOAvQv76Lf+qEEhZ9Rn2cfAlIHdyMSID1Y3uIruHG4rsTAavwKwg+4J0XFbEd2RZ13WAzbMN8c7xXZvLXj+16NHofRLwWOwdBnFPQMD7JTlbOnck440cwsoaNZ6JfwdEncfEIr2OTyW4oQ2vPDCLCtESa9AExCB5EsH/bD+0uHlcPiT9E/7Z+K5OVbVEQ3B/+Bv+ViPS4G3vkErbl6Sw2boW1B22S9/k3Yz1gOOVs/4dXpryLbs84htJ6kP6IgBhdhwdPJ3mF8S2gPzA1pXUhGQS1toOt+TduPbumiIveB0FkiIvwxjkOyDNw5DV0R5PYHC77Q7C2h/gxjgee6PKT70JFw6Gy10o5VefA6ffcyuemG9v8TIE8wsevaJOT6FZAdC+Y8UNyjhau/gaAFnB0q4fMdzWKZh6mW0HiC+sj4CPQIxFQZ3nKfDkoTMnjM3QOAvb4ZagYeztiehzEugiKl5hwe5VuRRA5hoPW+1Bfew7KZq3HuLJLqWIbNHAtWDiUbB5jJn0VPxa+SZ7MafroTyA60zOgyFiIbOPJ43uuXgiYLwI1jSxbVkEaXD6Lj4CPQE9CIBpem33vFqmzJJ23SJxKGCtsFBxLPNkE/aBiFIeebD3S0yheI+uOhyV/5ngg4JHHt7lofG2nHuG59TqDEAgcSJmaFG/LoM6jGD6OSlO6tqbg2C857rmfzjHatiRg/qxTMeV/w9DjDJXDb1RuTXxOBKQI7uYXOHIzwk+3/CLInb+tzzD0J87b0sM9HdF6qH6BPovbO05n9F1GwnkGFdqyrqcEFpbCCn4HexwV3J6cufe8+HDiHDD1wn0DgJFCsSmCyryZB98mRSDcZznbwrfp5/Veg+TZGEFnTfLlNw1Ygwtrukt+C+kh3RXjv+Oi9y3kcC/bioNQFxxJno6TbVsoCu/E90vcEGA7FodQ52CFXoeIpkwgU4ZYiAu0ajYbNIaUz+HgMVy8Net0o1PTxaTviCmbVXRK3s8Bveh3ICKVTf0mH33qcQhYYvRg9d75kv6IxU059mbrat+tN1wLghKKYcouL+2hV43O1hzT7r2JQGQBnIjHCZntSbudYSRmNnKk2ACgYcbOd/iih76QHJ1NZsNFkTMSIocxafexIPAloPfg02+XkK97kZkfxWLFbI8PpuAeedSPYeEBFBan/yE1I8wqmQ/aRM5gnINoSbockOsxfbuWp6Wh0YQK34PjtN0c0OgNSG+IHICBg9PqR/KkcMqfUVbzHMqqn09tqyYiXVNatR9Kqx5JHSfTFWsGoxXa7kEiX7RD0M/aEabrgqRVHvjuyipnoKyyOG1BSysvTqtMmPSBI9OO12f0EfARyBUCHh15Y5JKBVbjrX/JEwSml3yG7edfBStwPCBTODE3yn2006wHSy4Agn+HXZPii6nMUkjJHY2NAWRbeBkJ3AnzW01ePF5+apmv4byUFUUQnOoVhe+3miFgvqARlopU2Va4K+JShfX9fQR8BLonAjFnPbYOqZQBtXDi+TM3jMsvgPwMbxNgvraA+bLZm69n+Rrl1mZDxgHOpYBu7pG5rxHHlfgpZBYePNiy7FXnbAfFeAgCSBh9Cyh6ChMknnhs/c8WB5AHUowLt4SGD4FRXqIHmdKXBgPUfYm4/6wU2G8L7kMUj8BufYIC0je1dWtDZUMGsGhdSKh8xLd4lyXHhSNvnc2XduoMBXRNTxlVFiMeYdsi6snXGZ6JhcG4WUio9UxOMBAIrOfJs7p7Dq2rhzqmD0uxuIahiIdZ7/IYsNK5ayPuXALIAHRnszz6IMV/g9aNBsDCKbBnei/+uYVu7r5iz95wdE9A1ocxYi1DTB5HeERam6ZMkLTtP+aFANkbgs3QYKIsezMRtt6EqdMNbi3/f/zdD1B9mo4LaZNRAJbsCrV2YTtvJWPw3XoAAhr7mH2UWxloyqCwHO0CM9Zrcsm3q6kDkei+UOxG0RrHerzLlOrjRTB1VuH+gYKa8Yh8i3Dcu59EpxnqUuR7psaxK/+7kXJ8Xtu3t5t3l7jXzWX7qJOZ9pa0QtuWFObElNsQlRfxwITk4/a2ofLHZYv1N4biQgi8TohbAsjfsTxQjcQcBF1jYvVmM8a+qxKXSoQHPLnqudWd+eAHcjddF9EmJ8WuiNaOS+7Z0jVPOhqlMl/3YMOYhhVTcFvmwu3JwjDGyQYqnbhhFpDFLaq8c1d9P3OZ5J3Mw3RpiN35/tIpE+bdrZW+pNY26cWbKDdm0px+1D6nj4CPQPYREKyRMlLzdUdKJp+h0xGYwEHk1FHv46P5VyEaGwHBjZwMt3dg2YdhD0DEeR4XVG2dpbykjkblT2TixJ//k5HqNwg+6z5wSxamtVvF6DfZL5mvBFr7NHvWI7rkBIRmEvi3eYVAekpCSUzq8kpwXxgfAR+BHCMQgFkkDnqmoohBQrWePJ3p+e0KsxDJRTqvRMXoLjbAgLpCL64e5Xe48m0OPR1iPcAx1HDmzUXxqouhzhUoHPBAhzYkMoGMKYAjIcKF2MaQirtQMdx85dzokOSyZMlHUOHYSZIrVEWYTzkCqB6SJHT3dLI/Z7mNTYc52QrgfZJsmCP5VWchqpfjyjE/JOFI38my1oQg1cJnBCLLsfV7Xb84nn7OGjgXvBqA6rq0vRocXP9HEQxEXH072yMgPzHJVG0v5zyOr4ciUK707c5xtosL6J/qK9F+QHwQbNv0H2TPM5pSvQkkdgtEx+eZZJmL09/6HurcyzktF8xcgovsjljhSBff9J3DgXUh2JtphROB1HkXhfIcbJc+BS1MZg9f1Q9jvvZnoP60JP2J/dczsEe5L+yahbSQvED53mWA5KRYl+99H9SO8X+uLTlC3d815LCMWKk/ghFsgFHrNpTlvMu1Cr6q/x3L6mUUbW3a9pMV68M6sQ7rrvvam6jCUfbZRU77E8pmyHERKNhvp2hbLKyD0PKB2Uy5Q3GZjWVW7N+A7AkgSJucRP+D0IDrcXmx2QiQnCdfXSfVbAqNPkHxRtBatMnIgcotWLLkFlwzOtXYK1n47LiZr/8VpzIy1gH+B1bw/79hb23mv7x1oYLiV9j/POziC9al3hA5CfbM5PMKrDJuAK3i8O/yEwGVlzMWzEE6vz+TcbR+AB8BHwEfgRwjkHoDQMPOxVyJ8SwUt3aCpSI0V1no4njNJPiK8d9havEZgGUU2PdTIio2UwykydSCBOZvHQ5hX0dZ1bEwv2nVgiHThxT8ZVWUVcd6cClUplKp1MEJiigsuZbpxGmTkxhleOAcTpokOYPvulohUB9Ib4Kp8VSKURfYWCZdfBqczZFlzh3IddsIvQ0B+bghTf+/j4CPQFoIiGyQBh8VPSHvRdo0Iskay/IfoxB4K0FUBSL90EuKspZu7iIqQum8objguf6J32ZMnFpA+ZOlZ75yN/4T3+xNBc6aMAtCZvxRVnMGNqv6CJDrAJixsPDairSe44L3ATkFH3/3d9jbUGmKzjMXzt2ZiR1MGRpkU7yK2uX/oZs3DdpvGcc9z1OpZRbwkvMKtkZE983rr+KSS97Mle/clIHJVSWILHgBguMAjmL5LwlxoQAPAEXH47KSr5L4p+9kypQVH8z0UiwsKccIgZ9hd+CkgfSlyi5nwQ8WxyDMo3gvmpgTAKIw2GY3/fbGFkMtVGMpg4vn7/WmDN7jGWxRLhItZj5TKfT7cBw5CFtd3NBGMUCHScSCYBgGLtsBF9bsmJE1bfuU6t1RWnkC57L/gTrzKI/5CRWL1+5N9vgY4jIbireYEaVNRgOhsd/BnjkgmWfabhrbD6pbNfLHoNbtsNPdNNUYKp2LPTMIONQFWLuQ3ZQhZf5eQkHhbEDc8ogGM+Y7Xh+jdavvFiD7I1iU/keO8E23QuDdcWyf1Hw9nkrsdTC/vncqpoz8RXshWrBdRu1TU3tWNms0Jlf9hm3UNJRVf8y6dgfTHkzbMRKrF+vPWt6RiJkP/ICff+H41puzU3xtcWCZsZKqZ3qOFkEDgzx5cu1p2xbOnbkm39lhkFg1zGlTAmmbLPMEcA6olyMUOw32Nsva8uSpi8mjXT0YpTXHIOD8D5DNgGR5RJzl9hs4uADh0ZNww/5dWJ5UEFlRAqzUKzf0IxJP3Y/YfFcSvAmA+yZ51VGIhA9h/EneNUM2Ejucxjv/0n0QUEThFJpBVWYyx6OvMUCU1icfAR8BH4Hug4CKUXp6yysWFdneLO32FdyIiuKTcm6nFV/Zbhm7U8CKMW+gfvmJEOsUDlKo6NT2HNUXYthp6OVMQBq7HV3h8fIwx7ApJnM46TGQUg4qa//tFU3afsEoB7D4IAX/obi4ar0UPL736oBAMJS6XTQ4WNrOtlE9yr2JWH5BaOAZOW8XK0pOwaUlL5kUfesj4COQBgKm7wL6pcGZXyxbH24Ua96LdKZVUnDxPxDOL+GTSSM7QGr/RcX+dPTvOxmD6yZSWXUGF31Ob2Ojc87CWnXnoWDJXxEL3gDHeZRKKyqFnBsgsnGy2OlmFvpfpyL1WsDaD+XF93f60Z1m/KXRswBp2pBWgWBcyQAAEABJREFUBwu34Jp9f0EqY4sDjb9Mtrm0yTc/KvoDsg82GtaxL87QFUYlsQGkrOY3CBRNIy7/5VrRaBdJ6vi+Z/FdU0m5zrEpT09wiaSF8wLzdTwGQ9G7hXvrB7HqACfS2rlbPAcHWMSUZQTSLeRtEjIkxFzMl19NLm7XAW4evnsCASrQY8RSvcuvIgh1+mLhq1YiVHb+sR+SsxAIvsz2+rWMLPAKVF+AyG0U5TB092P/0coUFX3MvJmj77nw2cqv4ZEL6hiN+sD2xKF9ddcscKn8jumEElEq3kAk8hDSNZnw1ffqD5WDKeuqfk4CM4DH0+vn4Bhdx8euSQrWBfRQ2JrN8umanO/RyQg8IHGmuITWm1R7IRjs5c2Uqa9sCTjVGbVPjtPQniFQnRi3AJOY6iYQ/vGmwyRayJgKOxxPZ0egWASCSetOIiGOSfq4M+TQx8z9zObh6J5Hoyh0HUU1/csmLilSN6RPAPInLv6XwR7PfhT5b8xGabvmV4jseQJi1M9L/O8U2k0vyvelD8KyTsTPhdfBlhi60vzlxf6AdTAga8AYhdlw8RSCzrfmMaUN9XuDfdATgDhIZkQC9DoF9queG1D8TiYZeHnvph/jsuFmt2tmkvbCT1B8nVkgn9tHwEfAR6ALETC/NyWSejCs+lMXSuknnSkCV+6zHFNHP8wJyZkMeiIHNFxEVzMQ4mO6JEMBKUMsPBa23a7xDLzMJutuB9HdvVgg1gwU7clBtCdXep4L+i6Ho8TBg12pzI1Zv/Xg8L1WFwRUGyYQqfIbLEjnq4NUsfj+PgI+At0FgSE7UrkuVDR0F4Gb5LyENynHAQKR3ojVc/GF7PlNXLzj4jVwKlQuhqKCY4rrKf/f2lg41wAoB+QcNCxqbEWe5HlU1ELxKO3psOInIhyzUT7mS3SFiYZ3hcpBK5NWfRXBIBVUK128b3Yo+YH5fJxMycdRAiq0MAoB3Qn5vDhiFK/ma9bz566LSZV7oqxqIkprboXU3QGoUVCaIz/XQltDRZ7O5Rh4IhQn48P5/4S9UXYUsdv057hY+jFetgdtE17loiug0eykuSrSzrkL9LZYfgZ0TmJZTEVjXLDW+jRiTFZm0gi2GrFowGykMNY908J2xGwce+cHC77JPQL2cC78W0+wXfvCNTHFxggE9sMF1e1bLCsI/oZx70BrKAKVf+Kq8WnrgkygtK1Ed4VgPDj6SIRRfRkhPEPdA9vvhIv3v9DY+Zzf3wyFWQhOxhuEygTUzXHb7JcsjO/WrRBIZ5FPwgiFCrpVttojbCyh100+vm1PfJ0VRpzl4CAe3qYAivT0M97xePsmFsJfH4ApL22E0uoDUFp1IbYY9k+2ubcz4PWU4UhA2s4DNTHueBIqpwHh06mLfRT2+Bjy0djvhGEzj5MrN8eUqoM4rr4Ea9X9C1HnXxBcx7xSH8r5YGvZVc17MpuuTkZU/sI8PtfpP4vWWiYCjgJrOwh245wgDFOORD/k9Tnin9742946yvD/AZwf4GpkZ0Rr92Lf5DrWcfVwjdP3yAME5OX2CfGjGaS8076wfigfAR8BH4EuQODT2Nrs4FMPhoPyURdI5yfZUQSmjf0R5cWPI150KifHR8OB++/kJU1LN2W464B9hyX19nb09g3EjwLESzGxkOO3R2CL6VvRYXPz8CgsazYg8+FmRAo58NsLpc+v7cbiu68mCKizYRo5/QH2iOQLK2kE9ll8BHwEfATyDwENIR4K5J9cnSSRiBkTm2P3d4IT4P1aTiel3DKZK97sDcWxdGxSMtZC5GHgKw/lFLmb0wSJI4rnOM5/r7lzq/t1GO+vsezZPFEY64Ycqz1LZeRnK+1mwz5BLPQmQrG5CMi/AbkEguMA7EU7mPfCayOJw/x+BejlUCrr6q2D8VPRzZhW8hHMz2U1cnX48ss3wnQCLdPucKx+BJ2KgLJ+d2qC3S+xEGKsR12xgBEhWI8B1tmch/45I6t6CcPeD8UnrKNdITuTzykplqz1JlN4DpwoI5kR9mOqh6AgtHEyb08387vWgrPYtoUSfKrvIB7gu0g8pfMvfZ4zn2QddE5ngIav/xVchLFugD36F7qlRzbbfOAeyvuhawCR9SHxk9gviCuP79F9ERDUdonwqt+wnTk/o/bJ0T8DQqtm82IV26jF8E26CAQA4QIvsmhkNKLx2SvHm2VVn2Gt+o8QW/4mnGgVm9g72baUwsHRTLSYth9tc3L4Dt8D5AJoaCeErKPx8fz7UL7r1xBR5IMR7ImQzm2Rx+jiDxN5tORFluHb2TZeQFGPpB1J24u2OZl50BvMz1nQwHCEIicj9MJ/cXnxAhAgdLWZ+GwvBAL7UoxNaUlqxi0zsXTw+3xIk4S9j7wKlWoo/5KH4rzM+g2wW0N/lYTH3wCQBJT8d9J57ZJxq8PjELzTrrB+IB8BHwEfga5AIOhswXardSffWhIHEXGfVLXm9p/zDAEOaMypNtPGPoyCwl2hzkUUcCEgZjCHFEZYPn6FaOz/kPiNvhTcLbw9HibVcIAmY8kRpHWj2QjgUzfPdrmHIh8x/3Pcw6oAsi2swhHwzeqNgAMzAfLGQIQTC9Yvby7f10fAR6DnIZAfSp2c4CpRBKLxnESd3Ui/YH9+FpVyB3PJYBRgrYtQeGgbKwW/AuQQjmXOAbhIgMRc3ZzcsgSqSfKpFnnXof0T+ecguugrKs1KMWnmhji7pohuQptrEixavj0Vcmac1JCe6nuI6XOwJ0QyStwo6MS6h2HcxnwW0zkIRf2JE7m6nkIUYQjVb+uvtMB6fM/r0N1sRh1EeTlv4Xuiwyoy71KuRiiwASqKN0R5yQUwP4l15Zgfuv7rpFVSdq+77wg130L3EjoDaa36DJh9VjcEjLJcLMWg9bPXL6pG2T5X48fwjZhWcn1GtqLERnnx79gObMYCbNqRKyi62TiVPfkYYZfSDZvVc8HjNmjcfaFcsSkXsA4mjg19SDoCq5kHO/sA1paN7BG2AA+h9pcMfk6wMWQ6l959SwDZHSuN1mB+JP1TbprCmQ8eoHcxr0n6dMPE/sLCMZgyN1/6OSOUb7OGAIcxqeNSBJ0stwHyPcLRmzJqn0x7Vj7mepSXnI7y4hKEY2tCLS5eyuvMQmbjOwZoS1E6SZbzySjzgziO1Sxg1CIzhXwaxnbO9BWNVtfj8zqcB5gxp1nsLeS9kK85LYLIWYjrmqgo2RrlYy7HtJHvJTYvZXOzafMU23/PuYsky+O6jHIo7RqAFABolUf9DuaI//rl/VhWd8TUMTdg2ugPYI9fBNvmu0B+mFDf9SjIgbSNm0P0Z0DuheknkYHpFf2B3E8ThYW8JiMLoiMQDXLOafrKtixWWyffJc8RqONgqqZdMk4AK4G8xrCm1eXFJx8BHwEfgTxHQJwtKWFvWncyu1utaL07g+/TbRCwh69AxdhLgfivOUl+gYPbdHdMH4BIaHxG+XRjtm0LVmwcvTehdaMV9HgR7y74ntfskT3uew7qKhnhMlo3GsbefBzsmV6nE8A3PRgBe2YhLOyAVMbBe6lYfH8fAR+BHoZAyaYxqNNTT/5QLpgsQbDv8vx/a/o+xzNURhU/isuL56J89HzYu37Xxk7d5UOUj3kEU4uvoQLrKNptEbK2guMcC8jNAKo4FlrKqxuZE4HKYYXmoZdTgQtrdoA902vzols86bufX9WHSqaDOF5ZpyGQmDHR04iv+KThOdP/dY8wj+79lQgX3ONH4+R5ZvE908izzK/ma7j/MO+3tLHKxSEgljxBCUCcCYjFN8fhD/g6uOQgZeYa70vdVuJ9ZBauu3CrLuguonaZnFErzHoY9kxfEOe8aSkW1LH/8OTsfM+Ksd/i4wWTAT2V+fi68wXIYYrfhz8ArHuZAusp/7cmQYDt/rGYPHcDpGvKKodA4r8nXk1t6EeIy/O4Yf/0F93STeuCef1hyRkQKUoEUSxlv3wLbh9fl3jO9J8Gn2UQjw9WZDCc2BEwR1+T0acehIBq00lJ7pkSrQUcM5Zy5+kKH3NEfMXoZ6CB/Zj8XayzlJN37SUxpyFo++pQe9PsrHBq3iGyuxlJZD7F/zcELcec0H8C+j793GgAVCcwXP7/lJDiS+blXrTOo6r58p9+blmUtRGPHYyCvgPcOLrcvUGv/BvmbYuELMoapLgf4cg7iedM/pm6aKGSMZgTczVpUAXnS7IfJlUlxaSp40wa1nfMRwTkfQSC37ZPMmFxi38OiEclgm98BHwEfATyAwHzm5qA6Sx7eQpksU2rj/sbmzxB6mae5WPnwLFO42DpfkDSnNTLyWgoM0jHuPLU7z4IlmU2E6zhygOwH9b52GKdDTGlepPs2ZqN4ehXTNfs8OQlKQlx2R8rQusk9fUdez4CkfCmzGTSgT3dV5HIe6se/DsfAR+B1QKBCRKHqPtXdytB0BCc2gbF9kq3PL9pUHcsR30P3/Rpj/4F5lSk8BguPoSOobLnPCrynuHbcVeaCgaR50zEnX+gPnAIbM2dnicom8Bx9qQ8jeXH+QYafApX7r2CbplT+R7zWWZvYUC38Z4wb0dgzfqtydPFZP3AvP8V5cV/amMDoT9TuLl8Xw0llQ8tSK11oM4l2HzYjoBKC79sP/xcqxCLinpNvvjWlJ6DIOLBQNNjt7rGFjkcDy8ENVyecgtCkJj3IrFnBF3kKeYrsS5Ku7skG3R68/17fygA1MFyFmHhZ05eZst8jfnRt4+yXbmW8vUcfcbN5qftgjdA1X1jg8iGMEffp7NpzfRpEhgNke2Ik6F6xj0LlvktZZYC45KGTZvFqhvDfmnMSn7Rl4Fw9crnTG/C8glgPcVg9bTJyILgAEQWb57M03frxgioNSwN6ReitnB5Gnxdw1Ix4gfE4zeyjL5JAZS2faTBFYwjVT4FlhQgOMBCdzIi9bACi7MqsjofIhSY0ma8WV5yMnE8hy0f2xXXFMewDatAWXX6m6xco8qpx1uIyqQ2eQwXnwSV86Fw2wwpgHUAEL8UE6sH51TC9kYeLRkKtY5m8BAt+M6YF+tW2ONjaI9Zr/AzvtMnGTT5WEHMeBe7IxAwH1GSrSV1rwrVUvbV80mdV7AikqrBdMcmbH3HgZLZMeLO4/v4CPgI+AjkAwKbDhnITnJ9ihKgdSfFR+i7Vp07g+/T/RAQxfTRn1ChPIlK0mfSk193wqZDzIaRdNg9eAKbwtG9PBiM1zBAKtifPpp1a8nlABg//7vTpgjpHu7evk+PRiCgWzF/vWi9yZJXvBl8Xx8BH4EeiYBaVDCkyJlQuWYV9EvB1YneFwMq3gocgQPFjygIcGGzE0XrqqRscTB1xOf4ZMGtcIKnMu/XUJQltMlJJADBLhC5DvVVx2KGBpIzdsA18RW+swfTMP1QQ0QqM1EQeh3USqHdxnmIQd03rSnWZgE5IT9OAaCkyShQ+ybHj5cRhk+SeVN+AaxdOW48DxdUrZecJ0uug4riTGcZIN5KRg4Yj78AABAASURBVEFvBKUI3dHUDzbtwS8QuG0cacyVFiFYkD95lEAvytynUTi3SxQi37h5+u5EwLYtxKUfLBTyyZ1UV0BlEWYcnp8bAIzkZhOABGdA8bV57DE2OOJTmKOOwb4byYxaEByLeCj1/L3+5TUA3Q+KhnGC4mfeP41pxYuSxezilp6zPXMA2/HDABmABsN+Vx7HsDW/bXhsx/+LRywlCuYUAI96bW3GdPeEPSPcjhT8IPmIgP15IduoNL7Clu9RtLQ+H7PQIJMoYgM+4v1c2iht+0jr6zg28a6zigDU6YuQE2xfIjkIpdaabH8sz5hVFyHg/OzJk03P4Br/g2VNZ5RuHw4JxNqf7eTZOLtmIPm6F5k5UHjAwxC5DJDFSGaEZQVyBML6F+Td6ajKWWvBsSw3m6HBKPPyb+ww2uvkhgZOt/+nDI8yjodgTkB24wE2AZwDYDbNteLxLsCtmPPj0ZG05XCQPm/akXYhY8Pu5pdw2Zhl7ZbCHFcs8na7w+drQEcyeNesiPmaD18uHwEfgVUIBGUYFWnrrHJwuRN5Cdh6hYuv79ydEagY+T3Kiw8C9K2U2RCsASuwc0q+BIPLP3OCgKW/hshaLhxNzkW82ZTWKMCzbMUMEr2VWQCV+nIy088OiapnRCLi6Z9rz7gVT5mEcHqTkqmHMDjO1hz490qRm88Ri36Wgsf39hHwEeiZCFDpnjJjXICKmN9WTMnYKQzvPsB+RlO1aw7HA/OBnVevTZ9mcWj6qC8QHjMFgvOJgXefKBw/W7gOb9YckvV313/pIIj8gfGacRBF0Xpo7A6Yn3CiY7vpx97mI4V/M3zyxVyBMLFD8+MUAEqZjOzxMRR8+yzUupfeLmVUwxA5BAHr97ndzDDLgaVGKZxCb6S9KW9/jimEMncvGrYz60HALMjVewquMPlbw5OnMz0dMV+ChT2TVF2BOL725FntPXezEMBa1BX09YRCZBEEP0JEPfm62rN81AIIPuhqMbKavi0OHDGbu75wjVcwhO/wmBTtoQD1v2IfMJbxNCwKCl5CQWg2MnqvDJ0OxcI7Md7dyNqQlupH0OD/cNZm3m0NA7iSkbMw8jIczIWydidldPrTeT/E1k6t/yKjT90Agfqv+8KJp/4CW/AJ7Hb+vERnwXDl9ss5VjCLl3XtTlICjEO+9wwvCLL+rYl4tMCTrzM9G34KhO2QR6KC74DCRR4c2fWyt4kgFn8QKg+zbXTb7FkE0aPQyzkQ6Zy0kl0JOx6byWMInBvo44yMc0D+b0u9IDgO0eB+sO38WeOeMntDWPgjRAJoMKaPvx/mpL6G5/b9nzrmI8C60yNwkPX0D4i+0qYfITgrhfEIn09esiWmzD4zLWthPES8B9f5lLVUsgi+AuRz5qn9g1ejQFDzWyH6E3qO6Q0L+6VVJkzZATbpOVn3c+Ij0EMRMDvW4tie7Z1ZDHXPpIKKYHwKM8GEb3osAmqVcmDrPdlQ7QVx0mvf3YBat19/KI50884z9+0xpXr37MhkLU0RT68UipkUwTvobaHhaw+vaFTbfzqSV7z55nfezCFQaxe2jd6TYsUcFMRXj69k8+0d+fL4CHQ1ApZwzqg/pBDD4vxpWAqezvPeeOM+bNdStfUxiPU2Vtcxn8n31DE3U/FzFgDvPk+lD8cz1JnMSXoMJMO3g1QQDB3MgFvTNpK8AwkWoqyyuEN2rdoRfP/Mk/7SGHHbi8iaED2USszCtp554mJPiMBxruOY9QFKFKNNRmGIMxGD6g5ELk5pMCnatgMnaNqA5F9NGZ6Eld5cnBiICQ9QL5hw6D7/bJZwBwsocKov7gphvhRXFfJ2PSXGtNLLWxD5nmXE40th79Crhe/ighAcXQ8Qr/ZAAXyFWMr+kGxdTaJsA19BT9vQrAVcPMcTrK1RF4QDUNkDa65w76vOfDKMgOzH8BvRGjIL8ffDHrnEPKRt02E8v6ov1NmHrOvSGopB5GtY0cEd6uNMHxkNbgNL50OQfKMbPZjgTtDArjC6MD741M0RkAAX4Th28cqGIs46UO3Fkjd+4phNShyrtVOiaP8VUP2KoU0d5iUpCRwMQKFTgHww9jthyjyEogRo3UnlbWCpS912D9Yhn+klC2EFrgTkf3A3g6CYiHjhCHeWPPaxx/wAFeYRVWj4KDqZsEPpeA6ie+7Ea9eT2WzhWL8BhPUfTeZ9xHXNDvcjZVVj2FV8BqjHBl9ZH6g7GubjtqbUeeVAX+O8dicaCbWuT8/CKPGLulPmvGXVT2DF3XdPegdu5uu8A4jZLY0eYszxlX9Mr0yw7Ah2gG98BHwE8huB+hfWAKwDKaSp37y4kOBjxMQMIl0YfOcegYDj1ADyHLyMWGGorEPFcNCLzfi52l69JnBAtaGrf755OI6Nk/9hviTqoGTOjykjWGu5mfikZMsJg5Ni4gyh0sz6JSdp51OkRnkdDI1mGfWe3CQmR2KO6PPeNJNPefNl8RHwEcgeAsGCb7lQbr7S8Y7TwfZ587VEUV1vCus95lNdwsXK1PliRD2X2N+FAndTGfkI8xinTU7CnkJ0Rzix/XH9x9lRop5fPRQW/sQEw7QNJNiZKVHxKJVARy1uYBxD4G5MPvZGbTD1cdHuceTeZzoVsghMhsL8TqeTPEGOayxcjbdq9shZHSzQ76kg/DZ5+o2uit4IWOthSJ/2j50bo+r8C+sC4iZ/3l/KC/rAcobglFfzI49xHcD3UuCJl+AzxKye9MGOZ3bb6Uk9r2zJ9sdrHhSFyse0C9uZRucGc5zP4CDSuYnmOLWinc0m86eYins9FWwGyH44u4bvFG1N7wHrwZGDAWmqwy8hVvg0MjRpsQedzVhe9idvUx01aR5KN+ohJAv9nJzPuJPnkx6kQWwfjkDtq31571N3R0DUfNDUJ0U2vkMozMXjFFz54C34FooUGws9BN31wzoY/S3gsXjJ8IJhrO7ecwKydQotmx9iOmvQCq0bRaiOeguY5TLmcwuWBfepIz8GrAsZk3sZEmwNJ16B86vye/zMTCSl8HNvQQKX0u9TWheSXelxEcpq1uG1aykW2gQiB7EtN2WnSZY9IWZekI1+RO8EpA/cjDAltU7Bphts0pzFav7g3+cxAgoOXvEGfujd8Z3A4fhHEHkXkM5vnOAbHwEfAR+BNBCwem0BwV6enGa3rIN3UFjf8XbRMyHfs8sRKCxezMmG+b19d2U31Ax1emHZwKYJu5vYyd3PpIJc5bzknnnqKtYIrLXNqA5Lp9YPKePQIiokUnLlhkF0zZQRW07PV5Se/WJ/WLIvFN5fyQq+obL7bdjjYylx8xl8BHwEeh4C5jh2lXSU49sAu/XKCwDqdCjlMAo2XlxI8DoKemVhM7xL/N3F2Ta/JezcQXG9x7+KvhxL74nvF3j3GYwoJZmvEYPW4VBlmUnJnSsGAWRLBLEbznyyfWM9dJIpHz0fcZQytSpapU1GG0CdckR3H57Ms8Nui+o5LhKjLHUfCwjClGELFPULdzi9roigIPYDVM0Xie55BHqRZ0P0C3V9Hs3XWJYOgfdX6xHK+yqc5e1fZMFqYHo7a7F9S9UeLYelb6PXmPZ/sdqpUOpzlPckOM4rsG3t1KRzlZgtDjT0MkSrWa7d5vGmr/o1QtENk4oR0N/wXf+q0W8F76/DZcMzrR+NwVNcLDmI8W+RgiuX3qafG49AnVnMymU6fty5RiCx+VJLUiYjeANDxX2DDPLILA9Rx6B/goOp+OabSMaSTZgQh4Teg8L74xORIYiHBmccfy4C9OnTG5C1KbMFN6OyAI5Q92J3zRrb1FGvAHIGF5w/R3LDdgXFCOIanD+36XST5Jz56GoT12DdixCZSPHcyo55P/sCzlWwawaSr2voH/NCUIxn4mZ8YnDnbReQyLqQ6BHNP44zAHWBJH6SGSMg+AWwZuLm4dGMw7YOYH5bRp0nAfWaKLUO5T/7CPgI+Ah0HgLq/IltVH/PBIUDR8d5HqZN82T0Pbs9AgnlgfMTRFZ450XCCC5JoeBziaHf94dArKajBV2Y8s1ZQ1SonNRxqRy3gXSzqCObN3vo3FvVFGkrELd+7lyhuiC1osAwpmqUYMKrO4m8hKD1njuD7+Mj4CPQ4xGI6wtUti9Jkc9NgAIuRqXg6gzvgBhFSV/PpBzMgNnc4Mm0OniKwqKiUfUV5pYdIP8nI4FQCbU1HFk/mXdGbtE5QwHnKIgEGsP9CME0mC9MVE9hWcuOhf4ZKq81ppHs0g+Qg9Gvd34ohuFhip5/n2O0S6Ewi/AujLIjYJ0Pu2qYC0P7nef/aL68fYPpe/0ckECs1HWv/VLkOuQyiLwNqNcCbxAWtkCfZVTi51qcFPFv0L8QKhuSq4g2OSk438GruGLvFHOe5MFXH1fZlu3ORt751QVAvAZmHoluYCrGfoupJc/DXFmou4HE6Yk4bddfIIF7ITBtUrIwAshOnLuUtD6yuOFrTvYx4NIVEuY1BNvz9X8irPe/0ue5yGcdSaYQrXkDP7OMTaTNTv/W1FeKdQbj/CyRRvJ/A+j857zf6EYhffJA4IcF7NdZrj1Y6BWBOv/BKVlY52FkOadrRteyjXoR08bMwQNczG9PgorPIc5nSJxY6BKBo4PYXnCOomwbXHg6yzmCwRAZCuGfW5qCN2FZXnXaLWR23EUUoeeqoFYpI3Q78caCYC8EY2WY+GZv8nUvssfH8NH8xyFSxsa5zkV403b/FtH4eV3286VfRQazbJufrFmjUcYlULkOTe1/1q5yFuP/H63SJiG1iNNvUB/cuMmTDk23/jW/EdD5CAcqsyZjvP/j6MjRLVkTxI/IR8BHwEegGQK2baG0Zh+I/r6Zq8utfoZ47fMunr5zT0MgEKzjwMltsNeQW01jktDA2fK/+W0v4I8cJElLj7x/EmKyMy6o3L5Dkprj3FJHsHtqllxw8J0KdvOM2Qx7VX705OnunuZYTLEuBjgh9s7LUjiYDfN7ad58vq+PgI9AT0bguwgXxVDtmUXBWqiP7eDJ0xme5utyYDyT8uiD9XMUhJ4jj08GgaXLfgHELJR7Le4CkDUh2AwdMearZcQPZjxNX0U6HHu8ACd4HSpG34yKkuzZ8pLrAecfFNfloweOCYDhQGgkefKbbNvBwADroN4OyGIkNwFA90EEx2ZdIWuU8+rU8L39kDzpJldncwSdLZueutXVKIQtaxYg3j8Hp9gakfgQdLUJF/Xj+9gYCrfNyso58IcQ6y0YhX5Xy5uv6c/QADE8iBi5b6RIyG7VsGx/nrj1/3UhAqII8l1AazyE6AXI77DOZqs2AtrUC8E5FEDjAoYuh+B22DvXIlOTit/0cxo+he1xU3/pMK37UFF8NbLZx5m4po6+kXX8Rork0s/RRzAe/ft2/JQ/RuVTFyCgHKtoYFemvB6tO6l+gnCBV71wD9tdfcIjqLOxOJ5X9xMEBGzbnV1x+ANWF2dTYMmGbBfWgptRtktw5iJU37Vq7QGCAAAQAElEQVQfo9gcczrOM+wb76aobjrTIASHonDZBDToP8najciMa+PBB/k+7qfUbuXHjA+OwaD6g5t//U7+3JOp94iPhsCc/CGNCb6KaGR61vuRHUb/ne/6NqbB+sT/yUiwEcvv3k04dHVlSiai75YUAXkc9shUX3EkDZnUcfp2CyH4b1I/39FHwEfAR6CrEKjfYxcgfisgVIjB24h1N67cZ7k3k+/bYxBQLYBIQYr8LMPypZ5fzCQNH/nFDNK2TOq3yvFD3j7LAecznWfB9GDSZdIupFgfQdmraWDnwuXtLJb5Wty7LglGwZ5H5Yx3VFn3LZu9LuPcitadzGkgv4S/dGfo5j7mqOMi5y/MxeG0qegLWIF0jv5OFY/v7yPgI9CdEfjXuHog8B9mYRmtG4XYrxaz/yh0Y+gU90jl2lBnjEdaMSo5ZsAeMd+DZ/XyGrQfFw6UeGiKMY+GIZb5mrD9+Gyy3nqAHACgNy1Jv4cEHkW4NkfKzvpH+L7fYEJu1A9wjofZGOfGkS/u54yuRbjoH1C9jyLFaZOQ9OGi71koWPYbNGyGScLTTqc1g+9DhQvkXuGlPyAHwiyAoRuaZfiEUnvnUWVdiLUt+bqWnMC6FGALCP9405aklu/rOSxe/E1bP99lJQJv1+wE6L4rn5Pe6CI63w3/pEDCkAdkdNkSuI6SeOm0R6L38oPJ00DR/dfhe/59w4Mo68ZbiFvVYIOJDE1K9i2GbQqRE5vxfc70uLjCdJs5Zu027jwMxVse8fVCXE7F+VWrNkR4MPteeYZAaeWafL9so4T9q6ts9bCsx7BsEcdyrjw9z8MWB44+yfr2TYrMjcH6G6+Rgie33vbMAsTj23AMN9A1IZFP4WBmXvQ100sWworfQHmfhsJlzAnOufQc1C0eRz5xzVe+epgTZazgVRTvRcrvkkcZym7ifESDI5HtcTUTdqVLXuwNyBEAmuZddVD5J64Y/x2ybSZInHXI9IcvM2qHNglZpv/YD3WFZuwJKwmH75RvCJiK6+g9WRfLsv7OOF0KCn188hHwEeihCMxKXe+74quDSZUbcxB8ITuyYWkA/yUHY/9Jg89n6SkIqPaGNimfk2bKof8S3LA/Fz2S+hvHttYsrsLaBxD3nb2AUVZcC1gnICqdZ0M4nnJdD8AosZDUiJiB5jhEraFJ/dNxDMpCYvdOCta1Eas1O9lTsGXZW61jGGOK8ao+k5WfSGJCKclSTcmTbYZ+/SdAMIlWUkateBrBER7HDaeMwWfwEfAR6AkImHGcooZt++se2QnAHLtbF0xn3IWcGUt25uLc+h7xf4gAuIAqnd/+egjVpV62OLBkGSBReBmRABWT7d+8Z88MIhAbBSgX3GD6YjOHmMckZ1PZGeM1+1Sx5/eAXE/Fntd4rhhF2kUnEyEzYw//CbFgOfPzGAMqbRKSIYBzCSJz9kri2X4nswHBce4CZDHcjRlbHIot1tnQnSWPfcyRxI7zb0roMVZGEdvC/bt0k4PZYGHJliwH61HW5KTOJ6xlD2Y8l0keW890NQuiccdsivWat7Geyd0IRYxivGfi0B1ztVwqofKkq+hi6qmeidK5a8Ms2mjdHoA0bgLXWoDzvejSr5G58Q7xj3khOPJbiCQWSaCI8v5hhAs+8g7YAd/lQ+ZDnQeZFvPlEo+lo2HpCBffnuosKJ19FEqrjoQ9s0+3zKQpuxLYFYLdWGZDHnn4kPrMx1bLn3v5uehz4nMvsWFbzf/JSLApCuq7tvwvxwBYsgPEKkJyE4PIU6gLeG3mSR4yV66Xjv0EgdBULoC/6ZqEYmtYzsUonbOFK0/eenAu+ME378GxKiBI3kYLfaA7QTEFsdkbdFpWIiEzX2K/ZdJPpFqDFZHHE3e5+BeKzOd7Zvzi8vM6ajHZnWFFR2LGjIB54LNPeY7A0yh4wexszq6YgVHvIuWO8Owm6cfmI+AjkAcI2LaTWopOXuSaVLkGAnI2B8l7UrY0+iY5F9PGuh93w0h86kEImEV6MwkQcLECboaTdSx082xwT/K/b9FmLHdjacNJfJucOMh0ZqN89HxcXryg06zNtBz5H1S9TwGAUwwJpjrBoCkvSa7hOubfKPST+K10GgSVwzr1uLDzZg7h8Pn0lRK43liPuHp1d48ps3YHnMmAmB288DQKtonONTALQ/CNj4CPwGqPQMHQLwB9gpYLxUnRoIpEN6Nyazhsc9RuUp7cOl7/cQH7FvN1uVsfXMc+8AEEij7OrSA9NHZFPQLCvqHd+RsAWIcCMhgNhu/DmoHQqG8bHnP0PzaI/bq85hF7X6gzERfM8/q6ziN4J3tdPvIbSIH5vc5q95RlI4hzHSbV7OjO0w6fwjVq2AY85BnSLHypHtOlC+StBVQVlNbsgbKqizDprTVae7d4/mTgK8zjjBZurR9Ed8emQzdp7dxpz1tvXAA4VAxLv6RpKhRW4O/4cEH29X5JE+ymjiHZGyL7e0qv+BASuho98ev/yTW/4uLoIyirPr/b/Ua82awD/SdL+s+u709keyAyASuqhwByEICm+vIZAsGnceXeK+iWIaVg/2KF0QWwn0OjDko/p4yPIhc/NdAkyg2bRWAFn4fAY2xjrU2JDsUFz3WPfq4pbx25Xly5HsS6ANCtgL5OR6LqurBv9IOlRzL99WmTk6IWkEdR0Ot1mA276EHm5HkhlFVOQ1nVLJTO3ilpzm4eHoX2uo3v+bOk/g2OhfQ/AV15VH24aEOKMpxyNLYNfGpOim+A2B1ItG3NPbrw3pSnwIjXATmd7dgCJDMC/slIWPGbMPmlQclY8trN/BRAQX01HJnIPC5JKqtIACJ7wLEu75QTwya+2RuCiYBw3gQWGV0O0Wtw7Xj3zalk6xDZ42NQeRpw3PsRxWC2qb9FzVp9kxfiDkngB84yAnEErMlIa8Euw5RtcaB6JUN1046VkvvkI+AjkDkC5mueVKGU041UPNnyN8eKByy2RXoGoyygTUVPIBR5JBWT759lBGyzOKCS5VjTi67vQA6k1CwQePFTISDeSrPWoU2eJLgjB0U7tvZa+awaB+RlFJR8gK4w00abdOexv6YcbgJIfzh6CMxCihuLp/vjdYAyDagHW5B+eyK6cFdec0/m3RQET4XKYM/EVKMYtL771ySegdvh6Ujn1AGT/9Kq8dDAPwHZku8nVboO4JyFirG5XZSBb3wEfAS6DQL2RnWwnBlQvEqZXdp3trGC/YDdBpKn8+m7b7egfONcEjYyz6OSfAbs4ezjXbhWR2fzNbGjHBsh7Jl9wQqOD6ig9ORy94wUcDHGGU+GBr2R4nWEI4/Dph4BOTSX/2op+7SbmQKV5PyfjCwZjdDyMcm88tKtfNevqQy8hLLNp3WjzRBwLoA9b003hozd7W0iLAPToPgC7obvV4/HZuts687SyT4XVm8OiU/h+PdEyLK1PVN/gHnU0EXMY/KvwRKBZQgsPRO2Mq8Jh879V7vcbD7YkwpiSZKwAvosFq91O4xiOwmD70QE7DkbQp0/8T3351NyUiyFg6sQfOrL5Azd2NXmXNxyjmYOxhCDrzDI/BQMn7oTxQtfZj18DmwMkdwEIdaRCGJ/QHYGEjqpGCA1qF3yDq+sK8jMeHGbxUqx9mG93KyRLQrIC4gVvMlrdtNCC6Oo7/0u6/3zdK2nTUIahkgJrMJfJfHseU5mXBPTsyEYyDLwSrcc95k6Gq3dj22Q2VAS8HhJbyJu3dYt8+iRKXoJ1qwrZvt0IlSDcAZ8TrfkVL7Tl+SrIF99cgbjqgcgunikuet0a96lRvegjBu6pB2Fyl8xdbf3Xfy7ztnmGL28eC7bl8m0dS6CcCwk42BFJ+LMj9PRwbtE00XONhe/pxU/yfbiUkrAdpv/2xL7E/kNipxTcbh61ce2ITN1CS8dC8X4lcEEM/FD0f9WPufqpnwM65F1J8wm0mRpCITO49E3tC1fOG99ymME9CksWuQxkemg6Jb1JguKx070DsbvB/cR8BHwEXBDQFUwefZWiNb9lwOTE5Dom5DKfA8Rs6OfE8FUrL5/VhGo3WNblFbvCXNaA2AGEugcw3JixY4EZFd4m18gMQ503Zna+uzWD5DDWP7CcDMiP3ERNvfKbrf0jbvlPAgRr+NbyaW/xfc/DOFN5mTbDtT6lAG9lWWKLQDrZNjVg8mbO0ooAPYaR5n+AIH3YF2kGues575AgG5o7M8LEd3jMOrGzNF4G6SZg6dRXnJ/mrw+m4+Aj8DqgkBwJtt1MT/79otLli0q3/akstt84SIuPLlxNhtAYR3MPni9pAkIzIam/8OHX+duLpw04W7guEH/QohswneX6ojcLxDHV+3KkTl9Cc4xgDQtRteyR/477PG5+5oFzUy4aDafXqFV2rakCMEJHNs4Lm3rn48uwaIaLsZfTdF+ok1GRj+3P2J1Z8GeOSAZQ7vcCmKfQXAdPH8KgAvk4pyGc2c2vW90mZk8ey2W7TJAtoPgKnzyjfvXTWg0FSO/Z1tyGZ/c2jp6WYcjWs35BOcWfOo0MnUpEDiFeRmaJE1Tvt9DIDAZN2zmsRCSJOTq5GSzXEYdLg7KGOLo1ldxoUMfQDz6OGzObXoUPiyzkT13Y5b+yPnBh9D4W7DFQXczO++8jG3gI6yrP3iIviP9LwZ0nQYeXcT24D+4cp/lDc+Z/ffkHljL8YfuQ51440lr+h0gT+CynZcg1+bK7ZezLD/GZLw3hYnsjbNrisjXg0kFWwwtAazj+C4+BGQeupsxm8ti+4wFnGv4Xr3eF9+3MxXTR3ltyutuuW+Qt2zWusz76RAxp/Y8iOnbpTidM/QEIGwPEEMyI1IAh4u3DfrHZBy5c6vbZ30I3HRREZbTe1DQ/57cCZCFmJeGHmYsf2f7uZxXNzoJfb4/rtu2MaGg2Sx8JzPH/p//25LF9zgZm9UcnrNTc/5ixut6AtMJNSb/C/WY92PYzm4yNbJl6xKYwbTf8YhtIBycbnkw+F5djYAqK6ncjUEvu+1m6biEoX4/MpKHaHuW8pwZ8slHwEcgjxGwawZiypzfwbJuo5T70qZDpp26C9GCV9Nh9nmyjEBQfkWFw50ISAUurCrOcuzu0ZXNGcUB9hR3hpU+HyA41ixir3RoddP2sT68JQfEu7f1aOHyPiKLa1q4dPZD0GH6+p5nsiJrcYJkvgrxZHP1LIi+z4Gj2RDorkwSckAnIKJ/oWItd2PILYZtyndeSlk3pnUn1XoOZm93Z+hmPjM0gAtrtkVkPvMubBslzQ0d8gksTIeIdrMc++L6CPgI5BoBm4sg4cLHoXIDk0quiBAZBid+Ii77MNViMqPIEplNoNHlO0BwKER6tYlVUUuZr0Oo8CH/i9g26ACFfYZw/DKC2LlvYIQYBWUViqJUNieJI5VT/76jIDhwFZvOQiD4+KrnHN8tjn7LcQ2Vw0YnkiQtMyYR7AJLdyUWkoQj/5zs4StQUHQnFHdRuAhtMupL/z8ivhvJMwAAEABJREFUEpyQtSP57fFUrof/y3ifYoJuuiUqLuUgFIZ+h8TmHHJ2BZU+vzbE+isghyDxEyaB/6bdBqiaxbR7AaygTUZr0fGPOHuOWZzgbSdRv757sy79nqkJbUtSfMn3UoFA/dstPfynlQiYr7SjgRP4fARtIW0yctgOsI0quBpXjPs+GUO3dps0x5wgcTnz0B8iL6Eg3P6TXRhJl9EEiSMunNeK2bQfd5GjgHnkQqI0bAJXzEZ4wBwX3lTO7v5mw3lASmA+MhAIjHGEafXifLiT5lQRoV6LfatJO7k15f0A9I6um9y7h7iWzRkGR85mbvqzPaxCqL571eHEmLZqNBCfzjysTetGzJdMRrjE9MVuPN3T3Ywb1JpA4cfzHf6AcNj0x3z0oPCIHyHxG8nxBiDJdVAiuyEke8HoSdBJxuTFil8AyOZoaziekmcRtsphb+M2jmsbqitcbhi5BHHcCBGO3cVN1oEQ/AW94nvCnhlEdzP2iKVwEqcJm9NU+G6SZECxJovXJPTvW5L1cmTbFnoHSwBh/YcxLMdSg0isGrbwHrk3iQ2wzs2sd16bSPfOnfI291ns+SkI5rIxfJlK9twVmkSD5cwimOaIYV588hHwEfARyCUCKiirPBDR+B2Acy3MhAvpmETn+SIgf8NlwxfDN52PgAonxjKECf8RDv6JsqobUDZrPT7njia/OArqXM8EBtKmIHkSdqKcILlJ4mo5f4RI447/JP7GyXHuzckXBybudK09vg6a2CzjHcLS02HP7efN5OJrj/sZ0Efo6/ZVGr0SVADRMxHZoyInCuKymnU4eL0aqmMpj/c4VWQONGrGMAnBuvU/c6rCG1VTuNhxJwTnMS/e5ZIMjfQzhG1pIPpK47N/8RHwEfARaImAWXQMB65h23oL29ZkCncLIvtj6c/Hsd1lX98yeE6eJlcNgAb+xLi3oW2dJhU4cj2ielMPPB6V2e0gma/MgN35znZhTK2xo1MTOfN59xQwK/MvGc+v6gtH/sLwjQuluhSit8GmMpGOnULXjK6jLuRZqHjpKcyY4de45MX+nSIT1ELQEnTE2MN/QgxXQkGFrGtEQ+lThs2G7c5rdii0y3yKfy3TfcsjwsFsA85BfW320vVIrI3XxOrBkKJLITiBcvK9sx0IPWtOAmnDmtRhWslPCFnXMQ/m2NXkbZ1iP/Ry9us0Rffkml9R1lLaJHMZXUxZ/4+Lm4/CNps0yOVTSwTM4v+atecAwrGxeiys6acsMxcjsMv75FX0JHP+B31hxf/CerETy8sv7McrcTEXPbprHotiC6DyKMVn+ed/b1oGCzfB3maZN5ubr4f7uv3Yb+jvieuARq5a3t8Be4d05GoM0sHL5cVLGcM/aZfQupDuAA0c4OKZufOKRQKxJPOAOQpxdk0RNH4UzEKv+QhSrKe7VXtoFv9LZ40kplOhGO6KknIcBedqhCIPwM5QZ4VuYKKRHSA4kZKuwfHi08CKz3jvTQaH0KC5HG9eBjg/JGUWrA1Hj8drczZK6p9tRzPGjtT+jnkwH9YEWkWvfH4DTnw63v3mc97nmEQQi3Ssrha+8AViOg1w3nARVvjeNofKhYiEzXgFnWpM6qGOtEeiDadEBZhHvJtUdpOGYhuWs1K8/rLZTJeUrV2OS3cbCMUhxNBsMAWN6T+ewpLe3/C+80jxX0Bf90hwDcvD0/fqWgRWsBA9hqljv8i5GOGSN1hYubCGaM7T8hPwEfARWP0QUBXYlRujtGoyyqrfBuQ/tGYSM5idVJr9kBkQ9voTzG/cwDddgoCoGfCapM3O0E15cxoQmIeyyquyvhHA7D6dPPu3sIJ3s3+isoOpedNPWBK425OltefklwYBYgb28DBLUBC/y8O/87xqIndT4bPQO0FrCGKx33vzuPly8Lxk6aNM4303jlXu0gci5yJSdydK53oo41aFSOuutJoKDudljn/2gyCUIkwd1HkQheO+SsGXv95G4TGlem/WoTu50EVFNy6ksDvQmq89eElJMWL1bwTXuIWKkrqU3D6Dj4CPwOqLgFm4XR4sJwBUECDOa2vqBceZhtLKTlj447jQwrls539HIVq29YnfMHTux9JABRoU42TJQwrEiiiVGQ/x0skUm7cZVMuYaj9aN4rR4yms0Xd2uz4mCFi78f3swTgaSSpZajp7w50iDNM3VrKvc9NThGHJPohaW1FQoc019YdjcdGog8lcXrwA4cJTOA960zUmkfWpgOY4q9psknFlS9vDFgcVI18m/2Sm67aYRgxlI2J6K8qq3BcyGEnWyRz7X+BMp2zHMu46WFqKqWNezaz8cixrj/4UTuw8QJKfnCUylPXnbNQGt0auTWJDg8OxnbhhOQPhNW5ELhY3s5Y3h3N1ZbnIWoTpR2S+wlyzfhpEpmLVT5GgjVH9DE7415hW/DJslnN01CztxRgGQcC8o2uN2QAR/Pk4iBxJQSyofAIEXuZz07yczt2MzGYXiT7HevhRGpLXcJ5jfg4mDdYkLF5OvfqOh2IcAIExqs/g4wXs59iOmOfOsj8WvUQszKal5CmKFNDjVJjN4rzpMAXqwkxvSIfjyUYE5uvZQmc8RP7Etr8/5XoZofrXMoq6KChQJ5BRmGwyT64ZTp3VTXB0LKN1kSNxkhEX/+N/h/mwg4wdplDQ9GV9OhxPNiKYPo/vLn41IL/ie1wOse5nPmNIx9jbRPDJ/IdgFqAVkSRBAvTbA5ZzMq6uMWPvJCxZdIpUsjxiCiC90cbojwD7o4L/zUn7ZCJ0xOiakIDpj9ofiW07uKz4LVjWGYzEbaORBTGbV5yb0N6PiRh5O2kNiNOxcvzAhDimjprDdsDMjZKVIUDE1M1xsKIVsN8JIytGBUVB6u5kd0YXoiXp+5DgU7h5uNu8hTw5oJ/e/wmQB4GkdYjOyIMBTUIM/18bBBSvwwk+AnTC4MMWB/HgrQC6rxKdwvvkI+AjkCYCdlpfGIRx3swhGdnS2UMxac6GmDzzV5hUuTPK5hyYWPS/sPolROVDCCoooVG2mA5XeJ8OOWR6hwPqPVC+05e873xSrJERDpni5sZvFsE7P7fuKapIK0+Lz4MBOQew3uci5gxMqtmX735j2jVgds8iAzNjRgBG+WcWgetDd8Oy7mHojWmF1ovq6XkGzBFXvHGjNu5WdCLdGgdqvEtG5sjkbE3SksWfiduLrLdi3ZgyiIMTMfHNJBOWlCGBG/avh+q15KylTUVBMvwGiD7Len4kTP1XTfWuGKQVmQH4lDlbMo4KiPMiBMNoveNRmL9XoMHnYHMMg040ogMybg/sueti0sxNcWHV1iirHkt7NMqqbkeR8wnxfgYNG1E4+UlMTJCmqWW7eBciyy+AmTinGSh9thTjT0UAscVrZ4yFW3uXrrs9c8308+Bz+gj4CLRA4LqR38MJlLHdoaJNTd/ZwhtilF1yNSZXjkCujto0iqXSSiqh5HwmHqRtTiv4cAdqdXLGfToDdippwHzpke5mreyIZo4rnjx7K2j9fYxwA9rkZHpIxRzO76/AedsvT87k4WpXDWN/fBIgDUpHRS1U/odly9wUh8iZMWMw0UoIPPIhm0ADR2Pimw3ydkiY+j4Mvh6tG3HsG21dbt14vd3NSQAauID18WtAFEmNcJFG78lanRSmUzHmecRxONP9CqasJE0Xa9P/f5hcdWxiTA/1HpcljyM9V7MANLlyc1hyFSB/YLrfEY7DMLXkeRh5kbFRTBv/ARznt1DMZegobTNiXkR2QlAuwfkcnzXzyeqtPW9NhON/YV4OZryty8wSynYDQkV/4RjObTMGg3UxLS4IwbLW4jsxC5CdJ4yZA5dWb4NY3T+IH/sLBJMmbn4KDHgRsA7GtBEfgczIiukLRiXoamNwGBw5gLL8maKsQau8n8UF0u95372pfLevWbbuZNmKe2SEfY7+jXUk4sHj6eXqeT77OejFbGOayjbTsG7loh6vrqFy42EWigJyHyNv1VbRZRVtgaieDTNvXuXWvruCgCnbyetUU4wSMCeWGL4ml+xf7RlhRPbcAxZuYOQbsE1kWbDuhen36ZA2/fxtfwgGp82fFUYV2DUDqb+YgIA+wii3pwxueP1A/+mIrLiCecteex+LczzCmLuWiAPr0tK6f0FkBEUxGLwFJ/4279Mns4BbMeY2qJzBNuFbtB6bCEIQnI4f4sfDnlmYfsQZcCbG2Jz/iHURVDZsEdLIk9hoJiehfPSjsLmo3oKhnQ+OY/Qapm13iUAsxLOxuYVjv0tHz2MiZbRLaJORQKQY0dhDsKnTV5VkTJm7yTpQGP1/8qCCABDoeFk248Vw/BlALmd6yXWZAgFwGCKL2M68zPF14plO7SR7zhpw5A8wG3YbonCgUo31Ap379b9J++aTY7DkVd66nprVcZAZu09ZRkARheBuTB/1RZZjdo9u2sj32ND+y53B9/ER8BHoMQjYduq2X7A9wqHbMrJi3Y5A/EFYwf+xD68B4o+yLauAYhdi5z3JIEMSitLtBWj0aEwb+x7vu4rOyAiHTHFz468PG9y6Ks9t0w1YTlvHJhfhgrMcjoDzGCw8xvd/FaLVp3FReA+YxV27ejAH623LgJnETpm1EcpmjcVbw/4IK3BjQrFh4QjG7D5QpOdKEnkMoegjK5+T37R0tV8bxnJpvjxs6d78SXUhlYf/bO7U5fdi3UsZfqF1J8HGKFiytztDCp9PvnsMIleTK/UX5QL+yXYQ+ReESpOy6hMxuWY3vs/1YCZRjCQpmS96pszakuXjAA7Az4XG74fA/M5a/6T8rR1Fv4eFazFt1IfobCNyJtzqrJt7NH4PrNATcFDDsRYVlXonxT6WeR7Ga3toIaDXYdnSP3XZz1MIBkKdGzPGwg2jdN3rQ1QEZmvXdnug98P4CHRzBKaP/gROYDLbH9MOtVWQiGwFkWl4fU5xdjcBqCQW26KxC9hfTIUg1BJJoULK+Tti0QtxNRcGWnrm4ZN6LRJnX96Jz/TGpkP3Q8Ay45IdXRMwSkrBq4jrme3SJRjFakQO4/spZj9jJdIR/Y73b+H6/Tp/YcQIEI+9zgtl4P/kZFHeCR0a+zTFWxfpBdXCpsck10LWnzWTuLfPKVw/iwEvY5pmgYC3SUjQWCcri2EWBJOwZOYkiunFz/CdnkbcXuV4OJ40vPmJLNHrEJAKju1GZiftVimZn5qI7PlbWHIL5TiS9j1YOB1Ti19oxZn547SSjxCLn0BszRxheasIhO6/RjB6A0rnbsd7aeXfsUczt4nU2mzrzgKEcyQ0GGUOgc85vyhHeMBE2MNXNHjk6f/+VojY9IeI1TkSqsD8FFgsfBIEt3Hc/Hum27RAy9tmZOZpIrfAif6JCzLvNvPp+G1kBfsnYTsgLuXC4GEVYlihi3/HRcA/5oUQLTiU+E9jbGZDPKusLkEcDyO9jzmQ98apfRCKN5DcKJ1rmH+zaMXbdlHyQPbMPhyBnM20t4/LJCcAABAASURBVGnG8Dm06OVmz517Gyk0m5UWpEj0BER+GU9MOljuQmHqW9b0TEudrWF/nrzueQZMx5P13GzAiKxzCgv17QzRUL6Br2GFn+VzZhQMcE6KoswCdYA7oceoGY+IcxVjuYt2KG0yMif1VdLjLKxfdFnW5+xq9YZIiPEnJzGYDGirf0vOnbmr0SlPnr0lImJwOLgxgghUnsbSod76qkbmlheOTX4uYHlQs0HYtOmtdY+9EJByRAJncjziXX5bRpz66XANYLP19iae15C5GMI/3jRShE8zqa86DtOKHwUrD7JlxOrHuMU9OukDkX7u/pn4EN/QOrey/biJocxPj/CSlMYi6lyOspe2RaYfdCWLTtEXIL7J/IybKscYTh9z22Fr+sZB61dAcAfjqqV1IT0E0cg0TKnanHh44O8SvMk56hzI4nAI89c0RqqHFX8Rp3Ty1/8JeUQRc77g7ce0SalJyKSevmOXIfAyYrVmB2DnChAOmIbg7c5N1E/NR8BHoPMR2C2dtn8w5do/Q7s3+amUFA6CNcz7jlCMgZ9lB3ouwrO8fquSbDmn4UwhUyw6zm/piUw3f8hRTUOYIMziAfR48nIR2bqZg+V/Iqp3IBa8C6VVV6O0+mJMNj8bUHkzoou4AGH9Cxq4BQpOHvRwhstkkPs6B21XwR5Xz3Ae1MortvxQpNopLngahbFUE/FWEef4MSjzidPznqmoDqD/fjCTU95kTGYHdrCggrj+M/2wifq+PzG9HpbeAgRux+ZD70JZ1U20V6C00kZZ9aUoq7oGpVW3IFp/J9/5PwG5GcCltNvSptMukQ2cWKKMSttHwRE3Ot+0oz3QscRmc4raj1fhtSP0E8vAFYhSIW9ObOhITJ5hNZWcRjG6D6PYv1OthSMQXTySafrkI+Aj0F4EEpsAtAyQSTAKV/5rRkHes82KX4s3an4Pc/wwHTpMF9YMRzB2HeMxX3P247UZycds1/6MWK+puHz8N8088vPWbF5UlHSacFOqtkC476UQuY44uW8OVcTJ8z84gb/gspI3M5eP7X40OBxw/siwa9A2kMgvgPzIuBVdYZbgJyb7Fa0XDYTKX3FB5fZeTCn9rIItAFmVd7QxhVygbr5o1IYhIwd7fB20/t+AUOmNFUhuWCfNOMK6CvWBX8PWdMdLyWNLuIrip17PciH6L4CaBZdYwrn1P5H+UJzI4dbNiISvTOCbjfTN4uaFNbshqNdCcAOTHQNIDeBMxA9Fz/BekQ1z2dgPYDnnw1GmoYtaRCkSgMjBkNiNKK0+EGc+WdDCvz0Ppr0sqz4YGrgZAi5uofeqaMTh/WtwKE+09kbk5PQmppBdMm31BoySZZD/c0Zse0rnDeXi/5lA/DaocwWguxLDQJskzSYKRSX9TkHIujhx2oOItuHrkENgENNfhzZ5XRMuH4uzDmrrijqUjFvgslnr4as6G6qcS+uvyCa0JJmFwnXf5U3PoIo9v4dYf2M+65NkaCnxfwwb9DLtfxLvdJyS8JhNVHXWwYz7KJahRlwNn36LgsULzV2X2KKlZhPY1ynSXhMif8GF1WZOmYLVwzsig6GyrgcHmM6uiH67E9+NIJtm8vODUDrneITUnO7BhToZtjJ6wQv46IsFK5/Tv9mO8vZJn72dnDM0gAuqx1KPcR0U/wfB0bRuOs9vyXMFrMApKC++PycLgpbuzJy4pc8ijrVRX2u+ciZblsnUo8he+0HkesA5lLE3tJWKb/hcg+s3jdAtczKnYQwZ+gAs+TOg/2MEStucBrDNmIJo8BLqEzu2eNsUqzmBdPPqiyDxGyAYReeGvPCGZE5suAsWzkLFblV8zh6Zj2UCYsaTIfdIdSBU1s9aPbQ3qoMU3Mh8PsQ03d5REOocBMSuRLxmBPnaT2a+JDDjao/xlayJWHwd2Gl8pJiOJOesV4tQwRVQPEl2t7FtgH5HkOcylFVy/YJjEDpkRKVzdmIZLQNkAJqMYgUk8GXTY6dfpX4xZTL9SDxZ2s0LdjJ/362zETAD2oCcjcv24ovr5MTt0b/AiU9iJWjdyHayIH5yPgI+Aqs9AopbEMXJ+OvYdzgYMAqT1Q8SxQSY3wjPm5yb34DMSJgQB5dmR/dIhtoXKr/j85mAMxmW8Con0p0DL9mN7mYi20xJRp/U9BnUmoxQdB6omYSXae5njmtTa186eSltIhD5Dy7OcGMBI80tPb0MkMcALoLzX1ISCQAyEtEVZkKBdhnzRZJYl0PxaIbhialuxjC7A/J7ACfRcgKHSYCex/szIHIC7w/j/UjeD4OA8iI9Y8ZIkMn4qegu2OOTD+jTi6l7cqlyUm2dhXDhdbi8mIqx7pmNDkutemSH4/Aj8BFY3RGYNvZHhOb/HxejjiYURqm1aqwlbJdFtoc4f8Nadf/FBVVbs90WtMeUPr82Sqv+xnQeYXtPpTvMglJTTKYd/y808FvsMPouXDa88+e/TZJkcq1buC1E2c95BZKNUTr7AJxrfrakHYol8yX+ZC6Qllb9h33xCxDndOJnxlRu+huOW/RexOVUFIyc4yWZq5/96iCmNRUiZvyw6n07uhbf3xBoO/LhmlgGHoPRh3INThlCsDWCcjUmzdkwJW8yBoO55ZgyGkrm3ehm8D8Q9usDGp87fpm2588I42rm8WnXyMSM7bATLOtWRKqvbDiW35U7PQ+jaDe/YxstOIFpX0G7qg1oHoMgBME2ED0NAXkS0erHMXn2ITj7a475mjOmcW8wnjzrIHxZ+xTUmQGRYxhqTbYvM+AEjsHH372Y1d9NFVFM3e1zjomnwpFjmdbXtM2J5VxGw8It6Nv3Ktgz279YYn7qaa3afzAvtzCBPQBZtWieGL86dzCPv0VB/OGsfwmKHBhT3yOxLVkutmfsxIn/25CY+rAFYuHR7SqTZnGg9MXxKK25E6gzG0AqAGtvND81ASsNUdTP+HQyEDwcoRf+C6O/pENWyfwknci2gGwEd8N3K7sjUDca2dgQ05TOufPWRFnVRCDwBJ3OYZ1btUirGmc5vQNmAYeePYZq5VFAkn15/xYcfbFDC6dobdiHRQt2RSBwLlTWbukr6yESMD/t09K5s54i8f5MKlX6FhTj4eBstsGpeBmdC4mMg2CYi2+TM+fzzi2YUn06yl7YAOYn2MwCeJNvqqutQUysHoyy6g04FtqDi2yTaJ+CFM5jub4GKgcwij60TbQE4jwN8yFCk0s6V3OKjFhmHBtwZVf0hehvYFduDNPmuDIm8TDt4KSZG6K0+gK8UT0PQf0PoMfREh+0TVP5dhQzmL/9EV44FVNHvZ8k1o47TapcA3B+4xmRyBqw4qWY+GamOjb3aM0mNzM2jQY4ptfbITKetiARQKG8vsPF3A/pZu752A46a7N6BEe/iFDB0YzxSsZQR9ucOI+QP8KSBzFlzu9gvt5v7pvuvSkLZbMPo6xGt2ZOHdikRVDVxVDrdCyPTsSlY94jX/vz1CLixoetN1yL+SsBJAQ3kyi72AunvMo+x40pQ/epu3yDuHMZQ5l2N3meRAog2B1x/Tfr7WkwWDFAxlS3cF2ojmA+CzzCDoJljQF2C3vwZOj1xBeUfzrTfYMBk+cRMGPYAyDWDEyq+h1stcibHplTnhC7EZBN0dwI+iCGluWouX+u7wv6ME/CMTXatk1MO/0MktmnnCPAJtO5l43LKzlPyS2BgkGz2Tk+QW+3SkIvn3wEfAR8BHKCQBSQTzjQ2hcVxadxgWsBjOIGq6kRTlZ6xQ/Jm9xrIBv9QpDvtIB5MgPdjoxBFgFyPipGPwM79UIwmpuYMxJwtqeT0CYnhemH36Ws2chz8jTa42rbDuKxNwExg1m4GtWtYDYB2DODaK8pH/MlELgEgjmMIk7bHjKDzxDMJAIw753yZDC4bpliLQQ3Iqq3ZFU53DKN/HtSTl2AFZw8PYOwjED56Ptg5/mRsblGUWR3mC9Icp2OH7+PQE9HwJ4QwbSxsxAKHsD536XM7ve0HIvxP9jiQoxC+kAE8RLKau6g3TWhdLbn9qMyuJA2mPi5F/MVi+lvzOKe8ZtolL5Vw1FWdSub/o8Z0+mMbijYqdAaMie5fMGbsxGKHoeKkW9hgrS3n2E0nUEqiS+ES2fvBEuo9AEVkJ7pbkGl0uMoDP1AxfFilFbOpRLtTtqzcGH173Bh5Qi677DKVu3FRbCTUVb9D/K8gWhoISznRWJ3GFNZhzaJYkwcQJezf/iMMp2IH4tOhDndwTbuDJEOmfdmFOil1dsgUvs4RHZjsJbjM5H1EJBzqCDbju+8DxoWAYR8uSUjm1l0iAVPheBXaSRGmWQcAvG7iONYmLLYIKt3UKP0M+lEgueyx/2DN3PCdzgiy25C6eyhiWO6E04d/GeP+YHv8HzGYhY4eUlCQhSAQfx/Niy8hSlVZ8NsdrhgXv+EctYsVCQJ5ulkxpVX7Pod512lXBQbzbL0Ei3LFJFoG5DjOQyj836wrIdQ9OViltVXUVZVjsmVvyYe28Ec62wW0Y019xe+vDUV9OMwuaoUpVVVDeU6wIUD2YMpDGBc85ne8Sgv/kPiJysyXfxhBGnRlfssTxzdG8XWrDNXM4z52jbGKykxLh0MkdMRDX7C/NyAC2ZtC7N5x9QN066ZNq7JmnJpNmhf8Fx/lM5dmxgUM2+PwAq9B8jxANaiFbBRBVBL+ybjPpR5ZP0c9UXG8xZG0LmkAvvzQlw0ZwcILqNd3z39BHbFfIfPICC/oLRyEfF7hdcbUVZ9Fi6qYrtW3aydq9mH7qcRr5tQVvUuoouWQYL/I1Tma+wNkVj4Z/pYaaK8W0T7Ed/bGYjJDiyrt7LP+B6m7NIja2Ter3nfrw3dC+r8lfGa98iLK21Cn4cRrZ7F/E5FWdWpGdvSqgsZ5m6GryYmP6Gw7kfGeQXttrSFtKtI5B0E135+lUMPues/ahEs3MPcrKBtojrezEZB3L09JEMqavSXRDtt6mtZFReZ4jfQ3ZRt4bU5bQKx/pFoU80peqY8NPfNyT3Luj0jnFgoR+GFbBPNBxHeKYmY+fTJbINvwpRZmyX65HRkNf3cmRy7lVVzsdwx4z3OyT2TCtB3K1riVfAFoqEf8WZ1jOVV07LR6igKlGNK/YK4stzKNED2ZXvCeq79eG2Jv+rHiFhvI11j2mH75SEISzmD7EHrTiK9oXo9ovIp25wlrGufs849hMmzTb90JPMzEqWN7VRZFe+rjqT/xXR/hv3sQgRCn7ONmk6Zd2Aipl1Yhd3KeToW0O8OWLI526gjUDHmDdgHNi/T9O4omfIysxBTqjeBhX9AZCfvGNUC5CiEl3zAvNzOMeafec2snSqtPpN93D+Jx0MorfoYa9UtTYxNYZkTANcEKAn/JUjA9hyVWNSX7z3h0v5/tjiwE2OT8+HoOEY0m+9wOa9KayjMf+zPuX62WfVHHFcfh4sr109sRDP119QrUy+arDnhx/yUlvkQyJywUlpzAqIL34HLgLRqAAAQAElEQVRaDxDHEYyrqb3lXETZ58h9qA1sjYrRd+La8eZZyZMdMuNSI0MkUgrB3uzb+J5cohZyKP6AtWrPxV9eHwBT7l1Y03Y2+vVpY99je1PBtBd7hAswdfb/ciMii95kOTia88B1EmNrsxEEKq5hjZwTZ22EgFVGnnGMx50XCJDnVMRCxybiNu+MDh0imzrT8uJ5rLfXMx6+P/5PTqYub4KA3Ito9Wsoqz4MZmxvxgImD83zaN6bcT+/ajgSpzzJroxSaJuRhJnXUkyuHsV31TRfauafk1thWkGcXTMQkSjnjDrWLRX3guYWwnfPHQKKTxHUy3KXQBoxX7w1J3yBv7FxnZ8Gt8/iI+Aj0B0RuHgcBzZ5JLjCTO7fgMj1EHDxf/SzeSRdF4siR8BM2LpYikTygdingJqdombwnXDq9H8NZWUeFKegfMx/00x/FZsZ+DsYB1jrwN1EITILi5d84c7ShT5FzsfM/xxKEKNNTiIBqBwC9DYTs+Q86bhWjHoNMT2K7/1e2mXpBMkJj+pXjPcixAqnYPX68v1HCJ6Bg7MQjv0adrFRLsA3GAirYHcfBx8BH4EsIWCPXIKpJTbnoUaZcR0U7GOUSiHRxhR6sw84GurMhcVFiGj0fkTD5YiF/4Qt1jkhYaMhLs6Ep8L4FTjvMpwZL5wIkb68N2Ti+oVxV1IhMw1wRnNB7G+wx3dd32KkcrPm647Sqv2o8NwPU6r2R2n14ejf72KI9SAERlnoFrK1u5C/L0QYRo4G5DoqM++DI3OJw+urLJ6FOP8A9GTymE2KTcpItDAKM4Y3OL5K9/sBOQWRFdvh0jF3Z7Q5zhx5WlZ5KMx7C6np42saZEQyI3TciwqySr73m/FW1am4sOpwKgMPxKSall+/kLHdZMa7kys3T+BdVnk40zqL8s2AyiTGmRwPerQkNfqtYuL4FCKxe/FmzSmYXPUbTK4cAbNoaxR4F84Zg1Xv9iDE556AaPBuQC4G0I82DZIjINYz+LpuIhWih66Kb9buDYs5aUTRmmXqmE8hgVOg+llrrzbPIutCcTUC8XcRrH+UixplmFJzFJWXB1MerzFum6hWOlw66iXEi/ZhvGcC+git+VrelLeVLC1uBCFAzAJEKduFxyDWm3SZj2jo24QNYT6cyDvQ+ExYKIdgDADzHpfxvhpwpiMeHcfFkrsAMe0Dcm4uL16Kj749H5b8mjjfzPTmMb9mkZ63hrhQBJyBgPU6CkKvsE28E5HQX7H5sJMS7Zxp72KhM1FEnV2w6CEg+j5UZjM/B9ESDxMHTF5+BPQ5qDUZ9X335pyFeIpxTzDkzT/btriIuCWmVO7bUIarD+Ai0RGILOC7cV6knDvSpk/mJyOA4RA5jfm/jq0V2zV9na+30TpP0/1GCE5lpFvRhmhXkaopbz/T/006PkYkrwWsAxCKbo3ykr9nbQ5gFPtlVOI3tQPmZxvMOw7JvWznHmH66bVrgjDlLIZIGa/m51Qzs4K/MtwfGH40BIN4n5wMLiKPw95sSXKGbuxqc6EvAlPWTF1srCPyNeLyMMcHdWi/QWJBcErVofiKi2eJ+iovoKHNEiQ3B7JNnYto3fXYYtgfMbn6t4l6Yb5iT86fuavNPspsJCyrPBBlVb9DdOj5XCh/hu//dFo3uVqmI+QEDoMTeJOLZrdhs6Ensj8+FBdW7pxgtLlIbO5Xlm+mVV99PPrF2M/pbeTpQ5tPFIVIJXoFv00q1MQ3e6OsemziXZg8mRNoouHTEKn/L9vxPzFMgDZdKoBgQ4gcAssy/dI9AOasbKMS97iH/jbd9wakP9qaFWzHPqfzi0z/31Dnz0DhcEwdczztp3TPDtnz1mTbvGci32YcWlY9ge/7r1B9FpDfIj0jzMu6ZD0W4rA9RYZtlF4Ps7HN4CUw7WIB3I15f1UZjUXd41rlM63kJUTxa+bjdOb9EXoYHV2c1wYSbMxx9b/I8x4seS5Rf029Mv11k+3frwwFvW9E1JkNBD4mFqwHshmEfw2xGL3aB7y9GyoTEOp/HK4Z3fE1MbMRtbRqr8Q7nMy5xKTKQ/FG1alQ6yGmdQZtIa03CUKUqQK9lj+GSPhPrAsHJ+JL1IWqEtgzB3hH4OL78QJTji6k70JabxL8CpA7AOctRGP3Yc26v6C0+vcsn/vCbEgxshhryqkZc0eDZ6Eg+ASgxxPhlv08kpo+UPydcT8A0xdPrjoIpY1zsLLKYuaxfW3W2kM4hzBlGKyzSdNt7rg9y9cDEOsNhPQO1rUzE+OhyZW/TrSvb805FUHcTr/nAexJa9G2IjUlagTL1xMMf1viXZdWTeAc5Ncwp0SZ9r9ViHY/mn5pCmUrrTocUbaHvZwHALmCtj9cTBKBXTh95xwjoMsA52YEnI9znJB39GY3UNhiByj/JmMdrU8+Aj4CPQ2BS14M5EmW6mEG+5DJUDkePxRMbhw0N07+8kTKrhVjJ9RXmcF210phUg+8QOWNHAPFJXw0ihleOpV+ZGr/B1NWKopn8D5NasZm9V4Xgv3BUXcz15a3gm85aKvEDfvXt/TIkyd7fB3zMJPS/ETrTmaxIVa/gztDmj7TSz5DKHYaYF3MEF/SdjJJDfN7BpYsuaHbHA/dcYRY1vUuKM6AFp6AaWP+iUxPuui4DHkcg5qJTTESu8/zWExfNB+B7obAJSVf4aMFk+BYx8CRP0OEykJ5jdloWCATCCCDafcF9BwqSq6nvTlhgWt4PbfBT9YEDC+MWQHFK7y5ivYsmGO+gy/8FRVjv+Vz/pIGDqBwXGDH/ZTfXP/N62S6bUCba1rOBIxC7gdeXwfkfwCorNSpRHUisT8DARyLUOQElBff067jxAMBs7BnxlLXAFwMFfRFKpPg0d9D5W9wiAtkBqz4YamCpe//YhjChXXFPTBxQ02ZGQ+ggDZT6gUB36HeCAuUU8rQK7omXnyxiAvSlzEy807Nu30ATvwWQPYhfzpKSiSMkBvYFooKqNMQF4iJBm5CWHaEUgmYYMzwX7COC2EyjfGmu+GvF8uD+drnIl7vpJ0BOOMyTHUV+2XDF2Nayb8QjlFpK1RSwyiH/5uQR/l/FWemd7XE5DWoXEfk/oKgdRw++u6vmD6einzp3HmfOWXg0jGvMI9/Ztk4DuqcC4XZfPM+r9FExkQCEGyIhkWP8wDcBNWGtk5xNUTPpNt4+q+BhrLAR5jF2RfpxzJhxm/BY1Ax6npcub2px8Y/D+1vg9DAMVC5l8LdT9ln0N5Dy/YfaW6GYcjMyCG7wepnQN+DopLPd0CEi3EykW5nwnFOwIC+v0fFmAtQProGWR8DF7HeYBLTvT9hVc2VC2P6az6HafOM5Hu2U2bul2dyZUmcqPUV4DwBwdKGGJ3nUFjwbsN9O//btoVoYE/WW7aLbFOB8YwpnXe7NvlOZLn8P1j6AAR3sj5wzEPXrBD7OcsyY6UZgKl3cimA9s3XBezPMAEi/wDkfjhi2ipgRWggHFwA0yclrOmrcSufD6QN0eYZKdsCzMW7X5mxT1vZCmuHAcrxKEw9vR/myG4o+xIZTeYgbbZJGSEXDHURywH7Yn0ZKg8T5+sAazLbp7/w+RTO0Y/EtOI/cDx7KyqGf0t/ZbjsUax+OzjCsR/zrQl7HxTmHW8M4V/2UspOTKpvIlb7fnYiaxWL2bxXXnwHwmscDXXMBrKLiMWj5FrVv5qTHgRmE8yJAOvVyj5b2XcnxjLHAtiatmlMaTYRfEkk2f8J64uehBXWn1BR/BzsbSLk6zjFnQ0Yye2097M9uR8BmQGRv9HuDOEfPdIiwysoZlt0PfN/P8M0WEuvRDS0KZ8zJzMWWrI2x7+4gYG5Hsn/3iT0Hki7PyW/nPYevoN/QvVQujXIY8qp6AxAOH7XLQFyIW1jkXNv2puYT8bBMp9ov6zLuJjevo2t5uckYmtdBZVboKhn3N4kCXkHA3Io+a9mu3MvLBidHOuew/eG39CvP1IZ8/MbmmibGUb+DZF7EQgzLNv/VGHT9tf92D4YnP5NOdk2YXcIPNt3A3Da0fuMOUPAdBTViMXvg03Ffs6SSTNi8xVI3PoHC/wHaYbw2XwEfAR8BDJB4AcOFP7Jjur3QOxY7DD6Wk7w38j6btFMJMpf3gHEyftYs86S3bYdTC3+EEuXXA9HJwByOt+j+QLNKHKQY/MBBzR/RLzwQpaVdzJKqzlzQMbycQtad1J8hIhUuzPkgU8oMJN99OfekmgvqPzRmydNX/OF5pLFNwJCJSHu5JWTYuTWKL5g+ZrC934MFzgez9sNGciaMTvP57G+G2X7QXDCZ3MCOiOhUOAsKGup9IyIqOyRHTAwskHPyI6fCx+BPELAKISmj/6ESs070L/3hZDY7+DE9oDiCNppbKOeZ9v8Da9LAFnV/5vFQcXSRr+n6TeV9rewArvDkt8jVHgx2/J7Esd8m/EE8txoIRc9ZR8orSN78zqGipbRnWLjsfFclOPCBdOW+BEIOscjrudhydKpKB9zHcpL7sOlxe92SG8Q7PUyHGtsB/OzOxdR78nem5wVgYRvp5J03w7K1fo9lUACF2DJ8h/w44+1cIKn8X02vdtxHU5LZbeV8Vl6JJz4yxDRduFij48h0vc+IGgUe63zkc4z8+M82660mweyxy/imP9JhKNXQIOnAw7lie8JM/aHUJmur7KufwvoMrQ2qsvp9w0UlfS6Fhb+ADCsEzgC4cgUxnsbbLYxpq0hQ5eRwdrUo4qS/+P84gI4gd9AnfGU5yjaa6FaDRXT1rXNI7CQPJ9C8Sjf9RSo/BpR7AHIcRg85NKG8dvI7/ncvnKATjIXbx1FPMrFQ+zPPOzDurAnbTFtOmWtfTzAGNbx3WFZe8GJHUbcj0HIOgeLB1+6sn2rGPsaztuei4HtrEep4NtqxHKmad7bPol8q4zPaZ4daR9WTeEgByBSOzdVtrqt/zWja6GBZym/qVNmbHEX7A7+3JlNvQWsJ+CwXjbh2J4rrAMQjD5C2bJE7OeAv8KR3Wk7Vi5a5mc3CC5OCNlrwE+sVxfBjF+MzX5a2ZR7NKzA/ogVPg23PmG5w/5ETliZH9Pntsx7duVxYmMQl3Hs+/aA6fvEOpJ92GkIsv/66JsrMG3sLWzjn8v5HD0YeQ0B+W0i304nj0Pbg6/lXIDpe7L+Jkphbv7Z2yxDxdhn8GPhFXw3pyawiSv7EpwNgTnhwmxA+B6KaCsBjJ6FfbK+BugtUOsUiO4J03cHC/+M0OhrObau+n/27j5UsrKOA/jvzPWaki/YHyWCRZRhYlK5vt2Z2XUlNSsQiaR/QkJTjEoFlSyyW3tn7nVbUBKENHshNApMC4LIyrDUNMuyQMswzRJJ8wVfVvfuvdPzzLp19929M/dlz/O5zLN377yc83w/v5kzZ855zpnIy6KtHjjQn+uffziq2bQzOa3PzzZOSa/5lamNDdR6jdWp35veuxpxboxOPzjvPl592CvptZd2dMfqefVptzOfBwAADN1JREFUZub0iMZ3/9ef/vO00Z7XtOY+5+Zm7I2cHy82/jHvjGsPfz7WV19Oddh9+17VjBg5NarqxPlnimaMjJwc0xtujOgv/+cdZYsHju59S3JPn89TH+fa7eT/jS0m4I+lEng6LYCujrWr/7lUHdhmvnnDz+zMusgbc7a50RUECOzRAuNp41Ja80kZeovTevnMJt9LG1I+m5Z174tO603RbZ8dnfbNMbHq73FmNZP6sZSXRXJIS9Tdde/FflE1Wv3vnV1KobnzzkfGT7b/Gp3mNfGWfU9IK0SnpdrekdINt469mEnTvTvyhsPR6XelDYY/GvgI8Kp3dooyktqOar4xzfOWoZ1iMs1oQS55oF4VP0jT3lGOV6+vzohLf39Iut/gl1z3TvP26LbOimokH6lwU7JKNa/yvAaffn8K/Wm9lKY7Hr3RFdFtdzadFaR/ff8ei/BPL81jEVr1Uloc3JleN1fHbJwTVRwZndYx0WmnHTyt38Tkcf9J/VjKyyIYpPQJIYWcz7wOj5GZd6bHLsTFNAkQyAJ5B8zEqodi8sS7Ip95p9v6XFpGnRzd9qHp94FpPWAkOq2q37qtRnRbB8Sm205Lt30htZsin1Y8n9p80I35uT+L2fLRXJPNu2Ip2tSJv438FTzd5h/SevJDkc/MMNV+ZtNAuCG9H46/59mhZBsf4vaL8bTDpnPsY7GmffdQ+ja3dhMnPND3yzsYusffP/Tpb57XmvbvItdqkOfqurTjc6A+rnxykNlv8dj8mbGbdmR3V94fU6t+kV7T16R2Xnqdr0jtkOi094/Ny4DNv7vt/dJth0a3tTLddlGsad0YnVV3Rt6+lAeUbjGDZfBHHqxxRT7zwdiDacfOHanPN6R2UXTbreg287Ju//T3puXc5oyd1hvSdW+Pbuv0tJ7aSff7cfrscG+yeTTyEWfLINZr6kLOPrX6kZQlrXcu0vKuk9Zx8+tkzdh9Mbn6wU2DwsaejrwjIoa0fNtV+LztYTzXe5Eyb14+zPd3fi9Yd+qLu4q1R9+eM3Za702vqQP7z8dhhMnLrnz68Pm658d1xu6J8dVPDKM7/Wnk97l8QEWe9rBbnm6eyfiRG2JyD3p+52VBXgbnvm+v5Z2y+fkxbK8dTi+t8041X10PS+99eR02n7Uqv3/ldYjt9XEhrhtf/Wzks9XssJ/LbPmVt+vm95SFsNh6mteumI5ck/y8mGr/LCZbV6VthR+LTvuI6LQOTsuRvdPvue/bo+nvg6PTPjq1c6M7dm1MtH8Z3eafY3zFUzFe5QHNMfSfvNzurLwnFqqGa9p/TMunFwbqd37tdVr3zquP+fNKZ+xf83rsazXJ68N5GTBIyCvTOsZ865DfA/J6y2vt7/bulz8Lr02fl8bT55xBcsx97PixT+yuuwEAcwGX6v+9+FtUs8/FZbc3l1UbaTyaSB5OzYUAgboJNBpHx4K0maNjtnFETE+/OV6czhtHGmkl6x3RaX00uu0rojv282VFOdu7dWEchuQ7kqYz2/hiHPnGhVkpHbQY56WV74nmT1NtW1HNnJR2Yq6J/B2FEfkrAtJ7SO/ZNIv1qe2o/xvTLsB8ez6S5rGI3gNRxd2RTxE60mvGU/u2+xsOx/uDVtJkBryMNM7fab171TFx0AHfGnAui/Pw6fjaTrP0X9/Vijhor+w73D5NHP9Q2sj5kbTu8taYnT0rTfz61O6NXjwSEanm1YbY+U9+PuR+PZXu9mBq+fTGUzEz+4EYnc4DhL60JDvAe6M379o0vSb7trv5u7fXu6OqDotXqoNjdJ/9kt/ro9NuRrf1mfSh9frYvNEmYSz5Zb+j1sdMXLhgFvPx2/Yx74/Zve+MBfkxUQIECBAgQIAAAQIECBAgQKD+AhISqLeAAQDLob5VHBdR/SptaP31smq5T1W8LfwQIFA/gTzSdkHaqvsijzjOI9yuWp12BC5zunyUzoI4jN0Xw5puPmIm72hf5pTRWXV72ol5eey14YyYnv5Q2hn88ehVl0RU4xExFVF9Ne3g/2b0ejemHaHXRRVXpjaZ2uXpuoujN/uJtEP5w7HXPidFt3lh5CPArl2x9em7YqCfXdUkjyLORzwONJNFenD+LrRd5cm3X3ZUHlyxAJ2qeqnmj8Vk+zvRaZ0TMwecElV1VmqXRq83Eb24Ms3026n9MLXbohe3RhXfj02njv1K+p2fFxfGxpkz007/D6ZpXBZT7Z8MPIo5BvjJR91ns4VoefRyPoJgXfPfsdyPhM2j4KfaDw9tGbYQnnmauV4DlHuHD3UDAQIECBAgQIAAAQIECBAgUH8BCQnUXMAAgJoXWDwCBAgQILCoAvlo/TwApNu6Lbqtr0enuTbt3P18PPm6i2O6uiD2Hvl0vPzCRfGXxy/pDxjotNZFt/2NyN/nNbHqgWW/c3RRMfegmU0d9Uyq9e0x0bwuuq018VCq78b1F8R0fDJmGudGY+a8mJn9VIy+7qL0fLgs3Tc/L26IK1b9Ke30f3kPSqqrNRcQjwABAgQIECBAgAABAgQIEKi/gIQE6i5gAEDdKywfAQIECBBYDgL5aP581Pr42NORvw9rMb9DbTnkL60Pub5XnPxcrG093v/e1/y9cJMrn+wP8KiqXmkc8u4xAjpKgAABAgQIECBAgAABAgQI1F9AQgK1FzAAoPYlFpAAAQIECBAgQIAAgV0LuAcBAgQIECBAgAABAgQIECBQfwEJCdRfwACA+tdYQgIECBAgQIAAAQIEdiXgdgIECBAgQIAAAQIECBAgQKD+AhISKEDAAIACiiwiAQIECBAgQIAAAQI7F3ArAQIECBAgQIAAAQIECBAgUH8BCQmUIGAAQAlVlpEAAQIECBAgQIAAgZ0JuI0AAQIECBAgQIAAAQIECBCov4CEBIoQMACgiDILSYAAAQIECBAgQIDAjgXcQoAAAQIECBAgQIAAAQIECNRfQEICZQgYAFBGnaUkQIAAAQIECBAgQGBHAq4nQIAAAQIECBAgQIAAAQIE6i8gIYFCBAwAKKTQYhIgQIAAAQIECBAgsH0B1xIgQIAAAQIECBAgQIAAAQL1F5CQQCkCBgCUUmk5CRAgQIAAAQIECBDYnoDrCBAgQIAAAQIECBAgQIAAgfoLSEigGAEDAIoptaAECBAgQIAAAQIECGwr4BoCBAgQIECAAAECBAgQIECg/gISEihHwACAcmotKQECBAgQIECAAAECWwv4mwABAgQIECBAgAABAgQIEKi/gIQEChIwAKCgYotKgAABAgQIECBAgMCWAv4iQIAAAQIECBAgQIAAAQIE6i8gIYGSBAwAKKnashIgQIAAAQIECBAgMFfA/wkQIECAAAECBAgQIECAAIH6C0hIoCgBAwCKKrewBAgQIECAAAECBAj8X8D/CBAgQIAAAQIECBAgQIAAgfoLSEigLAEDAMqqt7QECBAgQIAAAQIECGwW8JsAAQIECBAgQIAAAQIECBCov4CEBAoTMACgsIKLS4AAAQIECBAgQIDAJgH/EiBAgAABAgQIECBAgAABAvUXkJBAaQIGAJRWcXkJECBAgAABAgQIEMgCGgECBAgQIECAAAECBAgQIFB/AQkJFCdgAEBxJReYAAECBAgQIECAAIEIBgQIECBAgAABAgQIECBAgED9BSQkUJ6AAQDl1VxiAgQIECBAgAABAgQIECBAgAABAgQIECBAgAABAvUXkJBAgQIGABRYdJEJECBAgAABAgQIlC4gPwECBAgQIECAAAECBAgQIFB/AQkJlChgAECJVZeZAAECBAgQIECAQNkC0hMgQIAAAQIECBAgQIAAAQL1F5CQQJECBgAUWXahCRAgQIAAAQIECJQsIDsBAgQIECBAgAABAgQIECBQfwEJCZQpYABAmXWXmgABAgQIECBAgEC5ApITIECAAAECBAgQIECAAAEC9ReQkEChAgYAFFp4sQkQIECAAAECBAiUKiA3AQIECBAgQIAAAQIECBAgUH8BCQmUKmAAQKmVl5sAAQIECBAgQIBAmQJSEyBAgAABAgQIECBAgAABAvUXkJBAsQIGABRbesEJECBAgAABAgQIlCggMwECBAgQIECAAAECBAgQIFB/AQkJlCtgAEC5tZecAAECBAgQIECAQHkCEhMgQIAAAQIECBAgQIAAAQL1F5CQQMECBgAUXHzRCRAgQIAAAQIECJQmIC8BAgQIECBAgAABAgQIECBQfwEJCZQsYABAydWXnQABAgQIECBAgEBZAtISIECAAAECBAgQIECAAAEC9ReQkEDRAgYAFF1+4QkQIECAAAECBAiUJCArAQIECBAgQIAAAQIECBAgUH8BCQmULWAAQNn1l54AAQIECBAgQIBAOQKSEiBAgAABAgQIECBAgAABAvUXkJBA4QL/BQAA//+qcKMHAAAABklEQVQDAMiyqDZqJSSBAAAAAElFTkSuQmCC',
    'Área Sanitaria de Vigo': 'iVBORw0KGgoAAAANSUhEUgAAASwAAAAqCAYAAAAJU2bcAAAQAElEQVR4Aex9CZxcRfH/t/rNscfMbrLJEhLDoUBAolxCOMSETSDIoYhKFEE5VDy4lSPZBB0hFwEROTy4kR8eoAgGQYHsJpzK5cVpAKOEADl3d2aPmXmv6/+t2Z3NbhKO8AcUmPfpev1eV3V1dXV3veru92ZcurlF32ZoS09racarHHXNCzImU/3UBZNfhayCqmigooH3mAbce6y+lepWNFDRwDtYA2/UYEWs81JCJVQ0UNFARQNvmwbeqMFaTgmvIywh/BeDytAz76yvbW7dqW76gnG1zXeNwGE3BCaQXQ+Z2rrlIMi0DsGp91czz+bl9Pqp9ww1emRuSFSf0TqauPrSffl03MPxmjPuGDX8jHvT6Wl3DTO+OPG2ZE3z3SPLPMox827ekLmtrpx1yPQFW6SmLxhfN711j/5yysj/QlwpsqKBd7oG3qjBSkAwilD131RAetqChtAF33HQH6q6Cx2CH6e3bjzADIqo+1oU6KWDoOg/lqrObxUFwXn96a54CY3KYfU9I2tjMX9KFHNfHFin+oaOHYJ4LFMICturuM8EcMfX16ZrYgg/38+jr5wwFszMh9XjzPilmlumRCoXgKCqF6orzqqb1rL7QN6V64oGKhrYOA24100ueIm0/6SRupbxZaqIw+PnqnIh7x8k/IuQJ7xtQSXYj/J8EpDbFHqhQF/i/fDGYdVxEd2ZxmIoZVxUBu9iS5wgBWBPGpHFonhEIJsB8pUw7kczdqzPkXWZPzSg7/Ci+5Fut0gkT34fVMFu7VUvdkbA38t81ePPJN+F95u5CEvT+cbJVOw0pg0jPA3o88y3lwdOrJuxaBumVUJFAxUNvAENcFy9Zq5lNALzFDgBIl/m4P0PBMMD+GsgWnSice9kuge+CtVvkNtThLcliOrmLLM7FPfzXLbnFiQKM+KJ7ptWYFEXIBHleyKe7PmJQTJK/Lgz8E8C3rMu3cz7q1ii53wakktZp02DMNoM6m8EMEzzib0Zo25q69Zw+KhXuRdFvCTMZOnITCl0JlbeU1MsXmoAkVVQdIrgckWwivFB5LtMvD8dGjtFE+HxonI+03eDj/YEiC0xqpwqGniLNPAuZft6DJZ5Innn8XV6MLuo4tNQfIoDcwQgB3DwHx2o306AY+FkCoDbCW9L0EAepLGojSG6LJWqOsjna4LVj+c6kcl4QGhLsWVYrP6EQU8iv9vQ7rAWvQcdLTc0H9VuIep3U6CoTtuyPTV/5fWzInq4Tet84HdWlZEa4JbOWrRh4EGj9fL5+3cWAkcvSo8h3fyOlXU0eIVG8iBffShWVXg6N2f8ilxm/+XeuUcAedZD3o/MwiQqR0UDFQ1stAZey2CF9FKyHIAxLxijHlUQ5FhKAKjljdF4ddIDGSqQkaLYFJAX8DYdueXp+yjfKZQhEsEF4sLzUtuOGNO78K6BKj5MONkAXj7l44mGkmiKOtZpXhBFv4fIvrRs10i8+BSWLi2wUj+HYNtU2LgX8+ztIM/F4J7Fqm5mKeXuP9kCexgEzUJl+Kh4Hi7bteicpHjLhXd5bjUeNF2V6KOCz1LW5Q46IoUC8aXkyqmigYoGNkID7jVoY7A1HkWSgzKAwxgAZpBsXrUlr82ALXaQUQLdmvf3MT6A8dsTaCCysybNjxD7Kgv8KQ3X7vD+0w1jU+ZJhQosYNqXDCIJzmtv71xGOtC45UXkCeLzonqrxt0vOzIfX40bD/NwwQM0wCpeT6AbNtaLLmyPvWx1tqxrIdMaKxSqD2TCR6FySde/2l7mNcIoVgSkQCM5uhETatB3xAKkoBjOtbauGOeufcmVqKKBigY2QgPuNWkV23HF5bOkq+aAs3hHXqdEcLjCPCpsyUF4MK9HemBPxh8j/m0J9TMWcXrVWtU1e/yLPghuVcjzAp+OetL0AKE0OG3ZOU1PGXTPmvA8Lj6wtCmgQBc8rqJRuU3hJmvRNyHTWgW6QEGxsFyARQAmky4u3v/D1qx4PyikQreNQI8kj78B0a24cUpkBBL5F8lmCQSTe/LYx3Ys05nW4dTX/kwbDcgTbY+vyKJyVDRQ0cBGa+C1DRZK3lWBnH8NkcsJvyVcQrhVRH7E+BeEXxIuFEEAgF4Zz29DUPXHpgv6w9S0BTOCKGymtzRKxf2lfXU1p2KcEgI7p6e1nFGG2rMWfRhwzkQT6AoXBJcA+g9ROaGuoBNs3WrNc2tyUGkhjRfgMcZPEAaFIZnWIc6Hn2PiOCeSgLjjrIzU9JZvauBGOvW/IC5Q0QzX1uZpXi9U6HFQ/YtzQWvZuJGmEioaqGhgIzRQGryvg344afaAYgJeHUaRjuOc57cjqP6JxYwQkf3pKXETQK9RjRbaWhIN6D0QPE7vcHQZRMNGrsPZ4vnvJNCVHbFWWwS/mDzuo3Xaaggaa8yYuNJuIq7yKq3Z5L1riAfym3iW8wjr31LoKtawPPOSuMGgK+mJjbQyRDHaBWG6PRHcQyXMJO0/RbA7YSte/0FdbF57rPXfJX6VU0UDFQ1stAZeyWBF5NRDKIchvNgB0N0GAL0V3WXAPXFoJB3HKs9AkWcDRm9N6EisvNMX5euB6JFhiC9VJd1POmdPWm6lxcLoSvWxadDYOWWoi3X9KZtc8ZwkCnPb4u4FZDK+M4HHYlF0bhDFf1GeqrXH3L9JM7O2WLjdaIwfLvtImIzjFtJe1tXWsCJI4ErV2GleY98t84f67+e6qh9CpqmnY2VdCxA7mfRTQgmmOI3PzMVan+7nh8rxv6GBihTvJA24DQhb4BrM7bQ6XyaunWDGi1F/sPsXFTiKXsUCpq5rlIjCSno315DHZcSb4fOM3/yQmVLoPK/ppbZZk/7dPa9p6YpME6eC9HNY0ppz92svvVJgrxX0wbLMJ7psParDFtgzTSHJQOMSGm373I+tMe+qnGY09tpC6b50EjX+RmseXFumqW1d/tk5+67CD/bqLpFzQ8DwpP+PrZ+V+GfsdYsStnKqaKCigTeggfUNlqBF1V2vCblDFfZW++JBfBX2Yug1gY/f6VSuA+ReDD46ReQnXMe50TvhuheuJ/pFwns0qIAL+vYt4ntUAZVqVzTwpmlgfYMFvAjVF2PdYdGJPMeSOghrg6BdgSWCsN47v4wL3aUpGNYeIfMvLvr4SvXRau7GGY9er2MtzZt3deJtSVtMt2/30s2tXyEcXT91weTej41pLPpKasjcVpee1nqwfZCMTKa/3nVTW7eua245qPRRcx+tRaMy82tS0xeMJz/j2Qctx6amtUwwPI1QLDXtru3T0xccmm5uOTbd3Hp07YzWfVOZP25SwvM09Mw761PTWg9NF3B8T1A41WjtQ2qiKqGigYoG3oAG+gduf17FMfSQTi7GYzsr9FymjyMMDHtxqndeKNhfVL5Nk2C7ZQPxQ2jQLg1c9EUnweEQzCLS3tFi9CaHTGusLl21j/PRXHI+k4byCEC/rAHmRq64H9P6Q+mjZOjcyAtpJiT6EQF2p7ynO3WDZGwr1NcL5Gjy+y6gxxCOVMjnnegOJWNV9HuJuLOhcpqlk983xetc5BOH8hojTvtjbRQEx4rgPOY9kfFXoTg3iMVOqJ6+4H1GU4GKBioa2DgNrG+wSvl1HAdfhpdxwoZCDUSPI4KL7jyvH6o4SA+mAfnU+qg3L2VID0YrcJaqxEXlHBqTb4vI6R5yBVR8f0mn3l/tFJMgeB/hc+l8ZC+W9qNf/UIXKvwMB/2WRnp6BL2hthBtz/K+LQCNs1xMGc5gcaeR5lon6DR+XfH4JIV+A4qboXIC9fVN0t8A4KC4x0TQM+R1JVQ0UNHARmjgFQwWBzZgU59Xwq9wkGcFamtThQ2UFwNkGwDbEd6yEAb6ETIf6eB/kH1mxfz22ZMe7ZjV9KdcR8+VsSTuIq4Uaqu7toFiRxqr3zCBxsp9lPHrDS+L4u/Gu/PciX/rnL3vy4HIWBrj9wNyTS6BmztnN/01N2vS3R3d1Ve4BG6FHaJ7QdGDAJdl5+zz++ysSfPF4WqilkBkj1QqVfk8h8qohIoGNkYDbmOIS7SKFgc92CcKJ3ivnCbhp0x/69aoyPyVAr2mHem1dLiEewCbpkprSqkZrfuk6qv3jIr40PAz7kvZepVoYL+QUAPvr6XxeRiCo3q/N3wlzgPT5UuAW5RubvlbunnBb+0HANXLZhBZI4KnuMvYUzf1rt1K5dbmd/c90Tb2434CbKGQx1wY5xqfqHG01yXodf2L19sGWkgzfieFiqwVDfzXNfBKBiuiZPY6AqNBoUNEz29fvOpv9gsEuTmTntQguFwF9wyi6r2x1wbWfeWhF/MmnWkFRnDe59uAnvradI1z7nPw+n3Ctap6fU+ysHld2LQVDcv+gC723DDwThey+B3TY4a93h/Te4L1+xUgP4Pid0Pq80k4GQqgX3fq3GEs83xOo6/k9XUO4SeYJ0VdRVIVKmn7wiKKgB4m1EexeLIvsRJVNFDRwOvUQP+gW4f+aQjsdQQzOgNRK0Mfe6n/fSUuzCQKhXb6D/bW90A6myY+yIQHCG9ZoEH4h3Adqb4Qfczec+qYNfG7ucUrx9Gw3Mzpaj4R9DzvNdqBXg2njvIhB/keILYonlRPI9P3c8p41UMfFvU/ys5u+n52zqSraRxzqkovSUdwPWsLcOE/O3viGVVhoomG6BqyWsW1tL+KuqWAbIkeDKGxFPCoLezdqCqjKdszkYYdTKqEigYqGtgIDbgN0P4dkKuiYvgdxnfCXnPgqS9sEgv89vbNXemeAz50sgWv30cohzwH6O84Qi+ll3ERE+1D4t7PW3jzZoZI3b00Ems83NfsFQT7Xfe6McOmiOAgGoaHCwiq4DGexvcpynMxRK4RLz9WxU3itCm97fCt++Sposxj6psX7GKQnta6XcJHZQ9oBOuxg6XXntmyIxfch5lBAtxLovr5urxOrm1u3Skf5D9PXocI8JdI9d9QnU+ojwL9Rl1z6+TU9AXjAwm4+I4PKtyC7q7q1aSvhIoG/ic18L8q1AYMlt7Oac1tQ2rybVBcRnhygPApehcn1RWHfSbV3PohTqsOFHEnKLDDAJosB+pFzsv98TB6SGzHrvcnaQaQvDmXnYkVTwByPoCRom6uqvxAFWfR8XtS4C/W7vhICPYW4I8dCbkhO2viLVwAv0U9roKilnAQ81oYLiInecgFBsw/TV1kmwbEyT4CN9PSJZDzAsiUVLzrMXptlygkrSIZp5yGQk4mv6eiCJfbwnxHcuVd4krlfIG0FwvkIuruCKEsofjb+9+IR+WoaKCigdergQ0YLDlAvD+wrSs5hIP9OMIH12E2TlXOEdWLOUDPVeBTHIQ1WHukIXKSd7pXMRbsxsH6FaIGemC8fZNCZkqhLpH7nYpOg/qrBfpLVfcd9XpmNhk8JNXFFynLPHpEt3BxvG96K1pf3blYBVPh/QOI8GdOLc9SyKWA/J8B+d2qLr6EsnOKJzaNvNrSjb9X+fuyzCe6RAzD4QAAEABJREFU6hKdtwbQZtb9Iq5m/V4dvisBzrGdRNhB2YLQX62KkwH3UxrJm2gIvx2F4SXdsya9YCQVqGigooGN08AGDBboLemxQTx2NqdJ+5HdSALXY/RmDr7r++BPSq9JVR5VxY2EUjrp7TOeOCCfJP54UZwEwF6PsEVqXr75wYyHvVKQTbqrO1bW/yQ3p+nXuTn7PmEGKpfZf3luVtOv2mdOenZgyb15Jv4mO2ffBzrmNj3TMWvSL7hGdUUZyO/GjpkTFjO+u5zWG0+8Kjdn4iLjZTzsVYeO2U3/l43Lj3L/XHmzL6LK/rU6NWPBRBz3cHxN6XvGpt9mE7i0Kkr8IDtr0m+75k0u/Yig8XhPQ6Y1xqn3p9LNrTfVT10w+fXv2r6ntfaer/yGDJYpZVsojuBFjEAnRVdA3ZVA7NRXhGTxW/QknjF6QoJgb8jvyfjtCZmm0D5KfnsKW6eUTFNP7TYNw11gBtptL/Yz0Zd9pN+jo/HsWTlv73U3JtZh8t66TeexNT3OL9JrfbEQxJ4s/drre0sFldq+AQ24V8gTMJ0L0TwzqMoW4vyWueRL6/0Cgv0igYEWgk3pYW1B8jJPM3b0tpjyVoVMxtmveaam3d1oYN/u2a5df3GH3RA0ZG6rM9xAGJWZP3AK20ueaa2y7wnJY3N7j2o9GpZVZ3//lbnBjHFvHp4bM60p5ql3ilCARRB/frKQWNaQuT1d8hpIn860DrfyLW6gPK8lo9GtVz7LejcFtX8XUlzpXXBJ96zx9OBF2X9k+Bn3pu17zzKk7NtM6tBwttlTc8Ydo8o4i6n7ze3duH7dsM1N18anlKcfMfjC8HVTW7dOT7trjH0qZe04iJ4eoPHpTR+Q99T7q00GDPxSgfKV0oiz9/7sBx6tDa2tjbfJuSEo1Y3yWj+xb18NBpQES7f6lfMar95+wQWNDPv+tLuGmYyD8kCl9vTWTXvpBmP67/rkXct30Wa9ZZNvH5HVoYwvxzY+rH59JDDdmEw2myinWUyZ60vpLMfyGC9LLwHlNpl7dd+6Xf3Uuz7ANhi2lgdlYL7q6Ys2S89o2ZZe+HZWfm8bK5eFS1xe89RAo/XZVH74ntZpBlFTABPOQU5kj9tyEO4tvmnAuBSKOk8kPEdcOCuMxWakiv7QUuehXOmtG4aExaqjiZsDCbkYT0A4oyNMmffXL51995cq4hDW4ZLQBdcGCH+SzacOG9jotYW9G30+eVa62Dgob575wljw5eyqIR0dsyf+jNPLB3vixQPDQvUx6bGNQ2sLjduj4M9zEp0tBZxTLFadmc7rgeUOsiEZteBnZPO1O/cL+K67YMfTwpYe+IJoNAaZ79HWs5LHPRLriec/zvTfRir/Z+AKsR+m8o2fxIm3J+rCEVsEsfglTP8V4VqD0AVX1yZ6xjB3KdRvO2Jz6w89QfELI067Y/0HE6nMWOWDwjfU6a0Q99sY5Ip8QT+9ll4lHepW7Dffz+f9oWY4mK0UUlU942Kx2M/SdcnJ5fS6nsbNg1hwabqqa9fhXfvW+oL/Itv8tGK+emLg5VyTkzva1/XV64bee7kW+fjXh4xtTFsf8EHYrEHxFOuLpYJ4MlwYBDeS/peEa03OXKH2a7XNCzYZhY9UKdzxrOs0kvaH9LQFDS6ul+SKNYf0J65zkc43fsDkJc9evj76WSTFTN201nHlh2mY94eHKjeRxtrh2sjLNU7cSVg1rt8J6S76XWLqLq9v6OAyEnoPGuDQua/GVS6o72kY7eBmG68SksskdcUJ4zgWL1Txtqb7ey9yM5ybU9PQ+aESTebGeLo47KCYRj+HB9tH51OGX6aq8l8akllY70pE659sUXgRkz2hHD4qDlOt8wyZ2rqlVay2+a4RtYXxTTGV7wByqADlDhLyaWXrWfZTNHirjkI+lYTC1svSUO1QxY5O5fvsUDNqOydsErp4DXfmPgaVMaJYWgLBUkTR8oEydScTH3aqUwGJK3QB67EUTsesKdb0r70FcPuL6NEs50jWvQp9B/l/hHVtQnK560uCqN9ZgQlBD1KB6iiIcBD6vIcW2Qh7QDDPu/C4BnphMa+1lHs8FFtLWUbGLgjetWtdozK3VqsLxotw2UFxQAp7Di/pbtRzAu8aVGUIdXS/Qv5AGCniZ1fXJYdHAupdNwNkGfV1WwlEbguqC8/DDg4WH/nxvPyEiH6p9y/YeLdO6In17Er+3xDgBlHl7rI+7VVre2I1vd4zBw3btYn8j/CQ/Wu22XSTMguxnWFgH6g7NbVN4wcBUGSTS8YIXDqsKQYARrAOmzvx/3GCO8jnNlbsbqZZvVbbfQmcLmwDesQF27BfHa2KA7viQf/3uVE+Ml7vZ/q/SvSK5QqcLnATupBKUH+2mbUFBh3ySd7uS7ovmGHm9XohUlcNyGgALxhf6uo+Efmwilxc1xOVHpSUe4RAq6mn+40GDvNdpIvK/4vAvEA8fIrjZYQP5LOle56qx4wYpeLs/xBekEDsBfStjBdRqGto24l6/b4AY2hLrqKM00hrn8tt5vp25evyw3eGytmkX06ZMgo5izI8TRlPCwv4uCNiA0EeVCcZIoqEcohB8XFmnBsF+p00MIQW9wgyuIAERxOGEcqhBxBaR7kZb/2RA2R+NpufLhqdwApytw8Hu5iOQ/lQfSg7Z+K8EsyeeF7OFuXLOMbcwdtUgRQVdGNuVf25Hd1Vp0XF8NLururSu1Lm+hL/aZI+w/qPq+lx6+6cEvUqQUE+7opcQs7QGGwj4iZAjw3zVWv5iKyVcc6k89tnTvjXq3B8R6O6ijWNtOr7shJPE8YiH1+rByYwPC+RXJ6b3TRLITdCpT5QV354dPLpfEepLa1NZzd9v8N+kJGZUltv2sD24c40XuLt5mHMfYQPlxiv1wnBCAUfSaJ/6Zgz6bpcNn9mddJd1/7si6WXeYd2D60WdQcy01Mi2DrQ0H5Nl7f9ocD2GwX4r9jyQX/qOhcFwcvmdZusRcGPif4PvC6yewPb1LG/j/OK/Ymzj+Y5wN2eGDjdBDqd4Fajdx72XuO/AdkmzAdxrHOUvDNx5jg8DJWti7GePdYhGXjL8uQPxrdj9qQZIp4PbBooFxy7lsg940QvM5rsrIkXtM+ddMdaHGCbWry/D6rjS9Nb3sThx4v6TWns7gzj4dpP9rjkooH7FElotHV2x4r6S3OzJ96QW7xiJss4IRK9jziKrYcx5vPDn9Uxq+l69oFfusifLcC/RfSLjsjBQXC1qv4wXgz/IpAzibQ31hn1h63gcVBQ9LWiYgoxdzDRjwXaBDg+8sF1XqNfQDEdwDOEtzZcfGCeu37/dD52mUJeUpHSk6JUqGCr9PSWQwxS0xeMR2bwOhQVaQNnmap+IzW845ihic5EaTev79dDuwvRbuQzVlQvUpXuwPnJvN/4kGnqyZ0z8fEIsYupoy4vsms/E11HRrrP/bh30wXXhkLIrqrYhH3jKhV0URe7N2ZaUwOqWSPwH6hvXrAL450gKAbOcYCVKGpF3eT0tJbSn4vUN9+1X//6RxCNVWBHB/kZIH8h3eEjcoUk1jnEKfs2XlSVGenpLd+qr0+OXpFpypW/4CgEAdtF388BcjnLXgWH3Rts7XEtn5dU8UeBNDkpHuJ8VL0WtXFX6WHVaYF+nLl+TmiB4KPpVGKg1xRTyEh771EdxiswCupXCwqe9INCLhHneNStVJQPbXkx0uCT6/b1QRkG3HR05P8G6MMK3cHWwHpRfmuv0vsHK9Nav2ZrT73pA84K85Deh2Jsgnl03tNREHlOXOwpyQc0wL20tZ0YQsfg/ZT/r1Ex+nP/BtmNUyLbxV/7qo/sBtH/DHQq2ufu+xzleBKKj7hedoPOIyEyMqyOxWnmPkBMHWG9kAcnM6L0pNZDxZh/m7grDhcXNMDBeLzhBl2P+2sk2Cc6AiyB6rZVYYGXsGMs7482EJWDG9E40MAiF3v5WXX6PaD06xPTwyC4yDqIZTRwcIdB8M+iiy1kfC/E722LhYZ7I2B/S8aG6xH4ndbm9x8y+QyEU4MRdasGybiW7p191QhUOcUUdrxnvPhrWddH2Ej79YRq05ty5bbzgvMVciWEHRi4tDbWUfrfRxJQL9SV4EAQFO5DGJUNGmnwJPJ7kFeeer2HA+9Wwg6d1fEtmWdQyMbkWZZ7soM8CMXxGslPaLgOKRu+APgM05d7H78eXu+Bx15RT1Xfi8RkpehiH7hZgUWicpSX0oPHEbPRQQpqDz/7maKfQ9Q+ZdtUJdhlgGdYz375TcI1LO9rAr2bsCAZVdHLG1AcvTLW/RDK/VKI2MPCNSIRHVfX07DjAKpXvhyWU4hbSYKEi2EISofYQ2UfUM8O2hRs4IN9gX+M5T5DPezL9cedRbAjoaW9M25ebomLnXzCxZlu09gliZp4Vx03sNLNC65OTW/5owH1f1XqrJaxlDkNlRcsz0CgF2Zfy9Sur2TFRBF/hCmSBRzFTGsbijelIHguHu/upKAbmrbU0lP5uhd3mPP6WfS+HjGylG/d01txzyc4O+om4tip+vnrwkBwioEvygUrHl+x1lU1msyUQi7m7vbQE9g432Ln2E1ET0txh6p+euuu7Eh7iEci5sMjSM66yK6eazCkE94zsLRh1b3XmYyDSO81MRsKxpf8kyJuQKNKi8ln4IvuwpdTDwyWcUOM3oFp3QVsqcDebCMXiBxOT2Vz3u/APrPDKFTH+qr0Apy8oAKbWlwhicJPl2UOLuujAyrXBKJHGSBRuJZ2o9CTx2gRsbWbuFfZVwRbAKhju30BICfe9IdMU9gxp+lBJGwZAScSm+aAO6KuIbtF/YwFW1GeiRDJO1c8AoKtCGOiALuD05p+HghWaAQaVKx2IscxfVPCxgXrq6KH8yHFGSP2B2Rvypp2wAR6JL3rekCW5f+THer9ABY6h0w2GTwnNUWKyZS+UJ9OjhV6rqDcXLD+PHUwBoLNVWQCrE/20b1i1D46UK/Wt7PxZPcyoxPFX+itnW56Ljr9dnvV6qWWPhCyuUIO0FuYtpNT+TzjQCGP4gd7DnJmgnhPnm3dY15WgWsCHe113VD3G+a5jtAOxUfFaz3lXk4eIwfJ3Ct/A9NXUjeMBocEVA5QsDEAWnfwgTOYgIpo6Xi8wwp5kBizfIz6g/BqOBRHk4c1ZBXvN1QOk///A8vg2EfS3FFbT0gXcCS5jmDlF/XEEkSDokh3Mh+uNKivzZXWKUjTH8zdr0O+rjOxeo3EZRFEfiuK0RpVjaCLO1lVhrLh3sf0w8lsZ4LQiO1de/rCERBQ8dghnceW4FSzPmzaQjy2ECAUDT1Kh0ogqDIZbUcVhcRhqkix8R4poXlSTjVNPoOSjI9vTxZEvJsCO54IjmKV6B1IlyrsaRxRh6sB+cqarpR1SgDCTqs/EtWF6uVQzce2xWE3lvuQ5+MhZ3oy4DZxT2nnSnQvNvb2zNvOtvkglJsdwHNQfLrUTlh7DD3zzjQ5xWEAAAjmSURBVHpLSyJZqC4WFynXQNlLGiKHVBTJEQpwRuBXmXyAJCjfCgCH1PVgNOPeEKh2Vt/9D7Y1p1+lpFTpvBGndNEfAJWxELcEoNfIxXqw7go/QWK6LZKB9YEC5ZkPYC7l2COK3L7Du+KUjynlwOUD0nwS0E0I3QR2Hh0KxVI4TEjnJ4wpkw6I2cWRtHWvIZnWIana/O4i2IO6e2h15oCs0alDAepWm56HxrvWHedGAlx8QMHDtfKGXiI49vTPQRyPgYww4MjhgZUQWcwy9orH3FhDZXPdd0aB+xvbiuNEHpfIsd3lHuLG1hUnjIM9IOg51hYm7KDAbmyP+xyRGwpmpMzQbAi3yqv/lb3o57w8Tka3k4gRz4NDnLcGjN66wBatpTGZ3BMrfitAdCWgM1XlrkSU7F/wF8hOXfH48QbZfM03a7cbwY69VqawUHWwzye+W1cY/nkU9Bio7McKPWoUKrIP41/EE/k9srMn7phNiK2NNTNt1yDud1b1d/G6i+0zM10YfqT66DQV7OpVH/NSVWp4NtRQ7/XwQiz/7ZiXnznomRB3S1USlpfZQRLd2eQzMBnTWzdug3fZUV1seh/bZz8BLqIuj8rOnnRMdvbEo9hh54limyAWbQPniWYLRu7JKIjNhvCBIG5e7QeGf6hPHTXE7mN6MsgWUt+oTVd/kPo/lLR3JMP4EdkS30nHeMgJzJMI4no44/5QjMU+zkF1UXfen9KdjH9ZoJ9U6PNOtVbopVGA67MJR9l65aNsJt+W7Gfbw6mNjV5emYzvmDPxNkCuAWBGjdHrDHy4sZ8dDMW/agqFw0syz5rIh7zJLMtFsH9Q9LXGzSnaYlF0mQC3iPhv2ysbdF/6x2d6WHYrNe8KuMH5+FdKvGZP+iLl/i68mGe2O7iDisEH1wm1yXQYFfQ7fMBfwX77LBe4fwoqE3YothT4I42m1CfzjZPX5yPqEnnzyH4PICuCR+wfpXg9OGQynuaXNPKERzCrrqbnq6l09aEuCs9muR9m3W7qqMJSrg//khlXqmJeqoAvp9PVttA+m2kpiL/S8WLjguKlAEFTunnhl0Mq1QExMogIb3uI+zyf1OCWLIawIW3qthyKHwWBu2BlzV2dXIXj00YehWgWDhNKQAMkUfRBDDi84gURrVHAPMLPclA9ysb+jfNhWqB8Svv5q/EgXV9myjSFCnlQvdyjKvHqRPAQB9zlUAwlD+5SYkcIfhMEer2tp0E9O7I+zLK3U7hdGD9PeeYhjnNXZJpyxWrfAwVlRCdxfTK68eS1BUt7FwUgHhVt2kRdBLcOrFigrpUd9nZwJ9B5+Q/196fQF7s6z5nwD3V6AdtjiQtkt3jR2wPgEeov2a8rYALbiB6ErFFg/sAvCjqz3U+ynFtUMfjB6fQxhX9OxO0FBT0TcOqDqylDSN7/9BF+hcw+EfOWAmW6HyK/oxzDFLpGRO7hCm6uhORJE7ieuJ+J4uVkrlAExKZwfyUN7QpKR++1/oley7OWMASNNRC0QeS6lzuGWT+GHbmV6acpxy2sk5eQPVPQ6p0ss8+8XCRXA/JrJ9gpKgRxgdVDHnWO3jr0GRX9Q3mnEzwkCJ6E4HcC7zF6dIJJpRCPhR28sD5XDY4LUYzi/U2qcqYtcPMaELeYeZdA5MNGAydNNNZ7jhpbncQ6RweSHQ7+doFeLYIHy+h40nShf4LxYmL7rKaHIaDxUfuk7zOkP0VVulXk7EQYvwWZpp72p1ctYX1nsB2WUA+fo16PEfh2tu336hLdf3bks3FBwJ0YvZCMLueg+okCU8ggRnjbw5rn1uQCLycHkRwfIvbNqrickJ0zcXa7vRKQyVjvbtNE8UrD94PnEyxRXDhQWPs+0PnEGc7jZCrqVEkUmvnk/HNV3P3dFyWTzRUeBfmV8+Sy3Yu9BLMlWbjX1sOyz6z6gYqc7MV9y3iwK82xnQ+jT1V1Px5E7sQgoozijq/JF4/vmDXpkmymaaXhc0is1mTxCsP3g8dJLiF/Nvy7CZJVwZMh9dkRay0N2nLd2uY2LWH7zXTJwsJYsufeSGMXdrU10NADuZmTWiLEzwjF/bGtasUyzrPP6dcTdRp4nCQJaSHf5tp88Q9lnqWY0xW25UwfCncNSymlk+3UZhevmu58dIJ3wYn0Xs60Vww4PXyMfDKda+qeADsC+g6TT30wD0k3vzruHi2Kn9kWX/ViHxqltlQ9rz1f9YStPYqPz1fEruZaU3s/TTJo91ybjIf+DktrW5bujIrhD6sS+DUuK3/GRcxluxaDKLoqKrpL1lSveSksyne4svcwMTA5goTMc+LndiVeWqUJ9wuu4V1elS88qUU3N9dV/VB5p9Po259++T9IyAU0srdh6dJ+o9jenl9KPZ5d1qNLytezibunds5u+qvlM3Bx/L6ML8dsl0uWLRtZNPwg4EO8PRHcg0TxAvtV3TJu9eO5Tquz8SqnZWc13RtP9JzhISdD/alA7GT73rf/QXPjlKgztqjFRXFuivhTxPtTvMZPyi1e+ZtlmU90uTKjd2TMylkjGnTPmvC8eSyD6pFpCnOZ/ZcbfiBY2iA63pg31D534iPZWZPu7+h7r8f42R+1DnpZjrR23zV7/IslOspgnSQ3u+mx3MymhcajLdPUZmQGpuS2uU1LDEzGwX/OSopXkHEgD1K9K4Lp0/7wdqDxL1esrM/VmQM77Lp/25sEdm+6AzdHbPvbdDkQzGAY3/V0S6NjbVRqQ/IZFNhu5k3kZk540rwXw1l+4zOwbEs3yM0Zv8LKWUGv2GQwWSy9DNnyn+hmMt76ktHTYwjLeLs2OcplWRld8yYvM36gnBhwGI3RWhkmTy9NL4H1i7ZZk/5tOJPH+rLJXaLvew2nl5Jn1tFoSrLxmim94eID81aHfh1m2F8pdy+y91wqZ25vvy3TWTuY3L0U65zpHZmurZ79GJZpchmv/jRerGYbd9I4Uq4HSnpi2qBAWUyH7fYfDXP3fahEc+OUyGje2QbLalCBigYqGvhf08BbJk/FYL1lqq0wrmigooE3WwP/DwAA//8LVNiXAAAABklEQVQDAPj5jmipzEHWAAAAAElFTkSuQmCC',
    'Área Sanitaria de Pontevedra e O Salnés': 'iVBORw0KGgoAAAANSUhEUgAAASwAAAAqCAYAAAAJU2bcAAAQAElEQVR4Aex9B4BcVfX379w3ZcvMbCoJQXoTIk2KgIGwmwAKYqP/RQQ+EKUoICWFMkCyoVeRpkCQJkVEIAImWUJXQEBFQHoLhBR2Z2bLzLx3z/c7szubzSZIUBNU9u097953z7n3nnvuuee292ZdetJsXckwPzN51lH4Bxf5uZKgqYkzN/4HZAOoAQkMSOAzJgH3GavvQHUHJDAggf9iCfyzBitknd8lDLgBCQxIYEACK00C/5TBEuA9gV5HLl8kfHoum3XpiTOHNkyYvaWBhcE47H1rUHfiA6MGTWhZqy8Mz7akDPrHWQVGZe+uq508Z/URx99fb8+9cPSMZO3kWasZPpW9f5WGCQ8PtnDtiS2f65uPhS0Oxz5WW0mbbYnVnzJnk9TJLTvVnzR7s6XyrRCt3NtAaQMS+G+XwD9lsBRIEFZj5bs7JwMr36lkusauIxJM8Q4XG0DkinRxzJfSo4cPjsVik6JAL+sLxa5oo46SfqlvXFdZf5o6eVZTa6mhIYbo9I5kbJe+dWlIJ3aMQU7LFWvXQjH2fXXlPSq0gR7bNx8Lx2M6IZXqWgfZlpp0UY9zkb/Yqb/ABbi4MxGbaEatb94D4QEJDEjgk0ngkxis9wB9mdlfQ8NwrUISfL5dFZcw7inC64QyYeW47INJdfothW4vgltF5HLy1Q4EQ8MO1CiwFXkL4DGnCoHqPMIoUWwCxfNQ/QuZ3dx5OTzutIbhEaqy17ATH0kzXHEebj9VrCGBi0SwpQo27EhEH6rIk+jJW1TeYn7bKRD3UWxRpuwPheAIZtDqvfwNonmF7OWdHmIzNMYPuAEJDEjgn5CAW44070DkLBUcDZWDST9PFIMCr9MB8eCl8Cc6rz9Q0R/x8Q3CCncjCqVAgDVV8H6osdtyuc7bEZfja5J4UGJBBKDoBH+IJ7uuMEhGics/rP3wfYij/cE8H+HGWOSniQqNL7YCQsf63cl0ny8GpU3po27CnC2gujUEvw1CKVpcBbKNXYUEflNXLl8WRnJTJQ54SZ3cnIyKcVXZl3GzIgRHArFjfckdxuffKnSvyIWcmZJrRgy4AQmsMAn8j2bsPqZeKorqPtX32de3IP2e7Hh7sHOOhGI3ERwiKmt5J0fQ/wo76z2kWeFuXqqNszl5QhQbBChfkqmvbXRdsWg+5nRIOVJA4lDZKCzX7mFgRmhEoSEOXjR0MQn8EB+4db3qlxj1QVh07arRY8R5cdgF2VsTgYTjICJakl+TZklHozXvzbYuLgNHUx5NFNT0wkvzHy67YG0SDofiro7E7HmFaTvObz+38X0a9JlQ8bSKqyF7Oosh1YAbkMCABD6RBD7OYEUqWuAsI6Bh2FAhCXbfdggChh19bu+gAMhwACMAWUXgV87pYXafUpKzHIGcQhjsnV4ZSXhqbcfYURo3djVJQ7K9Kn5sIE6+0h6vrQMvBdZlmss85Le0R2sp5OL2erTGksFcgd5P+i1T5aHb0ijvDMUDDfWFHJMt5Wo3GDFKgYmkfzIRL96E2/axmd1QEta4QJ9H9jSi+UQXOvcB5VXwNO7AxjFGDbgBCQxI4BNKwH0MfUwha6iKzUziorI+w29bGoFfi34SKi/AYT2G11D4P0Jkd4ZXipufbSzkmhtv4MzvcAFupDH9ViyG8TSyMT53kNfboDjQICyXLy8Uug0Pca2As6VrgbQ31CZwDzhjan1+fh6QxwmriOIwphum4u+Zi6e70P/6/lOcXEUHsMyMc/LTRdndKkbNAUYbwes6SyRxGMb8UhC0AcP9EriBhwEJDEhguSTA/vWP6QQYDdFvk6pWofvTH82O10BDth/DIxi3Lp+/xvBwQHYAsB1hxbtsS6zh5Dlr29Ktbcq4V8njPTQei2hI0zVhiWwjEuj8/LTGFw06ztllLobW2gyI7GIuxF9O/DM0TPt0lbG5vQphMyTx/gVRfQ2Qrwnk3ZgPXkc2u5SBSQ9t3QqCfdXj1tzUxidQvVT+Dqbz4r6TOvmhz1u+xqdEugcNexjA/x3ZnSp8YOAakMCABD6RBNxyUCcFlVnDrRC5mnAn4TLC3YTLuKSy+BsZvpgGoJb5LU+eJPvX3BB01qmPjs+Uhl2Qmjj7VKicyBlVUsU9U1QNmXuNCnZMT5x9YhUGlf1qUC/EkTz2ing5D5BFUJ2QWn/4Rjz1lLb20rtqJ4BAwPq01NUUPkC/q/6ElpGs7xFQrKaCkdX805NmHyIxFwmNIXHbikYXpNYbfp730UUi+lXyd5uL/KsAzSQGrgEJDEjgk0pgeY3LcAi2ZScc+w8Bwn0sUn5SLv4J+kULUVTIkzQYa4tgF5qAeqheFI93/jlRF++gsZlJeJv4z1WhrG4wZ4RvMP4+76W17dUFz0KjC2lIXob4jZDlZviluxWZzxOMmw6Rx+dmv9Zp7IUStLOcOUz/rEtqA8Tx9BT3QrlbJeguA1jNI0qk4h13Mc05oojTUI1huAGKq3xYvvbDs8fn+DzgKIHBJ/1+jfTEls+j+rIt4wbcgAT+kQQ+ymBFTGR7MfQqbjA73GacgWzdA1vR34TwRUI1jj6GVKi7bzzFg810up/+3fdLv1rKdyZ/FZXDw7jF/p2y04PSyY4bF2W/mm/FnJz4xPmq8UnQ2JlVaE/gxXwyeNIl3EUdCR4ocpM8v3DQEy5ZPDMe+geqm+T5QtdTjDstl+t8BqDZAVB45f1FsXjxmvpiOCPf1vWGxIvnqo+dwInYGdX8XaJ0WT624PW52T068p01N4dheCB52ysU/U5N0l3Rcc7O76EnP3zGL1smhy64FuIPRcMDxU9PHAMl/zdJYFkGqwTIDBE5FOynhIjQ19nzWypykAAPAjDDRK/XKY3bBwpcAcEvGdtF4CPv/1ZHQ3Lh9p22N9U6ddybnVPHvTuXhgJgfDbr287a4UN7paAvgBvrBq3Zxlb6Iey6aqtyLvuVRR+evTM3w5nW4jjLsjjQt8cK0LjZxvq883Ztt3jD983bwhYHnl5W6PvxNj/bWIDxhoHLJMDl/A4KzAtDdxGyWW9xKxX2vjWozOzMX6kFDxT2r0jALZ1YZqngRu+DB2iQrieeey68L3bPK+RaFy/+XlXMIP1hMaoSWiTOjJXeLpFwfws3M3YB4bPpsllX+Y4w21Lz2RTAsmsdRXjOCaZ3tqbnVSlGZe+ua5jcslVm0uxdDeonzR6XnjhzA/vUyWgyE1rWqz+5ZbzhqlA/aeamVTxn+2LfkGYmt2xbkbklWgaksy3DMusP3SVT13VYeoOhX2+YNOuLOHpGsi+plWvfgNqhSTW+yp99IwqeElfis7cmUhNnbmxgB0CDsi2D0pNbxtg+Z9+6ZHrqVPXTE2d/2fiwPOxb1fTkWdv35mmRhPpJLZtX6a3eVu7wbEuKKFg9U5Nn7ci6x+y5AnvfGqQmtXzBoH9eFbzdqI/23W0134ZJM3dmXbczfqt1TU18aHimH7/2+dpH5mn5riRYhsHS93gk/37cF0s0XGas+u+5tDro64mOVNkHeJdK0t8YlZj+FXWxDyXAAvX6Gg2fzbJWSJVMMSsNMHnW/ulJLYdS0AeaMphy9S0wk71vSGry7D1NofrG13crxe796YdkZ2SY1+6WZy9MbPlepYNYBlTYzMTZX6Jy7Gf4zKSWA6xcS2dozhqcLXvSxbHf7YjHf5wu6dGZSbN3rSpcheazeqPxdjH8mBPOaalRCwdXxdDeVTsyUjQr8Avq3tlUzosgwXmpUtR98uz0QOf1BuLPsbQGou6bg4DuwSB7W9wFsUNV9eqOeMDtimrOi3377ErKfj+FnKWKPambp0RwZzXUJ1arUlXayLlJLpDLU6Mbhlbjq/y5KLois0r7WtR9SRVHNjhxRxpY2Bd1Q+6BXh3E9IucNu5rPBqwrOsVuNXqVXkGfhh1udUt78DjYKhMTw1tXd+eq8B+djLT3GhpXOQvIFzWWdZdQSMZi8WbROXaITx8qtLXrT9yFQHO4PMFqaEdFAtDS7mxzjs5k/n+0vjw4s6Fc5dLDJMbNhy6lumtQ7g18XdyqXSp0RiI4tjho/LJpbJbyRHUiaVKPEQER5fjbksoziZ2K0JftyMr01yMlzcT1RMA+SaWvFZVkUudj/aj4hwkUhFOpWGWJPvXn8zItCcSe0ROpnmVn1BR9idvh0MwNV9OLaGwvpQYT6GfH5Z1XN+SnepuTHNCW3sq0zc+7KpdhfEnAp6g3wX0AOa7JxVzbRtpUkNze6jgbIH+wMol/ifEn1UqJba0fOo6th/pNTrNGpvx32fcj1X13M6yP4CjYncHY+Rn0aXL/ovswJuqyjquHFuiPUweAvyGbXUSVC4CdF1R97XFMpO3KNOTvbjjDAL1N7TaEh9AXWnkUBHsSXkPUwmamCbG6CVcKRGNZLnjSfOc8zjOq/6EZf0qjIKOKiHb6IuM24wddjOU4/aqThVV9bdXHx6OYx+vqUb098XJBzRslxuPBjx8mU2aF7nqmGDPGgvO7qjxLwyaPGtNERkPRcY5ty9p+jl5TFQnqrhmykIE2KO+NKTXyPcldhp9mTTrUye30SCyPWeS96XoDYcMPWR8GD+Ux23k76te3aFYuE2cOHNvC9xlRlMB1cnzszu141O+3EeUv614OZW4gLAsN4SzqDOtUZeFZFwdG+Db7PBfYXjFOE5tC+XUpmyg41mJVsJpXmi0RE4SyJ1UxN7RoDJiKsywDiLP9i4ZlvsSuY1pTuRod5wGOtFF8YfqV8mPFZEJqpir8KdbuULDRoN0P+MrG8hBLPYdymAsO8Z5AhxJ/0g4+ZOoHJmKdN3lLv9/jZAzU4HsSshRPs+oyr79l2M0SK/kEjInUMwSyhiC1CgU2MQUhmreRYlHClMaHzRoO2v8a4ytOCfl6onsDA5EY1KhW2LGYkQC8AwEXcx3VBSTdCFfeiw/rfEX9vmU4W0wEpWvUHfbSMv20v2W4g94DJA9U7Vde+AjrlARGm/Go4GKvAXIwtqw9LA9t5859i/INnZF3o0HMFgFM1Wxiy0P+dzH6dxkwj2siFoE8ldRTXqJJ/oQVILdKwc14/oq5foOZ6J7Ye/bumVWoeh747oH+o7xkTur8b7aBC5mfWdw4P16KpWqDNx8znvxzxiNQXvz+D8Dwmh8qtdHVAg2PR5LzoSwLLeQHXM+Fel9IvtvujMKZqU/z8CK65gcCRS6vRN1bOxLcomHftfe3PhsfupOj6YShauTia4nWX7FdRWjLURgynsNDdzGDbZnUcF8/E29vJVLLnimrXncnwpnNj1vm/nsDI0sO3BOLyokHplj5bY1j/99Ilm8JB3veNZyJU/jBfJkTVyuyjU33Zuf2nSXQH9GXFxC3Yb+Z9JlhuTXZMfk7F3vo4wupxA2TKcSS8jDe9kqVdIDI6enkGY9QF6bi56vDQQbeinfzuX9/ZlJLZfbkh49F1Vhf+rkg1yK3aQi9RL5JfI1sraYf887vUMFw8TrZZlMd0JttgAAEABJREFU7cn2O2iGM7BlGbvl1oD8Dio/BzA6U5/gwMhQjyP/91OPHhTgCHHRhj3Rn9jLZO8botBG5vVHlknd0KFxFet3ffKSDTqL+I6DnEBae5dvbiruP+xDUAlGRR1Nfsin3ENd/wXlsG39+kOGVZAfc5vPAyGmfZbWKOVceW0j5/P6ou48k7NBmvty5JPRhv30wH3iogX3efjdfaJ0lKocwPTXEroIK9UNq83YKGPLirnJcvw5tO2SNOVNnfzg2FyY2iYq124wwn6Mb2+eBtnnQoIFPsJ0QBZ5cfvD4vHxF6fKJ6eLQ59KT5r9XHry7KmVPSrVNRzk1Uj1zVHYsiY1edaO9kN9pbD2i22l2vVs01I8jb74x0wZqqXk4gtpzCSnAk7Xq7H/Nf6/zmg2y8HF09hjVaf6bCTB4+wkHSJuv77twQ63E3vGGcRtLpCLfBnTubdCO1RhoZP4F0XxvKp/laNGZZlibcAONZpp7g5c8BYnA6+o6C5L/QZZtrGr0FF7j1d/IHO7i3nsG/PRdZVN52zWAcEOcDrUef9M4PyDUPEaBPv05Y8DVU65vGdZac5KToBiKPP6hE5Fy/Evi8MXBNISOn0DkL8p8G0zZKheSkPkNMvH3QR6gwv0Z/Mxp4PPvW5U9u46PnwZIgE0eklc8BvyVO/gWEdqG5Ef50S1MvHwQE2FVtHF8l43ORt4jfrvZVfIVvaNDbTMIu3Vha5lYNhQ/pz2xCN/LWR3/aDQvNPzkQ+uIN0fCP2dCcDWyv3j/y3PkXRylSarMLNwQWeulGrIpwPgcKf+AqrYTV5xSVdNbGRm3SHcy5IxbJA/SAIhBI+yMcdy9BnNtMvh9AlAbiBczzwe6Swn0wBSCi/00dqRHOTUHeW8Xshyb3RwVzkpbwXRBPEhoY/7G5+1EyrGd5/4z0awtty4GpzYsmWYd+6YmPqL2SlqFbJ1/TrDvoCeiwbpF+CsirNngfjH2+vnfIDeS96OynJmPt81MZ9ceEku1vJqzzLu/0gymPkdpN6fDcVGAmwTCb5AY7dYz80ovfNOqb153F+Yx+nk4wjqxDpBzO/cEO6wNpWqUVWGMP6YSN2FEGUUltKX2mTwllOcwjI3pUn4Gv1P5Iad+GgKcDtCsYqHHBRTd6kCn2MmG2k5YTJikM7hd9CKDhKNR9s+aHiL9fHE9LpcmFmTPDTRYK9K+Z7ofXQekVbnr6cmPjyM4X/ozOCpExpOhLG449IPEKfv0ghcRRlNNGh/edHz+A+4rFJLs6Gwn5Sx33laQjAk/IACnrdYYKLxGFoBFAh9nRk7rvPxdN/If2e4trarTEV7nq24eiZTu4UZ0Fxz4w9zueJ2gDxFxe2Il5IfaOB2EGANpXJIpFOgQrwOD+DG9R018RGXqpuZTy64KN/ceH6uedzvOuML5zPuHSg2lMANs5dBWe4+5VD2IC8PquK1ssT+Cs7koLIR+rzFbR0CwBAReY7+Z8txRsuZzBZsF+sYd4j6mwl3CuRqdrR657C7F8S6haIfOsGPuBxvU5UTU2HjhtQ514NLxAJdLVNfu3qma/gaqWLT0Loh7V+g7LcW4HZ4d3tPvjbIzOdsaYeGrnEN3WmBNMYOSa03bEzD5Ae3HFRfuyrTcBAB1QNhpLI5eVgPIsabwZ2Ao/HUBodgZ+J6+APmA2Hgo0cVYstGrea/vH4xXtyEhmhb1v1eAW4ynlnn65k+J4qdK/uufCBnRaibyODLZPKE9PD8l+yUkM8VVy4m4s5HlCvWEpFbxONWUeUMS225/TlBeVnGlKQyyF4TSU2cuXGuXHcgy/k2+fgdO7P1XbIm9psnq1TkTFlzvTIK/wFXjxL04UTwHByuCUVPVcgsCGyfqkowEs5thGxLd8NRCSOEayuwapWAfkEUnGrjEgqA63I8xLg84d/q5iJVEvWzRTWu0EPSk2dtb++TpNM136ESbAORx8uuazgbcBvi7ZOaKxh3HXklT/oU/Z0WN4LUxeLRJg3c26rAhJnreIkSxjBHmjUyxWFbWHzqlNmjG7pWrWf+LcwrKZEenJ48ezt7vSEe84eKYLQT+UN9OWgVyM0soylTW/y+LRczk2bv6j0mME/vvM6k/5lyg0YPT0OwrULeUxdcnZs27pcGmpBfUhBPwOHLXtzqDFdcTbH0ggJZwlqIosnp4pj1DCHAGj7AFHW4wEAkOjAIov2oc0kP+Vl+2k7X5yzvBK5U5T6ZYlzowjUtrYHrijJsp//zqheHos3Un4uY52vEvco8vsTwOyLumkoezEc1MCPyBPFjNQiW6LT2srFodBug93l0L6lI9/Eu21Ij6rYmYQyq1+ebG6+z8hLlhG2v/JZ13pz7rpsQb06TUfzdyAens5yAdTq5oTR08WxUg8Ee2J1pno5KuMzyMQiS7ufg4Kki4ytbI5ZTFYQDg2LHivycO9+pHE7Uc94FlwFzSgyDsDoHi6O0R84xRKd1zwqJ+RTd0gZLMYOrnRmD4x0fQvRKVvqlPvylRHFcpox9U5NavpBeb/ge4t1RbOSN+tAsVCc/jVzwpHj/CFS40Y0P+uD/PcFsYxhLFB9lWVdAwdFKzoa4C9mox0NYB+5XeQk29oL1FHJ34eUFd9jGdyGB21Xdr6goawcIt+1mRkd5BKd7yAUVCIIf9iqn6t4qOMfiJZJpPijvGPfRLBGwcd324Mak4VXxbfLxQBCFdy44Z0xeNbqBfDxJY3myRP4SBX4KyPYqenlbjfxHTK+xEi+HTk85PS0elxXa2s1AVErnSNYq1DPK7y7wFM2ppwFzT8xLJYqZZPuTTnGqOPxJYy7u4Vsox8tpZP4g8H/qBn0PKi+wPc5JJDrfYh4kAU1dY5eK/42DTk9E4SLGVFxNGM2D+uvIy42MeBGKX3nRM1zc/V0hfxLFVbm2duZDLF0h+X4b7KN/xV3wWFDljyibmSH/yqJXWfa56tw99eXOjnIk77I+l6GM90nT65huJulunJdqK4+CnXjqiypykSTDXl0wvYHKnSq4CYFbpMCtfL5nQd3cYkfNvOchMllUHyz7mEaI/iyQSyIX5S2OtNf0nnSy1Na56Xanau+y3V9ORQGjetwcVlduFuj1QhlCYa9bnBtKcGwh1vISslmPIHhZoDTk+nshjQE8/l7ZhunJ5dPyljZYgt3gdLcPy3WDYZZXsGE/5ranlZ9CIV0K0bNoIHYnvpZQdUPF61GBj7ZmI44hzSFErJA9m0XZ3XJBFL/ZAyeTJ54M4Xo26sQgkjN4XPsKl2wvQdHsw/AR++kY8gEeJYca4hGFOwNen/ciMwSYRNw1gNxgwLo9GAsje+H1HMARhJ2IOMUdXv3rNrLG4l1Uej+Byn8tFfEOEZnIEfti4irKnk8+wmVNcBpUjgfkFgBXCXBcbdzdQB4q027GfWactVVhatMduWlNM+zTpt6KZxvD3NRxfyw0N13VzlPeHGc1Ofu5HsbPze7RkSN9fmrTBQWe0Baax7fkm5um5ZrHZRdD0015m6E0N11vZfTmy4Adxeeax11abRNGwT6tyk8b/zh5+VlhWtMZll9h6riHWrONrYXmxlusvCX526eUn9r4SH5a4y+I/2uuD3+Wn+lVjvy3T2mcaXl3ntP4Dp9/2td4GF0bT5FzUxtvRHafUqVezU33F6Y2/iqX/UqvMTW6/LTGFwtTm36Wn9L0UqG56VY+32NpDPIVPsad1XHW2Ge669Z4Mf15+Wnjrm5vbppl6Xvhqq3KbWeNeyDXXy7ZrM81N96Q65Fhvrnp3Fxz002dU8e+DeIsfW7K2JdzPfiqn5/WdA7l2Gb4TxOWNliKzWhND4mpnCHQcVCMVICdUO7kKHhjDzzKuHcZfopwG6ESz4q8QajnCPENEfyINEfweUdCmrBCnL1mYI1VWNjw88LfF1yZn9p0V+tZjcYHTPBUyl/b94Z9CzdlKkxrvL0wbfzf2q2TsFHzzY0/r0Kuuelee4fG/HyfeCrPdCpIZVPSOkeBip5PPvyLXK7rytyCzAPxyG2SmTQrmzp5TmXG2TZl7Ov55Jxf1pXLF+cTcmmOSjqfR8h9eRkID0hgQAJAOtsyLD1p1lmEqysn8R8hlKUNlhEK7B2q/2OwgufM4APxuAqIHftRwOn3T0j/DsGcHY1uz8CWhJXjOKLYaLdyCutTSpZT6Et3KzYMa9tEAz1agRrEuhYiexqDpMtmvY28n8VZFWs/4AYksHwSKOnXOcnZBYIZi1Db8VGJ3EcgbM1rRqeKXts7XaOwsK7VfpVgWaA+bhuStmlK+1ZJZi+Pdm/OVx5XwO37T8VTEx8aXoWKZeZBQG9J2ZbY4JN+31DFm5+eOHMoDw361q1CbhuTtZNnrUb6NeonzRyx1NvN2VsTlbeJmWclQeWmQvoGS+udtHMpfItzuDqGRGnE8Q/UcYrtDGflGljZw7P28SrnoJX0vLEO9q+/DF+FTPa+IUuVT9L/ekcZmmztcxQD+4DXjtR768W2S9tIO3HmBumJLZ+3bzErbdpDYLK0NlpCNtzArpv00Kp2Gpuh3Ngea1je/cFeDrW8BmVbBi2Nm7WapTXeqBuLdTabdZU2yd6/ivFpZfdNa89WbiXfybPWrOKMn+527mGcbWzveVXxls7q0oPt8agTPFE2OlBOPZFLeZau9sSWz1XzMr+ShmUsScz8jp6RNBlT7zZInzx7Q5MN60fdJ25JYoB1tfpnTp6zvsk9xTp301YJu3W9Oz7bx26omIxMflXKXp9tY3U1GfXGkU+Tj7WDlWm+1UHhX/feNavXl+pL0VBuNYnJsIEHYMa/tR/5b+hTcG+WywoM5WbkvpUjVRa4BAGVzDLjSQ1nXxi5BG4FPzQMaV1dJDzPSXSGuHBauVxzUnqDobuZAbCiB3eGo8IgOII0UyDhKRVw7jgKZAPDV8EUrjORODRQ+QXprw/UXZypr23sVRzWMcOTwqiMk1Il2OyzO+nRv0uEzh3WnkjsYXsOOe5v2M81hyUc2Z6M7TaisF2t/WNW8na28QhxUzuL/pj0xFnbUhkqHSM9NL+uBuFx3fhuHrWUOJzlm/HvLud/5J4Oh63tKINI5VeE6UE8dnWuXH+QKa3Jo2HdYZuj5M+lnH4H0XsjH91eKtX82BTcRNCZiI9h21ybztT0fvLV0KWjYxpemK7p2EqL8T2jILjG8o7U2d7mnQzfRJgeQ3QV9aMxLPv9QpU7InW/tHiDGOQ0X0zu7bhfmQ51XSvLYFjH+HqRcrMrJ462l5FJ93PSV/KjPz3wcm59qrxeWEx+I1T5NeNuIEwPEF3ZVdIfWr+wfAYNya8mcT2fuFsilhv3ch115qA0jbPhK5B9MEjXdY7n0cKlDV3DNqnELePWEU80BjFYHW9kftMNJIbJS33sTN1MpWu/ZTKGyD3c5rmHun1zuoQDKoNpv7xTpbEbU7YXq49+6310M4rxi+pDjEH1OvbxmtAFh0s5PiWF7YZVo2Gf/0h5svPK9/gAABAASURBVBYT9nNUvdEWyHRFm8RVpvOkfM9qX7I+G6B8iS/575p8q+0hPPQS0aNE3M+oIz9qmPDIoK6yft87dxdEfhNTP73s3NEfZbDeBTCHoIRuJ9iBVm8iO9juZhFNwcyyN6w7dFxMoyz3uvYgYZJgrsybvcv1Kv0V5hQYynK/oqJxVXTA3rFSudJL6XAbEUsIBgOyM4BhNLjvGKjXd3wYLGJcrwu7ardX6PECvMnN+5kcf9rh/PrdrzAw8dqj6ljG3lA9nELdpSr8EcmAOuyaGGf//qw3P4HuLJDNi6kE+XKbwutWzLNA+QUisruKu7ShyzeZDNX5VbzqV5l/0virAORt1BSX4LE38//igC+jhm20DqvwLhQzCNYmJ4YlfCXThbXUyUmAbAzgCspjggD2LeG+AcKDbLT1IoMo242ZbmKaszDSIQq03gvWE7i0j8WeoJwt3xkQ/buqDILgBdLbwQqNIF6EuhGimiD+YYsHcK8o7nLi32J4K0TSc3IMlF2JA5uM8YrnJdIM8atW87O0IpiZDNrfhMgwAVKA/sHi6b9H2oNiPjrGZi1RTJPEr0ve31Dgfg/MY/gIKekupKu4IfarCyoHQcVeN9gRHCQriH43gR8s6kcK8Cy6ZTgD7Ku26qHf6xrSydECPZM0IeFc0p8GAfujTulIxsf3EvYEHPQH7AMbU/8voT7fwrqVKZehPWjUxosWPpB57YhifLEhGz2cpLIW28tWWFXyiu9jQT3ru62HHpMuD9/GIqNYPAmR9VVkZFhXDrCM9lD17IPlDVjWYUzzmEDOYns9LiIZx4hluSfU6RlERISqizOD3ZjwbK9uAnfRB7mYHuidnE+CAwiDCFXXAcGvBbivGrHCfNGFEsnthc6ak2wPieVcB5GD20opKhuf6FTlATvlMChMG3e5neQwutfZcpcPYeSDK4ifEoui46Kyu62t5r12QKWYKI1SyE6k4ZG4btcQjliN4eV3gpdCuEvqSuExFPqxVI6XIyfH1XWUe09PuUd4e54nMQY5O8Xpd3q0/IX951NW2wPqJ5HbeewgW3iHrdhhNhXgsnxiwcU80LiN+LMp99+TZtdSyfe2J58Hi7hjKjMzPlSdfVDMQ5cLTIbwsP8l+ZYobrPn3NSmS+KxLjMkQKU9gissPt9cOSW7F969TMSb5GXr6gw9EuwOIKeIWuhTFdBaza8n7TWLeFINXip4Q1Wvsfh0ov1YKG5kObuLBhzoSVBx7t5CAuw/wfkKdHrI2pVo3sKuhJ3Gb840r1AOOw5ad/hHz7BFFkHkBivLwA6QmMUSzgP7QNCukMl5niLywOcmJ8FpJHoJ0KVmQyxzBMQVnLrf56aOuygZJo6KxXE/6SsuiOn2zG8kHxY6QZMtzxheHldkm9ar1x9UJjrLStGvPew0mH0ySdI02+OFVKL9drbrZCTknGUZrGtUcWm87J8mgxzx8BQT9nUbsHJfDbqQovHajgh7iS1Bv+reE9WjvQtopeU6NuIpZPjtKnKF+Rdu31k4s+n5CLFLWYaj6d6UfsU50c3Sk2d/Iz1x9tftt7MqkX1ugcif+BgGLppsdInIh3aSCB5BI3u6aCQ7UVAkw+WAjPKR3wz/xGWb77mpO/2BMvmZQNdxiWDx8lKwrZWdntjytdTEmTbL+CdK+O9KohKsSo5TYE+hPKzz2rt/L1bkDsBeDRGnD1JpM17dGoyiSsLzdqUCTWER+6qK7ZUaavlBsWEM0Q/SlX9QMut42xdK1RTmUp9nMt/NuETf0LYJRDCOmT7Kk+F59MFyB9Ew7W3pMpNmH1s/qWVzLOOam92jw6l1djVezRBVqNhvVq3rcptQz8azvnVQv6iC4E2d2xuQd9l3rgDciDDQjz6wUh0CVfv9tRMzk2cdlVqGvijcF5nXC4Xmxr+i52rjqTWN1ePwsnZP1GLPZmvQVdT5aWnqoNSV1V71qBIITNZ4QAS/oow2D11g/b6K/kf+h6zrHTRyoyOP/V0UZZYiXkZ7iNf3WeaLqvh+oVx/kr20nQda2Q/7J5dV4WRk2SUTHE1sPb90AUwiGtoLaEUG+7sE068nPhysEYaJk3VYQW709SdbMc8dC+oWMOcPWOZiYwBsB9WDqHAHchbWZ8QjJV1brvMv7BSnMjgCivM7kons4tFgrKPB25uCm41IZlAmbyswdon9ByZcfidkIaABlwwUvbMGhYw3HgV6gIN8tLIuf0H/uZSCPWmcfy7AVWQyAPxDEBkCSGcUymKdyma5WvY5VWlnR6oYJiUxO9wMxv2ObXYIO6UNmvhEl8iqlP04CHYD5KtdsbrEXBoZ9XgaQIKFblou1trrOKtBxGZqjKZT1MLrNmA6D+zqVCtGlJilXCS+E3D2UmxdH+QPY0H4Cz4fCpHHXBCbyTAyPCyg3wj1v5UAXALp2zR4jbacZPzSTqSextVmPLtRDjsHgtX6E1GPBitkqQ+WaXDegfCvX4Ig6e5kXhOgMoR5X9hZ8lNsw9vIGia3bEV/YxqdGwH5I4AA0O3ATXWGP8ZpyavYLPlucdhHHb4MtfR9ki2jPfKFEme7MgHi7lWPPRFhur2w7vok6wnqOHbK7zgX7aLAgYxcl9DPyWtRbVRg4W/2Q9jjEBbwA6jsxZ65DxT7M3IYYWm3AmJGjXrPFNv2trhn1F2AclnARj3GlnpqL252Ry++X7pbMZcr/tp5/z0BLqPifIPm+Hs27eXeynjKYSM4DGd992RTDxPxu6Pcs0QRjrmq1oBMyiyz2WXIlPG9TrmCjVZlntYxF/ZGs/Maj4GPThRN3NMn/n8uSEGFgMznTPM+D9mvJu5a4L3NYmoChxp2BpIAMFlqYHsnw8SLyYuRgA/RipjaN3zvi4i12efwSS7BH5numED0ezEfHdz26ns5S+6Spb+oVjrkWHbaHxBeiMc7/2y4HphPxs61dJEEhyWjREtPfD9PhaNcCqoZeCx+2VLxOvP/HPOYGZXDU3MvzXsNUNFi8lvMYBREVvERvkWNSlE/Gl1lQOMTkf3c++rdGcZHGMqRNaXosX54Puo7vI1YYi+M8lS4Na1M4pZwNpsqJOf/1gXBoURcI5A9NZCvW3qvfi8ohqhiM6/YlXgHh7HpUJkXnz7GqUer83KNQv4K1R+QfMl0soz2YJ9si8/5UywMpwD+MFHkqS9nLqtzJQCl5a4oRIqZB4S+jjxrSy6764fq5UkiWgl9nUBgP9VKo4XvEkEFZAwDK8aJo2FMDjvxkbQd9xZKqcMo3JDC7m1EWvZCshgu4FJvYb4t2V7pCH2YsRFucF28pm3RoLfh5W6IPCGCNVVig9V5+xVID6+c/uv/KcBljAwn3Zj2eC2n9RSlypZ1kx4eacft9Z07bsKsBwvATslQxbkg5su1tsxomDBrbarz/sLpv7qYHUxUKCC+rcpjW03Y2Z/HbqL/jTtleG8+15ktJBec1d6803PzsztxBiWvU+9GcNm0QyZ7/2DrKJmuseuwdb+mPITwsYD4nvrHIy0ED70gXrhER8TY3r1Ahj/eqRR5GraoKm+MTMUskek0Z9NPslPswOedoXLnoiXeCRLqvlTaKVPsWrSg7MJqOzkPF6irteXl4JNmrg6I7X9RlfA8ei+9RUSbFWiKxYI9R6zZUGN1Vaf7gPoA9hvCaFHh3ilqIW7bZZ3okTbSAK3Gf8aXPpyXSpgM0PdShc0WN0xtOHyHykwoe2siFe24EbzfVeDsQA19r8rrDxieaWtrn+tEfg/FK6qyRmaDoVsqZGcR/AWi6wv8WuRxAVS+wP6/BRZ2BpaPg8St7lWwX9CAj8RwCLy21sx5i2VfDQ5UAOoJi90y2sP2J2s7xo6K1caidLLzOerBXUyQdrwtyxkTZmiWxgk+EIdbQe6d4lkq2UzwRujv4oyoKAL9FecUPAnEHqVY6YRYTG+gwp9Adq6zX0msFqqKsR3x+JEGqZrOgwaVd1i9ijNfS8mDw1gwKTO8bS91/lDSj4bqE1HgRwGyHaV+EjcuN+UG7WbcEN4Cimsg+JYE4RC1UybBujGEp6TTtd91AU4B0OCBP0tnSDSf1K8H577HY/WTvHN3iaKJ4+bFtklMbLdT+ZrxZ5AuYt/60pjh3Yj/wbvSmHMERXafEqhHBlqSRwHcBZEfajl5Qmq9Yd9gW0ymrMfQiNzTHszjZjEpqi6b9blico5ALyPN/Gr0cvk97WGy7kjEj0hlar/ZnY4t4+UxZdvx+XXn/YPINi4eeEQzcPh6JR31KV3TeUBdx/YjSQu25xre6f4dicRRYRBcSh3aT0VuSyak1ziIuE7v41xW6b2AHMKy90U5+VWWupaHTsxPbTw43zzu4NpS6WAo7gTw7fZ4wNmIUgX5VHXqBzmve3WQB4NUCd+sGCUsvpy6W/jUCq/npcr+oHRx6HfgcQFoMaF6HXFLuCAWOxNlPS+dSh7ImcwhEIyCyvPq0cTCM6r+iHxz0yF58sf94SPYx55yqt8clEmOZBMGCmxmvFQhNax1jIPj5KenmGzW56eNf5h5XUl4l0t57cEwq+7+UUlr7ZGq2T0q+m/FY3pdZ9EfUyjXHeJVvssEb36UwerNaxmBeSxs5/SklkO9YBzxMYInrHTHZmyDw1MsmEss2ZxCfoVwgSTL1zAOMfHcp9MnGFdDurEVgOxADaQSGEUPiH/NlEZVjgFkjAhmIpCHGB7OtE9EZbEPRFG52Mk4o6PC4e8sPxaPwru94k4KcxMRz4aUIYBcXlNOPFBb21V20JfhhHtW2BwK26h8VB1Os+/IYJe6RYx/lOWMhEOFRxH9kgRupS2jjY0VD+htDyfeXiFYokg75FAXcATW28TrViIyAZAM5XqORPFfdhs3z1M+bYkngu7lIQ9ayhL7DeV1NdQvYbQEmMee8AQNWm98DaJwifYweQt2YtrefbDcovSb1IO7Ablaaha//tKdHzhAY9VqO8FxMIvFVyH962w/6p5sxM67I8tdBMg00ein87ONhbAkNmN6zDs/t/DK+4tCiV1MmnvJ2xpedR3mPcupZ97UQpDx83Zt94IZpPmrSpBA9nSSEEHnA7zLsp6j7q0N45/A/rh1XSmwgZsU3S5XgzdU1A7NXhKVAyDy/6DgzAin5JILnummWnynvGdDOW8Ttz9lPgKQq6iDDzGdkM8b48nS+4wjCmidOu5NUfyWPLQTYqLyDATWF8dWeXLiNlEFt2XkoZiLsx+iciXCxJ0eYDu7FzmjWWZ7sD23F4fnvcpfhDL26r5OAfyVhZ/kKrl8kptiM1W5FNCrmfEV6P4nFAE+hSvPjbkgkiMNwhBH5ONyVH5q0wXVD0rbaha948tyieGrEHM60cXLvScnxnZ+6rg7aZR+TGU71nn8OB7vmmovgroo/pAvyYT2N+b3Kr3RZ2IdT4ShTK4N5M0Pzx7/dm1SJrMj/NiLO45C/VE+gYsXnDMmPxdPd7mEzKiWHYoekV+QObowtekOy8egkMDf8GE7AAAC10lEQVQrPgrPrtKYXw6lORO0v274/yWotkciTD6yrHoVpox9oa4UnqmQoyE4LgrDowvNTVfZ96JGX18Mnw4ldmrr8/N7O4B9tKsJdxFnzn2WXkAq2f4XX3YXtXXWvmBpDfq3h8naIJLY+YavwFVbleM+vCUZxq9qxZzK3pbFW34hglOMvgp8Pi1VKr0c82FLNc78dKLjyHxz43WcUVT2KDtb0/NYl/NqY+6PuG1vbzwnwuTZ5dD9PB5F0zUhp9qJqJVThfbOmsdCQbY2jL+K7Glajbc8KINJVk4VwkguHJRoa6vSVPxsY1iIP/yIqP+R6bSBhfPJh+7pNv4Vqt5bbmrjjUHCHc2CfkRej84vSF+S+6D+Dfafaynfny7CHwu9xAzEksW7YlF0elvMvRvBX1blpep7H7u5rlx+NhQ9sy02710mqTjrF7UJuTDgQD8fczpcn/5RTRvSoOemjvsj99ROsjhxejTL+lF74qGWT26wKsX+h9y4rLAPnQ0q71ZlG7uW4IyzIRu5Dd8LHB2qBq0vrdHlp457rO2spqcX9bxbYx3F4vt/oziXJ0pW3nyOnqCVM7/ym+9TGh+sHCNX+chmfWu2sbX1rMY3DOyfvYIdAn0v0trH2YavguVtZfQl+58I97SHKe1H1WceZxf5aY0v5qc2PmJy6UtnOOvs/dsjn21c0F9e9lxpO87CevPo1x698p469u1eGgY+PHvntgqPpOdjxVl+VnY1jfn2bDwZvT1XwWgriao3trnVxfQE1BfwsvytnZn2LeMffcoiGiDfpi9Gh5404GV5WLnVssy3fJYqk7SWpxnNNuq0gYUtzlDLAtNV01/jtaKn5NtkuCz+FrGPGO+g/tprH609Ol717UVWk43Vob+BnM9+w7RtxouVWU1T9a1+Ff6oLxZnk4cq/X+3warUauA2IIEBCfyHSWCFsTNgsFaYaAcyHpDAgAT+3RL4/wAAAP//2tEz1wAAAAZJREFUAwBlQ5WrhXRb+QAAAABJRU5ErkJggg==',
}

def _logo_area_sanitaria_bytes(area):
    data = LOGOS_AREAS_SANITARIAS_B64.get(str(area or "").strip())
    if not data:
        return None
    try:
        return base64.b64decode(data)
    except Exception:
        return None


def get_empresa():
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT empresa FROM settings WHERE id=1")
    row = c.fetchone()
    conn.close()
    return row[0] if row else ""


def set_empresa(valor):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE settings SET empresa=? WHERE id=1", (valor,))
    conn.commit()
    conn.close()


def get_cif():
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT cif FROM settings WHERE id=1")
    row = c.fetchone()
    conn.close()
    return row[0] if row else ""


def set_cif(valor):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE settings SET cif=? WHERE id=1", (valor,))
    conn.commit()
    conn.close()


def get_logo_laboratorio():
    """Ruta al logotipo del laboratorio para la ficha de registro. Si
    no se ha subido uno propio, usa el que trae la app por defecto."""
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("SELECT logo_laboratorio_path FROM settings WHERE id=1")
    row = c.fetchone()
    conn.close()
    ruta = row[0] if row else None
    if ruta and os.path.exists(ruta):
        return ruta
    ruta_default = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "logo_laboratorio_default.png"
    )
    return ruta_default if os.path.exists(ruta_default) else None


def set_logo_laboratorio(ruta):
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    c.execute("UPDATE settings SET logo_laboratorio_path=? WHERE id=1", (ruta,))
    conn.commit()
    conn.close()



def _logo_predeterminado_sobre_blanco(logo_bytes):
    """Compone los logotipos predeterminados de las áreas sobre fondo blanco."""
    if not logo_bytes:
        return logo_bytes
    try:
        with Image.open(io.BytesIO(logo_bytes)) as img:
            img = img.convert("RGBA")
            fondo = Image.new("RGBA", img.size, (255, 255, 255, 255))
            fondo.alpha_composite(img)
            salida = io.BytesIO()
            fondo.convert("RGB").save(salida, format="PNG")
            return salida.getvalue()
    except Exception:
        return logo_bytes


def get_logo_informe(empresa_override=None):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT logo_path, logo_nombre FROM settings WHERE id=1")
    row = c.fetchone()
    conn.close()
    ruta, nombre = (row[0], row[1]) if row else (None, None)

    # Un logotipo personalizado siempre tiene prioridad, con independencia del área.
    if ruta and os.path.exists(ruta):
        with open(ruta, "rb") as f:
            return f.read(), nombre or os.path.basename(ruta)

    empresa_logo = str(
        empresa_override if empresa_override is not None else get_empresa()
    ).strip()

    if empresa_logo == AREA_SANITARIA_CORUNA:
        from utils_informe.assets import logo_por_defecto
        logo_bytes = logo_por_defecto()
        logo_bytes = _logo_predeterminado_sobre_blanco(logo_bytes)
        return (logo_bytes, "logo_uprl.png") if logo_bytes else (None, None)

    logo_bytes = _logo_area_sanitaria_bytes(empresa_logo)
    if logo_bytes:
        logo_bytes = _logo_predeterminado_sobre_blanco(logo_bytes)
        nombre_seguro = re.sub(r"[^A-Za-z0-9_-]+", "_", empresa_logo).strip("_")
        return logo_bytes, f"logo_{nombre_seguro}.png"

    # Respaldo: si el nombre guardado no coincide con ninguna de las siete áreas,
    # se conserva el comportamiento seguro de no inventar un logotipo.
    return None, None


def set_logo_informe(nombre_archivo, contenido_bytes):
    """Guarda un logotipo personalizado (o lo quita, si contenido_bytes
    es None) para el informe final y el Informe PDF de colocación."""
    data_dir = get_data_dir()
    conn = sqlite3.connect(get_db_path())
    c = conn.cursor()
    if contenido_bytes is None:
        c.execute("UPDATE settings SET logo_path='', logo_nombre='' WHERE id=1")
    else:
        ext = os.path.splitext(nombre_archivo)[1] or ".png"
        ruta = os.path.join(data_dir, f"logo_informe_personalizado{ext}")
        with open(ruta, "wb") as f:
            f.write(contenido_bytes)
        c.execute("UPDATE settings SET logo_path=?, logo_nombre=? WHERE id=1", (ruta, nombre_archivo))
    conn.commit()
    conn.close()


def _mostrar_previsualizacion(nombre_archivo, contenido_bytes):
    """Vista previa en pequeño de un anexo: miniatura real tanto si
    es una imagen como si es un PDF (se renderiza la primera
    página con PyMuPDF, que no depende de ningún programa externo
    del sistema como poppler, así que funciona igual en Streamlit
    Cloud que en Termux)."""
    extension = nombre_archivo.rsplit(".", 1)[-1].lower() if "." in nombre_archivo else ""
    if extension in ("png", "jpg", "jpeg"):
        st.image(contenido_bytes, width=160)
    elif extension == "pdf":
        try:
            import pymupdf
            doc_pdf = pymupdf.open(stream=contenido_bytes, filetype="pdf")
            num_paginas = doc_pdf.page_count
            pagina = doc_pdf.load_page(0)
            miniatura = pagina.get_pixmap(matrix=pymupdf.Matrix(0.3, 0.3))
            st.image(miniatura.tobytes("png"), width=160)
            st.caption(f"📄 {nombre_archivo} ({num_paginas} página{'s' if num_paginas != 1 else ''})")
            doc_pdf.close()
        except Exception:
            try:
                from pypdf import PdfReader
                num_paginas = len(PdfReader(io.BytesIO(contenido_bytes)).pages)
                st.caption(f"📄 {nombre_archivo} ({num_paginas} página{'s' if num_paginas != 1 else ''})")
            except Exception:
                st.caption(f"📄 {nombre_archivo}")
    else:
        st.caption(f"📎 {nombre_archivo}")

def _widget_archivo_con_eliminar(clave, etiqueta_subida, tipos, valor_por_defecto=None):
    """Campo para subir un archivo (anexo o logotipo) que, mientras
    haya algo cargado —recién subido, guardado de una vez anterior,
    o el que trae la app por defecto—, no vuelve a mostrar el
    campo de subida: solo se ve la miniatura y un botón para
    eliminarlo. El campo de subida solo reaparece si se elimina lo
    que hubiera (y, tras eliminar un valor por defecto, ya no se
    recupera solo: hay que subir uno nuevo). Devuelve
    (nombre, bytes) o None."""
    guardado_key = f"{clave}_guardado"
    eliminado_key = f"{clave}_eliminado"
    reset_key = f"{clave}_reset"

    actual = st.session_state.get(guardado_key)
    if actual is None and valor_por_defecto and not st.session_state.get(eliminado_key):
        actual = valor_por_defecto

    if actual:
        _mostrar_previsualizacion(*actual)
        st.markdown('<div class="marcador-btn-eliminar"></div>', unsafe_allow_html=True)
        if st.button("❌ Eliminar", key=f"{clave}_eliminar"):
            st.session_state.pop(guardado_key, None)
            st.session_state[eliminado_key] = True
            st.session_state[reset_key] = st.session_state.get(reset_key, 0) + 1
            st.rerun()
        return actual

    uploader_key = f"{clave}_{st.session_state.get(reset_key, 0)}"
    archivo = st.file_uploader(etiqueta_subida, type=tipos, key=uploader_key)
    if archivo is not None:
        nuevo = (archivo.name, archivo.getvalue())
        st.session_state[guardado_key] = nuevo
        st.session_state[eliminado_key] = False
        st.rerun()
    return None


def _acordeon_informe(numero, titulo, completo):
    """Apartado desplegable del informe final. Conserva exactamente el
    estilo existente y, si falta algún dato, añade únicamente un punto rojo
    al título para que el pendiente se vea de un vistazo."""
    clase = "marcador-acordeon-gris" if completo else "marcador-acordeon-rosa"
    st.markdown(f'<div class="{clase}"></div>', unsafe_allow_html=True)
    etiqueta = f"{numero}. {titulo}" if numero else titulo
    if not completo:
        etiqueta = f"🔴 {etiqueta}"
    return st.expander(etiqueta, expanded=False)


# ============================================================
# UTILIDADES
# ============================================================

def _slug(texto):
    texto = (texto or "centro").lower()
    texto = re.sub(r"[^a-z0-9]+", "_", texto).strip("_")
    return texto or "centro"


def _limpiar_para_nombre_archivo(texto):
    """Deja el texto tal cual (con acentos, mayúsculas y espacios) para
    que el nombre del archivo siga siendo legible, sustituyendo solo
    los caracteres que un nombre de archivo no puede llevar (sobre
    todo "/" y "\\", que crearían sin querer una "subcarpeta")."""
    texto = (texto or "").strip()
    texto = re.sub(r'[\\/:*?"<>|]+', "-", texto)
    return re.sub(r"\s+", " ", texto).strip()


def _nombre_foto_situacion(codigo_detector, centro, zona):
    partes = [p for p in (
        _limpiar_para_nombre_archivo(codigo_detector) or "detector",
        "situación",
        _limpiar_para_nombre_archivo(centro),
        _limpiar_para_nombre_archivo(zona),
    ) if p]
    return "-".join(partes)


def _nombre_foto_detector(codigo_detector, centro, zona):
    partes = [p for p in (
        _limpiar_para_nombre_archivo(codigo_detector) or "detector",
        _limpiar_para_nombre_archivo(centro),
        _limpiar_para_nombre_archivo(zona),
    ) if p]
    return "-".join(partes)


def _nombre_foto_plano(codigo_detector, centro, zona):
    partes = [p for p in (
        _limpiar_para_nombre_archivo(codigo_detector) or "detector",
        "PLANO",
        _limpiar_para_nombre_archivo(centro),
        _limpiar_para_nombre_archivo(zona),
    ) if p]
    return "-".join(partes)


def _nombre_documento(centro, tipo_documento, sufijo_extra=""):
    """Nombre de archivo para los documentos generados (Informe PDF,
    Excel, Registro para laboratorio, Informe final...), con el
    formato "CENTRO-TIPO-fecha-hora" (p.ej.
    "C.S. Carballo-INFORME-COLOCACIÓN-28082026-114500")."""
    centro_limpio = _limpiar_para_nombre_archivo(centro) or "Centro"
    marca = _ahora_espana().strftime("%d%m%Y-%H%M%S")
    nombre = f"{centro_limpio}-{tipo_documento}-{marca}"
    return nombre + sufijo_extra if sufijo_extra else nombre


def generar_plano_con_punto(plano_path, px, py, destino_path):
    """Copia la imagen de un plano dibujando encima, en rojo, el punto
    de un detector concreto (posición relativa 0-1). Se usa para poder
    descargar/enviar el plano YA con el punto marcado, en vez de la
    imagen del plano "en blanco". Devuelve True si se generó bien."""
    try:
        with Image.open(plano_path) as im:
            im = ImageOps.exif_transpose(im)
            im = im.convert("RGB")
            if px is not None and py is not None and px >= 0 and py >= 0:
                draw = ImageDraw.Draw(im)
                w, h = im.size
                cx, cy = px * w, py * h
                r = max(6, int(min(w, h) * 0.015))
                draw.ellipse([cx - r, cy - r, cx + r, cy + r],
                             fill=(220, 20, 20), outline=(120, 0, 0), width=2)
            im.save(destino_path, quality=90)
        return True
    except Exception:
        return False


def guardar_bytes_imagen(file_bytes, prefijo, ext=".jpg"):
    """Guarda bytes de una imagen (subida o capturada con la cámara) en la
    carpeta de datos y devuelve la ruta del archivo guardado."""
    data_dir = get_data_dir()
    nombre = f"{prefijo}_{_ahora_espana().strftime('%Y%m%d_%H%M%S_%f')}{ext}"
    destino = os.path.join(data_dir, nombre)
    with open(destino, "wb") as f:
        f.write(file_bytes)
    return destino



def convertir_plano_subido_a_jpg(uploaded_file, calidad=92):
    """Convierte un plano PNG/JPG o un PDF de una sola página a JPG."""
    if uploaded_file is None:
        raise ValueError("No se ha seleccionado ningún archivo.")

    nombre = (getattr(uploaded_file, "name", "") or "").lower()
    contenido = uploaded_file.getvalue()

    if nombre.endswith(".pdf"):
        try:
            import pymupdf
        except Exception as exc:
            raise RuntimeError(
                "No se puede convertir el PDF porque falta PyMuPDF en el entorno."
            ) from exc

        doc = pymupdf.open(stream=contenido, filetype="pdf")
        try:
            if doc.page_count != 1:
                raise ValueError("El plano en PDF debe tener una sola página.")

            pagina = doc.load_page(0)
            pix = pagina.get_pixmap(matrix=pymupdf.Matrix(2.0, 2.0), alpha=False)
            imagen = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
        finally:
            doc.close()
    else:
        try:
            imagen = Image.open(io.BytesIO(contenido)).convert("RGB")
        except Exception as exc:
            raise ValueError("El archivo seleccionado no es una imagen válida.") from exc

    salida = io.BytesIO()
    imagen.save(salida, format="JPEG", quality=calidad, optimize=True)
    return salida.getvalue()


def extension_de(uploaded_file, por_defecto=".jpg"):
    if uploaded_file is None:
        return por_defecto
    nombre = getattr(uploaded_file, "name", "") or ""
    ext = os.path.splitext(nombre)[1]
    return ext if ext else por_defecto


# ============================================================
# LOGOTIPO (para cabecera del informe PDF) - idéntico al original
# ============================================================

LOGO_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAABkAAAADeCAMAAABSbjycAAAAwFBMVEX9/v4Ce8QChMrK1+q0yOMvhsnl6/SIqtVrmc/a5PBSlM6Wtdr+/v51pNRJi8mku9ytwt7AzeW80eckfsU7kM6lvuCpvd6avOCaweFeoNN9sduewN7///+lvuFgj8qux+QDe7+AntGww9/V3u3Z"
    "5PAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAADdAvdLAAAAMHRSTlP+///+///+/v7///+9//////7/////vf//////i73/vf//vb29AAAAAAAAAAAAAABMCEqxAABi"
    "NklEQVR42u2diZriOrKgs9AutbxgmKzTfWef93/GUYQkW7JlY0ggyURxv9snC4wRYMev2D/+619VqlSpUqXK1fJfH//6qFKlSpUqVa6Wf1WAVKlSpUqVCpAqVapUqVIBUqVKlSpVKkCqVKlSpUoFSJUqVapUqVIBUqVKlSpVKkCqVKlSpUoFSJUqVapUqQCpUqVKlSoV"
    "IFWqVKlSpUoFSJUqVapUqQCpUqVKlSoVIFWqVKlSpQKkSpUqVapUgFSpUqVKlSqrAOFVqjxP6l1XpcqvAcj/JT2tUuWJQtRV1yhX6mrscPei7xR+zWeraujLwq/4zqvcFSD/5yyaKlWeJkybqwjCCaVHcp16UIR8KyL3rhc+GyFV9X1VFDlevS+pch+A/Dcj/1Sp8jw5"
    "CHrNra46Ywy9ljmt+VahZBcUFHXHdlXzfdn+8N9jBfE3AOS/f1aAVHmmsOsAQoTW0mnkaxQKNVZ/p7j1qgd9tiolA+Sv+9JtBUgFSJUKkLmStYcDs+11ABH68J2y9yPiZxNdBcgXAXI07PBH9xUg3+HCqgCp8voAob8YINUCuQdADhUg3wOQtgKkSgVIBUi1QKpUgFSp"
    "AKkAebsYyBcAUtOoK0CqVIBUgFSA3IICXnPgfhVAWNNI3QxVzVaAfBdAmNaasQqQdwAIV6Sv3/5vAkgjpRCSVYJUgHwbQCwkAbMKkDcACNSb1m//6QAZGGNRG7Dp73soFyla2gndVD1bAfItAGFMG2qsrAB5F4BUF9azATKwRjYeG6zRUt5P22ObC66oqHGZCpDvAog9"
    "OoVkdQXIGwBE1T4yTwcIQ2g4gqDHqZHW6sb9fPcJgFgKvydpBauKtgLk+QBx9oc1ClqM2K84sSpAfgpAKj+eDJABgxRChjCFs0UE/PM+ce9G+IJj5QgiHxNJH9DnVkMsFSBFcfg4gk7hhBpZAfJ+FsjK2IGNiQSbnaK5b+vZFzpmLl63et7w18ZohNK5ymtdHvZkgDAp"
    "TOvUu2ZBH8O/OxP//TXlrk1oZ8MdQTR7gBXizCfsCFsNnAqQkjh+hDtKHcULA+SH7Zv5Y5vW3wwQ0O/ZyjgZJdX4ycPzp6Dz86oZ487fGiutaDsy69TlTqnUfC0qZ1t8Qvk9zVLU2sdQ5Y8xM7g4+XIr6CsBAgZH2xphxyi3s0iAKGAxfFUps8bddvH2JdSd8s563q3e"
    "CmPwA3x9uVV+pQXS4l0JFoh9SYAopxuOR3qEe3+mDrYnk2wMLNmYpoEnTf7km2+2HIWCquzoW9yTRzmMbgUIdG3usl8KpwCA3u2zuTWKdv7hIHx2CrK2rq7r4KV+DM6lt+7a5C05id2FVe8sGD4ChOIbBllZg4ofg8w+Rt4lerGGxwNEio7QDBYQUXe3gLtlv6qRG5l2"
    "zHPfirivkgfWtXDfucvj7JZb/VgVIMsaEGvPypkf/7H69IIAgcbl1koQa/Mm99DS/DNtKf/p/tWNGsOpkk8UaHze0aNKb7XO/D2qsgp0Z/WfBf4ctSA/fhbb2H/+zc6ML4L1QodiXPBLAQS0llNe6SNCCAOqtxPCPcOnh01LSwABLZUcmC+rF9BqGfxEcMIuM2mobbRN"
    "9V1vJHRmjgepVrb+myQwHmAECO2MtHEt4YDWrUGmp4ePAcf01Ahp04+R/gLKrUEK8jyADEwLSvtutntnoPlp+1WCsFkPbLis72koAD6ot9g4Xte6OrIqQApVhNZtMf6KL9USPgogaBjpBuN4zOm+83HUZZz8tUsxNLrkjiZ5WOCEFTV+Xe45UyIIvp8wPjAJZ4gKTlFh"
    "y/KZfHK8z+JyIT9B0IdkPd0IEN63wpG4TV5GrbSGgJXn9LRTa+GpXkr8whZGFnx1s1MkX547fTAj3DfR0QwgrWwa3aVncvtyHaHhPpORwaEPNACA4FsT+EbdWcm0EHxl0xiVfQy3Vwbjj6YfgziYTPBza7BNIyl/GkCAH4QsSHGAzHkCSP2KQkYH1nL7c7cc4QHspOnL"
    "AwOnVptUgJTSeOnRqYVXrAPxyWFsyhjTScyGGH0adXWQ06gwOLVan05Bk8P0DDtOveKQtezOVHxD6Z4hcbuu7TkCxEqdv9Xp5M560gmI+NFAsDRLcROPGBx4I0BU60BqZ5rX+p+Ow5ciokukB88LL57C8VVkp0i/vMDeZZxbgYWTqPaIAS3iIwuAjJ+1lTaPbLg1uNea"
    "5CrqhTXjx5DjvIAcIMQtz72uU08DSGPdByYFnc4Qd19T9k38vnIb04H6Lr4m4FN+frghDlWdV4BE7QY6ELWg+eu2lRr/dbrNDnkMQJxCtkH9o5zACom6Sx1tgEoqNq6DU6fJp2e9QRDf2G1F4V9LDak6h53wWZyFA0WWatx5z4Zo+W9Km8QmMgiZ9BDGQpDpFQBChGzd"
    "lSISvYAA4REvo/7v5cqPRMBJZRpRogucoF35rKC83UVmepVcs+6im8LAawDhpJu5neDINv8YvXvnuDtopRBlgLiPera6/dIcrqsAAj4mTgqRCea+XwXlf7erelZ2JGI65T3cWI1sF2kQXa02qQAZXVfSejd+dP6giNvKQR4DEFDIkGgcwqMt6vAUIAxyIsujdREgenwW"
    "7Qcd91Tu1p3YkL+jZvoc9FgKEAj5zsIf+E0xOX5wcHlprKvx6+2BOSf4990nB94GEE6spKSz6UCvBCCg0cUFgDhtCKcoP7sFEAqmmDOApqiveztHECna/lqAuI/R+TXwEkA6dzHTAkA4laLP1/BggAwSAkIlT9UBPGng1mNf0CrluWwQNbyDEVLik9sh1a6NFSCwLYbg"
    "uekoPR6PqJetMP4f3meE8u0AcYhwq7CfwUcEDl5h7WcEiDsXmBHHTBJ1AQCJz1LjbZkQdeSggJldLuNoTywaMTlAlCLp+1D3TYGJoyfdBvBhJzBzgtESiMIkvTdBbgOIotKCR0VMkes5QKLmXQOI+9Jl7ygi2pL5BgAhKxeog4Byl9qklQAgKOEluwHCiZTkI19DChCw"
    "bEoAUdMangMQ5gw15QBS6nTYyA6jCrcCZFgaCEkgDoPpwxeV1pJPXLU1l/etAeKUmYUcdGdpfNIQl/SVX8fpH7i9htRvuR8ijwEItQetP0cPkLszjlNeJlog0AmonFOLALHjsxCydMpcBh+NwmcXa4aHx1ShHCAfPE/kJcCLgx7zYCDNCGIwaaYYhtTBDrp3+6mbAAIx"
    "ZKeJnXZJ7IQcICKzQHjhFB0cAsGOgqkBcY7iyzCEbhRmKHQZQFpK43tuAyTNI+7gY6g2XUMGkLIFMq2hnALwCIC4D+QMgpKZ0ehWuevZNHc0ENLr48tlhU3JbuadrGH0NwYIbsGPCjfT6eWRxTs5PAnHYD7RNwLE3evyIGeRCj4iAgCi11NqPEBUGgI02n2vIXKO9ski"
    "TqzOGphUBMhSJ54Oh38mB1YHt2yCuxDF8V4t8gIAUZ4c6thKUQRIl8VASkH0ECYHWBY+Eu+MbCCYvQztui2xexeIh5sUIO5YIkI+7l6ABEeZ+xh2Mkx2xEDAs2MgObhd/U3vDJAG4nUpQIbJFAGAwFputRMacSabF8gXs26HpmTgOGuq+rDeFiCQBmuOmMvKSxVuESPc"
    "V9Ipd4tmOUXPB0irD9aQdQcXAORjCyAiK/74674jTcfwCtNmtmhFRTOF4bcB4vQRm2DjT8j0IjdYUeTKnU2QmwASfFccPVnzLCyOWVimI2ka71jbzfNTeC8SL8ZvrS8rybOX4+vcW4/fGAKk544bXlXtBUj0XWUfY5aF1RYAQoL/THVafMGleBVAwNThI0CGpKmUAwg8"
    "191aCpKmAnCllsX66mtZtw4gpUuW1nr09wVIbHuV6CE1XXAfanHb8/1Neh8DEGcy2O5eAAGPGERFVNT/J2ZndDgiVWKe8CZAIAk1802Bh2xBJAxp4lnuGwW5CSBOYaNVkWl/rAPp4ac/djatA2mkY0GQSQnboJAhHF/Yg/hGJk6EafvUnRgjHVSKKU0ODR5OjY9q7waI"
    "tN20hggQqAPpfaqF1MU6kLgGd2GIXj0DIBr9dQAQFjvyOtFNw4bBl6Pw7sYo+iCz+hZCJyFjrjoWwN+o75kuA8RWgLwrQOZ1q6CC6WfXxRD1sWtbeswVHT9+2l2BkMcB5G4WCCQhOYCcVby99Ny15G71U/LFbwJEQcQ8S/HRK9FyjK3fOQpyC0B8HPtjpv0pJkJB/xEs"
    "4R4VMtT4yViFeR7DCxCFjwnBxUxeiLD4vgFJFb57XfgXyYLd6DEjnQBFy3cCBNfgOwUka+ix/5ZTnxC6m7iSAIS3MYe4l6IjTwQIcUqXwbdMaU9pBytsQgjjCwDpUoDQVggbed/SsZaT3txwq1ogFSBzB5bx9x1Xowo1vnbbb0vxH2PZWyz6xUza7wLIGZJ416IcXwMI"
    "xstneViqA1U/3phbAME8YJbYFarF05V0qgEsGfLdAFHdmEFghInVdAAQp9sgf9smKVS9jDv6xC5N8qGUsab4i6M7pXcKLRYohq8gqHrwI7U5QOASaxXfCRCupjW008dwAMG2f/gxEr/iBBBYg8p8WU8CCMC2aTC3m3NsVtwK3fjfgtPbAHIYnJWlRp8zga4wjc+uhEZb"
    "3WSE3FoTUo6BVIC8L0C0Hjs8eZsDfSvwxL/hCfTvwPv+9Q7pY2yCA66abwuiY56VuhNAPkjiwor/4jOr4jS93QZAoMLRvXnHU40O8Cnnkmm2zsFnAQTsOTO678b4MxTYS695RYKEYhYWhqHVGMneYKJ3n4hRj4noEIFTCJUBBJpgwVQ9sg8gpI2Z0yrNGkMOwsdwps9U"
    "aj4BRBHRhJ+LTB/jOQCB/jGQnIfN5iE7pZXB2XkrQHAQSDe6qyBxd7Q1mG9hFfZ/5Nauv40p1L9CFlatRX9PgPg9MB97Baqj0SE08hcBIsO//LPUfHpjhNA9TU4elsaLfqJiOykASHNzDATUlLMZ0m0WfCOpot8AiPrroGCTmofVykR0A0rwbn0zQKAPlA0VmRCnmADi"
    "w+XOajCTzijWgYAysp0/hdOJW6llcJ010aMHTZ9EfGvRhNeNAMEICVX7AIJ9FKePMdatWOmD9+AimgzDESCQ+xvX0Lk13N5RcT9AovUMuWcSDQafaI6eJW8kcHq7iwlcjJM/UKG3apgQMj1hbsqcKt6ltQ7kbQHCAjagQxTePdC0owCQkFhEsFMqmCqoq74HIFDu51S+"
    "LtcWBIDw/VlY8BE1newNx4fUtQSAkWMIfRMgBEo+0sA4MdgbpfgdYIGI/maAQGuwRgcvuZRNXM+UxuvsBNvG6HIZIND9MHjaoZvh5kdyujGUKnBQYXJ66xCiGAECUHC/w06ApGuQ4+/fi1AdyRWGZ9QcIKBfpzUweXNSw16AMNbETpyqN9gqmECj+x77C4dPyvvb49yM"
    "eVBEFx1tE4RI247WCb1khPiZUbNqR5Y0Sk5uyNmJwksrVd4AIHA9QQNCZ4z4pJciQEJ9NkZ+bXBRfxdAsIOhrx8nywFNsZCQq+I0unkdCMeiviRsrqAuJNlG8yP0UUyWtgoQCMDPgjOYv7XW5Bytk7P6XoC0mFUV+tUIHWtgZr2w2o1WJoCE9BRNu/mRiAwZVwoDIlMD"
    "mHH3HQECJX7u7cQOgPg1mLgGCdGTAJC451atNYs0XtUGF5d/obbdrQTZCZDBt7QZL7VALliFgHD/2CDH2JuTbYeAkDGm2Y45UuEJPgY72SY/JI5sZ/Mk5EUvrPnk9aHxL63Vhb8eICfq/VGC7QeIRgcBhIe/sRuv9n3cPxdNbX0vLHeXpimMag0givhQRNItNmBgPOI/"
    "zjhLI8OrAIHk5llYPCTx8rIhJTbqSZ4EEKyOGCs7KChrNQMIBzXXbwDEKfIuO8Xmbz4CxJEheV1nbRv8/2MYxV18sqVjm8MNgKi2EbTPPsYcINCzPsJ9BEj+8aFi8taevLsAAjm7YyDCcwI7R8KGHRPUVJpA5RTwcL35ESMeY5Np7gejT0+0/WidmNVZhXh3IVVlbqcs"
    "uvFiYGvIjghxJ5jwfmu6sJ+Yu8+MGaaDh/UTLeXCUPetlzbjN3370n4TQFo7AmTFheUjwb5Ph89Vp98HEBzDgS2nJJTQ8zlAcB5WMvUjHXAXW5n4wOXxrz2xvFADsrIS1xIxp7xeYw0gWF2el3yAg6hk878MQDju71U01qaawQQgH+ib56sAcdeEgBYFIRR8vtCVkEgZ"
    "ig6hviSEkKG0yL1uBhBkDDZovwAQ8KGkH6ONu/wEINh1vJ0DRCRrgJ20uDUtbg9AnGbNjWaOA2eFN0DabJ/jA+BXah3vOAqWziJmPoxgGJ9wvyzbWCoOhHEGTHoQDG9Pw48eTzlhgMl+CIu9zQgZgEEoOyZh4YDdePCweE6sCM532+DTtIbCa9G+uvzrbC3tl7iwuhQg"
    "PHhwCkF0n6ACEeaQ4wTJrd83kRAjjtpPAsl3vKGde9ZhPZ3NMevGa303XjXLutKfKsNJ2vukDBCn0LAERH0sAHLkq6GcbwYIxFhpNhYqJDClAPGJUGu1EsrpuTZtSbXIhuXJjHPuLpuYj5EVHYK+tGQOkBZuvMsAQcSpbA2G+iD6BBBsehUqXkLRIYc1JOkSoM0fBxCn"
    "kUy/nHQMWQrGtN08J+SGZNsGCxKHlAF89DLJ1I81tZ9uVs4U++mjAZMdxPQYSXHP9rNgPI5VVFN+3k07b+bnw7e+OubSp9aOWF3bdTBQni34AafJpfXinZ2rDPFrKL4SX2uRIWzH0uB1Rkj2KwHSTilV1m8MqV3NwkJtOrqwmm8ESBgqFed5qOsAckgnQGk5K2t3h2gW"
    "M4+gZRGz2bTXMkDmPUwCQE6vDBDcc6crdvaDNzZSgGAbKRorxmdRAujimw/DXXQlhAnkYduNESdBPsbeh+nabUNnAOHoY7kIEPwYdJZY3PqPMQEEfW3+XwEgcIpsmi5AjD4KILAzL6XAqryqJn3m2jmCTl0liiok7qrJJ5alY6kA1SLrtEnaFbuDhim+ghEW08VR8zgt"
    "N/EFpfOsOAwzbm5SjWO1kVNhF7HZjs5Lm2t0HD2/Ib0vMS1imiUO1pJA8SckXmxxQWdLG34jQGJKlTtOnn1A3efnZgAJ74sZr/7v7wyij0ZIZ7wjy3RJq9sxBnL24n7o46wb7yGW0aMFs6izhz5V4YtWOAkkw0IZIPxY6KJ1wYX17QBR3cxnA1UhuKXIAYKQQIUs05gB"
    "bEOh1WKbn6KdlVNwevbFCT7RyH8d+ZhCzwccCpgCxBclpQDpigBZfoywhtSFlRAouLCgjaRRM4/ejWlxlwGy3mh9O9Z3zb6962mbBMaZ2wHHqw9CFaPThUU17+7voXgmmjU6DnqZhbCA9ikL4x7eBwUGPFMeAlOX9f92tQn8zBcOlomtPBvSxeQFlyTguysTpFjxkr7U"
    "OyA3wzwbS/stAIlkAMeOxZQgZU5sBpCgQrGJSJi1p74dIH6LhpM1Tnr6sUvt3OdZWIcw1oSd/vl3YbRsMpgQ85vt38sA8SXn88TiyzGQz28FiJlHvHnr1TVo8aQeMtbY9d5dHwSCHdAKJHfboafoYx5l9b5gK6N/gyymTCmDv6K7Zk1SMy6cupqC6KYIECg5pLPmAf45"
    "kgOkCx8DZhASf8L8t5lq4+8OkOk6mOUF+gdm/xwNtivUDs4ZwXECadLX6McC629UlUPjr1Xe28IWms1GS6LyQ38+TnRA6UKCSvTpxHBF/gWumTiXAJI4HullgJBsoekHudzqgXvjbAnq5vKVwFX6nZYBMoUKfyVADt7Bwo9CRgdMeFEKkKBNwZyMDiPS2W8HyNSemk35"
    "l1gHoldTSWMMBE0XJg0tdicXzjTxzj20KzLGFAGijqLUGxG/sBWAqO8PonPazj8/MeHn7ZJfA+4xfwWgFzh6hJ1VAebFzPnidMbsR4eXjzJmANHZW/PeL4ZQmg64bU32kszhRNXKx8A1+HMRtfwYyr+3gg85Mz5vHTR8CSBDE1puYdeSzGPFvS03c2mp0ONlvwuogaYs"
    "0FQ51VQ46idN3G1iSYjvtgwX4VIzzVyxCuIcMazOi4LRdnfUMKt6cjuw4WsAuWSBHL4GkDHMM1wPkDFS9cYAYboLibwwFSPcXiIHSPgZ0KOvx7+/qxfW4gcUOo2D7CkktAT7Ca91w0XnEgt7Vc1m7UZKAPEOLLvIAcVCwtU6EDhP+71ZWHzpeOelZ8aHc/2xeYqPxa46"
    "3WoXXhf33VndTjJiYLlxXz3XB9/8GKrwTqtrvwdA3Pa9g1PHoDm2G5kmktCxF7F7Hg5wG3zU+mD77dUnwcQhM4UGUe3RCOnHssJGdKsAaeZOyDa07eLbahiC/rafoVw8GCB/vgwQTOGwcxuk2WmLqo1cBxnz338vQA4+XMzJWcYtNIfMVQ8QX6Z9Mn6z0jJwYB39zlyz"
    "lwAIRmN1OmN2Ry8sq6BrFaRfla8RzDCDpDTHycXYpxJA1LIEJASlL1Wid9/cyqTKfeQSQJwqgX06Fn5IDCFMWl3IkE/AcbJuPAIMOd/2f1/5YLh/IHCySEWiWVmh220PMZl6J0CkvGybOSUiGpFffar/CQDxhgS7CSA4aqZ5Y4D4PonqKGKROZgd+nCaAAJjNnGrjxF0"
    "//f3tXNf7nyOsLB4pewDiI+OM7mY9JQc5G46QrWex+EKAMEvTxeCbpu9sMwr9MKq8hSAYJtd6EyFWbbMh6HPJOQzu3/2fk9h3OPeRBhA7WNtzj7FM+YawD07zP1ndkqtJS2kRsMjqzGQBUBgINgOdQo2iJi7sH4EQNAHk782B8gicJWpC7OSo/sOAAFTg+CFqk/sBKn5"
    "sOsBgEA7EPIJnh7/oPuWYPar+5t04vsGShWu20+3MHEFQKAXFtb9nexn8T1DBTn1fRTVBYDARX5gsvBR1dEwdirWFvDX6MZb5TkAgcqLsbCDwayVVvjWJZCEjLWXGOUWput855FQx6F2RkGmuB8YLQsDyBo6RUKoMSYMEStmYS2D6PMshVV3jsjr1O8QRL8XQPL5Wv1i"
    "TB6Em4YyQNTipflrV7Pl3gIg2J2H0pARa92fDg+Nf7jDxiYnAQ/iBBBUq5D7zV4JIDQZ2LG7nbvyTbCKjl1HSIbVhjYZRbgKEOzMXt7rqBZCSxvzQI4fFSBvABDd9gTHGHvdBHFtSE0TQYv3PntAwD2nRo8IjpeC4/aF0KeM0YJVkVQPhji970ZsdNEdlu2awMm2s0kY"
    "MX6wcEIU/SIAwarUKNC8rc37ec/tiCwTTMjppfDaWfaF6sphkLcASKy6O7Hxbzb+6R884YPJsfvw8SyA4Gb+aoCADQL5WyUbBAqmmbVQqbggzAIgeB69MkiJQpfK0lMkzC6sAHkHgGCr+TZqGazCwHkgPpcWM8051HSjoWJCH8UBS+R2htETgBxLNWt+opTiqT+Gz4vM"
    "k5zjJBPOmSlyp6ry3WCSBDNzUyHhQwAiQsUKSuwew3Mn1lBeA8SNRvFVMLTPXlsuNH8TgEBZxCmWRgQZ/5weOxzSB18KIJJd7cIK26MVGwSMGo9SMY9szAESpuDScv7Oykz0j9eZiV7lOQChk7LGiAhYwE1yzQNfmhZHaEVVM8DztLNfdmGFvn6zPgJum7Q2y3yqJ8cW"
    "YdLuvFChsc2slQl7GYA0LJdGy7wJXps6opI14MfIX5rMVgkui+aNARLAcMhhUfrrcA0+nubCAj+RuR4gXouXv1OIkGDF4SICPgeId4WtBeugkITpRRTE9xO+7zSQCpDXBkjSZATS/5RTjY3MdsgwfR4ml4yqZpBttxMgEERXcVDVsJoLlrTcwulexTp00K2QQwLip4SJ"
    "3ZMW3D5fxJf2rdA729Y+ASBLWGJtS9oKLnVEZQDRpThRNn/evDVA8u61F0Sz7wWImu31UW2ezBV1IONAKfK5YiD4tpHgvFucaQYQvjU0aiRFu+ggZUtcqQBZ/F7B2fLzAdJO2gmy7xSmf0qbRC5gU4saM9FiVnT7AAJup45st8AdYOiWb4QYCrCLTaCCh8dAxSj2iroG"
    "IMbnKMcOgug2Gl4UIFNLlxjr0fsBAjRWqfXyzi4s3ZK9onaWgDwMIIp+0qNaqOibABIUeSGZl/i5KHIZ5c4BAg5inbcOmkVTwL8G0fJ8yVIf9HW//xsChEPzLOhTS3rCfzpAMgUCvirspj5FLnqJDpW8dBDGd+zNY8LeiZfmDPpBIR40xQPBseNbkJtx0NblLvdjvhI0"
    "NJ5eOo6pHF4TIGDjJXYEScPoFwDic9WSW7pk9r1LEJ1dlU1K9ek7AUKsBnUca4l9AcusEh2CC3xeM10ACA48h45Ux0JzXV3+vnOAqFAdQ9TKG/oh61glMi7Z+1Rvbbv0RgDx6qg70o7+bIDoNi8zgM68nCeTVsCdildRHnaGLfJegOAgp0sDNLCXLo57Kh8IvdqXTWjb"
    "GF9ZKYYYC/g5KXSwpSuh+lcAyJ90G5flFFwGSJqrBnHYoQJkH0DsNwPkhDNxnBwhpRgMotOUfV6aSEi7sSXSfCa6I4jVJ7uwQXwJJSvZJn+xW6Ma42tYS2Nmb9gl3R2Nt0HOsGJcstAw9PDz7rvqX2iBYAb+0Wmhn+3C0qYT+Vw/d8th848k9B0DIyzPndrfzAS7rV+e"
    "chSlqLptIbHD/Q6jwb3S1y024CeljBTSXjkM44kAScNQ2XiUiwBxv+o5sV5Kp68AKQJEfjdAMMrnwzZhJMh01TuLBD6RzsX+JyaVAECyIg5vgyzLyMlKqyqwQJoIENVCdczy/XTWH9j4RYZAEy758AB+/EKAEO6LFhT/2QBxFkeeWwvDpdznMqkFInoYq6XZLPlWyD9P"
    "kyHN302u+QkgxftZ9ePkqmLN07VjQZ4IkDQSyfuEFDsA0iTR02K79veJgZjj7iDI98dAwiQpX5HCThqaQmT5EM1CnY8mBtSMzPqQQOKIXr65MlqXYilQ76VtuO+VkbosybXlCILjEwNl/JIpUR8/HyDQ6xDH2ZYf/urZlW/X+BB8jEvnTwHIPDTB0IqeXFi8t7KbjY8N"
    "ANHP4wcrZ4OMAFE3AWR98uELAKSZUpSzUPgOgMyzhGsW1s/IwiJdtDz8vt6kNRiY27IQa8bW39bOOv3AgHVrF03XlbGyNJ8O0lxsNNmIsbIo1iQFijxAz+dDY5H/8RFK8XqAcB4HpxQVKb8AAUV6aPBOZzW9MOsLHs61iZreKlvBvF8unxR88tyCIhuA4iqbBrPidemx"
    "yfy0yPxVY1ditTJZ5iqAOFzMQw6Nu0aMSCcOSfdvncebIWQhm6cBZN7D5F4AgSLJVwVI0o8ua7yyByAbb/1mdSDsdDrpXXI6fXsdiMJuPmGokclzsmC/D3rhnIhXcTzeC7Q9zj28tKOLPCp36ZqOlN99PFrRc1kozeesU/oJfVYtNkIw9DE+mevbuRNC+yTDbjmwe8tQ"
    "Un6uJ3y77iwqVc04PhrHEE6mYU/p8p3AQ5X1ToqzMTiOzVCpjsr7D/V0LfUNftAscbD8wWGRMNSExrYe+cvUR+HBmwHSyPn8bCxkyyulpWzmIQzmzGn2NIA0KxP8vgyQy2Ohvgsgw1cAotO31u8MkAPbL4fvLyTkWOGEt3nBf6KWkuZhLfaRuMssjJRwj675PqZRFGpF"
    "OF9ZsR/K8xCH0tUAIb2Bvj44WM7A9L7sq0Fby6yFxyCHCPoJESw2M9EMjA/Dp21N0teCCGtwjh08Ok7BVb1JdRZ3C/GXDBht6cBAOquipkLatnzBwKTeNsmfWOg0bFvYdmHprT8zZqC2Ux5E6I0OD3bjg+pGgPyBoa9DYcOfVKIbWch2HW4qxLsdIK16BEA+Xhcgf4ap"
    "Tcv1ALEVIA+V51Si/6Sw8EPPfvVEQojnCK8enaKUIpsliM2Z1lSKx0tADkeAhMGznQlt+3AA3jQqkujGq2JsajdN6rbpvBXVxhHrMPIyrdecjc4lbbPmb/lQprFt0jx12UemG2nJybH1TSlIK7RMABIn7zpO3QEgpfSoPJNbrTQFfyI//rjvTZUu0wqQ6sKqAPn9cr0F"
    "4vbxLZbM+L24ScNJ7upxmn6tSYvTd2ZqRsejW8q9SLTRGFG+U3kIF2koV4XBAeFdo29jFSApImYAwVLMZqX+3wFEkKQwpxDSFenSvXMMACLIrJAHAGI2TrUbIGsekDR414nmibAo5qWG4uq5jZ4ChP8ugAzDV2IgNYheAfLuAKFJqF+kWlq1UJe2MrURY775djT0B2+S"
    "h7ERXzBhiJYxQxTO7H9s0A8LgPi9v1OnTTMaMDlA3NtAJZw9l3N+zOakYFATWYtMDwZkxSKX25lL7Z6v83qAYIN1lSQOAEEkG74XIMF/qO4GEA/d1wWIMwP7W7OwGvGWabyQ2OF2lheSqZjGbKKLR+FhumEVID8TIEmiGWrWLKpr6GLE3PhKKztVJFL6MJTKhTAIAGS8"
    "I2WwMzYBorFttioBBOIsOFtcXQ8Q0NMFswoBwpcPinsCZAposEaH6lcsl+TRozh24t0X+xh8XHIYFhvr8MQtAFlUe3wJINjh8XUB8oU03rTF3fsUEkIXMBiVdIEgUFwN27zNo6CiYeuoCpCfBBDvtErzSil4e0qzhDhtre1LF5wQ2cNgbBQAYjOA8DJA4DqNQJsBpMVW"
    "A+5gcgNATBEKz7BAIKdKh3kUwkeNsF16679jBcF9Kf3UCd1cJgjsBSGxz8o8AD9AghcInKwCZBMgycuzcfKXAZK2wHujVibQN+p4PFJjtwFisFWI2Zw9iJ1z8Fy6AuTnAwT9Ru1YGi2cbQBRjIKlAU8WlLdT7LO+e7yDFuIFgJgdAHGXoAjP5gBRBkYtgbnTXw8QB7lS"
    "atlTAIJ9CrHVLe3H/AOppQw0gRRoxxCLh13M38Wpt17yQncmYz6Zsdd0EXlDgNzeysTZLn2SAlH6nn8lQLSGJhpoLm8BxGLoECdmbJopPgMSO0JVgPx0gGAII1ZctqDked+WWrEqE22I2cNy/jCVIWC+4cJaAYh1kIjhdgeirG+EM5Q42kj8WoA4FJWW/iSAOHhh+VJ3"
    "9KF5gGADxQjgz+JYFkP9EXsAAl4w4hPE0966OKsCyzYhEa4CZAsgeTNFux8gQ4qed2rnHvpz+HazWwAJP8k2QMLd3toKkF8AEKdd42xgMD3gx5vl2U4avPTTOoDMwxIOIL7a2gEkpniBX6vbARDKHah8NWdmgaijj/2jIVJo3oeB/LXEKd42xasSAJJkYSVUIRdzsPa7"
    "sHAcNxTKYHkiNNZExY9NE+NjLVTJWnnZhZXODUznVUHuQdgGwNCo+wLk4xaAvGoQnWUjq69p5+6Yn9ixir7PQKkHAIRXgPw+gAhfHlHe5q8AhIgFQHoHkM4DRHfeqs3SeLcB4u5dbwClAAnGERxtTRFj0AS9xabIbTuvV+emXD8Cvbw1tkx2r0rqQDScqz37mhL1NYD4"
    "cRvCQEWln9U0Bs0PMCUVBzH5SUx6h77RYvpl/HwP390nmZK0sjW+kIX1LgAZmmyq9NpAqWUjlmHKoVvaLr8YIHh5+T5O6vipV+rK8SoMvwoxuikfhueKFxy1TemgCpAfBpCPLg6BJzJUWYxepPwHk2YfQNyR3gkEdSC4xe5bY5NCwi2AhLQungNExbgMJmOVHGk4Ui+I"
    "6PcDZHpZ8NtlD9qNuUq703iBIGFiE8TLJytjiOaJf2KPukkBwuF7hWi6EIamvUVvAEjBAonFPHSlWWJwban55C/+wgBhTZM1WlBrI21JPhPd50CkbRtS++8XA8SxA/KvPoPzgH46U1ku1D6DDj1jgiTUi8F1uYADw3OF6wl8rZDYxSpAfgFAgtM6qFnwIi1+xQIpLgOk"
    "CaNOoafJOD/iAkDg5oU7OQUIgeiIz3t1h9MSQJLq8UWHrw2AJJXoZKxEnx7sb+6FlUraNmj1icNOfZ/3EYOONF32ie8DkD6MUnfqfAUgod6f0P5lATJ7bgCUtyT/qlh5Dc6+aHz2HKTHwf6izRu4trI5vAFAIPfxGJvLKUUgf2qp9KHVODlGm1S5v+nS3QWTxI9jozoO"
    "5zp2i4B7BcgPdWHBLRETaLvCNn8tBlJ2YYUYCISQ/f46uTcvAARj3nDKFCA2lqBDqtbSv4YxkNXvYA0gCmMghSC62XMJ31KJfo+6v7yV8bIl5nUACY1r+Gx+F6cYs1HQ3WwFIF3XE2xmQEsAeYlmivMu3ZDKkH4akvecn9YA/b9NItAmNIe0e+Xw59cDhJ3cfaemDoHY"
    "B5AuSgUhsyppXO27ALYLgEiatbeGw2C4bAXIrwCIwgpCn8hj7DJjFwDS7gIIlSHY6C0QaN+LjURWAGIWAOmxG0pigXBM7MLkI9pa3RYAsp2F1RSvtGdXoj8AIKVPeyVAfMBpARAjwhzc8jti5pjB4P9ZzQACSiLrEfJNAFGQnJBAAJE4a8/JygBZdPrnM0bblYS5XwYQ"
    "vdw+wHwltigAmS+Lq07Oq9lLqTmmAuSHA0QFgGBEWeCtBuO2ZGFYijXXpvFOQXQTr4vLAMG2vG0CEFCJjQxrs7pgTlwACNW2rQBZCdAInMtO5wAR0oeCGrkCkDFUNPvmwYOhsNjluwEC5lnfrw4vQDfUnxWAbH/DdLWH2a8DSMFrXQLIMvWxs4sKwpL6qAD52QAZ60B6"
    "IWNEF6qZ9byWENpmFQsJRaGQkM4KCfnU3wRiLbLPAXLOAcLJGVpXjQnCqnPaLEabrSzk5F4qJLTFfXQFCMQFtGjbrmvzSwg358xX0JcBAn0g/fNzgDh2QMuA7x9p691pfKUv5rKV5T6A+OYBax1j3hUgy/7X9JUBwjl/jZHZ+9bB1cqgkW8BiKQzJQr60ql9iwFvimml"
    "dpGHBWZ7sZWJyYMSvnHhopVJZ0MdCAKEpomUMWI/AgTr1unU5xE6kQhcGyTcmkJZ+eVWJnEi8c8HiLorQHylipx/pQCQWDy3BhCWBlGyLYmDvWz2ZQU8FiCbX1Mn9PDnaoD4QlC2vrTf7sL62OnConMXFrPH13Bh+RjNo1TyVep710J8vgF5yHq/ChCY92t9hR74MqL0"
    "y/JAbKZYSsOitkkVrm+m+DEDCHQTCZtnTqxMElqmNK4JIFCSYoyM1yX8qxsXRwtR7ksAoQVS/EiAxPEdW5fEVYWEPre1kWIFIA4Q5daabQTIvPs/2LRaXqs7nwwQd092CxtpF0DAv7rx4X5jFtY4JxRyJtaysKaZrRxrY7tyFtYxy8KCfjzPBQgOHKVdC5Mp6Dwotji0"
    "XAMFS994HY5txdy1INBGrDD+FlLVIFVjcyEKjvKhvM/t5X6LCwt2ixjCUD53Ngqmcvcfi51lVm/AQ0vy7GFohBIvgRQgMHUqdnmHMpMxr6OXchxMNS5NUSysCymmFKP7cW3gXb8SIE4nWUGzCCnnPxMgWCV/aWd99Uh1JkXusuRuex6f6ooAiQ1T2BIg4oaJWM8FCM4x"
    "WKj3PQCBVk9bDWd+WyEhg6HmNhYS/sfCiHNWarHrJNxLTiHgJPRlISEcNfobHIh0ob3vYwECCUIwpP3EGIOR7jYdh7Q0sm3x2lf039KuTdfGHUZhMPzMlON+IT6DHxdyLC1EHeFbYvEwa6h6JYDAveobVPsq8OxbkHM1yjtrJz0zzjIHI2Z8GIkUb+EUIOBYmcpMxFQV"
    "0ko5ZejGpcHDUNqKz6i2SdU8LHoejLkAEEjsct/9pHX8huAHAsT3vNqUTkh2C5h630sL0/BgP9HEdzQ0TJJOZHoXFixXFSc49628YcLJEwHiu48VbKQMIGvBE2wpyNi7AATLx8dKdGdVrLQpgUKmsZXJWr16OFdsZVKuVn8gQHwSZ/KuDFTymp3hdgralJ5D3f+5DhCr"
    "C2Phs8vS2RVt5gjEhRznb+bW+xlMNPxfv1z13QCB/b+/3XswjrAIhJo8FI5VIQvd2hnhR6LjNMOYGQ9BRRNb/MHo8fgRU4AEOyIUAxrTheO7doSJU0pjjaDT+E0ACBYWZkmUVNrZBw4TCYs5l+FXbU2y9HYEiCDTi/jo1yJrg+6/GyAQsbDbcl079+m0Ii18SPbn0E/Y"
    "LGR6l6GxIjlAXBk9LwCEPhAgoXelLDEgAQhXKZNVlvyrqNBvBBDfzf2qXlh6sxfWBJAn98KChuLj0tg0nkSszRhi5S/PAeS0DZDC50q1KTQ8jQthyUJs7kdW42EMbKbYGeCuRsj1AOms222GjFgbe2DAtI1uBpDCYEJFzgb7jvvusUndLjyMyiOZHJsDBKLlwVCBTRyc"
    "BhsITqVdiQXio2v+WoIIeP7FEiln5dHQC2vSYm23/D7AsyNkLG3wliB2LZl0n09EStKZ8cFVr8b3AATnSTUbwm6bcMhYk5bcpX25WNPIhTTT22CrluSZm9TmYwCSZ2H5Kc4d0nHYXAPuVUeBhgRZ/WG3TpDaTNG963YzxU/y8R3NFL3zEb1tca/1D/rjYM506Q2Ee0qa"
    "Iy+6b/Vq4bI6W7DTpnfBd7JZA52wEDtfSOZG4UeINvkT/ds4DWvRmVW2ip4GEKCaV7XIgLFSz3SzabVUFL4i6IoY0mlFQhzwaHiq2DRKQkTKINWZOGeEOxBFDE0oUH2b5PC7dVq0ZcBeIbMf0Mw8kIradBvcFhKHONhPAikXsalgRlayefYOM2qyB18NIHEi4arcOCE3"
    "x1K6Py8RK3ub/AB2k9Z8BEBUT7vAAP9f3/MYEDdcdKOJFIpWpPfH+9SBjADB/fOeeSBOTx93zgN5MkDQ9QhIoNEhS46YkO2+IlI0JLy9dDVAwDkFnzIXNVtII9EZHTy/biEAh9RMUUeDawvOHe/egQaUdw2EXA2QzDafHDTLqQ+KlL4hPr0+M+vHh1XeqonnZxxv0OR4"
    "ni6NZ0eHoX1kkVG8WO3qr7W19NxPEV5VfPCVAPI75REAmbUy0ZfMs/V27gNrrOnTVrxrfrpfOpEQnbrHT3vaBAiqOkI35xaeIOkSx9Y8d6CUL6CHTLBJm7iP5H7oUwkgbvsvVwthtgHSAUAgLT660nOXerA/8oQkdXQ7Fn3K/Fww/3G2XrcozRZpK08FSGLW548Wjts4"
    "w76HF7VbfMdpFs+XOLb+qbbHePDFEgqv2neqCpC7AmS6IS8C5LC7maKcmU1sL8QW80CgY36Sa3g0kr0RQBw+wYCzF2aiO8P9bIzYnonuhyUYK5860lZhxYkWeadVcGpSWup9B4EubHVql6DYARC5quM5Ouexe9h8IW360RF47hvNkrMUFr9CgzL+fQCpch+pAHkQQNRl"
    "gJibuvFeYQUtB0rhCEmV6Jd3aKaYpE/5bNLDpmDOr75wkDtXo2FiyOGZACGfYBcVNa+itGhluM97Ykwu32IHQNZTOWDir/tR/hZenk1OUGdn7c1jvb5/dLHyvwKkAqQC5M4AsXcESDoNMvTRGt4HIM+TxwAE9vOHdOxa7v8uKHOo5bBCsoKy/gpAIBFhLRCezlRwOxQY"
    "vnVcVPhD/+KTvZsTqwKkAuT3AcQOF9To2BqTPxEgzgZJ8nzJWwyU+i0AwZ5bcv/G3QEHiy7cT2g/rwQI3QJIcGBd6mjnLkDJTnLFt7aSN1YBUgHyvgCZ9AYndhguqNEEIOxZAIFFTk2dVwr+3xkg3oOlLzm6vgMgRPgg/95vgZw1zK4kXUlZfwkg6MC6XOwKQ1eKnMEY"
    "ysnSCpAKkCorALnQC4Xlx/55GkBYUhvFSVvyYb0xQJzCgxR/azV7NYD4nFxD9qpJglXoVPEjGCLzUpAvAAQsm8OOSg4F0+dtsVgRPgvTlFeAVIBUKVkVF4PfbLoBIST5PIAMTNI0VFNICn5XgDCci44ds41DyIm9FEAIAISZ3d/dERKp4KeG7b5etJW9HSCKum9XX/6E"
    "HiDl48CI0feKoleAVID8DoAkI78uAGRIK8bprPHWQwHi3jozfgrFIG8KkNhUKrYKgvq81wEIB4Do/bt2sBN8/q4KpsidAMJJhwC5/OnMyQGv/BbQKkbvt6YqQCpAviaHHwGQJLNKmf+x1U6L6U6tFh0+FiDQtnha5llUgEz8IPnI+Fv9WI8BSKsPujCOZNXjpUNVOCZN"
    "2UUnjIsA+WcNINAMxpLLC3Zqfc1igoSAxhxVBUgFyMNlaDDn/icAZJpIotqtfowsGXEFjTafCpAhe/N2mYj1lgDBsmqVb/mN1a8DEGgjvD/1FdT82PORSghmXwOQM7YyabN+FuFY3894D0DEFkAsa+7VUrECpAJkkx9SQn/eWxtkPXWlk2Yep1mVATL1YgZ191SAuHdP"
    "Unn7pa/tLQFSqJVQxLwOQLALpG/iuO9woWMXRU6WIes9zRR1HMGNY7jjNYkGzWkPQODTrUyVQBPnXtPgK0DeHiAMjAytSx0CoUl76+eTv7wRkrQixVbf+5xd7dyL9GiAZDp8mpryzgDBaC8vqrlXAog4Z+3GeT7gZgYIPX5rkBJ4MupKgBxYk06TsilA8lZWs3XwYKlc"
    "AoitAKkAuZPilX6bI5ctxrH7Buekb23z6jYIS0sJt3Rvs6nBHw+QlF9LU+kNAaLLjZ/IbU6sxwDEfajsWH6czbchWQhdn8YYg4Jui7mF9TWAnLLKRJz9mojvV+4Bcq4AqQAp+2u+1nl9ZmP4ocqdWdoZjfSpTeoBuuwRibxJq6k1m4ll5eBLH9LDAfJHb0ZB3hEg5bpq"
    "vtZv9xUAos65eZQmRilzSviATQ1z99dOF9YouQtrBhBzyjoc+6KTaoFUgGzhI05gunF2RqrzcMal4r5Btm3mz4bcWNLeNCbwyWH0tFGIWbGZ8oZUVLKnA6TRadeVeU/F9wMIWwsu+Oq914mB5FZGPrP9RLMANiT6jW8EXU1Eysgb5oHwNReW+sztNB+7rwCpAFnVATDg"
    "TxgD041uHd+X+VOmbTs5z0rbGtHGgWGvHwUZkjwsuNNKK2Z5S9ze6OHPswGSd1JpZ47DGUDWowC/ByCrayLmVbKwFgD5gPoOlJOfNj4BBIDBZHIpHsWshmQPQMopX8sgOgzOCwvB4ksZAAJZWO1qFpbWNQvrPQECLifRdtTvS4p+p6vOll1Haj4o70cBZHS4+dBivwj9"
    "Dz4nIJl42S4LMZ4AkD+pETSfj54BBGZYbcq1TsyXBMi6KjXshdJ4ZdbTitAx+IAhC5rCBrNwwaqHMcgE++LORz1dAMhKznAhjTddCIsAuZDGC4CqdSDvCBDcQU9TGDkOZrtdtTez8OXcU/WzAJLjkPuvZsj5kc5HgGxf9h0AaWyXhvHXAKI6Ky/ItT/KzwIINzd0NHlg"
    "IWFazcH9KFknxzYDSCjjMPToBVqzwCSR5EPeDpBSIeG0EHzOf/qtQkLi3WAVIG8IkAb5kV3HOOf4RpWrbcfnd/+dAcLYyhxYp8/vXV6SRkGQIM5As+lc8tT8WBnJ8RSA6CwRq1kBCHyAOJO9LEY0PwIg0IAcyhm0LgNk9VyFRF52wnOJtSTfB7UykVDNUXgG9nBZdxHs"
    "t+7WOCVG4RciU8v3CwABWuljYR2OaB4uPFoZbAUS2Aurrb2w3hAggyxVXLU3WwfS0BlA8sLsrwPE7fplsZId7IG7Fyg2WUMMd1cRZ+BTULWd+y/Jh4DidM8/3wEQ1tikI1aeiJUAZD4xeSGqVMr+kgABj4nTtJ+laefXAgSr1glZs3Qe1guLlYPSHPobJnAB9eyHJo7C"
    "MNUstTpvBchmM0XVTdaJb6bYkbVPXrvxviNAwEdDSmbtrSlS8wlroM3uCBAGkBAYphmWiWRuj2blfdukQJeS2SeK9j1Rs1o1qMEo4OsZAMkTxmj2LhlALiq23v4IgDQSryNocn4VQM4FgGj/A/Gjsc9v516GC7i3yHikBH6cmPt/J8z/lx3SsPiXACKd+XAuvlZBA8UQ"
    "qIEsMTbfH04rZLrOA3lDgGhxLjf4L6rCi9p9eDBAQsya+Er2iSED7M+su3dJT42Q9yWIaOdfUdiuL257UcTuUwDyR68WM14JEPlTAILVCdTo62Ig7LAKEPJMgEA9+RhfWLDlNLXIhCg3OLCSPiQG/KhQbr8XIOdLA6XKKVQqDdZj9Ulq9SSHYZD/WAHydgBhKyOZsZzh"
    "elWr2bAEyKoL61pFj9nGBu0B1bs9FeQThfLHpoHRD3DrghNG6K9Xs2TQavc0qob3Lb7rcwCSNeXNAk+/EiBadhsAEWvxXGXKZYehS6Flzxxpy9jBmmNx+DmbLhpgA4bQM/mPU+1sKpf80kRCd92yYun+0WDbRhX/ZVkxGRiyNsonqAD5JQCB8HJTKPBYvdxLTfkYnqSk"
    "mr0Wl6aVjRTzGIgoAwTrKvQVysoXuPugA+z23b+gcAU3ZBD9Dd1XIQTZGnHHKsXGrbm/gBBOSLf2nk8CyGoU5NcBBC62UGuO2/NFacvqGyicI1s4l88zgp1/oU7mMQCBHT3EMUghAsKSZpBHrOGgJG9QRdAeGL2rXwGI8qYFXbwaOr8kTiv3XZ/Sqyx//b2qQCpAXhEg"
    "bhcNuUOLDKFm9ZqDzI8ZhHyxul5qyQG50tje6S0p82QMtZbGG2q7r4r2TyFrZ2e0BGMRvfs/4v6k07M+FfmOfixM1lWb13y3Xj3zHID8GXRaNp/w/7cBhMGg839GgNiTXsQ1ZGljD1/LX1sYmm7NCBAs43sOQD7Ipz0UZvy5vdEJo9o8hh5YydMFyDiNlPkKQDAh7HBa"
    "fEiFBlkCFgips0WbY2T44W7jpCpAXhAgsHnvuq6ddzdMB+kt4m654o8driDPc2GaaCxPEwRCJ8mM7o9COD4DCDVyv//K5gmzXUcwbcjXVnFO6Wx80B17NQ7Mh154USuhzSPk6vu5F/cxyWnZ6OQCQNBlBy8l9FKCLeiu8X0ygLhbnl/KvwpC6JXBr6cDxG12jTM3wynU"
    "kZpPM0/A1eXRq2oecmeYGxsvHUjq+vw0Cxw9CCBR89L8hgHXbjLpXrXOxBKkFJ/QuZ+r0ME+DaLrMlRxId6JNV8IGHdpdQc/et/aMU0eUd6zZe9mgFSAvBxAmAQXjK9gzdQc0+e1n513uXUQuuiibprbJsybN5ZwKMROIwboVi5XooOdcEWkBTrj5vdPbnJz2ubPt/dM"
    "xoJOphhkmTME03qNkFtRF/f1xFE+YBgN16nnlo4vvQQQR7nwPn2ah4WBo53S02UrltcCCOhWlWYw4O5hRhDGVrrxzg+zNDsXnIwsCPIggEDyEmzpLT1OtjNFGE6xauVr9FTBbnBkOf1zTAAi3IlKndhjLyyqVvrFQyeSsJAEH34h6UeHCvjDSYu/I1SUx0xzvwhIBcjL"
    "AWTyW/JZgUfaX2d518pDWq5uxmkEC00GvVAQIEgojRvhoF0XfpckBkKI6gXb7UXKU+Z5nwNE0Tw9hFN55zp3MEMM7cOUbcVD4fCm8TGljo1lelc6iMDu8y+1FwNG7quP9YCpQw3Sm9udYq7xKn4TQAralCxqAAvzQLiaB8mZLgxSUk8DCKj9E3wdY4QD7hf3yMmO7Q1A"
    "Z5d7C2Me8Ck8A2fCQPtxFPg7fAfEA+QzefYIz06mBZa262whlvmxwLmzC6JE2JYxzpqH32Ylu7cC5HcAJNkeoInAEvXUbQGEZeXq46ZjCQUHECuZ7GGf3GDGSB/6atm5Gp8AQqB4S+zdkM8B8kGyIj6IhfDHAuQPWiEQsjdRp0MMHwtPhosvvLFRyJ/0pZde627tUkcS"
    "rJ3ZLa/eymQnQPKJ6MHrOc/m+m6AcAyFQ5/1qQkWKugujXwxWy7SgPTbmLAMAGFa25kEigJAsN5+/mxizsSFfH6e4bqGEv9lZB0nyx/igj/dev3L7tXFpALkFQGSNn7ObZC9ABm06JKLrZu5yT1ABgmOEtOEBoNOSm0ZpzRe0tKvAAR2/9OaaEceDpD/5QtOWNZ68PUn"
    "8z5YvgEgZAdAoLBt3qHHNnMyNN8LkJDohIL15ezAWLbtd5exPolyl0IYEhLD6Dgqiy0kKHbk1KHwbBqrzBfCZgvJfNKwzumwk72n/VEB8nIAyTJrs6g2W+9fg3dtWq7Okztxppt9c5FBt87uaBs2+dwLbRkTC+RrAOHqmHRR77r5fvP+AElTltnd5m9VgFwJENCnaeAC"
    "43IlgGDMyrev9U1Plscw6QN7aTwFyiKeBhCnkD/BiI3Xk7bC0GNys7ldmKDlTBd8rvVwUe5vu5Dp2VbY5fOztBOoThwpNF9Iepyx8hQPA1uE3tP+qAB5bYBg+QVLsrBWg+iyOUzKm2zqZq9KIWPIAQS7M3IfJV8WCzb2TgDJwubEzD/IwwBS5ZsB4l59nPyXHHO5TakN"
    "IiQMffqtDPSvLbXhhdSkxBeKJ1vC6IEAiZr7lEBv1l6to6up9pR2QXm7s2CDtnMirXvpcf3Ztp2locBC/Ewr7NxIj4qX9Tsex3xCtbg3PipAXtqF5cNyYy1HYy6n8TIYIpjn9y0skEa7Uzaig7yqJB2QL9OORoDw3nwRIGSqkXAw4hUgbwEQTOOFTpYkqrMzbVf66EKJ"
    "hy8yFeVh6KCxz+eY0Ad69oxZRU8ECO6z6BkjIC2251z31C5fOcJvbNCWicowu3x2uZBPXMjnvE/o4rjOx2woXT2sAuT3BNFntX1ujzXEu3+9kDA6n5pZd5JlEJ1BbFli0yh37rSNBF8OlBoBQq8DiF40LiVdDJy7K5pUgLwJQLz73Qrs3cePf23sTluyQU5j+9qVI3RW"
    "SGhk6VyPBUh0nRXaqz1f9i7kkQuuAHkxgLC8tg9TLoJWZ6IrE2SarL1oTrIcuoe5pgLb93BiclotCglDShgYJ18FCO2j/e7kowLkTQASemHR9V5YV4q2vhyKnJ/ZC6tKBcjPAMgw7wY4zT0aZPneGJspDWw+L6RQSNhY01GD/dTAC5D98KrLB0qFgkQonxBfBAgnvW8k"
    "pI5LO7oC5JcDxN4TIGMrk0MFSAVIBchC+4q8ZduUo8uaohMLbJRoveQVWWrZJAsAQiH9SrpzqZ7M736Zm0NYlk2NlfZrMZAp9EEKYcYKkDe0QNhuOaUuLW0/SbDMK0AqQCpASvXMMwz0IcF2gDYnahEli1YGk2KeqmGWA4cygMycSfN27kPjx4dKLW13FUBK9WMIEF4B"
    "8n4AwSZNSh3TiYTgSt0rbgc1m0iIOYMVIHNVQAi5f5y8AuRnAQQ6SVE1vylZMUYec6dCAhZdOr9YQbl3dBUgixX5AgrWyK8DhII7TJFCMjLvKkB+N0BwQkbaRhGs5f1i0hf6aRvG6gqQ+R2GyWGqAuS9AfJn0a3GN/30VYDYajHZc8S+fe6OnAdPaHHoN06Z9S1E1Szp"
    "EF8Sy7Zn1YfXAaQtjUKjvgNgESBN1e6/GCDYSBc6biQcaMlsasa6kM/khVg2h205KkBmN5j7VgQl1QJ5d4AUAumRBdh5pIcOp9CZ3E+2wBD7vAAEI+LFfT0UgjR/igCBPQyWvc7bLMGw8asAcl5eUNAKt+toadoTTMytZeK/FyBhLEg6/2mjNWhBaEqLsZNIBchMlbuv"
    "RTyXHxUgrwiQWTsS71vCueI4KHZsGk76zvcmhw58C370K0P3BuYAcigDBPKyKD1jK9qsxd9VAGHzZOJ4rbWmbfuSj5aY6sP61QBZ6vnbAXLhxG8KEGx85Qw99VEB8vYA8f6oRahDQuWU+18RuoaDVyA+ZmY1ItgEhRXNG4cgiTMJlwD58EOflJo1NrkKIIi/4iVu3MqL"
    "OyTVG10J8k4AOegKkPtq8g4GGT7ZgVUB8qIAQSWcbdW5gvmCUcb0lCjtopXz6tBvMGEEdFGkBYBkcZfmFoAMTMuV0YnEWFGeZYOj3ZqKkLcCiLripPvrR94UIO6GZ+yOo2orQH42QAbMyZ39VFuiZv4rs6aPYeRe314GCCdtk0ZlKNkFEN/et3w5qfNqjgj3s56qin8f"
    "gHzuV3aqZRUgF76io9B2dQ5uBci7AQSL+G5OyOP9ekwByUT1ZYDEkeCHg1+MA8gOG2FRxZIpl3bdxMYxOVXFvw1AmNm9FFRSFSAXAGJg/vlHBUgFyEgQcyNBsADksAWQTkPVyCZAHIXENPpOQC8sI2VzyUho7PrAZU7PG2VOELWpKv59AGJ3T+bG6d4VINvfuzrevVV7"
    "BciPBshaMtMeftj1iarurB1pdTkLa+5VakOsBccR4BycS9lSjejUOiLOG24LwF7N5X0bgBzKw8NL4jbXhwqQy1/9t1g+FSAvCxCsO1c38UNvGTbSdCZWom+4sHxKbxRCcLJAoTdKLnrj/lW0vQCQGkZ/H4BsTNicXzZWV4C8quusAuR1AYIBabX2w+HsGr6In3PS2c1Y"
    "BdSMSLZSB8LT+t95nB4qT0xz0QIh64Roycd9LRCogfk9+n14J4Ac2M7tUXmE4W8BSBgfpRSvAKlyX4Bs2CCcnM++p1yXB6b5cgLIQuc2uhlWWpmM0o9/jIIMaS8CxFkZCwlnN6bwXBRyfQwE6ip/S/IvFok++tO8FkD2VS0oaq9p9/4ogOTNVfjqAMCibOHjCIMFKT1u"
    "hAf52vvxIndwG8hXl77mcbhq4RUgPwIgazaIb3vjLjsjLHTXTe0PsV2SB4WEaxYIRDnWxZxxjPol5IkWaJN4vxA9HNsGi/NGx7zr03gbOYb1h3H3PoC8rhGwZRgKK1d+vNVPlDwx/DCAHHblGV47RORBAMFB5SEi+HmGybolja+OZ4OHtaMYuG1W3QhQ3IXNg2Ccrzvl"
    "ysC4tnwKdaSFJyB4mc5mdzszvxC/Lu+PLryoi/HOUbrLX08FyEsDpNR/9yOMCIH+JdjeJGm9iPbHcElRCWFZMQZCjG+MUhaojr8IEDj9oh13C+UfAJCtFt5CXl2M3gjxv6Hexe/efbtJ5ltBQsnKkDu4/KcIZTZ4AD7kDxzC0+HV0IHYHwuPTUfHF0/1OqyJTy89amx6"
    "y+yvJvw9jCccwvcmQlOw9DxDsE3Cmxxii+TpCRa/APajAOLOfbFwgRNxjQPrYQCh2LwxvRvs3+OyExDCbn7fmG6t+Mlgg7Bx7knZJFPEllPW1NHognogZ6Hd6tSIn9maNDbE58tAE1us/HKeQwXIiwME+u8uG474yYHMezz0mKwF9selQDSDXlqm3M7dJ+qmGnGWodtf"
    "BgicfiH9DgvkfLUF4tDpW4RpKfzYEuCi/9MpYtzTj19HI/EprK+Ebi4IWu2PBBK7hyx2iYH6fugQI1GZ4wh5eB22IXMrRADL+P00/v2C5eAOsNOHwPcPy/IrhL9YfAWofL9M+A+2q/UrZ/geY3tJfK3BpcfThpPJ8cNCnjUuu/lRADnoS6ULkDvOXgAgnEqdTsGKbeXz"
    "W5OHgpU5QMofEko3Tsy3htQ4NKucuA+zF50qL57gpJfFg5iEOTXFUvQ0W/lJC0OPs9LjkKlQAfLLAII2yPznUa3XFIeDV+zhsltp4L7AABQSNiWAgIoHbaSn/W72ytYBRF88fT93F/soCFe9aTc6eJPZOMTL3ian23DGROPuB0JIZ2zTCAoOMyhrx54tYz+WAW6anlIc"
    "kQKOwdZBgInz2bivTLqdv3vobIxvU+n+gu4w7f8YhkY4O55QpxZRb7uzwiCL+DWD9deTHpaBsBHmbHRii3W0791a3GtM52dYuBV26OATjT/hgP+x7kTu8zviuPUMcX3xG+1gTdTIw4hodwpcYg8dNaHFGCycdkb/LIA4q3bbBrnW/ngcQEC/QvgQxVt9Jz2r3XMAOcVu"
    "wZOsOeqIBWy4k9h//9tiU1SHA1KKAWmn9AuqHCyQMkD0KQUIWiB+5dFehVVlpot/l9nSZQXILwDIn2bRWyobPTuM5SLgIxouu30cBhxAnFFcGCgVWvEaAIm7pA/zGkHS6Yu4oz5hK4x7GButcNIbcxwfn9ED9N+VAMGmjdAzUmN/Ytq1RjYGtbNjiX8YbC0PG2n8aCKw"
    "WKAUBiDgFEgPj9i2te6GoqDmET9eJQNsGpx7dISpvk5xO60OhOxNtAYs6nL3reF5LRhfUfEDToAabQt0QrOsBYIY0vuFoAGBnAMD40w8QKBPGTYLiJByP1iHLxYIEPhYXUc7hw0D8SX3sX0vNPgCxKtaIF4xLW0JnFW4XluaDgKZnevEng4QDZGK6HB1Kv/EZgQBbQqu"
    "3tw3u/IeTrH7OSlweVGIpEMkpAAQo8GEKD2zHyAsXbl7Ggd6pVoFXVg4aSV1PdcYyG8AyHLUB9bmDuM+PDwLs22bHXGDFirRIVsK1PYiFwMEFDo18/aGuwACDhcMwJ1DIO58hmBcr9A+smmILhMwBa50YUF4yBHPgqVxpqBPWynbHvSwdDutAWZvAQSGABtIPPBRI2Ao"
    "dgsTPXBLC1DC7lsBv1XXt+Adgm8Jvs8Gn2t9J8u2swz6HsMt6kMVosUnjINXMH/GyS3a+Agp+KvcK/EwMICMWyGc0DfFRIAYfHvS4iEOeWgkRkg1PjYC217f4cbga0HbOFCCIoMvwAElHvKCAEFF+W9bpIFdizFDXp4tvgI+vn7qQCkEiG2ParxBAOgnsBmSt0OA6EWq"
    "4ZrmBY8rxdk+cN9BmdXn38IkT/8dFOyYKwCiaXAF+Fvbon8qDZ8AQDTGRi4vvQLkZwHkzx9tW7UIv8UJUzpk+u4sxAOPSN+CsnMAWb9/uHcJ3Q6QTCi9BBA/geS6CIhoj6D/HUgcfwS4pK0z1loZAsvQvgtuJjYCBL1zDC0XAKSjKADEOOOkQ4BAdnOLpokG/7bTyQgQ"
    "6ywL3Od3nWjA5kAPWQQI6HanwL35Q6LidybPGVjWNBjCAD+W9G4qQ87SBiKNADFhSXAk/DYUFdQIEDiPD5GD9QRv0fjTOcMILCrTwtKaVw2iF0ba5tm8qpyitEIcPNdzR9p6gCRn9rNAGYxonzS406aNQ8qu93fbPUyFzD/xMloC+h8st/mU0k2ANHOA2OR9OARftL9B"
    "eAaQ9vpmvhUgPwEgzaK+m4PXnaV5vo4pds9YP9DwbgcCBX+EbN0/7hbJR5Xvc2HB7nnpogoxELMZA7muDgTUZut9QeaMriLaOQsEHFB+X27RteSx5Ky4DhMxwagAOji7yKlhcBU5td0GgIBPATMMwOHV+aNhm2+oM0ucHQFaWsBpHDYSgPghX2cB5g++YzAbe/CBMccD"
    "g2+GYBCwQnBr5BaIgQw3gEoLADH4aUJDTHBhubWECDmeTPghYsK7xYTEcC7g50WD6OioUhg0PqwQYV4Oq44darnF4Se44P3V8myA5GcOejgNcYBVcdKf+96fWnaazX/iqpBee3bvIXUp5/kKgPxDZmwG8ye1nipAfjNA/gxy6cRCr/vEj9g893LoudG6GRp092/fP5xk"
    "OcH7ALLaCwuA1PGtd+vsNXm8zNkG4EA2Tn+eKfZdaSVM2qZgLkCEAbS5CXEBJlv0NbtvTaOfy3SdRRXslDuNABkGtEAcEcDHZM7e49She8mOADFjcD4AZEAdCBBw7+hYMXhLoW/R0eX9XBp+xQ4B0rfxhClAMMRv3PvCjdxCgMMGgGAoHvbvYIwCjGygFyUYwJ9i6C8a"
    "AwnZPKslHd6PlW/v16LnUZG6q+V7AeLbf8JJxlvzKoAYuyvLCVQ9JBOypWVzowUyhjwgE7MC5C0AsmwxBerYQFw3bN3wQtsXUYFkDFCU5BJAsC7kaoDQDYDQTYDQawDiNuCgVsGHJbwF0nmAoAUC/GgjVQYfO3DqGkMRFowveKWAhCoIFUK4w+mStokAwSj22al0zI6C"
    "VpIRIBh3gWw0NA8iQDB3ucVItoMRfAymHUD8T+LY1aJZwmSwQFo4AZywBYDYAJAgLb55O4bRwQI5QzQFzQsNDsIRIGjwWEdQiGuZl03jvQSQaIXwMZUDHINrAPE7FEK/HSCowVnyBT0AIO6wE1jFWjNB7geQDyR08v4IkHMFyG8FCASMFwoZXSF8vND2RRCcPSwBID41"
    "dPv+4Z28I0DIfQECxgakdMG2vQVN2nqAdBaiyT4bDJO7bKgCbEJ4xkBrehBQ1RDQxgywAJAGACJDhoFjjPQuLOr1vdvtd5gy5uw99CB6gPjMWvDD+Gwy5n1mJAIEvE4pQLwLC07YjgAJoIJ3Mu48PYESCO+08vTwEY5GJP4zsD4g5xqD6EJe7oTybQDptgHic1+DzlNG"
    "ruZYwbnalwEIPgyeIH4bQE7SXAQIDKolaDDYo7ofQHiIrVeAvAdAZuG2eA30o9OUmJ2tbJmEKjkmO69+L94217uwyCpAWrrZPv6aZopYYAn7fcheBYBgFhYE0Z3VAFlYInYRpiFYhNPkIZrg9TR6kXxUm+YAMeDA6vGQ3ljMl2p9spWzcfBgfIFk0TsFZXztGcol/YLA"
    "cmDeCwYj66G4o41ONQyUhLQuSM6CAAbWLhp6Rk8UfJywPoO5XgAQ47OqQ6sYTLjSWnquQRYWWFHiJbOw3H4F88XiFsHJamVHuMRXmyfiuUTw2UIuUelcTwWIx8CorK8CSOdn0F56X3dCQTjaK91Mv38FIO7HsOkM3AqQXw6QRhQutSTNjoidAGksVh9IzOu70AKUZ4Ul"
    "e+tAyDz7D5IHsQ5ksxvvdQOlmI9bg+KFrfu5FW5r6lR9h8FkqLwg1BdIYI8t8Pdi5R1s86FiRAgfcXC4QXPD+CwsBIiG8Hs42oBjDMozaOfLOqhP6/XxBibOmE4FAGtpSKPyeViDRpsCAuWYINZiRxcIqjhQarRmfM4wws+n+Ep0K/r14UvgN20CvmTIBpDhBfiu8Bkt"
    "BoHgL/lyLiynokzXRUsXMlXdOtcIErQcTL5cqTrERgYhaZ34app5nP3JAIFZVyyqZw8Qs0uZusudXewn6TZVsD1QaLwtwuhfAghyevqyKkB+O0DshXDF9QC5/Jvno3H3pfHCSNux/nxMPcdWJpiFtdaMV5GzuA4gAgs2gnKGj2Rar8oBFaGyW3ofE86Yx8oobG4Hjh93"
    "NP5bYsCDCPA4AUAMpOrSWDiIYROIu2PxBVT0tdjxBIyGBgsJ0SeGZgqoM3eXt9S3IWlEbGEnJFoc+E/I/DINJj8Yn82Lhgq6sSCK2XcdJlb59/AAMaG0E+hw8PUlwWihHi3W4+b1srAgQHDE0tEYBnNibgaIIdO5PvCyAvPyOwHidDib1LPTpqe9abwcFfisnG9JGeHT"
    "vDhMoZ7bK18DiIIM97Fy0deBXD9+qALk3QDCsOfTPoDkjrG9AGlJQbAXFlY7rAndG8WJoRwLgYFG29gHy6ERlH7ozCiweAI7XHkXlvV9rrBxiIZnLBb5YQjhjPkIEFqAkzjtbjC04E/tdBSk4grfKcsHHUKXLRabZ/k+WtiXC9bSxK6VwneogkX6o7BvCqTPWXfesHZ4"
    "RzhhgzWD3tjAd8caFo8+OI8ONfUidNMSAS3wQUz7kgDRYtZzifu4wQ0AARjltW0wA2BBo+cCxDcviT0hA0DI7lNiOd/GKNqjgX0RIAbr1mcx9y8CJITRU4Acr55KUgHyfgBBd7rcdZuR1l4LEB16huQCACHo+6ddV+yl2LVXFRL6ZDJfCePbd2F+cuwviC29MAQdhmwM"
    "HgfYNdHBZoC29hoPGhiEhXyTe1y/9m2OY9wEj2baz4a3FpOl4ZnYPtHjy3dG9NxofIt5380RaOUD+PDmbIAzMf8zMN8XER5u8IThs8Bhf3yPxSGACOmj0yaNgKV5M8UXBEhBm5IbAVK6BdQ3A8S3KtThiVCJftwxfgNW8InlfBt6m3fuoohpy5aN73MXgHBIaWA6AqQL"
    "lehXTQOpAHlDgKBGHWSndg3zEVdbILbQjZd4gBjRrXfj7a6yQP6wUJo9NE2AhG9uLoMODilJ2DvdNzPRvsUJHDUkjd/xKdTx/m9/iO+23kz88RLi1MN4kkChho0vasaK8AGfhVeElrs+DO5fCf/xB4xrZ3/iZ4nrC+8lAyPHMVp+E9CMTzQjeV4OIOoKgKhwKRefd7vl"
    "lwSIngEEmt2M1zSl66vBVi1QZm6hJqtYjN+6zYyZUgtOuY/prgABewiSQZIbcs8XWQHybgBxCgkGLsluT8AM+lNf78KiZRcWueTCEvrrQ6CG0Im6MFRqnOOxzPXyGnvHWKllrSbbnKrr3vMwHjcUVg2/xp5phfnLWTKaZJpscrmQ9CUskONvAsjHHCBj49sgYqNhvfLd"
    "ULwjq1SFbthJd3zS8Pl6vwwQmQNk6jYc2ggfK0AqQNYqE1enrae5Xe4Sbm4ASD+TBCD9qtwOkCr7q/efH0TH7stxYA00IenkaSWNV4177WIMRHoPSzwXnOwoXhAgei9AuMI2z2AIlIc8iSlIgeGevBrsAQDRFSAVILsrE1ccr+nAS9VOXpG9ADErA6VCDOROLqwqPwAg"
    "vg3v55jG6/SlMauFIGMhoWZr5erGJGm8n+6fi0O/HSC+8mUUs52nq3zLYQyFzBKyOPm00Gp7/Oe8dP2LADnOAQJTQvKlV4BUgKyfcOzdHHd0IZ22Taxp8MJeWQciuoJ/Su3IwmqvnAdS5fUB4jTUP1hIyL3T38r1PuynCwA5eAUXTGflCwlPhxcDCNQVppG+CyMXcYI5"
    "jqo8lGY8nRJHATEnloXR722BLDZ/NQZSAbJVyo32ADiYACSE9D328xOpbcKThrw7K9FbsizyCJuttiN8VVR/TSV6lR8CEHY66dgL6wglNWtdSsbdtfrUKz4ud66TFLESXehTaUTVdwfRdWw5zS9kYaVWCHji2IwgAFKd3Ivwb5bN8HhwFtYeKFSAvCtAYjqoEX6EII0j"
    "uRuWVLxyKpvDSzZTrPIjALKvmWIGEE7t5mFnhVfL2bJv74UV1LOcALK7kHAe7bDeBlHZCmcAkSy9N79cB7IASK1ErwC5KnspZH9i28E2pr+m7UURIHdqpvgBFoiqAKkAKcs4MQl8/zsA0r5AM0VYhrOWZN7K5HqAQETHt2wniYLHXvFYye99vAI6sKfzq+5dSFjbuVeA"
    "XMuQgQ0wGZwQMWaEpsNHACDDdQBZnwfSm3YTIF0FyC8FiN+ibwNkitmqo9gCCB4H7QVfACDYymTWC+tT3aIOHB9OqaIHOhxg4OE0zBw63Gs58emLAOkWrUwqQCpAbhEcEZ5UDCZBEE6n4rS9Ewn7lYmDBKc23W0iYZUfAxBs+UTKI84ngNAp38isHxfP1Qr9AgCBZorT"
    "vfkVgGC3q1l/dWxAPIlmeWWNBwgtA+TzhmaKFSAVIDdplQVApovSN0O/Jo3XmnKaru973q4m8raipvH+UoD4/Cmp9QY/0qUps3UgNKyBJvbsFQBibDIr8CsAgbdIm5VAGzpMQkjEj5CfBiCuAMQAQP5eBMiynXsFyO8FCHY/XPuVOMcsvNu1Sg4Qt8cb36kX+jqAjB2a"
    "FoINp0KDwuLTunqwfilAdkiaXwTJ44cb5RsGSolbBkqVVqQdQOK3cPStFq3vwAYiEcFJGB3eTi8XhbND/k0vAAQeTrKCK0B+N0D+QBfvzT4gF/sgrcpBnGkGkEZOsYrEstkFkD+zdgiThCZTK8/C86xm8b4tQJhOdtLrDXtfDCDq9pG2RYCARo+paBJ7/ucCRZj6M0Y9"
    "oBUwk0YVfFPuoG2AQAj94HalvALkPQAyNL6r94qIG/kBvQa9U9lMTQCb5KK8GiBVKkBuEC3pbBP9+gDhGO1nYvIpfQkgRw+QwIYWU6RmccIjTEgfTRAMm5zmbED6yqmipAiQ0Et+mlJXAfLLARLblq/Jbdt33yq8G+tAPIbS8YfXurCqVIDcovbzvhk46O/VAQL3jM4a"
    "Dl8DEL6wZpzNMQbRIaVd28+52oeoRxJGxyEhs8k+flBVMj+9BBA8D0v74VWA/HaAoLPpsPL4rfoEWk93vqUJThu1oQW6jFfSLAurdxujqoYrQO4OEJiJkasl/eoAQXOAZe2p3LKbvXUgZDYExKftBkXvYyvLfqcckHGQxyTuPj/MD9kVJAt2HP4hM/8VdAFO0YOFhBUg"
    "vxwgd9cnje18pypoZQKlSr5scHCXElxb7tE2qUSHisO2ZkpVgNwdIEzPGgmionw1gNhkVBQmJftdPMm06dTKZLMjCCcdPZL4JGTB+HyuCBB5Kn7xi1QtsDamRr7Ym1GyLDASm61MSz/C7GVoAJy2RaGFViacV4BUgGxIo2NhBvbGVr0Jdecw2ZzAXXK2Yx3hH2yeZSpA"
    "KkDuDxA7T0Y93hpGfyBA9JShDj2FcZxg1nwkNFP8zCPfxXaK7nuy1h1Jok6HbljadlM/ydSKSJd9Okw2h2/Cru0ZzsOhlurTxzYonwFkCscbYbWznJjNjD6sS59n4B8vz0ivAHlrC0Tn7XGjBfInJoAYkaTWQsDE1GLxCpC7A6Qwd4oYq18PIDbkrFhQ0zjFoyO5Nj2F"
    "rvOJfJYI4p1UqLGPPr3q4NATprxhHKPsCiPgNxsDRtD+B8FjoOPJMZxHZx8LAcLimoS1vh7RzpyGfj350j/pRZ9WBchbAwT6KSYShq7iM1itled2QY1HUwFSAXJvgBSKGSAN6dUAks5bgro+Lc5ELbTpbKCU1rZEEI6uJh1nh2gs+og+JUXnoz/y8LcWYxSE/EWTIwzy"
    "kP48eawGAHLQ2dJxftXMZ4joma3c0gqQCpDt9GAvi6Gp/pmqcitAngEQcyyEi+0LAeSDgpLPZaGEwXKQhQqn0vBPTv4jpOcQ2DIMet7HtlYwek2vjXMCTS8nFxU//rXxNP48M6vIRzdmK7eLJUEMZrnyCpAKkL35XSBVw1aAfA9App6zUY7mtiDIgwBCKG0niRECvlCn"
    "6VFezuUoCMS7O4hH+K2+MUnEQYFXayX+4OhypsmyMW4elb8jQ7sYjTuu6Xw+h6UTsrT46Hzp5+x9KkAqQKpUgLwqQObyH/NaAOGzCZvl5CqnzUl5HmdZ+R7ppx8e+zdT+5C+sj7OCd5DzVBkfMsT8Vngzr6lK7V/4RUgFSBVKkBeByDLBmngz38lgHyUhmzuOGz7YMi/"
    "OtLjkZBcq/OtDNrFk3iaaFnwi2v6uGLpHxUgFSBVqrw2QPKes0nr2RcCSJUKkFcFyHfGHg5Z9CP+OT40i4vUQEkFyCMAckepAKkAeS+AwEzAgX2DjMlXK4+vPV/lZ0lItKsAqQCpAPltABmgtqJKlYeLhgqeCpAKkAqQ3wQQ6GZIqlR5tNBWNBUgFSAVIL8MINi9gFep"
    "8lgh2Ee5AqQCpALkV7mwoEFh3R9XeaD0Tmi1QCpAHgKQYwXI9wHkLAem16eEV6lyL/FtzW4CSFcBUqUC5AUB8t8/IYjeaKmrVHmkSOknzl8NEHZoxM+yQE5mP0B0BchdAHKoAPkWgPxPB5A/NVW2ynOSeW9J42WsEedrAfKdHxOAtxcgrKkAqQD5yQChoqaXVnmmXDe0"
    "VBlrxfkq5nDoy/SdIijZBxCz/9gq67+3gslYhlSAfANA/p8itEqVJ8qVN7qaN2nao1HUt6YM7F7vDZ+tSvEH39N0scoDAPKvmr9b5blSb7sqVX4PQKpUqVKlSpUKkCpVqlSpUgFSpUqVKlUqQKpUqVKlSgVIlSpVqlSpUgFSpUqVKlUqQKpUqVKlSgVIlSpVqlSpAKlS"
    "pUqVKhUgVapUqVKlSgVIlSpVqlSpAKlSpUqVKhUgVapUqVKlAqRKlSpVqlSAVKlSpUqVKhUgVapUqVKlAqRKlSpVqlSAVKlSpUqVnweQ//pXlSpVqlSpcrX81/8HMs7Nu19xnXIAAAAASUVORK5CYII="
)

# ============================================================
# GENERAR PDF (idéntico al original, sin dependencias de Tkinter)
# ============================================================

def _escapar_pdf(texto):
    """Escapa caracteres especiales de XML (&, <, >) antes de meter un
    texto en un Paragraph de ReportLab. Sin esto, cualquier campo que
    el usuario escriba con esos símbolos (p. ej. "Sala < 20m2" o "A&B")
    rompe el generador de PDF con un error de "unclosed tags"."""
    if texto is None:
        return ""
    texto = str(texto)
    return texto.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generar_pdf(centro_id, output_path):
    try:
        import base64
        import io
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.platypus import Image as RLImage
        from reportlab.lib.utils import ImageReader
        from PIL import Image as PILImage, ImageDraw
        
        centro = get_centro(centro_id)
        if not centro:
            raise ValueError("Centro no encontrado")
        
        _, nombre, zona, fecha, img_ext, tecnico, direccion = centro
        detectores = fetch_detectores(centro_id)
        
        # ---- Logotipo (para cabecera de todas las paginas): el que se
        # haya configurado en Datos de la empresa, o si no el de
        # UPRL/SERGAS que trae la app por defecto ----
        logo_bytes, _ = get_logo_informe()
        if not logo_bytes:
            logo_bytes = base64.b64decode(LOGO_PNG_B64)
        with PILImage.open(io.BytesIO(logo_bytes)) as _logo_im:
            logo_w_px, logo_h_px = _logo_im.size
        logo_aspect = logo_w_px / logo_h_px
        
        def _dibujar_cabecera(canvas_obj, doc_):
            canvas_obj.saveState()
            header_h = 1.3 * cm
            header_w = header_h * logo_aspect
            page_w, page_h = doc_.pagesize
            x = (page_w - header_w) / 2
            y = page_h - 0.5 * cm - header_h
            img_reader = ImageReader(io.BytesIO(logo_bytes))
            canvas_obj.drawImage(img_reader, x, y, width=header_w, height=header_h,
                                  mask='auto', preserveAspectRatio=True)
            canvas_obj.restoreState()
        
        doc = SimpleDocTemplate(output_path, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2.1*cm, bottomMargin=1.5*cm)
        styles = getSampleStyleSheet()
        centrado = ParagraphStyle('Centrado', parent=styles['Normal'], alignment=TA_CENTER)
        nombre_style = ParagraphStyle('NombreCentro', parent=styles['Normal'],
                                       fontName='Helvetica-Bold', fontSize=20, leading=24,
                                       alignment=TA_CENTER, spaceAfter=4)
        zona_style = ParagraphStyle('ZonaCentro', parent=styles['Normal'],
                                     fontName='Helvetica', fontSize=20, leading=24,
                                     alignment=TA_CENTER, textColor=colors.HexColor('#444444'),
                                     spaceAfter=10)
        story = []
        
        story.append(Paragraph("Informe de colocación de detectores de Rn", styles["Title"]))
        story.append(Spacer(1, 0.6*cm))
        story.append(Paragraph(_escapar_pdf(nombre) or '-', nombre_style))
        if zona:
            story.append(Paragraph(_escapar_pdf(zona), zona_style))
        story.append(Spacer(1, 0.3*cm))
        
        if img_ext and os.path.exists(img_ext):
            try:
                with PILImage.open(img_ext) as im_ext:
                    w, h = im_ext.size
                r = min(14*cm/w, 10.5*cm/h)
                img_portada = RLImage(img_ext, width=w*r, height=h*r)
                img_portada.hAlign = 'CENTER'
                story.append(img_portada)
                story.append(Spacer(1, 0.5*cm))
            except Exception:
                pass
        
        empresa_pdf = get_empresa()
        cif_pdf = get_cif()
        if empresa_pdf:
            story.append(Paragraph(f"<b>Empresa:</b> {_escapar_pdf(empresa_pdf)}", centrado))
        if cif_pdf:
            story.append(Paragraph(f"<b>CIF:</b> {_escapar_pdf(cif_pdf)}", centrado))
        if tecnico:
            story.append(Paragraph(f"<b>Técnico:</b> {_escapar_pdf(tecnico)}", centrado))
        story.append(Paragraph(f"<b>Fecha:</b> {_escapar_pdf(fecha) or '-'}", centrado))
        story.append(Paragraph(f"<b>Detectores:</b> {len(detectores)}", centrado))
        story.append(Spacer(1, 0.5*cm))

        categorias_centro_pdf = fetch_categorias_centro(centro_id)
        if categorias_centro_pdf:
            story.append(Paragraph("<b>Categorías profesionales</b>", centrado))
            story.append(Spacer(1, 0.2*cm))
            filas_cat = [["Categoría profesional", "Nº personas expuestas"]]
            for _, _, categoria_pdf, num_personas_pdf, _ in categorias_centro_pdf:
                filas_cat.append([_escapar_pdf(categoria_pdf), str(num_personas_pdf)])
            tabla_cat = Table(filas_cat, colWidths=[9*cm, 6*cm])
            tabla_cat.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F5A623")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ]))
            story.append(tabla_cat)
            story.append(Spacer(1, 0.5*cm))
        
        for idx, d in enumerate(detectores, 1):
            (did, _, planta, sala, fecha_det, codigo, _plano_antiguo, px, py, foto_sit, foto_det, _,
             codigo_sala, profesionales_sala, hora_colocacion, turno_trabajo, nivel, plano_centro_id, fecha_retirada_real, hora_retirada_real,
             resultado_bq_m3, incertidumbre) = d
            plano = None
            nombre_plano_actual = None
            if plano_centro_id:
                plano_info = get_plano_centro(plano_centro_id)
                if plano_info:
                    nombre_plano_actual = plano_info[2]
                    plano = plano_info[3]
            story.append(PageBreak())
            
            titulo_partes = [codigo or "-", nombre or "-"]
            if zona:
                titulo_partes.append(zona)
            titulo_detector = f"Detector {idx}: " + " - ".join(titulo_partes)
            titulo_detector = _escapar_pdf(titulo_detector)
            estilo_titulo_detector = ParagraphStyle(
                'TituloDetector', parent=styles['Heading2'], fontSize=13,
                spaceBefore=0, spaceAfter=6,
            )
            story.append(Paragraph(titulo_detector, estilo_titulo_detector))
            
            fecha_retirada_optima_pdf = calcular_fecha_retirada_optima(fecha_det)

            # Las celdas se envuelven en Paragraph para que el texto largo
            # (p.ej. varios profesionales, o "Mañana + tarde + noche") se
            # reparta en varias líneas en vez de desbordar la celda; la
            # fila crece de alto automáticamente según haga falta.
            estilo_label_celda = ParagraphStyle(
                'CeldaLabel', parent=styles['Normal'], fontSize=9, leading=11, fontName='Helvetica-Bold',
            )
            estilo_valor_celda = ParagraphStyle(
                'CeldaValor', parent=styles['Normal'], fontSize=9, leading=11,
            )

            fecha_y_hora_colocacion = _escapar_pdf(fecha_det) or "-"
            if hora_colocacion:
                fecha_y_hora_colocacion += f"<br/>{_escapar_pdf(hora_colocacion)}"

            fecha_y_hora_retirada_real = ""
            if fecha_retirada_real:
                fecha_y_hora_retirada_real = _escapar_pdf(fecha_retirada_real)
                if hora_retirada_real:
                    fecha_y_hora_retirada_real += f"<br/>{_escapar_pdf(hora_retirada_real)}"

            filas_texto = [
                ["Codigo del detector", _escapar_pdf(codigo) or "-", "Fecha y hora de colocacion", fecha_y_hora_colocacion],
                ["Planta", _escapar_pdf(planta) or "-", "Nivel", _escapar_pdf(nivel) or "-"],
                ["Sala", _escapar_pdf(sala) or "-", "Codigo de la sala", _escapar_pdf(codigo_sala) or "-"],
                ["Profesionales en la sala", _escapar_pdf(profesionales_sala) or "-", "Turno de trabajo", _escapar_pdf(turno_trabajo) or "-"],
                ["Fecha optima retirada", _escapar_pdf(fecha_retirada_optima_pdf) or "-", "Fecha y hora real de retirada", fecha_y_hora_retirada_real],
                ["Resultado (Bq/m3)", _escapar_pdf(str(resultado_bq_m3)) if resultado_bq_m3 is not None else "-",
                 "Incertidumbre", _escapar_pdf(incertidumbre) or "-"],
            ]
            filas_pdf = []
            for fila_txt in filas_texto:
                fila_pdf = []
                for i, val in enumerate(fila_txt):
                    if not val:
                        fila_pdf.append("")
                    else:
                        estilo = estilo_label_celda if i % 2 == 0 else estilo_valor_celda
                        fila_pdf.append(Paragraph(val, estilo))
                filas_pdf.append(fila_pdf)

            tabla = Table(filas_pdf, colWidths=[3.3*cm, 4.7*cm, 3.3*cm, 4.7*cm])
            tabla.setStyle(TableStyle([
                ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
                ("BACKGROUND", (0,0), (0,-1), colors.whitesmoke),
                ("BACKGROUND", (2,0), (2,-1), colors.whitesmoke),
                ("VALIGN", (0,0), (-1,-1), "TOP"),
            ]))
            story.append(tabla)
            story.append(Spacer(1, 0.2*cm))

            estilo_subtitulo = ParagraphStyle(
                'SubtituloCompacto', parent=styles['Normal'], fontSize=10,
                fontName="Helvetica-Bold", spaceBefore=2, spaceAfter=2,
            )

            if plano and os.path.exists(plano):
                nombre_plano_esc = _escapar_pdf(nombre_plano_actual)
                titulo_ubicacion = f"Ubicacion en el plano ({nombre_plano_esc}):" if nombre_plano_esc else "Ubicacion en el plano:"
                story.append(Paragraph(titulo_ubicacion, estilo_subtitulo))
                try:
                    with PILImage.open(plano) as im_plano:
                        im_plano = im_plano.convert("RGB")
                        w, h = im_plano.size
                        if px is not None and py is not None and px >= 0 and py >= 0:
                            draw = ImageDraw.Draw(im_plano)
                            cx, cy = px * w, py * h
                            r = max(6, int(min(w, h) * 0.012))
                            draw.ellipse([cx - r, cy - r, cx + r, cy + r],
                                         fill=(220, 20, 20), outline=(120, 0, 0), width=2)
                        tmp_plano_path = os.path.join(get_data_dir(), f"_tmp_plano_{did}.jpg")
                        im_plano.save(tmp_plano_path, quality=90)
                    r = min(16.5*cm/w, 8.5*cm/h)
                    story.append(RLImage(tmp_plano_path, width=w*r, height=h*r))
                    story.append(Spacer(1, 0.2*cm))
                except Exception:
                    pass
            
            if foto_sit and os.path.exists(foto_sit) and foto_det and os.path.exists(foto_det):
                story.append(Paragraph("Fotos:", estilo_subtitulo))
                titulo_foto_style = ParagraphStyle(
                    'TituloFoto', parent=styles['Normal'], alignment=TA_CENTER, fontName="Helvetica-Bold",
                )
                try:
                    with PILImage.open(foto_sit) as im:
                        w, h = im.size
                    r = min(8.2*cm/w, 7.5*cm/h)
                    img1 = RLImage(foto_sit, width=w*r, height=h*r)
                except:
                    img1 = Paragraph("(no disponible)", styles["Normal"])
                try:
                    with PILImage.open(foto_det) as im:
                        w, h = im.size
                    r = min(8.2*cm/w, 7.5*cm/h)
                    img2 = RLImage(foto_det, width=w*r, height=h*r)
                except:
                    img2 = Paragraph("(no disponible)", styles["Normal"])
                cap1 = Paragraph("Situación del detector", titulo_foto_style)
                cap2 = Paragraph("Detector", titulo_foto_style)
                story.append(Table([[cap1, cap2], [img1, img2]], colWidths=[8.5*cm, 8.5*cm]))
        
        doc.build(story, onFirstPage=_dibujar_cabecera, onLaterPages=_dibujar_cabecera)
        return True
    except Exception as e:
        raise Exception(f"Error: {str(e)}")


# ============================================================
# GENERAR EXCEL (hoja de cálculo del centro, con las fotos de cada
# detector incrustadas cada una en su propia celda)
# ============================================================

# ============================================================
# CÁLCULO DE LA FECHA DE RETIRADA ÓPTIMA (90 días laborables desde la
# colocación, saltando fines de semana y festivos de Galicia)
# ============================================================

def _domingo_de_pascua(anio):
    """Algoritmo de Gauss/Meeus para calcular el Domingo de Pascua de
    un año dado (necesario para Jueves y Viernes Santo, festivos
    movibles que dependen de la Pascua cada año)."""
    a = anio % 19
    b = anio // 100
    c = anio % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    mes = (h + l - 7 * m + 114) // 31
    dia = ((h + l - 7 * m + 114) % 31) + 1
    return date(anio, mes, dia)


def _festivos_galicia(anio):
    """Festivos de ámbito estatal/autonómico aplicables en Galicia para
    un año dado (no incluye fiestas locales de cada ayuntamiento, que
    varían por municipio)."""
    pascua = _domingo_de_pascua(anio)
    jueves_santo = pascua - timedelta(days=3)
    viernes_santo = pascua - timedelta(days=2)
    fijos = [
        date(anio, 1, 1),    # Año Nuevo
        date(anio, 1, 6),    # Reyes
        date(anio, 5, 1),    # Fiesta del Trabajo
        date(anio, 7, 25),   # Santiago Apóstol (patrón de Galicia)
        date(anio, 8, 15),   # Asunción
        date(anio, 10, 12),  # Fiesta Nacional de España
        date(anio, 11, 1),   # Todos los Santos
        date(anio, 12, 6),   # Día de la Constitución
        date(anio, 12, 8),   # Inmaculada Concepción
        date(anio, 12, 25),  # Navidad
    ]
    return set(fijos + [jueves_santo, viernes_santo])


def _parsear_fecha(texto):
    """Intenta interpretar una fecha escrita en varios formatos
    habituales. Devuelve un date, o None si no se pudo interpretar."""
    texto = (texto or "").strip()
    if not texto:
        return None
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
        try:
            return datetime.strptime(texto, fmt).date()
        except ValueError:
            continue
    return None


def calcular_fecha_retirada_optima(fecha_colocacion_texto):
    """90 días naturales desde la fecha de colocación; si ese día cae
    en sábado, domingo o festivo de Galicia, se pasa al siguiente día
    laborable. Devuelve el resultado como texto "DD/MM/AAAA", o ""
    si la fecha de colocación no se pudo interpretar."""
    fecha_col = _parsear_fecha(fecha_colocacion_texto)
    if fecha_col is None:
        return ""
    objetivo = fecha_col + timedelta(days=90)
    cache_festivos = {}
    while True:
        if objetivo.year not in cache_festivos:
            cache_festivos[objetivo.year] = _festivos_galicia(objetivo.year)
        if objetivo.weekday() >= 5 or objetivo in cache_festivos[objetivo.year]:
            objetivo += timedelta(days=1)
            continue
        break
    return objetivo.strftime("%d/%m/%Y")


def importar_centro_desde_excel(archivo_bytes):
    """Reconstruye un centro completo (datos, planos y detectores, con
    el punto exacto de cada uno sobre su plano) a partir de un Excel
    generado por generar_excel.

    Devuelve (centro_id, numero_de_detectores_creados). Lanza
    ValueError con un mensaje claro si el archivo no tiene la
    estructura esperada (no es un Excel exportado por esta app).
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(archivo_bytes))
    if "Detectores" not in wb.sheetnames:
        raise ValueError(
            "El archivo no parece un Excel exportado por esta app "
            "(falta la hoja 'Detectores')."
        )
    ws = wb["Detectores"]

    nombre_centro = ws.cell(row=1, column=1).value or "Centro importado"
    area_centro = ws.cell(row=1, column=5).value or ""
    tecnico_centro = ws.cell(row=1, column=10).value or ""
    fecha_centro = ws.cell(row=1, column=15).value or ""
    direccion_centro = ws.cell(row=1, column=17).value or ""

    header_row = 2
    headers = {}
    for col in range(1, ws.max_column + 1):
        val = ws.cell(row=header_row, column=col).value
        if val:
            headers[val] = col

    columnas_esperadas = ["Centro", "Área / Zona", "Planta", "Sala", "Código"]
    faltan = [c for c in columnas_esperadas if c not in headers]
    if faltan:
        raise ValueError(
            "El archivo no tiene el formato esperado (faltan columnas: "
            + ", ".join(faltan) + ")."
        )

    # Imágenes incrustadas de la hoja "Detectores", indexadas por la
    # posición exacta de celda (fila0, col0) en la que se anclaron al
    # generarlas, para poder recuperar la foto de cada detector.
    imagenes_por_celda = {}
    for img in ws._images:
        try:
            fr = img.anchor._from
            imagenes_por_celda[(fr.row, fr.col)] = img._data()
        except Exception:
            continue

    centro_id = crear_centro(str(nombre_centro).strip() or "Centro importado")
    imagen_exterior_path = None

    # --- Hoja "Planos": crea TODOS los planos del centro (incluidos los
    # que ningún detector tenga asignado) y construye el mapa
    # nombre -> id para poder vincular cada detector a su plano.
    # También trae la Empresa y el CIF (guardados en esta misma hoja,
    # en sus dos primeras filas). ---
    planos_por_nombre = {}
    datos_informe_importados = {}
    if "Planos" in wb.sheetnames:
        ws_p = wb["Planos"]

        empresa_importada = ws_p.cell(row=1, column=2).value
        cif_importado = ws_p.cell(row=2, column=2).value
        if empresa_importada:
            set_empresa(str(empresa_importada).strip())
        if cif_importado:
            set_cif(str(cif_importado).strip())

        # Restaurar metadatos adicionales del informe si existen en el Excel.
        etiquetas_meta = {
            "Superficie construida": "superficie_construida",
            "Superficie útil": "superficie_util",
            "Nº de plantas": "num_plantas",
            "Fecha comunicación trabajadores": "fecha_comunicacion_trab",
            "Medio de comunicación": "medio_comunicacion",
        }
        datos_informe_importados = {}
        for fila_meta in range(3, min(ws_p.max_row, 10) + 1):
            etiqueta_meta = ws_p.cell(row=fila_meta, column=1).value
            valor_meta = ws_p.cell(row=fila_meta, column=2).value
            campo_meta = etiquetas_meta.get(str(etiqueta_meta).strip()) if etiqueta_meta else None
            if campo_meta and valor_meta is not None:
                datos_informe_importados[campo_meta] = str(valor_meta).strip()

        imagenes_planos_por_celda = {}
        for img in ws_p._images:
            try:
                fr = img.anchor._from
                imagenes_planos_por_celda[(fr.row, fr.col)] = img._data()
            except Exception:
                continue
        orden = 0
        # Las filas 1-4 son Empresa, CIF, una fila en blanco y la
        # cabecera "Nombre"/"Imagen"; los planos empiezan en la 5.
        for fila_p in range(5, ws_p.max_row + 1):
            nombre_p = ws_p.cell(row=fila_p, column=1).value
            if not nombre_p:
                continue
            datos_img = imagenes_planos_por_celda.get((fila_p - 1, 1))
            if not datos_img:
                continue
            ruta_guardada = guardar_bytes_imagen(datos_img, f"plano_centro_{centro_id}")
            if str(nombre_p) == "(Foto exterior del centro)":
                imagen_exterior_path = ruta_guardada
            else:
                nuevo_id = insert_plano_centro(centro_id, str(nombre_p), ruta_guardada, orden)
                planos_por_nombre[str(nombre_p)] = nuevo_id
                orden += 1

    update_centro(
        centro_id, str(nombre_centro).strip() or "Centro importado",
        str(area_centro).strip(), str(fecha_centro).strip(), imagen_exterior_path,
        str(direccion_centro).strip(),
    )
    if tecnico_centro:
        set_tecnico_centro(centro_id, str(tecnico_centro).strip())
    if datos_informe_importados:
        set_datos_informe_centro(centro_id, **datos_informe_importados)

    # --- Hoja "Categorías profesionales" (si existe) ---
    if "Categorías profesionales" in wb.sheetnames:
        ws_cat = wb["Categorías profesionales"]
        for fila_cat in range(2, ws_cat.max_row + 1):
            categoria_val = ws_cat.cell(row=fila_cat, column=1).value
            num_val = ws_cat.cell(row=fila_cat, column=2).value
            if categoria_val:
                try:
                    num_final = int(num_val) if num_val is not None else 0
                except (TypeError, ValueError):
                    num_final = 0
                turno_val = ws_cat.cell(row=fila_cat, column=3).value if ws_cat.max_column >= 3 else ""
                insert_categoria_centro(centro_id, str(categoria_val).strip(), num_final, str(turno_val or "").strip())

    def _val(fila, nombre_col, default=""):
        col = headers.get(nombre_col)
        if not col:
            return default
        v = ws.cell(row=fila, column=col).value
        return v if v is not None else default

    detectores_creados = 0
    for fila_actual in range(header_row + 1, ws.max_row + 1):
        codigo_det = _val(fila_actual, "Código")
        sala_det = _val(fila_actual, "Sala")
        if not str(codigo_det).strip() and not str(sala_det).strip():
            continue  # fila vacía, se ignora

        planta = _val(fila_actual, "Planta")
        nivel = _val(fila_actual, "Nivel")
        codigo_sala = _val(fila_actual, "Código de la sala")
        profesionales_multilinea = str(_val(fila_actual, "Profesionales en la sala"))
        profesionales = ", ".join(
            linea.strip() for linea in profesionales_multilinea.split("\n") if linea.strip()
        )
        turno = _val(fila_actual, "Turno de trabajo")
        fecha_colocacion = _val(fila_actual, "Fecha de colocación")
        hora_colocacion = _val(fila_actual, "Hora de colocación")
        fecha_retirada_real = _val(fila_actual, "Fecha de retirada real")
        hora_retirada_real = _val(fila_actual, "Hora de retirada real")
        nombre_plano_fila = _val(fila_actual, "Nombre del plano")
        punto_x_fila = _val(fila_actual, "Punto X", None)
        punto_y_fila = _val(fila_actual, "Punto Y", None)

        plano_centro_id = planos_por_nombre.get(str(nombre_plano_fila)) if nombre_plano_fila else None

        try:
            punto_x_final = float(punto_x_fila) if punto_x_fila not in (None, "") else -1
            punto_y_final = float(punto_y_fila) if punto_y_fila not in (None, "") else -1
        except (TypeError, ValueError):
            punto_x_final = punto_y_final = -1

        foto_sit_path = None
        foto_det_path = None
        col_situacion = headers.get("Foto situación")
        col_detector_foto = headers.get("Foto detector")
        if col_situacion:
            datos = imagenes_por_celda.get((fila_actual - 1, col_situacion - 1))
            if datos:
                foto_sit_path = guardar_bytes_imagen(datos, "foto_situacion")
        if col_detector_foto:
            datos = imagenes_por_celda.get((fila_actual - 1, col_detector_foto - 1))
            if datos:
                foto_det_path = guardar_bytes_imagen(datos, "foto_detector")

        data = (
            centro_id, str(planta), str(sala_det), str(fecha_colocacion), str(codigo_det),
            None, punto_x_final, punto_y_final, foto_sit_path, foto_det_path,
            _ahora_espana().strftime("%Y-%m-%d %H:%M"),
            str(codigo_sala), profesionales, str(hora_colocacion),
            str(turno) if turno else TURNOS_TRABAJO_OPCIONES[0],
            str(nivel) if nivel else NIVEL_OPCIONES[0],
            plano_centro_id, str(fecha_retirada_real), str(hora_retirada_real),
        )
        nuevo_detector_id = insert_detector(data)

        # Restaurar también los resultados de laboratorio si estaban guardados en el Excel.
        resultado_importado = _val(fila_actual, "Resultado (Bq/m³/h)", "")
        incertidumbre_importada = _val(fila_actual, "Incertidumbre", "")
        resultado_final = None
        if resultado_importado not in (None, ""):
            try:
                resultado_final = float(resultado_importado)
            except (TypeError, ValueError):
                resultado_final = None
        actualizar_resultado_detector(
            nuevo_detector_id,
            resultado_final,
            "" if incertidumbre_importada is None else str(incertidumbre_importada),
        )

        detectores_creados += 1

    return centro_id, detectores_creados


def generar_excel(centro_id, output_path):
    """Genera un .xlsx en formato de tabla "plana" (sin ningún bloque de
    cabecera encima), pensado para poder ir pegando debajo los datos de
    otros centros más adelante. El nombre del centro y el área figuran
    como las dos primeras columnas de cada fila, junto con todos los
    datos de cada detector, la fecha de retirada óptima calculada
    automáticamente, una columna vacía para anotar la retirada real, y
    sus tres fotos (plano con el punto marcado, situación, detector)
    incrustadas cada una en su propia celda, centradas y manteniendo su
    proporción real."""
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.drawing.spreadsheet_drawing import OneCellAnchor, AnchorMarker
    from openpyxl.drawing.xdr import XDRPositiveSize2D
    from openpyxl.utils.units import pixels_to_EMU
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    centro = get_centro(centro_id)
    if not centro:
        raise Exception("Centro no encontrado")
    cid, nombre, zona, fecha, img_centro, tecnico, direccion = centro
    detectores = fetch_detectores(centro_id)

    FUENTE = "Arial"
    fuente_normal = Font(name=FUENTE, size=7)
    fuente_cabecera = Font(name=FUENTE, size=8, bold=True, color="FFFFFF")
    fuente_info_grande = Font(name=FUENTE, size=12, bold=True, color="FFFFFF")
    centrado = Alignment(horizontal="center", vertical="center", wrap_text=True)
    # Estilo especial para la primera columna ("Centro"): sin fondo
    # naranja en las filas de datos, alineado a la derecha.
    alineado_derecha = Alignment(horizontal="right", vertical="center", wrap_text=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "Detectores"

    borde_fino = Border(*[Side(style="thin", color="CCCCCC")] * 4)

    # --- Línea de cabecera superior con el resumen del centro (Centro,
    # Área, Día, Técnico), y debajo la fila de cabeceras de columna
    # (al doble de alto de lo normal). ---
    fila_info = 1
    header_row = 2
    headers = [
        "Centro", "Área / Zona", "ID", "Planta", "Nivel", "Sala", "Código de la sala",
        "Profesionales en la sala", "Turno de trabajo", "Código",
        "Resultado (Bq/m³/h)", "Incertidumbre",
        "Fecha de colocación", "Hora de colocación",
        "Fecha de retirada óptima", "Fecha de retirada real", "Hora de retirada real",
        "Nombre del plano", "Punto X", "Punto Y",
        "Plano", "Foto situación", "Foto detector",
    ]

    col_letra = {h: get_column_letter(i) for i, h in enumerate(headers, start=1)}
    col_idx0 = {h: i for i, h in enumerate(headers)}  # índice base-0, para los anclajes

    # Línea superior en forma de banner: el nombre del centro ocupa las
    # 4 primeras columnas fusionadas, el área las 4 siguientes, luego
    # una columna con la palabra "Técnico" y las 4 siguientes con su
    # nombre fusionadas. El resto de columnas quedan en naranja liso
    # para que la fila entera se vea como un único banner.
    ws.merge_cells(start_row=fila_info, start_column=1, end_row=fila_info, end_column=4)
    c_centro = ws.cell(row=fila_info, column=1, value=nombre or "")
    c_centro.font = fuente_info_grande
    c_centro.fill = PatternFill("solid", fgColor="F5A623")
    c_centro.alignment = centrado

    ws.merge_cells(start_row=fila_info, start_column=5, end_row=fila_info, end_column=8)
    c_area = ws.cell(row=fila_info, column=5, value=zona or "")
    c_area.font = fuente_info_grande
    c_area.fill = PatternFill("solid", fgColor="F5A623")
    c_area.alignment = centrado

    c_tecnico_label = ws.cell(row=fila_info, column=9, value="Técnico")
    c_tecnico_label.font = fuente_cabecera
    c_tecnico_label.fill = PatternFill("solid", fgColor="F5A623")
    c_tecnico_label.alignment = centrado

    ws.merge_cells(start_row=fila_info, start_column=10, end_row=fila_info, end_column=13)
    c_tecnico_val = ws.cell(row=fila_info, column=10, value=tecnico or "")
    c_tecnico_val.font = fuente_cabecera
    c_tecnico_val.fill = PatternFill("solid", fgColor="F5A623")
    c_tecnico_val.alignment = centrado

    # Fecha del centro: no se muestra a gran tamaño como Centro/Área,
    # pero se guarda en el banner para poder reconstruirla al importar.
    c_fecha_label = ws.cell(row=fila_info, column=14, value="Fecha")
    c_fecha_label.font = fuente_cabecera
    c_fecha_label.fill = PatternFill("solid", fgColor="F5A623")
    c_fecha_label.alignment = centrado
    c_fecha_val = ws.cell(row=fila_info, column=15, value=fecha or "")
    c_fecha_val.font = fuente_cabecera
    c_fecha_val.fill = PatternFill("solid", fgColor="F5A623")
    c_fecha_val.alignment = centrado

    # Dirección del centro: igual que la fecha, se guarda en el banner
    # (aunque no se muestre a gran tamaño) para poder recuperarla si se
    # reimporta este Excel más adelante.
    c_dir_label = ws.cell(row=fila_info, column=16, value="Dirección")
    c_dir_label.font = fuente_cabecera
    c_dir_label.fill = PatternFill("solid", fgColor="F5A623")
    c_dir_label.alignment = centrado
    c_dir_val = ws.cell(row=fila_info, column=17, value=direccion or "")
    c_dir_val.font = fuente_cabecera
    c_dir_val.fill = PatternFill("solid", fgColor="F5A623")
    c_dir_val.alignment = centrado

    # El resto de la fila (hasta donde llegan las columnas de la tabla
    # de abajo) también en naranja con texto centrado, aunque no haya
    # más datos que mostrar, para que quede como un único banner.
    for col in range(18, len(headers) + 1):
        c_resto = ws.cell(row=fila_info, column=col)
        c_resto.fill = PatternFill("solid", fgColor="F5A623")
        c_resto.alignment = centrado
        c_resto.font = fuente_cabecera
    ws.row_dimensions[fila_info].height = 22

    for i, h in enumerate(headers, start=1):
        c = ws.cell(row=header_row, column=i, value=h)
        if i == 1:
            # "Centro": mismo estilo que el resto de cabeceras (texto
            # blanco, fondo naranja), pero alineado a la derecha.
            c.font = fuente_cabecera
            c.fill = PatternFill("solid", fgColor="F5A623")
            c.alignment = alineado_derecha
        else:
            c.font = fuente_cabecera
            c.fill = PatternFill("solid", fgColor="F5A623")
            c.alignment = centrado
        c.border = borde_fino
    ws.row_dimensions[header_row].height = 30  # el doble de una fila normal (~15pt)

    ANCHO_FOTOS = 24
    for h in ("Plano", "Foto situación", "Foto detector"):
        ws.column_dimensions[col_letra[h]].width = ANCHO_FOTOS
    for h in ("Centro", "Planta", "Nivel", "Sala", "Código de la sala", "Profesionales en la sala",
              "Turno de trabajo", "Código"):
        ws.column_dimensions[col_letra[h]].width = 15
    for h in ("Fecha de colocación", "Fecha de retirada óptima", "Fecha de retirada real"):
        ws.column_dimensions[col_letra[h]].width = 16
    for h in ("Hora de colocación", "Hora de retirada real"):
        ws.column_dimensions[col_letra[h]].width = 10
    ws.column_dimensions[col_letra["Área / Zona"]].width = 16
    ws.column_dimensions[col_letra["Nombre del plano"]].width = 14
    for h in ("Punto X", "Punto Y"):
        ws.column_dimensions[col_letra[h]].width = 8

    TAM_IMG_PX = 150  # tamaño MÁXIMO (lado mayor) de cada foto dentro de su celda
    ALTO_FILA_PT = 115  # alto de fila (puntos) reservado para las fotos

    def _ancho_col_px(col_chars):
        # Aproximación estándar: píxeles ≈ caracteres * 7 + 5
        return round(col_chars * 7 + 5)

    def _alto_fila_px(alto_pt):
        # 1 punto ≈ 1.333 píxeles (96 dpi)
        return round(alto_pt * 96 / 72)

    def _anclaje_centrado(col_h, fila1, ancho_img_px, alto_img_px):
        """Ancla la imagen centrada (horizontal y verticalmente) dentro
        de su celda, calculando el hueco libre a cada lado."""
        ancho_celda_px = _ancho_col_px(ws.column_dimensions[col_letra[col_h]].width)
        alto_celda_px = _alto_fila_px(ALTO_FILA_PT)
        off_x = max(0, (ancho_celda_px - ancho_img_px) / 2)
        off_y = max(0, (alto_celda_px - alto_img_px) / 2)
        marker = AnchorMarker(
            col=col_idx0[col_h], row=fila1 - 1,
            colOff=pixels_to_EMU(off_x), rowOff=pixels_to_EMU(off_y),
        )
        size = XDRPositiveSize2D(cx=pixels_to_EMU(ancho_img_px), cy=pixels_to_EMU(alto_img_px))
        return OneCellAnchor(_from=marker, ext=size)

    fila = header_row + 1
    for d in detectores:
        (did, _, planta, sala, fecha_det, codigo, _plano_antiguo, punto_x, punto_y,
         foto_sit_p, foto_det_p, fecha_creacion, codigo_sala, profesionales_sala,
         hora_colocacion, turno_trabajo, nivel, plano_centro_id, fecha_retirada_real, hora_retirada_real,
         resultado_bq_m3_d, incertidumbre_d) = d
        plano_p = None
        nombre_plano_d = ""
        if plano_centro_id:
            plano_info = get_plano_centro(plano_centro_id)
            if plano_info:
                plano_p = plano_info[3]
                nombre_plano_d = plano_info[2]

        fecha_retirada_optima = calcular_fecha_retirada_optima(fecha_det)

        # Los profesionales se escriben separados por comas en el
        # formulario; en la celda del Excel se muestran uno debajo de
        # otro (salto de línea + texto ajustado a la celda).
        profesionales_multilinea = "\n".join(
            p.strip() for p in (profesionales_sala or "").split(",") if p.strip()
        )

        hay_punto_valido = (
            punto_x is not None and punto_y is not None and 0 <= punto_x <= 1 and 0 <= punto_y <= 1
        )

        valores = [
            nombre or "", zona or "", did, planta or "", nivel or "", sala or "",
            codigo_sala or "", profesionales_multilinea, turno_trabajo or "", codigo or "",
            resultado_bq_m3_d if resultado_bq_m3_d is not None else "", incertidumbre_d or "",
            fecha_det or "", hora_colocacion or "",
            fecha_retirada_optima, fecha_retirada_real or "", hora_retirada_real or "",
            nombre_plano_d, round(punto_x, 4) if hay_punto_valido else "",
            round(punto_y, 4) if hay_punto_valido else "",
        ]
        for col, val in enumerate(valores, start=1):
            c = ws.cell(row=fila, column=col, value=val)
            if col == 1:
                # "Centro": texto negro (normal), sin fondo, a la derecha.
                c.font = fuente_normal
                c.alignment = alineado_derecha
            else:
                c.font = fuente_normal
                c.alignment = centrado
            c.border = borde_fino
        # Las tres celdas de fotos también con borde, aunque queden vacías
        for h in ("Plano", "Foto situación", "Foto detector"):
            ws[f"{col_letra[h]}{fila}"].border = borde_fino

        ws.row_dimensions[fila].height = ALTO_FILA_PT

        hay_punto = hay_punto_valido

        for col_name, ruta in (
            ("Plano", plano_p),
            ("Foto situación", foto_sit_p),
            ("Foto detector", foto_det_p),
        ):
            if ruta and os.path.exists(ruta):
                try:
                    with Image.open(ruta) as im_orig:
                        # IMPORTANTE: las fotos de móvil llevan a menudo
                        # metadatos EXIF de rotación que PIL no aplica
                        # solo; sin esto, el ancho/alto "en bruto" no
                        # coincide con cómo se ve realmente la foto y la
                        # proporción calculada sale mal.
                        im = ImageOps.exif_transpose(im_orig)
                        im = im.convert("RGB")
                        if col_name == "Plano" and hay_punto:
                            # Dibujar el punto rojo del detector sobre el
                            # plano, igual que en el informe PDF.
                            draw = ImageDraw.Draw(im)
                            w, h = im.size
                            cx, cy = punto_x * w, punto_y * h
                            r = max(6, int(min(w, h) * 0.02))
                            draw.ellipse(
                                [cx - r, cy - r, cx + r, cy + r],
                                fill=(220, 20, 20), outline=(120, 0, 0), width=2,
                            )
                        # Se sube algo la resolución respecto a antes
                        # (se veía muy borrosa) sin pasarse: esta misma
                        # imagen va dentro del propio Excel, que a
                        # veces se comparte por WhatsApp, así que
                        # engordarla demasiado podía hacer fallar ese
                        # envío por límite de tamaño.
                        im.thumbnail((600, 600))
                        buf = io.BytesIO()
                        im.save(buf, format="JPEG", quality=85)
                        buf.seek(0)
                        ancho_real, alto_real = im.size

                    # Escalar manteniendo la proporción real (sin
                    # deformar), ajustando al lado mayor = TAM_IMG_PX.
                    escala = min(TAM_IMG_PX / ancho_real, TAM_IMG_PX / alto_real)
                    ancho_final = max(1, round(ancho_real * escala))
                    alto_final = max(1, round(alto_real * escala))
                    xl_img = XLImage(buf)
                    xl_img.width = ancho_final
                    xl_img.height = alto_final
                    xl_img.anchor = _anclaje_centrado(col_name, fila, ancho_final, alto_final)
                    ws.add_image(xl_img)
                except Exception:
                    c_err = ws[f"{col_letra[col_name]}{fila}"]
                    c_err.value = "(no se pudo incrustar la imagen)"
                    c_err.font = fuente_normal
                    c_err.alignment = centrado
        fila += 1

    # --- Hoja "Planos": TODOS los planos del centro, uno por fila,
    # tengan o no algún detector que los use (los del "Detectores" solo
    # aparecen si algún detector los tiene asignado; esta hoja evita
    # que un plano "huérfano" se pierda al reimportar). También incluye
    # la foto exterior del centro, que no aparece en ningún otro sitio
    # del Excel. ---
    ws_planos = wb.create_sheet("Planos")

    # Empresa y CIF (datos globales de la empresa que realiza las
    # mediciones), en las dos primeras filas de esta hoja.
    empresa_xl = get_empresa()
    cif_xl = get_cif()
    ws_planos.cell(row=1, column=1, value="Empresa").font = fuente_cabecera
    ws_planos.cell(row=1, column=1).fill = PatternFill("solid", fgColor="F5A623")
    ws_planos.cell(row=1, column=2, value=empresa_xl).font = fuente_normal
    ws_planos.cell(row=2, column=1, value="CIF").font = fuente_cabecera
    ws_planos.cell(row=2, column=1).fill = PatternFill("solid", fgColor="F5A623")
    ws_planos.cell(row=2, column=2, value=cif_xl).font = fuente_normal

    # Datos del informe final (superficie, nº de plantas, comunicación
    # a los trabajadores) que se rellenan en esa pantalla: se guardan
    # aquí para que no se pierdan si se vuelve a generar el Excel.
    datos_informe_xl = get_datos_informe_centro(centro_id)
    filas_datos_informe = [
        ("Superficie construida", datos_informe_xl["superficie_construida"]),
        ("Superficie útil", datos_informe_xl["superficie_util"]),
        ("Nº de plantas", datos_informe_xl["num_plantas"]),
        ("Fecha comunicación trabajadores", datos_informe_xl["fecha_comunicacion_trab"]),
        ("Medio de comunicación", datos_informe_xl["medio_comunicacion"]),
    ]
    for i, (etiqueta_di, valor_di) in enumerate(filas_datos_informe, start=3):
        ws_planos.cell(row=i, column=1, value=etiqueta_di).font = fuente_cabecera
        ws_planos.cell(row=i, column=1).fill = PatternFill("solid", fgColor="F5A623")
        ws_planos.cell(row=i, column=2, value=valor_di).font = fuente_normal

    header_row_planos = 9
    ws_planos.cell(row=header_row_planos, column=1, value="Nombre").font = fuente_cabecera
    ws_planos.cell(row=header_row_planos, column=1).fill = PatternFill("solid", fgColor="F5A623")
    ws_planos.cell(row=header_row_planos, column=2, value="Imagen").font = fuente_cabecera
    ws_planos.cell(row=header_row_planos, column=2).fill = PatternFill("solid", fgColor="F5A623")
    ws_planos.column_dimensions["A"].width = 20
    ws_planos.column_dimensions["B"].width = 28

    fila_planos = header_row_planos + 1

    def _incrustar_en_hoja_planos(nombre_fila, ruta_imagen, fila_actual):
        ws_planos.cell(row=fila_actual, column=1, value=nombre_fila).font = fuente_normal
        ws_planos.row_dimensions[fila_actual].height = 115
        if ruta_imagen and os.path.exists(ruta_imagen):
            try:
                with Image.open(ruta_imagen) as im_orig:
                    im = ImageOps.exif_transpose(im_orig)
                    im = im.convert("RGB")
                    # Resolución alta (aunque en la celda del Excel se
                    # vea pequeña, con este mismo dato se genera luego
                    # el plano automático del Anexo II a tamaño de
                    # página completa, así que hace falta buena
                    # definición de origen para que no se vea borroso).
                    im.thumbnail((2000, 2000))
                    buf = io.BytesIO()
                    im.save(buf, format="JPEG", quality=92)
                    buf.seek(0)
                    ancho_real, alto_real = im.size
                escala = min(180 / ancho_real, 150 / alto_real)
                xl_img = XLImage(buf)
                xl_img.width = max(1, round(ancho_real * escala))
                xl_img.height = max(1, round(alto_real * escala))
                ws_planos.add_image(xl_img, f"B{fila_actual}")
            except Exception:
                ws_planos.cell(row=fila_actual, column=2, value="(no se pudo incrustar la imagen)")

    if img_centro and os.path.exists(img_centro):
        _incrustar_en_hoja_planos("(Foto exterior del centro)", img_centro, fila_planos)
        fila_planos += 1

    for plano_c in fetch_planos_centro(centro_id):
        _, _, nombre_plano_c, ruta_plano_c, _ = plano_c
        _incrustar_en_hoja_planos(nombre_plano_c, ruta_plano_c, fila_planos)
        fila_planos += 1

    # --- Hoja "Categorías profesionales": una fila por categoría con
    # el número de personas expuestas de esa categoría en el centro. ---
    categorias_centro_xl = fetch_categorias_centro(centro_id)
    if categorias_centro_xl:
        ws_cat = wb.create_sheet("Categorías profesionales")
        ws_cat.cell(row=1, column=1, value="Categoría profesional").font = fuente_cabecera
        ws_cat.cell(row=1, column=1).fill = PatternFill("solid", fgColor="F5A623")
        ws_cat.cell(row=1, column=2, value="Nº personas expuestas").font = fuente_cabecera
        ws_cat.cell(row=1, column=2).fill = PatternFill("solid", fgColor="F5A623")
        ws_cat.cell(row=1, column=3, value="Turno").font = fuente_cabecera
        ws_cat.cell(row=1, column=3).fill = PatternFill("solid", fgColor="F5A623")
        ws_cat.column_dimensions["A"].width = 28
        ws_cat.column_dimensions["B"].width = 20
        ws_cat.column_dimensions["C"].width = 22
        for i, (_, _, categoria_xl, num_personas_xl, turno_xl) in enumerate(categorias_centro_xl, start=2):
            ws_cat.cell(row=i, column=1, value=categoria_xl).font = fuente_normal
            ws_cat.cell(row=i, column=2, value=num_personas_xl).font = fuente_normal
            ws_cat.cell(row=i, column=3, value=turno_xl or "").font = fuente_normal

    wb.save(output_path)
    return True


# ============================================================
# REGISTRO PARA LABORATORIO
# Ficha calcada del formulario "FICHA DE IDENTIFICACIÓN E INFORMACIÓN
# DE LOS DETECTORES DE TRAZAS" que exige el laboratorio de análisis,
# para que la acepten sin tener que rellenarla aparte.
# ============================================================

def generar_registro_laboratorio(centro_id, output_path, tipo_firma="digital"):
    """Genera la ficha de identificación de detectores en el formato
    que exige el laboratorio, con una tabla prácticamente idéntica a
    su propio formulario en papel, más un cuadro de firma del técnico.

    tipo_firma="manual": no se añade ningún cuadro de firma (se firma
    sobre el papel impreso).
    tipo_firma="digital": se añade un cuadro pequeño de firma digital
    en la esquina de TODAS las páginas del documento.
    """
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.platypus import Image as RLImage
    from reportlab.lib import colors
    from PIL import Image as PILImage

    centro = get_centro(centro_id)
    if not centro:
        raise ValueError("Centro no encontrado")
    _, nombre, zona, fecha, img_ext, tecnico, direccion = centro
    detectores = fetch_detectores(centro_id)
    empresa = get_empresa()

    def _planta_desde_nivel(nivel_valor):
        """El laboratorio no quiere el texto libre de "Nivel" tal cual
        se escribe en la app, sino el número de planta (con signo)
        que le corresponde."""
        return NIVEL_A_PLANTA_LABORATORIO.get(nivel_valor, nivel_valor or "")

    def _ubicado_en(zona_valor, nombre_centro):
        """"Ubicado en" es el área del centro (p.ej. \"Área Sanitaria da
        Coruña e Cee\"). Si esa área está vacía, o es un valor genérico
        que no identifica bien el centro (\"Atención Primaria\", \"PAC\",
        \"Atención Primaria + PAC\", \"Consultorio\"), se usa el nombre del
        propio centro en su lugar."""
        valores_genericos = (
            "atención primaria", "atencion primaria", "pac",
            "atención primaria + pac", "atencion primaria + pac",
            "consultorio",
        )
        if not zona_valor or zona_valor.strip().lower() in valores_genericos:
            return nombre_centro or ""
        return zona_valor

    # Con firma digital reservamos margen inferior de sobra para que el
    # cuadro de firma (dibujado directamente en cada página) no se
    # solape nunca con las últimas filas de la tabla. El margen superior
    # se reserva SIEMPRE para la cabecera (logo + datos + condiciones),
    # que también se dibuja directamente en cada página, para que
    # aparezca completa en todas las hojas, no solo en la primera.
    ALTURA_CABECERA = 2.3*cm
    SEPARACION_CABECERA = 0.3*cm
    margen_superior = 1.2*cm + ALTURA_CABECERA + SEPARACION_CABECERA
    margen_inferior = 3.4*cm if tipo_firma == "digital" else 1.2*cm
    doc = SimpleDocTemplate(
        output_path, pagesize=landscape(A4),
        topMargin=margen_superior, bottomMargin=margen_inferior, leftMargin=1.2*cm, rightMargin=1.2*cm,
    )
    styles = getSampleStyleSheet()
    story = []

    GRIS_CLARO = colors.HexColor("#D9D9D9")

    estilo_celda = ParagraphStyle(
        "CeldaLab", parent=styles["Normal"], fontSize=7, leading=8.5, alignment=TA_CENTER,
    )
    estilo_celda_izq = ParagraphStyle(
        "CeldaLabIzq", parent=styles["Normal"], fontSize=7.5, leading=9, alignment=TA_LEFT,
    )
    estilo_cab = ParagraphStyle(
        "CabLab", parent=styles["Normal"], fontSize=7, leading=8.5, alignment=TA_CENTER,
        textColor=colors.black, fontName="Helvetica-Bold",
    )

    col_widths = [0.8*cm, 2.4*cm, 2.6*cm, 1.2*cm, 2.6*cm, 3.4*cm, 1.1*cm,
                  1.7*cm, 1.2*cm, 3.4*cm, 1.7*cm, 1.2*cm, 3.4*cm]
    ancho_hasta_mitad_habitacion = sum(col_widths[:4]) + col_widths[4] / 2
    ancho_total = sum(col_widths)

    # --- Logotipo del laboratorio, a la IZQUIERDA (hasta la mitad de
    # la columna "Habitación/Estancia"), y a su derecha, en un único
    # bloque sin líneas internas, los datos de obra/inmueble/cliente y
    # la barra de condiciones del edificio. Ambos con la misma altura.
    # Todo este bloque se dibuja directamente en el lienzo de CADA
    # página (ver _dibujar_elementos_pagina más abajo), no como
    # contenido normal de la tabla, para que aparezca completo en
    # todas las hojas del documento. ---
    ruta_logo = get_logo_laboratorio()
    ruta_logo_default = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "logo_laboratorio_default.png"
    )
    es_logo_default = bool(ruta_logo) and os.path.normpath(ruta_logo) == os.path.normpath(ruta_logo_default)

    contenido_logo = ""
    if ruta_logo and os.path.exists(ruta_logo):
        try:
            with PILImage.open(ruta_logo) as im_logo:
                w_logo, h_logo = im_logo.size
            max_w = ancho_hasta_mitad_habitacion - 0.4*cm
            max_h = ALTURA_CABECERA - 0.4*cm
            r = min(max_w / w_logo, max_h / h_logo)
            contenido_logo = RLImage(ruta_logo, width=w_logo * r, height=h_logo * r)
        except Exception:
            contenido_logo = ""

    def _construir_caja_cabecera():
        """Se reconstruye igual en cada llamada (una por página) para
        evitar reutilizar el mismo objeto Table ya "usado" por
        reportlab al dibujarlo en una página anterior."""
        caja_logo = Table(
            [[contenido_logo]], colWidths=[ancho_hasta_mitad_habitacion], rowHeights=[ALTURA_CABECERA],
        )
        caja_logo.setStyle(TableStyle([
            ("LINEAFTER", (0, 0), (-1, -1), 1, colors.black),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))

        # El código de formulario solo tiene sentido con el logotipo
        # POR DEFECTO de la app (es el código de SU ficha); si se ha
        # cambiado por el logotipo de otro laboratorio, ese código ya
        # no pinta nada y se quita.
        if es_logo_default:
            celda_orden_trabajo = [
                Paragraph("CF-PE-CYE/39 V.3 (30/05/2025)", ParagraphStyle(
                    "CodigoDerecha", parent=styles["Normal"], fontSize=5.5, leading=7,
                    alignment=TA_RIGHT, textColor=colors.black,
                )),
                Paragraph("<b>Orden de trabajo:</b> ", estilo_celda_izq),
            ]
        else:
            celda_orden_trabajo = Paragraph("<b>Orden de trabajo:</b> ", estilo_celda_izq)

        filas_bloque_derecho = [
            [Paragraph("<b>Ref. de Obra:</b> ", estilo_celda_izq), celda_orden_trabajo],
            [Paragraph(f"<b>Inmueble:</b> {_escapar_pdf(nombre or '')}", estilo_celda_izq),
             Paragraph(f"<b>Cliente:</b> {_escapar_pdf(empresa or '')}", estilo_celda_izq)],
            [Paragraph(
                '<b><span backColor="#D9D9D9" color="black">CONDICIONES DEL EDIFICIO DURANTE LA EXPOSICIÓN:</span></b>',
                estilo_celda_izq,
            ), ""],
        ]
        ancho_bloque_derecho = ancho_total - ancho_hasta_mitad_habitacion
        caja_datos_y_condiciones = Table(
            filas_bloque_derecho, colWidths=[ancho_bloque_derecho * 0.42, ancho_bloque_derecho * 0.58],
            rowHeights=[ALTURA_CABECERA * 0.32, ALTURA_CABECERA * 0.32, ALTURA_CABECERA * 0.36],
        )
        caja_datos_y_condiciones.setStyle(TableStyle([
            ("SPAN", (0, 2), (1, 2)),
            ("BOX", (0, 0), (-1, -1), 1, colors.black),
            ("BACKGROUND", (0, 0), (-1, -1), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))

        cabecera = Table(
            [[caja_logo, caja_datos_y_condiciones]],
            colWidths=[ancho_hasta_mitad_habitacion, ancho_bloque_derecho],
        )
        cabecera.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return cabecera


    # --- Tabla principal de detectores. Cabecera de 3 filas: el
    # título de la ficha (integrado en la propia tabla, sin el nombre
    # de la empresa), y debajo las dos filas de cabeceras de columnas,
    # ahora en gris claro (como el modelo original) en vez de oscuro. ---
    estilo_titulo_tabla = ParagraphStyle(
        "TituloTabla", parent=styles["Normal"], fontSize=12, leading=15,
        alignment=TA_CENTER, fontName="Helvetica-Bold", textColor=colors.black,
    )
    fila_titulo = [Paragraph(
        "FICHA DE IDENTIFICACIÓN E INFORMACIÓN DE LOS DETECTORES DE TRAZAS", estilo_titulo_tabla,
    )] + [""] * 12

    headers_fila1 = [
        "Nº", "Código\nDetector", "Edificio", "Planta", "Habitación / Estancia",
        "Ubicado en", "Foto", "INSTALACIÓN", "", "", "DESINSTALACIÓN", "", "",
    ]
    headers_fila2 = [
        "", "", "", "", "", "", "", "Fecha", "Hora", "Técnico\nFirma", "Fecha", "Hora", "Técnico\nFirma",
    ]
    estilo_cab_pequeno = ParagraphStyle(
        "CabLabPequeno", parent=styles["Normal"], fontSize=5.8, leading=7,
        alignment=TA_CENTER, textColor=colors.black, fontName="Helvetica-Bold",
    )
    filas_tabla = [
        fila_titulo,
        [Paragraph(h.replace("\n", "<br/>"), estilo_cab_pequeno if h in ("Habitación / Estancia", "Planta") else estilo_cab)
         for h in headers_fila1],
        [Paragraph(h.replace("\n", "<br/>"), estilo_cab) for h in headers_fila2],
    ]

    # Suficientes filas en blanco como para llegar al final de la hoja
    # aunque haya pocos detectores (una sola página de sobra). Con
    # firma digital, menos filas para dejar sitio de sobra al cuadro
    # de firma sin que ninguna fila se le pueda solapar.
    n_filas_totales = max(len(detectores) + 3, 14 if tipo_firma == "digital" else 15)
    estilo_firma_celda = ParagraphStyle(
        "FirmaCelda", parent=styles["Normal"], fontSize=6, leading=7.5, alignment=TA_CENTER,
    )
    nombre_tecnico_pdf = _escapar_pdf(tecnico or "")
    celda_firma_tecnico = (
        Paragraph(f"{nombre_tecnico_pdf}<br/>* Ver firma digital abajo", estilo_firma_celda)
        if tipo_firma == "digital" else ""
    )

    for i in range(n_filas_totales):
        # Solo se numeran las filas que realmente tienen datos; las
        # filas en blanco de sobra no llevan número.
        numero_fila = str(i + 1) if i < len(detectores) else ""
        if i < len(detectores):
            d = detectores[i]
            did_d, _, planta_d, sala_d, fecha_d, codigo_d = d[0], d[1], d[2], d[3], d[4], d[5]
            codigo_sala_d, hora_colocacion_d, nivel_d = d[12], d[14], d[16]
            fecha_retirada_real_d, hora_retirada_real_d = d[18], d[19]
            tiene_foto = bool(d[9]) or bool(d[10])
            fila = [
                numero_fila,
                Paragraph(_escapar_pdf(codigo_d or ""), estilo_celda),
                Paragraph(_escapar_pdf(nombre or ""), estilo_celda),
                Paragraph(_escapar_pdf(_planta_desde_nivel(nivel_d)), estilo_celda),
                Paragraph(_escapar_pdf(codigo_sala_d or ""), estilo_celda),
                Paragraph(_escapar_pdf(_ubicado_en(zona, nombre)), estilo_celda),
                "\u2713" if tiene_foto else "",
                Paragraph(_escapar_pdf(fecha_d or ""), estilo_celda),
                Paragraph(_escapar_pdf(hora_colocacion_d or ""), estilo_celda),
                celda_firma_tecnico,
                Paragraph(_escapar_pdf(fecha_retirada_real_d or ""), estilo_celda),
                Paragraph(_escapar_pdf(hora_retirada_real_d or ""), estilo_celda),
                celda_firma_tecnico,
            ]
        else:
            fila = [numero_fila] + [""] * 12
        filas_tabla.append(fila)

    tabla = Table(filas_tabla, colWidths=col_widths, repeatRows=3)
    estilo_tabla = [
        ("SPAN", (0, 0), (12, 0)),
        ("BACKGROUND", (0, 0), (-1, 0), GRIS_CLARO),
        ("BACKGROUND", (0, 1), (-1, 2), GRIS_CLARO),
        ("SPAN", (0, 1), (0, 2)), ("SPAN", (1, 1), (1, 2)), ("SPAN", (2, 1), (2, 2)),
        ("SPAN", (3, 1), (3, 2)), ("SPAN", (4, 1), (4, 2)), ("SPAN", (5, 1), (5, 2)),
        ("SPAN", (6, 1), (6, 2)),
        ("SPAN", (7, 1), (9, 1)), ("SPAN", (10, 1), (12, 1)),
        ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 3), (0, -1), "CENTER"),
        ("ALIGN", (6, 3), (6, -1), "CENTER"),
        ("FONTSIZE", (0, 3), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, 0), 6),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("ROWBACKGROUNDS", (0, 3), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
    ]
    if tipo_firma == "manual":
        # Filas más altas para que las 15 lleguen igualmente hasta el
        # final de la hoja (aquí no hace falta reservar hueco para
        # ningún cuadro de firma al pie).
        estilo_tabla.append(("TOPPADDING", (0, 3), (-1, -1), 5))
        estilo_tabla.append(("BOTTOMPADDING", (0, 3), (-1, -1), 5))
    tabla.setStyle(TableStyle(estilo_tabla))
    story.append(tabla)
    story.append(Spacer(1, 0.01*cm))

    # --- Cabecera completa (logo + datos + condiciones) y cuadro de
    # firma digital: se dibujan en cada página con onFirstPage/
    # onLaterPages, el mecanismo normal de ReportLab. Por una
    # peculiaridad interna de ReportLab, este mecanismo deja de
    # dibujar en la ÚLTIMA página del documento cuando el contenido
    # ocupa más de una hoja; para compensarlo, después de generar el
    # PDF se comprueba y, si hace falta, se "parchea" esa última
    # página superponiéndole el cuadro que falta. ---
    def _dibujar_elementos_pagina(canvas_obj, doc_obj):
        cabecera_pagina = _construir_caja_cabecera()
        x_cabecera = doc_obj.leftMargin + 6
        y_cabecera = doc_obj.pagesize[1] - 1.2*cm - ALTURA_CABECERA
        cabecera_pagina.wrapOn(canvas_obj, ancho_total, ALTURA_CABECERA)
        cabecera_pagina.drawOn(canvas_obj, x_cabecera, y_cabecera)

        if tipo_firma == "digital":
            canvas_obj.saveState()
            ancho_caja = 4.0*cm
            alto_caja = 1.5*cm
            separacion_tabla = 0.7*cm
            x = doc_obj.leftMargin + 6 + ancho_total - ancho_caja
            y = doc_obj.bottomMargin - separacion_tabla - alto_caja
            canvas_obj.setLineWidth(0.8)
            canvas_obj.rect(x, y, ancho_caja, alto_caja)
            canvas_obj.setFont("Helvetica", 6.5)
            canvas_obj.drawCentredString(x + ancho_caja / 2, y + 5, "* Firma digital")
            canvas_obj.restoreState()

    doc.build(story, onFirstPage=_dibujar_elementos_pagina, onLaterPages=_dibujar_elementos_pagina)

    # --- Parche de la última página: por la citada peculiaridad de
    # ReportLab, si el documento tiene más de una página, la última no
    # recibe el cuadro de firma digital pese a llamarse correctamente
    # a onLaterPages para ella. Se detecta y se corrige aquí,
    # superponiendo justo ese cuadro sobre la última página. ---
    if tipo_firma == "digital":
        from pypdf import PdfReader, PdfWriter
        from reportlab.pdfgen.canvas import Canvas as _CanvasBase
        import io as _io

        buf_overlay = _io.BytesIO()
        c_overlay = _CanvasBase(buf_overlay, pagesize=landscape(A4))
        ancho_caja = 4.0*cm
        alto_caja = 1.5*cm
        separacion_tabla = 0.7*cm
        x_o = 1.2*cm + 6 + ancho_total - ancho_caja
        y_o = margen_inferior - separacion_tabla - alto_caja
        c_overlay.setLineWidth(0.8)
        c_overlay.rect(x_o, y_o, ancho_caja, alto_caja)
        c_overlay.setFont("Helvetica", 6.5)
        c_overlay.drawCentredString(x_o + ancho_caja / 2, y_o + 5, "* Firma digital")
        c_overlay.save()
        buf_overlay.seek(0)

        lector_principal = PdfReader(output_path)
        pagina_overlay = PdfReader(buf_overlay).pages[0]
        escritor = PdfWriter()
        for i, pagina in enumerate(lector_principal.pages):
            if i == len(lector_principal.pages) - 1:
                pagina.merge_page(pagina_overlay)
            escritor.add_page(pagina)
        with open(output_path, "wb") as f:
            escritor.write(f)





# ============================================================
# WIDGETS DE IMAGEN (Subir archivo / Cámara del navegador)
# Sustituyen a los botones "Seleccionar" / "Camara" de la app Tkinter.
# ============================================================

# Solo se permite un flujo de cámara abierto a la vez en TODO el
# formulario (plano, foto de situación, foto del detector, imagen del
# centro). Guarda qué campo "posee" la cámara ahora mismo; el resto no
# monta su propio st.camera_input hasta que ese campo la libere. Esto
# evita el error "no autorizado" que da el móvil cuando varios widgets
# intentan acceder a la cámara al mismo tiempo.
GLOBAL_CAM_OWNER_KEY = "_camara_activa_global"



def _crear_componente_foto_pantalla_completa():
    """Cámara local para fotos: visor completo y cámara trasera por defecto."""
    html = """
    <div class="photo-camera">
      <button class="open-photo" type="button">📷 Abrir cámara</button>
      <div class="photo-overlay" hidden>
        <video autoplay muted playsinline></video>
        <div class="photo-actions">
          <button class="cancel-photo" type="button">✕ Cancelar</button>
          <button class="take-photo" type="button">● Hacer foto</button>
        </div>
      </div>
    </div>
    """
    css = """
    .photo-camera{width:100%}
    .open-photo{width:100%;background:#111;color:#f5ad22;border:1px solid #f5ad22;
      border-radius:10px;padding:.65rem .85rem;font-size:1rem;cursor:pointer}
    .photo-overlay{position:fixed;inset:0;z-index:2147483000;background:#000;
      width:100vw;height:100vh;overflow:hidden}
    .photo-overlay video{position:absolute;inset:0;width:100%;height:100%;
      object-fit:cover;background:#000}
    .photo-actions{position:absolute;left:0;right:0;bottom:max(22px,env(safe-area-inset-bottom));
      z-index:2;display:flex;justify-content:center;gap:18px;padding:12px}
    .photo-actions button{border-radius:999px;padding:.8rem 1.1rem;font-size:1rem;
      font-weight:700;cursor:pointer;border:2px solid rgba(255,255,255,.9);
      box-shadow:0 2px 12px rgba(0,0,0,.45)}
    .cancel-photo{background:rgba(20,20,20,.72);color:#fff}
    .take-photo{background:#fff;color:#111}
    """
    js = r"""
    export default function(component){
      const {parentElement,setTriggerValue}=component;
      const root=parentElement.querySelector('.photo-camera');
      const openBtn=root.querySelector('.open-photo');
      const overlay=root.querySelector('.photo-overlay');
      const video=root.querySelector('video');
      const takeBtn=root.querySelector('.take-photo');
      const cancelBtn=root.querySelector('.cancel-photo');
      let stream=null;

      const stop=()=>{
        if(stream){stream.getTracks().forEach(t=>t.stop());stream=null;}
        video.srcObject=null; overlay.hidden=true; openBtn.disabled=false;
      };

      const start=async()=>{
        openBtn.disabled=true;
        try{
          try{
            stream=await navigator.mediaDevices.getUserMedia({
              video:{facingMode:{exact:'environment'},width:{ideal:1920},height:{ideal:1080}},
              audio:false});
          }catch(_){
            stream=await navigator.mediaDevices.getUserMedia({
              video:{facingMode:{ideal:'environment'},width:{ideal:1920},height:{ideal:1080}},
              audio:false});
          }
          video.srcObject=stream; overlay.hidden=false; await video.play();
        }catch(err){
          stop();
          setTriggerValue('camera_error',String(err&&err.message?err.message:err));
        }
      };

      const capture=()=>{
        if(!stream||video.videoWidth<2||video.videoHeight<2)return;
        const maxSide=1600;
        let w=video.videoWidth,h=video.videoHeight;
        const scale=Math.min(1,maxSide/Math.max(w,h));
        w=Math.max(1,Math.round(w*scale)); h=Math.max(1,Math.round(h*scale));
        const canvas=document.createElement('canvas');
        canvas.width=w; canvas.height=h;
        canvas.getContext('2d',{alpha:false}).drawImage(video,0,0,w,h);
        const dataUrl=canvas.toDataURL('image/jpeg',0.86);
        stop(); setTriggerValue('captured',dataUrl);
      };

      openBtn.onclick=start; takeBtn.onclick=capture; cancelBtn.onclick=stop;
      return()=>stop();
    }
    """
    return st.components.v2.component(
        "radon_fullscreen_photo_camera", html=html, css=css, js=js, isolate_styles=True
    )


@st.cache_resource(show_spinner=False)
def _componente_foto_pantalla_completa():
    return _crear_componente_foto_pantalla_completa()


def _captura_foto_pantalla_completa(key_prefix):
    try:
        componente = _componente_foto_pantalla_completa()
    except Exception:
        st.error("Esta versión requiere Streamlit 1.62 o superior para la cámara a pantalla completa.")
        return None

    result = componente(
        key=key_prefix + "_fullscreen_camera",
        on_captured_change=lambda: None,
        on_camera_error_change=lambda: None,
    )
    if result is None:
        return None

    error = getattr(result, "camera_error", None)
    if error:
        firma = "error:" + str(error)
        k = key_prefix + "__fullscreen_error_consumido"
        if st.session_state.get(k) != firma:
            st.session_state[k] = firma
            st.warning(f"No se pudo abrir la cámara: {error}")
        return None

    captured = getattr(result, "captured", None)
    if not captured or not isinstance(captured, str):
        return None
    try:
        datos = base64.b64decode(captured.split(",", 1)[1] if "," in captured else captured)
        import hashlib as _hashlib
        firma = _hashlib.sha1(datos).hexdigest()
        k = key_prefix + "__fullscreen_captura_consumida"
        if st.session_state.get(k) == firma:
            return None
        st.session_state[k] = firma
        return datos
    except Exception:
        st.warning("No se pudo procesar la fotografía.")
        return None


def widget_imagen(label, state_key, key_prefix, con_camara=True, ancho_miniatura=None,
                   tab_por_defecto="subir", titulo_amarillo=False):
    """Selector de imagen con pestañas 'Subir' y 'Cámara'.

    Guarda automáticamente el archivo elegido en la carpeta de datos y
    recuerda la ruta en st.session_state[state_key]. Devuelve esa ruta
    (o None si no hay imagen). Para quitar una imagen ya subida basta
    con pulsar la "x" del propio archivo en el selector, o subir/sacar
    una nueva foto que la sustituya.

    tab_por_defecto: "subir" (normal) o "camara" (la pestaña de
    cámara aparece primero y por tanto activa por defecto al abrir
    el formulario).
    """
    if titulo_amarillo:
        st.markdown(f'<p class="subtitulo-amarillo">{label}</p>', unsafe_allow_html=True)
    else:
        st.markdown(f"**{label}**")
    file_id_key = key_prefix + "__file_id"
    cam_nonce_key = key_prefix + "__cam_nonce"
    cam_activa_key = key_prefix + "__cam_activa"
    if cam_nonce_key not in st.session_state:
        st.session_state[cam_nonce_key] = 0
    if cam_activa_key not in st.session_state:
        st.session_state[cam_activa_key] = False  # la cámara NO se activa sola

    camara_primero = con_camara and tab_por_defecto == "camara"
    if camara_primero:
        tab_labels = ["📷 Cámara", "📁 Subir desde archivo"]
        idx_subir, idx_camara = 1, 0
    else:
        tab_labels = ["📁 Subir desde archivo"] + (["📷 Cámara"] if con_camara else [])
        idx_subir, idx_camara = 0, 1
    tabs = st.tabs(tab_labels)
    nuevo_bytes = None
    nueva_ext = ".jpg"
    vino_de_camara = False

    with tabs[idx_subir]:
        up = st.file_uploader(
            "Selecciona una imagen", type=["png", "jpg", "jpeg"],
            key=key_prefix + "_up", label_visibility="collapsed",
        )
        if up is not None:
            fid = ("up", getattr(up, "file_id", None) or f"{up.name}_{up.size}")
            if st.session_state.get(file_id_key) != fid:
                nuevo_bytes = up.getvalue()
                nueva_ext = extension_de(up)
                st.session_state[file_id_key] = fid

    if con_camara:
        with tabs[idx_camara]:
            # Cámara local a pantalla completa. El lector de códigos no se toca.
            foto_bytes = _captura_foto_pantalla_completa(key_prefix)
            if foto_bytes is not None:
                nuevo_bytes = foto_bytes
                nueva_ext = ".jpg"
                vino_de_camara = True

    if nuevo_bytes is not None:
        path = guardar_bytes_imagen(nuevo_bytes, key_prefix, nueva_ext)
        st.session_state[state_key] = path
        # Marca para que quien llame a este widget sepa que ACABA de
        # capturarse/subirse una foto nueva en esta misma interacción
        # (se usa, por ejemplo, para fechar automáticamente la
        # colocación en el momento de sacar la foto del detector).
        st.session_state[key_prefix + "__recien_capturada"] = True
        if vino_de_camara:
            # El stream se cierra en el navegador inmediatamente al capturar.
            st.rerun()

    path_actual = st.session_state.get(state_key)
    if path_actual and os.path.exists(path_actual):
        if ancho_miniatura:
            st.image(path_actual, width=ancho_miniatura)
        else:
            st.image(path_actual, use_container_width=True)
    else:
        st.caption("Sin imagen")

    return st.session_state.get(state_key)


def widget_seleccionar_plano_y_punto(cid, ns):
    px_key = ns + "_plano_px"
    py_key = ns + "_plano_py"
    sel_key = ns + "_plano_centro_id"
    selector_widget_key = ns + "_plano_selector_id"

    st.markdown('<p class="subtitulo-amarillo">Plano</p>', unsafe_allow_html=True)

    planos = fetch_planos_centro(cid)
    if not planos:
        st.caption(
            "Este centro todavía no tiene ningún plano cargado. "
            "Añade uno desde la pantalla del centro (sección «🗺️ Planos del centro»)."
        )
        st.session_state[sel_key] = None
        return None, None, None

    ids = [p[0] for p in planos]
    nombres = {p[0]: p[2] for p in planos}
    rutas = {p[0]: p[3] for p in planos}

    actual = st.session_state.get(sel_key)
    if actual not in ids:
        actual = ids[0]
        st.session_state[sel_key] = actual

    # Con varios planos usamos el ID real del plano como valor del selector.
    # Así evitamos que una posición 0/1 quede asociada al plano equivocado.
    if len(planos) > 1:
        if st.session_state.get(selector_widget_key) not in ids:
            st.session_state[selector_widget_key] = actual

        nuevo_id = st.selectbox(
            "Selecciona el plano",
            options=ids,
            format_func=lambda pid: nombres.get(pid, f"Plano {pid}"),
            key=selector_widget_key,
        )

        if nuevo_id != actual:
            st.session_state[sel_key] = nuevo_id
            st.session_state[px_key] = None
            st.session_state[py_key] = None
            st.rerun()
    else:
        st.session_state[selector_widget_key] = actual
        st.caption(f"Plano: {nombres[actual]}")

    plano_id_actual = st.session_state[sel_key]
    plano_path = rutas.get(plano_id_actual)

    if not (plano_path and os.path.exists(plano_path)):
        st.warning("No se pudo abrir la imagen de este plano.")
        return plano_id_actual, st.session_state.get(px_key), st.session_state.get(py_key)

    try:
        img_orig = Image.open(plano_path).convert("RGB")
    except Exception:
        st.warning("No se pudo abrir la imagen del plano.")
        return plano_id_actual, st.session_state.get(px_key), st.session_state.get(py_key)

    ancho_max = 680
    escala = min(1.0, ancho_max / img_orig.width)
    disp_w = max(1, int(img_orig.width * escala))
    disp_h = max(1, int(img_orig.height * escala))
    img_disp = img_orig.resize((disp_w, disp_h), Image.Resampling.LANCZOS)

    px = st.session_state.get(px_key)
    py = st.session_state.get(py_key)

    # Vista estática siempre disponible. El plano deja de depender del iframe
    # interactivo para poder verse dentro del expander.
    img_preview = img_disp.copy()
    if px is not None and py is not None:
        draw = ImageDraw.Draw(img_preview)
        cx, cy = px * disp_w, py * disp_h
        r = max(6, int(min(disp_w, disp_h) * 0.02))
        draw.ellipse(
            [cx - r, cy - r, cx + r, cy + r],
            fill=(220, 20, 20),
            outline=(120, 0, 0),
            width=2,
        )

    # El plano vuelve a ser directamente interactivo: basta con tocarlo.
    # Mantenemos la selección por ID real del plano para evitar desincronizaciones
    # cuando el centro tiene dos o más planos.
    if IMG_COORD_DISPONIBLE:
        st.caption("Toca sobre el plano para marcar la ubicación del detector")
        coords = streamlit_image_coordinates(
            img_preview,
            # En móvil el componente reduce automáticamente el plano hasta el
            # ancho disponible de la pantalla. En escritorio conserva su tamaño
            # natural (limitado arriba a 680 px), evitando desplazamientos
            # horizontales cuando el plano es grande.
            use_column_width="auto",
            key=(
                ns + "_plano_coords_" + str(plano_id_actual)
                + "_" + str(st.session_state.get(f"_plano_render_nonce_{cid}", 0))
            ),
        )

        if coords is not None:
            nuevo_px = max(0.0, min(1.0, coords["x"] / disp_w))
            nuevo_py = max(0.0, min(1.0, coords["y"] / disp_h))

            if (px, py) != (nuevo_px, nuevo_py):
                st.session_state[px_key] = nuevo_px
                st.session_state[py_key] = nuevo_py
                st.rerun()
    else:
        st.image(img_preview, use_container_width=True)
        st.warning(
            "Para marcar el punto sobre el plano instala el componente:\n\n"
            "`pip install streamlit-image-coordinates`"
        )

    return plano_id_actual, st.session_state.get(px_key), st.session_state.get(py_key)


def boton_compartir_whatsapp_archivo(ruta, nombre_archivo, mime_type, texto_mensaje,
                                      etiqueta_boton, id_sufijo, titulo_compartir,
                                      color_boton="#25D366"):
    """
    Comparte un archivo usando Web Share API desde el contexto principal
    de Streamlit, para evitar el bloqueo de Chrome Android al compartir
    desde el iframe de components.html.
    """
    with open(ruta, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    safe_id = re.sub(r"[^A-Za-z0-9_]", "_", str(id_sufijo))
    datos_json = json.dumps({
        "b64": b64,
        "name": nombre_archivo,
        "type": mime_type,
        "text": texto_mensaje,
        "title": titulo_compartir,
    }, ensure_ascii=False)
    etiqueta_json = json.dumps(etiqueta_boton, ensure_ascii=False)
    color_json = json.dumps(color_boton)

    html = f"""
    <div id="wrap-compartir-{safe_id}"></div>
    <script>
    (() => {{
      const frame = window.frameElement;
      const parentWin = (window.parent && window.parent !== window) ? window.parent : window;
      const parentDoc = parentWin.document;
      const dataKey = "__radonShareData_{safe_id}";
      const handlerKey = "__radonShareHandler_{safe_id}";
      const containerId = "inyectado-{safe_id}";
      const msgId = "msg-{safe_id}";

      parentWin[dataKey] = {datos_json};

      const anterior = parentDoc.getElementById(containerId);
      if (anterior) anterior.remove();

      const cont = parentDoc.createElement("div");
      cont.id = containerId;
      cont.style.fontFamily = "'Source Sans Pro', sans-serif";

      const boton = parentDoc.createElement("button");
      boton.type = "button";
      boton.textContent = {etiqueta_json};
      boton.style.cssText =
        "background-color:" + {color_json} + " !important;" +
        "color:white !important;border:none !important;" +
        "padding:12px 20px !important;border-radius:8px !important;" +
        "font-size:16px !important;font-weight:600 !important;" +
        "cursor:pointer !important;width:100% !important;";

      const msg = parentDoc.createElement("p");
      msg.id = msgId;
      msg.style.cssText = "margin-top:8px;font-size:13px;color:#666;";

      cont.appendChild(boton);
      cont.appendChild(msg);

      let destino = document.getElementById("wrap-compartir-{safe_id}");
      if (frame && frame.parentNode && parentDoc !== document) {{
        frame.style.display = "none";
        destino = frame.parentNode;
      }}
      destino.appendChild(cont);

      const codigo = `
        return async function () {{
          const d = window["__radonShareData_{safe_id}"];
          const msg = document.getElementById("msg-{safe_id}");
          try {{
            if (msg) msg.textContent = "Abriendo menú para compartir…";
            if (!d) throw new Error("No se encontraron los datos del archivo.");

            const binary = atob(d.b64);
            const bytes = new Uint8Array(binary.length);
            for (let i = 0; i < binary.length; i++) {{
              bytes[i] = binary.charCodeAt(i);
            }}

            const file = new File([bytes], d.name, {{ type: d.type }});

            if (!navigator.share) {{
              if (msg) msg.textContent =
                "Este navegador no dispone de la función de compartir archivos.";
              return;
            }}

            if (navigator.canShare && navigator.canShare({{ files: [file] }})) {{
              await navigator.share({{
                files: [file],
                title: d.title,
                text: d.text
              }});
              if (msg) msg.textContent = "";
            }} else {{
              const blob = new Blob([bytes], {{ type: d.type }});
              const url = URL.createObjectURL(blob);
              const a = document.createElement("a");
              a.href = url;
              a.download = d.name;
              a.style.display = "none";
              document.body.appendChild(a);
              a.click();
              a.remove();
              setTimeout(() => URL.revokeObjectURL(url), 60000);

              if (msg) {{
                msg.innerHTML = "";
                const aviso = document.createElement("span");
                aviso.textContent = "Archivo guardado en Descargas. ";
                msg.appendChild(aviso);
                const abrir = document.createElement("a");
                abrir.href = "https://wa.me/?text=" + encodeURIComponent(d.text || "");
                abrir.target = "_blank";
                abrir.rel = "noopener noreferrer";
                abrir.textContent = "Abrir WhatsApp";
                abrir.style.fontWeight = "600";
                msg.appendChild(abrir);
              }}
            }}
          }} catch (err) {{
            if (err && err.name === "AbortError") return;
            if (msg) {{
              if (err && err.name === "NotAllowedError") {{
                // Fallback para Chrome Android: guardar el archivo localmente
                // y ofrecer abrir WhatsApp. El usuario podrá adjuntarlo desde
                // Descargas sin depender de la Web Share API del iframe.
                try {{
                  const binary = atob(d.b64);
                  const bytes = new Uint8Array(binary.length);
                  for (let i = 0; i < binary.length; i++) {{
                    bytes[i] = binary.charCodeAt(i);
                  }}
                  const blob = new Blob([bytes], {{ type: d.type }});
                  const url = URL.createObjectURL(blob);
                  const a = document.createElement("a");
                  a.href = url;
                  a.download = d.name;
                  a.style.display = "none";
                  document.body.appendChild(a);
                  a.click();
                  a.remove();
                  setTimeout(() => URL.revokeObjectURL(url), 60000);

                  msg.innerHTML = "";
                  const aviso = document.createElement("span");
                  aviso.textContent = "Archivo guardado en Descargas. ";
                  msg.appendChild(aviso);

                  const abrir = document.createElement("a");
                  abrir.href = "https://wa.me/?text=" + encodeURIComponent(d.text || "");
                  abrir.target = "_blank";
                  abrir.rel = "noopener noreferrer";
                  abrir.textContent = "Abrir WhatsApp";
                  abrir.style.fontWeight = "600";
                  msg.appendChild(abrir);
                }} catch (fallbackErr) {{
                  msg.textContent =
                    "Chrome no permite compartir el archivo. Utiliza el botón de descarga.";
                }}
              }} else {{
                msg.textContent = "No se pudo abrir Compartir: " +
                  (err && err.message ? err.message : err);
              }}
            }}
          }}
        }}
      `;

      try {{
        parentWin[handlerKey] = parentWin.Function(codigo)();
        boton.addEventListener("click", parentWin[handlerKey], {{ passive: false }});
      }} catch (e) {{
        msg.textContent =
          "No se pudo activar el envío directo en este navegador. Puedes usar el botón de descarga.";
      }}
    }})();
    </script>
    """
    components.html(html, height=110)

def boton_compartir_whatsapp(pdf_path, nombre_archivo, texto_mensaje, id_sufijo="pdf",
                              etiqueta_boton="Enviar por WhatsApp"):
    """Botón que comparte un PDF por WhatsApp (informe o registro para
    laboratorio). id_sufijo debe ser distinto entre botones que puedan
    coexistir en la misma pantalla, o se pisan entre sí."""
    boton_compartir_whatsapp_archivo(
        pdf_path, nombre_archivo, "application/pdf", texto_mensaje,
        etiqueta_boton, id_sufijo, "Informe de detectores de Rn",
    )


def boton_compartir_whatsapp_excel(xlsx_path, nombre_archivo, texto_mensaje):
    """Botón que comparte la hoja de cálculo del informe por WhatsApp."""
    boton_compartir_whatsapp_archivo(
        xlsx_path, nombre_archivo,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        texto_mensaje, "Enviar por WhatsApp", "xlsx",
        "Hoja de cálculo de detectores de Rn",
    )


def boton_compartir_whatsapp_fotos(imagenes, texto_mensaje, etiqueta_boton="Enviar por WhatsApp",
                                    id_sufijo="fotos-wa", titulo_compartir="Fotos del informe de detectores de Rn",
                                    color_boton="#25D366"):
    """
    Intenta compartir las fotos individualmente. Si el navegador/WhatsApp no
    admite compartirlas todas juntas, prepara un ZIP y aplica el mismo fallback
    que para Excel/PDF: descarga del ZIP + enlace para abrir WhatsApp.
    """
    archivos = []
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for ruta, nombre in imagenes:
            if not ruta or not os.path.exists(ruta):
                continue

            try:
                # Para el envío directo por WhatsApp seguimos usando una versión
                # JPEG optimizada para mejorar la compatibilidad en móvil.
                with Image.open(ruta) as im_orig:
                    im = ImageOps.exif_transpose(im_orig)
                    im = im.convert("RGB")
                    im.thumbnail((1280, 1280))
                    buf = io.BytesIO()
                    im.save(buf, format="JPEG", quality=78)
                    datos_compartir = buf.getvalue()
                nombre_compartir = os.path.splitext(nombre)[0] + ".jpg"
                mime_compartir = "image/jpeg"
            except Exception:
                with open(ruta, "rb") as f:
                    datos_compartir = f.read()
                nombre_compartir = nombre
                ext = os.path.splitext(nombre)[1].lower()
                mime_compartir = "image/png" if ext == ".png" else "image/jpeg"

            archivos.append({
                "b64": base64.b64encode(datos_compartir).decode("utf-8"),
                "name": nombre_compartir,
                "type": mime_compartir,
            })

            # El ZIP conserva el archivo original, con su nombre original.
            try:
                with open(ruta, "rb") as f:
                    zf.writestr(nombre, f.read())
            except Exception:
                zf.writestr(nombre_compartir, datos_compartir)

    if not archivos:
        st.caption("Marca al menos una foto arriba para poder enviarla.")
        return

    zip_bytes = zip_buffer.getvalue()
    zip_name = f"Fotos_detectores_{_ahora_espana().strftime('%Y%m%d_%H%M%S')}.zip"

    safe_id = re.sub(r"[^A-Za-z0-9_]", "_", str(id_sufijo))
    etiqueta_final = etiqueta_boton.format(n=len(archivos))
    datos_json = json.dumps({
        "files": archivos,
        "zip_b64": base64.b64encode(zip_bytes).decode("utf-8"),
        "zip_name": zip_name,
        "zip_type": "application/zip",
        "text": texto_mensaje,
        "title": titulo_compartir,
    }, ensure_ascii=False)
    etiqueta_json = json.dumps(etiqueta_final, ensure_ascii=False)
    color_json = json.dumps(color_boton)

    html = f"""
    <div id="wrap-compartir-{safe_id}"></div>
    <script>
    (() => {{
      const frame = window.frameElement;
      const parentWin = (window.parent && window.parent !== window) ? window.parent : window;
      const parentDoc = parentWin.document;
      const dataKey = "__radonShareFotosData_{safe_id}";
      const handlerKey = "__radonShareFotosHandler_{safe_id}";
      const containerId = "inyectado-{safe_id}";
      const msgId = "msg-{safe_id}";

      parentWin[dataKey] = {datos_json};

      const anterior = parentDoc.getElementById(containerId);
      if (anterior) anterior.remove();

      const cont = parentDoc.createElement("div");
      cont.id = containerId;
      cont.style.fontFamily = "'Source Sans Pro', sans-serif";

      const boton = parentDoc.createElement("button");
      boton.type = "button";
      boton.textContent = {etiqueta_json};
      boton.style.cssText =
        "background-color:" + {color_json} + " !important;" +
        "color:white !important;border:none !important;" +
        "padding:12px 20px !important;border-radius:8px !important;" +
        "font-size:16px !important;font-weight:600 !important;" +
        "cursor:pointer !important;width:100% !important;";

      const msg = parentDoc.createElement("p");
      msg.id = msgId;
      msg.style.cssText = "margin-top:8px;font-size:13px;color:#666;";

      cont.appendChild(boton);
      cont.appendChild(msg);

      let destino = document.getElementById("wrap-compartir-{safe_id}");
      if (frame && frame.parentNode && parentDoc !== document) {{
        frame.style.display = "none";
        destino = frame.parentNode;
      }}
      destino.appendChild(cont);

      const codigo = `
        return async function () {{
          const d = window["__radonShareFotosData_{safe_id}"];
          const msg = document.getElementById("msg-{safe_id}");

          function bytesFromB64(b64) {{
            const binary = atob(b64);
            const bytes = new Uint8Array(binary.length);
            for (let i = 0; i < binary.length; i++) {{
              bytes[i] = binary.charCodeAt(i);
            }}
            return bytes;
          }}

          function descargarZipYAbrirWhatsapp() {{
            const zipBytes = bytesFromB64(d.zip_b64);
            const blob = new Blob([zipBytes], {{ type: d.zip_type }});
            const url = URL.createObjectURL(blob);

            const a = document.createElement("a");
            a.href = url;
            a.download = d.zip_name;
            a.style.display = "none";
            document.body.appendChild(a);
            a.click();
            a.remove();
            setTimeout(() => URL.revokeObjectURL(url), 60000);

            if (msg) {{
              msg.innerHTML = "";
              const aviso = document.createElement("span");
              aviso.textContent = "No se pudieron compartir todas las fotos directamente. ZIP guardado en Descargas. ";
              msg.style.setProperty("color", "white", "important");
              aviso.style.setProperty("color", "white", "important");
              msg.appendChild(aviso);

              const abrir = document.createElement("a");
              abrir.href = "https://wa.me/?text=" + encodeURIComponent(d.text || "");
              abrir.target = "_blank";
              abrir.rel = "noopener noreferrer";
              abrir.textContent = "Abrir WhatsApp";
              abrir.style.fontWeight = "600";
              msg.appendChild(abrir);
            }}
          }}

          try {{
            if (msg) msg.textContent = "Abriendo menú para compartir…";
            if (!d) throw new Error("No se encontraron las fotos.");

            const files = d.files.map(a => {{
              const bytes = bytesFromB64(a.b64);
              return new File([bytes], a.name, {{ type: a.type }});
            }});

            if (!navigator.share) {{
              descargarZipYAbrirWhatsapp();
              return;
            }}

            if (navigator.canShare && navigator.canShare({{ files }})) {{
              try {{
                // Primera opción: fotos individuales, sin ZIP.
                await navigator.share({{
                  files: files,
                  title: d.title,
                  text: d.text
                }});
                if (msg) msg.textContent = "";
                return;
              }} catch (shareErr) {{
                if (shareErr && shareErr.name === "AbortError") return;

                // Si la cantidad de fotos, WhatsApp o Chrome impiden compartirlas
                // juntas, pasamos automáticamente al ZIP.
                if (shareErr && shareErr.name !== "NotAllowedError") {{
                  try {{
                    const zipFile = new File(
                      [bytesFromB64(d.zip_b64)],
                      d.zip_name,
                      {{ type: d.zip_type }}
                    );
                    if (!navigator.canShare || navigator.canShare({{ files: [zipFile] }})) {{
                      await navigator.share({{
                        files: [zipFile],
                        title: d.title,
                        text: d.text
                      }});
                      if (msg) msg.textContent = "";
                      return;
                    }}
                  }} catch (zipShareErr) {{
                    if (zipShareErr && zipShareErr.name === "AbortError") return;
                  }}
                }}

                descargarZipYAbrirWhatsapp();
                return;
              }}
            }}

            // Demasiadas fotos o formato no admitido: ZIP.
            try {{
              const zipFile = new File(
                [bytesFromB64(d.zip_b64)],
                d.zip_name,
                {{ type: d.zip_type }}
              );

              if (!navigator.canShare || navigator.canShare({{ files: [zipFile] }})) {{
                await navigator.share({{
                  files: [zipFile],
                  title: d.title,
                  text: d.text
                }});
                if (msg) msg.textContent = "";
                return;
              }}
            }} catch (zipErr) {{
              if (zipErr && zipErr.name === "AbortError") return;
            }}

            descargarZipYAbrirWhatsapp();

          }} catch (err) {{
            if (err && err.name === "AbortError") return;
            descargarZipYAbrirWhatsapp();
          }}
        }}
      `;

      try {{
        parentWin[handlerKey] = parentWin.Function(codigo)();
        boton.addEventListener("click", parentWin[handlerKey], {{ passive: false }});
      }} catch (e) {{
        msg.textContent =
          "No se pudo activar el envío directo en este navegador.";
      }}
    }})();
    </script>
    """
    components.html(html, height=120)



CAMPOS_DETECTOR_TRACKEADOS = (
    "_planta", "_sala", "_fecha", "_codigo", "_codigo_sala", "_profesionales_sala",
    "_hora_colocacion", "_fecha_retirada_real", "_hora_retirada_real",
    "_turno_trabajo", "_nivel", "_plano_centro_id", "_plano_px", "_plano_py",
    "_foto_sit", "_foto_det",
)


def _snapshot_detector(ns):
    """Foto de los valores actuales de un detector en session_state,
    para poder compararla más tarde y saber si hay cambios sin
    guardar."""
    return {suf: st.session_state.get(ns + suf) for suf in CAMPOS_DETECTOR_TRACKEADOS}


def _detector_tiene_cambios(ns):
    """True si los campos del detector abierto en `ns` ahora mismo son
    distintos de los que había la última vez que se guardó (o se
    cargó, si no se ha guardado nunca en esta apertura)."""
    snapshot_guardado = st.session_state.get(ns + "__snapshot")
    if snapshot_guardado is None:
        return False
    return _snapshot_detector(ns) != snapshot_guardado


def _inicializar_ns_detector(cid, detector_id, ns):
    init_key = ns + "__cargado"
    if st.session_state.get(init_key):
        return
    if detector_id:
        d = get_detector(detector_id)
        (did, _, planta, sala, fecha, codigo, _plano_antiguo, px, py, foto_sit, foto_det, _,
         codigo_sala, profesionales_sala, hora_colocacion, turno_trabajo, nivel, plano_centro_id, fecha_retirada_real, hora_retirada_real,
         _resultado_ignorado, _incertidumbre_ignorada) = d
    else:
        centro = get_centro(cid)
        planta = sala = codigo = codigo_sala = profesionales_sala = hora_colocacion = turno_trabajo = nivel = ""
        fecha = centro[3] if centro else _ahora_espana().strftime("%d/%m/%Y")
        foto_sit = foto_det = None
        plano_centro_id = None
        fecha_retirada_real = hora_retirada_real = ""
        px = py = None

    st.session_state[ns + "_planta"] = planta or ""
    st.session_state[ns + "_sala"] = sala or ""
    st.session_state[ns + "_fecha"] = fecha or _ahora_espana().strftime("%d/%m/%Y")
    st.session_state[ns + "_codigo"] = codigo or ""
    st.session_state[ns + "_codigo_sala"] = codigo_sala or ""
    st.session_state[ns + "_profesionales_sala"] = profesionales_sala or ""
    from utils_informe.excel_parser import parse_profesionales_multiples
    _lista_prof_inicial = parse_profesionales_multiples(profesionales_sala)
    st.session_state[ns + "_profesionales_lista"] = _lista_prof_inicial
    st.session_state[ns + "_hora_colocacion"] = hora_colocacion or ""
    st.session_state[ns + "_fecha_retirada_real"] = fecha_retirada_real or ""
    st.session_state[ns + "_hora_retirada_real"] = hora_retirada_real or ""
    _turnos_distintos_inicial = sorted({t for _, _, t in _lista_prof_inicial if t})
    st.session_state[ns + "_turno_trabajo"] = " / ".join(_turnos_distintos_inicial)
    st.session_state[ns + "_nivel"] = (
        nivel if nivel in NIVEL_OPCIONES else None
    )
    st.session_state[ns + "_plano_centro_id"] = plano_centro_id
    st.session_state[ns + "_plano_px"] = px if (px is not None and px >= 0) else None
    st.session_state[ns + "_plano_py"] = py if (py is not None and py >= 0) else None
    st.session_state[ns + "_foto_sit"] = foto_sit
    st.session_state[ns + "_foto_det"] = foto_det
    st.session_state[ns + "_codigo_sala_bloqueado"] = bool((codigo_sala or "").strip())
    st.session_state[ns + "__snapshot"] = _snapshot_detector(ns)
    st.session_state[init_key] = True



@st.cache_resource(show_spinner=False)
def _motor_ocr_detector():
    """Carga una sola vez el motor OCR local usado para leer códigos de detectores."""
    try:
        from rapidocr import RapidOCR
    except Exception as exc:
        raise RuntimeError(
            "Falta la dependencia 'rapidocr'. Añade rapidocr==3.9.2 a requirements.txt."
        ) from exc
    return RapidOCR()


def _normalizar_codigo_detector_ocr(texto):
    """
    Busca un código con formato exacto LLNNNN (2 letras + 4 números).
    También corrige confusiones OCR frecuentes únicamente cuando la posición
    del carácter permite saber si debe ser letra o número.
    """
    if not texto:
        return None

    limpio = re.sub(r"[^A-Z0-9]", "", str(texto).upper())

    # Primero, coincidencia exacta sin correcciones.
    m = re.search(r"[A-Z]{2}\d{4}", limpio)
    if m:
        return m.group(0)

    # Correcciones prudentes según posición.
    numero_a_letra = {
        "0": "O",
        "1": "I",
        "2": "Z",
        "5": "S",
        "6": "G",
        "8": "B",
    }
    letra_a_numero = {
        "O": "0",
        "Q": "0",
        "D": "0",
        "I": "1",
        "L": "1",
        "Z": "2",
        "S": "5",
        "B": "8",
        "G": "6",
    }

    for i in range(max(0, len(limpio) - 5)):
        trozo = limpio[i:i + 6]
        if len(trozo) != 6:
            continue

        letras = "".join(numero_a_letra.get(c, c) for c in trozo[:2])
        numeros = "".join(letra_a_numero.get(c, c) for c in trozo[2:])
        candidato = letras + numeros

        if re.fullmatch(r"[A-Z]{2}\d{4}", candidato):
            return candidato

    return None


def _extraer_textos_resultado_rapidocr(resultado):
    """Compatibilidad con las versiones actuales y anteriores de RapidOCR."""
    textos = []

    txts = getattr(resultado, "txts", None)
    if txts:
        textos.extend(str(t) for t in txts if t)

    # Compatibilidad con salidas antiguas tipo (result, elapsed).
    if not textos and isinstance(resultado, tuple) and resultado:
        primera = resultado[0]
        if isinstance(primera, (list, tuple)):
            for item in primera:
                if isinstance(item, (list, tuple)) and len(item) >= 2:
                    textos.append(str(item[1]))

    return textos


def _preparar_imagen_ocr_rapida(imagen, max_lado_objetivo, recorte_central=False):
    """
    Prepara una imagen muy ligera para OCR.
    El primer intento usa solo la zona central, que es donde normalmente
    se encuadra la etiqueta del detector al hacer la foto.
    """
    if recorte_central:
        w, h = imagen.size
        # Conserva aproximadamente el 70 % central de la imagen.
        margen_x = int(w * 0.15)
        margen_y = int(h * 0.15)
        imagen = imagen.crop((margen_x, margen_y, w - margen_x, h - margen_y))

    # Escala de grises: reduce trabajo sin perjudicar un código negro sobre etiqueta clara.
    imagen = imagen.convert("L")

    max_lado = max(imagen.size)
    if max_lado > max_lado_objetivo:
        escala = max_lado_objetivo / max_lado
        imagen = imagen.resize(
            (
                max(1, int(imagen.width * escala)),
                max(1, int(imagen.height * escala)),
            ),
            Image.Resampling.BILINEAR,
        )

    buf = io.BytesIO()
    # Calidad baja intencionadamente: para seis caracteres grandes sigue siendo
    # suficiente y reduce mucho el tamaño procesado.
    imagen.save(buf, format="JPEG", quality=62, optimize=False)
    return buf.getvalue()


def _buscar_codigo_en_resultado_ocr(resultado):
    textos = _extraer_textos_resultado_rapidocr(resultado)

    for texto in textos:
        codigo = _normalizar_codigo_detector_ocr(texto)
        if codigo:
            return codigo

    return _normalizar_codigo_detector_ocr(" ".join(textos))


def _normalizar_codigo_detector_barcode(texto):
    """Acepta directamente LLNNNN o lo extrae del contenido del código de barras."""
    if not texto:
        return None
    limpio = re.sub(r"[^A-Z0-9]", "", str(texto).upper())
    m = re.search(r"[A-Z]{2}\d{4}", limpio)
    return m.group(0) if m else None


def _leer_codigo_barras_detector(imagen):
    """
    Intenta leer primero un código de barras lineal con zxing-cpp.
    Esta ruta es mucho más ligera que ejecutar el modelo OCR.
    """
    try:
        import zxingcpp
    except Exception:
        # Si la dependencia no está instalada, se continúa con RapidOCR.
        return None

    # El código de barras de la etiqueta suele ocupar la zona central.
    w, h = imagen.size
    margen_x = int(w * 0.10)
    margen_y = int(h * 0.10)
    recorte = imagen.crop((margen_x, margen_y, w - margen_x, h - margen_y)).convert("L")

    # Para barras no hace falta conservar una fotografía grande.
    max_lado = max(recorte.size)
    if max_lado > 900:
        escala = 900 / max_lado
        recorte = recorte.resize(
            (
                max(1, int(recorte.width * escala)),
                max(1, int(recorte.height * escala)),
            ),
            Image.Resampling.BILINEAR,
        )

    # zxing-cpp trabaja directamente con la imagen/PIL y evita JPEG intermedio.
    try:
        resultados = zxingcpp.read_barcodes(recorte)
    except Exception:
        resultados = []

    for resultado in resultados:
        contenido = getattr(resultado, "text", "") or ""
        codigo = _normalizar_codigo_detector_barcode(contenido)
        if codigo:
            return codigo

    return None


def _reconocer_codigo_detector_desde_foto(foto):
    """
    Reconoce un código LLNNNN.

    Primero intenta decodificar el código de barras con zxing-cpp. Solo si
    esa lectura falla ejecuta el RapidOCR de la v57 como respaldo.
    """
    if foto is None:
        raise ValueError("Primero haz una foto de la etiqueta del detector.")

    contenido = foto.getvalue()
    imagen = Image.open(io.BytesIO(contenido)).convert("RGB")

    # PRIORIDAD 1: código de barras. Si contiene LLNNNN, terminamos aquí
    # sin cargar ni ejecutar RapidOCR.
    codigo_barras = _leer_codigo_barras_detector(imagen)
    if codigo_barras:
        return codigo_barras

    # PRIORIDAD 2: OCR actual de la v57 como respaldo.
    motor = _motor_ocr_detector()

    # 1ª pasada ultrarrápida: recorte central + 640 px.
    # Es la ruta habitual cuando la etiqueta está bien encuadrada.
    datos_rapidos = _preparar_imagen_ocr_rapida(
        imagen,
        640,
        recorte_central=True,
    )
    codigo = _buscar_codigo_en_resultado_ocr(motor(datos_rapidos))
    if codigo:
        return codigo

    # 2ª pasada: imagen completa, pero todavía pequeña.
    # Solo se ejecuta si la primera no encuentra un código válido.
    datos_detalle = _preparar_imagen_ocr_rapida(
        imagen,
        900,
        recorte_central=False,
    )
    codigo = _buscar_codigo_en_resultado_ocr(motor(datos_detalle))
    if codigo:
        return codigo

    raise ValueError(
        "No se ha reconocido un código con el formato de 2 letras y 4 números. "
        "Acerca un poco más la cámara y procura que la etiqueta esté enfocada."
    )


def _crear_componente_camara_detector():
    """
    Cámara local en navegador, sin WebRTC entre móvil y servidor.

    Requiere Streamlit >= 1.62 por st.components.v2.component.
    La cámara se abre en el propio móvil, preferentemente la trasera,
    con alta definición. Solo se analiza y, si hace falta, se envía la zona central del objetivo.
    """
    html = """
    <div class="radon-camera">
      <button class="open-btn" type="button">📷 Leer código</button>
      <div class="camera-wrap" hidden>
        <video autoplay muted playsinline></video>
        <div class="hint">Centra un solo código dentro del círculo</div>
        <div class="target"></div>
        <div class="status">Buscando código…</div>
      </div>
    </div>
    """

    css = """
    .radon-camera { width: 100%; }
    .open-btn {
      background: #111;
      color: #FFFFFF;
      border: 1px solid #f5ad22;
      border-radius: 10px;
      padding: .55rem .85rem;
      font-size: 1rem;
      cursor: pointer;
    }
    .camera-wrap {
      position: fixed;
      inset: 0;
      width: 100vw;
      height: 100vh;
      max-width: none;
      margin: 0;
      overflow: hidden;
      border-radius: 0;
      background: #000;
      z-index: 2147483000;
    }
    video {
      width: 100vw;
      height: 100vh;
      object-fit: cover;
      display: block;
    }
    .target {
      position: absolute;
      left: 50%;
      top: 50%;
      width: 56vmin;
      height: 56vmin;
      transform: translate(-50%, -50%);
      border: 3px solid rgba(255,255,255,.96);
      border-radius: 50%;
      box-sizing: border-box;
      pointer-events: none;
      box-shadow: 0 0 0 9999px rgba(0,0,0,.20);
    }
    .hint, .status {
      position: absolute;
      left: 50%;
      transform: translateX(-50%);
      background: rgba(0,0,0,.62);
      color: white;
      padding: .3rem .55rem;
      border-radius: 6px;
      white-space: nowrap;
      font-size: .85rem;
    }
    .hint { top: .5rem; }
    .status { bottom: .5rem; }
    """

    js = r"""
    export default function(component) {
      const { parentElement, setTriggerValue } = component;
      const root = parentElement.querySelector('.radon-camera');
      const openBtn = root.querySelector('.open-btn');
      const wrap = root.querySelector('.camera-wrap');
      const video = root.querySelector('video');
      const status = root.querySelector('.status');

      let stream = null;
      let stopped = false;
      let timer = null;
      let scanTimer = null;

      const stopCamera = (cerrarVisualmente = true) => {
        stopped = true;
        if (timer) {
          clearTimeout(timer);
          timer = null;
        }
        if (scanTimer) {
          clearInterval(scanTimer);
          scanTimer = null;
        }
        if (stream) {
          stream.getTracks().forEach(t => t.stop());
          stream = null;
        }

        // Cerrar también la vista de cámara, no solo el stream.
        if (cerrarVisualmente) {
          video.srcObject = null;
          wrap.hidden = true;
          openBtn.hidden = false;
          openBtn.disabled = false;
        }
      };

      const sendFrame = async () => {
        if (!stream || stopped || video.videoWidth < 2 || video.videoHeight < 2) return;

        // Segunda fase de alta definición. Solo se utiliza si el lector
        // ultrarrápido del navegador no ha reconocido el código.
        //
        // IMPORTANTE: recortamos únicamente la zona central del círculo.
        // De este modo, si hay varios detectores/códigos en la imagen, el
        // servidor recibe solo el que el usuario ha colocado dentro del visor.
        const srcW = video.videoWidth;
        const srcH = video.videoHeight;
        const cropSide = Math.round(Math.min(srcW, srcH) * 0.56);
        const sx = Math.round((srcW - cropSide) / 2);
        const sy = Math.round((srcH - cropSide) / 2);

        const maxSide = 1280;
        const outSide = Math.min(maxSide, cropSide);

        const canvas = document.createElement('canvas');
        canvas.width = outSide;
        canvas.height = outSide;
        const ctx = canvas.getContext('2d', {alpha: false});
        ctx.drawImage(
          video,
          sx, sy, cropSide, cropSide,
          0, 0, outSide, outSide
        );

        // Menos compresión para conservar las barras finas.
        const dataUrl = canvas.toDataURL('image/jpeg', 0.80);

        // Al cumplirse el tiempo máximo cerramos inmediatamente la cámara,
        // aunque después el servidor no consiga reconocer el código.
        status.textContent = 'Procesando…';
        stopCamera(true);
        setTriggerValue('captured', dataUrl);
      };

      const scanBarcodeInBrowser = async () => {
        if (!('BarcodeDetector' in window) || !stream || stopped) return;

        try {
          const detector = new BarcodeDetector({
            formats: ['code_128', 'code_39', 'ean_13', 'ean_8', 'itf', 'codabar']
          });

          // Analizamos únicamente la zona central que corresponde al círculo
          // mostrado en pantalla. Así no puede "ganar" un código situado fuera
          // del objetivo cuando hay varios detectores juntos.
          const srcW = video.videoWidth;
          const srcH = video.videoHeight;
          if (srcW < 2 || srcH < 2) return;

          const cropSide = Math.round(Math.min(srcW, srcH) * 0.56);
          const sx = Math.round((srcW - cropSide) / 2);
          const sy = Math.round((srcH - cropSide) / 2);

          const scanSide = Math.min(900, cropSide);
          const canvas = document.createElement('canvas');
          canvas.width = scanSide;
          canvas.height = scanSide;
          const ctx = canvas.getContext('2d', {alpha: false});
          ctx.drawImage(
            video,
            sx, sy, cropSide, cropSide,
            0, 0, scanSide, scanSide
          );

          const codes = await detector.detect(canvas);
          for (const item of codes) {
            const raw = String(item.rawValue || '').toUpperCase().replace(/[^A-Z0-9]/g, '');
            const m = raw.match(/[A-Z]{2}\d{4}/);
            if (m) {
              status.textContent = 'Código detectado';

              // Cierre inmediato al reconocer el código en el navegador.
              stopCamera(true);
              setTriggerValue('barcode', m[0]);
              return;
            }
          }
        } catch (_) {}
      };

      const startCamera = async () => {
        openBtn.disabled = true;
        status.textContent = 'Abriendo cámara…';

        try {
          // Primero exigimos cámara trasera. Si el dispositivo no admite
          // "exact", repetimos con "ideal".
          try {
            stream = await navigator.mediaDevices.getUserMedia({
              video: {
                facingMode: {exact: 'environment'},
                width: {ideal: 1280, max: 1920},
                height: {ideal: 960, max: 1440}
              },
              audio: false
            });
          } catch (_) {
            stream = await navigator.mediaDevices.getUserMedia({
              video: {
                facingMode: {ideal: 'environment'},
                width: {ideal: 1280, max: 1920},
                height: {ideal: 960, max: 1440}
              },
              audio: false
            });
          }

          stopped = false;
          video.srcObject = stream;
          wrap.hidden = false;
          openBtn.hidden = true;
          await video.play();
          status.textContent = 'Buscando código…';

          // Si Chrome soporta BarcodeDetector, intentamos leer localmente
          // sin subir nada al servidor.
          scanTimer = setInterval(scanBarcodeInBrowser, 180);

          // Si no se detecta un código de barras, a los 4 segundos
          // cerramos la cámara y enviamos un único fotograma pequeño a zxing/OCR.
          // Así la cámara nunca queda abierta indefinidamente.
          timer = setTimeout(sendFrame, 4000);
        } catch (err) {
          status.textContent = 'No se pudo abrir la cámara';
          openBtn.disabled = false;
          openBtn.hidden = false;
          wrap.hidden = true;
          setTriggerValue('camera_error', String(err && err.message ? err.message : err));
        }
      };

      openBtn.onclick = startCamera;

      return () => {
        stopCamera(true);
      };
    }
    """

    return st.components.v2.component(
        "radon_detector_camera",
        html=html,
        css=css,
        js=js,
        isolate_styles=True,
    )


@st.cache_resource(show_spinner=False)
def _componente_camara_detector():
    return _crear_componente_camara_detector()


def _procesar_resultado_camara_local(ns, result):
    """
    Procesa una sola vez cada resultado devuelto por la cámara.

    El componente mantiene su último valor entre reruns. Si se vuelve a
    procesar continuamente, la página entra en un bucle de reruns y parece
    bloqueada: no deja cambiar de detector ni salir. Guardamos una firma del
    último resultado consumido para evitar ese comportamiento.
    """
    if result is None:
        return False

    consumido_key = ns + "_camara_resultado_consumido"

    barcode = getattr(result, "barcode", None)
    if barcode:
        codigo = _normalizar_codigo_detector_barcode(barcode)
        if codigo:
            firma = "barcode:" + codigo
            if st.session_state.get(consumido_key) == firma:
                return False

            st.session_state[consumido_key] = firma
            st.session_state[ns + "_codigo"] = codigo
            st.session_state[ns + "_ocr_codigo_mensaje"] = (
                "ok",
                f"Código reconocido: {codigo}",
            )
            return True

    camera_error = getattr(result, "camera_error", None)
    if camera_error:
        firma = "error:" + str(camera_error)
        if st.session_state.get(consumido_key) == firma:
            return False

        st.session_state[consumido_key] = firma
        st.session_state[ns + "_ocr_codigo_mensaje"] = (
            "error",
            f"No se pudo abrir la cámara: {camera_error}",
        )
        return False

    captured = getattr(result, "captured", None)
    if not captured or not isinstance(captured, str):
        return False

    try:
        contenido_b64 = captured.split(",", 1)[1] if "," in captured else captured
        datos = base64.b64decode(contenido_b64)

        # Firma estable del fotograma. Evita volver a procesar exactamente
        # la misma captura en cada rerun de Streamlit.
        import hashlib as _hashlib
        firma = "captured:" + _hashlib.sha1(datos).hexdigest()
        if st.session_state.get(consumido_key) == firma:
            return False

        # Marcar ANTES de hacer OCR para que incluso un error no provoque
        # un bucle de procesamiento.
        st.session_state[consumido_key] = firma

        foto = io.BytesIO(datos)

        # 1) Código de barras en servidor.
        imagen = Image.open(io.BytesIO(datos)).convert("RGB")
        codigo = _leer_codigo_barras_detector(imagen)

        # 2) OCR v57 como respaldo.
        if not codigo:
            foto.seek(0)
            codigo = _reconocer_codigo_detector_desde_foto(foto)

        st.session_state[ns + "_codigo"] = codigo
        st.session_state[ns + "_ocr_codigo_mensaje"] = (
            "ok",
            f"Código reconocido: {codigo}",
        )
        return True
    except Exception as exc:
        st.session_state[ns + "_ocr_codigo_mensaje"] = ("error", str(exc))
        return False


def _renderizar_lector_camara_local(ns):
    """Muestra una cámara local, sin STUN/TURN ni selección manual de dispositivo."""
    try:
        componente = _componente_camara_detector()
    except Exception:
        st.error(
            "Esta versión requiere Streamlit 1.62 o superior. "
            "Actualiza la dependencia «streamlit» en requirements.txt."
        )
        return

    result = componente(
        key=ns + "_camara_local",
        on_captured_change=lambda: None,
        on_barcode_change=lambda: None,
        on_camera_error_change=lambda: None,
    )

    # El evento del componente ya ha provocado el rerun necesario.
    # No hacemos otro st.rerun aquí: el campo "Código del detector" se crea
    # justo después y recoge el valor recién reconocido en esta misma ejecución.
    _procesar_resultado_camara_local(ns, result)



def _renderizar_campos_detector(cid, detector_id, ns):
    c1, c2 = st.columns(2)
    with c1:
        # El campo se crea después del lector para permitir que el resultado
        # automático se escriba antes en session_state.
        codigo_placeholder = st.empty()

        _renderizar_lector_camara_local(ns)

        codigo_placeholder.text_input(
            "Código del detector",
            key=ns + "_codigo",
            help="Formato esperado: 2 letras y 4 números, por ejemplo GJ4306.",
        )

        mensaje_ocr = st.session_state.get(ns + "_ocr_codigo_mensaje")
        if mensaje_ocr:
            tipo_ocr, texto_ocr = mensaje_ocr
            if tipo_ocr == "ok":
                st.success(texto_ocr)
            else:
                st.warning(texto_ocr)

        st.text_input("Planta (opcional)", key=ns + "_planta")
    with c2:
        st.text_input("Sala", key=ns + "_sala")
        st.selectbox(
            "Nivel", options=NIVEL_OPCIONES, key=ns + "_nivel",
            index=None, placeholder="Selecciona un nivel",
        )

        codigo_sala_ya_fijado = st.session_state.get(ns + "_codigo_sala_bloqueado", False)
        if not codigo_sala_ya_fijado:
            codigo_detector_actual = st.session_state.get(ns + "_codigo", "").strip()
            sala_actual = st.session_state.get(ns + "_sala", "").strip()
            nivel_actual = st.session_state.get(ns + "_nivel")
            if codigo_detector_actual and sala_actual and nivel_actual:
                centro_de_este_detector = get_centro(cid)
                zona_centro = centro_de_este_detector[2] if centro_de_este_detector else ""
                tipo_centro_actual = get_tipo_centro(cid)
                st.session_state[ns + "_codigo_sala"] = _generar_codigo_sala(
                    cid, detector_id, nivel_actual, zona_centro, tipo_centro_actual
                )
        st.text_input("Código de la sala", key=ns + "_codigo_sala")
        if not st.session_state.get(ns + "_codigo_sala_bloqueado", False):
            st.caption(
                "Se genera automáticamente cuando están cubiertos el código del detector, "
                "el nivel y el nombre de la sala. Mientras el detector no se haya guardado, "
                "se seguirá actualizando automáticamente si modificas alguno de esos datos."
            )

    st.markdown('<div class="marcador-profesionales-sala marcador-acordeon-detector"></div>', unsafe_allow_html=True)
    with st.expander("Profesionales en esta sala", expanded=False):
        st.markdown('<div class="marcador-profesionales-sala-sentence-case"></div>', unsafe_allow_html=True)
        categorias_centro_form = fetch_categorias_centro(cid)
        opciones_categoria_form = [c[2] for c in categorias_centro_form]
        opcion_profesionales_generica = "Profesionales (sin especificar)"
        if opcion_profesionales_generica not in opciones_categoria_form:
            opciones_categoria_form.append(opcion_profesionales_generica)

        lista_key = ns + "_profesionales_lista"
        lista_actual = st.session_state.get(lista_key, [])

        editar_prof_key = f"{ns}_prof_modo_anadir"
        modo_anadir_prof = bool(st.session_state.get(editar_prof_key, False))

        if not modo_anadir_prof:
            st.markdown('<div class="marcador-btn-anadir-categoria"></div>', unsafe_allow_html=True)
            if st.button("➕ Añadir", key=f"{ns}_prof_add", use_container_width=True):
                st.session_state[editar_prof_key] = True
                st.rerun()

        if lista_actual:
            seleccionadas_prof = []
            for idx_prof, (cat_linea, num_linea, turno_linea) in enumerate(lista_actual):
                etiqueta_linea = f"{num_linea} {cat_linea}" + (f" — {turno_linea}" if turno_linea else "")
                marcado = st.checkbox(etiqueta_linea, key=f"{ns}_prof_chk_{idx_prof}")
                if marcado:
                    seleccionadas_prof.append(idx_prof)
            if seleccionadas_prof:
                st.markdown('<div class="marcador-btn-eliminar"></div>', unsafe_allow_html=True)
                if st.button(f"❌ Eliminar seleccionadas ({len(seleccionadas_prof)})", key=f"{ns}_prof_del"):
                    st.session_state[lista_key] = [
                        entrada for i, entrada in enumerate(lista_actual) if i not in seleccionadas_prof
                    ]
                    st.rerun()
        else:
            st.caption("Todavía no has añadido ningún profesional a esta sala.")

        pend_add_cat_key = f"{ns}_add_categoria_pend"
        pend_add_num_key = f"{ns}_add_num_pend"
        pend_add_turno_key = f"{ns}_add_turno_pend"
        if pend_add_cat_key in st.session_state:
            st.session_state[f"{ns}_add_categoria"] = st.session_state.pop(pend_add_cat_key)
        if pend_add_num_key in st.session_state:
            st.session_state[f"{ns}_add_num"] = st.session_state.pop(pend_add_num_key)
        if pend_add_turno_key in st.session_state:
            st.session_state[f"{ns}_add_turno"] = st.session_state.pop(pend_add_turno_key)

        if modo_anadir_prof:
            categoria_nueva = st.selectbox(
                "Categoría profesional", options=opciones_categoria_form, key=f"{ns}_add_categoria",
                index=None, placeholder="Selecciona una categoría",
            )
            turno_nuevo = st.selectbox(
                "Turno de trabajo", options=TURNOS_TRABAJO_OPCIONES, key=f"{ns}_add_turno",
                index=None, placeholder="Selecciona un turno",
            )
            num_maximo = 20
            if categoria_nueva and categoria_nueva != "Profesionales (sin especificar)":
                fila_categoria = next((c for c in categorias_centro_form if c[2] == categoria_nueva), None)
                if fila_categoria:
                    num_maximo = max(int(fila_categoria[3] or 1), 1)
            if st.session_state.get(f"{ns}_add_num") not in range(1, num_maximo + 1):
                st.session_state[f"{ns}_add_num"] = 1
            etiqueta_num = "Nº de {} en esta sala en este {}".format(
                categoria_nueva if categoria_nueva else "personas",
                turno_nuevo if turno_nuevo else "turno",
            )
            num_nuevo = st.selectbox(
                etiqueta_num, options=list(range(1, num_maximo + 1)), key=f"{ns}_add_num",
            )

            st.markdown('<div class="marcador-btn-anadir-categoria"></div>', unsafe_allow_html=True)
            if st.button("💾 Guardar", key=f"{ns}_prof_guardar", use_container_width=True):
                if not categoria_nueva:
                    st.warning("Selecciona una categoría")
                else:
                    st.session_state[lista_key] = lista_actual + [
                        (categoria_nueva, num_nuevo, turno_nuevo or "")
                    ]
                    st.session_state[pend_add_cat_key] = None
                    st.session_state[pend_add_num_key] = 1
                    st.session_state[pend_add_turno_key] = None
                    st.session_state[editar_prof_key] = False
                    st.rerun()

        lista_final = st.session_state.get(lista_key, [])
        st.session_state[ns + "_profesionales_sala"] = ", ".join(
            f"{cat} ({num})" + (f" - {turno}" if turno else "")
            for cat, num, turno in lista_final
        )
        turnos_distintos = sorted({turno for _, _, turno in lista_final if turno})
        st.session_state[ns + "_turno_trabajo"] = " / ".join(turnos_distintos)

    st.markdown('<div class="marcador-acordeon-detector"></div>', unsafe_allow_html=True)
    with st.expander("Plano de ubicación", expanded=False):
        widget_seleccionar_plano_y_punto(cid, ns)

    st.markdown('<div class="marcador-acordeon-detector"></div>', unsafe_allow_html=True)
    with st.expander("Fotos del detector", expanded=False):
        fc1, fc2 = st.columns(2)
        with fc1:
            st.markdown('<div class="marcador-imagen-exterior"></div>', unsafe_allow_html=True)
            widget_imagen("Situación del detector", ns + "_foto_sit", key_prefix=ns + "_fsit",
                          tab_por_defecto="camara", titulo_amarillo=True)
        with fc2:
            st.markdown('<div class="marcador-imagen-exterior"></div>', unsafe_allow_html=True)
            widget_imagen("Detector", ns + "_foto_det", key_prefix=ns + "_fdet",
                          tab_por_defecto="camara", titulo_amarillo=True)

    if st.session_state.get(ns + "_fdet__recien_capturada"):
        st.session_state[ns + "_fdet__recien_capturada"] = False
        ahora = _ahora_espana()
        st.session_state[ns + "_fecha"] = ahora.strftime("%d/%m/%Y")
        st.session_state[ns + "_hora_colocacion"] = ahora.strftime("%H:%M")
        st.rerun()

    st.markdown('<div class="marcador-acordeon-detector"></div>', unsafe_allow_html=True)
    with st.expander("Fecha y hora de colocación", expanded=False):
        st.caption("Se rellenan solas al hacer la foto del detector, o se pueden escribir/corregir aquí.")
        fh1, fh2 = st.columns(2)
        with fh1:
            _date_input_texto("Fecha de colocación", ns + "_fecha")
        with fh2:
            st.text_input("Hora de colocación", key=ns + "_hora_colocacion")
    # Los cuatro acordeones de la ficha del detector son mutuamente exclusivos:
    # al abrir uno se cierran automáticamente los demás. Es solo comportamiento
    # visual del navegador; no altera ni borra los valores de los widgets.
    components.html(
        """
        <script>
        (() => {
          try {
            const doc = window.parent.document;
            const key = "__radon_detector_accordion_exclusive__";
            if (window.parent[key]) return;
            window.parent[key] = true;

            const activar = () => {
              const expanders = Array.from(doc.querySelectorAll('div[data-testid="stExpander"] details'));
              expanders.forEach((details) => {
                if (details.dataset.radonExclusive === "1") return;
                details.dataset.radonExclusive = "1";
                details.addEventListener("toggle", () => {
                  if (!details.open) return;
                  expanders.forEach((otro) => {
                    if (otro !== details && otro.open) otro.open = false;
                  });
                });
              });
            };
            activar();
            const obs = new MutationObserver(() => activar());
            obs.observe(doc.body, {childList: true, subtree: true});
          } catch (e) {}
        })();
        </script>
        """,
        height=0,
    )



def _guardar_detector_desde_ns(cid, detector_id, ns, mostrar_mensajes=True):
    """Guarda en la base de datos los valores que haya ahora mismo en
    st.session_state para ese detector (sirve tanto para el guardado
    manual con botón como para el guardado automático al cambiar de
    detector o salir de la pantalla). Devuelve el id del detector
    guardado (nuevo o existente), o None si no se pudo guardar por
    faltar algún campo obligatorio."""
    if not st.session_state.get(ns + "__cargado"):
        return None  # nunca se llegó a cargar/tocar este formulario

    sala_val = st.session_state.get(ns + "_sala", "").strip()
    codigo_val = st.session_state.get(ns + "_codigo", "").strip()
    if not codigo_val:
        if mostrar_mensajes:
            st.warning("El código es obligatorio")
        return None

    px = st.session_state.get(ns + "_plano_px")
    py = st.session_state.get(ns + "_plano_py")
    data = (
        cid,
        st.session_state.get(ns + "_planta", "").strip(),
        sala_val,
        st.session_state.get(ns + "_fecha", "").strip(),
        codigo_val,
        None,
        px if px is not None else -1,
        py if py is not None else -1,
        st.session_state.get(ns + "_foto_sit"),
        st.session_state.get(ns + "_foto_det"),
        _ahora_espana().strftime("%Y-%m-%d %H:%M"),
        st.session_state.get(ns + "_codigo_sala", "").strip(),
        st.session_state.get(ns + "_profesionales_sala", "").strip(),
        st.session_state.get(ns + "_hora_colocacion", "").strip(),
        st.session_state.get(ns + "_turno_trabajo") or "",
        st.session_state.get(ns + "_nivel") or "",
        st.session_state.get(ns + "_plano_centro_id"),
        st.session_state.get(ns + "_fecha_retirada_real", ""),
        st.session_state.get(ns + "_hora_retirada_real", ""),
    )
    if detector_id:
        update_detector(detector_id, data)
        nuevo_id = detector_id
    else:
        nuevo_id = insert_detector(data)

    if mostrar_mensajes:
        st.success("Detector guardado")
    return nuevo_id


def _guardar_y_actualizar_snapshot_detector(cid, abierto, ns, mostrar_mensajes=True):
    """Guarda el detector abierto (nuevo o existente) y, si sale bien,
    actualiza la "foto" de referencia usada para saber si hay cambios
    sin guardar (con esto, justo después de guardar, ya no aparecerán
    como pendientes). Gestiona también, si aplica, la transición de
    detector "nuevo" a uno ya guardado con id real: traslada el
    estado de sesión a la nueva clave y actualiza el desplegable.
    Devuelve (guardado_ok, ns_final)."""
    detector_id_abierto = None if abierto == "nuevo" else abierto
    guardado_id = _guardar_detector_desde_ns(cid, detector_id_abierto, ns, mostrar_mensajes=mostrar_mensajes)
    if not guardado_id:
        return False, ns

    ns_final = ns
    if abierto == "nuevo":
        # Se traslada el estado "de datos" de sesión de "det_nuevo_*" a
        # "det_{id}_*" (ver explicación detallada en el histórico de
        # cambios de esta función), para no perder lo ya escrito ni
        # las fotos ya subidas al pasar a tener un id real.
        ns_nuevo_real = f"det_{guardado_id}"
        sufijos_a_trasladar = (
            "_planta", "_sala", "_fecha", "_codigo", "_codigo_sala",
            "_codigo_sala_bloqueado", "_profesionales_sala", "_profesionales_lista",
            "_hora_colocacion",
            "_fecha_retirada_real", "_hora_retirada_real", "_turno_trabajo", "_nivel",
            "_plano_centro_id", "_plano_px", "_plano_py", "_foto_sit", "_foto_det",
            "__cargado",
        )
        for _sufijo in sufijos_a_trasladar:
            _k_viejo = ns + _sufijo
            if _k_viejo in st.session_state:
                st.session_state[ns_nuevo_real + _sufijo] = st.session_state.pop(_k_viejo)
        abierto_key = f"detector_abierto_{cid}"
        st.session_state[abierto_key] = guardado_id
        st.session_state["detector_form_ns"] = ns_nuevo_real
        detectores_tras_guardar = fetch_detectores(cid)
        ids_tras_guardar = [d[0] for d in detectores_tras_guardar]
        pend_sel_key = f"selector_detector_id_pend_{cid}"
        if guardado_id in ids_tras_guardar:
            # El selector trabaja con IDs persistentes, no con posiciones.
            # Así, al añadir un detector, no queda asociado a un índice que
            # pueda representar otro detector en el siguiente rerun.
            st.session_state[pend_sel_key] = guardado_id
        ns_final = ns_nuevo_real

    st.session_state[ns_final + "__snapshot"] = _snapshot_detector(ns_final)
    return True, ns_final


def _limpiar_namespace(ns):
    borrar = [k for k in st.session_state.keys()
              if k == ns or k.startswith(ns + "_") or k.startswith(ns + "__")]
    for k in borrar:
        del st.session_state[k]
    # Si la cámara global la tenía un campo de este formulario, la
    # liberamos para que no quede "atascada" al salir de la pantalla.
    owner = st.session_state.get(GLOBAL_CAM_OWNER_KEY)
    if owner and (owner == ns or owner.startswith(ns + "_") or owner.startswith(ns + "__")):
        st.session_state[GLOBAL_CAM_OWNER_KEY] = None




# ============================================================
# CONTROL DE CAMBIOS Y DESCARGA DEL ARCHIVO DE DATOS
# ============================================================

def _firma_estado_centro(centro_id):
    """Firma del estado persistido del centro para saber si cambió desde que se abrió."""
    import hashlib
    try:
        estado = {
            "centro": get_centro(centro_id),
            "categorias": fetch_categorias_centro(centro_id),
            "planos": fetch_planos_centro(centro_id),
            "detectores": fetch_detectores(centro_id),
            "informe": get_datos_informe_centro(centro_id),
        }
        bruto = json.dumps(estado, ensure_ascii=False, sort_keys=True, default=str)
        return hashlib.sha256(bruto.encode("utf-8")).hexdigest()
    except Exception:
        # Si algo no puede serializarse, conservamos igualmente una comparación estable.
        estado = (
            get_centro(centro_id),
            fetch_categorias_centro(centro_id),
            fetch_planos_centro(centro_id),
            fetch_detectores(centro_id),
        )
        return hashlib.sha256(repr(estado).encode("utf-8")).hexdigest()


def _registrar_entrada_centro(centro_id):
    st.session_state["_centro_snapshot_id"] = centro_id
    st.session_state["_centro_snapshot_firma"] = _firma_estado_centro(centro_id)


def _centro_modificado_desde_entrada(centro_id):
    if st.session_state.get("_centro_snapshot_id") != centro_id:
        return False
    firma_inicial = st.session_state.get("_centro_snapshot_firma")
    if not firma_inicial:
        return False
    return firma_inicial != _firma_estado_centro(centro_id)


def _obtener_excel_descarga_centro(centro_id):
    """Genera el Excel y lo mantiene en memoria mientras no cambien los datos."""
    centro = get_centro(centro_id)
    nombre_centro = centro[1] if centro else "Centro"
    firma = _firma_estado_centro(centro_id)
    cache_key = f"_excel_datos_cache_{centro_id}"
    cache = st.session_state.get(cache_key)

    if isinstance(cache, dict) and cache.get("firma") == firma:
        return cache["bytes"], cache["nombre"]

    nombre_archivo = _nombre_documento(nombre_centro, "DATOS") + ".xlsx"
    ruta_tmp = os.path.join(
        get_tmp_dir(),
        f"_descarga_datos_{centro_id}_{_ahora_espana().strftime('%Y%m%d%H%M%S%f')}.xlsx",
    )
    generar_excel(centro_id, ruta_tmp)
    try:
        with open(ruta_tmp, "rb") as f:
            contenido = f.read()
    finally:
        try:
            if os.path.exists(ruta_tmp):
                os.remove(ruta_tmp)
        except Exception:
            pass

    st.session_state[cache_key] = {
        "firma": firma,
        "bytes": contenido,
        "nombre": nombre_archivo,
    }
    return contenido, nombre_archivo


def _cerrar_aviso_guardado_salida(centro_id):
    st.session_state.pop("_preguntar_guardado_centro", None)
    st.session_state["_centro_snapshot_id"] = centro_id
    st.session_state["_centro_snapshot_firma"] = _firma_estado_centro(centro_id)




def _render_titulo_principal(texto):
    """Título principal robusto, fuera del renderizado Markdown de Streamlit."""
    st.html(
        '<div style="'
        'color:#F5A623;'
        'font-size:clamp(1.3rem, 4.5vw, 2.25rem);'
        'line-height:1.10;'
        'font-weight:800;'
        'margin:0.35rem 0 0.55rem 0;'
        'word-break:normal;'
        'overflow-wrap:anywhere;'
        '">'
        + html.escape(str(texto or ""))
        + '</div>'
    )


def _zona_para_titulo_centro(centro_id, zona):
    """Oculta área/zona en títulos de centros de AP, AP+PAC, PAC y consultorios."""
    tipo = str(get_tipo_centro(centro_id) or "").strip().casefold()
    tipo_compacto = re.sub(r"\s+", "", tipo)
    solo_nombre = (
        tipo in {"atención primaria", "atencion primaria", "pac", "consultorio"}
        or tipo_compacto in {
            "atenciónprimaria+pac",
            "atencionprimaria+pac",
            "primaria+pac",
        }
    )
    return "" if solo_nombre else str(zona or "")


# ============================================================
# PANTALLA: INICIO (lista de centros)
# ============================================================

def pantalla_inicio():
    st.markdown(
        '<p style="color:#999999; font-size:0.85rem; font-weight:700; margin:0;">'
        'Gestión de muestreo</p>',
        unsafe_allow_html=True,
    )
    _render_titulo_principal("☢️ Detectores de Radón")

    # Si se vuelve desde un centro y hubo cambios desde que se abrió,
    # se pregunta si se quiere guardar el archivo de datos.
    cid_guardado_pendiente = st.session_state.get("_preguntar_guardado_centro")
    if cid_guardado_pendiente:
        centro_pendiente = get_centro(cid_guardado_pendiente)
        if centro_pendiente:
            nombre_pendiente = centro_pendiente[1] or "este centro"
            st.warning(
                f"Has modificado datos de «{nombre_pendiente}» desde que entraste. "
                "¿Quieres guardar los datos del centro?"
            )
            try:
                excel_pendiente, nombre_excel_pendiente = _obtener_excel_descarga_centro(
                    cid_guardado_pendiente
                )
                cg1, cg2 = st.columns(2)
                with cg1:
                    st.download_button(
                        "Sí, guardar datos",
                        data=excel_pendiente,
                        file_name=nombre_excel_pendiente,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        type="primary",
                        key=f"guardar_salida_{cid_guardado_pendiente}",
                        on_click=_cerrar_aviso_guardado_salida,
                        args=(cid_guardado_pendiente,),
                    )
                with cg2:
                    if st.button(
                        "No guardar",
                        use_container_width=True,
                        key=f"no_guardar_salida_{cid_guardado_pendiente}",
                    ):
                        _cerrar_aviso_guardado_salida(cid_guardado_pendiente)
                        st.rerun()
            except Exception as e:
                st.error(f"No se pudo preparar el archivo de datos: {e}")
        else:
            st.session_state.pop("_preguntar_guardado_centro", None)

    if "mostrar_form_nuevo_centro" not in st.session_state:
        st.session_state.mostrar_form_nuevo_centro = False
    if "mostrar_form_importar" not in st.session_state:
        st.session_state.mostrar_form_importar = False

    with st.container(border=True):
        bc1, bc2 = st.columns(2)
        with bc1:
            marcadores_nuevo = '<div class="marcador-btn-nuevo-centro"></div>'
            if st.session_state.mostrar_form_nuevo_centro:
                marcadores_nuevo += '<div class="marcador-activo-naranja"></div>'
            st.markdown(marcadores_nuevo, unsafe_allow_html=True)
            if st.button("➕ Nuevo centro", use_container_width=True, type="secondary"):
                abrir = not st.session_state.mostrar_form_nuevo_centro
                st.session_state.mostrar_form_nuevo_centro = abrir
                if abrir:
                    st.session_state.mostrar_form_importar = False
                st.rerun()
        with bc2:
            marcadores_importar = '<div class="marcador-btn-importar"></div>'
            if st.session_state.mostrar_form_importar:
                marcadores_importar += '<div class="marcador-activo-naranja"></div>'
            st.markdown(marcadores_importar, unsafe_allow_html=True)
            if st.button("Importar centro", use_container_width=True, type="secondary",
                         icon=":material/folder_open:"):
                abrir = not st.session_state.mostrar_form_importar
                st.session_state.mostrar_form_importar = abrir
                if abrir:
                    st.session_state.mostrar_form_nuevo_centro = False
                st.rerun()

    if st.session_state.mostrar_form_nuevo_centro:
        # Fuera de un st.form: así el desplegable "Tipo de centro"
        # puede rellenar la casilla "Área / Zona :" al momento, sin
        # esperar a un envío conjunto (los campos de un st.form no
        # reaccionan entre sí hasta que se pulsa el botón).
        nombre_nuevo = st.text_input("Nombre del centro", key="nuevo_centro_nombre")
        tipo_centro_nuevo = st.selectbox(
            "Tipo de centro", options=TIPO_CENTRO_OPCIONES, key="nuevo_centro_tipo",
            index=None, placeholder="Selecciona un tipo de centro",
        )
        area_automatica = TIPO_CENTRO_A_AREA_AUTOMATICA.get(tipo_centro_nuevo, "")
        _sincronizar_valor_auto("nuevo_centro_zona", area_automatica)
        zona_nueva = st.text_input("Área / Zona :", key="nuevo_centro_zona")
        st.markdown('<div class="marcador-btn-crear-centro"></div>', unsafe_allow_html=True)
        crear = st.button("Crear centro", type="primary", key="btn_crear_centro_nuevo")
        if crear:
            if nombre_nuevo and nombre_nuevo.strip():
                cid = crear_centro(
                    nombre_nuevo.strip(),
                    zona_nueva.strip() if zona_nueva else "",
                    tipo_centro_nuevo or "",
                )
                st.session_state.centro_actual = cid
                _registrar_entrada_centro(cid)
                st.session_state.mostrar_form_nuevo_centro = False
                for _k in ("nuevo_centro_nombre", "nuevo_centro_tipo", "nuevo_centro_zona",
                           "nuevo_centro_zona__ultimo_auto"):
                    st.session_state.pop(_k, None)
                st.session_state.view = "centro"
                st.rerun()
            else:
                st.warning("Escribe un nombre para el centro")

    if st.session_state.mostrar_form_importar:
        st.caption(
            "Reconstruye un centro completo (datos, planos y detectores, con "
            "el punto exacto de cada uno) a partir de un Excel generado por "
            "esta misma app."
        )
        st.markdown('<div class="marcador-uploader-importar"></div>', unsafe_allow_html=True)
        archivo_importar = st.file_uploader(
            "Selecciona el archivo .xlsx", type=["xlsx"], key="importar_centro_file",
        )
        if archivo_importar:
            st.markdown('<div class="marcador-btn-confirmar-importar"></div>', unsafe_allow_html=True)
            if st.button("Importar", key="btn_confirmar_importar", type="primary"):
                try:
                    with st.spinner("Importando centro..."):
                        nuevo_cid, n_detectores = importar_centro_desde_excel(archivo_importar.getvalue())
                    st.session_state.mostrar_form_importar = False
                    st.session_state.centro_actual = nuevo_cid
                    _registrar_entrada_centro(nuevo_cid)
                    st.session_state.view = "centro"
                    st.success(f"Centro importado correctamente ({n_detectores} detector(es)).")
                    st.rerun()
                except ValueError as e:
                    st.error(str(e))
                except Exception as e:
                    st.error(f"No se pudo importar el archivo: {e}")

    centros = fetch_centros()
    if not centros:
        # Si se ha eliminado el último centro, limpiamos también el
        # valor visual que pudiera conservar el desplegable.
        st.session_state.pop("selector_centro_home", None)
        st.session_state.pop("_centro_a_seleccionar_tras_borrado", None)
        return

    st.markdown("---")
    st.markdown('<p class="subtitulo-amarillo">Centros registrados</p>', unsafe_allow_html=True)

    opciones = list(range(len(centros)))
    etiquetas = {
        i: f"{c[1] or '(sin nombre)'}"
           + (f" · {c[2]}" if c[2] else "")
        for i, c in enumerate(centros)
    }
    # Tras eliminar un centro seleccionamos explícitamente el siguiente
    # centro de la lista. Si el eliminado era el último, seleccionamos
    # el anterior. Esto evita que el selectbox conserve visualmente el
    # nombre del centro ya eliminado.
    centro_objetivo = st.session_state.pop(
        "_centro_a_seleccionar_tras_borrado", None
    )
    if centro_objetivo is not None:
        nuevo_indice = next(
            (i for i, c in enumerate(centros) if c[0] == centro_objetivo),
            0,
        )
        st.session_state["selector_centro_home"] = nuevo_indice
    elif st.session_state.get("selector_centro_home") not in opciones:
        st.session_state["selector_centro_home"] = 0

    idx_sel = st.selectbox(
        "Centro", options=opciones, format_func=lambda i: etiquetas[i],
        key="selector_centro_home", label_visibility="collapsed",
    )
    cid_sel, nombre_sel, zona_sel, fecha_sel, img_sel = centros[idx_sel]

    b1, b2 = st.columns([3, 1])
    with b1:
        st.markdown('<div class="marcador-btn-abrir-centro"></div>', unsafe_allow_html=True)
        if st.button("Abrir centro", type="primary", use_container_width=True,
                     icon=":material/folder_open:"):
            st.session_state.centro_actual = cid_sel
            _registrar_entrada_centro(cid_sel)
            st.session_state.view = "centro"
            st.rerun()
    with b2:
        st.markdown('<div class="marcador-btn-eliminar"></div>', unsafe_allow_html=True)
        if st.button("❌ Eliminar", key="btn_eliminar_centro_home"):
            st.session_state["confirmar_borrado_centro"] = cid_sel
            st.rerun()

    try:
        excel_datos_home, nombre_excel_home = _obtener_excel_descarga_centro(cid_sel)
        st.markdown('<div class="marcador-btn-descarga-amarillo"></div>', unsafe_allow_html=True)
        st.download_button(
            "💾 Guardar archivo de datos",
            data=excel_datos_home,
            file_name=nombre_excel_home,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key=f"guardar_archivo_datos_home_{cid_sel}",
        )
    except Exception as e:
        st.error(f"No se pudo preparar el archivo de datos: {e}")

    if st.session_state.get("confirmar_borrado_centro") == cid_sel:
        st.warning(f"¿Eliminar el centro «{nombre_sel}» y todos sus detectores? Esta acción no se puede deshacer.")
        cc1, cc2 = st.columns(2)
        with cc1:
            if st.button("Sí, eliminar", key=f"confirmar_del_{cid_sel}", type="primary"):
                # Elegimos qué centro debe mostrarse después del borrado:
                # el siguiente de la lista y, si era el último, el anterior.
                siguiente_cid = None
                if idx_sel + 1 < len(centros):
                    siguiente_cid = centros[idx_sel + 1][0]
                elif idx_sel > 0:
                    siguiente_cid = centros[idx_sel - 1][0]

                delete_centro(cid_sel)
                st.session_state["confirmar_borrado_centro"] = None
                st.session_state.pop("selector_centro_home", None)
                if siguiente_cid is not None:
                    st.session_state["_centro_a_seleccionar_tras_borrado"] = siguiente_cid
                else:
                    st.session_state.pop("_centro_a_seleccionar_tras_borrado", None)
                st.rerun()
        with cc2:
            if st.button("Cancelar", key=f"cancelar_del_{cid_sel}"):
                st.session_state["confirmar_borrado_centro"] = None
                st.rerun()


# ============================================================
# PANTALLA: CENTRO (datos, detectores, generar PDF)
# ============================================================

def pantalla_centro_datos():
    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.session_state.view = "inicio"
        st.rerun()
        return
    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro
    ns_centro = f"centro_{cid}"

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    volver_clic = st.button("← Volver")
    aviso_placeholder = st.container()

    _zona_titulo_actual = _zona_para_titulo_centro(cid, zona)
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + _zona_titulo_actual if _zona_titulo_actual and _zona_titulo_actual.strip() else ""}'
    )
    st.markdown('<p class="subtitulo-amarillo">Datos del centro</p>', unsafe_allow_html=True)

    with st.container(border=True):
        img_key = ns_centro + "_img"
        if img_key not in st.session_state:
            st.session_state[img_key] = img_path

        fecha_valor_original = fecha or _ahora_espana().strftime("%d/%m/%Y")

        c1, c2 = st.columns(2)
        with c1:
            st.text_input("Nombre", value=nombre or "", key=ns_centro + "_nombre")
            st.text_input("Área / Zona :", value=zona or "", key=ns_centro + "_zona")
            st.text_input("Dirección", value=direccion or "", key=ns_centro + "_direccion")
            _date_input_texto("Fecha", ns_centro + "_fecha", fecha_valor_original)
        with c2:
            st.markdown('<div class="marcador-imagen-exterior"></div>', unsafe_allow_html=True)
            widget_imagen(
                "Imagen exterior", img_key, key_prefix=ns_centro + "_imgw",
                con_camara=True,
                ancho_miniatura=140, titulo_amarillo=True
            )

        hay_cambios = (
            st.session_state[ns_centro + "_nombre"] != (nombre or "")
            or st.session_state[ns_centro + "_zona"] != (zona or "")
            or st.session_state[ns_centro + "_direccion"] != (direccion or "")
            or st.session_state[ns_centro + "_fecha"] != fecha_valor_original
            or st.session_state.get(img_key) != img_path
        )

        def _guardar_cambios_centro():
            nombre_in = st.session_state[ns_centro + "_nombre"].strip()
            if not nombre_in:
                return False
            update_centro(
                cid, nombre_in,
                st.session_state[ns_centro + "_zona"].strip(),
                st.session_state[ns_centro + "_fecha"].strip(),
                st.session_state.get(img_key),
                st.session_state[ns_centro + "_direccion"].strip(),
            )
            return True

    if volver_clic:
        # Al volver se guardan automáticamente los datos del centro.
        # Si el nombre quedase en blanco, se conserva el nombre anterior.
        if not st.session_state[ns_centro + "_nombre"].strip():
            st.session_state[ns_centro + "_nombre"] = nombre or ""
        _guardar_cambios_centro()
        st.session_state.view = "centro"
        st.rerun()



def pantalla_centro_categorias():
    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.session_state.view = "inicio"
        st.rerun()
        return
    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro
    ns_centro = f"centro_{cid}"

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    volver_clic = st.button("← Volver")
    aviso_placeholder = st.container()

    _zona_titulo_actual = _zona_para_titulo_centro(cid, zona)
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + _zona_titulo_actual if _zona_titulo_actual and _zona_titulo_actual.strip() else ""}'
    )
    st.markdown('<p class="subtitulo-amarillo">Categorías profesionales</p>', unsafe_allow_html=True)

    with st.container(border=True):
        pend_nombre_key = f"nueva_categoria_nombre_pend_{cid}"
        pend_num_key = f"nueva_categoria_num_pend_{cid}"
        pend_turno_key = f"nueva_categoria_turno_pend_{cid}"
        if pend_nombre_key in st.session_state:
            st.session_state[f"nueva_categoria_nombre_{cid}"] = st.session_state.pop(pend_nombre_key)
        if pend_num_key in st.session_state:
            st.session_state[f"nueva_categoria_num_{cid}"] = st.session_state.pop(pend_num_key)
        if pend_turno_key in st.session_state:
            st.session_state[f"nueva_categoria_turno_{cid}"] = st.session_state.pop(pend_turno_key)

        mostrar_form_key = f"mostrar_form_categoria_{cid}"
        if mostrar_form_key not in st.session_state:
            st.session_state[mostrar_form_key] = False

        nueva_categoria = ""

        if not st.session_state[mostrar_form_key]:
            st.markdown('<div class="marcador-btn-anadir-categoria"></div>', unsafe_allow_html=True)
            if st.button("➕ Añadir", key=f"mostrar_add_categoria_{cid}", type="primary"):
                st.session_state[mostrar_form_key] = True
                st.rerun()
        else:
            cnew1, cnew2 = st.columns([2.2, 1.8])
            with cnew1:
                nueva_categoria = st.text_input(
                    "Categoría profesional", key=f"nueva_categoria_nombre_{cid}",
                )
            with cnew2:
                st.markdown('<div class="marcador-num-personas"></div>', unsafe_allow_html=True)
                clave_num_prof = f"nueva_categoria_num_{cid}"
                if clave_num_prof not in st.session_state:
                    st.session_state[clave_num_prof] = 1
                nuevo_num_personas = st.number_input(
                    "Nº de profesionales", key=clave_num_prof,
                    min_value=0, max_value=999, step=1,
                )

            nuevo_turno_cat = st.selectbox(
                "Turno", options=TURNOS_CATEGORIA_OPCIONES, key=f"nueva_categoria_turno_{cid}",
                index=None, placeholder="Selecciona un turno (opcional)",
            )

            st.markdown('<div class="marcador-btn-anadir-categoria"></div>', unsafe_allow_html=True)
            if st.button("Guardar categoría", key=f"add_categoria_{cid}", type="primary"):
                if not nueva_categoria.strip():
                    st.warning("Escribe el nombre de la categoría")
                else:
                    insert_categoria_centro(
                        cid,
                        nueva_categoria.strip(),
                        int(nuevo_num_personas),
                        nuevo_turno_cat or "",
                    )
                    st.session_state[pend_nombre_key] = ""
                    st.session_state[pend_num_key] = 1
                    st.session_state[pend_turno_key] = None
                    st.session_state[mostrar_form_key] = False
                    st.rerun()

        st.markdown("---")

        categorias_centro = fetch_categorias_centro(cid)
        if categorias_centro:
            seleccionadas_cat = []
            for cat_id, _, categoria, num_personas, turno_cat in categorias_centro:
                st.markdown('<div class="marcador-checkbox-categoria"></div>', unsafe_allow_html=True)
                etiqueta_cat = f"{categoria}: {num_personas}" + (f" — {turno_cat}" if turno_cat else "")
                marcado = st.checkbox(
                    etiqueta_cat, key=f"chk_cat_{cat_id}",
                )
                if marcado:
                    seleccionadas_cat.append(cat_id)

            if seleccionadas_cat:
                st.markdown('<div class="marcador-btn-eliminar"></div>', unsafe_allow_html=True)
                if st.button(f"❌ Eliminar seleccionadas ({len(seleccionadas_cat)})",
                             key=f"del_cats_sel_{cid}"):
                    for cat_id_del in seleccionadas_cat:
                        delete_categoria_centro(cat_id_del)
                    st.rerun()
        else:
            st.caption("Todavía no se ha añadido ninguna categoría profesional.")

    hay_cambios = bool(nueva_categoria.strip())
    confirm_key = f"categorias_confirmar_salida_{cid}"
    if volver_clic:
        if hay_cambios:
            st.session_state[confirm_key] = True
        else:
            st.session_state.view = "centro"
        st.rerun()

    if st.session_state.get(confirm_key):
        with aviso_placeholder:
            st.warning("⚠️ Tienes escrito el nombre de una categoría nueva sin añadir. ¿Qué quieres hacer?")
            cc1, cc2, cc3 = st.columns(3)
            with cc1:
                if st.button("➕ Añadir y salir", key=f"cc_guardar_{cid}", type="primary", use_container_width=True):
                    insert_categoria_centro(cid, nueva_categoria.strip(), int(nuevo_num_personas), nuevo_turno_cat or "")
                    st.session_state[confirm_key] = None
                    st.session_state.view = "centro"
                    st.rerun()
            with cc2:
                if st.button("🗑️ Descartar y salir", key=f"cc_descartar_{cid}", use_container_width=True):
                    st.session_state[confirm_key] = None
                    st.session_state.view = "centro"
                    st.rerun()
            with cc3:
                if st.button("Cancelar", key=f"cc_cancelar_{cid}", use_container_width=True):
                    st.session_state[confirm_key] = None
                    st.rerun()



def pantalla_centro_planos():
    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.session_state.view = "inicio"
        st.rerun()
        return
    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro
    ns_centro = f"centro_{cid}"

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    volver_clic = st.button("← Volver")
    aviso_placeholder = st.container()

    _zona_titulo_actual = _zona_para_titulo_centro(cid, zona)
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + _zona_titulo_actual if _zona_titulo_actual and _zona_titulo_actual.strip() else ""}'
    )
    st.markdown('<p class="subtitulo-amarillo">Planos del centro</p>', unsafe_allow_html=True)

    detectores = fetch_detectores(cid)
    planos_centro = fetch_planos_centro(cid)

    with st.container(border=True):
        st.caption(
            "Carga aquí los planos de las plantas del centro. Cada "
            "detector elegirá luego, en su propia pantalla, sobre cuál de "
            "estos planos marcar su ubicación."
        )

        for plano_c in planos_centro:
            plano_id_c, _, nombre_plano_c, ruta_plano_c, _ = plano_c
            with st.container(border=True):
                pcol1, pcol2, pcol3 = st.columns([1, 2, 1])
                with pcol1:
                    if ruta_plano_c and os.path.exists(ruta_plano_c):
                        st.image(ruta_plano_c, width=140)
                    else:
                        st.caption("(imagen no encontrada)")
                with pcol2:
                    st.markdown(f"**{nombre_plano_c}**")
                with pcol3:
                    st.markdown('<div class="marcador-btn-eliminar"></div>', unsafe_allow_html=True)
                    if st.button("❌ Eliminar", key=f"del_plano_{plano_id_c}"):
                        st.session_state["confirmar_borrado_plano"] = plano_id_c
                        st.rerun()

            if st.session_state.get("confirmar_borrado_plano") == plano_id_c:
                st.warning(
                    f"¿Eliminar el plano «{nombre_plano_c}»? Los detectores que lo "
                    "tuvieran asignado se quedarán sin plano ni punto marcado."
                )
                cc1, cc2 = st.columns(2)
                with cc1:
                    if st.button("Sí, eliminar", key=f"conf_del_plano_{plano_id_c}", type="primary"):
                        delete_plano_centro(plano_id_c)
                        st.session_state["confirmar_borrado_plano"] = None
                        st.rerun()
                with cc2:
                    if st.button("Cancelar", key=f"cancel_del_plano_{plano_id_c}"):
                        st.session_state["confirmar_borrado_plano"] = None
                        st.rerun()

        mostrar_add_plano_key = f"mostrar_add_plano_{cid}"
        if mostrar_add_plano_key not in st.session_state:
            st.session_state[mostrar_add_plano_key] = False

        st.markdown('<div class="marcador-btn-plano-amarillo"></div>', unsafe_allow_html=True)
        if st.button("➕ Añadir plano", type="tertiary"):
            st.session_state[mostrar_add_plano_key] = not st.session_state[mostrar_add_plano_key]
            st.rerun()

        nombre_plano_nuevo = ""
        archivo_plano_nuevo = None
        if st.session_state[mostrar_add_plano_key]:
            nombre_plano_nuevo = st.text_input(
                "Nombre del plano (p.ej. «Planta baja», «Planta 1»...)",
                key=f"nuevo_plano_nombre_{cid}",
            )
            st.markdown('<div class="marcador-btn-plano-amarillo"></div>', unsafe_allow_html=True)
            archivo_plano_nuevo = st.file_uploader(
                "Selecciona el plano (PDF de una sola página, PNG o JPG)",
                type=["pdf", "png", "jpg", "jpeg"],
                key=f"nuevo_plano_file_{cid}",
            )
            if archivo_plano_nuevo:
                st.markdown('<div class="marcador-btn-plano-amarillo"></div>', unsafe_allow_html=True)
                if st.button("Guardar plano", key=f"guardar_plano_{cid}", type="primary"):
                    if not nombre_plano_nuevo.strip():
                        st.warning("Ponle un nombre al plano")
                    else:
                        plano_jpg = convertir_plano_subido_a_jpg(archivo_plano_nuevo)
                        ruta_nueva = guardar_bytes_imagen(
                            plano_jpg,
                            f"plano_centro_{cid}",
                            ".jpg",
                        )
                        insert_plano_centro(cid, nombre_plano_nuevo.strip(), ruta_nueva, len(planos_centro))
                        st.session_state[mostrar_add_plano_key] = False
                        st.success("Plano añadido")
                    st.rerun()

    # Se considera "cambio sin guardar" tener el panel de "Añadir
    # plano" abierto con un nombre escrito y/o un archivo ya
    # seleccionado, pero sin haber pulsado todavía "Guardar plano".
    hay_cambios = st.session_state[mostrar_add_plano_key] and bool(
        nombre_plano_nuevo.strip() or archivo_plano_nuevo is not None
    )
    confirm_key = f"planos_confirmar_salida_{cid}"
    if volver_clic:
        if hay_cambios:
            st.session_state[confirm_key] = True
        else:
            st.session_state.view = "centro"
        st.rerun()

    if st.session_state.get(confirm_key):
        with aviso_placeholder:
            st.warning("⚠️ Tienes un plano nuevo sin terminar de guardar. ¿Qué quieres hacer?")
            cc1, cc2, cc3 = st.columns(3)
            with cc1:
                if st.button("💾 Guardar y salir", key=f"cp_guardar_{cid}", type="primary", use_container_width=True):
                    if not nombre_plano_nuevo.strip() or archivo_plano_nuevo is None:
                        st.warning("Hace falta un nombre y un archivo de plano para guardarlo.")
                    else:
                        plano_jpg = convertir_plano_subido_a_jpg(archivo_plano_nuevo)
                        ruta_nueva = guardar_bytes_imagen(
                            plano_jpg,
                            f"plano_centro_{cid}",
                            ".jpg",
                        )
                        insert_plano_centro(cid, nombre_plano_nuevo.strip(), ruta_nueva, len(planos_centro))
                        st.session_state[mostrar_add_plano_key] = False
                        st.session_state[confirm_key] = None
                        st.session_state.view = "centro"
                        st.rerun()
            with cc2:
                if st.button("🗑️ Descartar y salir", key=f"cp_descartar_{cid}", use_container_width=True):
                    st.session_state[mostrar_add_plano_key] = False
                    st.session_state[confirm_key] = None
                    st.session_state.view = "centro"
                    st.rerun()
            with cc3:
                if st.button("Cancelar", key=f"cp_cancelar_{cid}", use_container_width=True):
                    st.session_state[confirm_key] = None
                    st.rerun()



def pantalla_centro_detectores():
    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.session_state.view = "inicio"
        st.rerun()
        return
    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro
    ns_centro = f"centro_{cid}"
    abierto_key = f"detector_abierto_{cid}"
    confirm_key = f"detector_confirmar_salida_{cid}"
    if abierto_key not in st.session_state:
        st.session_state[abierto_key] = None

    selector_key = f"selector_detector_id_{cid}"
    pend_sel_key = f"selector_detector_id_pend_{cid}"
    if pend_sel_key in st.session_state:
        st.session_state[selector_key] = st.session_state.pop(pend_sel_key)

    def _ejecutar_accion_detector(accion):
        abierto_actual = st.session_state.get(abierto_key)

        # Guardado automático antes de abandonar/cambiar el detector.
        if abierto_actual is not None:
            ns_actual = f"det_{abierto_actual}"

            # CASO ESPECIAL: ficha "Nuevo detector".
            #
            # Al abrirla, algunos widgets reciben valores iniciales y por eso
            # _detector_tiene_cambios() puede considerarla modificada aunque el
            # usuario todavía no haya creado ningún detector. Hasta ahora, al
            # intentar seleccionar un detector existente, se intentaba guardar
            # esa ficha vacía; como el código es obligatorio, el guardado fallaba
            # y la función hacía return. El resultado era que, desde el momento
            # de pulsar "Nuevo detector", ya no se podía abrir ningún detector
            # guardado.
            #
            # Si todavía no hay código, la ficha nueva se considera un borrador
            # no creado: se descarta al cambiar de detector o salir. Si ya tiene
            # código (escrito o leído con la cámara), se conserva el
            # comportamiento de autoguardado.
            es_nuevo = abierto_actual == "nuevo"
            codigo_nuevo = st.session_state.get(ns_actual + "_codigo", "").strip() if es_nuevo else ""

            if es_nuevo and not codigo_nuevo:
                _limpiar_namespace(ns_actual)
                st.session_state.pop(pend_sel_key, None)
            elif _detector_tiene_cambios(ns_actual):
                guardado_ok, ns_final = _guardar_y_actualizar_snapshot_detector(
                    cid, abierto_actual, ns_actual, mostrar_mensajes=False
                )
                if not guardado_ok:
                    return

                # Cuando el guardado es consecuencia de una navegación, no debe
                # quedar una selección pendiente que anule el detector elegido
                # por el usuario en el siguiente rerun.
                st.session_state.pop(pend_sel_key, None)
                _limpiar_namespace(ns_final)
            else:
                _limpiar_namespace(ns_actual)

        tipo = accion["tipo"]
        if tipo == "volver":
            st.session_state[abierto_key] = None
            st.session_state.view = "centro"
        elif tipo == "nuevo":
            ns_destino = "det_nuevo"

            # El aviso verde del OCR pertenece únicamente a la ficha en la
            # que se reconoció el código. Al abrir un detector nuevo no debe
            # heredarse un mensaje de una ficha anterior.
            st.session_state.pop(ns_destino + "_ocr_codigo_mensaje", None)

            # Al abrir "Nuevo detector" dejamos el selector superior vacío.
            # Así no vuelve automáticamente al detector que estaba seleccionado
            # antes y, después, el usuario puede seleccionar cualquier detector
            # existente para cambiar de ficha.
            st.session_state[pend_sel_key] = None

            st.session_state[abierto_key] = "nuevo"
            st.session_state["detector_form_ns"] = ns_destino
            # Fuerza una instancia limpia del plano al cambiar de ficha.
            st.session_state[f"_plano_render_nonce_{cid}"] = (
                st.session_state.get(f"_plano_render_nonce_{cid}", 0) + 1
            )
        elif tipo == "cambiar":
            ns_destino = f"det_{accion['detector_id']}"

            # Limpiar cualquier aviso OCR antiguo asociado a la ficha que
            # vamos a abrir. El código guardado permanece intacto.
            st.session_state.pop(ns_destino + "_ocr_codigo_mensaje", None)

            st.session_state[abierto_key] = accion["detector_id"]
            st.session_state["detector_form_ns"] = ns_destino

            # El componente interactivo del plano mantiene estado propio en el
            # navegador. Al cambiar entre detectores forzamos una instancia nueva
            # para que no reutilice ni siquiera durante un instante la imagen o
            # las coordenadas del detector anterior.
            st.session_state[f"_plano_render_nonce_{cid}"] = (
                st.session_state.get(f"_plano_render_nonce_{cid}", 0) + 1
            )

        st.session_state[confirm_key] = None
        st.rerun()

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    volver_clic = st.button("← Volver")
    aviso_placeholder = st.container()

    _zona_titulo_actual = _zona_para_titulo_centro(cid, zona)
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + _zona_titulo_actual if _zona_titulo_actual and _zona_titulo_actual.strip() else ""}'
    )
    st.markdown('<p class="subtitulo-amarillo">Detectores colocados</p>', unsafe_allow_html=True)

    detectores = fetch_detectores(cid)
    nuevo_clic = False
    cambiar_accion = None

    with st.container(border=True):
        st.markdown('<div class="marcador-btn-nuevo-detector"></div>', unsafe_allow_html=True)
        if st.button("➕ Nuevo detector", type="tertiary"):
            nuevo_clic = True

        if not detectores:
            st.info("Todavía no se han añadido detectores para este centro.")
        else:
            # IMPORTANTE: el selector usa directamente el ID del detector.
            # Antes usaba la posición 0, 1, 2... del listado. Al crear un nuevo
            # detector Streamlit podía conservar esa posición como estado del
            # widget y no detectar correctamente el cambio posterior. El ID es
            # estable y hace que cada selección regenere la ficha adecuada.
            opciones = [d[0] for d in detectores]
            etiquetas = {}
            for d in detectores:
                did, _, planta, sala, fecha_det, codigo = d[0], d[1], d[2], d[3], d[4], d[5]
                partes = [codigo or f"Detector {did}"]
                if sala:
                    partes.append(sala)
                if planta:
                    partes.append(f"Planta {planta}")
                etiquetas[did] = " · ".join(partes)

            # Si por cualquier motivo quedó en sesión un ID que ya no existe
            # (por ejemplo, tras borrar un detector), limpiamos solo el estado
            # del selector, sin tocar la ficha abierta.
            valor_selector = st.session_state.get(selector_key)
            if valor_selector is not None and valor_selector not in opciones:
                st.session_state.pop(selector_key, None)

            did_sel = st.selectbox(
                "Detector",
                options=opciones,
                format_func=lambda did: etiquetas.get(did, f"Detector {did}"),
                key=selector_key,
                label_visibility="collapsed",
                index=None,
                placeholder="Selecciona un detector",
            )
            if did_sel is not None:
                # Si el usuario elige un detector diferente al que está
                # abierto, cambiamos a él. Esto incluye el caso en el que está
                # abierta la ficha "Nuevo detector".
                if st.session_state.get(abierto_key) != did_sel:
                    cambiar_accion = {"tipo": "cambiar", "detector_id": did_sel}

                st.markdown('<div class="marcador-btn-eliminar"></div>', unsafe_allow_html=True)
                if st.button("❌ Eliminar", key="btn_eliminar_detector_home"):
                    st.session_state["confirmar_borrado_det"] = did_sel
                    st.rerun()

                if st.session_state.get("confirmar_borrado_det") == did_sel:
                    st.warning("¿Eliminar este detector? Esta acción no se puede deshacer.")
                    cc1, cc2 = st.columns(2)
                    with cc1:
                        if st.button("Sí, eliminar", key=f"conf_del_det_{did_sel}", type="primary"):
                            delete_detector(did_sel)
                            st.session_state["confirmar_borrado_det"] = None
                            if st.session_state.get(abierto_key) == did_sel:
                                st.session_state[abierto_key] = None
                            if st.session_state.get(selector_key) == did_sel:
                                st.session_state.pop(selector_key, None)
                            st.rerun()
                    with cc2:
                        if st.button("Cancelar", key=f"cancel_del_det_{did_sel}"):
                            st.session_state["confirmar_borrado_det"] = None
                            st.rerun()

    abierto = st.session_state.get(abierto_key)
    hay_cambios = False
    if abierto is not None:
        detector_id_abierto = None if abierto == "nuevo" else abierto
        ns = f"det_{abierto}"
        st.session_state["detector_form_ns"] = ns
        _inicializar_ns_detector(cid, detector_id_abierto, ns)

        st.markdown("---")
        titulo_abierto = "Datos del detector" if detector_id_abierto else "Nuevo detector"
        st.markdown(f'<p class="subtitulo-amarillo">{titulo_abierto}</p>', unsafe_allow_html=True)

        with st.container(border=True):
            _renderizar_campos_detector(cid, detector_id_abierto, ns)
            hay_cambios = _detector_tiene_cambios(ns)

    if volver_clic:
        _ejecutar_accion_detector({"tipo": "volver"})
    elif nuevo_clic:
        _ejecutar_accion_detector({"tipo": "nuevo"})
    elif cambiar_accion:
        _ejecutar_accion_detector(cambiar_accion)




def pantalla_centro_retirada():
    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.session_state.view = "inicio"
        st.rerun()
        return
    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro
    ns_centro = f"centro_{cid}"
    confirm_key = f"retirada_confirmar_salida_{cid}"
    pend_sel_key = f"selector_retirada_pend_{cid}"
    if pend_sel_key in st.session_state:
        st.session_state[f"selector_retirada_{cid}"] = st.session_state.pop(pend_sel_key)

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    volver_clic = st.button("← Volver")
    aviso_placeholder = st.container()

    _zona_titulo_actual = _zona_para_titulo_centro(cid, zona)
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + _zona_titulo_actual if _zona_titulo_actual and _zona_titulo_actual.strip() else ""}'
    )
    st.markdown('<p class="subtitulo-amarillo">Retirada de detectores</p>', unsafe_allow_html=True)

    detectores = fetch_detectores(cid)
    hay_cambios = False
    did_r = None
    cambiar_sel_accion = None

    def _guardar_retirada_actual(did, fr_val, hr_val):
        actualizar_retirada_detector(did, fr_val.strip(), hr_val.strip())
        st.session_state[f"retirada_snapshot_{did}"] = (fr_val.strip(), hr_val.strip())

    with st.container(border=True):
        if not detectores:
            st.info("Todavía no se han añadido detectores para este centro.")
        else:
            opciones_r = list(range(len(detectores)))
            etiquetas_r = {}
            for i, d in enumerate(detectores):
                did_i, _, planta_i, sala_i, _, codigo_i = d[0], d[1], d[2], d[3], d[4], d[5]
                partes_i = [codigo_i or f"Detector {did_i}"]
                if sala_i:
                    partes_i.append(sala_i)
                if planta_i:
                    partes_i.append(f"Planta {planta_i}")
                etiquetas_r[i] = " · ".join(partes_i)

            idx_sel_r = st.selectbox(
                "Detector a retirar", options=opciones_r, format_func=lambda i: etiquetas_r[i],
                key=f"selector_retirada_{cid}", label_visibility="collapsed",
                index=None, placeholder="Selecciona un detector",
            )
            if idx_sel_r is None:
                st.info("")
            else:
                d_r = detectores[idx_sel_r]
                did_r = d_r[0]
                codigo_r = d_r[5]
                sala_r = d_r[3]
                foto_sit_r = d_r[9]
                punto_x_r = d_r[7]
                punto_y_r = d_r[8]
                plano_centro_id_r = d_r[17]
                fecha_retirada_real_r = d_r[18]
                hora_retirada_real_r = d_r[19]

                abierto_retirada_key = f"retirada_abierto_{cid}"
                if abierto_retirada_key not in st.session_state:
                    st.session_state[abierto_retirada_key] = did_r
                if st.session_state[abierto_retirada_key] != did_r:
                    cambiar_sel_accion = {"tipo": "cambiar", "detector_id": did_r, "idx": idx_sel_r}
                    did_mostrar = st.session_state[abierto_retirada_key]
                    d_mostrar = next((d for d in detectores if d[0] == did_mostrar), d_r)
                    did_r = d_mostrar[0]
                    codigo_r = d_mostrar[5]
                    sala_r = d_mostrar[3]
                    foto_sit_r = d_mostrar[9]
                    punto_x_r = d_mostrar[7]
                    punto_y_r = d_mostrar[8]
                    plano_centro_id_r = d_mostrar[17]
                    fecha_retirada_real_r = d_mostrar[18]
                    hora_retirada_real_r = d_mostrar[19]

                st.markdown(f"**Código:** {codigo_r or '-'}")
                st.markdown(f"**Sala:** {sala_r or '-'}")

                fr_key = f"retirada_fecha_{did_r}"
                hr_key = f"retirada_hora_{did_r}"
                pend_fr_key = f"retirada_fecha_pend_{did_r}"
                pend_hr_key = f"retirada_hora_pend_{did_r}"
                snapshot_key = f"retirada_snapshot_{did_r}"

                if pend_fr_key in st.session_state:
                    st.session_state[fr_key] = st.session_state.pop(pend_fr_key)
                if pend_hr_key in st.session_state:
                    st.session_state[hr_key] = st.session_state.pop(pend_hr_key)

                if fr_key not in st.session_state:
                    st.session_state[fr_key] = fecha_retirada_real_r or ""
                if hr_key not in st.session_state:
                    st.session_state[hr_key] = hora_retirada_real_r or ""
                if snapshot_key not in st.session_state:
                    st.session_state[snapshot_key] = (fecha_retirada_real_r or "", hora_retirada_real_r or "")

                _date_input_texto("Fecha de retirada", fr_key)
                st.text_input("Hora de retirada", key=hr_key)

                hay_cambios = (
                    st.session_state[fr_key].strip(), st.session_state[hr_key].strip()
                ) != st.session_state[snapshot_key]

                st.markdown('<div class="marcador-btn-retirada-amarillo"></div>', unsafe_allow_html=True)
                if st.button("🕒 Capturar fecha y hora", key=f"capturar_retirada_{did_r}", use_container_width=True):
                    ahora = _ahora_espana()
                    st.session_state[pend_fr_key] = ahora.strftime("%d/%m/%Y")
                    st.session_state[pend_hr_key] = ahora.strftime("%H:%M")
                    st.rerun()

                st.markdown("---")
                st.markdown('<div class="marcador-tabla-resultado-fila"></div>', unsafe_allow_html=True)
                fcol1, fcol2 = st.columns(2)
                with fcol1:
                    st.caption("Plano")
                    if (plano_centro_id_r and punto_x_r is not None and punto_y_r is not None
                            and punto_x_r >= 0 and punto_y_r >= 0):
                        plano_info_r = get_plano_centro(plano_centro_id_r)
                        if plano_info_r and os.path.exists(plano_info_r[3]):
                            ruta_tmp_punto_r = os.path.join(get_data_dir(), f"_tmp_plano_punto_retirada_{did_r}.jpg")
                            if generar_plano_con_punto(plano_info_r[3], punto_x_r, punto_y_r, ruta_tmp_punto_r):
                                st.image(ruta_tmp_punto_r, use_container_width=True)
                            else:
                                st.caption("No se pudo generar el plano con el punto")
                        else:
                            st.caption("Sin plano asignado")
                    else:
                        st.caption("Este detector no tiene plano ni punto marcado")
                with fcol2:
                    st.caption("Foto de situación")
                    if foto_sit_r and os.path.exists(foto_sit_r):
                        st.image(foto_sit_r, use_container_width=True)
                    else:
                        st.caption("Sin foto de situación")

    # Guardado automático al volver o al cambiar de detector.
    if volver_clic:
        if did_r is not None and hay_cambios:
            _guardar_retirada_actual(
                did_r,
                st.session_state[f"retirada_fecha_{did_r}"],
                st.session_state[f"retirada_hora_{did_r}"],
            )
        st.session_state[f"retirada_abierto_{cid}"] = None
        st.session_state.view = "centro"
        st.rerun()

    elif cambiar_sel_accion:
        if did_r is not None and hay_cambios:
            _guardar_retirada_actual(
                did_r,
                st.session_state[f"retirada_fecha_{did_r}"],
                st.session_state[f"retirada_hora_{did_r}"],
            )
        st.session_state[f"retirada_abierto_{cid}"] = cambiar_sel_accion["detector_id"]
        st.rerun()




def _sincronizar_valor_auto(key, valor_calculado):
    """Para los cuadros de texto editables cuyo contenido se calcula
    automáticamente a partir de otros datos (punto 3, conclusiones):
    Streamlit solo tiene en cuenta el "value=" la primera vez que se
    crea el widget; en las siguientes ejecuciones, aunque cambien los
    datos de los que depende, seguiría mostrando lo mismo si no se
    hace esto. Aquí se actualiza st.session_state[key] con el nuevo
    valor calculado SOLO si el usuario no lo ha modificado
    (es decir, si el texto seguía siendo igual al último valor
    calculado); si el usuario ya escribió lo suyo, no se toca."""
    prev_key = key + "__ultimo_auto"
    if key not in st.session_state:
        st.session_state[key] = valor_calculado
        st.session_state[prev_key] = valor_calculado
    elif st.session_state.get(prev_key) == st.session_state[key]:
        st.session_state[key] = valor_calculado
        st.session_state[prev_key] = valor_calculado
    else:
        st.session_state[prev_key] = valor_calculado


def _miniatura_pdf(pdf_bytes, ancho_px=140):
    """Genera una miniatura (PNG en bytes) de la primera página de un PDF,
    para previsualizarlo en pequeño en la propia app. Devuelve None si no
    se puede generar (p.ej. si el paquete pymupdf no estuviera disponible,
    o el archivo no fuera un PDF válido), sin que eso rompa nada más."""
    try:
        import pymupdf
        with pymupdf.open(stream=pdf_bytes, filetype="pdf") as doc:
            pagina = doc[0]
            escala = ancho_px / pagina.rect.width
            pix = pagina.get_pixmap(matrix=pymupdf.Matrix(escala, escala))
            return pix.tobytes("png")
    except Exception:
        return None



def _normalizar_referencia_zona_punto2(texto, nombre_zona, tipo_zona, idioma="gl"):
    """Corrige solo la preposición/artículo del tipo de zona, sin duplicar Unidad/Servicio."""
    texto = str(texto or "")
    nombre = str(nombre_zona or "").strip()
    tipo = str(tipo_zona or "").strip().lower()
    if not texto or not nombre or tipo not in ("unidade", "servizo"):
        return texto

    if idioma == "es":
        correcto = (
            f"en la Unidad de {nombre}"
            if tipo == "unidade"
            else f"en el Servicio de {nombre}"
        )
        tipos = r"(?:Unidad|Unidade|Servicio|Servizo)"
        patron_tipo = re.compile(
            rf"\ben\s+(?:el|la)\s+{tipos}\s+de\s+{re.escape(nombre)}\b",
            flags=re.IGNORECASE,
        )
        texto_nuevo, n = patron_tipo.subn(correcto, texto, count=1)
        if n:
            return texto_nuevo

        patron_nombre = re.compile(
            rf"\ben\s+{re.escape(nombre)}\b", flags=re.IGNORECASE
        )
        texto_nuevo, n = patron_nombre.subn(correcto, texto, count=1)
        return texto_nuevo if n else texto

    correcto = (
        f"na Unidade de {nombre}"
        if tipo == "unidade"
        else f"no Servizo de {nombre}"
    )
    tipos = r"(?:Unidade|Unidad|Servizo|Servicio)"
    patron_tipo = re.compile(
        rf"\b(?:no|na|en)\s+{tipos}\s+de\s+{re.escape(nombre)}\b",
        flags=re.IGNORECASE,
    )
    texto_nuevo, n = patron_tipo.subn(correcto, texto, count=1)
    if n:
        return texto_nuevo

    patron_nombre = re.compile(
        rf"\b(?:no|na|en)\s+{re.escape(nombre)}\b", flags=re.IGNORECASE
    )
    texto_nuevo, n = patron_nombre.subn(correcto, texto, count=1)
    return texto_nuevo if n else texto


def _poner_salas_negrita_punto3(docx_bytes, nombres_salas):
    """Pone en negrita todos los nombres de sala del punto 3, aunque Word los divida en varios runs."""
    try:
        from docx import Document
    except Exception:
        return docx_bytes

    salas_limpias = sorted(
        {str(s).strip() for s in (nombres_salas or []) if str(s).strip()},
        key=len,
        reverse=True,
    )
    if not salas_limpias:
        return docx_bytes

    try:
        doc = Document(io.BytesIO(docx_bytes))

        def _es_inicio_p3(t):
            x = t.casefold()
            return (
                ("información" in x or "informacion" in x)
                and ("traballador" in x or "trabajador" in x)
            )

        def _es_inicio_p4(t):
            x = t.casefold()
            return (
                ("condición" in x or "condicion" in x or "condicións" in x or "condicions" in x)
                and ("exposición" in x or "exposicion" in x)
            )

        dentro_p3 = False
        for p in doc.paragraphs:
            texto_p = p.text
            if _es_inicio_p3(texto_p.strip()):
                dentro_p3 = True
                continue
            if dentro_p3 and _es_inicio_p4(texto_p.strip()):
                break
            if not dentro_p3 or not texto_p:
                continue

            coincidencias = []
            for sala in salas_limpias:
                for m in re.finditer(re.escape(sala), texto_p, flags=re.IGNORECASE):
                    coincidencias.append((m.start(), m.end()))
            if not coincidencias:
                continue

            coincidencias.sort()
            unidas = []
            for a, b in coincidencias:
                if not unidas or a >= unidas[-1][1]:
                    unidas.append([a, b])
                else:
                    unidas[-1][1] = max(unidas[-1][1], b)

            runs_previos = list(p.runs)
            base_bold = runs_previos[0].bold if runs_previos else None
            base_italic = runs_previos[0].italic if runs_previos else None
            base_underline = runs_previos[0].underline if runs_previos else None

            for run in runs_previos:
                p._element.remove(run._element)

            pos = 0
            for a, b in unidas:
                if a > pos:
                    r = p.add_run(texto_p[pos:a])
                    r.bold = base_bold
                    r.italic = base_italic
                    r.underline = base_underline
                r = p.add_run(texto_p[a:b])
                r.bold = True
                r.italic = base_italic
                r.underline = base_underline
                pos = b

            if pos < len(texto_p):
                r = p.add_run(texto_p[pos:])
                r.bold = base_bold
                r.italic = base_italic
                r.underline = base_underline

        salida = io.BytesIO()
        doc.save(salida)
        return salida.getvalue()
    except Exception:
        return docx_bytes


def pantalla_centro_informe_completo():
    """Genera el informe oficial completo (Word, con anexos opcionales) a
    partir de los datos ya cargados en la app para este centro: genera el
    Excel automáticamente (como si se hubiera pulsado "Generar" en Informes
    y descargas) y lo carga sin necesidad de subirlo."""
    import sys as _sys
    import os as _os
    import pandas as pd
    _dir_app = _os.path.dirname(_os.path.abspath(__file__))
    if _dir_app not in _sys.path:
        _sys.path.insert(0, _dir_app)

    from utils_informe.excel_parser import (
        COL_CODIGO_DETECTOR, ExcelFormatError, areas_muestreadas, categorias_resumen,
        categorias_turnos_bullets, extraer_resultados_pdf_laboratorio, filter_group,
        group_options, load_workbook, merge_resultados, postos_traballo_bullets,
        salas_medidas, traducir_es_gl,
    )

    # Vocabulario adicional solicitado para la traducción automática ES -> GL.
    # Se aplica solo al generar el informe; los textos originales guardados no cambian.
    from utils_informe import excel_parser as _excel_parser_mod
    _excel_parser_mod.ES_GL_VOCABULARIO.update({
        "prevencion": "prevención",
        "riesgo": "risco",
        "riesgos": "riscos",
        "laboral": "laboral",
        "laborales": "laborais",
        "trabajo": "traballo",
    })

    _traducir_es_gl_base = traducir_es_gl

    def traducir_es_gl(texto):
        """Traducción ampliada con categorías compuestas solicitadas."""
        resultado = _traducir_es_gl_base(texto)
        resultado = re.sub(
            r"\bMedicina\s+del\s+traballo\b",
            "Medicina do traballo",
            resultado,
            flags=re.IGNORECASE,
        )
        resultado = re.sub(
            r"\bEnfermar[ií]a\s+del\s+traballo\b",
            "Enfermería do traballo",
            resultado,
            flags=re.IGNORECASE,
        )
        return resultado

    _excel_parser_mod.traducir_es_gl = traducir_es_gl
    from utils_informe.docx_generator import (
        ReportContext, _area_grupo, _quitar_turno_de_puestos, generar_conclusion_automatica,
        generar_texto_objeto_automatico, generar_texto_punto3_automatico,
        generar_texto4_automatico, generar_texto5_automatico, generate_report,
    )
    from utils_informe.pdf_tools import PdfToolsError, construir_pdf_completo, libreoffice_disponible
    from utils_informe.anexo2 import Anexo2Error, extraer_datos_planos, generar_documento_anexo2
    from utils_informe.assets import anexo3_por_defecto, anexo4_por_defecto

    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.session_state.view = "inicio"
        st.rerun()
        return
    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    if st.button("← Volver"):
        if st.session_state.get(f"ic_hay_cambios_{cid}"):
            with st.spinner("Actualizando la hoja Excel con los últimos cambios..."):
                nombre_xlsx_auto = _nombre_documento(nombre, "HOJA-DATOS") + ".xlsx"
                ruta_xlsx_auto = os.path.join(get_data_dir(), nombre_xlsx_auto)
                generar_excel(cid, ruta_xlsx_auto)
                st.session_state["ultimo_excel"] = ruta_xlsx_auto
                st.session_state["ultimo_excel_nombre"] = nombre_xlsx_auto
                st.session_state["ultimo_excel_centro"] = cid
            st.session_state[f"ic_hay_cambios_{cid}"] = False
        st.session_state.view = "centro"
        st.rerun()

    _zona_titulo_actual = _zona_para_titulo_centro(cid, zona)
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + _zona_titulo_actual if _zona_titulo_actual and _zona_titulo_actual.strip() else ""}'
    )
    st.markdown(
        '<p class="subtitulo-amarillo">INFORME DE RESULTADOS DE MEDICIONES DE Rn</p>',
        unsafe_allow_html=True,
    )

    detectores = fetch_detectores(cid)
    if not detectores:
        st.info("Añade al menos un detector a este centro para poder generar el informe.")
        return

    # --- Generar y cargar el Excel automáticamente, sin pedirlo ---
    ruta_excel = os.path.join(get_data_dir(), f"_tmp_informe_completo_{cid}.xlsx")
    try:
        generar_excel(cid, ruta_excel)
        wb = load_workbook(ruta_excel)
    except ExcelFormatError as e:
        st.error(f"El Excel generado no tiene el formato esperado: {e}")
        return
    except Exception as e:
        st.error(f"No se ha podido generar/cargar el Excel de este centro: {e}")
        return

    df = wb["detectores"]
    det_meta = wb["detectores_meta"]
    planos_meta = wb["planos_meta"]
    categorias_df = wb["categorias"]

    opciones = group_options(df, "Centro")
    if not opciones:
        st.error("El Excel generado no tiene ningún valor en la columna 'Centro'.")
        return
    selected_value = opciones[0]  # un solo centro: el actual
    df_center = filter_group(df, selected_value, "Centro")

    areas = areas_muestreadas(df_center)
    salas = salas_medidas(df_center)
    total_personas, categorias_texto = categorias_resumen(categorias_df)
    tipo_centro_actual = (get_tipo_centro(cid) or "").strip().lower()
    # Solo cuenta como "Atención Primaria" para el horario de 14 a 21
    # h si el tipo de centro incluye ese texto exactamente; "PAC" a
    # secas (sin "Atención Primaria" delante) no cuenta.
    es_ap_centro = any(p in tipo_centro_actual for p in ("atención primaria", "atencion primaria"))
    postos_bullets_default = postos_traballo_bullets(df_center, es_atencion_primaria=es_ap_centro)
    total_traballadores, categorias_bullets_default = categorias_turnos_bullets(
        categorias_df, es_atencion_primaria=es_ap_centro,
    )

    st.success(f"Excel generado y cargado automáticamente: {len(df_center)} detector(es) en «{selected_value}».")

    logo_bytes, _logo_nombre = get_logo_informe()

    # =========================================================
    # 1. IDENTIFICACIÓN DEL CENTRO DE TRABAJO
    # =========================================================
    xerencia_previa = str(planos_meta.get("Empresa", ""))
    cif_previa = str(planos_meta.get("CIF", ""))
    enderezo_previo = str(det_meta.get("Dirección", ""))
    completo_1 = bool(
        st.session_state.get(f"ic_xerencia_{cid}", xerencia_previa).strip()
        and st.session_state.get(f"ic_cif_{cid}", cif_previa).strip()
        and st.session_state.get(f"ic_enderezo_{cid}", enderezo_previo).strip()
    )
    with _acordeon_informe("1", "IDENTIFICACIÓN DEL CENTRO DE TRABAJO", completo_1):
        st.caption("Autorrellenados a partir de los datos ya guardados en la app. Revísalos y complétalos.")
        col1, col2 = st.columns(2)
        with col1:
            xerencia = st.text_input("Xerencia (Empresa)", value=xerencia_previa, key=f"ic_xerencia_{cid}")
            cif = st.text_input("CIF", value=cif_previa, key=f"ic_cif_{cid}")
            centro_nombre = st.text_input("Nombre completo del centro", value=selected_value, key=f"ic_centro_nombre_{cid}")
            servizo_unidade = st.text_input("Servizo / Unidade mostrexada (Área)", value=", ".join(areas), key=f"ic_servizo_{cid}")
            tipo_zona_elegido = ""
            # Si el Área/Zona ya empieza por "Unidad"/"Servicio" (o su
            # forma en galego, "unidade"/"Servizo"), no hace falta
            # preguntar con el checklist: se entiende directamente de
            # lo que ya se ha escrito, y se separa la palabra
            # ("unidade"/"servizo", para las plantillas) del resto del
            # nombre del área (que es lo que luego se traduce y se
            # muestra en el informe).
            _match_zona_directa = re.match(
                r"^(unidad|unidade|servicio|servizo)\b\s*(?:de\s+)?(.*)$",
                servizo_unidade.strip(), re.IGNORECASE,
            ) if servizo_unidade.strip() else None
            if _match_zona_directa:
                _tipo_palabra = _match_zona_directa.group(1).lower()
                tipo_zona_elegido = "unidade" if _tipo_palabra in ("unidad", "unidade") else "servizo"
                servizo_unidade_efectivo = _match_zona_directa.group(2).strip()
            elif _area_grupo(servizo_unidade) == "B":
                st.caption(
                    "El Área/Zona no es Atención Primaria, PAC, Atención Primaria + PAC ni "
                    "está en blanco: indica si es una Unidad, un Servicio, u otro tipo de zona."
                )
                opcion_tipo_zona = st.radio(
                    "¿Qué es?", options=["Unidad", "Servicio", "Otro tipo de zona"],
                    key=f"ic_tipo_zona_{cid}", horizontal=True, label_visibility="collapsed",
                )
                tipo_zona_elegido = {"Unidad": "unidade", "Servicio": "servizo", "Otro tipo de zona": "outro"}[opcion_tipo_zona]
                if tipo_zona_elegido == "outro":
                    preview_zona = ""
                elif tipo_zona_elegido == "unidade":
                    preview_zona = f"na Unidade de {servizo_unidade}" if servizo_unidade else "na Unidade"
                else:
                    preview_zona = f"no Servizo de {servizo_unidade}" if servizo_unidade else "no Servizo"
                st.session_state[f"ic_preview_zona_{cid}"] = preview_zona
                st.text_input(
                    "Vista previa (así aparecerá en el informe; en blanco si es \"Otro tipo de zona\", para completarlo tú)",
                    disabled=True, key=f"ic_preview_zona_{cid}",
                )
                servizo_unidade_efectivo = servizo_unidade
            else:
                servizo_unidade_efectivo = servizo_unidade
            enderezo = st.text_input("Dirección (Enderezo)", value=enderezo_previo, key=f"ic_enderezo_{cid}")
        with col2:
            _datos_informe_guardados = get_datos_informe_centro(cid)

            def _guardar_dato_informe(campo, key_widget):
                set_datos_informe_centro(cid, **{campo: st.session_state.get(key_widget, "")})
                st.session_state[f"ic_hay_cambios_{cid}"] = True

            superficie_construida = st.text_input(
                "Superficie construida (m²)", value=_datos_informe_guardados["superficie_construida"],
                key=f"ic_sup_constr_{cid}",
                on_change=_guardar_dato_informe, args=("superficie_construida", f"ic_sup_constr_{cid}"),
            )
            superficie_util = st.text_input(
                "Superficie útil (m²)", value=_datos_informe_guardados["superficie_util"],
                key=f"ic_sup_util_{cid}",
                on_change=_guardar_dato_informe, args=("superficie_util", f"ic_sup_util_{cid}"),
            )
            num_plantas = st.text_input(
                "N.º de plantas", value=_datos_informe_guardados["num_plantas"],
                key=f"ic_num_plantas_{cid}",
                on_change=_guardar_dato_informe, args=("num_plantas", f"ic_num_plantas_{cid}"),
            )
            # Por defecto, la fecha del informe es la de hoy (el día en
            # que se genera), aunque se puede cambiar; se
            # muestra en formato día/mes/año.
            fecha_default = _ahora_espana().date()
            data_informe = st.date_input(
                "Fecha del informe", value=fecha_default, key=f"ic_fecha_informe_{cid}",
                format="DD/MM/YYYY",
            )

    # =========================================================
    # 2. OBJETO DEL INFORME (siempre gris: solo necesita el nombre
    # del centro, que siempre está disponible)
    # =========================================================
    with _acordeon_informe("2", "OBJETO DEL INFORME", True):
        _ctx_previa_objeto = ReportContext(
            centro=traducir_es_gl(centro_nombre), servizo_unidade=traducir_es_gl(servizo_unidade_efectivo),
            tipo_zona=tipo_zona_elegido,
        )
        _texto_objeto_auto = generar_texto_objeto_automatico(_ctx_previa_objeto)
        _texto_objeto_auto = _normalizar_referencia_zona_punto2(
            _texto_objeto_auto, traducir_es_gl(servizo_unidade_efectivo), tipo_zona_elegido, idioma="gl"
        )
        _sincronizar_valor_auto(f"ic_objeto_{cid}", _texto_objeto_auto)
        objeto_manual = st.text_area(
            "texto_objeto", label_visibility="collapsed", height=140, key=f"ic_objeto_{cid}",
        )
        objeto_es_auto = st.session_state.get(f"ic_objeto_{cid}__ultimo_auto") == objeto_manual

    # =========================================================
    # 3. INFORMACIÓN SOBRE LOS TRABAJADORES
    # =========================================================
    completo_3 = bool(str(total_traballadores or "0").strip() not in ("", "0"))
    with _acordeon_informe("3", "INFORMACIÓN SOBRE LOS TRABAJADORES", completo_3):
        st.caption("")
        col3, col4 = st.columns(2)
        with col3:
            st.markdown("**Puestos de trabajo por sala** (una línea por puesto)")
            _sincronizar_valor_auto(f"ic_postos_{cid}", "\n".join(postos_bullets_default))
            postos_text = st.text_area(
                "postos_text", label_visibility="collapsed", height=120, key=f"ic_postos_{cid}",
            )
            postos_bullets = [line.strip() for line in postos_text.splitlines() if line.strip()]
            num_traballadores_total = st.text_input(
                "N.º total de trabajadores adscritos", value=str(total_traballadores or ""), key=f"ic_num_trab_{cid}",
            )
        with col4:
            st.markdown("**Categorías profesionales y quendas** (una línea por categoría)")
            _sincronizar_valor_auto(f"ic_categorias_{cid}", "\n".join(categorias_bullets_default))
            categorias_text = st.text_area(
                "categorias_text", label_visibility="collapsed", height=120, key=f"ic_categorias_{cid}",
            )
            categorias_bullets = [line.strip() for line in categorias_text.splitlines() if line.strip()]
            data_informacion_traballadores = st.text_input(
                "Fecha de comunicación a los trabajadores", value=_datos_informe_guardados["fecha_comunicacion_trab"],
                key=f"ic_data_info_trab_{cid}",
                on_change=_guardar_dato_informe, args=("fecha_comunicacion_trab", f"ic_data_info_trab_{cid}"),
            )
            medio_informacion_traballadores = st.text_input(
                "Medio de comunicación", value=_datos_informe_guardados["medio_comunicacion"] or "correo electrónico",
                key=f"ic_medio_info_{cid}",
                on_change=_guardar_dato_informe, args=("medio_comunicacion", f"ic_medio_info_{cid}"),
            )

        st.markdown("**Punto 3 completo del informe**")
        _ctx_previa_punto3 = ReportContext(
            centro=traducir_es_gl(centro_nombre), servizo_unidade=traducir_es_gl(servizo_unidade_efectivo), postos_bullets=postos_bullets,
            num_traballadores_total=num_traballadores_total, categorias_bullets=categorias_bullets,
            medio_informacion_traballadores=medio_informacion_traballadores,
            data_informacion_traballadores=data_informacion_traballadores,
            tipo_zona=tipo_zona_elegido,
        )
        _texto_punto3_auto = generar_texto_punto3_automatico(_ctx_previa_punto3)
        _sincronizar_valor_auto(f"ic_punto3_{cid}", _texto_punto3_auto)
        texto_punto3_manual = st.text_area(
            "texto_punto3", label_visibility="collapsed", height=220, key=f"ic_punto3_{cid}",
        )
        # Si no se ha tocado (sigue igual que el último texto
        # calculado automáticamente, en gallego), se puede volver a
        # generar en el idioma que haga falta a la hora de generar el
        # informe; si el usuario ya escribió lo suyo, se
        # mantiene igual en los dos idiomas (solo hay un cuadro de
        # texto, no uno por idioma).
        punto3_es_auto = st.session_state.get(f"ic_punto3_{cid}__ultimo_auto") == texto_punto3_manual
        st.caption("Las líneas que empiecen por '-' se muestran como viñetas en el informe.")

    # =========================================================
    # 4. CONDICIONES DE LA EXPOSICIÓN (siempre gris: texto fijo)
    # =========================================================
    with _acordeon_informe("4", "CONDICIONES DE LA EXPOSICIÓN", True):
        _texto4_auto = generar_texto4_automatico()
        _sincronizar_valor_auto(f"ic_texto4_{cid}", _texto4_auto)
        texto4_manual = st.text_area(
            "texto4", label_visibility="collapsed", height=160, key=f"ic_texto4_{cid}",
        )
        texto4_es_auto = st.session_state.get(f"ic_texto4_{cid}__ultimo_auto") == texto4_manual

    # =========================================================
    # 5. PLANOS (siempre gris: texto fijo)
    # =========================================================
    with _acordeon_informe("5", "PLANOS", True):
        _texto5_auto = generar_texto5_automatico()
        _sincronizar_valor_auto(f"ic_texto5_{cid}", _texto5_auto)
        texto5_manual = st.text_area(
            "texto5", label_visibility="collapsed", height=120, key=f"ic_texto5_{cid}",
        )
        texto5_es_auto = st.session_state.get(f"ic_texto5_{cid}__ultimo_auto") == texto5_manual

    # =========================================================
    # 6. RESULTADO DE LAS MEDICIONES REALIZADAS
    # =========================================================
    df_working = df_center.copy()

    resultados_pdf_file_key = f"ic_resultados_pdf_{cid}"
    pendientes_previo = int(df_working["Resultado Bq/m3"].isna().sum()) if "Resultado Bq/m3" in df_working else len(df_working)
    completo_6 = pendientes_previo == 0 and len(df_working) > 0

    with _acordeon_informe("6", "RESULTADO DE LAS MEDICIONES REALIZADAS", completo_6):
        st.caption(
            "Rellena el resultado y la incertidumbre directamente en la tabla, o sube el PDF "
            "con los resultados del laboratorio para completarlos automáticamente."
        )

        resultados_pdf_file = st.file_uploader(
            "Subir PDF de resultados del laboratorio",
            type=["pdf"], key=resultados_pdf_file_key,
        )
        if resultados_pdf_file:
            try:
                resultados_pdf_df = extraer_resultados_pdf_laboratorio(resultados_pdf_file)
                if resultados_pdf_df.empty:
                    st.warning(
                        "No se ha encontrado ninguna fila de resultados con el formato esperado "
                        "en este PDF."
                    )
                else:
                    coincidencias = set(resultados_pdf_df["Código"]) & set(
                        df_working[COL_CODIGO_DETECTOR].astype(str).str.strip()
                    )
                    df_working = merge_resultados(df_working, resultados_pdf_df)

                    # Se incluyen los códigos de detector actuales en
                    # el identificador, para que corregir un código y
                    # volver a subir el mismo PDF sí se note.
                    firma_codigos_pdf = ",".join(sorted(df_working[COL_CODIGO_DETECTOR].astype(str)))
                    fid_resultados_pdf = f"{resultados_pdf_file.name}_{resultados_pdf_file.size}_{firma_codigos_pdf}"
                    if st.session_state.get(f"ic_resultados_pdf_fid_{cid}") != fid_resultados_pdf:
                        st.session_state[f"ic_resultados_pdf_fid_{cid}"] = fid_resultados_pdf
                        for idx_r, fila_r in df_working.reset_index(drop=True).iterrows():
                            resultado_r = fila_r.get("Resultado Bq/m3")
                            try:
                                resultado_r_val = float(resultado_r) if pd.notna(resultado_r) else None
                            except (TypeError, ValueError):
                                resultado_r_val = None
                            st.session_state[f"ic_resultado_{cid}_{idx_r}"] = resultado_r_val
                            incert_r = fila_r.get("Incerteza expandida e K")
                            incert_r_val = (
                                "" if incert_r is None or (isinstance(incert_r, float) and pd.isna(incert_r))
                                else str(incert_r)
                            )
                            st.session_state[f"ic_incerteza_{cid}_{idx_r}"] = incert_r_val
                            detector_id_r = fila_r.get("ID")
                            try:
                                detector_id_r = int(detector_id_r) if pd.notna(detector_id_r) else None
                            except (TypeError, ValueError):
                                detector_id_r = None
                            if detector_id_r:
                                actualizar_resultado_detector(detector_id_r, resultado_r_val, incert_r_val)
                        st.session_state[f"ic_hay_cambios_{cid}"] = True

                    if coincidencias:
                        st.success(
                            f"Encontrados {len(coincidencias)} detector(es) de este centro en el PDF: "
                            + ", ".join(sorted(coincidencias))
                        )
                    else:
                        st.warning(
                            "El PDF se ha leído correctamente, pero ninguno de sus códigos de "
                            "detector coincide con los de este centro."
                        )
            except Exception as e:
                st.error(f"No se ha podido leer el PDF de resultados: {e}")

        # Tabla propia (no st.data_editor: ese componente pinta la
        # cabecera y las celdas ya confirmadas dentro de un <canvas>,
        # con los colores fijados por el tema oscuro de la app -letra
        # blanca sobre fondo oscuro-, sin ninguna forma de forzarlo
        # por CSS). Aquí se dibuja una tabla real, lo más parecida
        # posible a la que sale en el informe final: una fila de
        # cabecera con las mismas columnas del punto 6, y debajo, una
        df_working_reset = df_working.reset_index(drop=True)

        def _texto_seguro(valor):
            if valor is None or (isinstance(valor, float) and pd.isna(valor)):
                return ""
            return str(valor)

        # ------------------------------------------------------------
        # Tabla editable tipo Excel. Solo Resultado e Incertidumbre
        # pueden modificarse; el resto de los datos son informativos.
        # ------------------------------------------------------------
        tabla_resultados = pd.DataFrame({
            "ID": df_working_reset["ID"],
            "Código zona": df_working_reset["Código de la sala"].apply(_texto_seguro),
            "Código detector": df_working_reset["Código"].apply(_texto_seguro),
            "Fecha inicio": df_working_reset["Fecha de colocación fmt"].apply(_texto_seguro),
            "Fecha fin": df_working_reset["Fecha de retirada real fmt"].apply(_texto_seguro),
            "Sala / Puestos": [
                (
                    _texto_seguro(fila.get("Sala", ""))
                    + (
                        " — " + _quitar_turno_de_puestos(
                            _texto_seguro(fila.get("Profesionales en la sala", ""))
                        )
                        if _quitar_turno_de_puestos(
                            _texto_seguro(fila.get("Profesionales en la sala", ""))
                        )
                        else ""
                    )
                )
                for _, fila in df_working_reset.iterrows()
            ],
            "Resultado Bq/m³": pd.to_numeric(
                df_working_reset["Resultado Bq/m3"], errors="coerce"
            ),
            "Incertidumbre": df_working_reset["Incerteza expandida e K"].apply(_texto_seguro),
        })

        # Si se ha cargado un nuevo PDF, se cambia la clave del editor para
        # que los valores importados se reflejen inmediatamente en la tabla.
        editor_version = st.session_state.get(f"ic_resultados_pdf_fid_{cid}", "manual")
        editor_key = f"ic_tabla_resultados_{cid}_{editor_version}"


        # Tabla de resultados con el mismo criterio visual que el resto de campos:
        # azul si tiene contenido y rosa + punto rojo si está pendiente.
        # Se construye con widgets nativos para que el estilo se aplique también
        # a las celdas editables (st.data_editor no permite estilos condicionales
        # en columnas editables).
        anchos_tabla = [1.0, 1.15, 1.0, 1.0, 2.0, 1.0, 1.15]
        cabeceras_tabla = [
            "Código zona",
            "Código detector",
            "Fecha inicio",
            "Fecha fin",
            "Sala / Puestos",
            "Resultado Bq/m³",
            "Incertidumbre",
        ]
        st.markdown('<div class="marcador-tabla-punto6-excel"></div>', unsafe_allow_html=True)
        cols_cabecera = st.columns(anchos_tabla, gap=None)
        for col_cab, titulo_cab in zip(cols_cabecera, cabeceras_tabla):
            with col_cab:
                st.markdown(
                    f"<div style='font-size:0.82rem;font-weight:600;"
                    f"padding:0;'>{html.escape(titulo_cab)}</div>",
                    unsafe_allow_html=True,
                )

        filas_editadas = []
        for idx_tabla, fila_tabla in tabla_resultados.reset_index(drop=True).iterrows():
            st.markdown('<div class="marcador-fila-punto6-excel"></div>', unsafe_allow_html=True)
            cols_fila = st.columns(anchos_tabla, gap=None)

            valores_fijos = [
                _texto_seguro(fila_tabla["Código zona"]),
                _texto_seguro(fila_tabla["Código detector"]),
                _texto_seguro(fila_tabla["Fecha inicio"]),
                _texto_seguro(fila_tabla["Fecha fin"]),
                _texto_seguro(fila_tabla["Sala / Puestos"]),
            ]
            for pos_col, valor_fijo in enumerate(valores_fijos):
                with cols_fila[pos_col]:
                    st.text_input(
                        cabeceras_tabla[pos_col],
                        value=valor_fijo,
                        key=f"{editor_key}_fijo_{idx_tabla}_{pos_col}",
                        disabled=True,
                        label_visibility="collapsed",
                    )

            resultado_inicial = fila_tabla["Resultado Bq/m³"]
            if pd.isna(resultado_inicial):
                resultado_inicial = None
            else:
                try:
                    resultado_inicial = float(resultado_inicial)
                except (TypeError, ValueError):
                    resultado_inicial = None

            with cols_fila[5]:
                resultado_editado = st.number_input(
                    "Resultado Bq/m³",
                    min_value=0.0,
                    value=resultado_inicial,
                    step=1.0,
                    format="%.0f",
                    key=f"{editor_key}_resultado_{idx_tabla}",
                    label_visibility="collapsed",
                )

            incertidumbre_inicial = _texto_seguro(
                fila_tabla["Incertidumbre"]
            )
            with cols_fila[6]:
                incertidumbre_editada = st.text_input(
                    "Incertidumbre",
                    value=incertidumbre_inicial,
                    key=f"{editor_key}_incertidumbre_{idx_tabla}",
                    label_visibility="collapsed",
                )

            filas_editadas.append({
                "ID": fila_tabla["ID"],
                "Código zona": valores_fijos[0],
                "Código detector": valores_fijos[1],
                "Fecha inicio": valores_fijos[2],
                "Fecha fin": valores_fijos[3],
                "Sala / Puestos": valores_fijos[4],
                "Resultado Bq/m³": resultado_editado,
                "Incertidumbre": incertidumbre_editada,
            })

        tabla_editada = pd.DataFrame(filas_editadas, columns=tabla_resultados.columns)

        # Guardado automático: cualquier cambio hecho en la tabla se persiste en SQLite
        # en el mismo ciclo de Streamlit, evitando que desaparezca al salir y volver a entrar.
        for _, fila_editada_auto in tabla_editada.iterrows():
            try:
                detector_id_auto = int(fila_editada_auto["ID"])
                resultado_auto = fila_editada_auto["Resultado Bq/m³"]
                resultado_auto = float(resultado_auto) if pd.notna(resultado_auto) else None
                incertidumbre_auto = _texto_seguro(
                    fila_editada_auto["Incertidumbre"]
                ).strip()

                detector_db_auto = get_detector(detector_id_auto)
                if detector_db_auto:
                    resultado_db_auto = detector_db_auto[20]
                    incertidumbre_db_auto = detector_db_auto[21] or ""
                    if resultado_db_auto != resultado_auto or incertidumbre_db_auto != incertidumbre_auto:
                        actualizar_resultado_detector(
                            detector_id_auto, resultado_auto, incertidumbre_auto
                        )
                        st.session_state[f"ic_hay_cambios_{cid}"] = True
            except Exception:
                pass

        if st.button(
            "Guardar resultados",
            key=f"ic_guardar_resultados_{cid}",
            type="primary",
        ):
            guardados = 0
            errores = 0
            for _, fila_editada in tabla_editada.iterrows():
                try:
                    detector_id = int(fila_editada["ID"])
                    resultado = fila_editada["Resultado Bq/m³"]
                    resultado = float(resultado) if pd.notna(resultado) else None
                    incertidumbre = _texto_seguro(
                        fila_editada["Incertidumbre"]
                    ).strip()
                    actualizar_resultado_detector(
                        detector_id, resultado, incertidumbre
                    )
                    guardados += 1
                except Exception:
                    errores += 1

            st.session_state[f"ic_hay_cambios_{cid}"] = True
            if errores:
                st.warning(
                    f"Se guardaron {guardados} resultado(s), pero hubo "
                    f"{errores} fila(s) que no se pudieron guardar."
                )
            else:
                st.success(f"Resultados guardados correctamente ({guardados}).")

        # Reconstruir df_final con lo que se ve actualmente en el editor.
        df_final = df_working_reset.copy()
        df_final["Resultado Bq/m3"] = pd.to_numeric(
            tabla_editada["Resultado Bq/m³"], errors="coerce"
        )
        df_final["Incerteza expandida e K"] = tabla_editada[
            "Incertidumbre"
        ].fillna("").astype(str)

        pendientes = int(
            df_final["Resultado Bq/m3"].isna().sum()
        ) if "Resultado Bq/m3" in df_final else 0
        if pendientes:
            st.info(f"ℹ️ Quedan {pendientes} detector(es) sin resultado.")

        exceeded = int(
            (df_final["Resultado Bq/m3"] > 300).sum()
        ) if "Resultado Bq/m3" in df_final else 0
        if exceeded:
            st.warning(
                f"⚠️ {exceeded} medición(es) superan el nivel de referencia de 300 Bq/m³."
            )
        elif not pendientes:
            st.success("Ninguna medición supera el nivel de referencia de 300 Bq/m³.")

    # =========================================================
    # 7. CONCLUSIONES
    # =========================================================
    completo_7 = pendientes == 0
    with _acordeon_informe("7", "CONCLUSIONES", completo_7):
        conclusion_default = generar_conclusion_automatica(df_final)
        _sincronizar_valor_auto(f"ic_conclusion_{cid}", conclusion_default)
        conclusion_manual = st.text_area(
            "Texto de conclusiones (editable)", height=200, key=f"ic_conclusion_{cid}",
        )
        conclusion_es_auto = st.session_state.get(f"ic_conclusion_{cid}__ultimo_auto") == conclusion_manual
        if not (punto3_es_auto and conclusion_es_auto and objeto_es_auto and texto4_es_auto and texto5_es_auto):
            st.caption(
                "✏️ Como has escrito alguno de los textos (objeto, punto 3, "
                "conclusiones...), ese texto se usará igual en gallego y en castellano si "
                "generas también la versión en castellano (no se traduce automáticamente)."
            )

    # =========================================================
    # 8. FIRMA
    # =========================================================
    tecnico_previo = str(tecnico_centro or det_meta.get("Técnico", "") or "")
    completo_8 = bool(st.session_state.get(f"ic_tecnico_{cid}", tecnico_previo).strip())
    with _acordeon_informe("8", "FIRMA", completo_8):
        tecnico_nome = st.text_input(
            "Nombre del/de la técnico/a que firma el informe",
            value=tecnico_previo, key=f"ic_tecnico_{cid}",
        )

    # =========================================================
    # 9. ANEXOS
    # =========================================================
    # Los anexos III y IV disponen de archivo por defecto salvo que el usuario
    # los haya eliminado. Los anexos I y II deben estar incorporados/generados.
    _anexo1_ok = bool(st.session_state.get(f"anexo1_{cid}_guardado"))
    _modo_anexo2_previo = st.session_state.get(
        f"ic_modo_anexo2_{cid}", "Generar automáticamente a partir del Excel"
    )
    if _modo_anexo2_previo == "Generar automáticamente a partir del Excel":
        _anexo2_ok = bool(st.session_state.get(f"ic_anexo2_guardado_{cid}"))
    else:
        _anexo2_ok = bool(st.session_state.get(f"anexo2_{cid}_guardado"))
    _anexo3_ok = bool(st.session_state.get(f"anexo3_{cid}_guardado")) or not st.session_state.get(f"anexo3_{cid}_eliminado", False)
    _anexo4_ok = bool(st.session_state.get(f"anexo4_{cid}_guardado")) or not st.session_state.get(f"anexo4_{cid}_eliminado", False)
    completo_9 = _anexo1_ok and _anexo2_ok and _anexo3_ok and _anexo4_ok

    with _acordeon_informe("9", "ANEXOS", completo_9):
        ANEXOS_INFO = [
            ("anexo1", "ANEXO I: FORMULARIOS TOMA DE DATOS"),
            ("anexo2", "ANEXO II: ESQUEMA GRÁFICO DO EDIFICIO E PLANOS DE CADA PLANTA"),
            ("anexo3", "ANEXO III: INFORME DE ENSAIO DO LABORATORIO ACREDITADO"),
            ("anexo4", "ANEXO IV: CERTIFICADO ENAC DO LABORATORIO ACREDITADO"),
        ]
        ANEXOS_POR_DEFECTO = {"anexo3": anexo3_por_defecto, "anexo4": anexo4_por_defecto}

        anexos_datos = {}

        def _titulo_anexo(texto):
            """Título de cada apartado de anexo (I, II, III, IV): mismo
            tamaño y en amarillo para los cuatro por igual.

            OJO: no vale con un "style" en línea con !important aquí (el
            propio st.markdown sanea el HTML y elimina los "!important" de
            los estilos en línea); por eso se usa el mismo patrón de
            "marcador" + regla CSS en el bloque <style> global de toda la
            app, que sí conserva el !important."""
            st.markdown('<div class="marcador-titulo-anexo"></div>', unsafe_allow_html=True)
            st.markdown(f'<p style="font-weight:700;">{html.escape(texto)}</p>', unsafe_allow_html=True)

        for key, label in ANEXOS_INFO:
            key_cid = f"{key}_{cid}"
            if key == "anexo2":
                _titulo_anexo(label)
                modo_anexo2 = st.radio(
                    "¿Cómo quieres aportar este anexo?",
                    options=["Generar automáticamente a partir del Excel", "Subir otro plano"],
                    key=f"ic_modo_anexo2_{cid}", horizontal=True, label_visibility="collapsed",
                )
                if modo_anexo2 == "Generar automáticamente a partir del Excel":
                    st.caption("El plano aprovechará siempre prácticamente toda la hoja. El logotipo y la foto exterior, si se seleccionan, se superponen abajo a la derecha: el logotipo abajo y la foto justo encima.")
                    col_anexo2_logo, col_anexo2_foto = st.columns(2)
                    with col_anexo2_logo:
                        mostrar_logo_anexo2 = st.checkbox(
                            "Mostrar logotipo (superpuesto)",
                            value=True,
                            key=f"ic_anexo2_mostrar_logo_{cid}",
                        )
                    with col_anexo2_foto:
                        mostrar_foto_anexo2 = st.checkbox(
                            "Mostrar foto exterior (superpuesta)",
                            value=True,
                            key=f"ic_anexo2_mostrar_foto_{cid}",
                        )

                    if st.button("🗺️ Generar plano automáticamente", key=f"ic_generar_anexo2_{cid}"):
                        try:
                            datos_planos = extraer_datos_planos(ruta_excel)

                            # Las superficies pertenecen al centro y están
                            # guardadas en SQLite. Se pasan expresamente al
                            # generador del Anexo II; antes no se estaban
                            # enviando y por eso no podían aparecer.
                            datos_informe_anexo2 = get_datos_informe_centro(cid)
                            sup_construida_anexo2 = datos_informe_anexo2.get("superficie_construida", "")
                            sup_util_anexo2 = datos_informe_anexo2.get("superficie_util", "")

                            anexo2_docx = generar_documento_anexo2(
                                datos_planos,
                                logo_bytes=logo_bytes,
                                centro=centro_nombre,
                                mostrar_logo=mostrar_logo_anexo2,
                                mostrar_foto_exterior=mostrar_foto_anexo2,
                                superficie_construida=sup_construida_anexo2,
                                superficie_util=sup_util_anexo2,
                            )
                            nombre_generado = f"ANEXO_II_planos_{selected_value}.docx"
                            st.session_state[f"ic_anexo2_guardado_{cid}"] = (nombre_generado, anexo2_docx)
                            # Miniaturas de cada plano YA con los puntos
                            # dibujados encima (la misma composición que se
                            # mete dentro del Anexo II), para poder verlas
                            # aquí mismo sin tener que abrir el .docx.
                            from utils_informe.anexo2 import componer_plano
                            miniaturas_planos = []
                            for clave_plano, info_plano in datos_planos.get("planos", {}).items():
                                puntos_plano = datos_planos.get("puntos", {}).get(clave_plano, [])
                                try:
                                    miniaturas_planos.append((
                                        info_plano.get("nombre", clave_plano),
                                        componer_plano(
                                            info_plano["imagen"],
                                            puntos_plano,
                                            logo_bytes=logo_bytes,
                                            foto_exterior=datos_planos.get("foto_exterior"),
                                            mostrar_logo=mostrar_logo_anexo2,
                                            mostrar_foto_exterior=mostrar_foto_anexo2,
                                            centro=centro_nombre,
                                            superficie_construida=sup_construida_anexo2,
                                            superficie_util=sup_util_anexo2,
                                        ),
                                    ))
                                except Exception:
                                    pass
                            st.session_state[f"ic_anexo2_miniaturas_{cid}"] = miniaturas_planos
                            st.success("Plano generado correctamente.")
                        except Anexo2Error as e:
                            st.error(f"No se ha podido generar el plano automáticamente: {e}")
                    if st.session_state.get(f"ic_anexo2_guardado_{cid}"):
                        nombre_guardado, _ = st.session_state[f"ic_anexo2_guardado_{cid}"]
                        anexos_datos["anexo2"] = st.session_state[f"ic_anexo2_guardado_{cid}"]
                        st.caption(f"📎 Plano disponible: **{nombre_guardado}**")
                        for nombre_plano_mini, imagen_mini in st.session_state.get(f"ic_anexo2_miniaturas_{cid}", []):
                            st.image(imagen_mini, width=200, caption=nombre_plano_mini)
                    else:
                        anexos_datos["anexo2"] = None
                    continue
                else:
                    st.session_state.pop(f"ic_anexo2_guardado_{cid}", None)
                    st.session_state.pop(f"ic_anexo2_miniaturas_{cid}", None)
                    anexos_datos[key] = _widget_archivo_con_eliminar(
                        f"ic_{key_cid}", "Sube el archivo",
                        ["pdf", "doc", "docx", "jpg", "jpeg", "png"],
                        valor_por_defecto=None,
                    )
                    continue

            default_fn = ANEXOS_POR_DEFECTO.get(key)
            valor_por_defecto_bytes = default_fn() if default_fn else None
            _titulo_anexo(label)
            anexos_datos[key] = _widget_archivo_con_eliminar(
                f"ic_{key_cid}", "Sube el archivo",
                ["pdf", "doc", "docx", "jpg", "jpeg", "png"],
                valor_por_defecto=valor_por_defecto_bytes,
            )

    # =========================================================
    # GENERACIÓN DEL INFORME (sin número)
    # =========================================================
    with _acordeon_informe("", "GENERACIÓN DEL INFORME", True):
        col_opt1, col_opt2 = st.columns(2)
        with col_opt1:
            st.markdown('<div class="marcador-checkbox-anexo2"></div>', unsafe_allow_html=True)
            generar_castellano = st.checkbox("Generar versión en castellano", key=f"ic_castellano_{cid}")
        with col_opt2:
            pdf_disponible = libreoffice_disponible()
            generar_pdf_completo = st.checkbox(
                "Generar PDF completo (informe + anexos)",
                disabled=not pdf_disponible, key=f"ic_pdf_completo_{cid}",
            )
            if not pdf_disponible:
                st.caption("⚠️ LibreOffice no está disponible en este entorno; no se puede generar el PDF completo.")

        st.markdown('<div class="marcador-btn-guardar-detector"></div>', unsafe_allow_html=True)
        if st.button("📄 Generar informe", type="primary", use_container_width=True, key=f"ic_generar_{cid}"):
            try:
                ctx_gl = ReportContext(
                    xerencia=xerencia, cif=cif, centro=traducir_es_gl(centro_nombre), servizo_unidade=traducir_es_gl(servizo_unidade_efectivo),
                    enderezo=enderezo, superficie_construida=superficie_construida, superficie_util=superficie_util,
                    num_plantas=num_plantas, postos_bullets=postos_bullets, num_traballadores_total=num_traballadores_total,
                    categorias_bullets=categorias_bullets, texto_punto3_manual=texto_punto3_manual,
                    data_informacion_traballadores=data_informacion_traballadores,
                    medio_informacion_traballadores=medio_informacion_traballadores,
                    data_informe=data_informe.strftime("%d/%m/%Y"), tecnico_nome=tecnico_nome,
                    incertezas_por_defecto="", conclusion_manual=conclusion_manual, logo=logo_bytes,
                    tipo_zona=tipo_zona_elegido, objeto_manual=objeto_manual,
                    texto4_manual=texto4_manual, texto5_manual=texto5_manual,
                )
                buffer_gl = generate_report(ctx_gl, df_final, idioma="gl")
                buffer_gl = io.BytesIO(_poner_salas_negrita_punto3(buffer_gl.getvalue(), salas))
                report_name_gl = _nombre_documento(selected_value, "INFORME-FINAL", "-gl") + ".docx"
                st.session_state[f"ic_report_gl_{cid}"] = (buffer_gl.getvalue(), report_name_gl)

                docx_buffers_para_zip = [(report_name_gl, buffer_gl.getvalue())]
                docx_buffers_para_pdf = [("gl", report_name_gl, buffer_gl.getvalue())]

                if generar_castellano:
                    postos_bullets_es = postos_traballo_bullets(
                        df_center, traducir_galego=False, es_atencion_primaria=es_ap_centro,
                    )
                    _, categorias_bullets_es = categorias_turnos_bullets(
                        categorias_df, traducir_galego=False, es_atencion_primaria=es_ap_centro,
                    )

                    # El objeto, el punto 3 y las conclusiones son
                    # cuadros de texto editables (uno solo, no uno por
                    # idioma): si el usuario no los ha tocado (siguen
                    # en modo automático), se recalculan aquí en
                    # castellano en vez de reutilizar el texto en
                    # gallego tal cual, que es lo que causaba que
                    # saliera parte del informe en el idioma que no
                    # tocaba. Si el usuario SÍ escribió algo, se
                    # deja igual (no hay forma de saber en qué idioma
                    # lo escribió).
                    objeto_es = objeto_manual
                    if objeto_es_auto:
                        _ctx_previa_objeto_es = ReportContext(
                            centro=centro_nombre, servizo_unidade=servizo_unidade_efectivo, tipo_zona=tipo_zona_elegido,
                        )
                        objeto_es = generar_texto_objeto_automatico(_ctx_previa_objeto_es, idioma="es")
                        objeto_es = _normalizar_referencia_zona_punto2(
                            objeto_es, servizo_unidade_efectivo, tipo_zona_elegido, idioma="es"
                        )

                    texto_punto3_es = texto_punto3_manual
                    if punto3_es_auto:
                        _ctx_previa_punto3_es = ReportContext(
                            centro=centro_nombre, servizo_unidade=servizo_unidade_efectivo, postos_bullets=postos_bullets_es,
                            num_traballadores_total=num_traballadores_total, categorias_bullets=categorias_bullets_es,
                            medio_informacion_traballadores=medio_informacion_traballadores,
                            data_informacion_traballadores=data_informacion_traballadores,
                            tipo_zona=tipo_zona_elegido,
                        )
                        texto_punto3_es = generar_texto_punto3_automatico(_ctx_previa_punto3_es, idioma="es")

                    conclusion_es = conclusion_manual
                    if conclusion_es_auto:
                        conclusion_es = generar_conclusion_automatica(df_final, idioma="es")

                    texto4_es = generar_texto4_automatico("es") if texto4_es_auto else texto4_manual
                    texto5_es = generar_texto5_automatico("es") if texto5_es_auto else texto5_manual

                    ctx_es = ReportContext(
                        xerencia=xerencia, cif=cif, centro=centro_nombre, servizo_unidade=servizo_unidade_efectivo,
                        enderezo=enderezo, superficie_construida=superficie_construida, superficie_util=superficie_util,
                        num_plantas=num_plantas, postos_bullets=postos_bullets_es, num_traballadores_total=num_traballadores_total,
                        categorias_bullets=categorias_bullets_es, texto_punto3_manual=texto_punto3_es,
                        data_informacion_traballadores=data_informacion_traballadores,
                        medio_informacion_traballadores=medio_informacion_traballadores,
                        data_informe=data_informe.strftime("%d/%m/%Y"), tecnico_nome=tecnico_nome,
                        incertezas_por_defecto="", conclusion_manual=conclusion_es, logo=logo_bytes,
                        tipo_zona=tipo_zona_elegido, objeto_manual=objeto_es,
                        texto4_manual=texto4_es, texto5_manual=texto5_es,
                    )
                    buffer_es = generate_report(ctx_es, df_final, idioma="es")
                    buffer_es = io.BytesIO(_poner_salas_negrita_punto3(buffer_es.getvalue(), salas))
                    report_name_es = _nombre_documento(selected_value, "INFORME-FINAL", "-es") + ".docx"
                    st.session_state[f"ic_report_es_{cid}"] = (buffer_es.getvalue(), report_name_es)
                    docx_buffers_para_zip.append((report_name_es, buffer_es.getvalue()))
                    docx_buffers_para_pdf.append(("es", report_name_es, buffer_es.getvalue()))
                else:
                    st.session_state.pop(f"ic_report_es_{cid}", None)

                anexos_subidos = {k: v for k, v in anexos_datos.items() if v is not None}
                numeros = {"anexo1": "I", "anexo2": "II", "anexo3": "III", "anexo4": "IV"}

                if anexos_subidos:
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                        for nombre_docx, contenido_docx in docx_buffers_para_zip:
                            zf.writestr(nombre_docx, contenido_docx)
                        for key, (nombre_original, contenido) in anexos_subidos.items():
                            ext = nombre_original.rsplit(".", 1)[-1] if "." in nombre_original else "pdf"
                            zf.writestr(f"ANEXO_{numeros[key]}.{ext}", contenido)
                    nombre_zip_final = _nombre_documento(selected_value, "INFORME-FINAL", "-con-anexos") + ".zip"
                    st.session_state[f"ic_zip_{cid}"] = (zip_buffer.getvalue(), nombre_zip_final)
                else:
                    st.session_state.pop(f"ic_zip_{cid}", None)

                if generar_pdf_completo and pdf_disponible:
                    anexos_para_pdf = [(n, c) for (n, c) in anexos_subidos.values()]
                    for idioma_pdf, nombre_docx, contenido_docx in docx_buffers_para_pdf:
                        try:
                            pdf_bytes = construir_pdf_completo(contenido_docx, anexos_para_pdf)
                            pdf_name = nombre_docx.rsplit(".", 1)[0] + "_completo.pdf"
                            st.session_state[f"ic_pdf_{idioma_pdf}_{cid}"] = (pdf_bytes, pdf_name)
                        except PdfToolsError as e:
                            st.error(f"No se ha podido generar el PDF completo ({idioma_pdf}): {e}")
                else:
                    st.session_state.pop(f"ic_pdf_gl_{cid}", None)
                    st.session_state.pop(f"ic_pdf_es_{cid}", None)

                st.success("Informe generado correctamente")
            except Exception as e:
                st.error(f"Error al generar el informe: {e}")

        for key_ss, etiqueta, mime in [
            (f"ic_report_gl_{cid}", "⬇️ Descargar informe en gallego (.docx)",
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            (f"ic_report_es_{cid}", "⬇️ Descargar informe en castellano (.docx)",
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            (f"ic_pdf_gl_{cid}", "⬇️ Descargar PDF completo en gallego", "application/pdf"),
            (f"ic_pdf_es_{cid}", "⬇️ Descargar PDF completo en castellano", "application/pdf"),
            (f"ic_zip_{cid}", "⬇️ Descargar informe(s) + anexos (.zip)", "application/zip"),
        ]:
            if key_ss in st.session_state:
                contenido, nombre_archivo = st.session_state[key_ss]
                st.markdown('<div class="marcador-btn-descarga-amarillo"></div>', unsafe_allow_html=True)
                st.download_button(etiqueta, data=contenido, file_name=nombre_archivo, mime=mime, key=f"dl_{key_ss}")


def pantalla_centro_informes():
    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.session_state.view = "inicio"
        st.rerun()
        return
    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro
    ns_centro = f"centro_{cid}"

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    if st.button("← Volver"):
        st.session_state.view = "centro"
        st.rerun()

    _zona_titulo_actual = _zona_para_titulo_centro(cid, zona)
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + _zona_titulo_actual if _zona_titulo_actual and _zona_titulo_actual.strip() else ""}'
    )
    st.markdown('<p class="subtitulo-amarillo">📄 Informes y descargas</p>', unsafe_allow_html=True)

    detectores = fetch_detectores(cid)
    planos_centro = fetch_planos_centro(cid)

    with st.container(border=True):
      if not detectores:
          st.caption("Añade al menos un detector para poder generar los archivos.")
      else:
          # Recopilar todas las fotos del informe (imagen exterior del
          # centro y plano-con-punto/situación/detector de cada uno).
          # Se usa tanto para el ZIP de descarga como para la checklist
          # de WhatsApp.
          fotos_disponibles = []
          if img_path and os.path.exists(img_path):
              ext = os.path.splitext(img_path)[1] or ".jpg"
              fotos_disponibles.append({
                  "ruta": img_path,
                  "nombre_archivo": f"exterior_{_slug(nombre)}{ext}",
                  "etiqueta": "Foto exterior del centro",
              })
          # Los planos se listan una sola vez cada uno (se comparten
          # entre varios detectores, así que no se repiten).
          for plano_c in planos_centro:
              _, _, nombre_plano_c, ruta_plano_c, _ = plano_c
              if ruta_plano_c and os.path.exists(ruta_plano_c):
                  ext = os.path.splitext(ruta_plano_c)[1] or ".jpg"
                  fotos_disponibles.append({
                      "ruta": ruta_plano_c,
                      "nombre_archivo": f"plano_{_slug(nombre_plano_c)}{ext}",
                      "etiqueta": f"Plano — {nombre_plano_c}",
                  })
          for d in detectores:
              did_d, codigo_d = d[0], d[5]
              foto_sit_d, foto_det_d = d[9], d[10]
              punto_x_d, punto_y_d = d[7], d[8]
              plano_centro_id_d = d[17]  # columna plano_centro_id (ver orden de init_db)
              base_id = codigo_d or f"Detector {did_d}"

              # Plano de ESTE detector, con su propio punto rojo dibujado
              # encima (no la imagen del plano "en blanco").
              if (plano_centro_id_d and punto_x_d is not None and punto_y_d is not None
                      and punto_x_d >= 0 and punto_y_d >= 0):
                  plano_info_d = get_plano_centro(plano_centro_id_d)
                  if plano_info_d and os.path.exists(plano_info_d[3]):
                      ruta_tmp_punto = os.path.join(
                          get_data_dir(), f"_tmp_plano_punto_det{did_d}.jpg"
                      )
                      if generar_plano_con_punto(plano_info_d[3], punto_x_d, punto_y_d, ruta_tmp_punto):
                          fotos_disponibles.append({
                              "ruta": ruta_tmp_punto,
                              "nombre_archivo": _nombre_foto_plano(codigo_d, nombre, zona) + ".jpg",
                              "etiqueta": f"Plano con punto — {base_id}",
                          })

              if foto_sit_d and os.path.exists(foto_sit_d):
                  ext = os.path.splitext(foto_sit_d)[1] or ".jpg"
                  fotos_disponibles.append({
                      "ruta": foto_sit_d, "nombre_archivo": _nombre_foto_situacion(codigo_d, nombre, zona) + ext,
                      "etiqueta": f"Foto situación — {base_id}",
                  })
              if foto_det_d and os.path.exists(foto_det_d):
                  ext = os.path.splitext(foto_det_d)[1] or ".jpg"
                  fotos_disponibles.append({
                      "ruta": foto_det_d, "nombre_archivo": _nombre_foto_detector(codigo_d, nombre, zona) + ext,
                      "etiqueta": f"Foto detector — {base_id}",
                  })

          # --- Acceso al generador de informe completo (Word/PDF con el
          # modelo oficial UPRL/SERGAS), en una pantalla aparte. ---
          st.markdown('<div class="marcador-btn-informe-completo"></div>', unsafe_allow_html=True)
          if st.button("📝 INFORME DE RESULTADOS", use_container_width=True):
              st.session_state.view = "centro_informe_completo"
              st.rerun()
          st.caption(
              "Genera el informe oficial completo (identificación del centro, "
              "trabajadores expuestos, resultados y anexos) a partir de los datos "
              "de este centro, en formato Word."
          )
          st.markdown("---")

          # --- Checklist de qué documentos generar ---
          st.markdown(
              '<p class="subtitulo-amarillo">OTROS DOCUMENTOS PARA DESCARGA</p>',
              unsafe_allow_html=True,
          )
          gen_pdf = st.checkbox("Informe de colocación de detectores (PDF)", value=False, key=f"gen_chk_pdf_{cid}")
          gen_excel = st.checkbox("Hoja de cálculo de datos (XLS)", value=False, key=f"gen_chk_excel_{cid}")
          gen_fotos = st.checkbox("Fotos (JPG o ZIP)", value=False, key=f"gen_chk_fotos_{cid}")
          gen_lab = st.checkbox("Informe de registro para laboratorio (PDF)", value=False, key=f"gen_chk_lab_{cid}")
          tipo_firma_lab = "digital"
          if gen_lab:
              tipo_firma_lab = st.radio(
                  "Tipo de firma del técnico",
                  options=["digital", "manual"],
                  format_func=lambda v: "Firma digital" if v == "digital" else "Firma",
                  key=f"gen_firma_tipo_{cid}", horizontal=True,
              )
              logo_lab_actual = get_logo_laboratorio()
              if logo_lab_actual and os.path.exists(logo_lab_actual):
                  st.image(logo_lab_actual, width=180)
              nuevo_logo_lab = st.file_uploader(
                  "Logotipo del laboratorio, sube un archivo para cambiarlo",
                  type=["png", "jpg", "jpeg"], key=f"logo_lab_up_{cid}",
              )
              if nuevo_logo_lab is not None:
                  fid_logo = getattr(nuevo_logo_lab, "file_id", None) or \
                      f"{nuevo_logo_lab.name}_{nuevo_logo_lab.size}"
                  if st.session_state.get(f"logo_lab_last_fid_{cid}") != fid_logo:
                      ruta_logo_guardada = guardar_bytes_imagen(
                          nuevo_logo_lab.getvalue(), "logo_laboratorio", extension_de(nuevo_logo_lab),
                      )
                      set_logo_laboratorio(ruta_logo_guardada)
                      st.session_state[f"logo_lab_last_fid_{cid}"] = fid_logo
                      st.success("Logotipo del laboratorio actualizado")
                      st.rerun()

          # --- Un único botón que genera lo marcado, pequeño y
          # alineado a la izquierda ---
          col_gen, _col_resto = st.columns([1, 3])
          with col_gen:
              st.markdown('<div class="marcador-btn-generar"></div>', unsafe_allow_html=True)
              generar_clic = st.button(
                  "📦 Generar", type="tertiary", use_container_width=True,
              )

          if generar_clic:
              if not (gen_pdf or gen_excel or gen_fotos or gen_lab):
                  st.warning("Marca al menos un documento para generar.")
              else:
                  try:
                      marca_tiempo = _ahora_espana().strftime('%Y%m%d_%H%M%S')
                      with st.spinner("Generando documentos..."):
                          if gen_pdf:
                              nombre_pdf = _nombre_documento(nombre, "INFORME-COLOCACIÓN") + ".pdf"
                              ruta_pdf = os.path.join(get_data_dir(), nombre_pdf)
                              generar_pdf(cid, ruta_pdf)
                              st.session_state["ultimo_pdf"] = ruta_pdf
                              st.session_state["ultimo_pdf_nombre"] = nombre_pdf
                              st.session_state["ultimo_pdf_centro"] = cid

                          if gen_excel:
                              nombre_xlsx = _nombre_documento(nombre, "HOJA-DATOS") + ".xlsx"
                              ruta_xlsx = os.path.join(get_data_dir(), nombre_xlsx)
                              generar_excel(cid, ruta_xlsx)
                              st.session_state["ultimo_excel"] = ruta_xlsx
                              st.session_state["ultimo_excel_nombre"] = nombre_xlsx
                              st.session_state["ultimo_excel_centro"] = cid

                          if gen_fotos:
                              nombre_zip = _nombre_documento(nombre, "FOTOS") + ".zip"
                              ruta_zip = os.path.join(get_data_dir(), nombre_zip)
                              modo_fotos_gen = st.session_state.get(f"modo_fotos_{cid}", "todas")
                              if modo_fotos_gen == "individual":
                                  fotos_a_incluir = [
                                      foto for i, foto in enumerate(fotos_disponibles)
                                      if st.session_state.get(
                                          f"wa_chk_{cid}_{i}_{_slug(foto['nombre_archivo'])}")
                                  ]
                              else:
                                  fotos_a_incluir = fotos_disponibles
                              with zipfile.ZipFile(ruta_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                                  for foto in fotos_a_incluir:
                                      if os.path.exists(foto["ruta"]):
                                          zf.write(foto["ruta"], arcname=foto["nombre_archivo"])
                              st.session_state["ultimo_zip_fotos"] = ruta_zip
                              st.session_state["ultimo_zip_fotos_nombre"] = nombre_zip
                              st.session_state["ultimo_zip_fotos_centro"] = cid

                          if gen_lab:
                              nombre_lab = _nombre_documento(nombre, "REGISTRO-LABORATORIO") + ".pdf"
                              ruta_lab = os.path.join(get_data_dir(), nombre_lab)
                              generar_registro_laboratorio(cid, ruta_lab, tipo_firma=tipo_firma_lab)
                              st.session_state["ultimo_lab"] = ruta_lab
                              st.session_state["ultimo_lab_nombre"] = nombre_lab
                              st.session_state["ultimo_lab_centro"] = cid

                      st.success("Documentos generados correctamente")
                  except Exception as e:
                      st.error(f"Error al generar los documentos: {e}")

          hay_pdf = st.session_state.get("ultimo_pdf_centro") == cid
          hay_excel = st.session_state.get("ultimo_excel_centro") == cid
          hay_zip = st.session_state.get("ultimo_zip_fotos_centro") == cid
          hay_lab = st.session_state.get("ultimo_lab_centro") == cid

          if hay_pdf or hay_excel or hay_zip or hay_lab:
              ultimo_pdf = st.session_state.get("ultimo_pdf")
              ultimo_excel = st.session_state.get("ultimo_excel")
              ultimo_zip = st.session_state.get("ultimo_zip_fotos")
              ultimo_lab = st.session_state.get("ultimo_lab")

              # 1) Descargar PDF
              if hay_pdf and ultimo_pdf and os.path.exists(ultimo_pdf):
                  with open(ultimo_pdf, "rb") as f:
                      pdf_bytes = f.read()
                  st.markdown('<div class="marcador-btn-descarga-amarillo"></div>', unsafe_allow_html=True)
                  st.download_button(
                      "Descargar informe de colocación de detectores (PDF)", data=pdf_bytes,
                      file_name=st.session_state.get("ultimo_pdf_nombre", "informe.pdf"),
                      mime="application/pdf", use_container_width=True,
                      icon=":material/picture_as_pdf:",
                  )

              # 2) Enviar PDF por WhatsApp
              if hay_pdf and ultimo_pdf and os.path.exists(ultimo_pdf):
                  texto_wa = f"Informe de colocación de detectores de Rn – {nombre or ''}"
                  boton_compartir_whatsapp(
                      ultimo_pdf,
                      st.session_state.get("ultimo_pdf_nombre", "informe.pdf"),
                      texto_wa,
                  )

              # 3) Descargar Excel
              if hay_excel and ultimo_excel and os.path.exists(ultimo_excel):
                  with open(ultimo_excel, "rb") as f:
                      excel_bytes = f.read()
                  st.markdown('<div class="marcador-btn-descarga-amarillo"></div>', unsafe_allow_html=True)
                  st.download_button(
                      "Descargar hoja de cálculo de datos (XLS)", data=excel_bytes,
                      file_name=st.session_state.get("ultimo_excel_nombre", "detectores.xlsx"),
                      mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                      use_container_width=True,
                      icon=":material/grid_on:",
                  )

              # 4) Enviar Excel por WhatsApp
              if hay_excel and ultimo_excel and os.path.exists(ultimo_excel):
                  texto_wa_excel = f"Hoja de cálculo de detectores de Rn – {nombre or ''}"
                  boton_compartir_whatsapp_excel(
                      ultimo_excel,
                      st.session_state.get("ultimo_excel_nombre", "detectores.xlsx"),
                      texto_wa_excel,
                  )

              # 5) Descargar Registro para laboratorio
              if hay_lab and ultimo_lab and os.path.exists(ultimo_lab):
                  with open(ultimo_lab, "rb") as f:
                      lab_bytes = f.read()
                  st.markdown('<div class="marcador-btn-descarga-amarillo"></div>', unsafe_allow_html=True)
                  st.download_button(
                      "Descargar informe de registro para laboratorio (PDF)", data=lab_bytes,
                      file_name=st.session_state.get("ultimo_lab_nombre", "registro_laboratorio.pdf"),
                      mime="application/pdf", use_container_width=True,
                      icon=":material/picture_as_pdf:",
                  )

              # 6) Enviar Registro para laboratorio por WhatsApp
              if hay_lab and ultimo_lab and os.path.exists(ultimo_lab):
                  texto_wa_lab = f"Registro para laboratorio – {nombre or ''}"
                  boton_compartir_whatsapp(
                      ultimo_lab,
                      st.session_state.get("ultimo_lab_nombre", "registro_laboratorio.pdf"),
                      texto_wa_lab,
                      id_sufijo="lab",
                      etiqueta_boton="Enviar por WhatsApp",
                  )

              # 7) Fotos: por defecto se incluyen todas; el botón
              # "Selección individual" muestra la checklist para
              # elegir solo algunas (y así no alargar la pantalla
              # con el listado completo salvo que se pida).
              seleccionadas = []
              if fotos_disponibles:
                  st.markdown(
                      '<p class="subtitulo-amarillo">📷 Fotos (JPG o ZIP)</p>',
                      unsafe_allow_html=True,
                  )

                  modo_key = f"modo_fotos_{cid}"
                  if modo_key not in st.session_state:
                      st.session_state[modo_key] = "todas"
                  modo_actual = st.session_state[modo_key]

                  st.markdown('<div class="marcador-btn-descarga-amarillo"></div>', unsafe_allow_html=True)
                  etiqueta_toggle = (
                      "🔎 Selección individual" if modo_actual == "todas" else "☑️ Seleccionar todos"
                  )
                  if st.button(etiqueta_toggle, key=f"toggle_modo_fotos_{cid}", use_container_width=True):
                      if modo_actual == "todas":
                          st.session_state[modo_key] = "individual"
                      else:
                          st.session_state[modo_key] = "todas"
                          for i, foto in enumerate(fotos_disponibles):
                              st.session_state[
                                  f"wa_chk_{cid}_{i}_{_slug(foto['nombre_archivo'])}"
                              ] = True
                      st.rerun()

                  if st.session_state[modo_key] == "individual":
                      for i, foto in enumerate(fotos_disponibles):
                          marcado = st.checkbox(
                              foto["etiqueta"],
                              key=f"wa_chk_{cid}_{i}_{_slug(foto['nombre_archivo'])}",
                          )
                          if marcado:
                              seleccionadas.append((foto["ruta"], foto["nombre_archivo"]))
                  else:
                      st.caption(
                          f"Se incluirán las {len(fotos_disponibles)} fotos disponibles. "
                          "Pulsa «Selección individual» para elegir solo algunas."
                      )
                      seleccionadas = [(f["ruta"], f["nombre_archivo"]) for f in fotos_disponibles]

              # 8) Descargar fotos
              if hay_zip and ultimo_zip and os.path.exists(ultimo_zip):
                  with open(ultimo_zip, "rb") as f:
                      zip_bytes = f.read()
                  st.markdown('<div class="marcador-btn-descarga-amarillo"></div>', unsafe_allow_html=True)
                  st.download_button(
                      "⬇️ Descargar fotos (ZIP)", data=zip_bytes,
                      file_name=st.session_state.get("ultimo_zip_fotos_nombre", "fotos.zip"),
                      mime="application/zip", use_container_width=True,
                  )

              # 9) Enviar fotos seleccionadas por WhatsApp
              if fotos_disponibles:
                  texto_wa_fotos = f"Fotos del informe – {nombre or ''}"
                  boton_compartir_whatsapp_fotos(seleccionadas, texto_wa_fotos)



def pantalla_centro():
    cid = st.session_state.centro_actual
    centro = get_centro(cid) if cid else None
    if not centro:
        st.error("Centro no encontrado.")
        st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
        if st.button("← Inicio"):
            st.session_state.view = "inicio"
            st.rerun()
        return

    cid, nombre, zona, fecha, img_path, tecnico_centro, direccion = centro

    # En Atención Primaria, Atención Primaria + PAC, PAC y Consultorio,
    # el título principal muestra únicamente el nombre del centro.
    tipo_centro_titulo = (get_tipo_centro(cid) or "").strip().casefold()
    tipos_solo_nombre = {
        "atención primaria",
        "atencion primaria",
        "atención primaria + pac",
        "atencion primaria + pac",
        "pac",
        "consultorio",
    }
    zona_titulo = _zona_para_titulo_centro(cid, zona)

    if st.session_state.get("_centro_snapshot_id") != cid:
        _registrar_entrada_centro(cid)

    # El botón "Datos de la empresa" también indica si quedan datos
    # obligatorios por cubrir. El logotipo no cuenta como obligatorio,
    # ya que la app utiliza el logotipo por defecto si no se carga uno.
    empresa_actual = get_empresa()
    cif_actual = get_cif()
    datos_empresa_completos = bool(
        str(empresa_actual or "").strip()
        and str(cif_actual or "").strip()
        and str(tecnico_centro or "").strip()
    )
    etiqueta_datos_empresa = (
        "Datos de la empresa"
        if datos_empresa_completos
        else "🔴 Datos de la empresa"
    )

    top1, top2 = st.columns([3, 1])
    with top1:
        st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
        if st.button("← Inicio"):
            if _centro_modificado_desde_entrada(cid):
                st.session_state["_preguntar_guardado_centro"] = cid
            else:
                st.session_state.pop("_preguntar_guardado_centro", None)
            st.session_state.view = "inicio"
            st.rerun()
    with top2:
        if st.button(etiqueta_datos_empresa, use_container_width=True):
            st.session_state.view = "ajustes"
            st.rerun()

    st.html(
        '<div style="color:#FFFFFF;font-size:1.15rem;line-height:1.2;font-weight:600;'
        'margin:0.15rem 0 0.35rem 0;">Gestión de muestreo</div>'
    )
    _render_titulo_principal(
        f'🏢 {nombre or ""}{" · " + zona_titulo if zona_titulo and zona_titulo.strip() else ""}'
    )

    # Estado de cada apartado para mostrar solamente un punto rojo en el
    # botón correspondiente cuando quedan datos por cubrir. El aspecto del
    # botón no se modifica.
    categorias_centro = fetch_categorias_centro(cid)
    planos_centro = fetch_planos_centro(cid)
    detectores_centro = fetch_detectores(cid)

    datos_centro_completos = bool(
        str(nombre or "").strip() and str(zona or "").strip()
        and str(fecha or "").strip() and str(direccion or "").strip()
        and str(img_path or "").strip()
    )
    categorias_completas = bool(categorias_centro) and all(
        str(fila[2] or "").strip() and int(fila[3] or 0) > 0
        for fila in categorias_centro
    )
    planos_completos = bool(planos_centro) and all(
        str(fila[2] or "").strip() and str(fila[3] or "").strip()
        for fila in planos_centro
    )

    def _detector_colocacion_completa(d):
        try:
            punto_ok = d[7] is not None and d[8] is not None and 0 <= float(d[7]) <= 1 and 0 <= float(d[8]) <= 1
            return bool(
                str(d[3] or "").strip() and str(d[4] or "").strip()
                and str(d[5] or "").strip() and punto_ok
                and str(d[9] or "").strip() and str(d[10] or "").strip()
                and str(d[12] or "").strip() and str(d[13] or "").strip()
                and str(d[14] or "").strip() and str(d[15] or "").strip()
                and str(d[16] or "").strip() and d[17] is not None
            )
        except Exception:
            return False

    detectores_completos = bool(detectores_centro) and all(
        _detector_colocacion_completa(d) for d in detectores_centro
    )
    retirada_completa = bool(detectores_centro) and all(
        len(d) > 19 and str(d[18] or "").strip() and str(d[19] or "").strip()
        for d in detectores_centro
    )
    informes_completos = bool(detectores_centro) and all(
        len(d) > 20 and d[20] is not None for d in detectores_centro
    ) and bool(str(tecnico_centro or "").strip())

    opciones_menu = [
        ("Datos del centro", "centro_datos", datos_centro_completos),
        ("Categorías profesionales", "centro_categorias", categorias_completas),
        ("Planos del centro", "centro_planos", planos_completos),
        ("Detectores colocados", "centro_detectores", detectores_completos),
        ("Retirada de detectores", "centro_retirada", retirada_completa),
        ("📄 Informes y descargas", "centro_informes", informes_completos),
    ]
    for etiqueta, vista, completo in opciones_menu:
        etiqueta_boton = etiqueta if completo else f"🔴 {etiqueta}"
        if st.button(etiqueta_boton, key=f"menu_{vista}_{cid}", use_container_width=True):
            st.session_state.view = vista
            st.rerun()




# ============================================================
# PANTALLA: DETECTOR (nuevo / editar)
# ============================================================

TURNOS_TRABAJO_OPCIONES = ["Mañana", "Tarde", "Noche", "PAC", "Rotatorio", "Rotatorio complejo"]
# Turno a nivel de categoría profesional (independiente del turno por
# sala/detector): solo se usa para la segunda parte del punto 3 del
# informe ("O número de traballadores... divididos nas seguintes
# categorías"), sin relación con los turnos de los detectores.
TURNOS_CATEGORIA_OPCIONES = [
    "Mañana", "Tarde", "Noche", "Mañana/Tarde", "Mañana/Tarde/Noche",
    "Horario PAC", "Rotatorio", "Rotatorio complejo",
]
NIVEL_OPCIONES = [
    "3 niveles bajo rasante (Sótano -3)",
    "2 niveles bajo rasante (Sótano -2)",
    "1 nivel bajo rasante (Sótano -1)",
    "Nivel de la rasante (Planta Baja)",
    "1 nivel sobre rasante",
    "2 niveles sobre rasante",
    "3 niveles sobre rasante",
]
# Segundo bloque del código de sala (p.ej. "S-1", "PB", "02"): según
# la opción elegida en "Nivel".
NIVEL_A_CODIGO = {
    "3 niveles bajo rasante (Sótano -3)": "S-3",
    "2 niveles bajo rasante (Sótano -2)": "S-2",
    "1 nivel bajo rasante (Sótano -1)": "S-1",
    "Nivel de la rasante (Planta Baja)": "PB",
    "1 nivel sobre rasante": "01",
    "2 niveles sobre rasante": "02",
    "3 niveles sobre rasante": "03",
}
# Columna "Planta" de la ficha de registro para laboratorio: número
# de planta (con signo) que corresponde a cada opción de "Nivel".
NIVEL_A_PLANTA_LABORATORIO = {
    "3 niveles bajo rasante (Sótano -3)": "-3",
    "2 niveles bajo rasante (Sótano -2)": "-2",
    "1 nivel bajo rasante (Sótano -1)": "-1",
    "Nivel de la rasante (Planta Baja)": "0",
    "1 nivel sobre rasante": "1",
    "2 niveles sobre rasante": "2",
    "3 niveles sobre rasante": "3",
}
TIPO_CENTRO_OPCIONES = [
    "Atención Primaria", "PAC", "Atención Primaria + PAC", "Consultorio",
    "Centro de especialidades", "Hospital", "Otro",
]
# Para estos cuatro, la casilla "Área / Zona :" se rellena sola; para el
# resto (incluido "Otro") se deja en blanco para que se escriba.
# "Consultorio" es un caso especial: su Área/Zona es "Atención
# Primaria" (no "Consultorio", que solo se usa como tipo de centro
# para el prefijo del código de sala).
TIPO_CENTRO_A_AREA_AUTOMATICA = {
    "Atención Primaria": "Atención Primaria",
    "PAC": "PAC",
    "Atención Primaria + PAC": "Atención Primaria + PAC",
    "Consultorio": "Atención Primaria",
}


def pantalla_detector():
    cid = st.session_state.centro_actual
    detector_id = st.session_state.detector_actual
    ns = f"det_{detector_id or 'nuevo'}"
    st.session_state["detector_form_ns"] = ns

    _inicializar_ns_detector(cid, detector_id, ns)

    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    if st.button("← Volver"):
        if _detector_tiene_cambios(ns):
            guardado_id = _guardar_detector_desde_ns(
                cid, detector_id, ns, mostrar_mensajes=False
            )
            if not guardado_id:
                return
        _limpiar_namespace(ns)
        st.session_state.view = "centro_detectores"
        st.rerun()

    texto_titulo_detector = "Datos del detector" if detector_id else "Nuevo detector"
    _render_titulo_principal(texto_titulo_detector)

    _renderizar_campos_detector(cid, detector_id, ns)


# ============================================================

def pantalla_ajustes():
    cid = st.session_state.get("centro_actual")
    centro = get_centro(cid) if cid else None

    _render_titulo_principal("Datos de la empresa")
    st.markdown('<div class="marcador-btn-azul-claro"></div>', unsafe_allow_html=True)
    volver_clic_ajustes = st.button("← Volver")

    if not centro:
        if volver_clic_ajustes:
            st.session_state.view = "inicio"
            st.rerun()
        st.error("No hay ningún centro seleccionado.")
        return

    _, nombre_centro, _, _, _, tecnico_actual, _ = centro
    st.caption(f"Centro: {nombre_centro}")

    empresa_actual = get_empresa()
    cif_actual = get_cif()

    # Compatibilidad con instalaciones anteriores y selección por defecto.
    if empresa_actual not in AREAS_SANITARIAS_GALICIA:
        empresa_actual = AREA_SANITARIA_CORUNA

    if "ajustes_empresa" not in st.session_state:
        st.session_state["ajustes_empresa"] = empresa_actual
    if "ajustes_cif" not in st.session_state:
        st.session_state["ajustes_cif"] = (
            cif_actual
            if cif_actual
            else AREAS_SANITARIAS_GALICIA.get(empresa_actual, "")
        )

    def _cambiar_area_sanitaria():
        area = st.session_state.get("ajustes_empresa", AREA_SANITARIA_CORUNA)
        st.session_state["ajustes_cif"] = AREAS_SANITARIAS_GALICIA.get(area, "")
        # Si se había pulsado Eliminar sobre un logo por defecto, al cambiar de
        # área mostramos de nuevo el logo por defecto correspondiente.
        st.session_state["ajustes_logo_eliminado"] = False

    st.selectbox(
        "Área sanitaria",
        options=list(AREAS_SANITARIAS_GALICIA.keys()),
        key="ajustes_empresa",
        on_change=_cambiar_area_sanitaria,
    )
    st.text_input(
        "CIF",
        key="ajustes_cif",
        help="Se actualiza automáticamente al seleccionar el área sanitaria, pero puedes modificarlo manualmente.",
    )
    st.text_input(
        "Técnico (aparece en el PDF de este centro)",
        value=tecnico_actual or "",
        key="ajustes_tecnico",
    )

    st.markdown("**Logotipo** (aparece en el informe final y en el Informe PDF de colocación de detectores)")
    logo_bytes_actual, logo_nombre_actual = get_logo_informe(
        empresa_override=st.session_state.get("ajustes_empresa", empresa_actual)
    )
    _widget_archivo_con_eliminar(
        "ajustes_logo", "Logo (si no subes uno, se usa el logotipo específico del área seleccionada)",
        ["png", "jpg", "jpeg"],
        valor_por_defecto=(logo_nombre_actual, logo_bytes_actual) if logo_bytes_actual else None,
    )
    # Se guarda en la base de datos en cuanto cambia de verdad (se sube
    # uno nuevo o se elimina), comparando contra lo último que se
    # sincronizó -para no reescribir el archivo en cada repintado-.
    _logo_guardado_sesion = st.session_state.get("ajustes_logo_guardado")
    _ultimo_logo_sincronizado = st.session_state.get("_ultimo_logo_sincronizado")
    if _logo_guardado_sesion and _logo_guardado_sesion != _ultimo_logo_sincronizado:
        set_logo_informe(*_logo_guardado_sesion)
        st.session_state["_ultimo_logo_sincronizado"] = _logo_guardado_sesion
    elif st.session_state.get("ajustes_logo_eliminado") and _ultimo_logo_sincronizado is not None:
        set_logo_informe(None, None)
        st.session_state["_ultimo_logo_sincronizado"] = None

    hay_cambios_ajustes = (
        st.session_state["ajustes_empresa"] != (empresa_actual or "")
        or st.session_state["ajustes_cif"] != (cif_actual or "")
        or st.session_state["ajustes_tecnico"] != (tecnico_actual or "")
    )

    def _guardar_ajustes():
        set_empresa(st.session_state["ajustes_empresa"].strip())
        set_cif(st.session_state["ajustes_cif"].strip())
        set_tecnico_centro(cid, st.session_state["ajustes_tecnico"].strip())

    # Al pulsar "Volver", si hay cambios sin guardar se guardan solos
    # (sin pedir confirmación), en vez de perderlos sin más.
    if volver_clic_ajustes:
        if hay_cambios_ajustes:
            _guardar_ajustes()
        st.session_state.view = "centro"
        st.rerun()


# ============================================================
# TRADUCCIÓN DE TEXTOS FIJOS EN INGLÉS (cámara, subida de archivos...)
# ============================================================
# Streamlit no permite traducir directamente los textos internos de
# algunos widgets nativos (p.ej. "Take Photo" del selector de cámara,
# o "Drag and drop file here" del subidor de archivos). Se inyecta un
# pequeño script que busca esos textos en la página y los sustituye por
# su equivalente en castellano, y los vuelve a aplicar cada vez que
# Streamlit repinta la interfaz. Si el navegador bloquea el acceso al
# documento padre (restricción de origen cruzado), simplemente no se
# traduce y el texto se queda en inglés, sin romper nada más.

def inyectar_traduccion_widgets():
    html = """
    <script>
    (function() {
      const traducciones = [
        ["Take Photo", "Sacar foto"],
        ["Clear photo", "Quitar foto"],
        ["Switch camera", "Cambiar cámara"],
        ["Drag and drop file here", "Arrastra el archivo aquí"],
        ["Drag and drop files here", "Arrastra los archivos aquí"],
        ["Browse files", "Buscar archivo"],
        ["Upload", "Subir archivo"],
        ["200MB per file", "200MB max."],
      ];

      function traducirNodo(nodo) {
        if (nodo.nodeType === Node.TEXT_NODE) {
          const original = nodo.nodeValue;
          for (const [en, es] of traducciones) {
            if (original.includes(en)) {
              nodo.nodeValue = original.split(en).join(es);
            }
          }
        } else if (nodo.nodeType === Node.ELEMENT_NODE) {
          nodo.childNodes.forEach(traducirNodo);
        }
      }

      // Campos de texto/número con algo escrito: fondo gris claro.
      // Vacíos: blanco (como siempre). Como escribir no cambia el
      // árbol del DOM (solo la propiedad "value"), hace falta mirar
      // cada campo con sus propios listeners de "input"/"change",
      // además de revisarlos todos cada vez que Streamlit vuelve a
      // dibujar la pantalla (rerun), que es cuando aparecen campos
      // nuevos todavía sin su listener enganchado.
      const ROSA_CAMPO_VACIO = "#FDECEF";
      const COLOR_PUNTO_PENDIENTE = "#E94B68";

      function quitarPuntoPendiente(campo) {
        const contenedor = campo.parentElement;
        if (!contenedor) return;
        const punto = contenedor.querySelector(":scope > .punto-campo-pendiente");
        if (punto) punto.remove();
        campo.style.removeProperty("padding-left");
      }

      function ponerPuntoPendiente(campo) {
        const contenedor = campo.parentElement;
        if (!contenedor) return;
        contenedor.style.setProperty("position", "relative", "important");
        let punto = contenedor.querySelector(":scope > .punto-campo-pendiente");
        if (!punto) {
          punto = document.createElement("span");
          punto.className = "punto-campo-pendiente";
          punto.setAttribute("aria-hidden", "true");
          contenedor.appendChild(punto);
        }
        punto.style.cssText = [
          "position:absolute",
          "left:15px",
          "top:50%",
          "transform:translateY(-50%)",
          "width:14px",
          "height:14px",
          "border-radius:9999px !important",
          "min-width:14px !important",
          "max-width:14px !important",
          "min-height:14px !important",
          "max-height:14px !important",
          "display:block !important",
          "background-color:" + COLOR_PUNTO_PENDIENTE + " !important",
          "background:" + COLOR_PUNTO_PENDIENTE + " !important",
          "border:0 !important",
          "box-shadow:none !important",
          "opacity:1 !important",
          "z-index:9999",
          "pointer-events:none"
        ].join(";");
        campo.style.setProperty("padding-left", "38px", "important");
      }

      function aplicarFondoSegunValor(campo) {
        // La clave de acceso queda blanca y sin aviso.
        if (campo.type === "password") {
          campo.style.setProperty("background-color", "#C9DFF2", "important");
          quitarPuntoPendiente(campo);
          campo.style.removeProperty("border");
          campo.style.removeProperty("outline");
          campo.style.removeProperty("outline-offset");
          campo.style.removeProperty("box-shadow");
          return;
        }

        const tieneValor = campo.value !== undefined && campo.value !== null
          && String(campo.value).trim() !== "";

        // Los campos se mantienen con su borde normal. El aviso de pendiente
        // es únicamente el fondo rosa claro + el punto rojo grande.
        campo.style.removeProperty("border");
        campo.style.removeProperty("outline");
        campo.style.removeProperty("outline-offset");
        campo.style.removeProperty("box-shadow");

        if (tieneValor) {
          campo.style.setProperty("background-color", "#C9DFF2", "important");
          if (campo.parentElement) campo.parentElement.style.setProperty("background-color", "#C9DFF2", "important");
          quitarPuntoPendiente(campo);
        } else {
          campo.style.setProperty("background-color", ROSA_CAMPO_VACIO, "important");
          ponerPuntoPendiente(campo);
        }
      }

      function engancharYAplicarCampos(raiz) {
        const campos = raiz.querySelectorAll(
          'input[type="text"], input[type="number"], input[type="password"], input:not([type]), textarea'
        );
        campos.forEach(function(campo) {
          // Las celdas de las tablas editables (st.data_editor) NO
          // llevan este fondo gris/rosa: son de "glide-data-grid"
          // (clase gdg-input) y se pintan aparte, en negro sobre
          // blanco, para que el texto se pueda leer bien siempre.
          if (campo.classList.contains("gdg-input")) return;
          if (!campo.dataset.fondoSegunValor) {
            campo.dataset.fondoSegunValor = "1";
            campo.addEventListener("input", function() { aplicarFondoSegunValor(campo); });
            campo.addEventListener("change", function() { aplicarFondoSegunValor(campo); });
          }
          aplicarFondoSegunValor(campo);
        });
      }

      try {
        const doc = window.parent.document;
        traducirNodo(doc.body);
        engancharYAplicarCampos(doc.body);
        const observer = new MutationObserver(function() {
          traducirNodo(doc.body);
          engancharYAplicarCampos(doc.body);
        });
        observer.observe(doc.body, { childList: true, subtree: true, characterData: true });
      } catch (e) {
        // Si el navegador no permite acceder al documento padre, se
        // deja tal cual (seguirá en inglés en ese caso concreto).
      }
    })();
    </script>
    """
    components.html(html, height=0)


# ============================================================
# MAIN
# ============================================================

def _leer_claves_acceso():
    """
    Lee varias cadenas desde «Asacec», situado junto a app.py.
    La contraseña de acceso son los caracteres 3.º a 6.º de cada cadena.

    Ejemplo:
        ADMIN=AB1234KLM9
        Coruña=ZX1111PQ7R

    Contraseñas resultantes:
        ADMIN -> 1234
        Coruña -> 1111
    """
    ruta_clave = os.path.join(_carpeta_script, "Asacec")
    try:
        with open(ruta_clave, "r", encoding="utf-8") as f:
            lineas = [ln.strip() for ln in f if ln.strip() and not ln.lstrip().startswith("#")]
    except FileNotFoundError:
        return None
    if not lineas:
        return None
    accesos = []
    for i, linea in enumerate(lineas):
        if "=" in linea:
            identificador, clave = linea.split("=", 1)
            identificador, clave = identificador.strip(), clave.strip()
            if not identificador or not clave:
                continue
            es_admin = identificador.upper() == "ADMIN"
            accesos.append({"usuario": "__admin__" if es_admin else identificador, "clave": clave, "admin": es_admin})
        else:
            accesos.append({"usuario": "__admin__" if i == 0 else f"usuario_{i+1}", "clave": linea, "admin": i == 0})
    return accesos or None


def _clave_visible_desde_cadena(cadena):
    """
    La contraseña válida son los caracteres 3.º a 6.º de la cadena
    guardada en «Asacec» (índices 2:6 en Python).
    """
    cadena = str(cadena or "")
    return cadena[2:6] if len(cadena) >= 6 else ""


def _identificar_clave(clave_introducida, accesos):
    clave_introducida = str(clave_introducida or "")
    for acceso in accesos or []:
        clave_real = _clave_visible_desde_cadena(acceso["clave"])
        if clave_real and hmac.compare_digest(clave_introducida, clave_real):
            return acceso
    return None


def _activar_fullscreen_primer_gesto():
    """Intenta activar pantalla completa con el primer gesto real del usuario.

    Los navegadores solo permiten requestFullscreen() durante una interacción
    del usuario. Se instala un listener en el documento principal y se elimina
    tras el primer intento. Si el navegador no lo permite, la app sigue normal.
    """
    components.html(
        """
        <script>
        (() => {
          try {
            const w = window.parent;
            const d = w.document;

            if (w.__radonFullscreenPrimerGestoV1) return;
            w.__radonFullscreenPrimerGestoV1 = true;

            const limpiar = () => {
              d.removeEventListener("pointerdown", intentar, true);
              d.removeEventListener("keydown", intentarTeclado, true);
            };

            const pedir = () => {
              try {
                if (d.fullscreenElement) {
                  limpiar();
                  return;
                }
                const el = d.documentElement;
                const fn =
                  el.requestFullscreen ||
                  el.webkitRequestFullscreen ||
                  el.msRequestFullscreen;

                if (!fn) {
                  limpiar();
                  return;
                }

                const r = fn.call(el);
                if (r && typeof r.catch === "function") {
                  r.catch(() => {});
                }
              } catch (e) {}
              limpiar();
            };

            function intentar(ev) {
              // Solo un gesto primario real. No intercepta ni cancela el clic.
              if (ev && ev.isTrusted === false) return;
              pedir();
            }

            function intentarTeclado(ev) {
              if (ev && ev.isTrusted === false) return;
              // Permite que Intro en el formulario de acceso cuente como gesto.
              if (ev.key === "Enter" || ev.key === " ") pedir();
            }

            d.addEventListener("pointerdown", intentar, true);
            d.addEventListener("keydown", intentarTeclado, true);
          } catch (e) {}
        })();
        </script>
        """,
        height=0,
    )


def pantalla_login():
    _render_titulo_principal("🔒 Acceso restringido")
    _activar_fullscreen_primer_gesto()
    accesos = _leer_claves_acceso()
    if accesos is None:
        st.error(
            "No se ha configurado ninguna clave de acceso: falta el archivo «Asacec» "
            "en la carpeta de la app (o está vacío). Ejemplo: "
            "ADMIN=AB1234KLM9 y Coruña=ZX1111PQ7R."
        )
        return
    st.markdown('<div class="marcador-btn-guardar-ajustes"></div>', unsafe_allow_html=True)

    # El fondo coral general se aplica también al contenedor BaseWeb del input.
    # Este estilo se inyecta DESPUÉS de crear la clave para que la pantalla
    # de acceso quede siempre blanca, tanto el contenedor como el input interior.
    st.markdown("""
    <style>
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"],
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"] > div,
    div[data-testid="stTextInput"]:has(input[type="password"]) input[type="password"],
    div[data-baseweb="input"]:has(input[type="password"]),
    div[data-baseweb="input"]:has(input[type="password"]) > div,
    input[type="password"] {
        background: #FFFFFF !important;
        background-color: #FFFFFF !important;
    }

    div[data-testid="stTextInput"]:has(input[type="password"]) input[type="password"] {
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        caret-color: #000000 !important;
    }

    div[data-testid="stTextInput"]:has(input[type="password"]) button,
    div[data-baseweb="input"]:has(input[type="password"]) button {
        background: #FFFFFF !important;
        background-color: #FFFFFF !important;
        color: #000000 !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # Formulario de acceso: además de pulsar el botón «Entrar»,
    # permite enviar la clave pulsando Intro desde el campo de contraseña.
    with st.form("form_login", clear_on_submit=False):
        clave_introducida = st.text_input(
            "Clave de acceso",
            type="password",
            key="login_clave",
        )
        enviar_login = st.form_submit_button(
            "Entrar",
            type="primary",
            use_container_width=True,
        )

    if enviar_login:
        acceso = _identificar_clave(clave_introducida, accesos)
        if acceso:
            st.session_state["autenticado"] = True
            st.session_state["usuario_acceso"] = acceso["usuario"]
            st.session_state["es_admin"] = acceso["admin"]
            st.session_state["centro_actual"] = None
            st.session_state["view"] = "inicio"
            st.rerun()
        else:
            st.error("Clave incorrecta.")


def _activar_atras_navegador(view_actual, centro_actual=None):
    """Convierte Atrás del navegador/móvil en el equivalente a «← Volver»."""
    destinos = {
        "centro": "inicio",
        "centro_datos": "centro",
        "centro_categorias": "centro",
        "centro_planos": "centro",
        "centro_detectores": "centro",
        "centro_retirada": "centro",
        "centro_informes": "centro",
        "centro_informe_completo": "centro",
        "detector": "centro_detectores",
        "ajustes": "inicio",
    }
    destino = destinos.get(view_actual)
    if not destino:
        return

    components.html(
        f"""
        <script>
        (() => {{
          try {{
            const w = window.parent;
            const vista = {view_actual!r};
            const destino = {destino!r};

            // El destino se actualiza en cada render para que el manejador nunca
            // conserve la pantalla anterior.
            w.__radonBackDestino = destino;

            // Creamos una entrada "guardia" por cada pantalla real. Al pulsar Atrás
            // se vuelve a la entrada base y popstate ejecuta el mismo destino que
            // el botón ← Volver.
            const estado = w.history.state || {{}};
            if (!(estado.radonGuard === true && estado.radonView === vista)) {{
              w.history.replaceState(
                {{...estado, radonGuard: false, radonView: vista}},
                "",
                w.location.href
              );
              w.history.pushState(
                {{radonGuard: true, radonView: vista}},
                "",
                w.location.href
              );
            }}

            if (!w.__radonBackHandlerV2) {{
              w.__radonBackHandlerV2 = true;
              w.addEventListener("popstate", () => {{
                try {{
                  const d = w.__radonBackDestino;
                  if (!d) return;
                  const u = new URL(w.location.href);
                  u.searchParams.set("view", d);
                  if (d === "inicio") {{
                    u.searchParams.delete("centro");
                  }}
                  // replace evita crear otra entrada extra durante el propio "Volver".
                  w.location.replace(u.toString());
                }} catch (e) {{}}
              }});
            }}
          }} catch (e) {{}}
        }})();
        </script>
        """,
        height=0,
    )


def main():

    st.markdown("""
    <style>
    /* Refuerzo superior del aviso: se dibuja sobre el recuadro visible */
    div[data-testid="stTextInput"]:has(input:placeholder-shown:not([type="password"])) div[data-baseweb="input"],
    div[data-testid="stTextArea"]:has(textarea:placeholder-shown) div[data-baseweb="textarea"],
    div[data-testid="stNumberInput"]:has(input:placeholder-shown) div[data-baseweb="input"],
    div[data-testid="stSelectbox"]:has(input:placeholder-shown) div[data-baseweb="select"] > div {
        position: relative !important;
        z-index: 2 !important;
        border: 3px solid #D32F2F !important;
        outline: 3px solid #D32F2F !important;
        outline-offset: -1px !important;
        box-shadow: inset 0 0 0 2px #D32F2F, 0 0 0 1px #D32F2F !important;
        background-color: #FFFFFF !important;
    }

    /* La clave sigue fuera del sistema de aviso */
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"] {
        border-color: transparent !important;
        outline: none !important;
        box-shadow: none !important;
    }
    </style>
    """, unsafe_allow_html=True)


    st.markdown("""
    <style>
    /* CAMPOS VACÍOS: fondo blanco y borde rojo visible */
    div[data-testid="stTextInput"] input:placeholder-shown:not([type="password"]),
    div[data-testid="stTextArea"] textarea:placeholder-shown,
    div[data-testid="stNumberInput"] input:placeholder-shown {
        background-color: #FFFFFF !important;
        border: 3px solid #D32F2F !important;
        outline: 3px solid #D32F2F !important;
        outline-offset: -1px !important;
        box-shadow: inset 0 0 0 2px #D32F2F, 0 0 0 1px #D32F2F !important;
    }

    /* Desplegables todavía sin selección */
    div[data-testid="stSelectbox"] div[data-baseweb="select"]:has(input:placeholder-shown) > div {
        background-color: #FFFFFF !important;
        border: 3px solid #D32F2F !important;
        outline: 3px solid #D32F2F !important;
        outline-offset: -1px !important;
        box-shadow: inset 0 0 0 2px #D32F2F, 0 0 0 1px #D32F2F !important;
    }

    /* La clave de acceso es la excepción: blanca y sin borde rojo */
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"],
    div[data-testid="stTextInput"]:has(input[type="password"]) input[type="password"] {
        background-color: #FFFFFF !important;
        border-color: transparent !important;
        box-shadow: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

    if not st.session_state.get("autenticado"):
        pantalla_login()
        return

    init_db()
    inyectar_traduccion_widgets()

    if "view" not in st.session_state:
        st.session_state.view = "inicio"
    if "centro_actual" not in st.session_state:
        st.session_state.centro_actual = None
    if "detector_actual" not in st.session_state:
        st.session_state.detector_actual = None

    # Para que el botón "atrás" del móvil (o del navegador) vuelva a la
    # pantalla anterior DENTRO de la app, en vez de salir directamente
    # a la de inicio: se guarda la vista y el centro actuales en la
    # propia URL (con st.query_params). El navegador trata cada
    # cambio de la URL como un paso más en su historial, así que su
    # botón de atrás nativo puede recuperarlo.
    #
    # OJO: no basta con comparar la URL contra el session_state.view
    # ACTUAL para saber si el cambio viene de fuera (botón atrás) o de
    # dentro (un botón que acaba de cambiar de pantalla y ha llamado a
    # st.rerun() a mitad de guión, sin llegar a la sincronización de
    # más abajo en ESTA vuelta): en ese caso la URL se queda "atrasada"
    # un instante, y si se comparara contra la vista nueva parecería
    # (por error) que el usuario ha pulsado atrás, deshaciendo la
    # navegación que se acababa de pedir. Por eso se guarda aparte cuál
    # fue la ÚLTIMA URL que esta misma app escribió con éxito, y solo
    # se considera "el usuario ha pulsado atrás" cuando la URL cambia
    # respecto a ESO, no respecto al session_state.view más reciente.
    qp_view = st.query_params.get("view")
    qp_centro = st.query_params.get("centro")
    url_actual = (qp_view, qp_centro)
    ultima_url_propia = st.session_state.get("_ultima_url_sincronizada")
    if qp_view and ultima_url_propia is not None and url_actual != ultima_url_propia:
        st.session_state.view = qp_view
        st.session_state.centro_actual = int(qp_centro) if qp_centro and qp_centro.isdigit() else None


    if st.session_state.get("centro_actual") and not usuario_puede_acceder_centro(st.session_state.centro_actual):
        st.session_state.centro_actual = None
        st.session_state.detector_actual = None
        st.session_state.view = "inicio"
        st.warning("No tienes permiso para acceder a ese centro.")


    # Estilo final de campos pendientes: rosa claro + punto rojo, sin borde rojo.
    st.markdown(r"""
    <style>
    /* Los campos vacíos conservan un borde neutro y usan un rosa claro de aviso. */
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) div[data-baseweb="input"]:has(input:placeholder-shown),
    div[data-testid="stTextArea"] div[data-baseweb="textarea"]:has(textarea:placeholder-shown),
    div[data-testid="stSelectbox"] div[data-baseweb="select"]:has(input:placeholder-shown) > div {
        background: #FBE9ED !important;
        background-color: #FBE9ED !important;
        border: 1px solid rgba(49,51,63,.25) !important;
        outline: none !important;
        box-shadow: none !important;
    }

    div[data-testid="stTextInput"]:not(:has(input[type="password"])) input:placeholder-shown,
    div[data-testid="stTextArea"] textarea:placeholder-shown {
        background: #FBE9ED !important;
        background-color: #FBE9ED !important;
        border: 0 !important;
        outline: none !important;
        box-shadow: none !important;
    }

    /* Punto rojo grande a la izquierda, integrado en el fondo del control. */
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) div[data-baseweb="input"]:has(input:placeholder-shown),
    div[data-testid="stTextArea"] div[data-baseweb="textarea"]:has(textarea:placeholder-shown),
    div[data-testid="stSelectbox"] div[data-baseweb="select"]:has(input:placeholder-shown) > div {
        background-image: radial-gradient(circle at 18px 50%, #E53935 0 6px, transparent 7px) !important;
        background-repeat: no-repeat !important;
    }
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) input:placeholder-shown,
    div[data-testid="stTextArea"] textarea:placeholder-shown {
        padding-left: 38px !important;
    }
    div[data-testid="stSelectbox"] div[data-baseweb="select"]:has(input:placeholder-shown) > div {
        padding-left: 28px !important;
    }

    /* Fuerza el punto dinámico a rojo: evita que estilos globales de Streamlit lo pinten blanco. */
    .punto-campo-pendiente {
        width: 14px !important;
        height: 14px !important;
        min-width: 14px !important;
        max-width: 14px !important;
        min-height: 14px !important;
        max-height: 14px !important;
        border-radius: 9999px !important;
        clip-path: circle(50% at 50% 50%) !important;
        display: block !important;
        background: #E53935 !important;
        background-color: #E53935 !important;
        border: 0 !important;
        box-shadow: none !important;
        opacity: 1 !important;
    }

    /* La clave nunca participa en el aviso. */
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"],
    div[data-testid="stTextInput"] input[type="password"] {
        background: #FFFFFF !important;
        background-image: none !important;
        border-color: transparent !important;
        outline: none !important;
        box-shadow: none !important;
        padding-left: inherit !important;
    }
    </style>
    """, unsafe_allow_html=True)


    # PRUEBA VISUAL: campos cubiertos en azul muy tenue, manteniendo pendientes en rosa.
    st.markdown(r"""
    <style>
    /* Controles normales/cubiertos: azul muy claro en consonancia con el fondo. */
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) div[data-baseweb="input"]:not(:has(input:placeholder-shown)),
    div[data-testid="stTextArea"] div[data-baseweb="textarea"]:not(:has(textarea:placeholder-shown)),
    div[data-testid="stNumberInput"] div[data-baseweb="input"],
    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div {
        background: #C9DFF2 !important;
        background-color: #C9DFF2 !important;
    }
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) input:not(:placeholder-shown),
    div[data-testid="stTextArea"] textarea:not(:placeholder-shown),
    div[data-testid="stNumberInput"] input {
        background: #C9DFF2 !important;
        background-color: #C9DFF2 !important;
    }

    /* Los campos pendientes conservan exactamente el rosa y el círculo rojo. */
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) div[data-baseweb="input"]:has(input:placeholder-shown),
    div[data-testid="stTextArea"] div[data-baseweb="textarea"]:has(textarea:placeholder-shown),
    div[data-testid="stSelectbox"] div[data-baseweb="select"]:has(input:placeholder-shown) > div,
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) input:placeholder-shown,
    div[data-testid="stTextArea"] textarea:placeholder-shown {
        background-color: #FBE9ED !important;
    }

    /* La clave sigue blanca. */
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"],
    div[data-testid="stTextInput"] input[type="password"] {
        background: #FFFFFF !important;
        background-color: #FFFFFF !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # Corrección final: algunos estilos anteriores de Streamlit devolvían los controles a blanco.
    # Este bloque se carga el último para que el azul tenue sea realmente visible.
    st.markdown(r"""
    <style>
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) div[data-baseweb="input"],
    div[data-testid="stTextArea"] div[data-baseweb="textarea"],
    div[data-testid="stNumberInput"] div[data-baseweb="input"],
    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div,
    div[data-testid="stDateInput"] div[data-baseweb="input"],
    div[data-testid="stTimeInput"] div[data-baseweb="input"] {
        background: #C9DFF2 !important;
        background-color: #C9DFF2 !important;
    }
    div[data-testid="stTextInput"]:not(:has(input[type="password"])) input,
    div[data-testid="stTextArea"] textarea,
    div[data-testid="stNumberInput"] input,
    div[data-testid="stDateInput"] input,
    div[data-testid="stTimeInput"] input {
        background: transparent !important;
    }
    /* El aviso pendiente lo pinta el script en rosa sobre el propio campo. */
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"],
    div[data-testid="stTextInput"] input[type="password"] {
        background: #FFFFFF !important;
        background-color: #FFFFFF !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # Ajuste final de controles: azul tenue visible y punto pendiente más pequeño.
    st.markdown(r"""
    <style>
    div[data-testid="stTextInput"] div[data-baseweb="input"],
    div[data-testid="stTextArea"] div[data-baseweb="textarea"],
    div[data-testid="stNumberInput"] div[data-baseweb="input"],
    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div,
    div[data-testid="stDateInput"] div[data-baseweb="input"],
    div[data-testid="stTimeInput"] div[data-baseweb="input"] {
        background-color: #C9DFF2 !important;
    }
    div[data-testid="stTextInput"] input,
    div[data-testid="stTextArea"] textarea,
    div[data-testid="stNumberInput"] input,
    div[data-testid="stDateInput"] input,
    div[data-testid="stTimeInput"] input {
        background-color: transparent !important;
    }
    .punto-campo-pendiente {
        width:14px !important; height:14px !important;
        min-width:14px !important; max-width:14px !important;
        min-height:14px !important; max-height:14px !important;
        border-radius:50% !important; clip-path:circle(50%) !important;
        background:#E53935 !important; background-color:#E53935 !important;
    }
    div[data-testid="stTextInput"]:has(input[type="password"]) div[data-baseweb="input"],
    div[data-testid="stTextInput"] input[type="password"] { background:#FFFFFF !important; }
    </style>
    """, unsafe_allow_html=True)

    view = st.session_state.view
    if view == "inicio":
        pantalla_inicio()
    elif view == "centro":
        pantalla_centro()
    elif view == "centro_datos":
        pantalla_centro_datos()
    elif view == "centro_categorias":
        pantalla_centro_categorias()
    elif view == "centro_planos":
        pantalla_centro_planos()
    elif view == "centro_detectores":
        pantalla_centro_detectores()
    elif view == "centro_retirada":
        pantalla_centro_retirada()
    elif view == "centro_informes":
        pantalla_centro_informes()
    elif view == "centro_informe_completo":
        pantalla_centro_informe_completo()
    elif view == "detector":
        pantalla_detector()
    elif view == "ajustes":
        pantalla_ajustes()
    else:
        st.session_state.view = "inicio"
        st.rerun()

    # Se sincroniza la URL con la vista que se acaba de mostrar (por si
    # ha cambiado durante esta ejecución), para que el botón de atrás
    # pueda volver a ella más adelante; y se recuerda como "la última
    # URL propia", para la comprobación de arriba en la próxima vuelta.
    nuevo_qp_view = st.session_state.view
    nuevo_qp_centro = str(st.session_state.centro_actual) if st.session_state.centro_actual else None
    if (st.query_params.get("view"), st.query_params.get("centro")) != (nuevo_qp_view, nuevo_qp_centro):
        st.query_params["view"] = nuevo_qp_view
        if nuevo_qp_centro:
            st.query_params["centro"] = nuevo_qp_centro
        elif "centro" in st.query_params:
            del st.query_params["centro"]
    st.session_state["_ultima_url_sincronizada"] = (nuevo_qp_view, nuevo_qp_centro)

    # Se instala al final, cuando la URL ya representa la pantalla actual.
    _activar_atras_navegador(st.session_state.view, st.session_state.get("centro_actual"))


if __name__ == "__main__":
    main()
