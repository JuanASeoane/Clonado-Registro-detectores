"""
Generación automática del Anexo II: esquema gráfico del edificio con la
posición de cada detector marcada sobre el plano correspondiente.

Los datos se leen directamente del Excel original:
  - Hoja "Planos": columna A = nombre del plano (o "(Foto exterior del
    centro)"), columna B = imagen embebida (una foto/plano por fila).
  - Hoja "Detectores": columnas "Nombre del plano", "Punto X", "Punto Y"
    (coordenadas relativas 0-1 sobre la imagen del plano), "Código"
    (código del detector) y "Código de la sala".

Por cada plano distinto que tenga al menos un detector asociado se genera
una página con:
  - el plano, con un punto rojo en cada posición exacta (las coordenadas
    relativas no se alteran; solo se recolocan las ETIQUETAS de texto para
    que no se solapen entre sí),
  - el código del detector junto al punto y, debajo, entre paréntesis, el
    código de la sala,
  - el logotipo y la foto exterior del centro en una columna aparte, sin
    tapar el plano.
"""

from __future__ import annotations

import io
import unicodedata

import openpyxl
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

from utils_informe.assets import logo_por_defecto

PUNTO_RADIO_FRAC = 0.010  # radio del punto, como fracción del lado menor de la imagen
MARGEN_FRAC = 0.012


def _normalizar(texto) -> str:
    texto = str(texto or "").strip().lower()
    return "".join(c for c in unicodedata.normalize("NFD", texto) if unicodedata.category(c) != "Mn")


class Anexo2Error(ValueError):
    """No se han podido extraer los datos necesarios del Excel para generar el Anexo II."""


def extraer_datos_planos(file) -> dict:
    """Lee el Excel original (ruta o fichero) y devuelve:
    {
        "planos": {nombre_normalizado: {"nombre": str, "imagen": bytes}},
        "foto_exterior": bytes | None,
        "puntos": {nombre_normalizado: [{"x": float, "y": float, "codigo": str, "sala": str}]},
    }
    """
    if hasattr(file, "seek"):
        file.seek(0)
    wb = openpyxl.load_workbook(file, data_only=True)

    if "Planos" not in wb.sheetnames or "Detectores" not in wb.sheetnames:
        raise Anexo2Error("El Excel no tiene las hojas 'Planos' y 'Detectores' necesarias.")

    ws_planos = wb["Planos"]

    # nombre (columna A) de cada fila que tiene una imagen anclada en columna B
    nombres_por_fila = {}
    for row in ws_planos.iter_rows(min_col=1, max_col=1):
        for cell in row:
            if cell.value is not None:
                nombres_por_fila[cell.row] = str(cell.value).strip()

    planos = {}
    foto_exterior = None
    for img in ws_planos._images:
        fila_excel = img.anchor._from.row + 1  # openpyxl es 0-index en el anchor
        nombre = nombres_por_fila.get(fila_excel, f"Plano fila {fila_excel}")
        data = img._data()
        clave = _normalizar(nombre)
        if "exterior" in clave:
            foto_exterior = data
        else:
            planos[clave] = {"nombre": nombre, "imagen": data}

    if not planos:
        raise Anexo2Error("No se ha encontrado ninguna imagen de plano en la hoja 'Planos'.")

    # puntos por plano, desde la hoja Detectores (fila 2 = cabecera real)
    ws_det = wb["Detectores"]
    raw = list(ws_det.iter_rows(values_only=True))
    header = raw[1]
    col_idx = {str(h).strip(): i for i, h in enumerate(header) if h is not None}

    requeridas = ["Nombre del plano", "Punto X", "Punto Y", "Código", "Código de la sala"]
    faltan = [c for c in requeridas if c not in col_idx]
    if faltan:
        raise Anexo2Error(
            "Faltan columnas en la hoja 'Detectores' para generar el Anexo II: " + ", ".join(faltan)
        )

    puntos: dict[str, list[dict]] = {}
    for fila in raw[2:]:
        if fila is None:
            continue
        nombre_plano = fila[col_idx["Nombre del plano"]]
        x = fila[col_idx["Punto X"]]
        y = fila[col_idx["Punto Y"]]
        if nombre_plano is None or x is None or y is None:
            continue
        codigo = fila[col_idx["Código"]] or ""
        sala_cod = fila[col_idx["Código de la sala"]] or ""
        clave = _normalizar(nombre_plano)
        puntos.setdefault(clave, []).append(
            {"x": float(x), "y": float(y), "codigo": str(codigo).strip(), "sala": str(sala_cod).strip()}
        )

    return {"planos": planos, "foto_exterior": foto_exterior, "puntos": puntos}


def _cargar_fuente(tamano: int):
    try:
        return ImageFont.load_default(size=tamano)
    except TypeError:
        # Pillow < 10.1: load_default() no admite tamaño variable
        return ImageFont.load_default()


def _bbox_multilinea(draw, xy, lineas, font, spacing=2):
    return draw.multiline_textbbox(xy, "\n".join(lineas), font=font, spacing=spacing, align="center")


def _rectangulos_solapan(a, b, margen=2):
    return not (
        a[2] + margen < b[0]
        or a[0] - margen > b[2]
        or a[3] + margen < b[1]
        or a[1] - margen > b[3]
    )


def _escalar_imagen_contener(img: Image.Image, max_w: int, max_h: int) -> Image.Image:
    """Escala una imagen manteniendo proporción para que quepa en max_w x max_h."""
    w, h = img.size
    if w <= 0 or h <= 0:
        return img
    factor = min(max_w / w, max_h / h)
    factor = max(factor, 0.01)
    nuevo = (max(1, int(w * factor)), max(1, int(h * factor)))
    return img.resize(nuevo, Image.LANCZOS)


def _caja_expandida(caja, margen):
    return (
        caja[0] - margen,
        caja[1] - margen,
        caja[2] + margen,
        caja[3] + margen,
    )


def _punto_mas_cercano_rectangulo(px, py, caja):
    """Punto del borde del rectángulo más cercano a (px, py)."""
    x1, y1, x2, y2 = caja
    x = min(max(px, x1), x2)
    y = min(max(py, y1), y2)

    # Si el punto queda proyectado dentro de la caja, usamos el borde
    # más próximo para que la línea termine claramente en la etiqueta.
    if x1 < px < x2 and y1 < py < y2:
        distancias = [
            (abs(px - x1), (x1, py)),
            (abs(px - x2), (x2, py)),
            (abs(py - y1), (px, y1)),
            (abs(py - y2), (px, y2)),
        ]
        _, (x, y) = min(distancias, key=lambda item: item[0])
    return x, y


def _segmento_interseca_rectangulo(p1, p2, rect, margen=2):
    """Comprobación ligera para penalizar líneas que atraviesen otras etiquetas."""
    x1, y1, x2, y2 = _caja_expandida(rect, margen)
    ax, ay = p1
    bx, by = p2

    # Rechazo rápido por cajas envolventes.
    if max(ax, bx) < x1 or min(ax, bx) > x2 or max(ay, by) < y1 or min(ay, by) > y2:
        return False

    # Muestreo del segmento: suficiente para la colocación gráfica de etiquetas.
    pasos = max(8, int(max(abs(bx - ax), abs(by - ay)) / 8))
    for i in range(pasos + 1):
        t = i / pasos
        x = ax + (bx - ax) * t
        y = ay + (by - ay) * t
        if x1 <= x <= x2 and y1 <= y <= y2:
            return True
    return False


def _preparar_overlay(
    imagen_bytes: bytes,
    max_w: int,
    max_h: int,
    fondo_blanco: bool = True,
    padding: int = 8,
) -> Image.Image:
    """Prepara logo/foto para superponerlos con un pequeño margen blanco."""
    overlay = Image.open(io.BytesIO(imagen_bytes)).convert("RGBA")
    overlay = _escalar_imagen_contener(
        overlay,
        max(1, max_w - 2 * padding),
        max(1, max_h - 2 * padding),
    )

    if not fondo_blanco:
        return overlay

    fondo = Image.new(
        "RGBA",
        (overlay.width + 2 * padding, overlay.height + 2 * padding),
        (255, 255, 255, 242),
    )
    fondo.alpha_composite(overlay, (padding, padding))
    return fondo


def _rectangulos_overlays(
    w: int,
    h: int,
    logo_bytes: bytes | None,
    foto_exterior: bytes | None,
    mostrar_logo: bool,
    mostrar_foto_exterior: bool,
):
    """Calcula overlays en la esquina inferior derecha.

    El logotipo queda abajo a la derecha y la foto exterior justo encima,
    con una separación pequeña. El logo se muestra aproximadamente un 50 %
    más grande que en la composición lateral anterior.
    """
    margen = max(10, int(min(w, h) * 0.015))
    separacion = max(8, int(min(w, h) * 0.012))
    overlays = []

    y_inferior = h - margen

    # Logo: ~25 % del ancho del plano. Equivale aproximadamente a un 50 %
    # más que el espacio visual que ocupaba en la antigua columna lateral.
    if mostrar_logo and logo_bytes:
        max_logo_w = int(w * 0.25)
        max_logo_h = int(h * 0.22)
        logo = _preparar_overlay(
            logo_bytes,
            max_logo_w,
            max_logo_h,
            fondo_blanco=True,
            padding=max(5, int(min(w, h) * 0.006)),
        )
        x = w - margen - logo.width
        y = y_inferior - logo.height
        overlays.append(("logo", logo, (x, y, x + logo.width, y + logo.height)))
        y_inferior = y - separacion

    if mostrar_foto_exterior and foto_exterior:
        max_foto_w = int(w * 0.23)
        max_foto_h = int(h * 0.24)
        foto = _preparar_overlay(
            foto_exterior,
            max_foto_w,
            max_foto_h,
            fondo_blanco=True,
            padding=max(4, int(min(w, h) * 0.005)),
        )
        x = w - margen - foto.width
        y = y_inferior - foto.height
        # Si no cabe sobre el logo, la mantenemos dentro del borde superior.
        y = max(margen, y)
        overlays.append(("foto", foto, (x, y, x + foto.width, y + foto.height)))

    return overlays


def componer_plano(
    imagen_bytes: bytes,
    puntos: list[dict],
    logo_bytes: bytes | None = None,
    foto_exterior: bytes | None = None,
    mostrar_logo: bool = False,
    mostrar_foto_exterior: bool = False,
) -> bytes:
    """Compone el plano con puntos, etiquetas y elementos superpuestos.

    - Los puntos permanecen exactamente en sus coordenadas originales.
    - Las etiquetas se desplazan para evitar colisiones entre sí y con otros
      puntos.
    - Cada etiqueta se une a su detector mediante una línea roja, por lo que
      sigue siendo inequívoco qué código corresponde a cada punto aunque el
      texto haya tenido que alejarse.
    - El texto rojo es deliberadamente más pequeño que en la versión anterior,
      ya que el plano ahora aprovecha prácticamente toda la hoja.
    - Si se solicitan, la foto exterior y el logotipo se superponen en la
      esquina inferior derecha: foto arriba y logotipo abajo.
    """
    img = Image.open(io.BytesIO(imagen_bytes)).convert("RGB")
    w, h = img.size
    draw = ImageDraw.Draw(img)

    lado_menor = min(w, h)
    radio = max(4, int(lado_menor * PUNTO_RADIO_FRAC))

    # Antes se utilizaba 0.032. Al ocupar el plano más superficie de página,
    # 0.022 da una etiqueta sensiblemente más pequeña pero todavía legible.
    font = _cargar_fuente(max(10, int(lado_menor * 0.022)))
    separacion_texto = max(1, int(lado_menor * 0.0025))
    margen_etiqueta = max(3, int(lado_menor * 0.004))
    grosor_linea = max(1, int(lado_menor * 0.0022))

    overlays = _rectangulos_overlays(
        w, h,
        logo_bytes=logo_bytes,
        foto_exterior=foto_exterior,
        mostrar_logo=mostrar_logo,
        mostrar_foto_exterior=mostrar_foto_exterior,
    )
    cajas_reservadas = [item[2] for item in overlays]

    # Reservamos todos los puntos desde el principio para impedir que el texto
    # de un detector termine encima del punto de otro.
    cajas_puntos = []
    for punto in puntos:
        px = float(punto["x"]) * w
        py = float(punto["y"]) * h
        r = radio + margen_etiqueta
        cajas_puntos.append((px - r, py - r, px + r, py + r))

    cajas_etiquetas: list[tuple[float, float, float, float]] = []
    ubicaciones = []

    def candidatos(px, py, bw, bh):
        """Posiciones alrededor del punto, de próximas a más alejadas."""
        pad = radio + margen_etiqueta + 3
        direcciones = [
            (1, 0), (-1, 0), (0, -1), (0, 1),
            (1, -1), (-1, -1), (1, 1), (-1, 1),
        ]
        # Varias coronas. La última permite separar bastante etiquetas
        # concentradas sin perder la asociación gracias a la línea guía.
        for escala in (1.0, 1.45, 1.9, 2.45, 3.1, 3.9, 4.8):
            distancia_x = pad + bw * 0.48 * escala
            distancia_y = pad + bh * 0.55 * escala
            for sx, sy in direcciones:
                if sx > 0:
                    cx = px + distancia_x
                elif sx < 0:
                    cx = px - distancia_x - bw
                else:
                    cx = px - bw / 2

                if sy > 0:
                    cy = py + distancia_y
                elif sy < 0:
                    cy = py - distancia_y - bh
                else:
                    cy = py - bh / 2
                yield cx, cy

    # Elegimos primero todas las posiciones; las líneas se dibujan después,
    # para poder penalizar trayectorias que crucen etiquetas ya colocadas.
    for idx, punto in enumerate(puntos):
        px = float(punto["x"]) * w
        py = float(punto["y"]) * h

        lineas = [str(punto.get("codigo") or "").strip()]
        lineas = [x for x in lineas if x]
        sala = str(punto.get("sala") or "").strip()
        if sala:
            lineas.append(f"({sala})")
        if not lineas:
            ubicaciones.append(None)
            continue

        bbox_0 = _bbox_multilinea(draw, (0, 0), lineas, font, spacing=separacion_texto)
        bw = bbox_0[2] - bbox_0[0]
        bh = bbox_0[3] - bbox_0[1]

        mejor = None
        mejor_puntuacion = None

        for cx, cy in candidatos(px, py, bw, bh):
            caja = (cx, cy, cx + bw, cy + bh)
            if caja[0] < 1 or caja[1] < 1 or caja[2] > w - 1 or caja[3] > h - 1:
                continue

            # Colisiones fuertes: etiquetas, cualquier punto y overlays.
            col_etiquetas = sum(_rectangulos_solapan(caja, otra, margen=margen_etiqueta)
                                for otra in cajas_etiquetas)
            col_puntos = sum(_rectangulos_solapan(caja, otra, margen=margen_etiqueta)
                             for otra in cajas_puntos)
            col_overlays = sum(_rectangulos_solapan(caja, otra, margen=margen_etiqueta)
                               for otra in cajas_reservadas)

            extremo = _punto_mas_cercano_rectangulo(px, py, caja)
            cruces_linea = sum(
                _segmento_interseca_rectangulo((px, py), extremo, otra, margen=1)
                for otra in cajas_etiquetas
            )

            distancia = ((extremo[0] - px) ** 2 + (extremo[1] - py) ** 2) ** 0.5
            puntuacion = (
                col_etiquetas * 100000
                + col_puntos * 80000
                + col_overlays * 100000
                + cruces_linea * 12000
                + distancia
            )

            if mejor_puntuacion is None or puntuacion < mejor_puntuacion:
                mejor_puntuacion = puntuacion
                mejor = (cx, cy, caja, extremo)

            if col_etiquetas == 0 and col_puntos == 0 and col_overlays == 0 and cruces_linea == 0:
                break

        if mejor is None:
            # Último recurso: posición dentro de la imagen, a la derecha.
            cx = min(max(1, px + radio + 4), max(1, w - bw - 1))
            cy = min(max(1, py - bh / 2), max(1, h - bh - 1))
            caja = (cx, cy, cx + bw, cy + bh)
            extremo = _punto_mas_cercano_rectangulo(px, py, caja)
            mejor = (cx, cy, caja, extremo)

        cx, cy, caja, extremo = mejor
        cajas_etiquetas.append(caja)
        ubicaciones.append({
            "px": px, "py": py,
            "cx": cx, "cy": cy,
            "caja": caja,
            "extremo": extremo,
            "lineas": lineas,
        })

    # Primero las líneas, para que el texto quede siempre por encima.
    for item in ubicaciones:
        if not item:
            continue
        px, py = item["px"], item["py"]
        ex, ey = item["extremo"]
        dx, dy = ex - px, ey - py
        longitud = max((dx * dx + dy * dy) ** 0.5, 1.0)

        # La línea comienza en el borde del punto, no en su centro.
        inicio_x = px + dx / longitud * (radio + 1)
        inicio_y = py + dy / longitud * (radio + 1)
        draw.line(
            [(inicio_x, inicio_y), (ex, ey)],
            fill=(190, 0, 0),
            width=grosor_linea,
        )

    # Puntos y textos.
    for punto in puntos:
        px = float(punto["x"]) * w
        py = float(punto["y"]) * h
        draw.ellipse(
            [px - radio, py - radio, px + radio, py + radio],
            fill=(220, 0, 0),
            outline=(90, 0, 0),
        )

    for item in ubicaciones:
        if not item:
            continue
        draw.multiline_text(
            (item["cx"], item["cy"]),
            "\n".join(item["lineas"]),
            font=font,
            fill=(200, 0, 0),
            spacing=separacion_texto,
            align="center",
        )

    # Los overlays se añaden al final para que se vean nítidos.
    img_rgba = img.convert("RGBA")
    for _, overlay, caja in overlays:
        x1, y1, _, _ = caja
        img_rgba.alpha_composite(overlay, (int(x1), int(y1)))

    buffer = io.BytesIO()
    img_rgba.convert("RGB").save(buffer, format="PNG", quality=95)
    return buffer.getvalue()

def _fit(img_bytes: bytes, max_width_cm: float, max_height_cm: float):
    img = Image.open(io.BytesIO(img_bytes))
    w_px, h_px = img.size
    aspect = w_px / h_px if h_px else 1
    max_w, max_h = Cm(max_width_cm), Cm(max_height_cm)
    height = max_h
    width = int(height * aspect)
    if width > max_w:
        width = max_w
        height = int(width / aspect)
    return width, height


def generar_documento_anexo2(
    datos: dict,
    logo_bytes: bytes | None = None,
    centro: str = "",
    titulo: str = "ANEXO II: ESQUEMA GRÁFICO DO EDIFICIO E PLANOS DE CADA PLANTA",
    mostrar_logo: bool = True,
    mostrar_foto_exterior: bool = True,
) -> bytes:
    """Genera una página A4 horizontal por plano usando todo el ancho útil.

    El plano siempre aprovecha prácticamente toda la página. Si se seleccionan,
    el logotipo se superpone abajo a la derecha y la foto exterior se coloca
    justo encima, ligeramente separada. No se reserva ya ninguna columna
    lateral para estos elementos.
    """
    planos = datos["planos"]
    puntos_por_plano = datos["puntos"]
    foto_exterior = datos.get("foto_exterior")
    logo_a_usar = logo_bytes or logo_por_defecto()

    nombres_con_puntos = [
        clave for clave in planos
        if clave in puntos_por_plano and puntos_por_plano[clave]
    ]
    if not nombres_con_puntos:
        raise Anexo2Error(
            "Ningún plano tiene detectores con coordenadas (Punto X / Punto Y) asociadas."
        )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)

    # Márgenes algo más ajustados para aprovechar mejor la hoja sin dejar el
    # contenido pegado al borde físico del papel.
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)

    ancho_util_cm = (
        section.page_width - section.left_margin - section.right_margin
    ) / 360000
    alto_util_cm = (
        section.page_height - section.top_margin - section.bottom_margin
    ) / 360000

    for i, clave in enumerate(nombres_con_puntos):
        if i > 0:
            doc.add_page_break()

        plano_info = planos[clave]

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{titulo} — {plano_info['nombre']}")
        run.bold = True
        run.font.size = Pt(11.5)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(0)

        compuesto = componer_plano(
            plano_info["imagen"],
            puntos_por_plano[clave],
            logo_bytes=logo_a_usar,
            foto_exterior=foto_exterior,
            mostrar_logo=mostrar_logo,
            mostrar_foto_exterior=mostrar_foto_exterior,
        )

        p_plano = doc.add_paragraph()
        p_plano.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_plano.paragraph_format.space_before = Pt(0)
        p_plano.paragraph_format.space_after = Pt(0)
        run_plano = p_plano.add_run()

        # Se deja solo la altura imprescindible para el título.
        w_plano, h_plano = _fit(
            compuesto,
            ancho_util_cm,
            alto_util_cm - 0.8,
        )
        run_plano.add_picture(
            io.BytesIO(compuesto),
            width=w_plano,
            height=h_plano,
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

